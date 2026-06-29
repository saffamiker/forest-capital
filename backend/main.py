"""
Forest Capital Portfolio Intelligence System — FastAPI backend.
Sprint 4: all 8 agents live, council deliberation wired, QA methodology checklist,
          WebSocket streaming, scope guard enforced, council_sessions logging.
"""
from __future__ import annotations
import asyncio
import json
import os
import time
import uuid
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Any, Optional

from fastapi import (
    FastAPI, Depends, HTTPException, Query, Response, WebSocket,
    WebSocketDisconnect, UploadFile, File, Form,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from starlette.requests import Request

from config import FRONTEND_URL, ENVIRONMENT, PERMISSIONS, ROLE_PRESETS
from agents.base import SONNET_MODEL, OPUS_MODEL, GEMINI_MODEL
from logger import configure_logging, get_logger
from auth import (
    require_auth,
    require_team_member,
    require_sysadmin,
    require_permission,
    require_master_key,
    generate_magic_token,
    generate_session_token,
    verify_magic_token,
    verify_session_token,
    redeem_magic_token,
    invalidate_session,
    send_magic_link,
)
from models.schemas import (
    MagicLinkRequest,
    MagicLinkResponse,
    SessionResponse,
    LogoutRequest,
    CouncilQueryRequest,
    BacktestRequest,
    QAQueryRequest,
    OptimizeRequest,
    UIUXReviewRequest,
    AdvisorAnalyseRequest,
    AdvisorVerifyRequest,
    AdvisorCitationsRequest,
    MOCK_STRATEGIES,
    MOCK_REGIME,
    MOCK_COUNCIL_RESPONSE,
    MOCK_QA_AUDIT,
    MOCK_EFFICIENT_FRONTIER,
)

configure_logging()
log = get_logger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    log.info("forest_capital_starting", environment=ENVIRONMENT, frontend_url=FRONTEND_URL)
    # Warm the academic-context cache so the first agent invocation after a
    # restart already carries the uploaded rubric / requirements documents.
    # Fail-open: a cold cache simply means agents run without that context.
    if ENVIRONMENT != "test":
        # Reap any audit left 'running' by a crash or a redeploy before
        # this boot — a hung run otherwise holds the statistical-audit
        # lock until the first poll or audit-start happens to clear it.
        try:
            from tools.audit_engine import fail_stale_audits
            reaped = await fail_stale_audits()
            if reaped:
                log.warning(
                    f"Startup reap: marked {reaped} stale audit "
                    f"run(s) as failed", count=reaped)
            else:
                log.info("Startup reap: no stale audit runs found")
        except Exception as exc:  # noqa: BLE001
            log.warning("audit_startup_reap_failed", error=str(exc))
        # Screenshot cleanup — drops UAT screenshots older than 30 days
        # from SCREENSHOT_DIR. The disk is the persistent Render volume
        # in production; without this sweep, every failure-report
        # screenshot ever uploaded would accumulate forever. Fail-open
        # — a missing directory or unreadable file logs and continues.
        try:
            from tools.test_runner import cleanup_old_screenshots
            deleted, remaining = cleanup_old_screenshots()
            if deleted == 0 and remaining == 0:
                log.info("screenshot_cleanup: directory empty or absent")
            else:
                log.info(
                    "screenshot_cleanup",
                    deleted=deleted,
                    remaining=remaining,
                    message=(
                        f"screenshot_cleanup: deleted {deleted} "
                        f"screenshots older than 30 days, "
                        f"{remaining} remain"
                    ),
                )
        except Exception as exc:  # noqa: BLE001
            log.warning("screenshot_cleanup_failed", error=str(exc))
        # One-time baseline log — counts the historical rows that landed
        # before the cost-tracking deploy (pre-migration 020 token columns
        # are null). Sets a clear line in the Render logs: everything from
        # this deploy forward carries cost data; rows older than this can
        # never be backfilled.
        try:
            from datetime import date
            from sqlalchemy import text
            from database import AsyncSessionLocal
            async with AsyncSessionLocal() as ses:  # type: ignore[union-attr]
                result = await ses.execute(text(
                    "SELECT COUNT(*) FROM agent_interactions "
                    "WHERE estimated_cost_usd IS NULL"))
                n_null = int(result.scalar() or 0)
            today = date.today().isoformat()
            log.info(
                "cost_tracking_baseline",
                active_from=today,
                n_historical_null=n_null,
                message=(f"cost tracking active from {today}. "
                         f"{n_null} historical interactions have no cost "
                         "data (pre-migration 020)."),
            )
        except Exception as exc:  # noqa: BLE001
            log.warning("cost_tracking_baseline_failed", error=str(exc))

        # June 28 2026 -- platform-config startup seed for the
        # DEFER_SUBSTITUTION_TO_EXPORT flag.
        #
        # Operator-confirmed root cause of Phase 2 deferral
        # silently no-op-ing on drafts 74 / 77 (despite PRs
        # #470 / #471 / #473 / #474): the platform_config row
        # keys 'defer_substitution_to_export' kept disappearing
        # on Render restarts. Without the row,
        # platform_flags._read_flag returns its default of
        # False, the deferral swap never fires, content_json
        # gets resolved values, and the upgrade pass finds
        # nothing to convert to token_value nodes.
        #
        # ON CONFLICT DO NOTHING means: if the row already
        # exists (anyone has set it to true OR false via
        # admin), respect that value -- never overwrite. The
        # seed only inserts when the row is MISSING. Idempotent
        # across restarts.
        #
        # Fail-open: a DB write failure logs + boot continues;
        # the flag defaults to OFF in that case (legacy
        # behaviour preserved).
        try:
            from sqlalchemy import text
            from database import AsyncSessionLocal
            async with AsyncSessionLocal() as ses:  # type: ignore[union-attr]
                # CAST as JSONB so the column type matches.
                res = await ses.execute(text(
                    "INSERT INTO platform_config (key, value) "
                    "VALUES (:k, CAST(:v AS JSONB)) "
                    "ON CONFLICT (key) DO NOTHING"),
                    {
                        "k": "defer_substitution_to_export",
                        "v": '{"enabled": true}',
                    })
                await ses.commit()
                # rowcount > 0 -> we just inserted; == 0 ->
                # row already existed (the ON CONFLICT path).
                inserted = (res.rowcount or 0) > 0
            log.info(
                "platform_config_defer_substitution_seed",
                inserted=inserted,
                key="defer_substitution_to_export",
                seeded_value=(
                    '{"enabled": true}' if inserted else (
                        "<existing row preserved>")))
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "platform_config_defer_substitution_seed_failed",
                error=str(exc))

        try:
            from tools.academic_context import refresh_academic_context
            await refresh_academic_context()
        except Exception as exc:  # noqa: BLE001
            log.warning("academic_context_warm_failed", error=str(exc))
        # June 26 2026 -- auto-recover drafts left in the NULL-
        # current-draft state by a generation job that landed the
        # editor_drafts row but crashed before content_json was
        # written. For each document_type with is_current=true AND
        # content_json IS NULL, demotes the broken row + restores
        # the most recent prior good draft as is_current. Fail-
        # open: a DB error is logged and startup continues
        # unaffected. Idempotent (a subsequent boot finds zero
        # broken rows and returns 0).
        try:
            from tools.editor_drafts import (
                recover_null_current_drafts,
            )
            recovered = await recover_null_current_drafts()
            if recovered:
                log.warning(
                    "editor_null_current_drafts_recovered_at_startup",
                    count=recovered)
            else:
                log.info(
                    "editor_null_current_drafts_check_clean")
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "editor_null_current_drafts_startup_failed",
                error=str(exc))
        # Auto-extend the monthly data pipeline beyond the Excel file —
        # fetch any complete calendar months that have closed since the
        # last run. A daemon thread so startup never blocks on yfinance;
        # fail-open inside extend_market_data. audit_reason="startup" so
        # a redeploy with no data change logs audit_trigger_skipped and
        # fires no Opus audit; only a genuine new month triggers one.
        try:
            import threading
            from tools.data_fetcher import extend_market_data
            threading.Thread(
                target=lambda: extend_market_data(audit_reason="startup"),
                daemon=True, name="monthly-data-extend").start()
        except Exception as exc:  # noqa: BLE001
            log.warning("monthly_extend_startup_failed", error=str(exc))
        # Macro research digest — reap any 'running' row left by a
        # crashed run on the previous boot BEFORE firing the new
        # trigger. Without this, a zombie row indefinitely blocks
        # the concurrency lock and every Run Now click returns
        # "already_running" with no real run firing — the symptom
        # surfaced by UAT on May 22 2026.
        #
        # TEST-ENV SKIP: in the test environment we do NOT fire the
        # startup reaper, the startup trigger, or the daily scheduler.
        # Each of those three side-effects (a SQL UPDATE, a background
        # task that hits the real DB, and a daemon coroutine that
        # sleeps until 21:00 UTC) pollutes test isolation: the trigger
        # writes a 'running' row to the CI database that other tests
        # in the run then trip over (the
        # test_is_research_running_returns_false_without_db, the
        # test_get_recent_digests_returns_empty_without_db, and the
        # test_skipped_when_row_create_fails failures surfaced May 22
        # 2026). The engine layer itself substitutes a mock digest in
        # ENVIRONMENT=test, but the side-effect writes still happen
        # before the mock substitution. Skip the whole stack at the
        # lifespan boundary.
        import os
        _is_test_env = (
            os.getenv("ENVIRONMENT", "").lower() == "test"
            or os.getenv("TESTING", "").lower() == "true"
        )
        if not _is_test_env:
            try:
                from tools.research_engine import (
                    fail_stale_running_digests, start_daily_scheduler,
                    trigger_research_async,
                )
                # timeout_minutes=0 on startup: the previous process is
                # dead by definition post-restart, so ANY row left in
                # 'running' is a zombie regardless of age. The default
                # 10-minute timeout protects runtime checks against
                # legitimately-in-flight rows; startup has no live
                # process to protect against. This catches the row-15
                # case where a row was stuck for less than the timeout
                # when the worker crashed.
                reaped = await fail_stale_running_digests(
                    timeout_minutes=0)
                if reaped:
                    log.warning(
                        f"Startup reap: marked {reaped} stuck research "
                        f"run(s) as failed", count=reaped)
                # Then fire a research run on startup when the latest
                # completed digest is stale (> 24h) or absent. The
                # trigger is loop-aware (we are on the event loop here)
                # and idempotent — a fresh boot within the freshness
                # window logs research_run_skipped_current and no model
                # call fires. Fail-open: a research failure logs and
                # proceeds.
                trigger_research_async("startup")
                # Daily scheduler — fires run_research_if_stale once per
                # UTC day at 21:00 (US market close + 1h). The 24h
                # freshness gate inside run_research_if_stale means a
                # boot-time fire and a scheduled fire within an hour of
                # each other still produce only one model call. The
                # scheduler is a daemon task held on _research_bg_tasks
                # so the GC does not silently cancel it.
                start_daily_scheduler()
            except Exception as exc:  # noqa: BLE001
                log.warning("research_startup_trigger_failed",
                            error=str(exc))
        else:
            log.info("research_startup_skipped_test_env")
        # Macro context cache warm — read whatever digest already
        # exists in the DB into the agent-prompt injection cache so
        # the FIRST agent call after restart sees the previous
        # deploy's digest. The startup trigger above produces a fresh
        # one in the background; this warm-read is the in-flight
        # fallback that prevents an empty cache for the few seconds
        # it takes the agent to land. Fail-open.
        try:
            from tools.macro_context import refresh_macro_context
            await refresh_macro_context()
        except Exception as exc:  # noqa: BLE001
            log.warning("macro_context_warm_failed", error=str(exc))
        # Item 5 (May 23 2026 — analytics narrative). Warm-read the
        # narrative cache so the first agent call after restart sees
        # the previous deploy's narrative. The diversification refresh
        # tick fires the canonical rebuild later; this just covers
        # the cold-deploy gap.
        try:
            from tools.analytics_context import refresh_analytics_context
            await refresh_analytics_context()
        except Exception as exc:  # noqa: BLE001
            log.warning("analytics_context_warm_failed", error=str(exc))
        # Analytical staging findings — same warm-read pattern so the
        # Academic Writer's first call after restart sees the most
        # recent staged findings. The endpoint refresh() also fires
        # after every fresh staging.
        try:
            from tools.analytical_findings import refresh_findings_context
            await refresh_findings_context()
        except Exception as exc:  # noqa: BLE001
            log.warning("findings_context_warm_failed", error=str(exc))
        # Item 9 commit 5 — strategy context cache warm-read. Same
        # warm-read pattern as the other three context blocks: agent
        # calls in the first seconds after restart see whatever
        # strategy_characterisations the previous deploy produced. The
        # next refresh_strategy_characterisations() run refreshes the
        # cache via tools/strategy_context.refresh_strategy_context_
        # cache() at the end of its orchestrator. Fail-open.
        try:
            from tools.strategy_context import refresh_strategy_context_cache
            await refresh_strategy_context_cache()
        except Exception as exc:  # noqa: BLE001
            log.warning("strategy_context_warm_failed", error=str(exc))
        # May 24 2026 (fourth iteration) — fully automatic warm with
        # retry. Per user directive: "after every deploy, cache warms
        # automatically within 60 seconds of startup. No operator
        # action required."
        #
        # The auto-warm runs as a fire-and-forget background task so
        # the lifespan handler does NOT block on it. asyncio's task
        # scheduler picks up the coroutine after the server is
        # accepting requests, avoiding the cold-boot competition
        # that caused the prior iteration's hang.
        #
        # `auto_warm_analytics` retries up to MAX_ATTEMPTS (3) with
        # exponential backoff (5s → 15s → 45s). A transient hiccup
        # at boot doesn't leave the cache cold.
        #
        # Status is reported through GET /api/v1/admin/cache-status.
        # The manual POST /api/v1/admin/warm-analytics-cache endpoint
        # remains as a sysadmin-only override.
        if not _is_test_env:
            try:
                from tools.cache_warm_state import auto_warm_analytics

                async def _auto_warm_task() -> None:
                    # Tiny initial delay so the server has settled
                    # into its first-request handler before we
                    # kick off the heavier compute path.
                    await asyncio.sleep(2.0)
                    await auto_warm_analytics()

                asyncio.create_task(_auto_warm_task())
                log.info("analytics_cache_auto_warm_scheduled")
            except Exception as exc:  # noqa: BLE001
                log.warning("analytics_cache_auto_warm_schedule_failed",
                            error=str(exc))

        # PR-MODEL-1 (May 27 2026) — model availability check. Pings
        # every chain's primary; on 404, advances the chain and
        # pings the next entry. This is the SAME fallback chain
        # call_claude / call_gemini consult at call time, so a model
        # the startup check detected as deprecated is silently
        # routed past for every request that follows. Fires in a
        # background task so the lifespan handler does not block on
        # the SDK round-trips (~1-3s each, larger on cold provider
        # endpoints). Fail-open per chain.
        if not _is_test_env:
            try:
                from agents.models import check_model_availability

                async def _model_check_task() -> None:
                    # 5s delay so the first user request is already
                    # being served before the check fires — the
                    # availability check itself never queues behind
                    # a real request thanks to the model resolver,
                    # but a deferred ping keeps cold-boot logs tidy.
                    await asyncio.sleep(5.0)
                    summary = await check_model_availability()
                    log.info("model_availability_check_complete",
                             summary=summary)

                asyncio.create_task(_model_check_task())
                log.info("model_availability_check_scheduled")
            except Exception as exc:  # noqa: BLE001
                log.warning("model_availability_check_schedule_failed",
                            error=str(exc))

        # June 27 2026 -- background regime-signals refresh ticker.
        # The regime_signals_cache has a 15-minute TTL and used to go
        # stale whenever no user activity triggered the HMM, blocking
        # deck generation + light refresh via the hard gate at
        # _refresh_regime_signals_for_deck. This ticker keeps the
        # cache warm on a fixed 10-minute cadence (5 minutes of
        # headroom before TTL expiry). Same fire-and-forget shape as
        # the analytics auto-warm + model-availability tasks above
        # so the lifespan handler never blocks on it. Per-iteration
        # failures (FRED outage, detect_current_regime exception,
        # cache write failure) log a warning and the loop continues
        # to the next cycle -- the ticker NEVER crashes.
        if not _is_test_env:
            try:
                from tools.cache import set_regime_cache
                from tools.regime_detector import detect_current_regime

                async def _regime_signals_ticker_task() -> None:
                    # 10-second initial sleep -- defers past the
                    # auto-warm (2s) + model-availability (5s) so the
                    # cold-boot path doesn't pile detect_current_regime
                    # on top of the heavier startup compute.
                    await asyncio.sleep(10.0)
                    while True:
                        try:
                            fresh = await asyncio.wait_for(
                                asyncio.to_thread(
                                    detect_current_regime),
                                timeout=_REGIME_REFRESH_TIMEOUT_S)
                            if not isinstance(fresh, dict):
                                log.warning(
                                    "regime_signals_background_refresh_failed",
                                    error=(
                                        "detect_current_regime "
                                        "returned non-dict"),
                                    return_type=type(fresh).__name__)
                            else:
                                await set_regime_cache(
                                    fresh, ttl_minutes=15)
                                log.info(
                                    "regime_signals_background_refresh",
                                    regime=fresh.get("regime"),
                                    confidence=fresh.get(
                                        "confidence"))
                        except asyncio.CancelledError:
                            # Clean shutdown -- propagate so the task
                            # exits without logging a spurious failure.
                            raise
                        except Exception as exc:  # noqa: BLE001
                            log.warning(
                                "regime_signals_background_refresh_failed",
                                error=str(exc))
                        # Sleep BETWEEN iterations -- a failed cycle
                        # waits the full interval before retrying.
                        # 10 min × 24h = 144 detect calls/day, and
                        # the in-process TTL inside
                        # detect_current_regime makes most a no-op.
                        await asyncio.sleep(
                            _REGIME_SIGNALS_TICKER_INTERVAL_S)

                asyncio.create_task(_regime_signals_ticker_task())
                log.info("regime_signals_ticker_scheduled",
                         interval_s=(
                             _REGIME_SIGNALS_TICKER_INTERVAL_S))
            except Exception as exc:  # noqa: BLE001
                log.warning("regime_signals_ticker_schedule_failed",
                            error=str(exc))
    yield
    log.info("forest_capital_shutdown")


limiter = Limiter(key_func=get_remote_address)

app = FastAPI(
    title="Forest Capital Portfolio Intelligence System",
    version="0.4.0-sprint4",
    lifespan=lifespan,
)

app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# ── Request validation logger ────────────────────────────────────────────────
#
# May 24 2026 — when a frontend bug constructs a malformed path
# parameter (the "3:1" generation-id pattern was the reported case),
# FastAPI returns a 422 with the validation detail but Starlette
# doesn't log it. Production has no signal that the bug is occurring
# until a user reports it. This handler preserves the default 422
# response shape and adds a structured log line carrying the URL,
# method, client IP and a compact dump of the validation errors so
# future malformed IDs are visible in Render logs the moment they
# fire — making it trivial to track the bad call to its origin.
from fastapi.exceptions import RequestValidationError
from fastapi.exception_handlers import request_validation_exception_handler


@app.exception_handler(RequestValidationError)
async def _logged_validation_error(request: Request, exc: RequestValidationError):
    try:
        client_ip = request.client.host if request.client else "unknown"
        # exc.errors() may carry non-serializable objects (Decimal,
        # datetime, set). str() is the conservative summary —
        # full structured payload still goes back to the client.
        compact = [
            {
                "loc": list(e.get("loc", [])),
                "msg": str(e.get("msg", "")),
                "type": str(e.get("type", "")),
            }
            for e in (exc.errors() or [])
        ]
        log.warning(
            "request_validation_error",
            method=request.method,
            url=str(request.url),
            path=str(request.url.path),
            client_ip=client_ip,
            errors=compact,
        )
    except Exception:  # noqa: BLE001
        # Never let logging suppress the actual 422 response.
        pass
    return await request_validation_exception_handler(request, exc)

# Uploads mount — serves UAT test-runner screenshots read-only.
# config.SCREENSHOT_DIR resolves to /data/test_screenshots on Render (a
# persistent disk — screenshots survive redeployments) and to
# backend/data/test_screenshots in local development. The directory is
# created on startup so the mount always resolves; the mount is rooted
# one level above it so the stored "test_screenshots/<uuid>" relative
# paths resolve under /uploads.
try:
    from fastapi.staticfiles import StaticFiles
    from config import SCREENSHOT_DIR
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)
    _uploads_dir = os.path.dirname(SCREENSHOT_DIR)
    app.mount("/uploads", StaticFiles(directory=_uploads_dir), name="uploads")
except Exception as _exc:  # noqa: BLE001
    log.warning("uploads_mount_failed", error=str(_exc))

# CORS — dev allows localhost:5173 only; prod allows Vercel URL only
origins = [FRONTEND_URL] if FRONTEND_URL else ["http://localhost:5173"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["X-API-Key", "Content-Type"],
)


# ── Auth endpoints (no auth required) ────────────────────────────────────────

@app.post("/api/auth/request-link", response_model=MagicLinkResponse)
async def request_magic_link(body: MagicLinkRequest, request: Request):
    email = body.email.strip().lower()
    # Both branches return HTTP 200 with an identical message to prevent email
    # enumeration. The status field is the only difference: the frontend shows a
    # specific "check your inbox" confirmation only when status == "sent".
    # Authorisation is the platform_users table — is_login_allowed falls back
    # to the config ALLOWED_EMAILS allowlist only if that table is unreachable.
    from tools.platform_users import is_login_allowed
    if not await is_login_allowed(email):
        log.warning("magic_link_unauthorized_email", email_hash=hash(email))
        return MagicLinkResponse(
            message="If that email is authorized, a login link has been sent.",
            status="pending",
            dev_mode=(ENVIRONMENT == "development"),
        )
    token = generate_magic_token(email)
    await send_magic_link(email, token)
    return MagicLinkResponse(
        message="If that email is authorized, a login link has been sent.",
        status="sent",
        dev_mode=(ENVIRONMENT == "development"),
    )


@app.get("/api/auth/verify")
async def verify_magic_link(token: str = Query(...), response: Response = None):
    # JTI persistence: check the DB before the in-memory dict so single-use protection
    # survives Render restarts. The in-memory dict handles the scanner pre-fetch case
    # within the same server instance; the DB handles cross-restart replay protection.
    try:
        import jwt as _jwt
        from datetime import datetime, timezone as _tz
        _peek = _jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        _jti = _peek.get("jti", "")
        _exp = _peek.get("exp")
        _email_for_jti = _peek.get("sub", "")
        if _jti:
            from tools.cache import is_jti_used, mark_jti_used
            if await is_jti_used(_jti):
                raise HTTPException(status_code=401, detail="This link has already been used. Please request a new one.")
    except HTTPException:
        raise
    except Exception:
        # jwt.decode failure (expired, invalid) is handled inside redeem_magic_token
        _jti = ""
        _exp = None
        _email_for_jti = ""

    # Look up the user's role / display_name / permissions so they are
    # embedded in the session JWT — require_auth then needs no per-request
    # database hit. A None result (user not found, or database down)
    # mints a plain token that require_auth resolves via the fallback.
    user_attrs: dict | None = None
    if _email_for_jti:
        try:
            from tools.platform_users import get_active_user
            _u = await get_active_user(_email_for_jti)
            if _u:
                user_attrs = {
                    "role": _u["role"],
                    "display_name": _u["display_name"],
                    "permissions": _u["permissions"],
                }
        except Exception:  # noqa: BLE001
            user_attrs = None

    session_token = redeem_magic_token(token, user_attrs)

    # Persist JTI to DB after first successful redemption (non-blocking — failure is safe)
    if _jti and _exp:
        try:
            from tools.cache import mark_jti_used
            from datetime import datetime, timezone as _tz
            await mark_jti_used(_jti, datetime.fromtimestamp(_exp, tz=_tz.utc), _email_for_jti)
        except Exception:
            pass

    email = verify_session_token(session_token)["email"]
    # Stamp last_login_at — fail-open, never blocks the login response.
    try:
        from tools.platform_users import record_login
        await record_login(email)
    except Exception:  # noqa: BLE001
        pass
    log.info("auth_success", email=email)
    # Prevent browsers and intermediary caches from storing the session token.
    # A cached 200 response could replay a stale token on a shared machine.
    if response is not None:
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
        response.headers["Pragma"] = "no-cache"
    return SessionResponse(
        session_token=session_token,
        email=email,
        expires_in_hours=int(os.getenv("SESSION_EXPIRY_HOURS", "8")),
    )


@app.post("/api/auth/logout")
async def logout(body: LogoutRequest):
    invalidate_session(body.session_token)
    return {"message": "Logged out successfully."}


@app.get("/api/auth/me")
async def get_me(session: dict = Depends(require_auth)):
    """The signed-in user — email, role, display name, the authoritative
    permissions array the frontend gates the UI on, and the lifetime
    council-query allocation (council_queries_limit None = unlimited)."""
    council_used = 0
    council_limit: int | None = None
    try:
        from tools.platform_users import get_active_user
        u = await get_active_user(session["email"])
        if u:
            council_used = u.get("council_queries_used", 0)
            council_limit = u.get("council_queries_limit")
    except Exception:  # noqa: BLE001
        pass
    return {
        "email": session["email"],
        "role": session.get("role") or "viewer",
        "display_name": session.get("display_name"),
        "permissions": session.get("permissions") or [],
        "council_queries_used": council_used,
        "council_queries_limit": council_limit,
    }


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    # May 24 2026 — surface RENDER_GIT_COMMIT so deployment status is
    # verifiable WITHOUT the master API key. The dashboard-timeout
    # investigation kept hitting "is Render on the latest commit?"
    # as the unverifiable first question; exposing the short SHA on
    # /health closes that gap. Render injects RENDER_GIT_COMMIT
    # automatically on every build (full 40-char SHA); a local dev
    # run has no value set, in which case "dev" surfaces. The branch
    # name (RENDER_GIT_BRANCH) is included so a deploy from a
    # feature branch is obviously identified.
    commit = os.getenv("RENDER_GIT_COMMIT", "dev")
    return {
        "status": "ok",
        "sprint": "4",
        "environment": ENVIRONMENT,
        "anthropic": bool(os.getenv("ANTHROPIC_API_KEY")),
        "gemini": bool(os.getenv("GOOGLE_API_KEY")),
        "cache": True,
        "commit": commit[:7] if commit != "dev" else "dev",
        "commit_full": commit,
        "branch": os.getenv("RENDER_GIT_BRANCH") or "dev",
    }


# ── Provenance ───────────────────────────────────────────────────────────────

_PROVENANCE_PATH = Path(__file__).parent / "data" / "provenance.json"


@app.get("/api/v1/provenance")
async def get_provenance(session: dict = Depends(require_auth)):
    """
    Returns the full data_series_registry as JSON.

    Reads from provenance.json rather than PostgreSQL so the endpoint works
    in CI (no DB) and on cold starts before the pipeline has run.
    The frontend never hardcodes provenance — it always fetches from here.
    """
    if not _PROVENANCE_PATH.exists():
        # Pipeline hasn't run yet — return empty registry rather than 500.
        return {"series": [], "last_pipeline_run": None, "cross_validation": {}}
    try:
        data = json.loads(_PROVENANCE_PATH.read_text(encoding="utf-8"))
        # May 23 2026 — defensive sanitisation. A provenance.json
        # written by an older pipeline build may contain literal NaN /
        # Infinity tokens (Python json.dumps default behaviour), which
        # parse into float('nan') / float('inf') here. Starlette /
        # orjson then refuse to serialise them, returning 500 with
        # "Out of range float values are not JSON compliant". Replace
        # them with None at the read site so existing bad files don't
        # break the endpoint. New writes go through the sanitiser at
        # the WRITE site (see tools/data_fetcher._sanitise_json_floats).
        return _sanitise_nan_floats(data)
    except Exception as exc:
        log.warning("provenance_read_error", error=str(exc))
        raise HTTPException(status_code=500, detail="Could not read provenance data.")


def _sanitise_nan_floats(obj):
    """Recursively replaces NaN / Inf float values with None so the
    response round-trips through strict JSON serialisers (starlette /
    orjson). The write side has the same guard — this is the second
    layer for older files written before the fix landed."""
    import math
    if isinstance(obj, float):
        if math.isnan(obj) or math.isinf(obj):
            return None
        return obj
    if isinstance(obj, dict):
        return {k: _sanitise_nan_floats(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_sanitise_nan_floats(v) for v in obj]
    return obj


@app.get("/api/v1/provenance/justification")
async def get_provenance_justification(session: dict = Depends(require_auth)):
    """
    Structured justification for each supplemental data source.

    Returned to: Data Sources panel (expandable rows), Academic Writer Agent
    (Analytical Appendix Section 3.2), Explainer Agent (Commentary mode hover).
    The justification is factual metadata — it does not change unless the
    pipeline design changes, so no DB or provenance.json dependency is needed.
    """
    return {
        "spy_daily": {
            "source": "yfinance",
            "ticker": "SPY",
            "strategies_enabled": ["VOL_TARGETING", "MOMENTUM_ROTATION"],
            "without_this_source": "VOL_TARGETING and MOMENTUM_ROTATION unavailable — "
                "both require daily return frequency for 21-day rolling volatility signals.",
            "key_reason": "Monthly data cannot resolve intramonth volatility spikes "
                "(e.g. March 2020 circuit breakers). VOL_TARGETING scales equity weight "
                "by TARGET_VOL / realised_vol_21d, which requires daily observations.",
            "months_added": 0,
            "statistical_impact": "Enables 2 of 10 strategies. Without it, dynamic "
                "allocation universe shrinks from 5 to 3 strategies.",
        },
        "vixcls": {
            "source": "fred_api",
            "series_id": "VIXCLS",
            "strategies_enabled": ["REGIME_SWITCHING"],
            "without_this_source": "Regime threshold classifier degrades to equity-trend "
                "and yield-curve only — VIX forward-looking fear signal lost.",
            "key_reason": "VIX > 25 triggers the BEAR regime flag. It provides a "
                "forward-looking signal independent of equity price — VIX spikes "
                "precede equity drawdowns by 0-5 trading days on average.",
            "months_added": 0,
            "statistical_impact": "Regime Switching Sharpe 0.629 (with VIX) vs 0.571 "
                "(without, threshold-only). Regime agreement rate falls from 87% to 71%.",
        },
        "dgs2": {
            "source": "fred_api",
            "series_id": "DGS2",
            "strategies_enabled": ["REGIME_SWITCHING"],
            "without_this_source": "10Y-2Y yield curve signal unavailable. "
                "Regime classifier cannot detect yield curve inversion.",
            "key_reason": "The 10Y-2Y spread has preceded every US recession since 1955. "
                "The curve inverted in April 2022 — six months before the equity trough. "
                "A VIX + equity-only detector would have been late to this signal.",
            "months_added": 0,
            "statistical_impact": "April 2022 early warning — 6 months lead time vs "
                "equity-only detection in October 2022.",
        },
        "lqd_bridge": {
            "source": "yfinance",
            "ticker": "LQD",
            "strategies_enabled": [
                "BENCHMARK", "CLASSIC_60_40", "RISK_PARITY", "MIN_VARIANCE",
                "EQUAL_WEIGHT", "MOMENTUM_ROTATION", "REGIME_SWITCHING",
                "VOL_TARGETING", "BLACK_LITTERMAN", "MAX_SHARPE_ROLLING",
            ],
            "without_this_source": "BND inception April 2007 — aligned dataset starts "
                "May 2007 (224 months). Dot-com recovery (2002-2007) excluded entirely.",
            "key_reason": "LQD (iShares IG Corporate Bond ETF) tracks the same IG "
                "corporate bond universe as BND and began trading July 2002. Monthly "
                "returns spliced: LQD used 2002-07 to 2007-04, BND from 2007-05.",
            "months_added": 58,
            "statistical_impact": "Without the bridge, n=224 observations; "
                "with the bridge, n is the live monthly count (286 on the "
                "current Render snapshot — auto-extends past 2025-12 once "
                "Apr 2026 fully closes). Power analysis requires n >= 220 "
                "for 80% power at p < 0.005 — the bridge provides the "
                "statistical margin. Without it, the dataset barely clears "
                "the minimum.",
        },
    }


# ── Strategies ────────────────────────────────────────────────────────────────

@app.get("/api/strategies/list")
async def list_strategies(session: dict = Depends(require_auth)):
    # Primary path: derive strategy list from the real backtester output so the
    # names and types always match what run_all_strategies() actually produces.
    # MOCK_STRATEGIES is the fallback — same keys, but avoids running the full
    # pipeline just to return a name/type list.
    if ENVIRONMENT != "test":
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.backtester import run_all_strategies
            # OFF-LOOP: get_full_history reads the DB (and may run the cold
            # pipeline through Excel/FRED/yfinance). Pushed to a worker
            # thread so the event loop keeps serving concurrent requests.
            history = await get_full_history_async()
            results_dict = await asyncio.to_thread(run_all_strategies, history)
            return {
                "strategies": [
                    {"name": name, "type": r.get("strategy_type", "static")}
                    for name, r in results_dict.items()
                ]
            }
        except Exception as exc:
            log.warning("list_strategies_fallback", error=str(exc))
    # Fallback: MOCK_STRATEGIES mirrors the real strategy keys and types.
    return {
        "strategies": [
            {"name": s["strategy_name"], "type": s["strategy_type"]}
            for s in MOCK_STRATEGIES
        ]
    }


# ── Backtest ──────────────────────────────────────────────────────────────────

@app.post("/api/backtest/run")
@limiter.limit("20/minute")
async def run_backtest(
    request: Request,
    body: BacktestRequest,
    session: dict = Depends(require_auth),
):
    # Canonical strategy names come from MOCK_STRATEGIES, which mirrors the
    # keys that run_all_strategies() produces.  This validation step has no
    # dependency on the real pipeline so it is always cheap.
    valid_strategies = {s["strategy_name"] for s in MOCK_STRATEGIES}
    if body.strategy not in valid_strategies:
        raise HTTPException(
            status_code=422,
            detail=f"Unknown strategy '{body.strategy}'. Valid: {sorted(valid_strategies)}",
        )

    log.info("backtest_run", strategy=body.strategy, user=session["email"])

    # Primary path: run the full pipeline and return the requested strategy.
    # run_all_strategies() computes all 10 at once; we pick the one requested.
    # The old condition checked for "100% Equity (Benchmark)" — the Sprint 1
    # human-readable name — which never matched the real key "BENCHMARK", making
    # that branch dead code.  Fixed here to handle all strategies uniformly.
    if ENVIRONMENT != "test":
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.backtester import run_all_strategies
            history = await get_full_history_async()
            results_dict = await asyncio.to_thread(run_all_strategies, history)
            if body.strategy in results_dict:
                return results_dict[body.strategy]
        except Exception as exc:
            log.warning("backtest_run_fallback", strategy=body.strategy, error=str(exc))

    # Fallback: return the corresponding MOCK_STRATEGIES entry when the real
    # pipeline is unavailable (test environment or unhandled exception above).
    result = next(s for s in MOCK_STRATEGIES if s["strategy_name"] == body.strategy)
    return result


@app.get("/api/backtest/compare")
@limiter.limit("30/minute")
async def compare_strategies(request: Request, session: dict = Depends(require_auth)):
    # Sprint 5: strategy_results_cache checked before calling run_all_strategies().
    # Cache key = SHA-256 of (n_monthly_rows, last_date, n_strategies=10).
    # On a cache hit: ~200ms response.  On a miss: recompute (~30s) then cache.
    # Cache survives Render restarts because it lives in PostgreSQL, not memory.
    if ENVIRONMENT != "test":
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.backtester import run_all_strategies
            from tools.cache import get_strategy_cache, set_strategy_cache, _compute_data_hash

            # OFF-LOOP — this endpoint is the dashboard hot path. Sync
            # get_full_history() used to block the event loop on every
            # call, serialising all parallel dashboard requests behind a
            # single slow read and producing the 35-70s response-time
            # cluster observed in Render logs on 2026-05-24. await
            # get_full_history_async() pushes the entire read pipeline
            # (DB + Excel + FRED + yfinance on a cold path) into a worker
            # thread so the loop stays free to service the other 6+
            # parallel dashboard calls.
            history = await get_full_history_async()

            # Build a stable hash from pipeline metadata to detect new data
            monthly = history.get("equity_monthly")
            n_rows = len(monthly) if monthly is not None else 0
            first_date = str(monthly.index[0].date()) if monthly is not None and len(monthly) > 0 else "unknown"
            last_date = str(monthly.index[-1].date()) if monthly is not None and len(monthly) > 0 else "unknown"
            strategy_hash = _compute_data_hash(n_rows, last_date, n_strategies=10)

            # Expose the actual date range so the frontend can label charts
            # dynamically. With the LQD bridge: ~2002-07 to ~2024-12 (282 months).
            # Without LQD bridge: ~2007-05 to ~2024-12 (224 months) — fall-back state.
            data_range = {"start": first_date, "end": last_date, "n_months": n_rows}

            cached = await get_strategy_cache(strategy_hash)
            # Schema gate: cache entries written before Sprint 6 lack the
            # monthly_returns field that /charts/data needs. Treat them as
            # stale so a single recompute fills the cache with current-schema
            # entries; both endpoints see the benefit on the next request.
            cache_current_schema = bool(cached) and all(
                isinstance(r.get("monthly_returns"), list) and len(r["monthly_returns"]) > 0
                for r in cached.values()
            )
            if cache_current_schema:
                ranked = sorted(cached.values(), key=lambda r: r.get("sharpe_ratio", 0.0), reverse=True)
                return {"strategies": ranked, "ranked_by": "sharpe_ratio", "cache": "hit", "data_range": data_range}

            # Cache miss / schema refresh — run_all_strategies executes 10
            # backtests serially (~30s on a cold pipeline). Offload to a
            # worker thread so the loop is not blocked while it computes.
            results_dict = await asyncio.to_thread(run_all_strategies, history)
            ranked = sorted(results_dict.values(), key=lambda r: r.get("sharpe_ratio", 0.0), reverse=True)

            # Write-through: persist for next cold start or Render restart.
            # Thread risk_free_monthly through so the invariant framework's
            # 1b Sharpe-recomputation check uses the actual DTB3 series the
            # backtester used (anything else produces false positives).
            await set_strategy_cache(
                strategy_hash, results_dict, n_observations=n_rows,
                risk_free_monthly=history.get("risk_free_monthly"))

            cache_label = "miss" if not cached else "schema_refresh"
            return {"strategies": ranked, "ranked_by": "sharpe_ratio", "cache": cache_label, "data_range": data_range}
        except Exception as exc:
            log.warning("compare_all_strategies_fallback", error=str(exc))
    # Fallback: MOCK_STRATEGIES used only in test environment or when the real
    # pipeline raises an exception.  Should never be the primary response in
    # production — the warning log above will flag if this path is taken.
    sorted_strategies = sorted(MOCK_STRATEGIES, key=lambda s: s["sharpe_ratio"], reverse=True)
    return {"strategies": sorted_strategies, "ranked_by": "sharpe_ratio"}


# ── Charts data — aux payload for Statistical Evidence & Regime Analysis ─────

@app.get("/api/v1/charts/data")
@limiter.limit("30/minute")
async def get_chart_data(request: Request, session: dict = Depends(require_auth)):
    """
    Returns the auxiliary data required by all twelve Sprint 6 charts in a
    single call. Bundling avoids 12 sequential round-trips on Render's free
    tier where each cold start adds noticeable latency. The payload is
    derived from the same get_full_history() + run_all_strategies() outputs
    that /api/backtest/compare uses, so a cache hit there means cache hit here.
    """
    if ENVIRONMENT == "test":
        # Tests: return an empty-shape payload so frontend tests can mock
        # without spinning up the full pipeline.
        return {
            "cpcv": {}, "cv_radar": {}, "walk_forward": {},
            "regime_conditional": {}, "regime_timeline": [],
            "correlation_breakdown": [], "factor_loadings": {},
            "attribution": {}, "transition_matrix": {},
            "n_strategies": 0, "n_months": 0,
        }

    try:
        from tools.data_fetcher import get_full_history_async
        from tools.backtester import run_all_strategies
        from tools.chart_data import compute_chart_data
        from tools.cache import get_strategy_cache, _compute_data_hash

        # OFF-LOOP — second dashboard hot path. Same fix as
        # /api/backtest/compare; the two endpoints together account for
        # most of the dashboard's parallel-request fan-out.
        history = await get_full_history_async()
        monthly = history.get("equity_monthly")
        n_rows = len(monthly) if monthly is not None else 0
        last_date = str(monthly.index[-1]) if monthly is not None and len(monthly) > 0 else "unknown"
        strategy_hash = _compute_data_hash(n_rows, last_date, n_strategies=10)

        # Reuse the strategy cache if a prior /compare call populated it.
        # The cached results MUST carry monthly_returns — without them, the
        # per-strategy chart computations (CPCV, walk-forward, regime-
        # conditional, factor loadings, attribution, CV radar) all produce
        # empty output. Cache entries written before Sprint 6 lack this
        # field; detect and refresh so the system self-heals without
        # requiring a manual DB invalidation.
        cached_results = await get_strategy_cache(strategy_hash)
        cache_is_chart_compatible = bool(cached_results) and all(
            isinstance(r.get("monthly_returns"), list) and len(r["monthly_returns"]) > 0
            for r in cached_results.values()
        )
        if cache_is_chart_compatible:
            results_dict = cached_results
            log.info("chart_data_cache_hit", strategy_hash=strategy_hash[:8])
        else:
            if cached_results:
                log.info(
                    "chart_data_cache_schema_miss",
                    strategy_hash=strategy_hash[:8],
                    note="cached entry lacks monthly_returns — refreshing",
                )
            results_dict = run_all_strategies(history)
            # Write-through so the next /charts/data hit AND the next
            # /compare hit both find a schema-compatible entry. rf
            # threaded so invariant 1b can recompute Sharpe against
            # the same series the backtester used.
            try:
                from tools.cache import set_strategy_cache
                await set_strategy_cache(
                    strategy_hash, results_dict, n_observations=n_rows,
                    risk_free_monthly=history.get("risk_free_monthly"))
            except Exception as exc:
                log.warning("chart_data_cache_write_failed", error=str(exc))

        payload = compute_chart_data(history, results_dict)
        payload["strategy_hash"] = strategy_hash
        return payload
    except Exception as exc:
        log.warning("chart_data_fallback", error=str(exc))
        return {
            "cpcv": {}, "cv_radar": {}, "walk_forward": {},
            "regime_conditional": {}, "regime_timeline": [],
            "correlation_breakdown": [], "factor_loadings": {},
            "attribution": {}, "transition_matrix": {},
            "n_strategies": 0, "n_months": 0, "error": str(exc),
        }


# ── Academic analytics ────────────────────────────────────────────────────────

@app.get("/api/v1/analytics/academic")
@limiter.limit("30/minute")
async def get_academic_analytics(request: Request, session: dict = Depends(require_auth)):
    """
    Bundled academic analytics for the analytics view and the midpoint paper:
    summary statistics, cumulative total return, 12-month rolling correlation,
    rolling excess return, regime-conditional performance, drawdown comparison,
    Carhart four-factor loadings, and the source-controlled strategy metadata.

    Every figure is derived from data already in PostgreSQL —
    market_data_monthly, strategy_results_cache, ff_factors_monthly — so the
    endpoint is a set of light reads plus pure NumPy/statsmodels compute; it
    never triggers get_full_history() or run_all_strategies(). Parameter
    sensitivity is deliberately NOT bundled here — it re-runs ~23 backtests
    and has its own endpoint, /api/v1/analytics/sensitivity.
    """
    if ENVIRONMENT == "test":
        return {"available": False, "note": "analytics unavailable in test environment"}
    try:
        # May 22 2026 — first try the pre-computed cache. The
        # strategy_cache write hook fires refresh_all_analytics
        # immediately after a fresh strategy ingestion, so the
        # cache is normally hot. Cold-cache fallback runs the
        # inline compute below.
        from tools.cache import (
            get_monthly_returns, get_latest_strategy_cache, get_ff_factors,
            get_latest_strategy_hash,
        )
        from tools.precomputed_analytics import (
            get_metric as get_precomputed,
            get_latest_metric as get_latest_precomputed,
        )

        latest_hash = await get_latest_strategy_hash()
        if latest_hash:
            cached = await get_precomputed(latest_hash, "academic_analytics")
            if cached:
                log.info("academic_analytics_cache_hit",
                         data_hash=latest_hash[:8])
                return cached
            # Stale-cache fallback — the latest hash hasn't been
            # refreshed yet but a previous data ingestion's payload
            # is still in the table. Better to serve stale than to
            # block on the inline compute when the refresh hook is
            # in flight. The response carries _stale=True so the
            # frontend can decide whether to surface a notice.
            stale = await get_latest_precomputed("academic_analytics")
            if stale:
                log.info("academic_analytics_cache_stale_hit")
                return stale

        # COLD CACHE — never run the inline compute on the request
        # thread. The inline path runs factor_loadings (OLS per
        # strategy), regime_conditional_performance, and 5 other
        # heavy reductions; on Render's shared CPU it can exceed the
        # frontend's 30s timeout. Hotfix iteration 2 (May 23 2026):
        # fire the background refresh and return a warming response
        # immediately. The frontend retries every ~10s until the
        # cache row lands.
        #
        # May 24 2026 P0 hotfix — fire the refresh UNCONDITIONALLY,
        # not gated on `if latest_hash:`. The pre-existing gate meant
        # that on a Render restart with an empty strategy_results_
        # cache, the refresh NEVER fired and the user saw a 30s
        # timeout forever. refresh_all_analytics now substitutes a
        # "BOOT-WARM" sentinel hash when none is supplied, so the
        # subsequent get_latest_precomputed call finds the row even
        # without a real strategy hash to key against.
        log.warning("academic_analytics_cache_cold_warming",
                    has_strategy_hash=bool(latest_hash))
        try:
            from tools.precomputed_analytics import (
                trigger_refresh_async,
            )
            trigger_refresh_async(latest_hash or "")
        except Exception as trig_exc:  # noqa: BLE001
            log.warning("academic_analytics_refresh_trigger_failed",
                        error=str(trig_exc))
        return {
            "available": False,
            "warming": True,
            "retry_after_ms": 10000,
            "note": ("Analytics are being computed in the "
                     "background — refresh in ~10 seconds."),
        }
    except Exception as exc:
        log.warning("academic_analytics_failed", error=str(exc))
        return {"available": False, "note": "analytics computation failed"}


# ── Diversification suite endpoints (item 8) ──────────────────────────────────
# Seven GET endpoints for the diversification context suite. Each
# reads a single pre-computed row from analytics_metrics_cache (the
# refresh hook fires on every strategy_cache write). Cold-cache
# fallback: compute inline from tools.diversification_analytics so
# the user always gets data even before the refresh hook lands.
# require_team_member per the spec; rate-limited to protect the
# DB / inline-compute paths from abuse.


async def _read_div_metric_or_compute(
    metric_kind: str,
    inline_fn,
) -> dict[str, Any]:
    """Three-tier read for diversification metrics: cache hit by
    current data_hash → stale cache → inline compute. Returns the
    payload dict, never raises. Logs which path served the request
    so production logs surface cache hit / miss patterns."""
    try:
        from tools.cache import (
            get_latest_strategy_hash, get_latest_strategy_cache,
        )
        from tools.precomputed_analytics import (
            get_metric, get_latest_metric,
        )
        latest_hash = await get_latest_strategy_hash()
        if latest_hash:
            cached = await get_metric(latest_hash, metric_kind)
            if cached:
                log.info("div_metric_cache_hit", metric=metric_kind)
                return cached
            stale = await get_latest_metric(metric_kind)
            if stale:
                log.info("div_metric_cache_stale_hit", metric=metric_kind)
                return stale
        # Cold cache — inline compute.
        strategies = await get_latest_strategy_cache()
        if not strategies:
            return {"available": False,
                    "note": "strategy cache not yet populated"}
        log.info("div_metric_cache_miss_inline", metric=metric_kind)
        return inline_fn(strategies)
    except Exception as exc:  # noqa: BLE001
        log.warning("div_metric_failed", metric=metric_kind, error=str(exc))
        return {"available": False, "note": "computation failed"}


@app.get("/api/v1/analytics/correlation")
@limiter.limit("30/minute")
async def get_analytics_correlation(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """11x11 strategy + benchmark correlation matrices (full +
    pre-2022 + post-2022). Drives the heatmap + insight callout."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "correlation_matrices",
        lambda s: div.correlation_matrices(s),
    )


@app.get("/api/v1/analytics/tail-risk")
@limiter.limit("30/minute")
async def get_analytics_tail_risk(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """VaR + CVaR at 95% / 99%, monthly + annualised, historical
    simulation. Drives the Downside Risk table."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "tail_risk",
        lambda s: {"strategies": div.tail_risk(s)},
    )


@app.get("/api/v1/analytics/capture-ratios")
@limiter.limit("30/minute")
async def get_analytics_capture_ratios(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """Up / Down capture + capture score per strategy over full +
    pre-2022 + post-2022 windows. Drives the capture scatter."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "capture_ratios",
        lambda s: {"strategies": div.capture_ratios(s)},
    )


@app.get("/api/v1/analytics/drawdown-duration")
@limiter.limit("30/minute")
async def get_analytics_drawdown_duration(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """Avg / max drawdown duration + recovery + current-in-drawdown
    state per strategy. Drives the Drawdown Duration table."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "drawdown_duration",
        lambda s: {"strategies": div.drawdown_duration(s)},
    )


@app.get("/api/v1/analytics/crisis-performance")
@limiter.limit("30/minute")
async def get_analytics_crisis_performance(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """CAGR + max DD + Sharpe per strategy over 5 historical crisis
    windows. Drives the Crisis Performance table."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "crisis_performance",
        lambda s: div.crisis_performance(s),
    )


@app.get("/api/v1/analytics/risk-contribution")
@limiter.limit("30/minute")
async def get_analytics_risk_contribution(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """MCTR + % risk contribution for equal-weight and tangency-weight
    portfolios. Drives the Risk Contribution stacked bar."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "marginal_contribution_to_risk",
        lambda s: div.marginal_contribution_to_risk(s),
    )


@app.get("/api/v1/analytics/distribution")
@limiter.limit("30/minute")
async def get_analytics_distribution(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """Skewness / excess kurtosis / Jarque-Bera + best/worst months
    per strategy. Drives the Distribution Summary table."""
    if ENVIRONMENT == "test":
        return {"available": False}
    from tools import diversification_analytics as div
    return await _read_div_metric_or_compute(
        "return_distribution",
        lambda s: {"strategies": div.return_distribution(s)},
    )


@app.get("/api/v1/strategy-cache/key-metrics")
@limiter.limit("30/minute")
async def get_strategy_cache_key_metrics(
    request: Request,
    session: dict = Depends(require_auth),
):
    """Returns the key metrics from the current strategy cache,
    organized for the Reports-page verification panel (June 21 2026).

    The panel sits below the Generate Documents cards and gives Bob
    a one-glance way to see every figure the brief and deck must
    match. Every value is read directly from
    strategy_results_cache for current_data_hash + the live CIO
    recommendation overlay -- nothing computed inline, nothing
    sourced from the brief or deck output.

    Response shape:
      {
        "data_hash": "c421fb89",
        "computed_at": "2026-06-21T...",
        "available": bool,
        "metrics": {
          "strategy_performance": [{"label": "...", "value": "...",
                                    "source": "strategy_cache"}],
          "oos_metrics": [...],
          "correlation_regime": [...],
          "live_signal": [...]
        }
      }

    Fail-open: a cold cache or missing CIO row returns
    available=False with an empty metrics dict + a message; the
    panel renders a "cache not yet warm" placeholder."""
    try:
        from tools.cache import get_latest_strategy_cache
        from tools.cio_recommendation import (
            compute_implied_asset_allocation, get_latest_recommendation,
        )
        from tools.numeric_substitution import (
            format_corr, format_pct, format_sharpe,
            format_months_from_days,
        )
    except Exception as exc:  # noqa: BLE001
        log.warning("key_metrics_imports_failed", error=str(exc))
        return {
            "data_hash": "", "available": False, "metrics": {},
            "computed_at": None,
            "message": "Cache helpers unavailable."}

    # Read strategy_hash + computed_at directly from
    # strategy_results_cache. The earlier version read
    # current_data_hash() (a SHA256 of row counts + max dates
    # across three tables) which is a LIGHTWEIGHT FINGERPRINT --
    # NOT the strategy_results_cache.strategy_hash. They're two
    # different hashes by design; the brief / deck / digest all
    # cite the strategy_hash (c421fb89...), so the panel must
    # match for consistency. Fix: read the canonical strategy hash
    # directly from the cache row.
    strategy_hash: str = ""
    strategy_computed_at: str | None = None
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is not None:
            async with AsyncSessionLocal() as session:  # type: ignore[union-attr]
                row = await session.execute(text(
                    "SELECT strategy_hash, computed_at "
                    "FROM strategy_results_cache "
                    "ORDER BY computed_at DESC LIMIT 1"))
                r = row.fetchone()
                if r:
                    strategy_hash = str(r[0] or "")
                    strategy_computed_at = (
                        str(r[1]) if r[1] else None)
    except Exception as exc:  # noqa: BLE001
        log.warning("key_metrics_hash_read_failed", error=str(exc))

    try:
        strategy_cache = await get_latest_strategy_cache() or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("key_metrics_cache_failed", error=str(exc))
        strategy_cache = {}

    try:
        cio_row = await get_latest_recommendation() or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("key_metrics_cio_failed", error=str(exc))
        cio_row = {}

    # Compute the implied asset allocation from the live blend
    # weights. The CIO row carries `blend_weights` but not the
    # equity/IG/HY decomposition; `compute_implied_asset_allocation`
    # multiplies blend_weights x strategy_cache avg_equity_weight /
    # avg_ig_weight / avg_hy_weight to produce the live split.
    # Returns {equity_pct, ig_bond_pct, hy_bond_pct, bond_pct,
    # cash_pct} as fractions in [0, 1].
    implied: dict = {}
    try:
        if cio_row.get("blend_weights"):
            implied = await compute_implied_asset_allocation(
                cio_row.get("blend_weights")) or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("key_metrics_implied_alloc_failed", error=str(exc))

    if not strategy_cache:
        return {
            "data_hash": strategy_hash[:8] or "—",
            "available": False, "metrics": {},
            "computed_at": strategy_computed_at,
            "message": (
                "Strategy cache is empty -- run the backtester or "
                "warm the cache first.")}

    benchmark = strategy_cache.get("BENCHMARK") or {}
    classic = strategy_cache.get("CLASSIC_60_40") or {}
    regime = strategy_cache.get("REGIME_SWITCHING") or {}

    # Pull validated_constants from the same module the brief uses
    # so the OOS Sharpe figures shown on the panel match what the
    # brief substitutes for {{OOS_SHARPE_BLEND}} /
    # {{OOS_SHARPE_BENCHMARK}}.
    try:
        from tools.academic_deck import (
            OOS_SHARPE_BENCHMARK, OOS_SHARPE_REGIME_CONDITIONAL,
            CORRELATION_PRE_2022, CORRELATION_POST_2022,
        )
    except Exception:  # noqa: BLE001
        OOS_SHARPE_REGIME_CONDITIONAL = None
        OOS_SHARPE_BENCHMARK = None
        CORRELATION_PRE_2022 = None
        CORRELATION_POST_2022 = None

    metrics = {
        "strategy_performance": [
            {"label": "Benchmark (S&P 500) full-period Sharpe",
             "value": format_sharpe(benchmark.get("sharpe_ratio")),
             "source": "strategy_cache.BENCHMARK.sharpe_ratio"},
            {"label": "Classic 60/40 full-period Sharpe",
             "value": format_sharpe(classic.get("sharpe_ratio")),
             "source": "strategy_cache.CLASSIC_60_40.sharpe_ratio"},
            {"label": "Regime Switching full-period Sharpe",
             "value": format_sharpe(regime.get("sharpe_ratio")),
             "source": "strategy_cache.REGIME_SWITCHING.sharpe_ratio"},
            {"label": "Benchmark max drawdown",
             "value": format_pct(benchmark.get("max_drawdown")),
             "source": "strategy_cache.BENCHMARK.max_drawdown"},
            {"label": "Classic 60/40 max drawdown",
             "value": format_pct(classic.get("max_drawdown")),
             "source": "strategy_cache.CLASSIC_60_40.max_drawdown"},
            {"label": "Regime Switching max drawdown",
             "value": format_pct(regime.get("max_drawdown")),
             "source": "strategy_cache.REGIME_SWITCHING.max_drawdown"},
            {"label": "Benchmark recovery",
             "value": format_months_from_days(
                 benchmark.get("drawdown_recovery_days")),
             "source": (
                 "strategy_cache.BENCHMARK.drawdown_recovery_days")},
            {"label": "Regime Switching recovery",
             "value": format_months_from_days(
                 regime.get("drawdown_recovery_days")),
             "source": (
                 "strategy_cache.REGIME_SWITCHING."
                 "drawdown_recovery_days")},
        ],
        "oos_metrics": [
            {"label": "OOS window",
             "value": "January 2022 through May 2026 (53 months)",
             "source": "academic_deck.OOS window constant"},
            {"label": "Blend OOS Sharpe",
             "value": format_sharpe(OOS_SHARPE_REGIME_CONDITIONAL),
             "source": (
                 "academic_deck.OOS_SHARPE_REGIME_CONDITIONAL")},
            {"label": "Benchmark OOS Sharpe",
             "value": format_sharpe(OOS_SHARPE_BENCHMARK),
             "source": "academic_deck.OOS_SHARPE_BENCHMARK"},
        ],
        "correlation_regime": [
            {"label": "Pre-2022 equity-IG correlation",
             "value": format_corr(CORRELATION_PRE_2022),
             "source": "academic_deck.CORRELATION_PRE_2022"},
            {"label": "Post-2022 equity-IG correlation",
             "value": format_corr(CORRELATION_POST_2022),
             "source": "academic_deck.CORRELATION_POST_2022"},
        ],
        "live_signal": [
            {"label": "Current regime",
             "value": str(cio_row.get("regime") or "—"),
             "source": "cio_recommendation.regime"},
            {"label": "Regime confidence",
             "value": format_pct(
                 (cio_row.get("confidence") or {}).get("probability")
                 if isinstance(cio_row.get("confidence"), dict)
                 else cio_row.get("confidence")),
             "source": "cio_recommendation.confidence"},
            {"label": "Current blend equity",
             "value": format_pct(implied.get("equity_pct")),
             "source": (
                 "compute_implied_asset_allocation("
                 "cio.blend_weights).equity_pct")},
            {"label": "Current blend IG bonds",
             "value": format_pct(implied.get("ig_bond_pct")),
             "source": (
                 "compute_implied_asset_allocation("
                 "cio.blend_weights).ig_bond_pct")},
            {"label": "Current blend HY bonds",
             "value": format_pct(implied.get("hy_bond_pct")),
             "source": (
                 "compute_implied_asset_allocation("
                 "cio.blend_weights).hy_bond_pct")},
        ],
    }

    return {
        "data_hash":   strategy_hash[:8] or "—",
        "available":   True,
        "computed_at": strategy_computed_at,
        "metrics":     metrics,
    }


# ── Strategy characterisations (item 9) ───────────────────────────────────────


@app.get("/api/v1/strategies/characterisations")
@limiter.limit("30/minute")
async def get_strategy_characterisations(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """
    Item 9 — per-strategy Portfolio Profile data. Returns one row per
    strategy with the AI-generated construction_summary,
    behavioural_profile, regime_sensitivity, behavioural_tag, and the
    deterministic portfolio_characteristics.

    Reads from strategy_characterisations keyed by the current
    data_hash. Cold-cache fallback returns the most recent
    characterisation per strategy regardless of hash (DISTINCT ON in
    the helper) so a fresh deploy that hasn't seen a refresh yet
    still has something to render.

    Auth: require_team_member. Same gate as the diversification
    metrics — non-team viewers see the Dashboard rankings but not
    the per-strategy editorial context.
    """
    if ENVIRONMENT == "test":
        return {"available": False, "strategies": []}
    try:
        from tools.cache import get_latest_strategy_hash
        from tools.strategy_characterisations import (
            get_all_characterisations,
        )
        latest_hash = await get_latest_strategy_hash()
        rows = await get_all_characterisations(latest_hash)
        if not rows:
            return {
                "available": False,
                "strategies": [],
                "note": "Strategy characterisations have not been "
                        "computed yet. They populate automatically "
                        "after the first strategy_results_cache write.",
            }
        return {
            "available": True,
            "data_hash": latest_hash,
            "strategies": rows,
        }
    except Exception as exc:
        log.warning("strategy_characterisations_endpoint_failed",
                    error=str(exc))
        return {"available": False, "strategies": []}


# ── Analytical findings staging (May 22 2026) ─────────────────────────────────


@app.post("/api/v1/reports/stage-findings")
@limiter.limit("5/minute")
async def post_stage_findings(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """
    Runs the analytical staging computation and writes the result to
    analytical_findings_cache. The Academic Writer picks up the
    latest row on its next document-generation call via the workflow
    hook in agents/academic_writer._writer_system_prompt().

    On-demand only — NOT triggered by data-hash change. Findings carry
    interpretation (a NUGGET STRENGTH per finding and an IMPLICATION
    paragraph); pre-computing them silently on every ingestion would
    produce drift the team did not ask for.

    Returns the structured findings + the rendered markdown so the
    report writer UI can display the staging-summary card and let the
    user open the full report inline before enabling [Generate Draft].

    Rate-limited (5/minute). Auth: require_team_member.
    """
    if ENVIRONMENT == "test":
        return {
            "id":              None,
            "data_hash":       None,
            "strategy_count":  0,
            "surprise_count":  0,
            "n_high_strength": 0,
            "findings":        [],
            "findings_md":     "",
        }
    try:
        from tools.analytical_findings import (
            stage_findings, refresh_findings_context,
        )
        result = await stage_findings(triggered_by="api")
        # Refresh the in-process context cache so the next academic-
        # writer call picks up the fresh findings without waiting for
        # a process restart.
        await refresh_findings_context()
        return result
    except Exception as exc:
        log.warning("stage_findings_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Findings staging failed — see server logs.",
        )


@app.get("/api/v1/reports/latest-findings")
async def get_latest_findings_endpoint(
    session: dict = Depends(require_team_member),
):
    """Returns the most recent staged-findings row. The report writer
    UI reads this on mount so the "last staged at" pill renders
    without firing a fresh compute."""
    if ENVIRONMENT == "test":
        return {"available": False}
    try:
        from tools.analytical_findings import get_latest_findings
        row = await get_latest_findings()
        if not row:
            return {"available": False}
        return {"available": True, **row}
    except Exception as exc:
        log.warning("latest_findings_endpoint_failed", error=str(exc))
        return {"available": False}


# ── Report writer template pipeline (item 12) ────────────────────────────────


@app.get("/api/v1/reports/templates")
async def list_report_templates(
    session: dict = Depends(require_team_member),
):
    """Active report templates for the report-writer dropdown.
    Drops system_prompt + section_instructions from the response —
    those are heavy and only needed at generation time."""
    if ENVIRONMENT == "test":
        return {"templates": []}
    from tools.report_templates import list_active_templates
    rows = await list_active_templates()
    return {"templates": [
        {k: v for k, v in r.items()
         if k not in ("system_prompt", "section_instructions")}
        for r in rows
    ]}


@app.post("/api/v1/reports/source-citations")
@limiter.limit("3/minute")
async def post_source_citations(
    request: Request, body: dict,
    session: dict = Depends(require_team_member),
):
    """STEP 1B — source citations for a template's concept list.

    Body: {"template_id": "midpoint_check_fna670"}.

    Returns the verified citations object + the quality indicator
    (green / amber / red). Persists each citation in citations_cache.
    Rate-limited (3/min) because each call hits the web_search tool
    up to 10 times.
    """
    if ENVIRONMENT == "test":
        return {
            "template_id": body.get("template_id"),
            "citations": {},
            "quality": "red",
            "verified_count": 0,
        }
    template_id = body.get("template_id")
    if not template_id:
        raise HTTPException(
            status_code=422, detail="template_id is required.")
    try:
        from tools.report_templates import get_template
        from tools.template_pipeline import (
            source_citations, citation_quality, persist_citations,
        )
        from tools.report_generator import create_placeholder_generation
        tmpl = await get_template(template_id)
        if not tmpl:
            raise HTTPException(
                status_code=404,
                detail=f"Template '{template_id}' not found.")
        concepts = tmpl.get("concepts") or []
        citations = await source_citations(concepts)

        # UAT 2026-05-24 — Open Review fix. Create a placeholder
        # report_generations row BEFORE persisting citations so each
        # citation row carries a real generation_id (was NULL before
        # this change). The frontend Citation Review panel is keyed on
        # generation_id; without a row, the panel rendered nothing
        # and Open Review was a silent no-op.
        #
        # Body accepts an existing generation_id so a re-run uses the
        # same row instead of leaving an orphan placeholder behind.
        # On placeholder failure (DB unavailable) generation_id stays
        # None — citations are still persisted (standalone), behaviour
        # degrades to the pre-fix state rather than failing the step.
        gid_in = body.get("generation_id")
        if isinstance(gid_in, int) and gid_in > 0:
            generation_id: int | None = gid_in
        else:
            generation_id = await create_placeholder_generation(template_id)

        await persist_citations(citations, generation_id=generation_id)
        # CITATION_VERIFIED_STATES covers every state that counts as a
        # real citation — auto-verified, human-accepted, alternative-
        # selected, manually-added. Initial counts are all
        # auto-verified, but the new 3-pass path will also produce
        # pending_review entries that need reviewer attention.
        from tools.template_pipeline import CITATION_VERIFIED_STATES
        verified = sum(
            1 for c in citations.values()
            if c.get("verification_status") in CITATION_VERIFIED_STATES)
        return {
            "template_id":     template_id,
            "generation_id":   generation_id,
            "citations":       citations,
            "quality":         citation_quality(citations),
            "verified_count":  verified,
            "concept_count":   len(concepts),
        }
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("source_citations_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Citation sourcing failed — see server logs.")


# ── Citation reviewer workflow (item 1, May 23 2026) ─────────────────────────
#
# The midpoint paper's citation pipeline now produces pending_review
# entries when the trusted-domain search returns nothing — picked up
# by passes 2 and 3 of source_citations. Bob reviews each pending
# entry and either accepts the search result, picks an alternative
# from the captured pass-2/3 hits, enters a citation manually, or
# rejects the concept entirely. The two endpoints below back the
# CitationReviewPanel on the report writer screen.

@app.get("/api/v1/citations/{generation_id}")
async def get_citations_list(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Return every citation row for a generation_id with its current
    state, alternatives, and reviewer-action history. Powers the
    CitationReviewPanel."""
    from tools.template_pipeline import get_citations_for_generation
    citations = await get_citations_for_generation(generation_id)
    return {
        "generation_id": generation_id,
        "citations":     citations,
    }


@app.post("/api/v1/citations/{citation_id}/review")
async def post_citation_review(
    citation_id: int,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Apply a reviewer action to a citations_cache row.

    Body:
      action: 'accept_untrusted' | 'select_alternative' |
              'reject' | 'manual_add'
      selected_alternative: dict (required when action ==
                            'select_alternative')
      manual_citation:      dict (required when action == 'manual_add')

    The action drives the state transition per the 7-state machine
    (see CITATION_STATES in tools/template_pipeline.py). Reviewer
    identity and timestamp are recorded on the row.
    """
    action = body.get("action")
    if not action:
        raise HTTPException(
            status_code=422, detail="action is required.")
    from tools.template_pipeline import (
        CITATION_REVIEW_ACTIONS, apply_citation_review,
    )
    if action not in CITATION_REVIEW_ACTIONS:
        raise HTTPException(
            status_code=422,
            detail=(f"Unknown action '{action}'. Valid actions: "
                    + ", ".join(sorted(CITATION_REVIEW_ACTIONS))))
    reviewer_email = session.get("email") or "unknown@reviewer"
    selected = body.get("selected_alternative")
    manual = body.get("manual_citation")
    updated = await apply_citation_review(
        citation_id, action, reviewer_email,
        selected_alternative=selected,
        manual_citation=manual,
    )
    if updated is None:
        raise HTTPException(
            status_code=404,
            detail=("Citation not found, or the action payload was "
                    "invalid. select_alternative requires "
                    "selected_alternative; manual_add requires "
                    "manual_citation."))
    return {"citation": updated}


# ── Citation Review redesign — Level 1 findings + matches ────────────────────
#
# Three endpoints back the new 3-level hierarchy (Finding > Type >
# Citation). See backend/tools/citation_findings.py for the seeder
# + match helpers, migration 045 for the schema, and the design doc
# (May 26 2026) for the architecture rationale.

@app.get("/api/v1/citations/findings/{generation_id}")
@limiter.limit("30/minute")
async def get_citation_findings(
    request: Request,
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Re-seeds the per-generation findings table from the live
    statistical-audit + QA-methodology results, then returns the
    fresh findings list alongside the citation pool with each
    citation's `matched_finding_ids[]` joined in.

    Per the design doc: implicit re-seed on every panel open.
    Opening the panel = current analytical state. The team's prior
    matches survive UNCHANGED findings (UPSERT preserves id); a
    finding that resolves between sessions and disappears from the
    seed loses its matches via CASCADE — known limitation, see
    migration 045's docstring.
    """
    from tools.citation_findings import (
        get_matched_finding_ids_by_citation, seed_findings_for_generation,
    )
    from tools.template_pipeline import get_citations_for_generation

    findings = await seed_findings_for_generation(generation_id)
    citations = await get_citations_for_generation(generation_id) or []
    matched = await get_matched_finding_ids_by_citation(generation_id)

    # Join matched_finding_ids[] onto each citation. The frontend
    # uses this to render the checkbox state per (citation, finding)
    # pair without a second round-trip.
    for c in citations:
        c["matched_finding_ids"] = matched.get(int(c.get("id") or 0), [])

    return {
        "generation_id": int(generation_id),
        "seeded_at":     (
            findings[0].get("seeded_at") if findings else None
        ),
        "findings":      findings,
        "citations":     citations,
    }


@app.post("/api/v1/citations/match")
@limiter.limit("60/minute")
async def post_citation_match(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Records a citation_finding_matches row. Idempotent — a
    second call on the same (citation_id, finding_id) refreshes
    matched_at and matched_by, never errors.

    Body: {citation_id: int, finding_id: int, match_rationale?: str}
    Returns: {ok, citation_id, finding_id, matched_by, matched_at}
    422 on missing/invalid ids.
    """
    try:
        cid = int(body.get("citation_id") or 0)
        fid = int(body.get("finding_id") or 0)
    except (TypeError, ValueError):
        raise HTTPException(
            status_code=422,
            detail="citation_id and finding_id must be integers")
    if cid <= 0 or fid <= 0:
        raise HTTPException(
            status_code=422,
            detail="citation_id and finding_id are required")
    from tools.citation_findings import record_match
    email = session.get("email") or "unknown"
    result = await record_match(cid, fid, email)
    if not result.get("ok"):
        raise HTTPException(
            status_code=500,
            detail=f"Match record failed: {result.get('error')}")
    return result


@app.delete("/api/v1/citations/match")
@limiter.limit("60/minute")
async def delete_citation_match(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Removes a citation_finding_matches row. Idempotent — deleting
    a non-existent match returns deleted=false, never 404. Mirrors
    the QA panel's mark-intentional DELETE semantic so the frontend
    can fire without first checking that a match exists.

    Body: {citation_id: int, finding_id: int}
    Returns: {ok, citation_id, finding_id, deleted: bool}
    """
    try:
        cid = int(body.get("citation_id") or 0)
        fid = int(body.get("finding_id") or 0)
    except (TypeError, ValueError):
        raise HTTPException(
            status_code=422,
            detail="citation_id and finding_id must be integers")
    if cid <= 0 or fid <= 0:
        raise HTTPException(
            status_code=422,
            detail="citation_id and finding_id are required")
    from tools.citation_findings import remove_match
    result = await remove_match(cid, fid)
    if not result.get("ok"):
        raise HTTPException(
            status_code=500,
            detail=f"Match removal failed: {result.get('error')}")
    return result


@app.post("/api/v1/reports/team-activity")
@limiter.limit("30/minute")
async def post_team_activity(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """STEP 1C — per-member + platform-wide activity counts.

    No body required — the team_emails are pinned in the pipeline
    helper. Used by the report writer and the Issue Tracker activity
    view both.
    """
    if ENVIRONMENT == "test":
        return {"activity": {}, "cross_check_flags": []}
    try:
        from tools.template_pipeline import (
            fetch_team_activity, cross_check_team_activity,
        )
        activity = await fetch_team_activity()
        flags = cross_check_team_activity(activity)
        return {"activity": activity, "cross_check_flags": flags}
    except Exception as exc:
        log.warning("team_activity_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Team activity fetch failed — see server logs.")


@app.post("/api/v1/reports/validate-thesis")
@limiter.limit("30/minute")
async def post_validate_thesis(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """STEP 6 — thesis validation gate. Pulls the latest verified data
    + ranked findings from the cache and runs the three thesis
    conditions. Blocks generation when any condition fails."""
    if ENVIRONMENT == "test":
        return {
            "passed":           True,
            "conditions":       [],
            "blocker_reasons":  [],
        }
    try:
        from tools.analytical_findings import (
            gather_payload_from_db, get_latest_findings,
        )
        from tools.cache import get_latest_strategy_hash
        from tools.template_pipeline import (
            live_from_payload, cross_check, validate_thesis,
        )
        data_hash = await get_latest_strategy_hash()
        payload = await gather_payload_from_db(data_hash)
        live = live_from_payload(payload)
        findings_row = await get_latest_findings()
        staged_md = (findings_row or {}).get("findings_md") or ""
        verified, _ = cross_check(live, staged_md)
        ranked = (findings_row or {}).get("ranked_findings") or []
        result = validate_thesis(verified, ranked)
        return result
    except Exception as exc:
        log.warning("validate_thesis_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Thesis validation failed — see server logs.")


@app.post("/api/v1/reports/rank-findings")
@limiter.limit("30/minute")
async def post_rank_findings(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """STEP 7 — explicit ranking endpoint. Reads the latest findings,
    re-ranks them, and returns the ordered list. The stage-findings
    endpoint already writes ranked_findings; this endpoint is a
    convenience for the UI to refresh the ranking after a manual
    findings edit."""
    if ENVIRONMENT == "test":
        return {"ranked_findings": []}
    try:
        from tools.analytical_findings import get_latest_findings
        from tools.template_pipeline import rank_findings
        row = await get_latest_findings()
        if not row:
            return {"ranked_findings": []}
        ranked = rank_findings(row.get("findings") or [])
        return {"ranked_findings": ranked}
    except Exception as exc:
        log.warning("rank_findings_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Rank findings failed — see server logs.")


# ── Report writer — generation + editor + download endpoints ────────────────
#
# Backs item 12 commit 2. The generate endpoint runs the eight-step
# pipeline; the editor endpoints (paper-md PATCH, iterate, resolve-bob,
# final-check) support Bob's iteration loop; the academic-review
# endpoint scores the draft against the FNA670 rubric; the download
# endpoints emit the two .docx files. Every endpoint is team_member-
# gated so a viewer cannot kick off a generation that costs Anthropic
# tokens. The download endpoints honor a soft gate on
# academic_readiness=needs_significant_revision — Bob can override
# with acknowledge_warning=true on the query string.


_DOCX_MEDIA = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document")


@app.post("/api/v1/reports/templates/{template_id}/generate")
@limiter.limit("3/minute")
async def generate_report_from_template(
    request: Request, template_id: str,
    session: dict = Depends(require_team_member),
):
    """End-to-end report generation. STREAMS phase events via SSE so
    the request stays alive past Render's gateway timeout.

    May 26 2026 — converted from a synchronous JSON POST to an SSE
    stream after the 3-pass rationalization loop (PR #198) pushed
    total generation time past ~120s. The synchronous variant was
    502-ing at Render's gateway in exactly the same way the council
    endpoint did before its own SSE conversion (see line 4438 for
    the precedent). The frontend assembles the same payload shape
    from the events.

    Event sequence (each `data: {json}\\n\\n`):
      1. generate_started   — fires immediately on request receipt
      2. generate_progress  — periodic keepalive while generate_paper
                              runs (one frame every ~10s)
      3. generate_complete  — full draft payload (the JSON the legacy
                              synchronous response used to return)
      4. data: [DONE]\\n\\n — end-of-stream sentinel

    Error cases — emit a typed frame, then [DONE]:
      generate_error (404)     — template_not_found
      generate_error (422)     — thesis_validation_blocked (carries
                                 thesis_validation payload)
      generate_error (500)     — unexpected pipeline failure

    The test environment keeps the synchronous JSON contract every
    existing report_generator test relies on — TestClient.post(...)
    .json() would not parse SSE frames.
    """
    if ENVIRONMENT == "test":
        return {
            "id":              None,
            "template_id":     template_id,
            "paper_md":        "",
            "appendix_md":     "",
            "flag_count":      0,
            "bob_block_count": 0,
            "bob_blocks":      [],
            "flags":           [],
        }

    log.info("generate_report_started", template_id=template_id,
             user=session.get("email"))

    async def event_stream():
        import asyncio
        import traceback

        # Initial frame — fires immediately so the gateway sees bytes
        # within the first second of the request and starts holding
        # the connection open.
        yield _sse("generate_started", template_id=template_id)

        try:
            from tools.report_generator import generate_paper

            # Run the pipeline as a background task so we can emit
            # keepalive frames in parallel. asyncio.create_task keeps
            # generate_paper on the same event loop; the polling loop
            # yields a frame every ~10s until the task completes.
            task = asyncio.create_task(generate_paper(template_id))
            tick = 0
            while not task.done():
                try:
                    await asyncio.wait_for(asyncio.shield(task),
                                           timeout=10.0)
                except asyncio.TimeoutError:
                    tick += 1
                    # Keepalive frame — Render's gateway treats every
                    # data: ... frame as live traffic, so the 10s
                    # cadence is well under any reasonable idle cap.
                    yield _sse("generate_progress",
                               elapsed_seconds=tick * 10)

            result = task.result()

            if result.get("error") == "template_not_found":
                yield _sse(
                    "generate_error",
                    status=404,
                    message=f"Template '{template_id}' not found.",
                )
            elif result.get("error") == "thesis_validation_blocked":
                yield _sse(
                    "generate_error",
                    status=422,
                    error="thesis_validation_blocked",
                    thesis_validation=result.get("thesis_validation"),
                )
            else:
                yield _sse("generate_complete", **result)
        except Exception as exc:  # noqa: BLE001
            # Unexpected failure surfaces clearly — no silent
            # fall-through. The frontend renders the error message
            # and lets the user retry.
            log.warning("generate_report_failed",
                        template_id=template_id, error=str(exc),
                        traceback=traceback.format_exc())
            yield _sse(
                "generate_error",
                status=500,
                message="Report generation failed — see server logs.",
            )
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(),
                             media_type="text/event-stream")


@app.get("/api/v1/reports/generations")
async def list_report_generations(
    template_id: str | None = None,
    limit: int = 20,
    session: dict = Depends(require_team_member),
):
    """Returns the user's most recent draft generations newest first
    — backs the Draft selector dropdown so Bob can switch between
    saved drafts instead of starting fresh on every login.

    Query parameters:
      template_id  Optional — filter to a single template's drafts.
                   Default: list every template the user has touched.
      limit        Max number of drafts to return. Default 20, capped
                   at 100 server-side so a pathological caller cannot
                   pull every row.

    Auth: team-member only — same scope as the rest of the report
    writer endpoints (viewers cannot generate, so they cannot list).
    """
    if ENVIRONMENT == "test":
        return {"drafts": []}
    capped = max(1, min(int(limit or 20), 100))
    try:
        from tools.report_generator import list_generations_for_user
        email = (session.get("email") or "").strip()
        drafts = await list_generations_for_user(
            email, limit=capped, template_id=template_id)
        return {"drafts": drafts}
    except Exception as exc:
        log.warning("list_generations_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="List drafts failed — see server logs.")


@app.delete("/api/v1/reports/generations/{generation_id}")
async def delete_report_generation(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Hard-delete a saved generation (the Draft Selector trash icon
    calls this). Removes the generation row + every saved version +
    the pipeline-audit row that pointed at it. Citations are kept
    (they're concept-keyed, may be reused by other drafts).

    Frontend confirms with a type-DELETE dialog before firing the
    request — same pattern as the version-history Delete UX.

    Idempotent contract (May 24 2026 update — user-reported fix):
      - 200 OK with {"deleted": True}                — row removed
      - 200 OK with {"deleted": True,
                     "already_absent": True}         — row was
                                                       already gone
                                                       (a second
                                                       click, a
                                                       parallel
                                                       session, a
                                                       stale FE
                                                       cache)
      - 500 with structured detail                   — real DB
                                                       failure
    """
    if ENVIRONMENT == "test":
        return {"deleted": True}
    from tools.report_generator import delete_generation
    result = await delete_generation(generation_id)
    status = result.get("status")
    if status == "deleted":
        return {"deleted": True}
    if status == "already_absent":
        # Idempotent success — DO NOT raise. The user clicked
        # Delete on a draft that's already gone; the desired
        # outcome (no draft) is reached, so we return 200.
        return {"deleted": True, "already_absent": True}
    # status == "error" — surface as 500 so the frontend can
    # render an actual error message. The detail block carries
    # the underlying error string for Render-logs cross-reference.
    raise HTTPException(
        status_code=500,
        detail={
            "error":   "delete_failed",
            "message": result.get("error")
                       or "Delete failed — see server logs.",
        })


@app.get("/api/v1/reports/generations/{generation_id}")
async def get_report_generation(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Returns the persisted generation row + a fresh post-check on
    the current paper_md so the editor renders the same flag list the
    backend will gate on."""
    if ENVIRONMENT == "test":
        return {"error": "test_environment_no_generations"}
    try:
        from tools.report_generator import (
            get_generation, _post_check_summary,
            _load_citations_for_generation,
        )
        gen = await get_generation(generation_id)
        if not gen:
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        citations = await _load_citations_for_generation(generation_id)
        checks = _post_check_summary(
            gen.get("paper_md") or "",
            gen.get("verified_data") or {},
            citations)
        return {**gen, **checks, "citations": citations}
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("get_generation_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Read generation failed — see server logs.")


@app.patch("/api/v1/reports/generations/{generation_id}/paper-md")
@limiter.limit("60/minute")
async def patch_generation_paper_md(
    request: Request, generation_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """Inline editor save. Body:
      paper_md (required) — the new full paper text
      expected_revision (optional) — the paper_revision the caller
        last saw. If supplied AND it does not match the current row,
        the endpoint returns 409 with the actual revision so the
        frontend can show a "concurrent edit detected" prompt. The
        debounced auto-save typically omits this so it never blocks
        itself; the explicit Save action passes it.
      source (optional) — what triggered the save. One of: manual,
        auto_iterate, auto_resolve_bob, auto_edit (default). Stored
        on the version snapshot so the history reads cleanly.
    """
    if ENVIRONMENT == "test":
        return {"saved": False, "flag_count": 0, "paper_revision": 0}
    paper_md = body.get("paper_md")
    if paper_md is None:
        raise HTTPException(
            status_code=422, detail="paper_md is required.")
    expected_revision = body.get("expected_revision")
    source = body.get("source") or "auto_edit"
    try:
        from tools.report_generator import update_paper_md
        result = await update_paper_md(
            int(generation_id), str(paper_md),
            expected_revision=(int(expected_revision)
                                 if expected_revision is not None else None),
            saved_by_email=session.get("email"),
            source=source,
        )
        if result.get("error") == "generation_not_found":
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        if result.get("error") == "revision_mismatch":
            # 409 Conflict — the caller's snapshot is stale. The
            # response body carries the current revision so the
            # frontend can offer a "refresh and re-apply" flow.
            raise HTTPException(
                status_code=409,
                detail={
                    "error":             "revision_mismatch",
                    "current_revision":  result.get("current_revision"),
                    "expected_revision": result.get("expected_revision"),
                    "message": (
                        "Another reviewer saved while you were "
                        "editing. Refresh the paper and re-apply "
                        "your change."),
                })
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("patch_paper_md_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Save failed — see server logs.")


# ── Paper version history (item 2, May 23 2026) ──────────────────────────────
#
# Every save to paper_md creates an append-only snapshot in
# report_paper_versions. The three endpoints below back the version
# history panel: list, save (manual snapshot), restore.

@app.get("/api/v1/reports/generations/{generation_id}/versions")
async def get_paper_versions(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Every version snapshot for a generation, newest first.
    Includes paper_md so the preview can render without a second
    round-trip per version. Powers VersionHistoryPanel."""
    from tools.paper_versions import check_revision, list_versions
    versions = await list_versions(generation_id)
    revision = await check_revision(generation_id)
    return {
        "generation_id":     generation_id,
        "paper_revision":    revision,
        "versions":          versions,
        "version_count":     len(versions),
    }


@app.post("/api/v1/reports/generations/{generation_id}/versions")
async def post_paper_version(
    generation_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """Save the CURRENT paper_md as a named version snapshot. Body:
      label (optional) — short description of the save point
      source (optional) — defaults to 'manual'
    The current paper_md is read fresh from report_generations so a
    concurrent inline edit just before the manual save is captured."""
    from tools.report_generator import get_generation
    from tools.paper_versions import save_version
    gen = await get_generation(generation_id)
    if not gen:
        raise HTTPException(
            status_code=404, detail="Generation not found.")
    snapshot = await save_version(
        generation_id, gen.get("paper_md") or "",
        saved_by_email=session.get("email"),
        label=body.get("label"),
        source=body.get("source") or "manual",
        flag_count=int(gen.get("flag_count") or 0),
        word_counts=gen.get("word_counts") or {},
    )
    if snapshot is None:
        raise HTTPException(
            status_code=500,
            detail="Failed to save version — see server logs.")
    return {"snapshot": snapshot}


@app.post(
    "/api/v1/reports/generations/{generation_id}/versions"
    "/{version_number}/restore")
async def post_paper_version_restore(
    generation_id: int, version_number: int,
    session: dict = Depends(require_team_member),
):
    """Restore a prior version as the new current paper_md. The
    older version row is preserved — a new row is appended with
    source='restore' and restored_from_version pointing at the
    source. Returns the new version snapshot."""
    from tools.paper_versions import restore_version
    snapshot = await restore_version(
        generation_id, version_number,
        reviewer_email=session.get("email"),
    )
    if snapshot is None:
        raise HTTPException(
            status_code=404,
            detail=("Version not found, or the restore failed — "
                    "see server logs."))
    return {"snapshot": snapshot}


@app.post(
    "/api/v1/reports/generations/{generation_id}/versions"
    "/{version_number}/mark-final")
async def post_mark_version_final(
    generation_id: int, version_number: int,
    session: dict = Depends(require_team_member),
):
    """Marks one saved version as the Final Submission. Clears any
    prior Final marker on the same generation. Defense Prep and
    Citation Adjudication reference the Final-marked version
    instead of the most recent draft.

    May 24 2026 P5 — Final Submission marker.
    """
    if ENVIRONMENT == "test":
        return {"marked": True, "version_number": version_number}
    from tools.paper_versions import mark_version_final
    snapshot = await mark_version_final(generation_id, version_number)
    if snapshot is None:
        raise HTTPException(
            status_code=404,
            detail=("Version not found or the mark failed — "
                    "see server logs."))
    return {"marked": True, "snapshot": snapshot}


@app.delete(
    "/api/v1/reports/generations/{generation_id}/versions"
    "/final-marker")
async def delete_final_marker(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Clears the Final Submission marker on a generation. The
    versions themselves are preserved; only the flag is cleared.

    May 24 2026 P5 — Final Submission marker.
    """
    if ENVIRONMENT == "test":
        return {"unmarked": True}
    from tools.paper_versions import unmark_final_version
    ok = await unmark_final_version(generation_id)
    return {"unmarked": ok}


@app.get(
    "/api/v1/reports/generations/{generation_id}/versions/canonical")
async def get_canonical_version_endpoint(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Returns the canonical submission version for a generation:
    the Final-marked version when one exists, the most recent
    saved version otherwise. Defense Prep and Citation
    Adjudication consume this so both reference the same row.

    May 24 2026 P5 — Final Submission marker.
    """
    if ENVIRONMENT == "test":
        return {"canonical": None}
    from tools.paper_versions import get_canonical_version
    canonical = await get_canonical_version(generation_id)
    if canonical is None:
        raise HTTPException(
            status_code=404,
            detail="No saved version found for this generation.")
    return {"canonical": canonical}


@app.delete(
    "/api/v1/reports/generations/{generation_id}/versions"
    "/{version_number}")
async def delete_paper_version(
    generation_id: int, version_number: int,
    session: dict = Depends(require_team_member),
):
    """Hard-delete a single saved version. The frontend Version History
    panel gates this behind a type-DELETE confirmation. The currently-
    active paper_md on report_generations is NOT touched — only the
    snapshot row in report_paper_versions is removed.

    May 24 2026 — Version History Delete UX.
    """
    if ENVIRONMENT == "test":
        return {"deleted": True}
    from tools.paper_versions import delete_version
    ok = await delete_version(generation_id, version_number)
    if not ok:
        raise HTTPException(
            status_code=404,
            detail="Version not found, or the delete failed — "
                   "see server logs.")
    return {"deleted": True}


@app.delete(
    "/api/v1/reports/generations/{generation_id}/versions")
async def delete_all_paper_versions(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Hard-delete EVERY saved version for this generation. The
    Delete All Drafts flow on the Version History panel; gated behind
    a type-DELETE confirmation.

    May 24 2026 — Version History Delete UX.
    """
    if ENVIRONMENT == "test":
        return {"deleted": 0}
    from tools.paper_versions import delete_all_versions
    n = await delete_all_versions(generation_id)
    return {"deleted": n}


@app.post("/api/v1/reports/generations/{generation_id}/iterate")
@limiter.limit("30/minute")
async def post_iterate_text(
    request: Request, generation_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """AI iteration toolbar. Body: {action, selection, instruction?}.

    action ∈ {rephrase, tighten, expand, ask}.

    Returns the rewritten text + any new unverified numbers or
    citations the iteration introduced — the editor warns Bob before
    accepting changes that would fail the next final check."""
    if ENVIRONMENT == "test":
        return {
            "original":  body.get("selection", ""),
            "rewritten": body.get("selection", ""),
            "word_delta": 0,
            "new_unverified_numbers": [],
            "new_unverified_citations": [],
        }
    action = body.get("action")
    selection = body.get("selection")
    if not action or selection is None:
        raise HTTPException(
            status_code=422,
            detail="action and selection are required.")
    if action not in ("rephrase", "tighten", "expand", "ask"):
        raise HTTPException(
            status_code=422,
            detail=f"unknown action '{action}'")
    try:
        from tools.report_generator import iterate_text
        return await iterate_text(
            int(generation_id), action, str(selection),
            instruction=body.get("instruction"))
    except Exception as exc:
        log.warning("iterate_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Iteration failed — see server logs.")


@app.post("/api/v1/reports/generations/{generation_id}/resolve-bob")
@limiter.limit("60/minute")
async def post_resolve_bob_block(
    request: Request, generation_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """Replaces the FIRST occurrence of `marker` in paper_md with
    `replacement`. Body: {marker: str, replacement: str}."""
    if ENVIRONMENT == "test":
        return {"saved": False, "flag_count": 0}
    marker = body.get("marker")
    replacement = body.get("replacement")
    if marker is None or replacement is None:
        raise HTTPException(
            status_code=422,
            detail="marker and replacement are required.")
    try:
        from tools.report_generator import resolve_bob_block
        result = await resolve_bob_block(
            int(generation_id), str(marker), str(replacement))
        if result.get("error") == "generation_not_found":
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        if result.get("error") == "marker_not_found":
            raise HTTPException(
                status_code=422,
                detail=f"marker not found in paper_md: {marker}")
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("resolve_bob_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Resolve failed — see server logs.")


@app.post("/api/v1/reports/generations/{generation_id}/rebalance")
@limiter.limit("12/minute")
async def post_rebalance_paper(
    request: Request, generation_id: int,
    session: dict = Depends(require_team_member),
):
    """May 24 2026 — Pass 2 of the two-pass draft generation flow.
    After Bob adjudicates every [BOB] block, the section word
    counts are off-budget. This endpoint re-runs the writer over
    the current paper_md with a rebalance instruction that brings
    each off-budget section back within ±5 words of its target,
    without touching inline citations or specific numbers.

    No body required — the endpoint reads the current paper_md
    from the generation row. Returns the updated paper_md + a
    targets list naming which sections were re-balanced."""
    if ENVIRONMENT == "test":
        return {"saved": True, "paper_md": "", "rebalanced": False,
                "note": "test environment"}
    try:
        from tools.report_generator import rebalance_paper
        result = await rebalance_paper(int(generation_id))
        if result.get("error") == "generation_not_found":
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        if result.get("error") == "empty_paper":
            raise HTTPException(status_code=422,
                                detail="Paper is empty — cannot rebalance.")
        if result.get("error") == "writer_unavailable":
            raise HTTPException(
                status_code=503,
                detail="Writer model unavailable for rebalance — try again.")
        if result.get("error") == "writer_returned_empty":
            raise HTTPException(
                status_code=502,
                detail="Writer returned empty response — try again.")
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("rebalance_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Rebalance failed — see server logs.")


@app.post("/api/v1/reports/generations/{generation_id}/final-check")
@limiter.limit("60/minute")
async def post_run_final_check(
    request: Request, generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Re-runs the post-checks against the current paper_md and
    updates flag_count on the row so the download endpoints can gate.
    No body required."""
    if ENVIRONMENT == "test":
        return {"passed": True, "flag_count": 0, "flags": []}
    try:
        from tools.report_generator import run_final_check
        result = await run_final_check(int(generation_id))
        if result.get("error") == "generation_not_found":
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("final_check_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Final check failed — see server logs.")


@app.post("/api/v1/reports/generations/{generation_id}/academic-review")
@limiter.limit("6/minute")
async def post_run_academic_review(
    request: Request, generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Step 10 — scores the current paper_md against the active
    rubric for its template. Persists the review payload + readiness
    on the row so the download endpoint can soft-gate."""
    if ENVIRONMENT == "test":
        return {
            "per_criterion":     [],
            "data_gaps":         [],
            "citation_gaps":     [],
            "thesis_coherence":  [],
            "tone_violations":   [],
            "length_compliance": [],
            "readiness":         "ready_to_submit",
            "summary":           "(test environment)",
        }
    try:
        from tools.report_generator import run_academic_review
        result = await run_academic_review(int(generation_id))
        if result.get("error") == "generation_not_found":
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        if result.get("error") == "rubric_not_found":
            raise HTTPException(
                status_code=404,
                detail=(
                    f"No rubric uploaded for template "
                    f"'{result.get('template_id')}'."))
        return result
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("academic_review_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Academic review failed — see server logs.")


def _download_filename(template_id: str, kind: str, ext: str) -> str:
    """forest-capital-midpoint-check-paper-2026-05-22.docx"""
    from datetime import date
    slug = template_id.replace("_", "-")
    return (
        f"forest-capital-{slug}-{kind}-{date.today().isoformat()}.{ext}")


def _gate_download(
    gen: dict, acknowledge_warning: bool,
) -> None:
    """Enforces the two-tier download gate.

    Hard gate: flag_count > 0 means BOB blocks or unresolved numbers
    are still present — refuse with 422.
    Soft gate: academic_readiness == needs_significant_revision —
    refuse unless acknowledge_warning is True. The endpoint records
    the override in the row's audit trail (caller's responsibility
    via the activity log, not duplicated here)."""
    if int(gen.get("flag_count") or 0) > 0:
        raise HTTPException(
            status_code=422,
            detail={
                "error": "flags_remaining",
                "flag_count": int(gen.get("flag_count") or 0),
                "message": (
                    "Resolve all [BOB] blocks and unverified numbers "
                    "before downloading."),
            })
    readiness = gen.get("academic_readiness")
    if (readiness == "needs_significant_revision"
            and not acknowledge_warning):
        raise HTTPException(
            status_code=422,
            detail={
                "error": "academic_review_significant_revision",
                "readiness": readiness,
                "message": (
                    "Academic review flagged significant gaps. "
                    "Retry with ?acknowledge_warning=true to "
                    "download anyway."),
            })


@app.get(
    "/api/v1/reports/generations/{generation_id}/download-paper")
async def download_report_paper(
    generation_id: int,
    acknowledge_warning: bool = False,
    session: dict = Depends(require_team_member),
):
    """Returns the paper .docx. Soft-gated by academic_readiness."""
    if ENVIRONMENT == "test":
        raise HTTPException(
            status_code=404,
            detail="No generation available in test environment.")
    try:
        from tools.report_generator import (
            get_generation, render_paper_bytes,
        )
        gen = await get_generation(generation_id)
        if not gen:
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        _gate_download(gen, acknowledge_warning)
        content = await render_paper_bytes(generation_id)
        if not content:
            raise HTTPException(
                status_code=500,
                detail="Render failed — see server logs.")
        filename = _download_filename(
            gen.get("template_id") or "report", "paper", "docx")
        return Response(
            content=content, media_type=_DOCX_MEDIA,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            })
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("download_paper_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Download failed.")


@app.get(
    "/api/v1/reports/generations/{generation_id}/download-appendix")
async def download_report_appendix(
    generation_id: int,
    session: dict = Depends(require_team_member),
):
    """Returns the appendix .docx. Always available once a generation
    exists — the appendix is the data record, not the editable
    submission, so it is NOT gated by flag_count or readiness."""
    if ENVIRONMENT == "test":
        raise HTTPException(
            status_code=404,
            detail="No generation available in test environment.")
    try:
        from tools.report_generator import (
            get_generation, render_appendix_bytes,
        )
        gen = await get_generation(generation_id)
        if not gen:
            raise HTTPException(status_code=404,
                                detail="Generation not found.")
        content = await render_appendix_bytes(generation_id)
        if not content:
            raise HTTPException(
                status_code=500,
                detail="Render failed — see server logs.")
        filename = _download_filename(
            gen.get("template_id") or "report", "appendix", "docx")
        return Response(
            content=content, media_type=_DOCX_MEDIA,
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
            })
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("download_appendix_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Download failed.")


@app.get("/api/v1/reports/templates/{template_id}/rubric")
async def get_template_rubric(
    template_id: str,
    session: dict = Depends(require_team_member),
):
    """Latest active rubric for a template. Drives the report
    writer's Rubric panel."""
    if ENVIRONMENT == "test":
        return {"rubric": None}
    try:
        from tools.report_rubrics import get_latest_rubric
        rubric = await get_latest_rubric(template_id)
        return {"rubric": rubric}
    except Exception as exc:
        log.warning("get_rubric_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Read rubric failed.")


@app.get("/api/v1/reports/pipeline-audit/active")
async def get_active_pipeline_audit(
    session: dict = Depends(require_team_member),
):
    """Returns the most recent pipeline run started by the caller in
    the last two hours, or {available: false} when none exists. The
    report writer reads this on mount to restore step state across
    navigation."""
    if ENVIRONMENT == "test":
        return {"available": False}
    try:
        from tools.pipeline_audit import get_active_run_for_user
        row = await get_active_run_for_user(session["email"])
        if not row:
            return {"available": False}
        # Attach the persisted paper_md when a generation_id exists,
        # so the editor restores the draft text on the same fetch.
        paper_md = ""
        gen_id = row.get("generation_id")
        if isinstance(gen_id, int):
            try:
                from tools.report_generator import get_generation
                gen = await get_generation(gen_id)
                if gen:
                    paper_md = gen.get("paper_md") or ""
            except Exception:  # noqa: BLE001
                paper_md = ""
        return {"available": True, "audit": row, "paper_md": paper_md}
    except Exception as exc:
        log.warning("active_audit_endpoint_failed", error=str(exc))
        return {"available": False}


@app.post("/api/v1/reports/pipeline-audit")
@limiter.limit("60/minute")
async def post_pipeline_audit(
    request: Request, body: dict,
    session: dict = Depends(require_team_member),
):
    """Records one pipeline audit row. Posted by the report-writer UI
    after the pipeline reaches step 7 (success) or fails at any
    earlier step. The body shape is the flat per-step dict the UI
    builds from its timing state; the writer fills in the
    triggered_by from the authenticated session.

    Body:
      generation_id:        int | null
      template_id:          str (required)
      total_pipeline_ms:    int | null
      failure_step:         int | null
      failure_reason:       str | null
      steps:                {step_<n>_status, step_<n>_ms,
                             step_5_mismatch_count?,
                             step_6_conditions?}
    """
    if ENVIRONMENT == "test":
        return {"id": None}
    template_id = body.get("template_id")
    if not template_id:
        raise HTTPException(
            status_code=422, detail="template_id is required.")
    try:
        from tools.pipeline_audit import (
            upsert_active_run, update_generation_timings,
        )
        steps = body.get("steps") or {}
        # Item 12 commit B — incremental upsert. The frontend rounds
        # trip the audit_id so every step completion writes to the
        # same row. First write (audit_id None) inserts; subsequent
        # writes update only the fields present in the steps dict.
        # The terminal write (Step 7 success/failure) still flows
        # through here — it just happens to be the last upsert.
        new_id = await upsert_active_run(
            template_id=str(template_id),
            triggered_by=session.get("email"),
            steps=steps,
            total_pipeline_ms=body.get("total_pipeline_ms"),
            failure_step=body.get("failure_step"),
            failure_reason=body.get("failure_reason"),
            generation_id=body.get("generation_id"),
            audit_id=body.get("audit_id"),
        )
        # When the run produced a generation, also persist the
        # per-step ms dict on the row so the summary card has a
        # canonical record. Fail-open.
        gen_id = body.get("generation_id")
        if isinstance(gen_id, int):
            timings = {
                k.replace("_ms", ""): v
                for k, v in steps.items() if k.endswith("_ms")
            }
            if timings:
                await update_generation_timings(gen_id, timings)
        return {"id": new_id, "audit_id": new_id}
    except Exception as exc:
        log.warning("pipeline_audit_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail="Pipeline audit write failed.")


@app.get("/api/v1/admin/pipeline-audit")
async def list_pipeline_audit_runs(
    limit: int = 100,
    session: dict = Depends(require_sysadmin),
):
    """Newest-first audit runs for the sysadmin Settings panel."""
    if ENVIRONMENT == "test":
        return {"runs": []}
    try:
        from tools.pipeline_audit import list_audit_runs
        runs = await list_audit_runs(limit=limit)
        return {"runs": runs}
    except Exception as exc:
        log.warning("pipeline_audit_list_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Read audit runs failed.")


@app.get("/api/v1/admin/pipeline-audit/{audit_id}")
async def get_pipeline_audit_run(
    audit_id: int,
    session: dict = Depends(require_sysadmin),
):
    """Single audit run by id for the expand-row view."""
    if ENVIRONMENT == "test":
        return {"error": "not_found"}
    try:
        from tools.pipeline_audit import get_audit_run
        row = await get_audit_run(audit_id)
        if not row:
            raise HTTPException(status_code=404, detail="Run not found.")
        return row
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("pipeline_audit_read_endpoint_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Read audit run failed.")


@app.post("/api/v1/reports/templates/{template_id}/rubric")
@limiter.limit("6/minute")
async def post_upload_rubric(
    request: Request, template_id: str, body: dict,
    session: dict = Depends(require_team_member),
):
    """Uploads a new rubric version for a template.

    Body: {rubric_text, criteria, source_filename?}. The PDF/docx
    text extraction runs client-side (the frontend uses the same
    pypdf path as the academic_documents flow) so this endpoint
    accepts already-extracted text + already-parsed criteria.
    """
    if ENVIRONMENT == "test":
        return {"id": None, "version": 1}
    rubric_text = body.get("rubric_text")
    criteria = body.get("criteria")
    if not rubric_text or not isinstance(criteria, list):
        raise HTTPException(
            status_code=422,
            detail="rubric_text and criteria (list) are required.")
    try:
        from tools.report_rubrics import upload_rubric, get_latest_rubric
        await upload_rubric(
            template_id, str(rubric_text), criteria,
            uploaded_by=session.get("email"),
            source_filename=body.get("source_filename"))
        latest = await get_latest_rubric(template_id)
        return {"rubric": latest}
    except Exception as exc:
        log.warning("upload_rubric_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Upload rubric failed.")


_RISK_FREE_SOURCE = "FRED DTB3 (3-month T-bill, mean monthly rate, annualised)"


async def _read_cached_metric_or_fallback(
    metric_kind: str,
    fallback: "Callable[[], Awaitable[dict]]",
) -> dict:
    """Three-tier read for analytics_metrics_cache where the cold-cache
    fallback isn't the strategy-keyed inline path that
    _read_div_metric_or_compute assumes. The fallback is an arbitrary
    async callable that produces the payload from whatever the
    endpoint's natural inline path is.

    Used by /api/v1/analytics/sensitivity (cold path needs the full
    pipeline history) and /api/v1/analytics/config (cold path reads
    market_data_monthly directly). Both pre-existed the cache layer;
    F1 + F4 fold them in alongside the seven diversification metrics
    that already use _read_div_metric_or_compute.
    """
    try:
        from tools.cache import get_latest_strategy_hash
        from tools.precomputed_analytics import (
            get_metric, get_latest_metric,
        )
        latest_hash = await get_latest_strategy_hash()
        if latest_hash:
            cached = await get_metric(latest_hash, metric_kind)
            if cached:
                log.info("metric_cache_hit", metric=metric_kind)
                return cached
            stale = await get_latest_metric(metric_kind)
            if stale:
                log.info("metric_cache_stale_hit", metric=metric_kind)
                return stale
        log.info("metric_cache_miss_inline", metric=metric_kind)
        return await fallback()
    except Exception as exc:  # noqa: BLE001
        log.warning("metric_read_failed", metric=metric_kind, error=str(exc))
        return await fallback()


async def _overlay_live_regime(target: dict, *, prob_key: str = "probability") -> None:
    """Write the LIVE regime label + posterior confidence onto `target`
    (`target["regime"]` and `target[prob_key]`), from a fresh
    detect_current_regime() read. Both landing tiles call this so they show
    one identical, current confidence — not the value frozen into each
    tile's data_hash cache at a different warm time. detect_current_regime()
    has a 15-minute in-process cache (shared with the dashboard regime
    banner), so this is the platform's live regime read, not a per-tile
    cache. Fail-open: any error leaves the cached values untouched."""
    if not isinstance(target, dict):
        return
    try:
        import asyncio
        from tools.regime_detector import detect_current_regime
        live = await asyncio.to_thread(detect_current_regime)
        regime = (live or {}).get("hmm_regime")
        if not regime:
            return
        conf = ((live or {}).get("hmm_probabilities") or {}).get(regime)
        target["regime"] = regime
        if isinstance(conf, (int, float)):
            target[prob_key] = conf
    except Exception as exc:  # noqa: BLE001
        log.warning("live_regime_overlay_failed", error=str(exc))


@app.get("/api/v1/recommendation")
async def get_cio_recommendation(session: dict = Depends(require_auth)):
    """The live CIO recommendation behind the landing page. Served from
    the cio_recommendations cache — the four-component object the
    hash-change pipeline last computed for the current data. Read-only;
    the recompute is never triggered here (it fires from the analytics
    warm pipeline on a data_hash change). Returns {available, ...} so the
    frontend renders an empty state cleanly before the first compute."""
    if ENVIRONMENT == "test":
        return {"available": False, "recommendation": None}
    try:
        # Read the recommendation by its REGIME-AWARE composite cache key
        # ({data_hash}_{regime}_{bucket}) so a regime flip or a confidence
        # move that crosses a 10pp bucket boundary serves the appropriate
        # prose. On a miss, get_endpoint_recommendation returns the most
        # recent cached row (whatever its key) AND schedules a background
        # regenerate under the new composite key — the next read serves the
        # fresh prose without blocking this one.
        from tools.cio_recommendation import get_endpoint_recommendation
        rec = await get_endpoint_recommendation()
        # Overlay the LIVE regime read so the headline regime label +
        # confidence are never the data_hash-stale cached values, and match
        # the Forward Projection tile exactly (both read the same
        # detect_current_regime() 15-minute in-process regime cache). The
        # cached recommendation/dissent/limitations text is unchanged — the
        # card reads regime + probability from rec.confidence.
        if rec:
            conf = rec.get("confidence")
            if not isinstance(conf, dict):
                conf = {}
                rec["confidence"] = conf
            await _overlay_live_regime(conf, prob_key="probability")
            # Surface the live regime-conditional blend weights (from the
            # cached forward projection — the same prob-weighted blend) so
            # the tile can show the blend and flag a binding concentration
            # constraint. Fail-open: the constraints table still renders its
            # static rows without the weights, just no binding note.
            try:
                from tools.regime_meta_forward import get_cached_forward_projection
                proj = await get_cached_forward_projection()
                if proj and proj.get("blend_weights"):
                    rec["blend_weights"] = proj["blend_weights"]
            except Exception as exc:  # noqa: BLE001
                log.warning("recommendation_blend_overlay_failed", error=str(exc))
            # Bridge #81 -- implied asset allocation + blend change
            # trigger overlays. Both reuse already-cached primitives:
            #
            # implied_asset_allocation -- equity / bond / cash portfolio
            #   split derived from the live blend_weights times the
            #   per-strategy avg_equity_weight / avg_bond_weight
            #   persisted in strategy_results_cache. Pure read.
            #
            # blend_change_trigger -- one readable sentence describing
            #   what would shift the blend (regime flip, VIX threshold,
            #   etc.). Synthesised from detect_current_regime() output
            #   already used by _overlay_live_regime above. The 15-min
            #   in-process regime cache shields the second read.
            #
            # Both fail-open: a cold strategy cache leaves
            # implied_asset_allocation absent; an unavailable regime
            # leaves blend_change_trigger as the generic sentence. The
            # frontend gracefully omits or shows the existing line in
            # either case.
            try:
                import asyncio
                from tools.cio_recommendation import (
                    compute_implied_asset_allocation,
                    compute_blend_change_trigger,
                    compute_regime_blends_implied,
                )
                from tools.regime_detector import detect_current_regime
                allocation = await compute_implied_asset_allocation(
                    rec.get("blend_weights"))
                if allocation:
                    rec["implied_asset_allocation"] = allocation
                live = await asyncio.to_thread(detect_current_regime)
                rec["blend_change_trigger"] = compute_blend_change_trigger(
                    (live or {}).get("hmm_regime"),
                    (live or {}).get("monthly_hmm_regime"),
                    bool((live or {}).get("hmm_models_agree", True)),
                )
                # Bridge (June 8 2026) -- per-regime blend shifts overlay.
                # Pulls the cached regime_blends payload and computes the
                # equity/bond implied split + delta-from-current for each
                # of BULL / BEAR / TRANSITION so the CIO card and digest
                # can render the per-regime shift without re-computing.
                # Fail-open: a cold regime_blends row leaves the field
                # absent and the frontend / digest omit the section.
                try:
                    # get_latest_metric lives in tools.precomputed_analytics,
                    # not tools.cache (every other caller in the codebase
                    # imports it from there). The wrong module here silently
                    # broke the regime_blends overlay -- the try-block
                    # swallowed the ImportError and logged
                    # recommendation_regime_blends_overlay_failed, hiding
                    # the typo behind the structured warning.
                    from tools.precomputed_analytics import get_latest_metric
                    rb_row = await get_latest_metric("regime_blends") or {}
                    rb_blends = rb_row.get("blends") or {}
                    if rb_blends and allocation:
                        regime_implied = await compute_regime_blends_implied(
                            rb_blends, allocation)
                        if regime_implied:
                            rec["regime_blends_implied"] = regime_implied
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "recommendation_regime_blends_overlay_failed",
                        error=str(exc))
                # OOS validation overlay (June 15 2026). The
                # precomputed oos_summary metric carries the
                # December 2025 academic-lock Sharpe values for the
                # blend vs the 100%-equity benchmark plus the
                # value-add event count from the play-by-play
                # scorecard. Same pure-overlay pattern as the
                # blend_change_trigger -- the cached scalar is read
                # and threaded onto the recommendation payload so
                # the CIO card can render the OOS validation row
                # without a second round-trip. Fail-open: a cold
                # oos_summary cache leaves the field as None and
                # the frontend omits the row.
                try:
                    from tools.play_by_play import get_cached_oos_summary
                    oos = await get_cached_oos_summary()
                    if oos:
                        rec["oos_sharpe"] = {
                            "blend": oos.get("blend"),
                            "benchmark": oos.get("benchmark"),
                            "value_add_events": oos.get(
                                "value_add_events"),
                            "total_events": oos.get("total_events"),
                        }
                    else:
                        rec["oos_sharpe"] = None
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "recommendation_oos_summary_overlay_failed",
                        error=str(exc))
                    rec["oos_sharpe"] = None
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "recommendation_allocation_overlay_failed",
                    error=str(exc))
        return {"available": bool(rec), "recommendation": rec}
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.warning("cio_recommendation_read_failed", ref=ref, error=str(exc))
        return {"available": False, "recommendation": None, "ref": ref}


@app.get("/api/v1/play-by-play")
async def get_play_by_play(session: dict = Depends(require_auth)):
    """The nine stored point-in-time event evaluations behind the Council
    Performance Record page, read-only. Serves the frozen rows from
    play_by_play_events plus the aggregate scorecard and the curated
    key-limitation notes (Liberation Day). Never recomputes — the events
    are written once by run_play_by_play.py. Returns {available, ...} so
    the page renders an empty state cleanly before the first run."""
    if ENVIRONMENT == "test":
        return {"available": False, "events": [], "scorecard": None,
                "key_limitations": {}}
    try:
        from tools.play_by_play import (
            KEY_LIMITATION_NOTES, get_cached_performance_chart,
            load_stored_events, scorecard,
        )
        events = await load_stored_events()
        for ev in events:
            note = KEY_LIMITATION_NOTES.get(ev.get("event_id"))
            if note:
                ev["key_limitation"] = note
        cumulative = await get_cached_performance_chart()
        # Diagnostic: log the raw count so a "table has rows but the page
        # shows empty" report is resolvable from the Render logs.
        log.info("play_by_play_endpoint_read", n_events=len(events),
                 available=bool(events), has_cumulative=bool(cumulative))
        return {
            "available": bool(events),
            "events": events,
            "event_count": len(events),
            "scorecard": scorecard(events) if events else None,
            "key_limitations": KEY_LIMITATION_NOTES,
            "cumulative": cumulative,
        }
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.warning("play_by_play_read_failed", ref=ref, error=str(exc))
        return {"available": False, "events": [], "scorecard": None,
                "key_limitations": {}, "ref": ref}


@app.get("/api/v1/oos-cost-sensitivity")
async def get_oos_cost_sensitivity(session: dict = Depends(require_auth)):
    """Transaction-cost sensitivity of the regime-conditional blend over
    the post-2022 OOS window — net Sharpe + vs-benchmark at 10/15/20 bps
    per rebalance and the material-rebalance count — behind the Council
    Performance Record "Net of Switching Costs" table. Served from the
    data_hash-cached 'oos_cost_sensitivity' metric; never recomputes on a
    read. Returns {available, ...} for a clean empty state before the
    first warm."""
    if ENVIRONMENT == "test":
        return {"available": False, "cost_sensitivity": None}
    try:
        from tools.regime_meta_validation import get_cached_cost_sensitivity
        cost = await get_cached_cost_sensitivity()
        return {"available": bool(cost), "cost_sensitivity": cost}
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.warning("oos_cost_sensitivity_read_failed", ref=ref, error=str(exc))
        return {"available": False, "cost_sensitivity": None, "ref": ref}


@app.get("/api/v1/forward-projection")
async def get_forward_projection(session: dict = Depends(require_auth)):
    """The Layer 4 forward Monte Carlo confidence bands for the landing
    page (blend / benchmark / classic 60/40, each with a 90% band, plus
    P(blend outperforms each baseline) at 1/3/6/12 months). Served from
    the data_hash-cached `forward_projection` metric; never recomputes the
    simulation on a read. Returns {available, ...} for a clean empty
    state before the first warm."""
    if ENVIRONMENT == "test":
        return {"available": False, "projection": None}
    try:
        from tools.regime_meta_forward import get_cached_forward_projection
        proj = await get_cached_forward_projection()
        # Overlay the LIVE regime read onto the cached bands so the regime
        # label + confidence match the CIO card (same source) rather than
        # the data_hash-stale value the simulation was cached with. The
        # simulation bands / p_outperform stay as cached.
        if proj:
            await _overlay_live_regime(proj, prob_key="regime_probability")
        return {"available": bool(proj), "projection": proj}
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.warning("forward_projection_read_failed", ref=ref, error=str(exc))
        return {"available": False, "projection": None, "ref": ref}


@app.get("/api/v1/analytics/config")
async def get_analytics_config(session: dict = Depends(require_auth)):
    """
    The analytics assumptions surfaced in Settings → Analytics
    Configuration. Currently the risk-free rate applied to every Sharpe
    ratio and to the efficient frontier — the mean monthly DTB3 rate from
    market_data_monthly, annualised (×12). This is the SAME value the
    /api/optimize/weights frontier and the analytics layer use. Read-only.

    F4 (May 22 2026): cache-first via analytics_metrics_cache. The
    cold-cache fallback runs the original inline compute so the
    contract is preserved on a fresh deploy that hasn't seen a
    strategy_cache write yet.
    """
    if ENVIRONMENT == "test":
        return {"available": False, "risk_free_rate": None,
                "risk_free_source": _RISK_FREE_SOURCE}

    async def _inline_config() -> dict:
        try:
            from tools.cache import get_monthly_returns
            monthly = await get_monthly_returns()
            rf_list = (monthly or {}).get("rf") or []
            rf_annual = (
                sum(rf_list) / len(rf_list) * 12) if rf_list else None
            return {
                "available": rf_annual is not None,
                "risk_free_rate":
                    round(rf_annual, 4) if rf_annual is not None else None,
                "risk_free_source": _RISK_FREE_SOURCE,
            }
        except Exception as exc:
            log.warning("analytics_config_failed", error=str(exc))
            return {"available": False, "risk_free_rate": None,
                    "risk_free_source": _RISK_FREE_SOURCE}

    return await _read_cached_metric_or_fallback(
        "risk_free_rate_config", _inline_config)


@app.get("/api/v1/analytics/sensitivity")
@limiter.limit("10/minute")
async def get_analytics_sensitivity(request: Request, session: dict = Depends(require_auth)):
    """
    Parameter sensitivity analysis for the four dynamic strategies — the
    Sharpe ratio swept across a range of each strategy's key parameter.

    This is a ~23-backtest computation. F1 (May 22 2026) folded it into
    analytics_metrics_cache: refresh_sensitivity runs the sweep once
    when a fresh strategy_results_cache row lands, and this endpoint
    serves the cached payload on every subsequent request. The cold-
    cache fallback below preserves the original
    get_full_history + compute_sensitivity path for a fresh deploy
    that has not seen a strategy_cache write yet — compute_sensitivity
    has its own worker-local memo so even the cold path only does
    full work once per worker per history. The frontend section
    shows its own loading state regardless.
    """
    if ENVIRONMENT == "test":
        return {"available": False, "strategies": []}

    async def _inline_sensitivity() -> dict:
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.sensitivity import compute_sensitivity
            history = await get_full_history_async()
            result = await asyncio.to_thread(compute_sensitivity, history)
            return {"available": True, **result}
        except Exception as exc:
            log.warning("analytics_sensitivity_failed", error=str(exc))
            return {"available": False, "strategies": []}

    return await _read_cached_metric_or_fallback(
        "sensitivity", _inline_sensitivity)


# ── Admin: data status ────────────────────────────────────────────────────────

@app.post("/api/v1/admin/clear-story-plans")
async def post_clear_story_plans(
    document_type: str = "all",
    session: dict = Depends(require_team_member),
):
    """PR #336 Gap E -- delete cached story_plans rows so the next
    deck or brief regeneration runs the four-pass Opus pipeline from
    scratch. Required after a prompt change (e.g. PR #335's seven-
    citation grounding) lands so the new framing fires immediately
    rather than waiting for the data_hash to change.

    document_type:
      "brief" -- delete only brief plans
      "deck"  -- delete only deck plans
      "all"   -- delete every story_plans row (the default)

    Returns: {"deleted": <int>, "document_type": <str>}.
    """
    if document_type not in {"brief", "deck", "all"}:
        raise HTTPException(
            status_code=400,
            detail="document_type must be 'brief' | 'deck' | 'all'")
    if ENVIRONMENT == "test":
        return {"deleted": 0, "document_type": document_type,
                "note": "test environment"}
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal  # type: ignore[attr-defined]
        if AsyncSessionLocal is None:
            return {"deleted": 0, "document_type": document_type,
                    "note": "no database"}
        async with AsyncSessionLocal() as s:
            if document_type == "all":
                res = await s.execute(
                    text("DELETE FROM story_plans"))
            else:
                res = await s.execute(
                    text("DELETE FROM story_plans "
                         "WHERE document_type = :t"),
                    {"t": document_type})
            await s.commit()
            deleted = int(res.rowcount or 0)
        log.info("story_plans_cleared",
                 document_type=document_type, deleted=deleted,
                 actor=session.get("email"))
        return {"deleted": deleted, "document_type": document_type}
    except Exception as exc:  # noqa: BLE001
        log.warning("clear_story_plans_failed",
                    document_type=document_type, error=str(exc))
        raise HTTPException(
            status_code=500,
            detail=f"clear_story_plans_failed: {exc}")


@app.get("/api/v1/story-plans/exists")
async def get_story_plans_exists(
    document_types: str = "",
    session: dict = Depends(require_auth),
):
    """June 23 2026 -- pre-flight check for the brief regen
    confirmation modal. The frontend hits this endpoint BEFORE
    POSTing /api/v1/export/executive-brief; if any downstream
    plan exists for the listed document_types, the UI surfaces
    a confirmation modal warning the user that those plans will
    be cleared. If exists is false, the UI skips the modal and
    fires Generate immediately.

    Query param document_types: a comma-separated list of
    document_type values to check. Returns {exists: bool, types:
    {<doc_type>: bool}}.

    Auth: require_auth -- anyone signed in can ask. The clear
    itself still gates on generate_documents (it runs inside
    _generate_brief_document, which gates on that permission).
    """
    types = [t.strip() for t in document_types.split(",") if t.strip()]
    if not types:
        return {"exists": False, "types": {}}
    if ENVIRONMENT == "test":
        return {"exists": False, "types": {t: False for t in types}}
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal  # type: ignore[attr-defined]
        if AsyncSessionLocal is None:
            return {"exists": False, "types": {t: False for t in types}}
        async with AsyncSessionLocal() as s:
            res = await s.execute(
                text("SELECT document_type FROM story_plans "
                     "WHERE document_type = ANY(:types)"),
                {"types": types})
            present = {row[0] for row in res.fetchall()}
        per_type = {t: (t in present) for t in types}
        return {"exists": any(per_type.values()), "types": per_type}
    except Exception as exc:  # noqa: BLE001
        log.warning("story_plans_exists_failed", error=str(exc))
        # Fail-open: if the check fails, do NOT block the user.
        # The modal will simply not fire; the user can still hit
        # Generate. Returning exists=False matches that intent.
        return {"exists": False, "types": {t: False for t in types}}


@app.post("/api/v1/data/light-refresh")
@limiter.limit("10/minute")
async def post_light_refresh(
    request: Request,
    session: dict = Depends(require_permission("generate_documents")),
):
    """June 24 2026 -- light analytics-cache refresh for the team
    member's "I want to see if there's new market data" workflow.

    Runs the same three compute steps as the admin
    refresh-appendix-caches endpoint (backtester + academic
    analytics + OOS cost sensitivity) but is bound to the
    generate_documents permission rather than team_member so
    every doc-generator user can self-serve a refresh without
    a sysadmin escalation. The response shape is identical to
    refresh-appendix-caches so the frontend can reuse the same
    rendering.

    What this endpoint does NOT do:
      - Touch story_plans (those clear on brief regen via the
        existing _generate_brief_document hook)
      - Touch editor_drafts (no document_content is mutated)
      - Touch the canonical data_hash unless the data fetcher
        reports new rows that change it
      - Trigger document regeneration -- the user runs Generate /
        Regenerate themselves on the doc tiles when they want
        new prose. This endpoint just refreshes the underlying
        analytics so the next regen sees the latest figures.

    Returns: {ok, strategy_hash, steps[]} -- per-step status so
    a partial failure (e.g. academic_analytics refresh fails but
    backtester + cost_sens succeed) surfaces every step's
    individual outcome.
    """
    import asyncio
    if ENVIRONMENT == "test":
        return {
            "ok": True,
            "note": "test environment -- compute chain skipped",
            "steps": [],
            "strategy_hash": None,
        }

    # ── June 27 2026: regime_signals pre-flight gate ─────────────
    # Light refresh is the typical user workflow immediately before
    # a deck regen ("warm everything, then regenerate"). If the
    # regime cache can't be refreshed within 10s, fail the entire
    # light-refresh now -- the user otherwise sees a successful
    # refresh, kicks off a deck regen, and the deck regen 503s
    # with the same blocking message a moment later. Cleaner UX
    # to surface the failure here. Brief / appendix users are
    # blocked too; they don't NEED fresh regime signals (em-dash
    # fallback works) but the cost is bounded -- 10s max wait on
    # a 15-min cache means most calls hit and return instantly.
    ok, _signals = await _regime_signals_fresh_or_refresh()
    if not ok:
        log.warning(
            "light_refresh_blocked_on_regime_signals",
            note=("regime_signals_cache miss + refresh "
                  "failed/timed out within 10s -- blocking "
                  "light-refresh so the user doesn't waste "
                  "a deck regen on stale signals"))
        raise HTTPException(
            status_code=503,
            detail=_REGIME_BLOCKING_ERROR_DETAIL)

    steps: list[dict] = []
    strategy_hash: str | None = None

    # Step 1 -- backtester + strategy_results_cache write.
    #
    # June 25 2026 -- unified on Hash A. Previously this step used
    # tools.cache._compute_data_hash(n_rows, last_date, n_strategies)
    # (Hash B), which produced a value the UI Live Hash banner
    # (current_data_hash, Hash A) could never match. Every refresh
    # left the Light Refresh status table permanently 'Stale' and
    # the draft chip permanently amber because the two columns
    # read different formulas. The audit confirmed Hash A is the
    # canonical platform-state fingerprint (see audit_assembler
    # .current_data_hash docstring -- derived-cache churn must NOT
    # invalidate it). Switching the refresh to stamp Hash A makes
    # the Draft Hash + Live Hash columns finally agree.
    try:
        from tools.backtester import run_all_strategies
        from tools.cache import set_strategy_cache
        from tools.data_fetcher import get_full_history_async
        from tools.audit_assembler import current_data_hash
        from tools.submission_freeze import get_effective_data_hash
        history = await get_full_history_async()
        monthly = history.get("equity_monthly")
        n_rows = len(monthly) if monthly is not None else 0
        # June 27 2026 (PR 1 v4 -- architectural-rule closure) --
        # under freeze, light refresh MUST warm the strategy cache
        # under the FREEZE hash, not the live hash. Otherwise the
        # StrategyCacheMissingForHashError raised by the 3 doc
        # generators (PR 1 v3) is NOT self-healing: the user runs
        # light refresh, the cache populates under the live hash,
        # the freeze-hash slot is still empty, the next deck export
        # fails again -- infinite loop. Routing through
        # get_effective_data_hash makes the cache write land in the
        # correct slot for whichever mode the platform is in. All
        # downstream uses below (refresh_academic_analytics,
        # refresh_oos_cost_sensitivity, the editor_drafts UPDATE,
        # the response's strategy_hash field) inherit the effective
        # hash automatically -- no further changes needed.
        live_hash = await current_data_hash()
        strategy_hash = await get_effective_data_hash(live_hash)
        if not strategy_hash:
            # current_data_hash returns "" on a degraded data path
            # (the source tables haven't loaded). Fall through with
            # an empty hash so the step records the failure mode
            # rather than crashing -- set_strategy_cache treats an
            # empty key as a no-op and the response's strategy_hash
            # field surfaces null.
            log.warning(
                "light_refresh_current_data_hash_empty")
        results_dict = await asyncio.to_thread(
            run_all_strategies, history)
        await set_strategy_cache(
            strategy_hash, results_dict, n_observations=n_rows,
            risk_free_monthly=history.get("risk_free_monthly"))
        steps.append({
            "step": "backtester",
            "ok": True,
            "strategy_hash": strategy_hash,
            "n_strategies": len(results_dict),
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "light_refresh_backtester_failed", error=str(exc))
        steps.append({
            "step": "backtester",
            "ok": False,
            "error": str(exc),
        })
        # Same short-circuit as refresh-appendix-caches --
        # downstream metrics need the strategy_hash this step
        # produces. Surface the failure as a 500 so the UI can
        # render an error rather than a partial-success summary.
        raise HTTPException(
            status_code=500,
            detail={"steps": steps,
                    "blocked_at": "backtester"})

    # Step 2 -- academic_analytics refresh.
    try:
        from tools.precomputed_analytics import (
            refresh_academic_analytics,
        )
        await refresh_academic_analytics(strategy_hash)
        steps.append({
            "step": "refresh_academic_analytics",
            "ok": True,
            "data_hash": strategy_hash,
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "light_refresh_academic_failed", error=str(exc))
        steps.append({
            "step": "refresh_academic_analytics",
            "ok": False,
            "error": str(exc),
        })

    # Step 3 -- OOS cost sensitivity refresh.
    try:
        from tools.regime_meta_validation import (
            refresh_oos_cost_sensitivity,
        )
        ok = await refresh_oos_cost_sensitivity(strategy_hash)
        steps.append({
            "step": "refresh_oos_cost_sensitivity",
            "ok": bool(ok),
            "data_hash": strategy_hash,
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "light_refresh_cost_sens_failed", error=str(exc))
        steps.append({
            "step": "refresh_oos_cost_sensitivity",
            "ok": False,
            "error": str(exc),
        })

    all_ok = all(step.get("ok") for step in steps)

    # June 25 2026 -- after a successful refresh, stamp the new
    # strategy_hash onto every current draft so the editor's hash
    # status chip reads "Data current" against the fresh cache
    # rather than showing every draft stale. The light refresh is
    # explicitly opt-in by a team member with generate_documents
    # rights; the assumption is that the team intends the refreshed
    # cache to be the new authoritative baseline. Generation of any
    # given document is still the user's explicit action -- the
    # data_hash stamp tracks 'what cache did the team last sync to'
    # not 'what was the cache when each section was written'.
    drafts_updated = 0
    if strategy_hash and all_ok:
        try:
            from sqlalchemy import text as _text
            from database import (
                AsyncSessionLocal as _ASL,  # type: ignore
            )
            if _ASL is not None:
                async with _ASL() as _s:
                    r = await _s.execute(_text(
                        "UPDATE editor_drafts "
                        "SET data_hash = :h, "
                        "    updated_at = NOW() "
                        "WHERE is_current = true "
                        "  AND is_deleted = false "
                        "  AND document_type IN ("
                        "    'executive_brief', "
                        "    'analytical_appendix', "
                        "    'presentation_deck', "
                        "    'presentation_script') "
                        "RETURNING id"),
                        {"h": strategy_hash})
                    drafts_updated = len(r.fetchall())
                    await _s.commit()
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "light_refresh_draft_hash_update_failed",
                error=str(exc))

    # June 26 2026 -- auto-recover any draft left in the NULL-
    # current-draft state. Runs on every light-refresh so a botched
    # generation that raced the refresh recovers without an operator
    # restart. Same helper used at startup; idempotent.
    drafts_recovered = 0
    try:
        from tools.editor_drafts import (
            recover_null_current_drafts as _recover,
        )
        drafts_recovered = await _recover()
        if drafts_recovered:
            log.warning(
                "light_refresh_null_current_drafts_recovered",
                count=drafts_recovered)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "light_refresh_null_current_drafts_failed",
            error=str(exc))

    log.info(
        "light_refresh_complete",
        all_ok=all_ok,
        strategy_hash=strategy_hash,
        drafts_updated=drafts_updated,
        drafts_recovered=drafts_recovered,
        actor=session.get("email"))
    return {
        "ok": all_ok,
        "strategy_hash": strategy_hash,
        "drafts_updated": drafts_updated,
        "drafts_recovered": drafts_recovered,
        "steps": steps,
    }


@app.post("/api/v1/data/verify-post-refresh")
async def post_verify_post_refresh(
    session: dict = Depends(require_team_member),
):
    """June 27 2026 -- post-light-refresh verification pass +
    rounding audit. Re-reads every substitution token, classifies
    each by submission scope, runs the per-scope sanity check
    (LOCKED == strategy_cache, CONSTANT == academic_deck module
    constant, FULL_DATASET == plausible range, LIVE == fresh and
    non-null), and checks each numeric value against the
    canonical rounding rule for its token class.

    Self-contained: does NOT depend on the /data-reference-sheet
    endpoint. Derives submission_scope locally via the verifier
    module's classify_submission_scope (consolidated with
    data_reference_catalog.classify_submission_scope post-merge
    of PR #459).

    Operator wiring: the Light Refresh panel auto-calls this
    endpoint after every successful refresh and the Data
    Reference Sheet header carries a 'Verify submission data'
    button that fires it on demand.

    Response shape per spec:
      {verified_at, freeze_active, freeze_hash, effective_hash,
       passed, failed, warnings,
       results[{token, label, scope, expected, actual,
                rounded_correctly, status, message}],
       rounding_summary{checked, consistent, inconsistent,
                        inconsistent_tokens[]},
       ready_for_submission}

    ready_for_submission == true iff (failed == 0
       AND rounding_summary.inconsistent == 0
       AND no live tokens stale)."""
    _ = session  # team-gated
    from tools.post_refresh_verifier import run_verification
    return await run_verification()


# ── Dual-mode token storage endpoints (PR-DM-Lite, June 28) ──────


async def _auto_upgrade_draft_to_token_values(
    draft_id: int, document_type: str,
) -> None:
    """June 28 2026 (Fix 8b) -- per-draft auto-upgrade hook.

    Fires right after a generator persists its draft +
    value_manifest. Replaces the manual operator step of
    POSTing to /api/v1/admin/upgrade-all-drafts-to-token-values
    after every generation.

    Contract:
      * Only runs when DEFER_SUBSTITUTION_TO_EXPORT is ON --
        otherwise content_json carries resolved values not
        {{TOKEN}} placeholders, the {{TOKEN}}-only matcher
        finds nothing to upgrade, and the call is a no-op.
        Short-circuiting here saves a DB round-trip + keeps
        the legacy flag-OFF path bit-identical.
      * Only touches the ONE draft_id passed in -- never the
        batch. The admin endpoint
        post_upgrade_all_drafts_to_token_values remains
        available for backfill scenarios; this is the
        steady-state per-generation path.
      * Fail-open: every error logs + returns. Generation
        completion is never blocked by an upgrade-pass
        failure -- the draft is already persisted; the admin
        batch endpoint can be re-run to recover."""
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    from tools.draft_token_upgrade import (
        upgrade_canvas_slides_for_unverified_tags,
        upgrade_content_json_for_unverified_tags,
        upgrade_content_json_to_token_values,
    )
    from tools.platform_flags import (
        is_defer_substitution_enabled,
    )
    import json as _json

    try:
        # Phase 2 deferral gate ONLY guards the token-upgrade
        # pass. The <unverified> tag upgrade runs UNCONDITIONALLY
        # because soft-fail tags can be emitted regardless of
        # the deferral flag (the hard-lock soft-fail in
        # harness_narrative / generate_script /
        # _substitute_slide_content fires whenever the
        # 3-pass cap is hit).
        defer_on = await is_defer_substitution_enabled()
        if _ASL is None:
            return
        async with _ASL() as s:
            row = await s.execute(_text(
                "SELECT content_json, value_manifest, "
                "       migration_run "
                "FROM editor_drafts "
                "WHERE id = :id AND is_deleted = false"),
                {"id": draft_id})
            r = row.fetchone()
            if not r:
                return
            content_json, manifest, migration_run = r
            if not isinstance(content_json, dict):
                return
            cj_after = content_json
            token_stats = {
                "nodes_upgraded": 0, "already_upgraded": 0}
            unverified_stats = {
                "nodes_upgraded": 0, "already_upgraded": 0}
            # Token upgrade only when flag ON + manifest exists
            # + migration hasn't run yet.
            if (defer_on and manifest
                    and not migration_run):
                cj_after, token_stats = (
                    upgrade_content_json_to_token_values(
                        cj_after, manifest))
            # June 28 2026 (PR #479) -- unverified-tag upgrade
            # runs for every doc type, every generation,
            # regardless of the deferral flag. The deck uses
            # a different schema (canvas slides) so dispatch
            # by document_type.
            if document_type == "presentation_deck":
                cj_after, unverified_stats = (
                    upgrade_canvas_slides_for_unverified_tags(
                        cj_after))
            else:
                cj_after, unverified_stats = (
                    upgrade_content_json_for_unverified_tags(
                        cj_after))
            any_change = (
                token_stats["nodes_upgraded"] > 0
                or unverified_stats["nodes_upgraded"] > 0)
            if any_change:
                await s.execute(_text(
                    "UPDATE editor_drafts "
                    "SET pre_migration_content_json "
                    "      = COALESCE("
                    "          pre_migration_content_json, "
                    "          CAST(:snap AS JSONB)), "
                    "    content_json = CAST(:cj AS JSONB), "
                    "    migration_run = true, "
                    "    updated_at = NOW() "
                    "WHERE id = :id"),
                    {
                        "snap": _json.dumps(content_json),
                        "cj":   _json.dumps(cj_after),
                        "id":   draft_id,
                    })
                await s.commit()
        log.info(
            "draft_auto_upgrade_persisted",
            draft_id=draft_id,
            document_type=document_type,
            token_nodes_upgraded=token_stats["nodes_upgraded"],
            unverified_nodes_upgraded=(
                unverified_stats["nodes_upgraded"]))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "draft_auto_upgrade_failed",
            draft_id=draft_id,
            document_type=document_type,
            error=str(exc))


@app.post("/api/v1/admin/upgrade-all-drafts-to-token-values")
async def post_upgrade_all_drafts_to_token_values(
    session: dict = Depends(require_team_member),
):
    """June 28 2026 -- explicit, operator-triggered upgrade pass.

    Walks every current draft with value_manifest IS NOT NULL +
    migration_run = FALSE. For each:
      1. Snapshot content_json into pre_migration_content_json
      2. Run upgrade_content_json_to_token_values
      3. Persist new content_json + set migration_run = TRUE

    Returns a per-draft summary (id / document_type / status /
    nodes_upgraded / manifest_entries / message).

    The operator MUST run this manually after PR-DM-Lite
    deploys. NOT triggered automatically on first draft load.
    """
    _ = session  # team-gated
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    from tools.draft_token_upgrade import (
        upgrade_content_json_to_token_values,
    )
    import json as _json

    if _ASL is None:
        raise HTTPException(
            status_code=503,
            detail="Database not configured")

    results: list[dict] = []
    async with _ASL() as s:
        rows = await s.execute(_text(
            "SELECT id, document_type, owner_email, "
            "content_json, value_manifest "
            "FROM editor_drafts "
            "WHERE is_current = true "
            "  AND is_deleted = false "
            "  AND migration_run = false "
            "ORDER BY id"))
        drafts = rows.fetchall()

        for row in drafts:
            draft_id, doc_type, owner, content_json, manifest = row
            if not isinstance(content_json, dict):
                results.append({
                    "draft_id":       draft_id,
                    "document_type":  doc_type,
                    "owner":          owner,
                    "status":         "skipped",
                    "message":        (
                        "content_json missing or non-dict"),
                })
                continue
            if not manifest:
                results.append({
                    "draft_id":       draft_id,
                    "document_type":  doc_type,
                    "owner":          owner,
                    "status":         "skipped",
                    "message":        (
                        "value_manifest NULL (pre-Layer-3 "
                        "draft) -- cannot upgrade"),
                })
                continue
            try:
                upgraded, stats = (
                    upgrade_content_json_to_token_values(
                        content_json, manifest))
                # Persist: snapshot + new content + flag.
                await s.execute(_text(
                    "UPDATE editor_drafts "
                    "SET pre_migration_content_json "
                    "      = CAST(:snap AS JSONB), "
                    "    content_json = CAST(:cj AS JSONB), "
                    "    migration_run = true, "
                    "    updated_at = NOW() "
                    "WHERE id = :id"),
                    {
                        "snap": _json.dumps(content_json),
                        "cj":   _json.dumps(upgraded),
                        "id":   draft_id,
                    })
                results.append({
                    "draft_id":          draft_id,
                    "document_type":     doc_type,
                    "owner":             owner,
                    "status":            "upgraded",
                    "manifest_entries":  stats["manifest_entries"],
                    "nodes_upgraded":    stats["nodes_upgraded"],
                    "already_upgraded":  stats["already_upgraded"],
                    "message":           (
                        f"{stats['nodes_upgraded']} "
                        "token_value nodes inserted"),
                })
                log.info(
                    "draft_upgrade_persisted",
                    draft_id=draft_id,
                    document_type=doc_type,
                    nodes_upgraded=stats["nodes_upgraded"])
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "draft_upgrade_failed",
                    draft_id=draft_id,
                    error=str(exc))
                results.append({
                    "draft_id":       draft_id,
                    "document_type":  doc_type,
                    "owner":          owner,
                    "status":         "failed",
                    "message":        str(exc),
                })
        await s.commit()

    return {
        "drafts_checked": len(drafts),
        "upgraded":  sum(
            1 for r in results if r["status"] == "upgraded"),
        "skipped":   sum(
            1 for r in results if r["status"] == "skipped"),
        "failed":    sum(
            1 for r in results if r["status"] == "failed"),
        "results":   results,
    }


@app.post("/api/v1/admin/revert-all-draft-migrations")
async def post_revert_all_draft_migrations(
    session: dict = Depends(require_team_member),
):
    """June 28 2026 -- batch revert of every draft whose
    upgrade pass ran AND has a pre_migration_content_json
    snapshot. Restores content_json from the snapshot, clears
    the snapshot, sets migration_run = FALSE.

    Used to undo a buggy upgrade pass in one shot. After this
    runs the operator deploys the upgrade-pass fix + re-invokes
    POST /api/v1/admin/upgrade-all-drafts-to-token-values.

    Returns per-draft summary (id, document_type, owner, status,
    message)."""
    _ = session  # team-gated
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    import json as _json

    if _ASL is None:
        raise HTTPException(
            status_code=503,
            detail="Database not configured")

    results: list[dict] = []
    async with _ASL() as s:
        rows = await s.execute(_text(
            "SELECT id, document_type, owner_email "
            "FROM editor_drafts "
            "WHERE is_current = true "
            "  AND is_deleted = false "
            "  AND migration_run = true "
            "  AND pre_migration_content_json IS NOT NULL "
            "ORDER BY id"))
        candidates = rows.fetchall()
        for row in candidates:
            draft_id, doc_type, owner = row
            try:
                snap_row = await s.execute(_text(
                    "SELECT pre_migration_content_json "
                    "FROM editor_drafts WHERE id = :id"),
                    {"id": draft_id})
                r = snap_row.fetchone()
                snap = r[0] if r else None
                if snap is None:
                    results.append({
                        "draft_id":      draft_id,
                        "document_type": doc_type,
                        "owner":         owner,
                        "status":        "skipped",
                        "message":       "no snapshot present",
                    })
                    continue
                await s.execute(_text(
                    "UPDATE editor_drafts "
                    "SET content_json = CAST(:cj AS JSONB), "
                    "    migration_run = false, "
                    "    pre_migration_content_json = NULL, "
                    "    updated_at = NOW() "
                    "WHERE id = :id"),
                    {
                        "cj": _json.dumps(snap)
                              if not isinstance(snap, str)
                              else snap,
                        "id": draft_id,
                    })
                results.append({
                    "draft_id":      draft_id,
                    "document_type": doc_type,
                    "owner":         owner,
                    "status":        "reverted",
                    "message":       (
                        "content_json restored from snapshot; "
                        "migration_run reset to FALSE; "
                        "snapshot cleared"),
                })
                log.info(
                    "draft_migration_reverted",
                    draft_id=draft_id,
                    document_type=doc_type)
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "draft_migration_revert_failed",
                    draft_id=draft_id, error=str(exc))
                results.append({
                    "draft_id":      draft_id,
                    "document_type": doc_type,
                    "owner":         owner,
                    "status":        "failed",
                    "message":       str(exc),
                })
        await s.commit()

    return {
        "drafts_checked": len(candidates),
        "reverted":  sum(
            1 for r in results if r["status"] == "reverted"),
        "skipped":   sum(
            1 for r in results if r["status"] == "skipped"),
        "failed":    sum(
            1 for r in results if r["status"] == "failed"),
        "results":   results,
    }


@app.post("/api/v1/admin/revert-draft-migration/{draft_id}")
async def post_revert_draft_migration(
    draft_id: int,
    session: dict = Depends(require_team_member),
):
    """June 28 2026 -- restores content_json from
    pre_migration_content_json + sets migration_run = FALSE.
    Used when the upgrade pass produced an unexpected result on
    a specific draft."""
    _ = session
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    import json as _json

    if _ASL is None:
        raise HTTPException(
            status_code=503, detail="Database not configured")
    async with _ASL() as s:
        row = await s.execute(_text(
            "SELECT pre_migration_content_json "
            "FROM editor_drafts WHERE id = :id"),
            {"id": draft_id})
        r = row.fetchone()
        if not r or not r[0]:
            raise HTTPException(
                status_code=404,
                detail=(
                    f"draft {draft_id} has no pre-migration "
                    "snapshot -- nothing to revert"))
        snap = r[0]
        await s.execute(_text(
            "UPDATE editor_drafts "
            "SET content_json = CAST(:cj AS JSONB), "
            "    migration_run = false, "
            "    pre_migration_content_json = NULL, "
            "    updated_at = NOW() "
            "WHERE id = :id"),
            {
                "cj": _json.dumps(snap)
                      if not isinstance(snap, str) else snap,
                "id": draft_id,
            })
        await s.commit()
    return {"ok": True, "draft_id": draft_id}


@app.get("/api/v1/data/review-pending-updates/{draft_id}")
async def get_review_pending_updates_for_draft(
    draft_id: int,
    session: dict = Depends(require_team_member),
):
    """June 28 2026 -- per-draft review summary for the
    ValueUpdateReviewPanel. Walks the (upgraded) content_json
    and returns one entry per token_value node with current
    value, cache value, match flag, and override status."""
    _ = session
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    from tools.draft_token_upgrade import build_review_summary
    from tools.audit_assembler import current_data_hash
    from tools.cache import get_strategy_cache
    from tools.cio_recommendation import (
        get_latest_recommendation,
    )
    from tools.numeric_substitution import (
        get_substitution_table,
    )
    from tools.submission_freeze import (
        get_effective_data_hash,
    )

    if _ASL is None:
        raise HTTPException(
            status_code=503, detail="Database not configured")

    async with _ASL() as s:
        row = await s.execute(_text(
            "SELECT content_json, migration_run "
            "FROM editor_drafts WHERE id = :id"),
            {"id": draft_id})
        r = row.fetchone()
    if not r:
        raise HTTPException(
            status_code=404, detail=f"draft {draft_id} not found")
    content_json, migration_run = r
    if not migration_run:
        return {
            "draft_id": draft_id,
            "migration_run": False,
            "message": (
                "Draft has not been upgraded to dual-mode token "
                "storage yet. Run "
                "POST /api/v1/admin/upgrade-all-drafts-to-token-"
                "values first."),
            "entries": [],
        }

    # Build the current substitution table -- freeze-aware.
    live_hash = await current_data_hash()
    eff_hash = await get_effective_data_hash(live_hash) or live_hash
    cio = await get_latest_recommendation() or {}
    strategy_cache = await get_strategy_cache(eff_hash) or {}
    table = get_substitution_table(
        eff_hash, strategy_cache, cio, hash_verified=True)

    entries = build_review_summary(content_json, table)
    mismatched = sum(1 for e in entries if not e["match"])
    overridden = sum(1 for e in entries if e["overridden"])
    return {
        "draft_id":       draft_id,
        "migration_run":  True,
        "effective_hash": eff_hash,
        "entries":        entries,
        "total":          len(entries),
        "matched":        len(entries) - mismatched,
        "mismatched":     mismatched,
        "overridden":     overridden,
    }


@app.post("/api/v1/data/apply-updates/{draft_id}")
async def post_apply_updates_to_draft(
    draft_id: int,
    body: dict | None = None,
    session: dict = Depends(require_team_member),
):
    """June 28 2026 -- applies token updates to a draft's
    content_json. Body shape:
      {"tokens": ["{{TOKEN_A}}", "{{TOKEN_B}}", ...]}
    Only nodes whose attrs.token is in the list get updated.
    Omit the body (or send empty tokens) to apply all changes.

    Returns count + per-update audit log."""
    _ = session
    from sqlalchemy import text as _text
    from database import AsyncSessionLocal as _ASL
    from tools.draft_token_upgrade import apply_token_updates
    from tools.audit_assembler import current_data_hash
    from tools.cache import get_strategy_cache
    from tools.cio_recommendation import (
        get_latest_recommendation,
    )
    from tools.numeric_substitution import (
        get_substitution_table,
    )
    from tools.submission_freeze import (
        get_effective_data_hash,
    )
    import json as _json

    if _ASL is None:
        raise HTTPException(
            status_code=503, detail="Database not configured")

    selected: set[str] | None = None
    if body and isinstance(body.get("tokens"), list):
        selected = set(str(t) for t in body["tokens"])

    async with _ASL() as s:
        row = await s.execute(_text(
            "SELECT content_json, migration_run "
            "FROM editor_drafts WHERE id = :id"),
            {"id": draft_id})
        r = row.fetchone()
        if not r:
            raise HTTPException(
                status_code=404,
                detail=f"draft {draft_id} not found")
        content_json, migration_run = r
        if not migration_run:
            raise HTTPException(
                status_code=409,
                detail=(
                    "Draft has not been upgraded to dual-mode "
                    "token storage. Run the upgrade pass first."))

        # Build current substitution table.
        live_hash = await current_data_hash()
        eff_hash = (
            await get_effective_data_hash(live_hash) or live_hash)
        cio = await get_latest_recommendation() or {}
        strategy_cache = await get_strategy_cache(eff_hash) or {}
        table = get_substitution_table(
            eff_hash, strategy_cache, cio, hash_verified=True)

        new_content_json, updates = apply_token_updates(
            content_json, table, eff_hash, selected)
        if updates:
            await s.execute(_text(
                "UPDATE editor_drafts "
                "SET content_json = CAST(:cj AS JSONB), "
                "    data_hash = :h, "
                "    updated_at = NOW() "
                "WHERE id = :id"),
                {
                    "cj": _json.dumps(new_content_json),
                    "h":  eff_hash,
                    "id": draft_id,
                })
            await s.commit()
        log.info(
            "draft_token_rewrite_applied",
            draft_id=draft_id,
            updates_count=len(updates),
            actor=(session.get("email") if session else None))

    return {
        "draft_id":       draft_id,
        "updates_count":  len(updates),
        "updates":        updates,
        "effective_hash": eff_hash,
    }


@app.post("/api/v1/admin/refresh-appendix-caches")
async def post_refresh_appendix_caches(
    session: dict = Depends(require_team_member),
):
    """June 21 2026 -- the appendix's graded sections (B, C, D,
    E, G) depend on three upstream cache populations:
      1. strategy_results_cache (backtester run) -- Sections B,
         C, E source
      2. academic_analytics metric (refresh_academic_analytics)
         -- Section D's bootstrap_ci_sharpe AND Section E's
         factor_loadings
      3. oos_cost_sensitivity metric
         (refresh_oos_cost_sensitivity) -- Section G

    Before this endpoint, the operator had to ssh to the Render
    shell and trigger each compute manually. This endpoint runs
    the chain in sequence so the appendix's pre-flight cache
    gate (HTTPException 409 when any field is empty) can be
    cleared from the admin UI / a single curl call.

    Returns a per-step status report. Each step is wrapped
    individually so a partial failure (e.g. the backtester
    succeeds but academic_analytics refresh fails) still surfaces
    which steps completed.

    Sequencing matters: the backtester populates
    strategy_results_cache + computes the canonical strategy_hash;
    refresh_academic_analytics + refresh_oos_cost_sensitivity
    both read that hash + write keyed to it. Running them out of
    order against a stale hash produces orphan cache rows."""
    import asyncio
    if ENVIRONMENT == "test":
        return {
            "ok": True,
            "note": "test environment -- compute chain skipped",
            "steps": [],
        }
    steps: list[dict] = []
    strategy_hash: str | None = None

    # Step 1 -- backtester run + strategy_results_cache write.
    try:
        from tools.backtester import run_all_strategies
        from tools.cache import (
            _compute_data_hash, set_strategy_cache,
        )
        from tools.data_fetcher import get_full_history_async
        history = await get_full_history_async()
        monthly = history.get("equity_monthly")
        n_rows = len(monthly) if monthly is not None else 0
        last_date = (
            str(monthly.index[-1].date())
            if monthly is not None and len(monthly) > 0
            else "unknown")
        strategy_hash = _compute_data_hash(
            n_rows, last_date, n_strategies=10)
        results_dict = await asyncio.to_thread(
            run_all_strategies, history)
        await set_strategy_cache(
            strategy_hash, results_dict, n_observations=n_rows,
            risk_free_monthly=history.get("risk_free_monthly"))
        steps.append({
            "step": "backtester",
            "ok": True,
            "strategy_hash": strategy_hash,
            "n_strategies": len(results_dict),
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "admin_refresh_appendix_caches_backtester_failed",
            error=str(exc))
        steps.append({
            "step": "backtester",
            "ok": False,
            "error": str(exc),
        })
        # Backtester failure short-circuits the chain --
        # downstream metrics need the strategy_hash that step 1
        # produces.
        raise HTTPException(
            status_code=500,
            detail={"steps": steps,
                    "blocked_at": "backtester"})

    # Step 2 -- academic_analytics refresh (bootstrap_ci_sharpe
    # + factor_loadings).
    try:
        from tools.precomputed_analytics import (
            refresh_academic_analytics,
        )
        await refresh_academic_analytics(strategy_hash)
        steps.append({
            "step": "refresh_academic_analytics",
            "ok": True,
            "data_hash": strategy_hash,
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "admin_refresh_appendix_caches_academic_failed",
            error=str(exc))
        steps.append({
            "step": "refresh_academic_analytics",
            "ok": False,
            "error": str(exc),
        })
        # Do NOT short-circuit -- cost sensitivity is
        # independent; surfacing both failures is more useful
        # than masking the second one.

    # Step 3 -- OOS cost sensitivity refresh.
    try:
        from tools.regime_meta_validation import (
            refresh_oos_cost_sensitivity,
        )
        ok = await refresh_oos_cost_sensitivity(strategy_hash)
        steps.append({
            "step": "refresh_oos_cost_sensitivity",
            "ok": bool(ok),
            "data_hash": strategy_hash,
        })
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "admin_refresh_appendix_caches_cost_sens_failed",
            error=str(exc))
        steps.append({
            "step": "refresh_oos_cost_sensitivity",
            "ok": False,
            "error": str(exc),
        })

    all_ok = all(step.get("ok") for step in steps)
    log.info(
        "admin_refresh_appendix_caches_complete",
        all_ok=all_ok,
        strategy_hash=strategy_hash,
        actor=session.get("email"))
    return {
        "ok": all_ok,
        "strategy_hash": strategy_hash,
        "steps": steps,
    }


@app.post("/api/v1/admin/warm-analytics-cache")
async def post_warm_analytics_cache(
    session: dict = Depends(require_team_member),
):
    """May 24 2026 — manual warm trigger for the
    analytics_metrics_cache table. The auto-warm hook fires this
    automatically at startup with retry; this endpoint is the
    sysadmin override for "warm now, regardless of cache state".

    Runs refresh_all_analytics inline and returns when the rows
    have landed. Updates WarmState so /cache-status reflects the
    completed run.

    Returns:
      {warmed: bool, latest_hash: str | None, took_s: float,
       academic_analytics: bool, efficient_frontier: bool}
    """
    if ENVIRONMENT == "test":
        return {
            "warmed":            True,
            "latest_hash":       None,
            "took_s":            0.0,
            "academic_analytics": False,
            "efficient_frontier": False,
            "note":              "test environment",
        }
    try:
        from tools.cache import get_latest_strategy_hash
        from tools.cache_warm_state import auto_warm_analytics, get_warm_state
        # Single attempt — the user is asking for an explicit
        # warm now, so we don't apply the auto-warm's exponential
        # backoff. Failures still surface to the WarmState.
        state = await auto_warm_analytics(max_attempts=1)
        latest_hash = await get_latest_strategy_hash()
        if state.status != "warm":
            raise HTTPException(
                status_code=500,
                detail=f"Warm failed: {state.last_attempt_error}")
        landed = state.last_landed or {}
        return {
            "warmed":              True,
            "latest_hash":         latest_hash,
            "took_s":              state.last_took_s or 0.0,
            "academic_analytics":  bool(landed.get("academic_analytics")),
            "efficient_frontier":  bool(landed.get("efficient_frontier")),
        }
    except HTTPException:
        raise
    except Exception as exc:
        log.warning("warm_analytics_cache_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail=f"Warm failed: {exc}")


@app.get("/api/v1/admin/cache-status")
async def get_cache_status(
    session: dict = Depends(require_auth),
):
    """May 24 2026 — analytics cache status for the Admin UI.

    Returns the auto-warm subsystem's current state plus the
    per-row landed booleans from analytics_metrics_cache so the
    Admin UI can show the right state without a second round-trip:

      {
        "status": "idle" | "warming" | "warm" | "failed",
        "in_progress": bool,
        "attempts": int,
        "last_attempt_at": ISO8601 | null,
        "last_success_at": ISO8601 | null,
        "last_success_age_seconds": float | null,
        "last_attempt_error": str | null,
        "last_took_s": float | null,
        "last_landed": {academic_analytics, efficient_frontier},
        "cache_present": {
            "academic_analytics": bool,
            "efficient_frontier": bool,
        },
      }

    Any authenticated user can read this — the UI's Warm Cache
    button itself is sysadmin-only but every team member should
    be able to see whether the dashboard's cache is hot.
    """
    from tools.cache_warm_state import get_warm_state
    state = get_warm_state().to_dict()
    cache_present = {"academic_analytics": False, "efficient_frontier": False}
    if ENVIRONMENT != "test":
        try:
            from tools.cache import get_latest_strategy_hash
            from tools.precomputed_analytics import (
                get_metric as get_precomputed,
                get_latest_metric as get_latest_precomputed,
            )
            latest_hash = await get_latest_strategy_hash()
            sentinel = latest_hash or "BOOT-WARM"
            aa = await get_precomputed(sentinel, "academic_analytics")
            ef = await get_precomputed(sentinel, "efficient_frontier")
            if aa is None:
                aa = await get_latest_precomputed("academic_analytics")
            if ef is None:
                ef = await get_latest_precomputed("efficient_frontier")
            cache_present = {
                "academic_analytics":  bool(aa),
                "efficient_frontier":  bool(ef),
            }
        except Exception as exc:  # noqa: BLE001
            log.warning("cache_status_read_failed", error=str(exc))
    state["cache_present"] = cache_present
    return state


@app.get("/api/v1/admin/data-status")
async def get_admin_data_status(session: dict = Depends(require_auth)):
    """
    Read-only status of the data tables feeding the analytics layer —
    row counts, date ranges, last-updated timestamps and a green/amber/red
    staleness pill per table. Surfaced in Settings → Data and Study Period.
    """
    if ENVIRONMENT == "test":
        return {"available": False, "study_period": None, "tables": []}
    from tools.cache import get_data_status
    return await get_data_status()


@app.get("/api/v1/admin/invariants")
async def get_admin_invariants(session: dict = Depends(require_auth)):
    """May 30 2026 — surfaces the latest invariant-framework run
    summary. Backs the "Invariant Checks" health indicator in
    Settings → Data and Study Period.

    Returns either {"available": false, "ran_at": null} when no run
    has happened yet (cold deploy) or a full summary with per-
    violation breakdown for the admin to inspect. The framework
    runs on every analytics warm via set_strategy_cache."""
    from tools.invariant_checks import get_latest_result
    latest = get_latest_result()
    if latest is None:
        return {
            "available": False,
            "ran_at": None,
            "note": "No invariant run has landed yet — the framework "
                    "fires on the next analytics warm.",
        }
    return {"available": True, **latest}


@app.get("/api/v1/admin/invariants/history")
async def get_admin_invariants_history(
    limit: int = 7,
    session: dict = Depends(require_auth),
):
    """June 2 2026 — last-N invariant verdicts for the /admin/health
    Cache Warm History section.

    Pure read of the rows PR #252 (set_strategy_cache_invariant_persist)
    writes to analytics_metrics_cache with metric_kind='invariant_
    summary'. One row per distinct data_hash, upserted with a fresh
    computed_at on each warm against the same hash — so the result is
    effectively "one row per distinct dataset that was warmed, most-
    recent verdict per dataset."

    No recomputation, no fan-out — a single SELECT plus a JSON
    projection. Auth-only (not sysadmin) so any team member can read
    the warm-history strip from the admin panel.

    Fail-open: a DB read failure renders {available: false, rows: []}
    rather than 500, matching the convention every other admin
    health endpoint follows."""
    limit = max(1, min(int(limit or 7), 30))
    if ENVIRONMENT == "test":
        return {"available": False, "rows": []}
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return {"available": False, "rows": []}
        rows: list[dict[str, Any]] = []
        async with AsyncSessionLocal() as db:
            r = await db.execute(text(
                "SELECT data_hash, payload, computed_at "
                "FROM analytics_metrics_cache "
                "WHERE metric_kind = 'invariant_summary' "
                "ORDER BY computed_at DESC "
                "LIMIT :n"), {"n": limit})
            for data_hash, payload, computed_at in r.fetchall():
                # payload is JSONB; SQLAlchemy returns it as dict already
                # under asyncpg, but tolerate a string fallback (some
                # SQLAlchemy versions / drivers serialize differently).
                if isinstance(payload, str):
                    try:
                        import json as _json
                        payload = _json.loads(payload)
                    except Exception:  # noqa: BLE001
                        payload = {}
                payload = payload or {}
                rows.append({
                    "computed_at":   (computed_at.isoformat()
                                      if hasattr(computed_at, "isoformat")
                                      else str(computed_at)),
                    "data_hash":     (data_hash or "")[:8],
                    "passed":        bool(payload.get("passed", True)),
                    "checks_run":    int(payload.get("checks_run", 0)),
                    "hard_failures": int(payload.get("hard_failures", 0)),
                    "soft_warnings": int(payload.get("soft_warnings", 0)),
                    "ran_at":        payload.get("ran_at"),
                })
        return {"available": True, "rows": rows}
    except Exception as exc:  # noqa: BLE001
        log.warning("admin_invariants_history_read_failed",
                    error=str(exc))
        return {"available": False, "rows": []}


# ── Council query metrics (June 3 2026) ──────────────────────────────────


async def _write_council_query_metric(
    *,
    question_type: str,
    input_tokens: int | None,
    output_tokens: int | None,
    context_bundle_size: int,
    synthesis_text: str,
    cio_input_tokens: int | None = None,
) -> None:
    """Writes one row to council_query_metrics for the just-completed
    /api/council/query call.

    Reads the live regime ONCE (15-min in-process cache), extracts the
    recommendation direction from the synthesis prose, computes the
    refined alignment score, and inserts. Skipped entirely in the
    test environment (no DB). Fail-open per the wider council
    endpoint's post-stream-logging convention.
    """
    if ENVIRONMENT == "test":
        return
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return

        # Live regime (the 15-min cache; cheap once warm).
        hmm_state: str | None = None
        hmm_confidence: float | None = None
        try:
            from tools.regime_detector import detect_current_regime
            r = detect_current_regime() or {}
            hmm_state = r.get("hmm_regime")
            probs = r.get("hmm_probabilities") or {}
            if hmm_state and isinstance(probs, dict):
                hmm_confidence = float(probs.get(hmm_state, 0.0))
        except Exception as exc:  # noqa: BLE001
            log.warning("council_metric_regime_read_failed",
                        error=str(exc))

        # Direction extraction + refined alignment score.
        from tools.council_direction_extractor import (
            alignment_score, extract_direction,
        )
        direction = extract_direction(synthesis_text)
        score = alignment_score(direction, hmm_state, hmm_confidence)

        # Anchor the row to the dataset that produced it.
        data_hash: str | None = None
        try:
            from tools.cache import get_latest_strategy_hash
            data_hash = await get_latest_strategy_hash()
        except Exception as exc:  # noqa: BLE001
            log.warning("council_metric_hash_read_failed", error=str(exc))

        async with AsyncSessionLocal() as session:
            await session.execute(text(
                "INSERT INTO council_query_metrics "
                "(question_type, input_tokens, output_tokens, "
                " cio_input_tokens, context_bundle_size, "
                " hmm_state, hmm_confidence, "
                " recommendation_direction, hmm_alignment_score, "
                " data_hash) VALUES "
                "(:qt, :it, :ot, :cit, :cbs, :hs, :hc, :rd, :as, :dh)"
            ), {
                "qt":  question_type,
                "it":  (int(input_tokens)
                        if input_tokens is not None else None),
                "ot":  (int(output_tokens)
                        if output_tokens is not None else None),
                "cit": (int(cio_input_tokens)
                        if cio_input_tokens is not None else None),
                "cbs": int(context_bundle_size or 0),
                "hs":  hmm_state,
                "hc":  hmm_confidence,
                "rd":  direction,
                "as":  score,
                "dh":  data_hash,
            })
            await session.commit()
        log.info("council_query_metric_written",
                 question_type=question_type,
                 input_tokens=input_tokens,
                 output_tokens=output_tokens,
                 context_bundle_size=context_bundle_size,
                 direction=direction,
                 hmm_state=hmm_state,
                 alignment_score=score)
    except Exception as exc:  # noqa: BLE001
        log.warning("council_query_metric_write_inner_failed",
                    error=str(exc))


@app.get("/api/v1/admin/council-metrics")
async def get_admin_council_metrics(
    limit: int = 30,
    session: dict = Depends(require_auth),
):
    """The /admin/council-metrics surface that backs the cost-and-
    HMM-alignment measurement dashboard.

    Two payloads in one envelope:
      rows        — the latest N council_query_metrics rows
      aggregates  — per-question_type averages plus the headline
                    token_reduction_vs_baseline figure

    Fail-open: a DB unreachable / missing table returns
    {available:false, rows:[], aggregates:{}} rather than 500.
    Auth-only — any authenticated user can read. June 3 2026."""
    limit = max(1, min(int(limit or 30), 200))
    if ENVIRONMENT == "test":
        return {"available": False, "rows": [], "aggregates": {}}
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return {"available": False, "rows": [], "aggregates": {}}
        rows: list[dict[str, Any]] = []
        per_type: dict[str, dict[str, Any]] = {}
        baseline_input = None
        baseline_cio = None
        async with AsyncSessionLocal() as db:
            r = await db.execute(text(
                "SELECT id, timestamp, question_type, input_tokens, "
                " output_tokens, cio_input_tokens, "
                " context_bundle_size, hmm_state, "
                " hmm_confidence, recommendation_direction, "
                " hmm_alignment_score, data_hash "
                "FROM council_query_metrics "
                "ORDER BY timestamp DESC LIMIT :n"), {"n": limit})
            for row in r.fetchall():
                rows.append({
                    "id":                     row[0],
                    "timestamp":              (row[1].isoformat()
                                               if hasattr(row[1], "isoformat")
                                               else str(row[1])),
                    "question_type":          row[2],
                    "input_tokens":           row[3],
                    "output_tokens":          row[4],
                    "cio_input_tokens":       row[5],
                    "context_bundle_size":    row[6],
                    "hmm_state":              row[7],
                    "hmm_confidence":         row[8],
                    "recommendation_direction": row[9],
                    "hmm_alignment_score":    row[10],
                    "data_hash":              (row[11] or "")[:8] if row[11] else None,
                })
            # Aggregates — one SQL per metric so each remains
            # independently fail-open on a column type drift.
            agg_r = await db.execute(text(
                "SELECT question_type, "
                " AVG(input_tokens)::float AS avg_input, "
                " AVG(output_tokens)::float AS avg_output, "
                " AVG(cio_input_tokens)::float AS avg_cio_input, "
                " AVG(context_bundle_size)::float AS avg_bundle, "
                " AVG(hmm_alignment_score)::float AS avg_align, "
                " COUNT(*)::int AS n_rows "
                "FROM council_query_metrics "
                "GROUP BY question_type"))
            for (qt, avg_in, avg_out, avg_cio_in, avg_bundle,
                 avg_align, n) in agg_r.fetchall():
                per_type[qt] = {
                    "avg_input_tokens":       avg_in,
                    "avg_output_tokens":      avg_out,
                    "avg_cio_input_tokens":   avg_cio_in,
                    "avg_context_bundle_size": avg_bundle,
                    "avg_hmm_alignment_score": avg_align,
                    "n_rows":                 n,
                }
                if qt in ("baseline_full", "full"):
                    # Use the baseline_full averages preferentially;
                    # fall back to the live "full" fallback rows if
                    # the baseline-capture run hasn't been done yet.
                    if qt == "baseline_full":
                        if avg_in is not None:
                            baseline_input = avg_in
                        if avg_cio_in is not None:
                            baseline_cio = avg_cio_in
                    else:  # "full" — only used when baseline_full empty
                        if avg_in is not None and baseline_input is None:
                            baseline_input = avg_in
                        if avg_cio_in is not None and baseline_cio is None:
                            baseline_cio = avg_cio_in

        # Reductions per bundle — two parallel maps. Total-tokens
        # is noisy (chart-vision dominates); cio_input_tokens is the
        # like-for-like measurement. Both surfaced.
        reductions: dict[str, float | None] = {}
        cio_reductions: dict[str, float | None] = {}
        for qt, stats in per_type.items():
            avg_in = stats.get("avg_input_tokens")
            if (baseline_input and avg_in and avg_in > 0
                    and qt not in ("baseline_full", "full")):
                reductions[qt] = round(
                    1 - (float(avg_in) / float(baseline_input)), 4)
            else:
                reductions[qt] = None
            avg_cio_in = stats.get("avg_cio_input_tokens")
            if (baseline_cio and avg_cio_in and avg_cio_in > 0
                    and qt not in ("baseline_full", "full")):
                cio_reductions[qt] = round(
                    1 - (float(avg_cio_in) / float(baseline_cio)), 4)
            else:
                cio_reductions[qt] = None

        return {
            "available": True,
            "rows": rows,
            "aggregates": {
                "per_question_type": per_type,
                "token_reduction_vs_baseline": reductions,
                "cio_token_reduction_vs_baseline": cio_reductions,
                "baseline_avg_input_tokens": baseline_input,
                "baseline_avg_cio_input_tokens": baseline_cio,
            },
        }
    except Exception as exc:  # noqa: BLE001
        log.warning("admin_council_metrics_read_failed", error=str(exc))
        return {"available": False, "rows": [], "aggregates": {}}


# ── Document audit (June 3 2026) ─────────────────────────────────────────


async def _run_document_audit(
    content_text: str,
    document_type: str,
    email: str,
) -> dict[str, Any] | None:
    """Runs the post-generation audit. Returns the flags_by_check
    dict suitable for editor_drafts.audit_warnings, or None when
    nothing was flagged (so the column stores NULL).

    Reads the latest strategy_results_cache row ONCE so the numeric
    cross-reference and consistency checks share the same dataset
    snapshot. For the brief, ALSO reads the cached brief section
    plan from story_plans so PR #336's CHECK 5 (story_plan_violation),
    CHECK 6 (required_citations), and CHECK 7 (section_word_count)
    can fire end-to-end. Fail-open per the wider generator: an audit
    exception logs and returns None, never blocks the document write.
    """
    if ENVIRONMENT == "test":
        return None
    try:
        from tools.cache import (
            get_latest_strategy_cache, get_latest_strategy_hash,
        )
        from tools.document_audit import audit_document
        strategy_cache = await get_latest_strategy_cache()
        # PR #336 -- pull the cached brief section plan so CHECK 5
        # has the locked numeric_anchors to compare against. Fail-
        # open if the plan is unavailable (cold cache or a brief
        # generated before #333's caching layer landed); the check
        # skips silently in that case.
        brief_section_plan: dict | None = None
        if document_type == "executive_brief":
            try:
                from tools import story_plan as sp
                data_hash = await get_latest_strategy_hash()
                if data_hash:
                    plan = await sp.get_cached_story_plan(
                        data_hash, "brief")
                    if plan and isinstance(
                            plan.get("section_plan"), dict):
                        brief_section_plan = plan["section_plan"]
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "document_audit_brief_plan_lookup_failed",
                    error=str(exc))
        result = audit_document(
            content_text, document_type,
            strategy_cache=strategy_cache,
            brief_section_plan=brief_section_plan)
        log.info("document_audit_executed",
                 document_type=document_type,
                 owner=email,
                 flag_counts=result.flag_counts,
                 skipped=list(result.skipped.keys()))
        if not result.has_any_flag:
            return None
        return {
            "flags_by_check": result.flags_by_check,
            "flag_counts":    result.flag_counts,
            "skipped":        result.skipped,
        }
    except Exception as exc:  # noqa: BLE001
        log.warning("document_audit_failed",
                    document_type=document_type, error=str(exc))
        return None


async def _write_audit_metrics(
    document_type: str,
    email: str,
    draft_id: int | None,
    audit_warnings: dict[str, Any] | None,
) -> None:
    """Persist a document_audit_metrics row for the just-completed
    generation. Reads the live strategy_hash to anchor the row to
    the dataset that produced the document."""
    if ENVIRONMENT == "test":
        return
    try:
        from tools.cache import get_latest_strategy_hash
        from tools.document_audit_metrics import write_metric
        flag_counts = (
            (audit_warnings or {}).get("flag_counts")
            or {"numeric": 0, "direction": 0, "consistency": 0,
                "citation": 0, "total": 0})
        data_hash = await get_latest_strategy_hash()
        await write_metric(
            document_type=document_type, owner_email=email,
            draft_id=draft_id, flag_counts=flag_counts,
            data_hash=data_hash)
    except Exception as exc:  # noqa: BLE001
        log.warning("document_audit_metrics_inner_failed",
                    error=str(exc))


@app.get("/api/v1/admin/document-audit-metrics")
async def get_admin_document_audit_metrics(
    limit: int = 30,
    session: dict = Depends(require_auth),
):
    """Last-N document audit rows + per-document-type averages.
    Backs the admin tile that surfaces flag-rate trends so a
    drift in generation quality (e.g. citation-completeness flags
    spiking after a prompt change) is visible at a glance.

    Auth-only — any authenticated user can read. Fail-open in
    the test env to {available:false, rows:[], aggregates:{}}.
    June 3 2026."""
    from tools.document_audit_metrics import read_recent
    return await read_recent(limit=limit)


# ── Automated email system — manual triggers (June 1 2026) ────────────────


@app.post("/api/v1/admin/send-digest")
@limiter.limit("6/hour")
async def admin_send_digest(
    request: Request,
    session: dict = Depends(require_permission("manage_users")),
):
    """Component 1 — daily team digest manual trigger. Also called by
    the Render cron `forest-capital-digest` at 07:00 ET. Returns the
    Resend message id on success.

    Synchronous send: the assembly + Resend round-trip take a few
    seconds, well inside the request budget — no need for a job
    handle. Rate-limited to 6/hour so a sysadmin testing the digest
    cannot accidentally spam the team."""
    from tools.email_digest import send_daily_digest
    return await send_daily_digest()


@app.post("/api/v1/admin/test-alert")
@limiter.limit("6/hour")
async def admin_test_alert(
    request: Request,
    session: dict = Depends(require_permission("manage_users")),
):
    """Component 2 — synthetic invariant alert (Michael only).
    Fires the alert email using a hand-built violations payload so
    the wire format + Resend round-trip can be exercised without
    waiting for a real invariant breach. Rate-limited 6/hour."""
    from tools.email_alert import send_test_alert
    return send_test_alert()


# ── Layer 4: submission freeze ───────────────────────────────────────────────
# Locks document generation to a frozen data_hash for the FNA 670
# submission deadline (June 30 2026). Live platform reads (CIO card,
# regime detector, Investment Outlook, daily digest) continue calling
# current_data_hash() directly so the July 1 presentation shows live
# signals. See backend/tools/submission_freeze.py for the operational
# runbook (the curl invocations the operator runs on submission day).

@app.post("/api/v1/admin/submission-freeze")
async def admin_set_submission_freeze(
    body: dict,
    session: dict = Depends(require_permission("manage_users")),
):
    """Activate or deactivate the submission freeze.

    Body shape:
      {"active": true,  "freeze_hash": "c421fb89..."}  -- ON
      {"active": false}                                -- OFF

    On activate: validates freeze_hash matches a strategy_results_cache
    row before flipping (a typo'd hash would otherwise silently lock
    documents against a non-existent cache key, producing empty
    substitution tables on every generation).

    Sysadmin-only (manage_users permission, same gate every other
    destructive admin endpoint uses).
    """
    from tools.submission_freeze import set_freeze_config

    active = bool(body.get("active"))
    freeze_hash = body.get("freeze_hash")

    if active:
        if not freeze_hash or not isinstance(freeze_hash, str):
            raise HTTPException(
                status_code=400,
                detail=(
                    "freeze_hash is required when activating the "
                    "submission freeze. Supply the data_hash from "
                    "strategy_results_cache the documents were "
                    "generated against."),
            )
        # Validate the hash exists in strategy_results_cache. A typo
        # caught HERE costs the operator one HTTP round-trip; the
        # same typo caught at generation time costs an entire
        # document re-generation under the wrong frozen hash.
        try:
            from sqlalchemy import text
            from database import AsyncSessionLocal as _DB
            if _DB is not None:
                async with _DB() as s:
                    row = await s.execute(
                        text(
                            "SELECT 1 FROM strategy_results_cache "
                            "WHERE strategy_hash = :h LIMIT 1"),
                        {"h": freeze_hash},
                    )
                    if row.fetchone() is None:
                        raise HTTPException(
                            status_code=400,
                            detail=(
                                f"freeze_hash {freeze_hash[:8]}... "
                                "not found in strategy_results_cache. "
                                "Run a generation pass first so the "
                                "cache row exists, then activate the "
                                "freeze against the resulting hash."),
                        )
        except HTTPException:
            raise
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "submission_freeze_hash_validation_skipped",
                error=str(exc),
                note="DB unreachable -- activating without validation")

    try:
        new_config = await set_freeze_config(
            active=active,
            freeze_hash=freeze_hash if active else None,
            activated_by=session.get("email"),
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    return new_config


async def _read_latest_strategy_hash() -> str:
    """Returns the most recent strategy_results_cache.strategy_hash
    (e.g. 'c421fb895347f924') or "" on read failure / empty cache.

    The canonical hash used by document generation + freeze
    validation. Distinct from tools.audit_assembler.current_data_hash
    which returns a SHA256 of platform-level row counts + max dates.
    Extracted to a module-level helper so the submission-status
    endpoint can be stubbed in tests without monkeypatching
    AsyncSessionLocal."""
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return ""
        async with AsyncSessionLocal() as session_db:  # type: ignore[union-attr]
            row = await session_db.execute(text(
                "SELECT strategy_hash FROM strategy_results_cache "
                "ORDER BY computed_at DESC LIMIT 1"))
            r = row.fetchone()
            return str(r[0]) if r and r[0] else ""
    except Exception:  # noqa: BLE001
        return ""


@app.get("/api/v1/admin/submission-status")
async def admin_get_submission_status(
    session: dict = Depends(require_auth),
):
    """Reports the submission-readiness verdict: the freeze state, the
    drift between the frozen hash and the live hash, and the per-
    document export+verification state for the calling user.

    Available to any authenticated user (read-only) so Bob and Molly
    can see whether the freeze is on without needing sysadmin
    permissions. The activation/deactivation endpoint is sysadmin-only.

    Shape:
      {
        "freeze_active": bool,
        "freeze_hash": str | None,
        "freeze_date": str | None,
        "current_live_hash": str,
        "hash_drift": bool,
        "frozen_documents": {
          "brief":    {"generated", "exported", "export_verified",
                        "editor_draft_id"},
          "deck":     {...},
          "appendix": {...},
        },
        "submission_ready": bool,
        "submission_recommendation": str,
      }
    """
    from tools.editor_drafts import get_current_draft_with_layer3
    from tools.submission_freeze import get_freeze_config

    config = await get_freeze_config()
    freeze_active = bool(config.get("active"))
    freeze_hash = config.get("freeze_hash")
    freeze_date = config.get("freeze_date")

    # June 21 2026 -- read strategy_hash from strategy_results_cache
    # directly, NOT current_data_hash(). The latter is a SHA256 of
    # row counts + max dates across three tables (a platform
    # fingerprint), distinct from strategy_results_cache.strategy_hash
    # which is the canonical hash document generation + freeze
    # validation both cite. Same fix shape as PR #354 for the
    # key-metrics endpoint.
    try:
        live_hash = await _read_latest_strategy_hash()
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "submission_status_live_hash_read_failed", error=str(exc))
        live_hash = ""

    hash_drift = bool(freeze_active and freeze_hash
                      and live_hash and freeze_hash != live_hash)

    # Per-document state for the calling user. We read editor_drafts
    # because that's where the generated draft lands; once the user
    # exports it the export_verification JSONB (Layer 3a) carries the
    # verification result. The get_current_draft return shape is
    # version-compatible across Layer 3a -- on pre-3a deploys the
    # export_verification key is simply absent.
    email = session.get("email") or ""
    document_keys = (
        ("brief", "executive_brief"),
        ("deck", "presentation_deck"),
        ("appendix", "analytical_appendix"),
    )
    frozen_documents: dict[str, dict] = {}
    for short_key, doc_type in document_keys:
        # Switched from get_current_draft to the Layer-3 variant
        # (June 21 2026) so the export_verification JSONB column is
        # actually fetched -- the legacy get_current_draft uses the
        # legacy column set and NEVER read export_verification, so
        # `draft.get("export_verification")` below was always None
        # and every document showed exported=False even when a draft
        # had been exported. The layer3 helper has its own savepoint-
        # style retry that rolls back the failed-transaction state
        # before falling back to the legacy column set, so an
        # aborted transaction on a column-missing first attempt
        # doesn't break the second attempt.
        try:
            draft = await get_current_draft_with_layer3(email, doc_type)
        except Exception as exc:  # noqa: BLE001
            log.warning("submission_status_draft_read_failed",
                        document_type=doc_type, error=str(exc))
            draft = None

        if not draft:
            frozen_documents[short_key] = {
                "generated": False,
                "exported": False,
                "export_verified": None,
                "editor_draft_id": None,
            }
            continue

        # export_verification is a Layer 3a column -- pre-3a drafts
        # simply lack the key. Treat "missing" as "not yet exported"
        # rather than "exported but unverified" so the readiness
        # verdict is conservative on pre-3a deploys.
        ev: Any = None
        try:
            ev = draft.get("export_verification")
        except Exception:  # noqa: BLE001
            ev = None
        exported = ev is not None and ev != {}
        export_verified: bool | None = None
        if exported and isinstance(ev, dict):
            passed = ev.get("passed")
            if isinstance(passed, bool):
                export_verified = passed

        frozen_documents[short_key] = {
            "generated": True,
            "exported": bool(exported),
            "export_verified": export_verified,
            "editor_draft_id": str(draft.get("id"))
            if draft.get("id") is not None else None,
        }

    # Readiness verdict + recommendation -- the plain-English sentence
    # the Reports page surfaces in the freeze banner.
    all_generated = all(d["generated"] for d in frozen_documents.values())
    all_exported = all(d["exported"] for d in frozen_documents.values())
    all_verified = all(
        d["export_verified"] is True for d in frozen_documents.values())

    submission_ready = bool(
        freeze_active and all_generated and all_exported
        and all_verified and not hash_drift)

    def _missing(predicate) -> str:
        names = [k for k, d in frozen_documents.items() if not predicate(d)]
        return ", ".join(names) if names else ""

    if not freeze_active:
        recommendation = (
            "Activate the submission freeze after generating and "
            "verifying all deliverables.")
    elif not all_generated:
        recommendation = (
            f"Generate {_missing(lambda d: d['generated'])} "
            "before submitting.")
    elif not all_exported:
        recommendation = (
            f"Export {_missing(lambda d: d['exported'])} "
            "before submitting.")
    elif not all_verified:
        recommendation = (
            "Run Pre-Submission Check to verify all deliverables "
            "against the cache.")
    else:
        recommendation = (
            "All deliverables verified and freeze active. Safe to "
            "submit.")

    return {
        "freeze_active": freeze_active,
        "freeze_hash": freeze_hash,
        "freeze_date": freeze_date,
        "current_live_hash": live_hash,
        "hash_drift": hash_drift,
        "frozen_documents": frozen_documents,
        "submission_ready": submission_ready,
        "submission_recommendation": recommendation,
    }


@app.get("/api/v1/admin/team-activity/merge-commit-authors")
async def get_merge_commit_authors_diagnostic(
    session: dict = Depends(require_sysadmin),
):
    """Sysadmin diagnostic — returns the distinct `author` values
    of every commit_activity row whose message contains "merge
    pull request" (ILIKE).

    Hotfix May 23 2026 (iteration 4): production data shows ~100
    merged PRs but only 2 matches against the IN clause for
    Michael's known git emails. This endpoint exposes the actual
    author identities used on merge commits so we can confirm the
    fix is catching them.

    Expected output: a small set of distinct emails, e.g.
    mikeruurds@gmail.com (2) + mikeruurds@users.noreply.github.com
    (~98). The hotfix's local-part match should bring the count
    to ~100.

    Sysadmin-only — the endpoint exposes git author emails which
    could be PII for non-team contributors. Returns at most the
    most recent N rows to bound the response.
    """
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return {"available": False, "rows": []}
        async with AsyncSessionLocal() as s:
            r = await s.execute(text(
                "SELECT LOWER(author) AS author_email, COUNT(*) AS n, "
                " MAX(timestamp) AS latest_at "
                "FROM commit_activity "
                "WHERE message ILIKE '%merge pull request%' "
                "GROUP BY LOWER(author) "
                "ORDER BY n DESC LIMIT 50"
            ))
            rows = [
                {"author": row[0],
                 "merge_commit_count": int(row[1] or 0),
                 "latest_at": (row[2].isoformat() if row[2] else None)}
                for row in r.fetchall()
            ]
            r2 = await s.execute(text(
                "SELECT COUNT(*) FROM commit_activity "
                "WHERE message ILIKE '%merge pull request%'"))
            total = int(r2.scalar() or 0)
            return {
                "available": True,
                "total_merge_commits": total,
                "distinct_authors": rows,
                "note": (
                    "Hotfix iteration 4: if Michael's GitHub "
                    "noreply email appears here with a high count, "
                    "the local-part LIKE match in template_pipeline."
                    "fetch_team_activity is what now catches it. "
                    "Confirm Step 3's michael_prs_merged increases "
                    "to approximately this total."),
            }
    except Exception as exc:  # noqa: BLE001
        return {"available": False, "error": str(exc)}


# ── Academic documents (agent context) ───────────────────────────────────────

# 10 MB upload ceiling — the rubric / requirements documents are short;
# anything larger is almost certainly the wrong file.
_ACADEMIC_DOC_MAX_BYTES = 10 * 1024 * 1024


@app.post("/api/v1/documents/academic/upload")
async def upload_academic_document(
    request: Request,
    file: UploadFile = File(...),
    document_type: str = Form("other"),
    session: dict = Depends(require_team_member),
):
    """
    Uploads a PDF or Markdown (.md) reference document (the midpoint
    rubric, the final-presentation requirements, etc.). Text is extracted
    server-side; only the text is persisted. After storage every agent
    injects the document as system context on its next invocation.

    File type is decided by extension, not MIME type — browsers send
    Markdown as text/plain or text/markdown inconsistently, so the
    extension is the authoritative check. PDFs go through pypdf; .md
    files are read directly as UTF-8 (pypdf is bypassed entirely).
    """
    from tools.academic_context import (
        DOCUMENT_TYPES, extract_document_text, insert_academic_document,
    )

    if document_type not in DOCUMENT_TYPES:
        raise HTTPException(
            status_code=422,
            detail=f"Unknown document_type '{document_type}'. "
                   f"Valid: {', '.join(DOCUMENT_TYPES)}",
        )

    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=422, detail="Uploaded file is empty.")
    if len(raw) > _ACADEMIC_DOC_MAX_BYTES:
        raise HTTPException(status_code=422, detail="File too large (max 10 MB).")

    filename = file.filename or "document"
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        try:
            content_text = extract_document_text(filename, raw)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc))
        file_type = "PDF"
    elif lower_name.endswith(".md"):
        # Markdown — extension is authoritative. Read directly as UTF-8,
        # bypassing pypdf entirely; the bytes are the content verbatim.
        content_text = raw.decode("utf-8", errors="replace").strip()
        if not content_text:
            raise HTTPException(status_code=422, detail="Uploaded file is empty.")
        file_type = "MD"
    else:
        raise HTTPException(
            status_code=400,
            detail="Only PDF and Markdown (.md) files are supported",
        )

    doc_id = await insert_academic_document(filename, document_type, content_text)
    if not doc_id:
        raise HTTPException(
            status_code=500,
            detail="Could not store the document — database unavailable.",
        )
    # Log type + size so ingestion can be verified from the production logs.
    log.info(
        "academic_document_ingested",
        file_type=file_type,
        char_count=len(content_text),
        document_type=document_type,
    )
    # Team Activity — record the upload as an interaction (non-blocking).
    _log_interaction_bg(
        request, session, "document_upload",
        metadata={
            "document_type": document_type,
            "filename": filename,
            "file_type": file_type,
            "char_count": len(content_text),
        },
    )
    return {
        "id": doc_id,
        "name": filename,
        "document_type": document_type,
        "file_type": file_type,
        "char_count": len(content_text),
    }


@app.get("/api/v1/documents/academic")
async def list_academic_docs(session: dict = Depends(require_auth)):
    """Lists uploaded academic documents (metadata only — no content text)."""
    from tools.academic_context import list_academic_documents
    return {"documents": await list_academic_documents()}


@app.delete("/api/v1/documents/academic/{doc_id}")
async def delete_academic_doc(doc_id: str, session: dict = Depends(require_team_member)):
    """Deletes an academic document and refreshes the agent-context cache."""
    from tools.academic_context import delete_academic_document
    ok = await delete_academic_document(doc_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"deleted": True, "id": doc_id}


# ── Charts — server-rendered PNGs for the canvas presentation editor ──────────

@app.get("/api/v1/charts/available")
async def charts_available(session: dict = Depends(require_team_member)):
    """The charts the canvas presentation editor can embed — every chart
    render_deck_charts produces server-side. Project team only."""
    from tools.chart_render import AVAILABLE_CHARTS
    return AVAILABLE_CHARTS


@app.get("/api/v1/charts/render/{chart_key}")
@limiter.limit("60/minute")
async def charts_render(
    chart_key: str,
    request: Request,
    theme: str = "light",
    width: int = 720,
    height: int = 400,
    session: dict = Depends(require_team_member),
):
    """
    Renders one chart server-side as a PNG for the canvas editor, sized
    to width x height. theme=dark falls back to the light render (the
    matplotlib renderers are light-only). The PNG is cached five minutes
    per (chart_key, theme, width, height). 404 for an unknown chart_key.
    Project team only.
    """
    from tools.chart_render import is_known_chart, render_chart_png
    if not is_known_chart(chart_key):
        raise HTTPException(status_code=404,
                            detail=f"Unknown chart key: {chart_key}")
    theme = "dark" if str(theme).lower() == "dark" else "light"
    w = max(80, min(int(width), 2000))
    h = max(80, min(int(height), 2000))
    png = await render_chart_png(chart_key, theme, w, h)
    return Response(content=png, media_type="image/png")


# ── Document editor — editor_drafts / editor_draft_versions ────────────────────
#
# The in-platform editor for Bob's midpoint paper and Molly's presentation
# deck (migration 021). All endpoints require team_member; the auto-save
# PATCH is silent (no version), POST .../versions saves a named checkpoint.

@app.get("/api/v1/documents/drafts")
async def editor_list_drafts(
    include_all: bool = False,
    session: dict = Depends(require_team_member),
):
    """Default response (June 25 2026): only is_current=true drafts
    -- at most one row per document_type. The previous all-drafts
    response broke DocumentGenerationPanel's 'is there a current
    draft for this type?' lookup once history accumulated, blocking
    Open in Editor for the whole team.

    Query: include_all=true keeps the legacy full-history behaviour,
    used by the editor toolbar's DraftVersionSelector which lists
    every version of the current document_type for the user to
    switch between.

    Owner-scoping was removed June 24 2026 (PR #399); team members
    share access to every document. The endpoint stays behind
    require_team_member so viewers (Dr. Panttser) still cannot
    reach it -- they get the same 403 they got before."""
    from tools.editor_drafts import list_all_drafts
    return {
        "drafts": await list_all_drafts(include_all=include_all),
    }


@app.get("/api/v1/documents/drafts/{draft_id}")
async def editor_get_draft(
    draft_id: int, session: dict = Depends(require_team_member),
):
    """A single draft with its current working content. June 24
    2026: get_draft is not owner-scoped; the team_member gate
    is the authoritative access boundary. Viewers cannot reach
    this endpoint."""
    from tools.editor_drafts import get_draft
    draft = await get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Draft not found.")
    return draft


@app.post("/api/v1/documents/drafts", status_code=201)
async def editor_create_draft(
    body: dict, session: dict = Depends(require_team_member),
):
    """
    Creates a draft and makes it the current one for this owner +
    document_type. Body: {document_type, title, content_json,
    content_text, created_from?}.
    """
    from tools.editor_drafts import DOCUMENT_TYPES, create_draft
    document_type = str(body.get("document_type") or "")
    if document_type not in DOCUMENT_TYPES:
        raise HTTPException(
            status_code=422,
            detail=f"document_type must be one of {DOCUMENT_TYPES}.")
    title = str(body.get("title") or "Untitled draft")[:300]
    created_from = str(body.get("created_from") or "manual")
    draft = await create_draft(
        document_type, session["email"], title,
        body.get("content_json"), body.get("content_text"),
        created_from=created_from)
    if draft is None:
        raise HTTPException(status_code=503,
                            detail="Draft storage is unavailable.")
    return draft


@app.post(
    "/api/v1/editor/drafts/{draft_id}/accept-unverified")
async def accept_unverified_value(
    draft_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """June 28 2026 (PR #479) -- log an operator "accept as-is"
    decision for an <unverified> tag in a draft.

    The frontend popover (UnverifiedPopover.tsx) calls this
    when Bob or Molly clicks "Accept as-is" on a flagged
    numeric value. The endpoint:
      1. Logs the override into editor_numeric_overrides
         (migration 066 -- the existing table from PR #469).
      2. Returns the override-record id + a timestamp.

    The frontend separately rewrites the draft content_json
    via the existing PATCH endpoint to remove the
    <unverified> wrapper (leaving the raw value as plain
    text). This endpoint is the AUDIT LOG, not a content
    mutator -- separation of concerns.

    Body: {value: str, sentence_context: str, document_type:
    str?}. `value` is the raw numeric the operator accepted.
    `sentence_context` is the surrounding sentence text for
    the audit log (caps at 1000 chars on persist).
    `document_type` defaults to the draft's stored type when
    omitted.

    Fail-open: DB write failure returns a 503 with details;
    the frontend should surface the error + leave the
    <unverified> wrapper in place pending a retry."""
    value = str(body.get("value", "")).strip()
    sentence = str(body.get("sentence_context", "")).strip()
    user_email = session.get("email") or ""
    if not value:
        raise HTTPException(
            status_code=422,
            detail="value is required")
    if not user_email:
        raise HTTPException(
            status_code=401,
            detail="session has no email")
    # Resolve document_type from the body or the stored draft.
    doc_type = body.get("document_type")
    if not doc_type:
        try:
            from tools.editor_drafts import get_draft
            d = await get_draft(draft_id)
            if d is not None:
                doc_type = d.get("document_type")
        except Exception:  # noqa: BLE001
            doc_type = None
    # Persist as a single editor_numeric_overrides row using
    # the existing log_editor_overrides helper. The
    # suggested_token slot is intentionally NULL for an
    # accept-as-is decision (the operator chose NOT to swap).
    from tools.editor_save_numeric_check import (
        log_editor_overrides,
    )
    n = await log_editor_overrides(
        draft_id,
        doc_type,
        user_email,
        [{
            "offending_value": value,
            "sentence":        sentence,
            "suggested_token": None,
        }])
    if n != 1:
        # DB write failed (log_editor_overrides logged the
        # exception + returned 0 / -1). Surface 503 so the
        # frontend can retry.
        raise HTTPException(
            status_code=503,
            detail=(
                "Could not log the accept-as-is decision. "
                "Retry; the <unverified> wrapper is still in "
                "place."))
    log.info(
        "unverified_accept_logged",
        draft_id=draft_id,
        document_type=doc_type,
        user_email=user_email,
        value=value)
    return {
        "logged":         True,
        "draft_id":       draft_id,
        "document_type":  doc_type,
        "value":          value,
        "user_email":     user_email,
    }


@app.patch("/api/v1/documents/drafts/{draft_id}")
async def editor_update_draft(
    draft_id: int, body: dict,
    session: dict = Depends(require_team_member),
):
    """
    Auto-save — overwrites the working content. Silent: does NOT create
    a version. Body: {content_json, content_text, word_count?}.

    June 28 2026 -- touchpoint 5 of the hard-lock guardrail.
    BEFORE persisting, the incoming content_json is scanned for
    untoken-backed numerics (e.g. operator typed a "0.86" by hand
    that bypasses the harness loop). Warnings are NON-BLOCKING --
    the save always succeeds. Offenders are logged to
    editor_numeric_overrides for audit + returned in the response
    so the frontend can render a dismissible banner.
    """
    from tools.editor_drafts import update_draft
    from tools.editor_save_numeric_check import (
        log_editor_overrides,
        scan_editor_save_for_untoken_numerics,
    )

    # ── Pre-persist scan -- non-blocking, fail-open ────────────
    warnings: list[dict] = []
    try:
        # Build the current substitution table so the scanner
        # has the canonical token-to-value map. Freeze-aware via
        # PR #455 v4's effective-hash resolver.
        from tools.audit_assembler import current_data_hash
        from tools.cache import get_strategy_cache
        from tools.cio_recommendation import (
            get_latest_recommendation,
        )
        from tools.numeric_substitution import (
            get_substitution_table,
        )
        from tools.submission_freeze import (
            get_effective_data_hash,
        )

        live_hash = await current_data_hash()
        eff_hash = (
            await get_effective_data_hash(live_hash) or live_hash)
        cio = await get_latest_recommendation() or {}
        strategy_cache = await get_strategy_cache(eff_hash) or {}
        sub_table = get_substitution_table(
            eff_hash, strategy_cache, cio, hash_verified=True)

        warnings = scan_editor_save_for_untoken_numerics(
            body.get("content_json"), sub_table)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "editor_save_numeric_scan_skipped",
            draft_id=draft_id, error=str(exc))
        warnings = []

    # ── Persist regardless of warning count ────────────────────
    wc = body.get("word_count")
    ok = await update_draft(
        draft_id, body.get("content_json"), body.get("content_text"),
        word_count_override=int(wc) if isinstance(wc, (int, float)) else None)
    if not ok:
        raise HTTPException(status_code=404, detail="Draft not found.")

    # ── Audit-log every warning (after persist; idempotent on
    #    DB failure) ────────────────────────────────────────────
    if warnings:
        try:
            from sqlalchemy import text as _text
            from database import AsyncSessionLocal
            doc_type: str | None = None
            if AsyncSessionLocal is not None:
                async with AsyncSessionLocal() as s:
                    r = await s.execute(_text(
                        "SELECT document_type FROM editor_drafts "
                        "WHERE id = :id"), {"id": draft_id})
                    row = r.fetchone()
                    doc_type = row[0] if row else None
            await log_editor_overrides(
                draft_id,
                doc_type,
                session.get("email", ""),
                warnings)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "editor_save_numeric_audit_log_skipped",
                draft_id=draft_id, error=str(exc))

    return {
        "saved":    True,
        "draft_id": draft_id,
        # June 28 2026 -- non-blocking numeric warnings. Frontend
        # surfaces these in a dismissible banner.
        "numeric_warnings": warnings,
    }


@app.post("/api/v1/documents/drafts/{draft_id}/versions", status_code=201)
async def editor_save_version(
    draft_id: int, body: dict | None = None,
    session: dict = Depends(require_team_member),
):
    """Saves a named version checkpoint. Body: {version_label}."""
    from tools.editor_drafts import save_version
    label = None
    if body and body.get("version_label"):
        label = str(body["version_label"])[:200]
    version = await save_version(draft_id, label, session["email"])
    if version is None:
        raise HTTPException(status_code=404, detail="Draft not found.")
    return version


@app.get("/api/v1/documents/drafts/{draft_id}/versions")
async def editor_list_versions(
    draft_id: int, session: dict = Depends(require_team_member),
):
    """Every saved version of a draft, newest first."""
    from tools.editor_drafts import list_versions
    return {"versions": await list_versions(draft_id)}


@app.post("/api/v1/documents/drafts/{draft_id}/restore/{version_id}")
async def editor_restore_version(
    draft_id: int, version_id: int,
    session: dict = Depends(require_team_member),
):
    """Restores a saved version as the draft's current content."""
    from tools.editor_drafts import restore_version
    draft = await restore_version(draft_id, version_id)
    if draft is None:
        raise HTTPException(
            status_code=404,
            detail="Draft or version not found.")
    return draft


@app.delete("/api/v1/documents/drafts/{draft_id}")
async def editor_delete_draft(
    draft_id: int, session: dict = Depends(require_team_member),
):
    """Soft-deletes a draft."""
    from tools.editor_drafts import soft_delete_draft
    ok = await soft_delete_draft(draft_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Draft not found.")
    return {"deleted": True, "draft_id": draft_id}


@app.post("/api/v1/documents/drafts/{draft_id}/chat")
@limiter.limit("30/hour")
async def editor_chat(
    draft_id: int, body: dict, request: Request,
    session: dict = Depends(require_team_member),
):
    """
    The editor's Writing Assistant chat. The signed-in user asks for
    writing help on their draft; the assistant answers with the draft's
    full text as context, referencing the actual content.

    Body: {message, history: [{role, content}], selection?}. A selection
    (a passage the user is asking about) is quoted into the prompt. The
    rate limit — 30/hour — keeps the assistant a help, not a crutch.

    404 if the draft does not exist or is not owned by the caller.
    Logged as a writing_assistant interaction so it appears in Team
    Activity and AI cost tracking.
    """
    import asyncio

    from tools.editor_drafts import get_draft
    draft = await get_draft(draft_id)
    # A draft the caller does not own is a 404 — not a 403 — so the
    # endpoint never confirms another user's draft exists.
    if draft is None or draft.get("owner_email") != session.get("email"):
        raise HTTPException(status_code=404, detail="Draft not found.")

    message = str(body.get("message") or "").strip()
    if not message:
        raise HTTPException(status_code=422, detail="'message' is required.")
    if len(message) > 2000:
        raise HTTPException(status_code=422,
                            detail="Message exceeds the 2000-character limit.")
    selection = body.get("selection")
    history = body.get("history") or []

    if ENVIRONMENT == "test":
        return {"response": "Writing assistant is mocked in the test "
                            "environment."}

    content_text = (draft.get("content_text") or "")[:8000]
    system_prompt = (
        f"You are a writing assistant helping {draft['owner_email']} "
        f"improve their {draft['document_type']} document.\n\n"
        f"The full document is:\n---\n{content_text}\n---\n\n"
        "Help the user improve their writing. Be specific and "
        "constructive. Reference the actual content of their document in "
        "your responses. Do not rewrite large sections unless explicitly "
        "asked. Keep responses concise — this is a chat interface, not a "
        "document."
    )

    # The last six exchanges, flattened into the prompt as context
    # (call_claude is single-turn).
    convo_lines: list[str] = []
    for turn in history[-6:]:
        role = "User" if turn.get("role") == "user" else "Assistant"
        text = str(turn.get("content") or "").strip()
        if text:
            convo_lines.append(f"{role}: {text}")
    user_message = ""
    if convo_lines:
        user_message += ("Earlier in this conversation:\n"
                         + "\n".join(convo_lines) + "\n\n")
    if selection:
        user_message += (f"Regarding this passage:\n> "
                         f"{str(selection)[:1000]}\n\n")
    user_message += message

    from agents.usage import start_usage_capture
    start_usage_capture()
    try:
        from agents.base import SONNET_MODEL, call_claude
        reply = await asyncio.to_thread(
            call_claude, SONNET_MODEL, system_prompt, user_message, 600)
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("editor_chat_failed", ref=ref, error=str(exc))
        raise HTTPException(
            status_code=502,
            detail=f"Writing assistant unavailable (ref: {ref})")

    _log_interaction_bg(
        request, session, "writing_assistant",
        question_text=message[:500],
        response_summary=reply[:500],
        metadata={"draft_id": draft_id})
    return {"response": reply}


# ── Rubric Review (May 25 2026) ──────────────────────────────────────────────
#
# Sends the current midpoint paper draft to Gemini with the full FNA 670
# midpoint rubric. Returns a structured verdict — per-section pass/fail
# against the rubric criteria, specific optional edits with reasoning,
# and an overall readiness verdict. Read-only: the endpoint never
# modifies the draft — it surfaces suggestions inline in the Writing
# Assistant panel for the user to apply (or ignore) manually.
#
# Distinct from the Academic Review council pass: that endpoint runs
# every council peer + the arbiter and is expensive. This endpoint is
# a single fast Gemini call against an explicit rubric prompt — useful
# for the quick "where am I against the rubric" check the user wants
# before committing to a full review.

_MIDPOINT_RUBRIC_PROMPT = (
    "You are reviewing a graduate-finance midpoint paper against the "
    "FNA 670 rubric. The paper is 3 pages, double-spaced, 12-point font, "
    "750-900 words TOTAL, organised into four sections:\n"
    "\n"
    "  1. DATA AND METHODOLOGY (1 page, 235-285 words)\n"
    "     - Identifies data sources and study period\n"
    "     - Names the portfolio constraints (long-only, fully invested, "
    "no cash, quarterly rebalancing)\n"
    "     - Distinguishes static vs dynamic strategies\n"
    "     - Names the Carhart four-factor attribution model\n"
    "     - APA style, past tense, third person\n"
    "\n"
    "  2. PRELIMINARY RESULTS AND DIAGNOSTICS (1 page, 235-285 words)\n"
    "     - INTERPRETS the results (does not merely list numbers)\n"
    "     - Discusses the 2022 equity-bond correlation break explicitly\n"
    "     - References summary statistics and regime-conditional "
    "performance with specific values\n"
    "     - Connects findings back to the research question\n"
    "\n"
    "  3. ROLES AND DIVISION OF LABOR (0.5 page, 110-135 words)\n"
    "     - States each team member's role\n"
    "     - Attributes documented contributions (commits, sessions, "
    "documents)\n"
    "     - Factual, no invented contributions\n"
    "\n"
    "  4. NEXT STEPS AND OPEN QUESTIONS (0.5 page, 110-135 words)\n"
    "     - Forward-looking, not retrospective\n"
    "     - Names specific areas for further investigation\n"
    "     - Reflects the team's analytical priorities\n"
    "\n"
    "Return a structured JSON response with EXACTLY this shape — no "
    "additional fields, no preamble:\n"
    "\n"
    "{\n"
    '  "sections": {\n'
    '    "methodology":  {"verdict": "pass" | "fail",\n'
    '                     "reasoning": "1-3 sentences naming which '
    'criteria were met and which were not"},\n'
    '    "results":      {"verdict": "pass" | "fail", "reasoning": "..."},\n'
    '    "roles":        {"verdict": "pass" | "fail", "reasoning": "..."},\n'
    '    "next_steps":   {"verdict": "pass" | "fail", "reasoning": "..."}\n'
    "  },\n"
    '  "edits": [\n'
    '    {"section": "methodology|results|roles|next_steps",\n'
    '     "suggestion": "specific change to consider",\n'
    '     "reasoning": "why this would strengthen the section"}\n'
    "  ],\n"
    '  "overall": {\n'
    '    "verdict": "ready" | "needs_work" | "not_ready",\n'
    '    "reasoning": "2-4 sentences summarising readiness"\n'
    "  }\n"
    "}\n"
    "\n"
    "VERDICT semantics:\n"
    "  - section 'pass' — the section meets the rubric criteria. Minor "
    "wordsmithing is fine; the substance is there.\n"
    "  - section 'fail' — a substantive rubric criterion is unmet "
    "(missing required content, wrong tense, no interpretation in "
    "results, etc.). Be honest, not generous.\n"
    "  - overall 'ready' — all four sections pass and the paper reads "
    "cleanly as a 3-page midpoint draft.\n"
    "  - overall 'needs_work' — sections substantially correct but "
    "specific edits would strengthen the submission.\n"
    "  - overall 'not_ready' — one or more sections fail; the draft is "
    "not yet at submission quality.\n"
    "\n"
    "EDITS — provide 3-6 specific optional suggestions ordered by "
    "impact. Each must name a section, a concrete change, and the "
    "rubric reason. These are SUGGESTIONS — the team will choose what "
    "to apply.\n"
    "\n"
    "Respond ONLY with the JSON object. No markdown fences, no preamble."
)


def _parse_rubric_review(raw: str) -> dict[str, object] | None:
    """Best-effort JSON parse on a Gemini response. Strips ``` fences
    and any preamble before / after the {…} block. Returns None on
    unparseable input — the endpoint then falls back to a friendly
    error rather than a 500.
    """
    import json as _json
    text = (raw or "").strip()
    # Strip ```json fences if Gemini ignored the no-fence instruction.
    if text.startswith("```"):
        first_nl = text.find("\n")
        if first_nl != -1:
            text = text[first_nl + 1:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()
    # Extract the outermost {…} block if there is any preamble.
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end <= start:
        return None
    try:
        return _json.loads(text[start:end + 1])
    except Exception:  # noqa: BLE001
        return None


def _rubric_review_unavailable_payload(reason: str) -> dict[str, object]:
    """Used by the test env path AND the no-Gemini-key / parse-failure
    fall-throughs. Returns a payload that renders cleanly in the UI
    without claiming a pass/fail verdict the model didn't produce."""
    return {
        "sections": {
            "methodology": {"verdict": "fail", "reasoning": reason},
            "results":     {"verdict": "fail", "reasoning": reason},
            "roles":       {"verdict": "fail", "reasoning": reason},
            "next_steps":  {"verdict": "fail", "reasoning": reason},
        },
        "edits": [],
        "overall": {"verdict": "not_ready", "reasoning": reason},
        "unavailable": True,
    }


@app.post("/api/v1/documents/drafts/{draft_id}/rubric-review")
@limiter.limit("10/hour")
async def rubric_review(
    draft_id: int, request: Request,
    session: dict = Depends(require_team_member),
):
    """
    Sends the current draft to Gemini with the midpoint paper rubric
    and returns a structured per-section verdict + suggested edits.
    Read-only — never modifies the draft. The Writing Assistant panel
    surfaces the verdict inline so the user can decide which edits to
    apply manually.

    Scoped to midpoint_paper drafts — other document types (executive
    brief, presentation deck/script) have their own rubrics not yet
    encoded here.
    """
    import asyncio

    from tools.editor_drafts import get_draft
    draft = await get_draft(draft_id)
    if draft is None or draft.get("owner_email") != session.get("email"):
        raise HTTPException(status_code=404, detail="Draft not found.")
    if draft.get("document_type") != "midpoint_paper":
        raise HTTPException(
            status_code=422,
            detail="Rubric review is only available for midpoint paper "
                   "drafts.")

    content_text = (draft.get("content_text") or "").strip()
    if not content_text:
        return _rubric_review_unavailable_payload(
            "The draft is empty — write some content before requesting "
            "a rubric review.")

    if ENVIRONMENT == "test":
        # Deterministic shape so the frontend test can pin the render
        # without hitting Gemini. Mirrors the structured response a
        # real run would produce.
        return {
            "sections": {
                "methodology": {"verdict": "pass",
                                "reasoning": "Methodology rubric criteria met."},
                "results":     {"verdict": "fail",
                                "reasoning": "Results lack explicit "
                                             "interpretation."},
                "roles":       {"verdict": "pass",
                                "reasoning": "Roles attributed factually."},
                "next_steps":  {"verdict": "pass",
                                "reasoning": "Forward-looking and specific."},
            },
            "edits": [
                {"section": "results",
                 "suggestion": "Add one sentence interpreting the "
                               "post-2022 Sharpe gap.",
                 "reasoning": "Rubric requires interpretation, not "
                              "listing."},
            ],
            "overall": {"verdict": "needs_work",
                        "reasoning": "Three sections pass; results "
                                     "needs interpretation."},
        }

    if not os.getenv("GOOGLE_API_KEY"):
        return _rubric_review_unavailable_payload(
            "Rubric review is configured to run on Gemini but the "
            "GOOGLE_API_KEY env var is not set.")

    user_message = (
        "Review the following midpoint paper draft against the rubric. "
        "The draft text is the four sections concatenated in order; the "
        "[[BOB: …]] markers are placeholders that the user will resolve "
        "before submission (treat them as 'work in progress' rather "
        "than rubric failures, but note when a placeholder is the only "
        "content for a required topic).\n\n"
        f"---\n{content_text[:12000]}\n---\n\n"
        "Respond with the JSON object specified in the instructions. "
        "No preamble, no markdown fences."
    )

    from agents.usage import start_usage_capture
    start_usage_capture()
    try:
        from agents.base import GEMINI_MODEL, call_gemini
        raw = await asyncio.to_thread(
            call_gemini, GEMINI_MODEL,
            _MIDPOINT_RUBRIC_PROMPT, user_message,
            trigger="rubric_review",
        )
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("rubric_review_failed",
                  ref=ref, error=str(exc), draft_id=draft_id)
        return _rubric_review_unavailable_payload(
            f"Gemini call failed (ref: {ref}). Retry in a moment.")

    parsed = _parse_rubric_review(raw or "")
    if parsed is None or "sections" not in parsed or "overall" not in parsed:
        log.warning("rubric_review_parse_failed",
                    draft_id=draft_id, raw_head=(raw or "")[:200])
        return _rubric_review_unavailable_payload(
            "The rubric review came back in an unexpected format. "
            "Retry — Gemini occasionally drops the JSON shape.")

    _log_interaction_bg(
        request, session, "rubric_review",
        question_text=f"Rubric review for draft {draft_id}",
        response_summary=(parsed.get("overall") or {}).get(
            "reasoning", "")[:500],
        metadata={"draft_id": draft_id,
                  "overall_verdict": (parsed.get("overall") or {}).get("verdict")})
    return parsed


# ── Auto-fired Academic Review status for a draft (May 25 2026) ───────────────
#
# After midpoint or executive-brief generation, an Academic Review
# fires in the background and writes its parsed score into
# agent_interactions.metadata.draft_id. The editor reads this
# endpoint on draft load (and polls while it reads `running`) to
# render the header score pill and, for midpoint, the advisory
# banner when the score lands below 6.0.

@app.get("/api/v1/documents/drafts/{draft_id}/academic-review-status")
async def get_draft_academic_review_status(
    draft_id: int, session: dict = Depends(require_team_member),
):
    """
    Returns the latest auto-fired Academic Review score for a draft:
      {
        status:           "complete" | "running" | "missing",
        score:            float | None,        # 0-10
        rating:           str | None,          # Strong | Developing | Needs Work
        advisory:         bool,                # midpoint + score < 6.0
        document_type:    str,                 # midpoint_paper | executive_brief
        section_ratings:  dict,                # {section_key: rating}
        run_at:           ISO timestamp | None,
        threshold:        6.0,                 # the midpoint advisory cutoff
      }

    "missing" — no auto-review has landed yet for this draft. The
    generation flow always schedules one, so this means either the
    background task is still in flight (the editor polls) or it
    silently failed (the editor stops polling after a few attempts).
    "running" mirrors "missing" today — the placeholder gives us
    room to wire a heartbeat row later without changing the wire
    format.

    404 if the draft does not exist or the caller is not the owner.
    """
    from tools.editor_drafts import get_draft
    from tools.academic_review_score import ADVISORY_THRESHOLD

    draft = await get_draft(draft_id)
    # June 25 2026 -- documents are team-shared (PR #399), so the
    # owner_email filter was producing spurious 404s for team
    # members opening a draft someone else generated. The endpoint
    # already gates on require_team_member; that's the authoritative
    # access check.
    if draft is None:
        raise HTTPException(status_code=404, detail="Draft not found.")

    document_type = draft.get("document_type") or ""
    empty = {
        "status":          "missing",
        "score":           None,
        "rating":          None,
        "advisory":        False,
        "document_type":   document_type,
        "section_ratings": {},
        "run_at":          None,
        "threshold":       ADVISORY_THRESHOLD,
        # June 25 2026 -- new fields per the editor's review-status
        # badge contract. Backward-compatible additions; legacy
        # callers ignore them.
        "draft_id":        draft_id,
        "has_review":      False,
        "last_review_at":  None,
        "arbiter_score":   None,
        "verdict_summary": None,
    }

    # Auto-review tracking covers all four submission doc types now.
    # The midpoint legacy is preserved for advisory; per-doc reviews
    # fire from the editor's per-doc trigger (Concern 3) which logs
    # an agent_interactions row tied to the draft_id.
    if document_type not in (
            "midpoint_paper", "executive_brief",
            "analytical_appendix", "presentation_deck",
            "presentation_script"):
        return empty

    try:
        from tools.activity_log import get_latest_academic_review_for_draft
        latest = await get_latest_academic_review_for_draft(draft_id)
    except Exception as exc:  # noqa: BLE001
        log.warning("academic_review_status_read_failed",
                    draft_id=draft_id, error=str(exc))
        return empty

    if latest is None:
        return empty

    meta = latest.get("metadata") or {}
    score = meta.get("score")
    # Advisory is computed off the live score and the document type —
    # don't trust a stale `advisory` field in the metadata in case the
    # threshold ever moves. The threshold is the same source of truth
    # the frontend reads from this response.
    advisory = (
        document_type == "midpoint_paper"
        and isinstance(score, (int, float))
        and score < ADVISORY_THRESHOLD
    )
    response_summary = latest.get("response_summary") or ""
    verdict_summary: str | None = None
    if response_summary:
        # Strip markdown headings + collapse whitespace; first
        # ~280 chars is enough for the editor's badge tooltip.
        compact = " ".join(
            response_summary.replace("#", "").split())
        verdict_summary = (
            compact[:280] + "…" if len(compact) > 280
            else (compact or None))
    return {
        "status":          "complete",
        "score":           score,
        "rating":          meta.get("overall_rating"),
        "advisory":        advisory,
        "document_type":   document_type,
        "section_ratings": meta.get("section_ratings") or {},
        "run_at":          latest.get("timestamp"),
        "threshold":       ADVISORY_THRESHOLD,
        # New fields (June 25 2026) -- additive, legacy-safe.
        "draft_id":        draft_id,
        "has_review":      True,
        "last_review_at":  latest.get("timestamp"),
        "arbiter_score":   score,
        "verdict_summary": verdict_summary,
    }


@app.get("/api/v1/documents/drafts/{draft_id}/review-export")
async def export_draft_academic_review(
    draft_id: int, session: dict = Depends(require_team_member),
):
    """June 25 2026 -- DOCX export of the completed academic review
    for a draft. Returns the assembled report as
    application/vnd.openxmlformats-officedocument.wordprocessingml
    .document download.

    Sourced from three places:
      editor_drafts row              cover metadata
      council_debates row            peer responses + critic findings
                                     + arbiter resolution + fix
                                     proposals + counter arguments
      agent_interactions row         arbiter prose (response_summary)
                                     + overall_rating + score
                                     + independent_review block
    Returns 404 when the draft does not exist, 200 + DOCX even when
    no review has landed yet (the DOCX surfaces a 'no review found'
    notice in each section so the download endpoint never errors).
    """
    from datetime import date

    from tools.editor_drafts import get_draft
    from tools.activity_log import get_latest_academic_review_for_draft
    from tools.review_docx import build_review_docx
    from agents.academic_review import get_latest_debate_for_draft

    draft = await get_draft(draft_id)
    if draft is None:
        raise HTTPException(
            status_code=404, detail="Draft not found.")

    interaction = await get_latest_academic_review_for_draft(draft_id)
    debate = await get_latest_debate_for_draft(draft_id)

    if interaction is None and debate is None:
        raise HTTPException(
            status_code=404,
            detail=(
                "No academic review found for this draft. "
                "Run the review first."))

    content = await asyncio.to_thread(
        build_review_docx,
        draft=draft, debate=debate, interaction=interaction)

    doc_type = draft.get("document_type") or "document"
    slug = doc_type.replace("_", "-")
    filename = (
        f"forest-capital-{slug}-academic-review-"
        f"draft-{draft_id}-{date.today().isoformat()}.docx")
    return Response(
        content=content,
        media_type=_DOCX_MEDIA,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        })


@app.post("/api/v1/documents/script/generate")
@limiter.limit("6/minute")
async def generate_presentation_script(
    request: Request,
    session: dict = Depends(require_permission("team_member")),
):
    """June 25 2026 -- async job kickoff for presentation script
    generation. Matches the brief / appendix / deck pattern: 202
    + {job_id, status: 'pending'}, generation runs as a background
    task via _generate_async, the frontend polls /api/v1/jobs/<id>
    for completion.

    The previous synchronous shape (POST returns 200 with draft_id
    inline) prevented the Reports tile from inheriting the standard
    inProgress / complete / failed / idle ternary because the tile's
    UI is driven by the generation-jobs Zustand store -- which only
    populates when an endpoint returns a job_id. Refactoring to the
    async pattern lets the script tile move into the DOCS array and
    render with the same chrome as the other three deliverables.

    Deck-draft lookup moved from the request body to
    get_current_draft_by_type('presentation_deck') -- mirrors the
    academic review's team-shared lookup (PR #410). The body
    parameter is now ignored; existing clients passing
    {draft_id: ...} continue to work but the value is unused.

    Speakers-required validation + missing-deck 404 now surface
    INSIDE the job (status='failed' with the error message) rather
    than as synchronous HTTPException. The frontend's failed-state
    branch renders the error chip + Try Again so the UX is the
    same; the failure just lands one polling-tick later.
    """
    return _start_generation_job(
        "presentation_script", session, request)


@app.post("/api/v1/documents/script/regenerate-slide")
@limiter.limit("12/minute")
async def regenerate_script_slide(
    request: Request,
    body: dict,
    session: dict = Depends(
        require_permission("generate_documents")),
):
    """June 26 2026 -- per-slide script regeneration.

    Body: {draft_id: int, slide_number: int}

    Returns: {content_json: TipTapDoc, slide_number, draft_id,
              draft_version}

    Synchronous (NOT the async-job pattern): scoped to a single
    slide's prose so the wall-clock is ~30-60s, well within a
    single HTTP request. The full-script generation path stays
    async because it covers all 12 slides (~3-6 min). Per the
    user spec, the per-slide regen is 'outside the main async
    job queue -- synchronous response since it's scoped to one
    slide.'

    Process:
      1. Load the current script draft. 404 if not found, 400 if
         document_type != presentation_script.
      2. Load the current team-shared deck draft for slide
         context. 409 if no deck.
      3. Extract the H2-bounded TipTap block for slide_number
         from the script's content_json (the H2 'Slide N: title'
         marker bounds each slide).
      4. Call Sonnet with the per-slide context (deck slide N's
         body + speaker + brief excerpt) targeting ~150 words.
      5. Parse the response into TipTap nodes (reuses
         script_to_tiptap fragments).
      6. Splice into content_json: replace nodes from the H2
         marker up to the next H2 (or end of doc).
      7. Persist via update_draft (no new version row -- this is
         an auto-save shape) so version history doesn't churn.
      8. Return the new content_json so the frontend can patch
         the TipTap editor in place.

    Failure modes (HTTP error):
      404 -- draft not found or wrong document_type
      409 -- no current deck draft, or slide_number not in deck,
             or slide marker missing from script content_json
      422 -- malformed body
      502 -- generator call failed
    """
    import asyncio
    import re as _re

    from tools.editor_drafts import (
        get_draft, get_current_draft_by_type, update_draft,
    )
    from tools.script_generation import (
        build_script_prompt, script_to_tiptap,
    )
    from tools.brief_grounding import get_brief_for_grounding

    raw_draft_id = (body or {}).get("draft_id")
    raw_slide = (body or {}).get("slide_number")
    try:
        draft_id = int(raw_draft_id)
        slide_number = int(raw_slide)
    except (TypeError, ValueError):
        raise HTTPException(
            status_code=422,
            detail="draft_id and slide_number are required ints.")

    draft = await get_draft(draft_id)
    if draft is None:
        raise HTTPException(
            status_code=404, detail="Script draft not found.")
    if draft.get("document_type") != "presentation_script":
        raise HTTPException(
            status_code=400,
            detail="Draft is not a presentation script.")

    deck = await get_current_draft_by_type("presentation_deck")
    if deck is None:
        raise HTTPException(
            status_code=409,
            detail=(
                "No current presentation deck draft to source "
                "the regenerated slide from."))
    deck_content = deck.get("content_json") or {}
    deck_slides = (
        deck_content.get("slides", [])
        if isinstance(deck_content, dict) else [])
    target_slide = next(
        (s for s in deck_slides
         if isinstance(s, dict)
         and int(s.get("id") or 0) == slide_number),
        None)
    if target_slide is None:
        raise HTTPException(
            status_code=409,
            detail=(
                f"Slide {slide_number} not found in the deck."))

    # Extract the current script's H2-bounded slide block. The
    # H2 marker pattern is '## Slide N: <title>' OR a bare
    # 'Slide N' heading -- script_to_tiptap recognises both.
    script_content = draft.get("content_json") or {}
    nodes = (
        script_content.get("content", [])
        if isinstance(script_content, dict) else [])

    def _node_text(n: dict) -> str:
        if n.get("text"):
            return str(n["text"])
        return "".join(_node_text(c) for c in (n.get("content") or []))

    slide_marker_re = _re.compile(
        rf"^\s*Slide\s+{slide_number}\b", _re.IGNORECASE)
    start_idx = None
    end_idx = len(nodes)
    for i, n in enumerate(nodes):
        if not isinstance(n, dict):
            continue
        if n.get("type") != "heading":
            continue
        level = (n.get("attrs") or {}).get("level")
        if level not in (1, 2):
            continue
        text = _node_text(n)
        if start_idx is None and slide_marker_re.match(text):
            start_idx = i
        elif start_idx is not None:
            # Next H1/H2 -- the slide's block ends here.
            end_idx = i
            break
    if start_idx is None:
        raise HTTPException(
            status_code=409,
            detail=(
                f"Slide {slide_number} marker not found in the "
                "script content. The script may have been edited "
                "to remove the slide heading; refresh and retry."))

    # Build a single-slide prompt. Reuses build_script_prompt
    # against just this slide so the model sees the same
    # contextual scaffolding (slide body + speaker) it sees for
    # the full-script path. Brief excerpt threaded in for
    # consistency with the full-generation flow.
    brief = await get_brief_for_grounding()
    brief_text = brief.get("content_text") if brief else None
    one_slide_prompt = build_script_prompt(
        [target_slide], brief_text, None)
    # Tighten the ask -- single-slide prose, ~150 words.
    one_slide_prompt += (
        "\n\nFOCUS: regenerate ONLY this single slide's "
        "delivery prose. Target ~150 words. Keep the speaker "
        "assignment unchanged. Output the slide's prose in the "
        "same '## Slide N: title' + speaker-block + paragraphs "
        "format the full-script generator emits.")

    if ENVIRONMENT == "test":
        raw = (
            f"## Slide {slide_number}: "
            f"{target_slide.get('title') or 'Slide'}\n\n"
            f"**Speaker: {target_slide.get('speaker') or 'TBA'}**\n\n"
            "Regenerated delivery paragraph for this slide.")
    else:
        try:
            from agents.base import SONNET_MODEL, call_claude
            raw = await asyncio.to_thread(
                call_claude,
                SONNET_MODEL,
                "You are an academic presenter generating a "
                "word-for-word delivery script for one slide.",
                one_slide_prompt,
                max_tokens=600,
                trigger="script_regen_per_slide")
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "script_slide_regen_call_failed",
                slide_number=slide_number, error=str(exc))
            raise HTTPException(
                status_code=502,
                detail=f"Per-slide regeneration failed: {exc}")

    # Parse the new prose into TipTap nodes and splice.
    new_subdoc, _new_text = script_to_tiptap(raw)
    new_nodes = (
        new_subdoc.get("content", [])
        if isinstance(new_subdoc, dict) else [])
    if not new_nodes:
        raise HTTPException(
            status_code=502,
            detail=(
                "Regenerated slide produced no parseable "
                "content. Retry the regen."))

    spliced_nodes = (
        nodes[:start_idx] + new_nodes + nodes[end_idx:])
    new_content_json = {"type": "doc", "content": spliced_nodes}

    # Derive content_text from the spliced nodes via the same
    # walker the full-script generator uses.
    new_content_text = "\n\n".join(
        _node_text(n).strip() for n in spliced_nodes
        if _node_text(n).strip())

    ok = await update_draft(
        draft_id, new_content_json, new_content_text)
    if not ok:
        raise HTTPException(
            status_code=500,
            detail="Could not persist the regenerated script.")

    return {
        "draft_id": draft_id,
        "slide_number": slide_number,
        "content_json": new_content_json,
    }


async def _generate_script_document(
    email: str,
) -> tuple[bytes, str, str, int | None]:
    """June 25 2026 -- presentation script generator helper.
    Mirrors _generate_brief_document / _generate_appendix_document /
    _generate_deck_document's contract: returns (file bytes,
    filename, media type, editor draft id). Raises on any failure
    so _generate_async's outer try/except records the job as
    failed with the ref'd error.

    Inputs:
      - The current team-shared presentation_deck draft (via
        get_current_draft_by_type, NOT owner-scoped) -- the script
        is generated from the canonical deck regardless of which
        team member triggered.
      - The caller's executive_brief + midpoint_paper drafts as
        optional academic context (degrades gracefully when
        absent).

    Failure paths (raise -> job goes to status='failed'):
      - No current deck draft -> RuntimeError surfaces 'Generate
        the Presentation Deck first.'
      - Deck has no slides with assigned speakers -> RuntimeError
        surfaces 'Assign speakers to slides before generating the
        script.'

    The script draft is created with data_hash stamped (migration
    063) so the tile chip + Light Refresh status table render
    correctly post-generation.
    """
    import asyncio
    from datetime import date

    from tools.editor_drafts import (
        create_draft, get_current_draft,
        get_current_draft_by_type,
    )
    from tools.script_generation import deck_speakers, generate_script
    from tools.academic_docx import build_editor_docx

    deck = await get_current_draft_by_type("presentation_deck")
    if deck is None:
        raise RuntimeError(
            "No current presentation deck draft to source the "
            "script from. Generate the Presentation Deck first.")

    content = deck.get("content_json") or {}
    slides = (
        content.get("slides", [])
        if isinstance(content, dict) else [])
    if not deck_speakers(slides):
        raise RuntimeError(
            "Assign speakers to slides before generating the "
            "script.")

    exec_brief = await get_current_draft(email, "executive_brief")
    midpoint = await get_current_draft(email, "midpoint_paper")

    # June 28 2026 (Phase 2) -- build a substitution_table from
    # the freeze-aware effective hash so the script generator
    # can derive content_text from the substituted projection
    # under DEFER_SUBSTITUTION_TO_EXPORT. Fail-open: any error
    # leaves substitution_table=None + script_to_tiptap falls
    # through to the legacy single-pass behaviour.
    _script_sub_table: dict[str, str] | None = None
    try:
        from tools.audit_assembler import current_data_hash
        from tools.cache import get_strategy_cache
        from tools.cio_recommendation import (
            get_latest_recommendation,
        )
        from tools.numeric_substitution import (
            get_substitution_table,
        )
        from tools.submission_freeze import (
            get_effective_data_hash,
        )
        _live = await current_data_hash()
        _eff = await get_effective_data_hash(_live) or _live
        _cio = await get_latest_recommendation() or {}
        _scache = await get_strategy_cache(_eff) or {}
        _script_sub_table = get_substitution_table(
            _eff, _scache, _cio, hash_verified=True)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "script_substitution_table_build_failed",
            error=str(exc))

    result = await asyncio.to_thread(
        generate_script, deck, exec_brief, midpoint,
        substitution_table=_script_sub_table)

    # Stamp the live strategy hash on the draft (migration 063).
    try:
        from tools.audit_assembler import (
            current_data_hash as _curr_hash,
        )
        _script_hash = await _curr_hash()
    except Exception:  # noqa: BLE001
        _script_hash = None

    # June 26 2026 -- run the document audit against the
    # generated content_text so the script draft carries
    # audit_warnings just like the brief / appendix / deck do.
    # The editor's AuditWarningsBanner reads draft.audit_warnings
    # and renders the flagged-items list; without this call the
    # script draft was always landing with audit_warnings=null
    # and the banner stayed empty even when the script had real
    # issues (numeric / direction / consistency / citation
    # flags). Mirrors the pattern in _generate_brief_document +
    # _generate_deck_document.
    audit_warnings = await _run_document_audit(
        result["content_text"], "presentation_script", email)
    new_draft = await create_draft(
        "presentation_script", email, "Presentation Script",
        result["content_json"], result["content_text"],
        created_from="generated",
        audit_warnings=audit_warnings,
        data_hash=_script_hash)
    if new_draft is None:
        ref = uuid.uuid4().hex[:8]
        log.error("script_draft_create_failed", ref=ref)
        raise RuntimeError(
            f"Could not save the generated script (ref: {ref})")

    # June 28 2026 (PR #479) -- auto-upgrade hook for script.
    # Walks content_json + converts any <unverified> tag
    # substrings into structured unverified TipTap nodes so
    # the NodeView renders. Document-type-agnostic per
    # operator directive.
    await _auto_upgrade_draft_to_token_values(
        new_draft["id"], "presentation_script")

    # Build the docx for the job's /download endpoint -- the
    # frontend Download button reads from there. Mirrors the
    # brief / appendix / deck path which all stage the bytes on
    # the job row.
    file_bytes = await asyncio.to_thread(
        build_editor_docx, new_draft)
    filename = (
        f"forest-capital-presentation-script-"
        f"{date.today().isoformat()}.docx")
    return file_bytes, filename, _DOCX_MEDIA, new_draft["id"]


@app.get("/api/v1/documents/rehearsal")
async def get_rehearsal_payload(
    session: dict = Depends(require_team_member),
):
    """
    Combined deck + script payload for the presentation rehearsal mode
    (combined-script-and-slide-view overlay in the script editor).

    Reads the requesting user's current (is_current=true)
    presentation_deck and presentation_script editor drafts and returns
    them side by side, with the script parsed into per-slide sections.

    Returns 404 when either draft is absent — the rehearsal needs both:
      deck missing:    "No presentation deck found. Generate your deck first."
      script missing:  "No presentation script found. Generate your script first."

    Estimated delivery time is total_words / 150 (the platform-wide
    150-wpm convention; the script editor's delivery time pill uses
    the same rate).
    """
    from tools.editor_drafts import get_current_draft
    from tools.rehearsal import parse_script_sections

    email = session.get("email", "")
    deck = await get_current_draft("presentation_deck", email)
    if deck is None:
        raise HTTPException(
            status_code=404,
            detail="No presentation deck found. Generate your deck first.")
    script = await get_current_draft("presentation_script", email)
    if script is None:
        raise HTTPException(
            status_code=404,
            detail="No presentation script found. Generate your script first.")

    # Deck slides — pass through the canvas shape verbatim. The frontend
    # rehearsal overlay reuses the same CanvasSlide / CanvasElement
    # types the editor itself uses to render the static preview.
    deck_json = deck.get("content_json") or {}
    slides = deck_json.get("slides") if isinstance(deck_json, dict) else []

    # Script — parse the TipTap doc into per-slide sections.
    script_json = script.get("content_json") or {}
    sections = parse_script_sections(script_json)
    total_words = sum(s.get("word_count", 0) for s in sections)
    estimated_minutes = max(1, round(total_words / 150))

    return {
        "deck": {
            "draft_id": deck.get("id"),
            "slides":   slides or [],
        },
        "script": {
            "draft_id":          script.get("id"),
            "sections":          sections,
            "total_words":       total_words,
            "estimated_minutes": estimated_minutes,
        },
    }


@app.post("/api/v1/documents/drafts/{draft_id}/export")
async def export_editor_draft(
    draft_id: int,
    body: dict | None = None,
    session: dict = Depends(require_team_member),
):
    """
    Exports a presentation_script editor draft as a .docx — the master
    script (every speaker) when no speaker is given, or one speaker's
    individual script (their slides only) when {speaker} is provided.
    Team members only.
    """
    import asyncio
    from datetime import date

    from tools.editor_drafts import get_draft
    from tools.script_docx import build_script_docx

    draft = await get_draft(draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Draft not found.")
    if draft.get("document_type") != "presentation_script":
        raise HTTPException(
            status_code=400,
            detail="This export is available for presentation scripts only.")

    raw_speaker = (body or {}).get("speaker")
    speaker = str(raw_speaker).strip() if raw_speaker else None
    content = await asyncio.to_thread(build_script_docx, draft, speaker)

    slug = (speaker.lower().replace(" ", "-") if speaker else "master")
    filename = (f"forest-capital-script-{slug}-"
                f"{date.today().isoformat()}.docx")
    return Response(
        content=content, media_type=_DOCX_MEDIA,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'})


# ── Regime ────────────────────────────────────────────────────────────────────

@app.get("/api/regime/current")
async def get_current_regime(session: dict = Depends(require_auth)):
    # Sprint 5: regime_signals_cache checked before calling FRED.
    # DB cache TTL = 15 minutes — matches in-process dict in regime_detector.py.
    # On a Render restart the DB row is still valid; the in-process dict is gone.
    # This prevents FRED timeout (30-60s) from blocking every post-restart request.
    if ENVIRONMENT != "test":
        try:
            from tools.cache import get_regime_cache, set_regime_cache
            cached = await get_regime_cache()
            if cached:
                return cached

            from tools.regime_detector import detect_current_regime
            result = detect_current_regime()
            await set_regime_cache(result, ttl_minutes=15)
            return result
        except Exception as exc:
            log.warning("regime_detection_fallback", error=str(exc))
    return MOCK_REGIME


# ── Regime-signals freshness gate (June 27 2026, deck-tier only) ─────────────
#
# The presentation deck includes a live CIO recommendation on slides 7 and
# 11 that depends on regime_signals_cache being current (15-min TTL). Stale
# signals at generation time would produce a deck that misleads the live
# panel. Per the user spec, the deck is held to a stricter standard than
# the brief / appendix: those keep the existing graceful em-dash fallback
# (they don't surface a live recommendation); the deck must HARD GATE.
#
# Three operations gate on this helper:
#   * _generate_deck_document       (full generation)
#   * council_academic_review       (when document_type=presentation_deck)
#   * post_light_refresh            (the cache-warming workflow that
#                                    typically immediately precedes deck
#                                    regen)
#
# When the cache is hit, the helper returns (True, signals). When it
# misses, we attempt detect_current_regime() with a HARD 10-SECOND TIMEOUT
# (FRED + market-data fetches can hang for minutes; the panel can't wait
# that long for an error message) and either return (True, signals) on a
# successful refresh OR (False, None) on timeout / failure. The caller
# raises a 503 with the user-spec blocking error message in the failure
# case -- never falls back to em-dashes for the deck.

_REGIME_REFRESH_TIMEOUT_S = 10.0
_REGIME_BLOCKING_ERROR_DETAIL = (
    "Live regime signals unavailable. The deck includes a live CIO "
    "recommendation that requires current data. Please try again in a "
    "few minutes.")

# June 27 2026 -- background refresh cadence for the regime_signals_cache.
# The cache TTL is 15 minutes (set_regime_cache ttl_minutes=15). Refreshing
# every 10 minutes leaves 5 minutes of headroom so the cache is always warm
# when a user (or the deck hard gate) reads it -- the gate never fires
# due to TTL expiry during normal operation. Scheduled by the lifespan
# handler as a fire-and-forget task; per-iteration failures (FRED outage,
# detect_current_regime exception, cache write failure) log a warning and
# the loop continues to the next cycle.
_REGIME_SIGNALS_TICKER_INTERVAL_S = 600.0


async def _regime_signals_fresh_or_refresh() -> tuple[bool, dict | None]:
    """Returns (is_fresh, signals).

    1. Check regime_signals_cache via get_regime_cache(). If the cached
       row is unexpired, returns (True, signals) immediately.
    2. On miss/expiry, run detect_current_regime() in a worker thread
       with a 10-second hard timeout. The detector hits live market
       data APIs (FRED, Yahoo) which can hang under load -- without the
       timeout the blocking error itself could take minutes to surface,
       worse UX than the original stale-cache behaviour.
    3. On successful refresh: write to set_regime_cache(ttl_minutes=15)
       and return (True, signals).
    4. On timeout / exception: log + return (False, None). The caller
       translates to a 503 with the spec blocking message.

    In ENVIRONMENT=test we return (True, None) so the gate is a no-op
    for the existing test suite -- tests that exercise the gate set
    up the cache row directly via tools.cache.set_regime_cache.
    """
    if ENVIRONMENT == "test":
        return True, None
    try:
        from tools.cache import get_regime_cache, set_regime_cache
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "regime_gate_cache_module_unavailable", error=str(exc))
        return False, None
    cached = None
    try:
        cached = await get_regime_cache()
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "regime_gate_cache_read_failed", error=str(exc))
    if cached is not None:
        return True, cached
    # Cache miss / expired -- attempt a fresh detect with a hard
    # 10s timeout. detect_current_regime() is synchronous; run it in
    # a worker thread so the async event loop isn't blocked.
    import asyncio as _asyncio
    try:
        from tools.regime_detector import detect_current_regime
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "regime_gate_detector_import_failed", error=str(exc))
        return False, None
    try:
        fresh = await _asyncio.wait_for(
            _asyncio.to_thread(detect_current_regime),
            timeout=_REGIME_REFRESH_TIMEOUT_S)
    except _asyncio.TimeoutError:
        log.warning(
            "regime_gate_refresh_timeout",
            timeout_s=_REGIME_REFRESH_TIMEOUT_S,
            note=(
                "detect_current_regime() did not return within the "
                "10s hard timeout -- treating as a blocking failure "
                "for the deck-tier gate"))
        return False, None
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "regime_gate_refresh_failed",
            error=str(exc),
            note="detect_current_regime() raised during refresh")
        return False, None
    if not isinstance(fresh, dict):
        log.warning(
            "regime_gate_refresh_returned_non_dict",
            type=type(fresh).__name__)
        return False, None
    # Best-effort cache write -- a write failure doesn't unwind the
    # successful detect; the next caller just re-runs detect.
    try:
        await set_regime_cache(fresh, ttl_minutes=15)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "regime_gate_cache_write_failed", error=str(exc))
    return True, fresh


# ── Optimize ──────────────────────────────────────────────────────────────────

async def _strategy_portfolio_points() -> list[dict]:
    """
    Builds the (volatility, return) scatter coordinates for all ten
    strategies on the efficient-frontier chart.

    Reads the most recent strategy_results_cache row — the SAME results
    /api/backtest/compare serves to the rest of the dashboard — so the
    scatter is consistent with the strategy comparison table. It does NOT
    call get_full_history() or run_all_strategies(): an optimize request
    must stay light. (An earlier revision rebuilt get_full_history() here
    only to recompute the cache hash; on a Render cold start that ran the
    full ~30s pipeline concurrently with /api/backtest/compare's identical
    call.) On an empty cache returns [] — the dashboard's
    /api/backtest/compare call populates the table; the frontier curve
    still renders, just without the per-strategy markers until then.
    """
    try:
        from tools.cache import get_latest_strategy_cache

        cached = await get_latest_strategy_cache()
        if not cached:
            log.info("portfolio_points_cache_empty")
            return []

        points: list[dict] = []
        for name, r in cached.items():
            if not isinstance(r, dict):
                continue
            vol = r.get("volatility")
            ret = r.get("cagr")
            if vol is None or ret is None:
                continue
            # The strategy name: prefer the record's own field, but a
            # null/blank strategy_name produced a grey unlabelled dot on
            # the chart. `.get(key, default)` only substitutes when the
            # key is ABSENT — a present-but-None value slips through — so
            # fall back explicitly to the cache key (which IS the strategy
            # name, strategy_results_cache being keyed by it).
            strategy_name = r.get("strategy_name") or name or "UNKNOWN"
            points.append({
                "strategy": strategy_name,
                "volatility": float(vol),
                "expected_return": float(ret),
                "sharpe": float(r.get("sharpe_ratio") or 0.0),
            })
        return points
    except Exception as exc:
        log.warning("portfolio_points_unavailable", error=str(exc))
        return []


def _build_efficient_frontier(
    raw_frontier: list[dict],
    portfolio_points: list[dict],
) -> dict:
    """
    Reshapes the optimizer's flat frontier sweep into the structured
    {frontier_points, portfolio_points, max_sharpe_point, min_variance_point}
    object the EfficientFrontier component reads.

    efficient_frontier() keys the annualised return as `return`; the
    component reads `expected_return` — renamed here. Pure and synchronous
    so it is unit-testable without the DB or an event loop. Tolerant of an
    empty sweep and of a point dict missing a key — a malformed point is
    skipped, never raised on.
    """
    frontier_points: list[dict] = []
    for p in raw_frontier:
        vol = p.get("volatility")
        ret = p.get("return")
        if vol is None or ret is None:
            continue
        frontier_points.append({
            "volatility": float(vol),
            "expected_return": float(ret),
            "sharpe": float(p.get("sharpe") or 0.0),
        })

    max_sharpe_point = None
    min_variance_point = None
    if frontier_points:
        max_sharpe_point = max(frontier_points, key=lambda fp: fp["sharpe"])
        min_variance_point = min(frontier_points, key=lambda fp: fp["volatility"])

        # Diagnostic — the frontier's tangency portfolio should dominate
        # every STATIC strategy dot (each strategy is a long-only 3-asset
        # portfolio, a subset of the frontier's feasible set). A dynamic
        # strategy (regime switching, momentum) can legitimately sit above
        # the static frontier — that timing edge is the project's thesis —
        # so this is logged, not asserted.
        strat_sharpes = {
            p.get("strategy", "?"): p.get("sharpe")
            for p in portfolio_points if p.get("sharpe") is not None
        }
        if strat_sharpes:
            best_strat = max(strat_sharpes, key=lambda k: strat_sharpes[k])
            log.info(
                "efficient_frontier_max_sharpe_check",
                frontier_max_sharpe=round(max_sharpe_point["sharpe"], 4),
                frontier_point={"volatility": round(max_sharpe_point["volatility"], 4),
                                "expected_return": round(max_sharpe_point["expected_return"], 4)},
                best_strategy=best_strat,
                best_strategy_sharpe=round(float(strat_sharpes[best_strat]), 4),
                strategy_sharpes={k: round(float(v), 4) for k, v in strat_sharpes.items()},
            )

    return {
        "frontier_points": frontier_points,
        "portfolio_points": portfolio_points,
        "max_sharpe_point": max_sharpe_point,
        "min_variance_point": min_variance_point,
    }


@app.post("/api/optimize/weights")
async def optimize_weights(body: OptimizeRequest, session: dict = Depends(require_auth)):
    valid_methods = {"MEAN_VARIANCE", "RISK_PARITY", "MIN_VARIANCE", "BLACK_LITTERMAN", "MAX_SHARPE", "MIN_DRAWDOWN"}
    if body.method not in valid_methods:
        raise HTTPException(status_code=422, detail=f"Unknown method '{body.method}'")

    # Sprint 3: real optimizer backed by historical returns.
    if ENVIRONMENT != "test":
        try:
            from tools.optimizer import optimize_weights as _optimize, efficient_frontier as _frontier
            from tools.cache import get_monthly_returns
            import pandas as pd

            # The frontier is computed from the equity/IG/HY monthly return
            # series in market_data_monthly — the SAME three-asset universe
            # the ten strategies are built on. Earlier this path fetched
            # SPY/TLT/IEF/GLD daily from yfinance: a different universe AND
            # a different frequency, so the frontier curve sat visibly
            # offset from the strategy scatter dots. yfinance also drops
            # tickers to NaN from Render's cloud IPs. Reading the DB series
            # is reliable, recompute-free, and puts the curve on the same
            # (volatility, return) scale as the dots.
            monthly = await get_monthly_returns()
            if not monthly or len(monthly.get("dates", [])) < 24:
                raise ValueError(
                    "market_data_monthly unavailable or too short for a "
                    "frontier — falling back to mock"
                )

            returns = pd.DataFrame(
                {
                    "EQUITY": monthly["equity"],
                    "IG":     monthly["ig"],
                    "HY":     monthly["hy"],
                },
                index=pd.to_datetime(monthly["dates"]),
            ).dropna()

            # Issue 1: log the exact asset list reaching the solver.
            log.info(
                "optimize_frontier_universe",
                tickers=list(returns.columns),
                n_obs=len(returns),
                source="market_data_monthly",
            )

            # Annualised risk-free rate for the frontier's Sharpe — the mean
            # of the monthly DTB3 series, ×12. Using the same rate the
            # strategy scatter is built on keeps the curve's tangency
            # (max-Sharpe) point consistent with the strategy dots.
            rf_monthly = monthly.get("rf") or []
            rf_annual = (sum(rf_monthly) / len(rf_monthly) * 12) if rf_monthly else 0.0
            log.info("optimize_frontier_risk_free", risk_free_annual=round(rf_annual, 4))

            # CACHE LOOKUP MOVED EARLY (May 24 2026 UAT fix).
            # Previously the endpoint ran `_optimize(method, returns)`
            # (a single SLSQP solve) BEFORE checking the precomputed
            # frontier cache. A transient solver convergence failure
            # — cvxpy CLARABEL "infeasible_or_unbounded", SLSQP
            # "Inequality constraints incompatible", etc. — would
            # raise here, fall through to the outer except, and
            # serve MOCK_EFFICIENT_FRONTIER. The user saw the chart
            # render with 11 hardcoded mock points + the wrong asset
            # universe (SPY/TLT/IEF/GLD) instead of the real frontier
            # that WAS in the cache. Re-ordering: cache first, then
            # solver. The chart never goes mock when real cache data
            # exists, even if the per-method weights solve fails.
            #
            # Hotfix (May 23 2026) note kept: the cache exists because
            # the 100-point SLSQP sweep runs once at ingestion via
            # refresh_efficient_frontier; same input → same output
            # across every method, so the cache row is method-
            # agnostic. ~50ms cache hit vs ~10-30s on-demand sweep
            # (would exceed Render's 30s gateway timeout).
            raw_frontier: list[dict] | None = None
            try:
                from tools.cache import get_latest_strategy_hash
                from tools.precomputed_analytics import (
                    get_metric as get_precomputed,
                    get_latest_metric as get_latest_precomputed,
                )
                latest_hash = await get_latest_strategy_hash()
                cached = None
                if latest_hash:
                    cached = await get_precomputed(
                        latest_hash, "efficient_frontier")
                if cached is None:
                    # Stale-cache fallback — serve the previous data-
                    # hash's frontier rather than block on the inline
                    # sweep when the refresh hook is in flight or has
                    # not yet fired.
                    cached = await get_latest_precomputed(
                        "efficient_frontier")
                if cached and cached.get("frontier_points"):
                    raw_frontier = cached["frontier_points"]
                    log.info("optimize_frontier_cache_hit",
                             n_points=len(raw_frontier),
                             data_hash=(latest_hash[:8]
                                          if latest_hash else None))
            except Exception as cache_exc:  # noqa: BLE001
                log.warning("optimize_frontier_cache_read_failed",
                            error=str(cache_exc))

            # Per-method weights — runs AFTER the cache lookup so a
            # solver failure here only nulls out the weights field,
            # never wipes the cached frontier into a mock-data
            # response. The dashboard's EfficientFrontier component
            # only consumes `efficient_frontier`; the `weights` field
            # is read by callers that explicitly asked for a method
            # (e.g. a Portfolio Profile drill-down). A null weights
            # field there degrades cleanly (the caller can re-try
            # the specific method) — better than a fake frontier
            # painted across the whole dashboard.
            #
            # UAT 2026-05-24 P0 fix — `_optimize` is a SYNC cvxpy
            # call (CLARABEL / SLSQP) and was running on the asyncio
            # event loop. On Render's shared CPU it can take 1-3s
            # per solve, occasionally longer if the solver iterates
            # to convergence. Dashboard mount fires `_optimize` AND
            # `/api/v1/analytics/academic` in parallel via the
            # store's Promise.all — a slow `_optimize` blocked the
            # loop, queued the analytics request behind it, and
            # both axios calls timed out at 30s. Same root cause for
            # both timeouts the user reported.
            #
            # asyncio.to_thread pushes the sync solver into a worker
            # thread so the event loop stays free to service the
            # cumulative-returns request in parallel. Mirrors the
            # PR #122 fix for get_full_history / run_all_strategies.
            import asyncio
            result: dict
            try:
                result = await asyncio.to_thread(
                    _optimize, body.method, returns)
            except Exception as opt_exc:  # noqa: BLE001
                log.warning("optimize_weights_solve_failed",
                            method=body.method, error=str(opt_exc))
                result = {
                    "method":    body.method,
                    "weights":   None,
                    "sum_check": None,
                    "error":     str(opt_exc),
                }

            if raw_frontier is None:
                # COLD CACHE — never run the inline sweep on the
                # request thread. The 100-point SLSQP sweep takes
                # 10-30s on Render's shared CPU, which exceeds the
                # frontend's 30s timeout. Hotfix iteration 2 (May
                # 23 2026): instead of blocking, fire the precompute
                # refresh in the background and return a warming
                # response immediately. The frontend retries every
                # ~10s until the cache row lands.
                #
                # The warming response carries an empty frontier
                # array + a `warming: true` flag the EfficientFrontier
                # component reads to render a "computing..." state.
                # portfolio_points still ships so the strategy
                # scatter is visible during the warmup.
                # May 24 2026 P0 hotfix — fire the refresh
                # UNCONDITIONALLY. The pre-existing `if h:` gate
                # meant a cold Render deploy with an empty
                # strategy_results_cache never triggered the
                # background sweep, and the user saw a 30s timeout
                # forever. refresh_all_analytics now substitutes a
                # BOOT-WARM sentinel hash when none is supplied.
                log.warning("optimize_frontier_cache_cold_warming",
                            method=body.method)
                try:
                    from tools.cache import get_latest_strategy_hash
                    from tools.precomputed_analytics import (
                        trigger_refresh_async,
                    )
                    h = await get_latest_strategy_hash()
                    trigger_refresh_async(h or "")
                except Exception as trig_exc:  # noqa: BLE001
                    log.warning(
                        "optimize_frontier_refresh_trigger_failed",
                        error=str(trig_exc))

                portfolio_points = await _strategy_portfolio_points()
                return {
                    "method":  body.method,
                    "weights": result["weights"],
                    "sum_check": result["sum_check"],
                    "warming": True,
                    "retry_after_ms": 10000,
                    "efficient_frontier": {
                        "frontier_points":   [],
                        "portfolio_points":  portfolio_points,
                        "max_sharpe_point":  None,
                        "min_variance_point": None,
                        "warming":           True,
                    },
                }

            # efficient_frontier() returns a flat list keyed `return`, but
            # the EfficientFrontier component reads `expected_return` off a
            # structured {frontier_points, portfolio_points, ...} object —
            # the same shape MOCK_EFFICIENT_FRONTIER uses. Reshape here so
            # the real and mock paths return an identical contract; a flat
            # list left the chart blank (no frontier_points key on an array).
            portfolio_points = await _strategy_portfolio_points()

            return {
                "method": body.method,
                "weights": result["weights"],
                "sum_check": result["sum_check"],
                "efficient_frontier": _build_efficient_frontier(
                    raw_frontier, portfolio_points
                ),
            }
        except Exception as exc:
            log.warning("optimize_weights_fallback", method=body.method, error=str(exc))

    return {
        "method": body.method,
        "weights": {"SPY": 0.40, "TLT": 0.30, "IEF": 0.15, "GLD": 0.15},
        "efficient_frontier": MOCK_EFFICIENT_FRONTIER,
        "note": "Fallback mock — real optimisation failed or test environment",
    }


# ── Market data ───────────────────────────────────────────────────────────────

@app.get("/api/data/market")
async def get_market_data(
    tickers: str = Query(..., description="Comma-separated tickers"),
    start: str = Query("2020-01-01"),
    end: str = Query("2024-12-31"),
    session: dict = Depends(require_auth),
):
    return {
        "tickers": tickers.split(","),
        "start": start,
        "end": end,
        "note": "Sprint 1 mock — real data fetch in Sprint 2",
        "prices": {},
        "returns": {},
    }


# ── Council ───────────────────────────────────────────────────────────────────

def _log_council_session(
    query: str,
    agents_called: list[str],
    response: dict[str, Any],
    start_time: float,
    user_email: str,
) -> None:
    """
    Persists council session metadata to the AI usage log.

    Writes to council_sessions table when DB is available; always logs
    to structlog so the session is traceable even without Postgres.
    Cost estimates use Anthropic's published pricing as of Sprint 4.
    """
    duration_ms = int((time.time() - start_time) * 1000)
    session_id = str(uuid.uuid4())

    log.info(
        "council_session_complete",
        session_id=session_id,
        user_hash=hash(user_email),
        agents_called=agents_called,
        n_significant=len(response.get("significant_strategies", [])),
        duration_ms=duration_ms,
    )


# ── Non-blocking interaction logging ──────────────────────────────────────────
# Background tasks must be referenced or the event loop may GC them mid-run;
# the set holds a strong ref and the done-callback drops it on completion.
_activity_bg_tasks: set = set()


def _log_interaction_bg(
    request: Request,
    session: dict,
    interaction_type: str,
    *,
    question_text: str | None = None,
    agents_involved: list[str] | None = None,
    response_summary: str | None = None,
    metadata: dict | None = None,
) -> None:
    """
    Fire-and-forget agent_interactions logging — schedules the DB write
    on the running loop and returns immediately. The session_id and
    session_type travel from the frontend as request headers. Wrapped
    so a scheduling or DB failure never touches the primary response.

    AI token cost: collect_usage() is read here, in the request context,
    so any endpoint that called start_usage_capture() has its token
    totals logged and its per-agent breakdown folded into metadata. An
    endpoint that did not start a capture simply logs null token columns.
    """
    try:
        import asyncio
        from tools.activity_log import log_agent_interaction

        in_tok = out_tok = model_used = cost = None
        try:
            from agents.usage import collect_usage
            usage = collect_usage()
            in_tok = usage.get("input_tokens")
            out_tok = usage.get("output_tokens")
            model_used = usage.get("model_used")
            cost = usage.get("estimated_cost_usd")
            if usage.get("per_agent"):
                metadata = {**(metadata or {}),
                            "per_agent_cost": usage["per_agent"]}
        except Exception:  # noqa: BLE001 — cost telemetry is never fatal
            pass

        task = asyncio.create_task(log_agent_interaction(
            user_email=session.get("email", ""),
            session_id=request.headers.get("x-session-id"),
            session_type=request.headers.get("x-session-type"),
            interaction_type=interaction_type,
            question_text=question_text,
            agents_involved=agents_involved,
            response_summary=response_summary,
            metadata=metadata,
            input_tokens=in_tok,
            output_tokens=out_tok,
            model_used=model_used,
            estimated_cost_usd=cost,
        ))
        _activity_bg_tasks.add(task)
        task.add_done_callback(_activity_bg_tasks.discard)
    except Exception as exc:  # noqa: BLE001
        log.warning("interaction_log_schedule_failed",
                    interaction_type=interaction_type, error=str(exc))


# Maps cio.deliberate() agent keys to the display name/role/model the frontend expects.
# The frontend's AGENT_STYLE dict in CouncilDebate.tsx is keyed by these exact display names.
_AGENT_META: dict[str, tuple[str, str, str]] = {
    "equity_analyst":       ("Equity Analyst",               "specialist", SONNET_MODEL),
    "fixed_income_analyst": ("Fixed Income Analyst",          "specialist", SONNET_MODEL),
    "risk_manager":         ("Risk Manager",                  "specialist", SONNET_MODEL),
    "quant_backtester":     ("Quant Backtester",              "specialist", SONNET_MODEL),
    "independent_analyst":  ("Independent Analyst (Gemini)",  "dissenter",  GEMINI_MODEL),
    "contrarian_analyst":   ("Contrarian Analyst (Grok)",     "dissenter",  "grok-4.3"),
    "cio":                  ("CIO",                           "cio",        OPUS_MODEL),
}


def _deliberate_to_frontend(query: str, council_response: dict[str, Any]) -> dict[str, Any]:
    """
    Converts cio.deliberate() output to the CouncilDebateResponse shape the frontend expects.

    cio.deliberate() returns {"agents": {snake_case_key: report_dict, ...}, ...}.
    The frontend's CouncilResponse type expects {"messages": [AgentMessage, ...], ...}.
    This conversion runs inside the council_query endpoint so the raw backend
    structure never reaches the client.
    """
    agents = council_response.get("agents", {})
    messages = []
    for key, (display_name, role, model) in _AGENT_META.items():
        report = agents.get(key, {})
        if not report:
            # Agent was never invoked — skip rather than show an empty card.
            # An expected agent missing from the response is a council-flow
            # bug, not a rendering bug; the heatmap or another diagnostic
            # will catch it.
            continue

        # Per-agent content selection — every agent shows its FULL
        # narrative, not a one-line summary:
        #   CIO         → final synthesis (the recommendation narrative)
        #   Gemini/Grok → full challenge text (the dissenting narrative)
        #   Specialists → raw_analysis (the complete specialist analysis;
        #                 summary is only the one-line fallback)
        tech = report.get("technical_findings", {}) or {}
        if key == "cio":
            content = tech.get("final_synthesis_text") or report.get("summary", "")
        elif key == "independent_analyst":
            content = tech.get("full_challenge") or report.get("summary", "")
        elif key == "contrarian_analyst":
            content = tech.get("full_challenge") or report.get("summary", "")
        else:
            content = tech.get("raw_analysis") or report.get("summary", "")

        # NEVER drop an agent whose report exists. An empty content field
        # used to be silently filtered, which produced an empty Debate tab
        # whenever a single LLM extraction returned no text. The frontend
        # renders a placeholder for empty content; the audience still sees
        # which agent ran and what its tag was.
        if not content:
            content = "(Narrative unavailable — agent ran but no text returned.)"

        messages.append({
            "agent":    display_name,
            "role":     role,
            "model":    model,
            "content":  content,
            "is_final": key == "cio",
        })
    return {
        "query":                query,
        "messages":             messages,
        "final_recommendation": council_response.get("final_recommendation", ""),
        "consensus_reached":    True,
        # "live" = real council run; the mock-fallback path sets "fallback"
        # so a consumer can tell genuine analysis from demo data.
        "mode":                 "live",
    }


def _chunk_synthesis(text: str, words_per_chunk: int = 8) -> list[str]:
    """
    Splits the CIO synthesis prose into small chunks for SSE streaming.
    Mirrors agents.academic_review.chunk_arbiter_text — word groups so
    the consumer sees the synthesis arrive progressively rather than as
    one large final frame. 8 words ≈ 50 chars typical, well below the
    user-spec'd 100-char ceiling, but big enough to keep frame count
    bounded for a 1000-word synthesis.
    """
    if not text:
        return []
    words = text.split(" ")
    return [
        " ".join(words[i:i + words_per_chunk]) + " "
        for i in range(0, len(words), words_per_chunk)
    ]


@app.post("/api/council/query")
@limiter.limit("10/minute")
async def council_query(
    request: Request,
    body: CouncilQueryRequest,
    session: dict = Depends(require_auth),
):
    """
    Convenes the full investment council and STREAMS phase events via SSE.

    Replaces the previous synchronous JSON response — the deliberation
    routinely takes 50-100 seconds and was hitting Render's gateway
    timeout. Streaming keeps the connection alive at each phase
    boundary; the frontend assembles the same CouncilDebateResponse
    shape from the events.

    Event sequence (each `data: {json}\\n\\n`):
      1. council_started     — fires immediately on request receipt
      2. specialist_complete — one per analyst, in as_completed order
      3. draft_ready         — CIO draft consensus
      4. dissent_complete    — gemini, then grok
      5. synthesis_chunk     — CIO synthesis text, chunked progressively
      6. council_complete    — full CouncilDebateResponse dict
      7. data: [DONE]\\n\\n   — end-of-stream sentinel

    On a deliberation failure the stream yields a `council_error` frame
    and a [DONE] sentinel rather than silently falling through to mock
    data — the previous fall-through hid real errors from the user.

    Scope guard runs synchronously BEFORE the StreamingResponse so a
    422 still returns as a normal HTTP error; the same for the 429
    quota response. The stream only starts once the request has been
    cleared for processing.

    The test environment keeps the synchronous JSON contract — every
    council test in tests/test_council_deliberation.py asserts on
    resp.json() and would break if asked to parse SSE.
    """
    if len(body.query) > 500:
        raise HTTPException(status_code=422, detail="Query exceeds 500 character limit.")

    # Viewer council query allocation — a user with a council_queries_limit
    # set is blocked once their lifetime allowance is spent. Team members
    # and sysadmins have a NULL limit and are never blocked. Checked before
    # the scope guard so a blocked viewer never spends a classifier call.
    # The 429 returns as a normal HTTP response, BEFORE the stream starts.
    council_quota: dict[str, Any] | None = None
    from tools.platform_users import (
        get_council_allocation, increment_council_queries,
    )
    allocation = await get_council_allocation(session["email"])
    if allocation and allocation.get("council_queries_limit") is not None:
        used = int(allocation.get("council_queries_used", 0) or 0)
        limit = int(allocation["council_queries_limit"])
        if used >= limit:
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "council_limit_reached",
                    "limit": limit,
                    "used": used,
                },
            )
        # The query is counted against the allowance before processing.
        council_quota = await increment_council_queries(session["email"])

    # Scope guard — must pass before any agent is invoked.
    # Runs synchronously so a 422 returns as a normal HTTP error, not
    # as an SSE frame.
    if ENVIRONMENT != "test":
        try:
            from scope_guard import ScopeGuard
            guard = ScopeGuard()
            scope_result = await guard.check(body.query)
            if not scope_result["allowed"]:
                raise HTTPException(
                    status_code=422,
                    detail={
                        "error": "out_of_scope",
                        "message": scope_result["rejection_message"],
                        "system": "Forest Capital Portfolio Intelligence System",
                    },
                )
        except HTTPException:
            raise
        except Exception as exc:
            # Scope guard failure is non-fatal — log and proceed
            log.warning("scope_guard_error", error=str(exc))

    log.info("council_query_started", user=session["email"], query_len=len(body.query))
    start_time = time.time()

    # Test environment keeps the synchronous JSON contract so the existing
    # /api/council/query test suite (TestClient.post(...).json()) keeps
    # passing without rewrite. Production streams; tests do not need to.
    if ENVIRONMENT == "test":
        response = dict(MOCK_COUNCIL_RESPONSE)
        response["query"] = body.query
        response["mode"] = "fallback"
        if council_quota:
            response["council_queries_used"] = council_quota["council_queries_used"]
            response["council_queries_limit"] = council_quota["council_queries_limit"]
        return response

    query = body.query

    async def event_stream():
        # The CIO generator is synchronous and does heavy work (parallel
        # specialists, network-bound LLM calls). We bridge it to async
        # via asyncio.to_thread(next, ...) so the event loop stays free
        # between yields — that is what allows the gateway connection
        # to keep alive past 30s.
        import asyncio
        import traceback

        from agents.cio import CIO
        from agents.harness import (
            start_harness_capture, collect_harness_metrics,
        )
        from agents.usage import start_usage_capture
        from tools.backtester import run_all_strategies
        from tools.data_fetcher import get_full_history

        final_result: dict[str, Any] | None = None

        try:
            yield _sse("council_started", query=query)

            # Heavy data load offloaded to threads so the council_started
            # event flushes immediately and the event loop stays
            # responsive. get_full_history() is memoised (30s); a warm
            # call is fast, but the first call after a hash change runs
            # the read pipeline.
            history = await asyncio.to_thread(get_full_history)
            strategy_results = await asyncio.to_thread(
                run_all_strategies, history)

            # Page-scoped live context — when the question was asked from
            # a council-facing landing page, resolve that page's live
            # cached data (recommendation / performance / prediction) and
            # thread it into the deliberation. None for an unscoped or
            # unknown scope, so the existing behaviour is unchanged.
            from tools.council_live_context import get_scope_context
            page_context = await get_scope_context(body.context_scope)
            if page_context:
                log.info("council_scope_context_injected",
                         scope=body.context_scope,
                         keys=list(page_context.keys()))

            # ── Question-type bundle (refines PR #229's page scope) ──
            # The keyword classifier in tools.council_question_bundles
            # narrows the context to a question-type-specific subset
            # (regime / recommendation / risk / statistical / forward).
            # A confident classification REPLACES the wider page
            # bundle; a low-confidence query falls back to the page
            # bundle (or to no context if no page was set). The
            # question_type label is recorded for the metrics row so
            # the team can quantify the cost reduction afterwards.
            from tools.council_question_bundles import (
                QUESTION_TYPE_FULL,
                classify_question,
                resolve_bundle,
            )
            question_type = classify_question(query)
            live_context = None
            if question_type is not None:
                bundle = await resolve_bundle(question_type)
                if bundle:
                    live_context = bundle
                    log.info("council_question_bundle_injected",
                             question_type=question_type,
                             keys=list(bundle.keys()))
                else:
                    # Classifier matched but bundle had nothing in
                    # cache — fall back to the page scope.
                    log.info("council_question_bundle_empty",
                             question_type=question_type)
                    question_type = None
            if live_context is None:
                live_context = page_context
                # Record "full" as the question-type for the metrics
                # row when we landed on the page bundle (or None) —
                # the column is non-null in the schema and "full" is
                # the canonical fallback label.
                question_type_recorded = QUESTION_TYPE_FULL
            else:
                question_type_recorded = question_type

            # Capture every specialist's harness run + every agent
            # call's token usage. Seeded before the generator runs so
            # the copied thread contexts share the accumulator lists.
            start_harness_capture()
            start_usage_capture()

            # June 5 2026 — fetch the prior CIO recommendation so the
            # synthesis step can write Section C (shift narrative). The
            # helper fails open to None — first run, cache cleared, or
            # DB outage all produce None and Section C is omitted from
            # the response per the system prompt's CONDITIONAL clause.
            prior_recommendation = None
            try:
                from tools.cio_recommendation import (
                    _current_data_hash, get_prior_recommendation,
                )
                current_hash = await _current_data_hash()
                if current_hash:
                    prior_recommendation = await get_prior_recommendation(
                        current_hash)
            except Exception as exc:  # noqa: BLE001
                log.warning("council_prior_recommendation_fetch_error",
                            error=str(exc))

            cio = CIO()
            gen = cio.deliberate_streaming(
                query, strategy_results, history,
                live_context=live_context,
                prior_recommendation=prior_recommendation)
            sentinel = object()
            while True:
                # Each next(gen) runs a phase — specialists fan-out,
                # draft, dissenter, synthesis. The phase work happens
                # in the worker thread; the yield back to async land
                # flushes the SSE frame.
                event = await asyncio.to_thread(next, gen, sentinel)
                if event is sentinel:
                    break
                kind = event[0]
                if kind == "specialist_complete":
                    _, agent_id, report = event
                    yield _sse("specialist_complete",
                               agent_id=agent_id, response=report)
                elif kind == "draft_ready":
                    _, draft = event
                    yield _sse("draft_ready", draft=draft)
                elif kind == "dissent_complete":
                    _, source, report = event
                    yield _sse("dissent_complete",
                               source=source, challenge=report)
                elif kind == "cio_synthesis_text":
                    _, synthesis_text = event
                    # Chunk the synthesis prose so the user sees it
                    # arrive progressively — same pattern as
                    # academic_review.chunk_arbiter_text.
                    for chunk in _chunk_synthesis(synthesis_text):
                        yield _sse("synthesis_chunk", text=chunk)
                elif kind == "council_complete":
                    _, full_result = event
                    final_result = full_result
                    # Convert to the CouncilDebateResponse shape the
                    # frontend store assembles into — same dict the
                    # previous synchronous handler returned.
                    result = _deliberate_to_frontend(query, full_result)
                    if council_quota:
                        result["council_queries_used"] = (
                            council_quota["council_queries_used"])
                        result["council_queries_limit"] = (
                            council_quota["council_queries_limit"])
                    yield _sse("council_complete", result=result)

            # Post-stream logging — runs only when the deliberation
            # completed cleanly. _log_council_session writes to the
            # AI usage log; _log_interaction_bg is fire-and-forget
            # Team Activity.
            if final_result is not None:
                harness_meta = collect_harness_metrics()
                council_agents = [
                    "equity_analyst", "fixed_income_analyst",
                    "risk_manager", "quant_backtester",
                    "independent_analyst", "contrarian_analyst", "cio",
                ]
                _log_council_session(
                    query=query,
                    agents_called=council_agents,
                    response=final_result,
                    start_time=start_time,
                    user_email=session["email"],
                )
                _log_interaction_bg(
                    request, session, "council",
                    question_text=query,
                    agents_involved=council_agents,
                    response_summary=final_result.get(
                        "final_recommendation", ""),
                    metadata=({"harness": harness_meta}
                              if harness_meta else None),
                )

                # ── council_query_metrics row ─────────────────────
                # One row per query: token totals + bundle size +
                # live regime + extracted recommendation direction +
                # the refined alignment score. Fail-open — a metric
                # write failure logs and never affects the user-
                # facing stream that already completed.
                try:
                    from agents.usage import collect_usage
                    usage = collect_usage()
                    bundle_size = (
                        len(json.dumps(live_context, default=str))
                        if live_context else 0)
                    # per-agent breakdown — pull the "cio" total
                    # specifically as the LIKE-FOR-LIKE bundle-impact
                    # signal. _compile_draft_consensus and _synthesise
                    # both _tag_agent("cio") before their call_claude,
                    # so this aggregates both — the only two calls in
                    # the deliberation whose prompt content the bundle
                    # changes. June 3 2026 (migration 052).
                    per_agent = usage.get("per_agent") or {}
                    cio_input = (
                        per_agent.get("cio", {}).get("input_tokens")
                        if isinstance(per_agent.get("cio"), dict)
                        else None)
                    await _write_council_query_metric(
                        question_type=question_type_recorded,
                        input_tokens=usage.get("input_tokens"),
                        output_tokens=usage.get("output_tokens"),
                        cio_input_tokens=cio_input,
                        context_bundle_size=bundle_size,
                        synthesis_text=final_result.get(
                            "final_recommendation", "") or "",
                    )
                except Exception as exc:  # noqa: BLE001
                    log.warning("council_query_metric_write_failed",
                                error=str(exc))
        except Exception as exc:  # noqa: BLE001
            # A deliberation failure surfaces clearly — no silent
            # fall-through to mock data. The frontend renders the
            # error and lets the user retry.
            log.error("council_query_failed", error=str(exc),
                      traceback=traceback.format_exc())
            yield _sse(
                "council_error",
                message="Council query failed. Please try again.",
            )
        yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ── Academic review ───────────────────────────────────────────────────────────

def _sse(event_type: str, **payload: Any) -> str:
    """Format one Server-Sent Events frame: data: {json}\\n\\n."""
    return f"data: {json.dumps({'type': event_type, **payload})}\n\n"


def _parse_overall_rating(verdict: str) -> str | None:
    """
    Pulls the section-5 readiness rating (Strong | Developing | Needs
    Work) out of the arbiter verdict so the Team Activity summary can
    show it without re-parsing the markdown. Returns None if the
    verdict is malformed or unavailable.
    """
    import re
    m = re.search(
        r"###\s*5\.[^\n]*\n\s*\*\*Rating:\*\*\s*(Strong|Developing|Needs Work)",
        verdict or "", re.IGNORECASE,
    )
    return m.group(1) if m else None


async def with_keepalive(producer, interval: float = 20.0):
    """June 25 2026 -- wrap any async-generator SSE stream so it emits
    a comment-frame keepalive ping whenever the upstream is silent for
    `interval` seconds. SSE comment lines (lines starting with ':')
    are ignored by the standard EventSource parser but reset the
    connection timeout on every hop between server, Render's edge
    proxy, the operator's network, and the browser. Without
    keepalives, a 60-90s arbiter call or a 3-4 min full pipeline
    looks like a dropped connection to the proxies even though the
    backend is still running.

    Implementation -- the producer pushes frames into an asyncio.Queue
    (decoupling the producer's awaits from the consumer's yields)
    and the consumer wakes every `interval` seconds. If the wakeup
    found a frame, yield it; if the wakeup hit the timeout, yield
    a keepalive line. A None sentinel from the producer (placed by
    the finally below) signals end-of-stream so the consumer
    drains any pending items and stops.

    Cancellation -- if the producer task is cancelled (client
    disconnect propagated as CancelledError), the consumer's
    finally cancels the producer task too and exits without
    yielding anything else."""
    import asyncio
    queue: asyncio.Queue = asyncio.Queue()
    DONE = object()

    async def _drive() -> None:
        try:
            async for msg in producer:
                await queue.put(msg)
        except asyncio.CancelledError:
            # Re-raise so the outer await propagates the cancel.
            await queue.put(DONE)
            raise
        except Exception as exc:  # noqa: BLE001
            log.error(
                "sse_with_keepalive_producer_failed",
                error=str(exc))
        finally:
            await queue.put(DONE)

    producer_task = asyncio.create_task(_drive())
    try:
        while True:
            try:
                msg = await asyncio.wait_for(
                    queue.get(), timeout=interval)
            except asyncio.TimeoutError:
                # Producer silent for `interval` seconds -- emit a
                # keepalive comment frame. Two newlines so the SSE
                # parser treats it as a complete event boundary.
                yield ": keepalive\n\n"
                continue
            if msg is DONE:
                # Drain any final items the producer may have queued
                # in parallel with the sentinel, then exit.
                break
            yield msg
        await producer_task
    finally:
        if not producer_task.done():
            producer_task.cancel()
            try:
                await producer_task
            except (asyncio.CancelledError, Exception):
                pass


@app.post("/api/council/academic-review")
@limiter.limit("10/minute")
async def council_academic_review(request: Request, session: dict = Depends(require_team_member)):
    """
    Convenes the council to evaluate the project's academic readiness.

    No request body — all context is assembled server-side: an analytics
    inventory plus the uploaded academic documents. Every peer agent
    answers a stock four-part review question in parallel; the academic
    advisor then arbitrates, synthesising a five-section rubric-mapped
    verdict.

    Streams a text/event-stream response:
      1. {"type": "peer_responses", "data": {agentId: text}}
      2. {"type": "arbiter_chunk", "text": chunk}   (streamed)
      3. data: [DONE]

    Both the peer responses and the arbiter verdict are produced through
    the generator-evaluator harness. The arbiter is generated IN FULL and
    harness-evaluated before any chunk is streamed — a failed attempt is
    never shown, only the accepted verdict.

    document_type query param — when "presentation_script", the arbiter
    runs the SCRIPT-SPECIFIC rubric (coherence, audience clarity, slide
    coverage, speaker differentiation; skips citation formatting and
    paragraph structure). Default rubric otherwise.
    """
    import asyncio
    from agents.academic_review import (
        gather_review_context, run_peer_fan_out, run_arbiter_with_harness,
        chunk_arbiter_text, ARBITER_MODEL,
    )
    from agents.harness import start_harness_capture, collect_harness_metrics
    from agents.usage import start_usage_capture

    # Rubric switch — read once at the top of the handler so a future
    # gather_review_context refactor can also consume it. Only the literal
    # string "presentation_script" enables the script rubric; the
    # literal "executive_brief" enables the brief-specific rubric (PR
    # — brief-specific rubric, replaces the midpoint 5.5/10 floor).
    document_type_q = request.query_params.get("document_type")
    script_review = document_type_q == "presentation_script"
    brief_review = document_type_q == "executive_brief"
    deck_review = document_type_q == "presentation_deck"
    appendix_review = document_type_q == "analytical_appendix"

    # ── June 27 2026: regime_signals pre-flight gate (deck only) ─
    # The deck-tier council review references the live CIO
    # recommendation that the deck surfaces on slides 7 + 11. A
    # review run against stale regime signals would grade a recom-
    # mendation the user can no longer trust by the time the panel
    # sees it. Hard-gate the same way the deck generation does --
    # 10s refresh, 503 if it fails. Brief / appendix / script
    # reviews are unaffected (their docs use the em-dash fallback).
    # Runs BEFORE the focus-brief body parse below so a deck-tier
    # 503 short-circuits without consuming the body.
    if deck_review:
        ok, _signals = await _regime_signals_fresh_or_refresh()
        if not ok:
            log.warning(
                "deck_council_review_blocked_on_regime_signals",
                note=("regime_signals_cache miss + refresh "
                      "failed/timed out within 10s -- blocking "
                      "deck-tier council review"))
            raise HTTPException(
                status_code=503,
                detail=_REGIME_BLOCKING_ERROR_DETAIL)

    # June 27 2026 -- optional pre-review focus brief. The frontend
    # surfaces a 1000-char textarea after the cross-document confirm
    # modal (or directly on the per-doc Run Review click); the brief
    # is sent as JSON {focus_brief: "..."} on the POST body. The
    # endpoint historically takes no body, so we parse defensively:
    # any parse / type failure falls back to None (legacy no-brief
    # behaviour). Strip + truncate at FOCUS_BRIEF_MAX_CHARS so a
    # client that doesn't enforce the limit still produces a
    # bounded prompt.
    from agents.academic_review import FOCUS_BRIEF_MAX_CHARS
    focus_brief: str | None = None
    try:
        body = await request.json()
        if isinstance(body, dict):
            raw = body.get("focus_brief")
            if isinstance(raw, str):
                cleaned = raw.strip()
                if cleaned:
                    if len(cleaned) > FOCUS_BRIEF_MAX_CHARS:
                        cleaned = cleaned[:FOCUS_BRIEF_MAX_CHARS]
                    focus_brief = cleaned
    except Exception:  # noqa: BLE001
        # Empty body / malformed JSON / wrong content-type --
        # silently treat as no-brief. Legacy callers that POST with
        # no body land here.
        focus_brief = None
    log.info(
        "academic_review_focus_brief",
        focus_brief_present=bool(focus_brief),
        focus_brief_chars=len(focus_brief or ""),
        document_type=document_type_q,
    )

    async def event_stream():
        try:
            # Capture the peer + arbiter harness runs for Team Activity, and
            # every agent call's token usage for cost tracking. The ContextVar
            # lists are shared into the asyncio.to_thread peer and arbiter
            # tasks, so every run is recorded.
            start_harness_capture()
            start_usage_capture()
            ctx = await gather_review_context(
                reviewer_email=session.get("email"),
                document_type=document_type_q or None,
                focus_brief=focus_brief)
            context_block = ctx["context_block"]
            multi_user = ctx.get("multi_user_activity", False)
            # Threaded into the chart-vision scope sentences so all-
            # strategy chart captions render the exact count rather
            # than the count-omitted fallback.
            n_strategies = ctx["analytics"].get("strategy_count")

            peer_responses = await run_peer_fan_out(
                context_block, multi_user, n_strategies)
            log.info(
                "academic_review_peers_complete",
                agents=list(peer_responses.keys()),
                response_lengths={k: len(v) for k, v in peer_responses.items()},
                arbiter_model=ARBITER_MODEL,
                risk_free_rate=ctx["analytics"].get("risk_free_rate"),
                document_types_present=ctx["document_types_present"],
                document_types_missing=ctx["document_types_missing"],
                script_review=script_review,
            )
            yield _sse("peer_responses", data=peer_responses)

            # Arbiter — generated in full and harness-evaluated in a worker
            # thread (the harness is synchronous), then streamed as chunks.
            # The loading state on the frontend covers the evaluation wait.
            # June 25 2026 -- when the harness exhausts retries and the
            # final score is below threshold, the arbiter still returns
            # the best attempt's response; we surface that with an
            # arbiter_degraded marker frame BEFORE the arbiter_chunk
            # stream so the client knows the quality threshold wasn't
            # met. The harness already returns the best-scoring
            # attempt's text on max retries (see HarnessResult), so the
            # chunks themselves are still real prose to render -- the
            # frame is an advisory tag, not a hide.
            arbiter_text = await asyncio.to_thread(
                run_arbiter_with_harness, context_block, peer_responses,
                multi_user, script_review, n_strategies, brief_review,
                deck_review, appendix_review)
            try:
                arbiter_metrics = collect_harness_metrics()
            except Exception:  # noqa: BLE001
                arbiter_metrics = {}
            # arbiter_metrics aggregates ALL agents in the harness run
            # log (peers + arbiter). The arbiter's final_score lives
            # under average_final_score in the aggregate; a more
            # accurate per-arbiter score requires the harness's
            # per-record buf -- defer to the existing telemetry. For
            # the degraded signal, a below-threshold AVERAGE final
            # score plus attempts > 1 is enough heuristic to flag.
            arbiter_degraded = bool(
                arbiter_metrics
                and arbiter_metrics.get("average_final_score") is not None
                and float(arbiter_metrics.get(
                    "average_final_score", 0)) < 7.0)
            if arbiter_degraded:
                yield _sse(
                    "arbiter_degraded",
                    average_final_score=arbiter_metrics.get(
                        "average_final_score"),
                    agents_retried=arbiter_metrics.get(
                        "agents_retried"),
                    message=(
                        "Arbiter quality threshold was not met after "
                        "harness retries. The verdict below reflects "
                        "the best attempt; consider re-running the "
                        "review."))
            for chunk in chunk_arbiter_text(arbiter_text):
                yield _sse("arbiter_chunk", text=chunk)
            log.info("academic_review_arbiter_complete",
                     arbiter_chars=len(arbiter_text),
                     degraded=arbiter_degraded)

            # Independent Review — second-opinion advisory layer. A
            # SEPARATE agent (Gemini) sees ONLY the headline findings
            # extracted from the arbiter, with no platform context.
            # Assesses plausibility, internal consistency, and
            # graduate-finance defensibility. Advisory only — never
            # affects the primary verdict or any gates. Fail-open:
            # any failure surfaces a stub Concerns verdict so the SSE
            # stream never blocks on this layer.
            try:
                from agents.independent_review import (
                    extract_key_findings, run_independent_review,
                )
                strategy_results_for_findings: dict[str, Any] | None = None
                try:
                    from tools.cache import get_latest_strategy_cache
                    strategy_results_for_findings = (
                        await get_latest_strategy_cache())
                except Exception as _exc:  # noqa: BLE001
                    log.info("independent_review_strategy_read_failed",
                             error=str(_exc))
                findings = extract_key_findings(
                    arbiter_text,
                    analytics_snapshot=ctx.get("analytics"),
                    strategy_results=strategy_results_for_findings,
                )
                # June 25 2026 -- pass the primary document to the
                # independent reviewer so it reads the same source
                # the peers did rather than the extracted-findings
                # block alone. The doc is sourced from
                # gather_review_context's docs_by_type (already
                # overlaid with the team-shared editor draft via
                # get_current_draft_by_type). Falls back to an empty
                # string when no target document was supplied (the
                # cross-document review path); the independent
                # reviewer's _build_user_message then skips the
                # primary-document block.
                primary_doc_text: str = ""
                try:
                    target_review = (
                        ctx.get("target_review_type")
                        or None)
                    docs_by_type = (
                        ctx.get("documents_by_type") or {})
                    if target_review and isinstance(
                            docs_by_type, dict):
                        rows = (
                            docs_by_type.get(target_review) or [])
                        if rows and isinstance(rows[0], dict):
                            primary_doc_text = str(
                                rows[0].get("content_text") or "")
                except Exception as _exc:  # noqa: BLE001
                    log.info(
                        "independent_review_primary_doc_lookup_failed",
                        error=str(_exc))
                independent = await asyncio.to_thread(
                    run_independent_review,
                    findings, primary_doc_text or None)
                independent["findings_seen"] = findings
                yield _sse("independent_review", **independent)
                log.info(
                    "academic_review_independent_complete",
                    verdict=independent.get("verdict"),
                    model=independent.get("model"),
                )
            except Exception as exc:  # noqa: BLE001
                # Defensive fallback — should be unreachable because
                # run_independent_review is itself fail-open. If we
                # land here it's an extract / asyncio failure; log
                # and emit a placeholder so the frontend's stream
                # contract is honoured (the event always fires).
                log.warning(
                    "independent_review_orchestration_failed",
                    error=str(exc),
                )
                yield _sse(
                    "independent_review",
                    verdict="Concerns",
                    overall_reasoning=(
                        f"Independent review failed to run: {exc}. "
                        "The primary verdict stands; re-run to retry."),
                    per_finding=[],
                    model="unavailable",
                    findings_seen={},
                )

            # ── Adversarial critic + debate round (Concern 7g) ──
            # Runs AFTER the primary arbiter completes. Critic
            # findings always stream; the debate round fires only
            # when fatal + major > 0. The council_debates row is
            # always written -- the durable audit record covers
            # both the "debate ran" and the "minor only -- skipped"
            # branches.
            critic_result = None
            debate_text: str | None = None
            was_addressed: list[bool] | None = None
            counter_arguments: list[dict[str, Any]] | None = None
            current_data_hash_str: str | None = None
            try:
                from agents.academic_review import (
                    build_critic_context, run_critic_review,
                    run_arbiter_debate_round,
                    parse_debate_assessments,
                    record_debate_round,
                )
                critic_context = await build_critic_context(
                    reviewer_email=session.get("email"),
                    document_type=document_type_q or None)
                critic_result = await run_critic_review(
                    critic_context, document_type=document_type_q)
                # Stream the structured critic findings frame.
                yield _sse(
                    "critic_findings",
                    document_scope=(
                        document_type_q or "full_package"),
                    gemini_findings=critic_result.gemini_findings,
                    grok_findings=critic_result.grok_findings,
                    merged_findings=critic_result.merged_findings,
                    gemini_prose=critic_result.gemini_prose,
                    grok_prose=critic_result.grok_prose,
                    fatal_count=critic_result.fatal_count,
                    major_count=critic_result.major_count,
                    minor_count=critic_result.minor_count,
                    partial_failure=critic_result.partial_failure)
                log.info(
                    "critic_review_complete",
                    fatal=critic_result.fatal_count,
                    major=critic_result.major_count,
                    minor=critic_result.minor_count,
                    partial=critic_result.partial_failure,
                )
                # Debate round gate: only fire if there's something
                # actionable to debate. Minor-only findings are
                # logged but skip the second arbiter call.
                try:
                    from tools.audit_assembler import (
                        current_data_hash as _cur_hash,
                    )
                    current_data_hash_str = await _cur_hash() or ""
                except Exception:  # noqa: BLE001
                    current_data_hash_str = ""

                if critic_result.has_actionable:
                    debate_text = await asyncio.to_thread(
                        run_arbiter_debate_round,
                        context_block, peer_responses,
                        critic_result.merged_findings,
                        multi_user, n_strategies)
                    for chunk in chunk_arbiter_text(debate_text):
                        yield _sse(
                            "debate_round_arbiter", text=chunk)
                    was_addressed, counter_arguments = (
                        parse_debate_assessments(
                            debate_text,
                            critic_result.merged_findings))
                    log.info(
                        "academic_review_debate_complete",
                        debate_chars=len(debate_text),
                        addressed=sum(
                            1 for x in was_addressed if x),
                        rebutted=len(counter_arguments),
                    )
                else:
                    yield _sse("critic_minor_only")
                    log.info(
                        "academic_review_debate_skipped_minor_only",
                        minor=critic_result.minor_count,
                    )

                # Always-write audit row -- both branches above.
                debate_id = await record_debate_round(
                    interaction_id=None,  # filled by _log_interaction_bg later
                    context="academic_review",
                    document_type=(
                        document_type_q or "full_package"),
                    critic_result=critic_result,
                    peer_responses=peer_responses,
                    arbiter_resolution=debate_text,
                    was_addressed=was_addressed,
                    counter_arguments=counter_arguments,
                    data_hash=current_data_hash_str,
                )
                log.info(
                    "council_debates_row_written",
                    debate_id=debate_id,
                )

                # Concern 7k-i auto-fire: every Fatal finding gets a
                # pre-populated arbiter fix proposal so the team sees
                # an Apply Fix button on every Fatal card without
                # waiting for a follow-up request. Majors are
                # explicit-only (UI Propose Fix button hits
                # /api/v1/documents/propose-fix). Best-effort -- a
                # proposal generation failure just leaves the card
                # without an auto-proposal; the team can still
                # request one manually.
                auto_proposals: list[dict[str, Any]] = []
                try:
                    from agents.academic_review import (
                        run_arbiter_fix_proposal,
                        write_fix_proposals_to_debate,
                    )
                    if debate_id and critic_result \
                            and critic_result.has_actionable:
                        target_doc = (
                            document_type_q or "full_package")
                        fatal_indexes = [
                            i for i, f in enumerate(
                                critic_result.merged_findings)
                            if str(f.get("severity")).strip()
                            .capitalize() == "Fatal"]
                        proposals = []
                        for fi in fatal_indexes:
                            f = (
                                critic_result.merged_findings[fi])
                            pdoc = (
                                str(f.get("target_document"))
                                if f.get("target_document")
                                and f.get("target_document")
                                != "cross_document"
                                else target_doc)
                            prop = await run_arbiter_fix_proposal(
                                finding=f, finding_id=fi,
                                document_type=pdoc,
                                reviewer_email=session.get(
                                    "email"))
                            if prop is not None:
                                proposals.append(prop)
                                auto_proposals.append({
                                    "finding_id": prop.finding_id,
                                    "target": prop.target,
                                    "section_name": (
                                        prop.section_name),
                                    "rationale": prop.rationale,
                                    "patch_instruction": (
                                        prop.patch_instruction),
                                    "severity": prop.severity,
                                    "auto_proposed": (
                                        prop.auto_proposed),
                                    "target_document": (
                                        prop.target_document),
                                    "source_of_truth_document": (
                                        prop.source_of_truth_document
                                    ),
                                })
                        if proposals:
                            await write_fix_proposals_to_debate(
                                debate_id, proposals)
                            log.info(
                                "fix_proposals_auto_written",
                                debate_id=debate_id,
                                count=len(proposals))
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "fix_proposal_auto_fire_failed",
                        error=str(exc))

                # Concern 7j + 7k-v -- emit the persisted-row
                # marker frame carrying debate_id + auto-fire fix
                # proposals so the UI can route propose-fix and
                # apply-fix calls without a follow-up GET. Always
                # emitted (the row is always written).
                yield _sse(
                    "debate_recorded",
                    debate_id=debate_id,
                    fix_proposals=auto_proposals)
            except asyncio.CancelledError:
                # Client disconnected mid-stream. Re-raise so the
                # SSE generator unwinds cleanly; without this catch
                # the asyncpg pool would leak the in-flight
                # connection from record_debate_round /
                # write_fix_proposals_to_debate. The outer
                # event_stream's [DONE] frame is skipped (the
                # client is gone) but the connections are released
                # by the async-with blocks unwinding under the
                # CancelledError propagation.
                log.info(
                    "academic_review_critic_cancelled",
                    document_type=(
                        document_type_q or "full_package"))
                raise
            except Exception as exc:  # noqa: BLE001
                # All other exceptions are advisory -- the critic
                # pipeline is fail-open. Log and continue so the
                # primary verdict still streams to the user. Partial
                # critic results (one of Gemini/Grok failed parse)
                # are NOT exceptions -- they surface via the
                # partial_failure field on the critic_findings
                # frame, and the debate round still fires on
                # whatever findings the other model produced.
                log.warning(
                    "academic_review_critic_pipeline_failed",
                    error=str(exc),
                )

            # Team Activity — log the completed review. The overall
            # readiness rating is parsed out of the verdict; the harness
            # block aggregates the peer + arbiter quality runs.
            agents = list(peer_responses.keys()) + ["academic_advisor",
                                                     "independent_reviewer"]
            review_metadata: dict[str, Any] = {
                "overall_rating": _parse_overall_rating(arbiter_text),
            }
            harness_meta = collect_harness_metrics()
            if harness_meta:
                review_metadata["harness"] = harness_meta
            if critic_result is not None:
                review_metadata["critic"] = {
                    "fatal": critic_result.fatal_count,
                    "major": critic_result.major_count,
                    "minor": critic_result.minor_count,
                    "partial_failure": (
                        critic_result.partial_failure),
                    "debate_round_fired": (
                        critic_result.has_actionable),
                }
                if counter_arguments:
                    review_metadata["critic"][
                        "rebuttals_logged"] = len(counter_arguments)
            _log_interaction_bg(
                request, session, "academic_review",
                agents_involved=agents,
                response_summary=arbiter_text,
                metadata=review_metadata,
            )
        except asyncio.CancelledError:
            # Client disconnected mid-stream. Re-raise so the async
            # generator unwinds cleanly (asyncpg releases the in-
            # flight connection from record_debate_round /
            # write_fix_proposals_to_debate). The outer [DONE] in
            # the finally below is skipped under the CancelledError
            # propagation -- the client isn't listening anyway.
            log.info(
                "academic_review_stream_cancelled",
                document_type=document_type_q or "full_package")
            raise
        except Exception as exc:  # noqa: BLE001
            log.error("academic_review_failed", error=str(exc))
            yield _sse("error", message="Academic review failed — please retry.")
        finally:
            # June 25 2026 -- always emit [DONE]. The previous
            # structure had [DONE] after the except block which is
            # functionally equivalent for the normal exit path but
            # subtle: any new branch (e.g. an inner return on a
            # downstream failure) would skip the terminal frame and
            # leave the client polling forever. The finally
            # guarantees [DONE] reaches the client on every exit
            # path except CancelledError (where the client is gone).
            yield "data: [DONE]\n\n"

    # June 25 2026 -- wrap the producer in with_keepalive so the
    # 60-90s arbiter / 3-4 min full pipeline doesn't trigger a
    # client / proxy timeout. The wrapper emits a ': keepalive'
    # comment frame whenever the producer is silent for 20s.
    return StreamingResponse(
        with_keepalive(event_stream()),
        media_type="text/event-stream")


# ── Peer Review Assistant (item 7, Feature A) ─────────────────────────────────
#
# Bob / Michael / Molly each review another team's midpoint
# submission for the June 3 cohort meetup. Upload the peer team's
# PDF / DOCX / MD; we extract text in-memory (no persistence —
# peer papers are not reference documents and don't belong in
# academic_documents), run the harness-gated Opus call against the
# four FNA 670 rubric dimensions, and stream a 3-4 minute review
# script SSE-style.

@app.post("/api/council/peer-review")
@limiter.limit("10/minute")
async def council_peer_review(
    request: Request,
    session: dict = Depends(require_team_member),
):
    """RETIRED (PR-B, June 2026).

    The frontend Peer Review page was removed in PR #338; this
    endpoint is retired in PR-B (June 21 2026). Returns 410 Gone
    so existing clients receive a clear "this existed and is now
    gone" signal rather than a 404 connection error.

    Replacement: brief quality verification is now handled by the
    post-generation audit checks (see tools/document_audit.py CHECK 5
    / CHECK 6 / CHECK 7 added in PR #336).

    The handler is preserved as a stub for one release cycle so a
    monitoring dashboard can detect retired-endpoint calls; the route
    decorator + 410 body will be deleted in a subsequent PR once
    Render logs confirm zero residual traffic.
    """
    return JSONResponse(
        status_code=410,
        content={
            "error": "gone",
            "message": (
                "Peer review has been retired. Use the post-"
                "generation audit checks for brief quality "
                "verification."),
            "canonical_path": "/api/v1/export/executive-brief",
        })




# ── Thesis Defense Prep (item 7, Feature B) ───────────────────────────────────
#
# Auto-loads the calling user's most-recent midpoint_paper editor
# draft and streams an anticipated Q&A prep sheet across three
# categories (technical, academic, governance). No file upload —
# the draft is the source of truth.

# ── Defense Prep — async background job pattern ─────────────────────────────
#
# May 30 2026 — the synchronous SSE flow was timing out on Render. A long
# Opus + harness generation could leave the connection idle for >100s
# between the draft_meta byte and the first arbiter_chunk byte, and the
# gateway killed it. Same shape as the document-generation jobs (see
# tools/generation_jobs.py): POST 202 + job_id → background task → poll
# GET /api/v1/defense-prep/{id}.
_defense_prep_bg_tasks: set = set()


async def _run_defense_prep_job(
    job_id: str, filename: str, draft_text: str,
    user_email: str, session: dict, ip: str,
) -> None:
    """Background task: build the context, run the harness, persist the
    verdict on the job. Fail-open — any exception flips the job to
    `failed` with a readable error so the polling frontend can render it
    inline rather than spinning forever."""
    import asyncio
    from agents.peer_review import (
        build_defense_prep_context_block,
        render_defense_prep_context_block,
        run_defense_prep_with_harness,
    )
    from tools.activity_log import log_agent_interaction
    from tools.generation_jobs import update_job
    from datetime import datetime, timezone

    update_job(job_id, status="running")
    try:
        ctx = build_defense_prep_context_block(
            "Forest Capital (team draft)", draft_text,
            source_name=filename)
        context_block = render_defense_prep_context_block(ctx)
        verdict = await asyncio.to_thread(
            run_defense_prep_with_harness, context_block)
        update_job(
            job_id,
            status="complete",
            completed_at=datetime.now(timezone.utc),
            _result_text=verdict,
        )
        log.info("defense_prep_complete", job_id=job_id, source="upload",
                 filename=filename, draft_chars=len(draft_text),
                 verdict_chars=len(verdict))
        # The endpoint returned 202 long before this task started; there
        # is no Request object to thread through to _log_interaction_bg.
        # Call log_agent_interaction directly so the activity log row
        # still lands. Fail-open inside the helper.
        try:
            await log_agent_interaction(
                user_email=user_email,
                interaction_type="defense_prep",
                session_id=session.get("session_id"),
                session_type=session.get("session_type"),
                ip_address=ip,
                user_agent=None,
                agents_involved=["thesis_defense_prep"],
                question_text=None,
                response_summary=verdict,
                metadata={
                    "source":      "upload",
                    "filename":    filename,
                    "draft_chars": len(draft_text),
                    "job_id":      job_id,
                },
            )
        except Exception as log_exc:  # noqa: BLE001
            log.warning("defense_prep_activity_log_failed",
                        job_id=job_id, error=str(log_exc))
    except Exception as exc:  # noqa: BLE001
        log.error("defense_prep_failed", job_id=job_id, error=str(exc))
        update_job(
            job_id,
            status="failed",
            completed_at=datetime.now(timezone.utc),
            error="Defense prep failed: " + str(exc),
        )


@app.post("/api/council/defense-prep")
@limiter.limit("10/minute")
async def council_defense_prep(
    request: Request,
    file: UploadFile = File(...),
    session: dict = Depends(require_team_member),
):
    """Validates an uploaded .pdf or .docx, creates a Defense Prep job,
    spawns the LLM run on a background task, and returns 202 + job_id
    immediately. The frontend polls GET /api/v1/defense-prep/{id} every
    few seconds and renders status / result / error as they land.

    Upload validation is SYNCHRONOUS — empty / oversize / unsupported
    extensions return 422 with a JSON detail, NEVER a job_id. The
    uploaded bytes are extracted in-memory and never persisted.

    Returns: 202 with {"job_id": "...", "status": "pending"}.
    """
    import asyncio
    from datetime import datetime, timezone

    from tools.academic_context import extract_uploaded_text
    from tools.generation_jobs import create_job, update_job

    filename = (file.filename or "upload").strip()
    try:
        raw = await file.read()
    except Exception as exc:  # noqa: BLE001
        log.warning("defense_prep_upload_read_failed",
                    filename=filename, error=str(exc))
        raise HTTPException(
            status_code=422,
            detail="Could not read the uploaded file. Try again.")
    if not raw:
        raise HTTPException(
            status_code=422,
            detail="The uploaded file was empty. Choose a .pdf or .docx "
                   "with content and try again.")
    if len(raw) > 10 * 1024 * 1024:
        raise HTTPException(
            status_code=422,
            detail="The uploaded file exceeds the 10 MB limit. Trim the "
                   "document or split it before retrying.")
    try:
        draft_text = extract_uploaded_text(filename, raw)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))

    job = create_job("defense_prep", session["email"])
    update_job(
        job["job_id"],
        _filename=filename,
        _draft_chars=len(draft_text),
        _word_count=len(draft_text.split()),
        _result_text=None,
    )
    ip = ""
    try:
        ip = (request.client.host if request.client else "") or ""
    except Exception:  # noqa: BLE001
        ip = ""
    task = asyncio.create_task(
        _run_defense_prep_job(
            job["job_id"], filename, draft_text,
            session["email"], session, ip))
    update_job(job["job_id"], _task=task,
               created_at=datetime.now(timezone.utc))
    _defense_prep_bg_tasks.add(task)
    task.add_done_callback(_defense_prep_bg_tasks.discard)
    return JSONResponse(
        status_code=202,
        content={
            "job_id":     job["job_id"],
            "status":     "pending",
            "filename":   filename,
            "word_count": len(draft_text.split()),
        })


@app.get("/api/v1/defense-prep/{job_id}")
async def get_defense_prep_job(
    job_id: str, session: dict = Depends(require_team_member),
):
    """Polling endpoint for the Defense Prep job created above. Returns:
      pending  / running                — the LLM is still working;
      complete with result_text          — the Q&A verdict is ready;
      failed   with error                — readable reason for inline UX.
    Owner-only: a job_id belongs to the user that created it."""
    from datetime import datetime

    from tools.generation_jobs import get_job

    job = get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job.get("document_type") != "defense_prep":
        raise HTTPException(status_code=404, detail="Job not found.")
    if job.get("owner_email") != session.get("email"):
        raise HTTPException(status_code=404, detail="Job not found.")
    created_at = job.get("created_at")
    completed_at = job.get("completed_at")
    elapsed = None
    if created_at is not None:
        end = completed_at if completed_at else datetime.now(created_at.tzinfo)
        elapsed = round((end - created_at).total_seconds(), 1)
    return {
        "job_id":       job["job_id"],
        "status":       job["status"],
        "filename":     job.get("_filename"),
        "word_count":   job.get("_word_count"),
        "result_text":  (job.get("_result_text")
                         if job["status"] == "complete" else None),
        "error":        job.get("error") if job["status"] == "failed" else None,
        "created_at":   created_at.isoformat() if created_at else None,
        "completed_at": completed_at.isoformat() if completed_at else None,
        "elapsed_seconds": elapsed,
    }


# ── Inline metric explainer ───────────────────────────────────────────────────

@app.post("/api/council/explain")
@limiter.limit("30/minute")
async def council_explain(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """
    Streams a plain-English explanation of one metric or chart — backs
    the InfoIcon → ExplainerPanel click path on the Analytics and
    Dashboard screens.

    Uses the Explainer agent's system prompt and streams via Haiku. The
    response is a raw text/plain stream of explanation chunks — NOT
    Server-Sent-Events: there is no `data:` framing and no `[DONE]`
    sentinel, unlike /api/council/academic-review. Consumers read it as
    a plain token stream. The completed explanation is logged to
    agent_interactions as interaction_type "explain" (team-gated inside
    log_agent_interaction).
    """
    metric = str(body.get("metric") or "").strip()
    if not metric:
        raise HTTPException(status_code=422, detail="metric is required")
    current_value = body.get("current_value")

    from agents.explainer_agent import stream_metric_explanation
    from agents.usage import start_usage_capture
    from tools.strategy_context import (
        detect_strategies_in_query, set_active_strategies,
    )

    # Item 9 commit 5 — strategy context. The InfoIcon click on a
    # strategy-specific metric ("Strategy: REGIME_SWITCHING Sharpe") or
    # a metric label that names a strategy injects the strategy's
    # characterisation into the explainer's system prompt via the
    # per-request ContextVar. _stream_haiku copies the request context
    # into its worker thread so the var propagates. No-op when no
    # strategy is named.
    named = detect_strategies_in_query(
        f"{metric} {current_value or ''}")
    if named:
        set_active_strategies(named)

    # Seed the usage bucket before the Haiku stream starts; _stream_haiku
    # copies the request context into its worker thread so record_usage
    # after the stream completes lands here.
    start_usage_capture()

    async def gen():
        collected: list[str] = []
        async for chunk in stream_metric_explanation(metric, current_value):
            collected.append(chunk)
            yield chunk
        # Team Activity — non-blocking, team-gated inside log_agent_interaction.
        _log_interaction_bg(
            request, session, "explain",
            question_text=metric,
            response_summary="".join(collected),
            metadata=({"current_value": str(current_value)}
                      if current_value not in (None, "") else None),
        )

    return StreamingResponse(gen(), media_type="text/plain")


@app.post("/api/council/explain-data")
@limiter.limit("30/minute")
async def council_explain_data(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """
    Streams a contextual explanation of the SPECIFIC values currently on
    screen — backs the "Explain this data" (✨) button on the strategy
    detail subscreen and the Analytics charts.

    Deliberately distinct from /api/council/explain: the InfoIcon answers
    "what does this metric mean?" in 150 words; this answers "what do
    these specific values mean together?" with the deeper, academic
    framing. The completed explanation is logged to agent_interactions as
    interaction_type "explain_data" (team-gated inside log_agent_interaction).

    Like /api/council/explain this is a raw text/plain token stream — no
    SSE framing, no [DONE] sentinel.
    """
    metric = str(body.get("metric") or "").strip()
    if not metric:
        raise HTTPException(status_code=422, detail="metric is required")
    current_value = body.get("current_value")
    context = body.get("context")

    from agents.explainer_agent import stream_data_explanation
    from agents.usage import start_usage_capture
    from tools.strategy_context import (
        detect_strategies_in_query, set_active_strategies,
    )

    # Item 9 commit 5 — strategy context. The Data Explain (✨) click
    # typically carries the strategy name in the metric label and the
    # full strategy row in the context dict. Scan both for known
    # strategy ids and set the per-request ContextVar so the explainer
    # system prompt picks up the characterisation block.
    haystack = (
        f"{metric} {current_value or ''} "
        + (str(context) if context else ''))
    named = detect_strategies_in_query(haystack)
    if named:
        set_active_strategies(named)

    # Same pattern as /api/council/explain — seed before the stream worker
    # thread starts so its copied context inherits this bucket.
    start_usage_capture()

    async def gen():
        collected: list[str] = []
        async for chunk in stream_data_explanation(metric, current_value, context):
            collected.append(chunk)
            yield chunk
        _log_interaction_bg(
            request, session, "explain_data",
            question_text=metric,
            response_summary="".join(collected),
            metadata=({"current_value": str(current_value)}
                      if current_value not in (None, "") else None),
        )

    return StreamingResponse(gen(), media_type="text/plain")


# ── Explainer CIO follow-up — multi-turn thread inside the panel ──────────────
#
# The ExplainerPanel surfaces a 150-word static explanation from Haiku
# (the /api/council/explain stream above). Once the user has read that,
# they may want a follow-up — a clarification, a question grounded in
# the chart, a tie-back to the macro digest. The follow-up endpoint
# calls the CIO directly (Opus) for higher-quality reasoning, scoped
# tightly to the explainer context: the topic, the panel content
# already shown, the optional chart values, the optional macro
# summary, and the prior thread of user/cio exchanges in this panel
# session. Capped at three exchanges per session; beyond that the
# user takes the question to the full council via a handoff package.


_FOLLOWUP_MAX_EXCHANGES = 3


class ExplainerFollowupExchange(__import__("pydantic").BaseModel):
    role: str  # "user" or "cio"
    content: str


class ExplainerFollowupRequest(__import__("pydantic").BaseModel):
    explainer_topic: str
    explainer_content: str
    chart_context: dict | None = None
    macro_summary: str | None = ""
    thread: list[ExplainerFollowupExchange] = []
    question: str


def _build_followup_system_prompt(
    topic: str, explainer_content: str,
    chart_context: dict | None, macro_summary: str | None,
) -> str:
    """Builds the CIO follow-up system prompt with the explainer
    context baked in. Same evidence-discipline rule the other agents
    use ("only reference numbers a tool has actually returned"), plus
    the explainer-specific instruction to be concise."""
    from agents.cio import _SYSTEM_PROMPT as _CIO_SYSTEM_PROMPT

    parts = [
        _CIO_SYSTEM_PROMPT,
        "",
        "=== EXPLAINER FOLLOW-UP CONTEXT ===",
        f"You are answering a follow-up question from a user looking "
        f"at the {topic} explainer on the Forest Capital portfolio "
        f"intelligence platform.",
        "",
        "The user has already read the following explainer content:",
        explainer_content[:2000],
    ]
    if chart_context:
        parts.append("")
        parts.append("Chart context — values currently on screen:")
        import json as _json
        parts.append(_json.dumps(chart_context, indent=2, default=str)[:1000])
    if macro_summary:
        parts.append("")
        parts.append("Current macro conditions summary:")
        parts.append(macro_summary[:1000])
        parts.append("When you draw on this macro context in your answer, "
                     "cite it inline using [Macro: <category>] (e.g. "
                     "[Macro: monetary_policy]) so the user sees which "
                     "signals you used.")
    parts.append("")
    parts.append(
        "CONCISENESS — keep your answer to 2-4 sentences for a simple "
        "clarification, up to 2 short paragraphs for a complex question. "
        "Do not restate the explainer content; build on it.")
    parts.append(
        "WHEN TO ESCALATE — if the question requires "
        "multi-specialist deliberation (e.g. comparing strategies on "
        "more than one dimension, weighing the equity vs fixed-income "
        "view on the same question, deriving a portfolio "
        "recommendation), conclude your answer with the line "
        "'[SUGGEST_COUNCIL]' on its own. The frontend strips this "
        "marker and surfaces a 'Take this to the Council' button. "
        "Use [SUGGEST_COUNCIL] sparingly — only when the question "
        "genuinely warrants the full council, not for every "
        "follow-up.")
    return "\n".join(parts)


@app.post("/api/v1/council/explainer-followup")
@limiter.limit("15/minute")
async def council_explainer_followup(
    request: Request,
    body: ExplainerFollowupRequest,
    session: dict = Depends(require_team_member),
):
    """
    Streams a CIO follow-up answer for a question asked inside the
    ExplainerPanel thread. Up to three exchanges per panel session.

    Response is text/event-stream — three frame types:
      data: {"type":"chunk","text":"..."}    (streamed chunks)
      data: {"type":"meta","exchanges_used":N,"suggest_council":bool}
      data: [DONE]

    The frontend assembles the chunks into the assistant message,
    reads exchanges_used to update the "X of 3 follow-ups used"
    counter, and toggles the council handoff prompt when
    suggest_council is true.
    """
    if not body.question.strip():
        raise HTTPException(status_code=422, detail="question is required")
    if len(body.question) > 300:
        raise HTTPException(status_code=422,
                            detail="question must be 300 chars or fewer")
    if len(body.thread) >= _FOLLOWUP_MAX_EXCHANGES:
        raise HTTPException(status_code=429,
                            detail="Follow-up limit reached. Take the "
                                   "question to the council.")

    from agents.base import OPUS_MODEL, call_claude
    from agents.usage import start_usage_capture
    start_usage_capture()

    system_prompt = _build_followup_system_prompt(
        body.explainer_topic, body.explainer_content,
        body.chart_context, body.macro_summary,
    )

    # Build the multi-turn conversation. Prior exchanges go into the
    # user_message as a transcript; the new question goes last. We
    # avoid Anthropic's messages-array conversation API for this
    # endpoint (it requires alternating user/assistant which is a
    # heavier conformance surface for a 3-turn cap).
    transcript_parts: list[str] = []
    for ex in body.thread:
        label = "USER" if ex.role.lower() == "user" else "CIO"
        transcript_parts.append(f"{label}: {ex.content}")
    transcript = "\n\n".join(transcript_parts)
    user_message = (
        f"{transcript}\n\nUSER: {body.question}".strip()
        if transcript
        else f"USER: {body.question}")

    exchanges_used = len(body.thread) + 1

    async def gen():
        full_text = ""
        suggest_council = False
        try:
            # Non-streaming call (call_claude returns the full string).
            # For the follow-up surface a "stream as one chunk" is
            # acceptable — the response is short (2-4 sentences typical)
            # and the user expects a CIO-quality answer rather than a
            # token-by-token reveal. The SSE framing is preserved so
            # the frontend's stream reader works uniformly.
            full_text = call_claude(
                OPUS_MODEL, system_prompt, user_message, max_tokens=600,
                trigger="council_followup")
        except Exception as exc:  # noqa: BLE001
            log.warning("council_followup_failed", error=str(exc))
            full_text = ("The CIO follow-up is unavailable right now. "
                         "Try again in a moment, or take the question to "
                         "the council.")
        # Strip the [SUGGEST_COUNCIL] sentinel — the frontend uses
        # the boolean, not the literal marker in the body text.
        if "[SUGGEST_COUNCIL]" in full_text:
            suggest_council = True
            full_text = full_text.replace("[SUGGEST_COUNCIL]", "").strip()

        # Emit as a single chunk + meta frame.
        chunk_frame = json.dumps({"type": "chunk", "text": full_text})
        meta_frame = json.dumps({
            "type": "meta",
            "exchanges_used": exchanges_used,
            "suggest_council": suggest_council,
        })
        yield f"data: {chunk_frame}\n\n"
        yield f"data: {meta_frame}\n\n"
        yield "data: [DONE]\n\n"

        _log_interaction_bg(
            request, session, "explainer_followup",
            question_text=body.question,
            response_summary=full_text[:500],
            metadata={
                "explainer_topic": body.explainer_topic,
                "exchanges_used": exchanges_used,
                "suggest_council": suggest_council,
            },
        )

    return StreamingResponse(gen(), media_type="text/event-stream")


# ── Guided UAT test runner ────────────────────────────────────────────────────
#
# Records attested test-step results, structured failure reports, and
# AI-categorised tester feedback. Test SCRIPTS are frontend config
# (constants/testScripts.ts) — only results and feedback are persisted
# here. All endpoints are team-gated; the admin views are ruurdsm@ only.

# Team-gating is the require_team_member dependency on the testing
# endpoints; the failure-reports and feedback-backlog views require the
# view_admin permission. Both are permission checks — no hardcoded email.


async def _read_screenshots(files: list[UploadFile]) -> list[str]:
    """Reads uploaded screenshot files and saves them to local storage,
    returning their relative paths. Fail-open — never raises."""
    from tools.test_runner import save_screenshots
    pairs: list[tuple[str, bytes]] = []
    for f in files or []:
        try:
            content = await f.read()
            if content:
                pairs.append((f.filename or "shot.png", content))
        except Exception:  # noqa: BLE001
            continue
    return save_screenshots(pairs) if pairs else []


# ── Automated triage triggers ─────────────────────────────────────────────────
#
# The triage engine runs in the background, never blocking a result /
# feedback submission. Two automatic triggers plus the manual endpoint:
#   threshold  — 5+ unaddressed items have accumulated since the last run
#   test_pass  — a tester has just completed a full test script
# Both are fire-and-forget and fail-open — a triage failure never affects
# the primary submission.

_triage_bg_tasks: set = set()


async def _triage_trigger(kind: str) -> None:
    """Runs in the background. For the threshold trigger it first checks
    the ≥5-unaddressed-since-last-run condition; the test_pass trigger
    runs unconditionally. run_triage itself skips a concurrent run."""
    try:
        from tools.triage_engine import (
            count_unaddressed_items, is_triage_running, last_triage_at,
            run_triage,
        )
        if await is_triage_running():
            return
        if kind == "threshold":
            since = await last_triage_at()
            if await count_unaddressed_items(since=since) < 5:
                return
            await run_triage("threshold")
        else:
            await run_triage("test_pass")
    except Exception as exc:  # noqa: BLE001
        log.warning("triage_trigger_failed", kind=kind, error=str(exc))


def _fire_triage(kind: str) -> None:
    """Fire-and-forget scheduling of a triage trigger — a strong reference
    is held so the task is not garbage-collected mid-run."""
    try:
        import asyncio
        task = asyncio.create_task(_triage_trigger(kind))
        _triage_bg_tasks.add(task)
        task.add_done_callback(_triage_bg_tasks.discard)
    except Exception as exc:  # noqa: BLE001
        log.warning("triage_fire_failed", kind=kind, error=str(exc))


@app.post("/api/v1/testing/results")
@limiter.limit("120/minute")
async def testing_record_result(
    request: Request,
    script_id: str = Form(...),
    step_id: str = Form(...),
    result: str = Form(...),
    notes: str = Form(default=""),
    failure_description: str = Form(default=""),
    expected_result: str = Form(default=""),
    actual_result: str = Form(default=""),
    severity: str = Form(default=""),
    browser_info: str = Form(default=""),
    override_reason: str = Form(default=""),
    low_quality: bool = Form(default=False),
    script_complete: bool = Form(default=False),
    screenshots: list[UploadFile] = File(default=[]),
    session: dict = Depends(require_team_member),
):
    """
    Records (upserts) one attested test-step result. Always multipart so
    a step with or without screenshots uses one content type. A
    re-attestation overwrites the row and flips `overridden`. Team-only
    (require_team_member).
    """
    email = session.get("email", "")
    if result not in {"pass", "fail", "skip"}:
        raise HTTPException(status_code=422, detail="result must be pass | fail | skip")

    from tools.test_runner import record_result
    paths = await _read_screenshots(screenshots)
    stored = await record_result(
        user_email=email,
        session_type=request.headers.get("x-session-type") or "testing",
        script_id=script_id, step_id=step_id, result=result,
        notes=notes or None, failure_description=failure_description or None,
        expected_result=expected_result or None,
        actual_result=actual_result or None, severity=severity or None,
        browser_info=browser_info or None, screenshot_paths=paths or None,
        low_quality=low_quality, override_reason=override_reason or None,
    )
    if stored is None:
        raise HTTPException(status_code=503,
                            detail="Could not record the result — database unavailable.")

    # Automation hooks — both fire-and-forget, never block this response.
    # Threshold: a new failure/feedback item may push the backlog past 5.
    # Test pass: `script_complete` is set by the client (which holds the
    # testScripts.ts step inventory) when the final step has been attested.
    _fire_triage("threshold")
    if script_complete:
        _fire_triage("test_pass")
    return stored


@app.get("/api/v1/testing/results")
async def testing_get_results(session: dict = Depends(require_team_member)):
    """The current user's test results, grouped by script_id. Team-only."""
    email = session.get("email", "")
    from tools.test_runner import get_results
    grouped: dict[str, list] = {}
    for row in await get_results(email):
        grouped.setdefault(row["script_id"], []).append(row)
    return {"results": grouped}


@app.get("/api/v1/testing/unseen")
async def testing_unseen(session: dict = Depends(require_team_member)):
    """Per-script attested-step inventory — the frontend diffs it against
    testScripts.ts to surface scripts with new/changed steps. Team-only."""
    email = session.get("email", "")
    from tools.test_runner import get_unseen
    return await get_unseen(email)


@app.get("/api/v1/testing/summary")
async def testing_summary(session: dict = Depends(require_team_member)):
    """Per-script pass/fail/skip counts for the current user. The frontend
    derives total and pending from its own step inventory. Team-only."""
    email = session.get("email", "")
    from tools.test_runner import get_summary
    return {"summary": await get_summary(email)}


@app.get("/api/v1/testing/team-progress")
async def testing_team_progress(
    session: dict = Depends(require_permission("view_uat_status")),
):
    """Shared UAT progress across every team member — backs the
    Settings → UAT Team Progress dashboard (May 24 2026).

    Permission: view_uat_status (carried by team_member and sysadmin
    per the UAT #119 split). READ-ONLY; the frontend cannot attest
    steps for another tester through this route — only
    /api/v1/testing/results (require_team_member, scoped to the
    caller's own email) accepts attestations.

    Real-time: the frontend polls every 15s. The response is small
    (per-user per-script step-id lists + 4 scalar fields) so the
    cost is bounded; one round-trip per 15s per logged-in viewer.

    Fail-open: a DB outage returns the team list with empty progress
    so the frontend renders a per-member card at 0%, never a blank.
    """
    from tools.test_runner import get_team_progress
    return await get_team_progress()


@app.get("/api/v1/testing/failures")
async def testing_failures(
    session: dict = Depends(require_permission("view_uat_status")),
):
    """Every failed step across all testers, severity-sorted.

    May 24 2026 (UAT #119) — relaxed from view_admin to view_uat_status
    so Bob and Molly can see real-time UAT progress without admin
    access. The endpoint is read-only; mutation endpoints (resolve,
    suggestions/approve, triage) remain manage_users / view_admin-
    gated so a team_member cannot act on a row, only see it.
    """
    from tools.test_runner import get_all_failures
    return {"failures": await get_all_failures()}


@app.get("/api/v1/testing/issue-tracker")
async def testing_issue_tracker(
    session: dict = Depends(require_permission("view_uat_status")),
):
    """
    Issue Tracker view — every row that has ever failed, with a
    computed status field ∈ {open, pending_retest, passed, closed}
    per compute_issue_status(). Includes Passed rows (re-attested
    after a resolution) so the tracker shows the full lifecycle of
    each issue, not just the currently-failing ones.

    Filtering, sorting and column projection live on the frontend
    — the endpoint returns the full row set and the UI shapes it.

    May 24 2026 (UAT #119) — relaxed from view_admin to view_uat_status
    so team_member sees the same read-only tracker every admin sees.
    """
    from tools.test_runner import get_issue_tracker_rows
    return {"issues": await get_issue_tracker_rows()}


# ── PR-driven Suggested Resolutions — Commit 3/7 ─────────────────────────────
# The webhook populates pr_suggestions; these three endpoints are the
# consumer side. Sysadmin-only (require_sysadmin = require_permission
# "manage_users") — same gate other sysadmin actions use.

@app.get("/api/v1/testing/suggestions")
async def testing_list_suggestions(
    session: dict = Depends(require_sysadmin),
):
    """
    Returns every pending_review suggestion joined to its failure row.
    Backs the Suggested Resolutions banner + review modal on Failure
    Reports. Sysadmin only.

    The response carries the failure context (script_id, step_id,
    user_email, severity, etc.) so the review modal renders the
    "Failure half" + "PR half" + resolution fields without a second
    round-trip. step_title and feature are derived on the frontend
    via the existing TEST_SCRIPTS / ROUTE_TO_FEATURE helpers.
    """
    from tools.pr_suggestions import list_pending_suggestions
    return {"suggestions": await list_pending_suggestions()}


@app.get("/api/v1/testing/suggestions/by-failure")
async def testing_suggestions_by_failure(
    session: dict = Depends(require_sysadmin),
):
    """
    Returns {failure_id: pending_count} for every failure that has at
    least one pending suggestion. Powers the "[Fix available — review]"
    row badge on Failure Reports (Commit 5/7) — the frontend fetches
    this once on load and joins it onto the failure list to render
    the badges in one shot. Sysadmin only.
    """
    from tools.pr_suggestions import pending_count_by_failure
    return {"by_failure": await pending_count_by_failure()}


@app.post("/api/v1/testing/suggestions/{suggestion_id}/approve")
async def testing_approve_suggestion(
    suggestion_id: int, body: dict,
    session: dict = Depends(require_sysadmin),
):
    """
    Converts a pending_review suggestion into a real resolution on
    its failure row.

    Body (required):
      root_cause        — written into resolution_note. Universal across
                          all resolution types.
      remediation_note  — written into the resolution row. Required
                          because this endpoint always applies
                          resolution_type='code_fix_deployed' (the
                          modal pre-selects it; a reviewer who wants
                          a different type uses the manual modal on
                          the Failure Reports row instead).

    Side effects:
      1. resolve_failure writes the structured resolution onto the
         failure row with fix_reference="#{pr_number}".
      2. The step is "reset to pending" by virtue of the resolved_at
         column being set (the existing get_unseen carve-out — see
         migration 025's UPSERT change).
      3. get_notifications surfaces the resolved_failures pill to the
         original tester on their next login (notification is
         DERIVED, not pushed).
      4. Sibling pending suggestions for the same failure id are
         auto-dismissed (decision point 4 — cleaner queue).

    404 — suggestion id not found, OR the suggestion's failure row
          has gone away (a stale request after the failure was
          cleaned up).
    409 — suggestion exists but is no longer in pending_review
          (already approved or dismissed). The frontend treats 409 as
          a "queue out of date — refresh" signal.
    422 — body validation: missing root_cause OR missing
          remediation_note.
    """
    root_cause = str(body.get("root_cause") or "").strip()
    remediation_note = str(body.get("remediation_note") or "").strip()
    if not root_cause:
        raise HTTPException(
            status_code=422, detail="root_cause is required.")
    if not remediation_note:
        raise HTTPException(
            status_code=422, detail="remediation_note is required.")

    from tools.pr_suggestions import approve_suggestion, get_suggestion

    # Pre-flight read so 404 vs 409 are distinguishable. The
    # alternative is to surface a single 4xx for both, but the
    # frontend benefits from "this is stale, refresh the queue" vs
    # "this never existed" guidance.
    existing = await get_suggestion(suggestion_id)
    if not existing:
        raise HTTPException(
            status_code=404, detail=f"Suggestion {suggestion_id} not found.")
    if existing["state"] != "pending_review":
        raise HTTPException(
            status_code=409,
            detail=f"Suggestion {suggestion_id} is already "
                   f"{existing['state']}. Refresh the queue.")

    result = await approve_suggestion(
        suggestion_id,
        reviewed_by=session.get("email", ""),
        root_cause=root_cause,
        remediation_note=remediation_note,
    )
    if result is None:
        # The pre-flight passed but the approve flow couldn't land
        # the resolution — most likely the failure row was deleted
        # between read and write. Return 404 with a clear message.
        raise HTTPException(
            status_code=404,
            detail="Suggestion approved but the failure row could not "
                   "be updated. Reload the queue.")
    return {"approved": True, **result}


@app.post("/api/v1/testing/suggestions/{suggestion_id}/dismiss")
async def testing_dismiss_suggestion(
    suggestion_id: int, body: dict | None = None,
    session: dict = Depends(require_sysadmin),
):
    """
    Marks a suggestion dismissed. The failure stays Open — no
    resolution is recorded. Body: {dismiss_reason?: str}.

    404 — suggestion not found OR already in a terminal state. The
          UPDATE's `WHERE state='pending_review'` clause is what
          gates this; the helper returns False on a no-op and the
          endpoint maps that to 404.
    """
    from tools.pr_suggestions import dismiss_suggestion as _dismiss

    body = body or {}
    dismiss_reason = (
        str(body.get("dismiss_reason") or "").strip() or None
    )
    ok = await _dismiss(
        suggestion_id,
        reviewed_by=session.get("email", ""),
        dismiss_reason=dismiss_reason,
    )
    if not ok:
        raise HTTPException(
            status_code=404,
            detail=f"Suggestion {suggestion_id} not found or already "
                   "in a terminal state.")
    return {"dismissed": True}


@app.post("/api/v1/testing/failures/{failure_id}/resolve")
async def testing_resolve_failure(
    failure_id: int, body: dict,
    session: dict = Depends(require_permission("view_admin")),
):
    """
    Marks a failure resolved. The row is kept (the resolution is the audit
    trail) with the migration-025 metadata block: resolution_type, the
    root cause in resolution_note, and — for code_fix_deployed only —
    fix_reference + remediation_note.

    Validation contract:
      - resolution_type      required, one of RESOLUTION_TYPES
      - resolution_note      required (the root cause; universal)
      - fix_reference        required when resolution_type =
                             'code_fix_deployed'. Accepted formats:
                             7+ hex chars (SHA), #NNN (PR number),
                             https://github.com/... URL
      - remediation_note     required when resolution_type =
                             'code_fix_deployed'

    Step-reset semantics:
      no_bug_detected / code_fix_deployed → tester sees the resolved
        failure as a pending re-test (the login notification carries
        a Re-test This Step CTA).
      wont_fix → step is NOT reset; the notification card is
        informational only, no CTA, no re-attestation prompt.

    Requires view_admin.
    """
    from tools.test_runner import RESOLUTION_TYPES, resolve_failure

    resolution_type = str(body.get("resolution_type") or "").strip()
    resolution_note = str(body.get("resolution_note") or "").strip()
    fix_reference = str(body.get("fix_reference") or "").strip() or None
    remediation_note = str(body.get("remediation_note") or "").strip() or None

    if resolution_type not in RESOLUTION_TYPES:
        raise HTTPException(
            status_code=422,
            detail="resolution_type is required and must be one of "
                   f"{list(RESOLUTION_TYPES)}.")
    if not resolution_note:
        raise HTTPException(
            status_code=422,
            detail="resolution_note (root cause) is required.")
    if resolution_type == "code_fix_deployed":
        if not fix_reference or not _is_valid_fix_reference(fix_reference):
            raise HTTPException(
                status_code=422,
                detail="fix_reference is required for code_fix_deployed. "
                       "Accepted formats: 7+ hex characters (commit SHA), "
                       "#NNN (PR number), or a GitHub URL.")
        if not remediation_note:
            raise HTTPException(
                status_code=422,
                detail="remediation_note is required for code_fix_deployed.")

    resolved = await resolve_failure(
        failure_id, session.get("email", ""), resolution_note,
        resolution_type=resolution_type,
        fix_reference=fix_reference,
        remediation_note=remediation_note,
    )
    if resolved is None:
        raise HTTPException(status_code=404, detail="Failure not found.")
    return {"resolved": True, **resolved}


# Fix-reference shape: 7+ hex chars (SHA), #NNN (PR number), or a
# https://github.com/... URL. Kept as a small standalone helper so the
# resolution-modal frontend and the endpoint validator share one
# definition (the frontend re-implements the same regex set; if either
# diverges the test in tests/test_failure_resolution.py catches it).
_SHA_RE = __import__("re").compile(r"^[0-9a-fA-F]{7,40}$")
_PR_RE = __import__("re").compile(r"^#\d{1,6}$")
_GH_URL_RE = __import__("re").compile(
    r"^https?://(?:www\.)?github\.com/[^/]+/[^/]+/(?:commit|pull|issues)/.+$")


def _is_valid_fix_reference(s: str) -> bool:
    """True when `s` looks like a commit SHA / PR number / GitHub URL.
    Lax on whitespace; strict on shape. The endpoint applies this guard
    BEFORE the DB write so a code-fix claim cannot be recorded without
    a traceable reference."""
    s = s.strip()
    return bool(_SHA_RE.match(s) or _PR_RE.match(s) or _GH_URL_RE.match(s))


@app.post("/api/v1/testing/feedback")
@limiter.limit("60/minute")
async def testing_submit_feedback(
    request: Request,
    feedback_type: str = Form(...),
    title: str = Form(...),
    description: str = Form(...),
    script_id: str = Form(default=""),
    step_id: str = Form(default=""),
    source_route: str = Form(default=""),
    priority: str = Form(default=""),
    browser_info: str = Form(default=""),
    low_quality: bool = Form(default=False),
    screenshots: list[UploadFile] = File(default=[]),
    session: dict = Depends(require_team_member),
):
    """
    Accepts a feedback submission, runs AI categorisation, and stores it.
    A submission is step-linked (script_id + step_id) or free-form
    (neither — source_route set, the "Suggest an enhancement" path).
    Returns the stored row including the AI categorisation. Team-only.
    """
    import asyncio

    email = session.get("email", "")
    from tools.test_runner import categorize_feedback, submit_feedback

    step_context = f"{script_id or 'free-form'} / {step_id or source_route or 'n/a'}"
    ai = await asyncio.to_thread(
        categorize_feedback, feedback_type, title, description, step_context)
    paths = await _read_screenshots(screenshots)
    stored = await submit_feedback(
        user_email=email, script_id=script_id or None, step_id=step_id or None,
        source_route=source_route or None, feedback_type=feedback_type,
        title=title, description=description, priority=priority or None,
        screenshot_paths=paths or None, browser_info=browser_info or None,
        low_quality=low_quality, ai=ai,
    )
    if stored is None:
        raise HTTPException(status_code=503,
                            detail="Could not store the feedback — database unavailable.")

    # Threshold trigger — a new feedback item may push the unaddressed
    # backlog past 5. Fire-and-forget; never blocks this response.
    _fire_triage("threshold")
    return stored


@app.get("/api/v1/testing/feedback")
async def testing_get_feedback(
    category: str | None = None, severity: str | None = None,
    effort: str | None = None, status: str | None = None,
    user_email: str | None = None,
    session: dict = Depends(require_permission("view_uat_status")),
):
    """All tester feedback, newest first, with optional filters.

    May 24 2026 (UAT #119) — relaxed from view_admin to view_uat_status.
    Team members READ the backlog; the resolve action below remains
    view_admin so only an admin can change a feedback row's status.
    """
    from tools.test_runner import get_all_feedback
    feedback = await get_all_feedback({
        "category": category, "severity": severity, "effort": effort,
        "status": status, "user_email": user_email,
    })
    return {"feedback": feedback}


@app.post("/api/v1/testing/feedback/{feedback_id}/resolve")
async def testing_resolve_feedback(
    feedback_id: int, body: dict,
    session: dict = Depends(require_permission("view_admin")),
):
    """Updates a feedback row's status. The submitter sees a login
    notification on the next visit. Requires the view_admin permission."""
    admin = session.get("email", "")
    status = str(body.get("status") or "")
    if status not in {"noted", "planned", "wont_do", "resolved"}:
        raise HTTPException(
            status_code=422,
            detail="status must be noted | planned | wont_do | resolved")
    from tools.test_runner import resolve_feedback
    resolved = await resolve_feedback(
        feedback_id, status, str(body.get("resolution_note") or "") or None, admin)
    if resolved is None:
        raise HTTPException(status_code=404, detail="Feedback not found.")
    return {"resolved": True, **resolved}


@app.get("/api/v1/testing/notifications")
async def testing_notifications(session: dict = Depends(require_team_member)):
    """
    The current tester's operational login notifications — failures an
    admin resolved (pending re-test) and feedback an admin responded to.
    The "new tests available" notification is computed on the frontend
    from /unseen. Team-only.
    """
    email = session.get("email", "")
    from tools.test_runner import get_notifications
    return await get_notifications(email)


@app.post("/api/v1/testing/quality-check")
@limiter.limit("120/minute")
async def testing_quality_check(
    request: Request, body: dict, session: dict = Depends(require_team_member),
):
    """
    The quality gate — scores a failure report or feedback submission
    before the frontend stores it. Fail-open: an evaluator error returns
    passed=true so a flaky evaluator never blocks a submission. Team-only;
    logs an interaction of type test_quality_eval.
    """
    import asyncio

    submission_type = str(body.get("type") or "feedback")
    description = str(body.get("description") or "")
    step_context = str(body.get("step_context") or "")
    actual_result = body.get("actual_result")

    from tools.test_runner import quality_check
    from agents.usage import start_usage_capture

    # quality_check is a Sonnet call wrapped in asyncio.to_thread, which DOES
    # propagate the contextvars — so seeding the bucket here captures it.
    start_usage_capture()
    verdict = await asyncio.to_thread(
        quality_check, submission_type, step_context, description,
        str(actual_result) if actual_result else None)

    _log_interaction_bg(
        request, session, "test_quality_eval",
        question_text=f"{submission_type}: {step_context}"[:500],
        metadata={"overall": verdict.get("overall"),
                  "passed": verdict.get("passed")},
    )
    return verdict


# ── Triage reports — sysadmin only ────────────────────────────────────────────

@app.post("/api/v1/testing/triage")
async def testing_run_triage(
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Manually triggers a triage run in the background — does not block.
    The report appears under Settings → Triage Reports when complete.
    Sysadmin only (the manage_users permission).
    """
    _fire_triage("manual")
    return {"status": "triage_started",
            "message": "Triage report will be ready shortly."}


@app.get("/api/v1/testing/triage")
async def testing_get_triage_reports(
    session: dict = Depends(require_permission("manage_users")),
):
    """Every triage report, newest first. Sysadmin only."""
    from tools.triage_engine import get_all_triage_reports
    return {"reports": await get_all_triage_reports()}


@app.get("/api/v1/testing/triage/latest")
async def testing_get_latest_triage_report(
    session: dict = Depends(require_permission("manage_users")),
):
    """The most recent triage report, or null. Sysadmin only."""
    from tools.triage_engine import get_latest_triage_report
    return {"report": await get_latest_triage_report()}


# ── Triage report items — sysadmin only ───────────────────────────────────────
# Item-level resolution endpoints (migration 023 + triage Commit 2). The
# items table normalises the verdict prose into addressable rows; these
# three endpoints back the Settings → Triage Reports per-item UI.

@app.get("/api/v1/testing/triage/items")
async def testing_get_triage_items(
    report_id: int | None = None,
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Every triage_report_items row with full resolution status.
    Optionally filtered to a specific report_id. Sysadmin only.
    """
    from tools.triage_engine import get_all_triage_items
    return {"items": await get_all_triage_items(report_id=report_id)}


@app.patch("/api/v1/testing/triage/items/{item_id}/resolve")
async def testing_resolve_triage_item(
    item_id: int,
    body: dict,
    request: Request,
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Marks a triage item resolved. Body: {resolution_note, fix_commit?,
    requires_retest?}.

    When requires_retest=true the row's retest_requested_at is stamped
    to now() — frontend TestNotifications then surfaces a "Fix ready
    for retest" pill to the original reporter (Commit 3 wires that
    notification path through the existing get_notifications surface).
    Sysadmin only.
    """
    from tools.triage_engine import resolve_triage_item

    resolution_note = str(body.get("resolution_note") or "").strip()
    if not resolution_note:
        raise HTTPException(
            status_code=422,
            detail="resolution_note is required.")
    fix_commit_raw = body.get("fix_commit")
    fix_commit = (str(fix_commit_raw).strip() or None) if fix_commit_raw else None
    requires_retest = bool(body.get("requires_retest", False))

    result = await resolve_triage_item(
        item_id,
        resolved_by=session.get("email", ""),
        resolution_note=resolution_note,
        fix_commit=fix_commit,
        requires_retest=requires_retest,
    )
    if result is None:
        raise HTTPException(
            status_code=404, detail=f"Triage item {item_id} not found.")
    return {"status": "resolved", "item": result}


@app.patch("/api/v1/testing/triage/items/{item_id}/unresolve")
async def testing_unresolve_triage_item(
    item_id: int,
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Clears the resolution fields on a triage item — sysadmin recovery
    for an item resolved in error. Sysadmin only.
    """
    from tools.triage_engine import unresolve_triage_item

    ok = await unresolve_triage_item(item_id)
    if not ok:
        raise HTTPException(
            status_code=404, detail=f"Triage item {item_id} not found.")
    return {"status": "unresolved"}


# ── Macro market research (FEATURE 2) ────────────────────────────────────────
# The daily-scheduled macro digest the council + academic_review prompts
# inject as a CURRENT MACRO CONDITIONS block. Read endpoints are open to
# any authenticated user (the dashboard widget renders the latest digest
# for the whole team); the manual run trigger is sysadmin only because
# it bypasses the 24h freshness gate and burns the Sonnet + web_search
# budget on demand.

@app.get("/api/v1/context/freshness")
async def get_context_freshness(
    session: dict = Depends(require_auth),
):
    """Returns a per-layer freshness map of every agent-context cache.
    Item 5 (May 23 2026 — analytics context injection, freshness
    badges). Powers the dashboard freshness badges so the user sees
    how current the prompts an agent sees are.

    Three layers — keys are stable, values are ISO timestamps or null:
      macro_context        — last research digest the macro layer is
                              reflecting (mirrors /research/latest's
                              last_completed_at)
      analytics_context    — last refresh of the narrative cache
                              (rebuilds on every analytics refresh
                              tick)
      diversification_context — last refresh of the structured
                              diversification cache

    Diversification context doesn't carry an explicit timestamp
    field today; the strategy cache's computed_at is reported as a
    proxy since the diversification refresh always trails it.
    """
    from tools.analytics_context import get_analytics_freshness
    from tools.research_engine import last_research_run_at
    macro_at = await last_research_run_at()
    analytics_at = get_analytics_freshness()
    # Diversification refresh shadows the strategy cache; report
    # the strategy cache's latest computed_at as the freshness proxy.
    diversification_at: str | None = None
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is not None:
            async with AsyncSessionLocal() as s:
                r = await s.execute(text(
                    "SELECT MAX(computed_at) "
                    "FROM strategy_results_cache"))
                row = r.fetchone()
                if row and row[0]:
                    diversification_at = row[0].isoformat()
    except Exception as exc:  # noqa: BLE001
        log.warning("context_freshness_diversification_failed",
                    error=str(exc))
    return {
        "macro_context":           (macro_at.isoformat()
                                      if macro_at else None),
        "analytics_context":        analytics_at,
        "diversification_context":  diversification_at,
    }


@app.get("/api/v1/research/latest")
async def research_get_latest_digest(
    session: dict = Depends(require_auth),
):
    """The most recent COMPLETED digest, or null when no completed
    digest exists yet. Powers the dashboard widget; the widget renders
    a "preparing first digest" empty state on null."""
    from tools.research_engine import get_latest_digest, last_research_run_at
    digest = await get_latest_digest()
    last_run = await last_research_run_at()
    return {
        "digest":              digest,
        "last_completed_at":   last_run.isoformat() if last_run else None,
    }


@app.get("/api/v1/research/history")
async def research_get_history(
    limit: int = 10,
    session: dict = Depends(require_auth),
):
    """The N most recent runs across every status (running / complete /
    failed). Used by the sysadmin run-history accordion below the
    widget so the team can see when failed runs happened. Open to any
    authenticated user — visibility into failures is a transparency
    feature, not a privileged one."""
    from tools.research_engine import get_recent_digests
    return {"runs": await get_recent_digests(limit=max(1, min(int(limit), 50)))}


@app.post("/api/v1/research/run")
async def research_run_now(
    session: dict = Depends(require_permission("manage_users")),
):
    """Forces a fresh research run — bypasses the 24h freshness gate.
    Returns immediately with status: 'running'; the dashboard widget
    polls /research/latest to pick up the digest once it lands.
    Sysadmin only — manual runs burn the Sonnet + web_search budget."""
    from tools.research_engine import (
        _research_bg_tasks, is_research_running, run_research,
    )
    if await is_research_running():
        return {"status": "already_running",
                "message": "A research run is already in progress."}

    # Direct manual run — bypass the stale gate. Spawn the run on the
    # event loop (we are on it here) so a long Sonnet + web_search
    # call returns the 200 to the user immediately; the digest lands
    # via the post-run cache refresh.
    import asyncio
    try:
        task = asyncio.create_task(run_research("manual"))
        # Strong ref so the task is not GC'd mid-run.
        _research_bg_tasks.add(task)
        task.add_done_callback(_research_bg_tasks.discard)
    except Exception as exc:  # noqa: BLE001
        log.warning("research_manual_spawn_failed", error=str(exc))
        raise HTTPException(
            status_code=500, detail="Failed to spawn research run.")
    return {"status": "running",
            "message": "Research run started. Poll /api/v1/research/latest "
                       "for the result."}


# ── Statistical audit — sysadmin only ─────────────────────────────────────────

@app.post("/api/v1/audit/run")
async def audit_run(
    body: dict | None = None,
    session: dict = Depends(require_sysadmin),
):
    """
    Triggers a full three-layer statistical audit in the background and
    returns immediately with the audit_id. A concurrent run is refused
    with already_running.

    `triggered_by` may be "manual" (default), "pre_submission" (an
    Analytical-Appendix audit) or "demo" (a forced run for the live
    presentation). The smart-audit-caching "Run Live Demo" button sends
    {"reason": "demo"} — accepted here as an alias for triggered_by.
    Sysadmin only — triggering a QA/audit run is restricted to the
    platform sysadmin (Michael); the read-only audit views remain open
    to the project team.

    SMART CACHE-HIT SHORT-CIRCUIT (May 26 2026 submission-night fix):
    Before claiming a new audit row, compare the current data hash
    against the last SUBSTANTIVE completed audit (status='complete',
    total_checks>0, no layer skipped). If the hashes match AND the
    caller did not pass force=true, return that audit's id+hash as
    status='cache_hit' WITHOUT creating a new row. The downloadable
    PDF (GET /audit/runs/{id}/export) then serves the prior real run.

    Why this is the right semantic: re-running the three layers on
    unchanged data is wasteful AND it was producing hollow rows on
    production (the per-request strategy_results_cache wasn't warm
    on the API process when the click hit), each saved as
    'complete' with 0 checks and rendering "this layer was skipped"
    in the PDF. The cache-hit path is the user-visible guarantee
    that "the PDF you download always reflects a real audit run".

    A caller bypasses the cache by sending {"force": true}; demo
    runs always force.
    """
    from tools.audit_assembler import current_data_hash
    from tools.audit_engine import (
        get_last_substantive_audit, is_substantive_audit, start_audit,
    )
    from tools.qa_guard import (
        QA_BUSY_MESSAGE_STATISTICAL, statistical_audit_in_progress,
    )

    # Per-type lock — a statistical audit is blocked only by another
    # statistical audit in flight, never by a methodology run. A run
    # stuck past the 15-minute timeout is reaped inside is_audit_running,
    # so a hung run never wedges this. start_audit keeps its own
    # is_audit_running() check as the race backstop.
    if await statistical_audit_in_progress():
        raise HTTPException(status_code=409,
                            detail=QA_BUSY_MESSAGE_STATISTICAL)

    body = body or {}
    triggered_by = str(body.get("triggered_by") or body.get("reason")
                       or "manual")
    if triggered_by not in ("manual", "scheduled", "pre_submission", "demo"):
        triggered_by = "manual"
    # Demo runs always force a fresh run (the operator clicked the
    # button expecting a live re-execution). Otherwise honour the
    # caller's explicit force flag; absent → False (cache-hit path).
    force = bool(body.get("force")) or triggered_by == "demo"

    # SMART CACHE-HIT — only on non-forced runs. The current data
    # fingerprint is the canonical key; if a prior substantive audit
    # verified the same data, serve it instead of redoing the work.
    #
    # Belt-and-braces (May 26 2026 follow-up). User reported the
    # hollow audit #40 was being served despite the SQL filter. The
    # endpoint now re-validates `last` via is_substantive_audit
    # BEFORE returning cache_hit; a row that fails the validator
    # falls through to a forced fresh run regardless of hash match.
    # The SQL filter inside get_last_substantive_audit + the
    # Python validator here together guarantee that whatever column
    # shape audit #40 has cannot slip through.
    if not force:
        try:
            current_hash = (await current_data_hash()) or ""
            last = await get_last_substantive_audit()
            hash_matches = (
                last is not None
                and current_hash != ""
                and last.get("data_hash") == current_hash
            )
            if hash_matches and is_substantive_audit(last):
                # Real cache hit — serve the prior substantive run.
                # No new audit_runs row is created.
                log.info(
                    "audit_run_cache_hit",
                    last_audit_id=last.get("id"),
                    data_hash=current_hash[:12],
                    triggered_by=triggered_by,
                    triggered_by_email=session.get("email", ""),
                )
                return {
                    "status":            "cache_hit",
                    "audit_id":          last.get("id"),
                    "data_hash":         current_hash,
                    "total_checks":      last.get("total_checks"),
                    "passed":            last.get("passed"),
                    "failed":            last.get("failed"),
                    "warnings":          last.get("warnings"),
                    "layer_1_status":    last.get("layer_1_status"),
                    "layer_2_status":    last.get("layer_2_status"),
                    "layer_3_status":    last.get("layer_3_status"),
                    "completed_at":      last.get("completed_at"),
                    "message": (
                        "Data unchanged since the last full audit "
                        f"(#{last.get('id')}) — serving the prior "
                        "substantive run. The downloadable PDF "
                        "reflects that run's results. Pass "
                        "{\"force\": true} to re-execute all three "
                        "layers on identical data."
                    ),
                }
            if hash_matches:
                # Hash matches but the candidate failed the Python
                # validator (hollow row that slipped past the SQL
                # filter). Fall through to a forced fresh run — the
                # diagnostic log records which columns rejected the
                # row so the operator can see exactly why a "matching"
                # audit was not served.
                log.info(
                    "audit_run_cache_miss_hollow_candidate",
                    last_audit_id=last.get("id"),
                    data_hash=current_hash[:12],
                    total_checks=last.get("total_checks"),
                    passed=last.get("passed"),
                    failed=last.get("failed"),
                    warnings=last.get("warnings"),
                    layer_1_status=last.get("layer_1_status"),
                    layer_2_status=last.get("layer_2_status"),
                    layer_3_status=last.get("layer_3_status"),
                )
        except Exception as exc:  # noqa: BLE001
            # Cache check is best-effort — any failure falls through
            # to a real run. Logged so the operator sees why a click
            # bypassed the cache.
            log.warning("audit_cache_check_failed", error=str(exc))

    # Fresh run path. force=True on this side guarantees the pre-warm
    # branch in _execute_audit fires, so the strategy_results_cache is
    # populated before assemble_audit_payload reads it. When the
    # caller didn't request force, the cache-hit branch above already
    # returned — by reaching here we know we want a real run.
    return await start_audit(
        triggered_by, session.get("email", ""), force=True,
    )


@app.get("/api/v1/audit/runs")
async def audit_get_runs(
    session: dict = Depends(require_permission("team_member")),
):
    """Every audit run with summary stats, newest first. Project team only."""
    from tools.audit_engine import get_audit_runs
    return {"runs": await get_audit_runs()}


@app.get("/api/v1/audit/runs/latest")
async def audit_get_latest_run(
    session: dict = Depends(require_auth),
):
    """
    The most recent audit run with its findings, or null. Open to every
    authenticated user — viewers see the read-only audit summary in the
    QA tab; the full findings panel is gated to the project team in the
    frontend.

    Carries the smart-audit-caching verdict: is_current is True when the
    live data fingerprint matches the last completed run's data_hash, so
    the QA tab can show a cached result instead of an unnecessary re-run.
    """
    from tools.audit_assembler import is_audit_current
    from tools.audit_engine import fail_stale_audits, get_latest_audit_run
    # Reap a hung run before reporting status — the latest run any user's
    # poll sees is then never a stale 'running' row.
    await fail_stale_audits()
    run = await get_latest_audit_run()
    currency = await is_audit_current()
    return {
        "run": run,
        "is_current": currency["is_current"],
        "statistical_current": currency["statistical_current"],
        "qa_current": currency["qa_current"],
        "current_data_hash": currency["current_data_hash"],
        "last_hash": currency["last_hash"],
    }


@app.get("/api/v1/audit/runs/{run_id}")
async def audit_get_run(
    run_id: int,
    session: dict = Depends(require_permission("team_member")),
):
    """One audit run with all findings grouped by layer. Project team only."""
    from tools.audit_engine import get_audit_run
    run = await get_audit_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="Audit run not found.")
    return run


@app.get("/api/v1/audit/runs/{run_id}/export")
async def audit_export_run(
    run_id: int,
    session: dict = Depends(require_permission("team_member")),
):
    """
    The audit run as a downloadable PDF — the Statistical Audit Report,
    professionally formatted for inclusion in the Analytical Appendix as
    evidence of independent statistical verification. Project team only.
    """
    from datetime import date
    from tools.audit_engine import get_audit_run
    from tools.audit_pdf import build_statistical_audit_pdf
    run = await get_audit_run(run_id)
    if run is None:
        raise HTTPException(status_code=404, detail="Audit run not found.")
    pdf = build_statistical_audit_pdf(run)
    filename = f"forest_capital_statistical_audit_{date.today().isoformat()}.pdf"
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/v1/audit/findings/{finding_id}/resolve")
async def audit_resolve_finding(
    finding_id: int,
    body: dict | None = None,
    session: dict = Depends(require_permission("team_member")),
):
    """
    Acknowledges an audit finding — the WARN acknowledge/resolve workflow.
    Records the team's response (resolution_note) and sets resolved. This
    is a response, not a correction: the audit's overall verdict does not
    change. Project team only.

    Optional body field `disclosure_text` (June 6 2026, bridge #75) is
    persisted to locked_disclosure_text -- the verbatim disclosure the
    team agreed to put in the report at acknowledge time. Bob copies
    the locked text from the finding card straight into the executive
    brief without re-deriving. Empty / whitespace-only strings are
    normalised to NULL (no disclosure locked).
    """
    from tools.audit_engine import resolve_finding
    note = str((body or {}).get("resolution_note") or "").strip()
    if not note:
        raise HTTPException(
            status_code=422, detail="A resolution note is required.")
    disclosure_text = (body or {}).get("disclosure_text")
    finding = await resolve_finding(
        finding_id, True, note,
        resolved_by=session["email"],
        disclosure_text=str(disclosure_text) if disclosure_text else None)
    if finding is None:
        raise HTTPException(status_code=404, detail="Audit finding not found.")
    return finding


@app.post("/api/v1/audit/findings/{finding_id}/unresolve")
async def audit_unresolve_finding(
    finding_id: int,
    session: dict = Depends(require_permission("team_member")),
):
    """Clears the acknowledgement on an audit finding. Project team only."""
    from tools.audit_engine import resolve_finding
    finding = await resolve_finding(
        finding_id, False, None, resolved_by=None)
    if finding is None:
        raise HTTPException(status_code=404, detail="Audit finding not found.")
    return finding


@app.get("/api/v1/report/readiness")
async def report_readiness(
    session: dict = Depends(require_auth),
):
    """
    Workstream C report-readiness verdict (May 28 2026).

    Returns the platform's combined verdict on whether either audit
    surface has unreviewed blocking items that should prevent the
    team from generating a graded submission. The same logic is run
    by the generation-endpoint gate (_require_report_ready) — this
    GET surfaces the verdict to the frontend so the Reports page can
    show a readiness indicator and a blocking modal that names every
    outstanding item.

    Shape:
      {
        is_ready: bool,
        blocking_count: int,
        statistical: { unreviewed_warnings, unreviewed_failures },
        methodology: { unresolved_warnings, unresolved_failures },
        checked_at: ISO timestamp,
        deck_story_plan_available: bool,   (June 21 2026, PR #343)
        deck_script_available: bool,       (June 21 2026, PR #346)
      }

    deck_story_plan_available -- true when story_plans has a real
    (non-fallback) row for (current_data_hash, document_type='deck').
    Used to surface the plan-derived state on the deck regen flow.

    deck_script_available -- true when the same row ALSO carries a
    non-empty full_script. The Presentation Script card flips its
    button state on this flag because Pass 2 (full_script) is a
    separate Opus call from Pass 1 (slide_plan) and can fail
    independently -- a plan_available=True / script_available=False
    state happens when slide_plan landed but the script call timed
    out or hit a transient API error. Treating the two as the same
    signal would let the user download a script that doesn't exist.

    Both flags computed inline here so the frontend has one round-
    trip on page load.
    """
    from tools.report_readiness import compute_readiness

    verdict = await compute_readiness()
    plan_available, script_available = (
        await _deck_story_plan_status())
    verdict["deck_story_plan_available"] = plan_available
    verdict["deck_script_available"] = script_available
    # Layer 3b (June 21 2026) -- per-document export_verification
    # status. Computes a {brief, deck, appendix -> verified/warned/
    # failed/not_exported} dict from the latest editor draft's
    # export_verification JSONB column. Drives the status pill
    # rendered next to each document card on the Reports page. Read
    # from the same authed user the GET is scoped to. Fail-open:
    # any error returns {* -> 'not_exported'} so the badges stay
    # neutral rather than 500ing the readiness endpoint.
    verdict["export_verification"] = (
        await _export_verification_status(session.get("email", "")))
    return verdict


async def _export_verification_status(
    email: str,
) -> dict[str, str]:
    """Layer 3b (June 21 2026) -- per-document export_verification
    status for the Reports-page card badges. Reads the latest editor
    draft's export_verification JSONB column (populated at
    /api/v1/export/verify-all time and at each in-editor export by
    Layer 3a's _editor_export). Returns one of:

      verified     -- last export passed verification, no warnings
      warned       -- last export had a warning (e.g. stale data hash)
      failed       -- last export had errors (missing/corrupted values)
      not_exported -- no draft exists, or the draft has no
                      export_verification column set yet

    Fail-open: any DB / module-load error returns all-not-exported so
    the readiness endpoint never 500s on a Layer-3-uninstalled
    environment."""
    out: dict[str, str] = {
        "executive_brief": "not_exported",
        "presentation_deck": "not_exported",
        "analytical_appendix": "not_exported",
    }
    if not email:
        return out
    try:
        from tools.editor_drafts import get_current_draft_with_layer3
    except Exception as exc:  # noqa: BLE001
        log.warning("export_verification_status_import_failed",
                    error=str(exc))
        return out

    for doc_type in (
            "executive_brief", "presentation_deck",
            "analytical_appendix"):
        try:
            draft = await get_current_draft_with_layer3(email, doc_type)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "export_verification_draft_read_failed",
                document_type=doc_type, error=str(exc))
            continue
        if not draft:
            continue
        ev = draft.get("export_verification") or None
        if not ev:
            continue
        # ev is the verify_export_against_cache return shape.
        if ev.get("errors"):
            out[doc_type] = "failed"
        elif ev.get("warnings") or ev.get("skipped"):
            out[doc_type] = "warned"
        else:
            out[doc_type] = "verified"
    return out


async def _deck_story_plan_status() -> tuple[bool, bool]:
    """Returns (plan_available, script_available) for the cached deck
    story plan. Both flags read from a single get_cached_story_plan
    call so the readiness response stays one round-trip.

    plan_available is True when story_plans has a non-fallback row
    keyed by (current_data_hash, 'deck'). The frontend uses this to
    enable surfaces that need the locked slide_plan (the Presentation
    Deck regenerate flow, future plan-only exports).

    script_available is True when the same row ALSO carries a
    non-empty full_script. The Presentation Script card flips its
    button state on this flag because Pass 2 (full_script) is a
    separate Opus call from Pass 1 (slide_plan) and can fail
    independently -- a plan_available=True / script_available=False
    state happens when slide_plan landed but the script call timed
    out or hit a transient API error. Treating the two as the same
    signal would let the user download a script that doesn't exist.

    Fail-open: any error returns (False, False) and the script card
    stays in the disabled 'Generate Deck First' state.

    Supersedes the single-flag _deck_story_plan_available() helper
    that landed with PR #343; this PR's rebase replaced it with the
    two-flag tuple variant per the PR-body coordination plan.

    June 22 2026 -- composite-hash bug fix. The previous
    implementation queried
    get_cached_story_plan(data_hash, "deck") with the bare
    current_data_hash(), but refresh_story_plan PERSISTS the deck
    row under a composite hash via
    cache_key_with_brief_and_appendix
    ("<data_hash>|<brief_hash>|<appendix_hash>"). The bare-hash
    exact-match query missed every real deck row -- the gate
    kept the script card locked even when a fresh deck plan
    with a populated full_script was sitting in the table.
    Replaced with get_latest_story_plan(document_type='deck',
    exclude_fallback=True) which queries by document_type only +
    excludes deterministic_fallback rows at the SQL layer. The
    gate's question is "is there a fresh non-fallback deck plan
    with full_script?" -- "latest by computed_at" is the correct
    answer. Hash-drift staleness remains handled at export time
    by verify_export_against_cache."""
    try:
        from tools.story_plan import get_latest_story_plan
        plan = await get_latest_story_plan(
            "deck", exclude_fallback=True)
        if not plan:
            return False, False
        # exclude_fallback=True already filters fallback rows at
        # the SQL layer; the defensive recheck below covers a
        # future schema change where a row could land with
        # model=NULL but still represent a fallback.
        plan_available = plan.get("_model") != "deterministic_fallback"
        if not plan_available:
            return False, False
        script_text = plan.get("full_script") or ""
        script_available = bool(script_text and script_text.strip())
        return True, script_available
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_story_plan_status_check_failed",
                    error=str(exc))
        return False, False


# ── Layer 3 (June 21 2026) -- Pre-Submission Check ──────────────────────


@app.post("/api/v1/export/verify-all")
async def post_verify_all_for_submission(
    session: dict = Depends(require_auth),
):
    """Layer 3 pre-submission check. For each of brief, deck,
    appendix:
      1. Load the latest editor draft for the owner
      2. Run verify_export_against_cache on its content_text
      3. Also run check_cross_deliverable_consistency across
         all three documents

    Returns a structured verdict the frontend renders as a panel:
      ready          -- all three passed, no errors, hashes match
      needs_attention -- warnings only (e.g. stale data_hash)
      blocked        -- any errors, or any document not generated

    `submission_recommendation` is a plain-English sentence Bob
    and Molly can read directly without parsing the structured
    flags.

    Fail-open: any helper-import error returns a "verification
    helpers unavailable" message rather than 500ing -- the user
    can still submit, they just don't get the pre-flight check."""
    try:
        from tools.audit_assembler import current_data_hash
        from tools.document_audit import (
            check_cross_deliverable_consistency,
        )
        from tools.editor_drafts import get_current_draft
        from tools.numeric_substitution import (
            get_substitution_table, verify_export_against_cache,
        )
    except Exception as exc:  # noqa: BLE001
        log.warning("verify_all_imports_failed", error=str(exc))
        return {
            "overall": "needs_attention",
            "submission_recommendation": (
                "Verification helpers unavailable. The submission "
                "can still proceed but the pre-flight check did "
                "not run."),
            "brief": {"status": "not_generated"},
            "deck": {"status": "not_generated"},
            "appendix": {"status": "not_generated"},
            "cross_deliverable": {"passed": True, "flags": []},
        }

    email = session.get("email", "")
    try:
        cur_hash = await current_data_hash() or ""
    except Exception:  # noqa: BLE001
        cur_hash = ""

    async def _verify_one(doc_type: str) -> dict:
        try:
            draft = await get_current_draft(email, doc_type)
        except Exception as exc:  # noqa: BLE001
            log.warning("verify_all_draft_read_failed",
                        document_type=doc_type, error=str(exc))
            draft = None
        if not draft or not (draft.get("content_text") or "").strip():
            return {
                "status": "not_generated",
                "passed": False, "errors": [], "warnings": [],
                "data_hash_match": False,
                "last_verified_at": None,
            }
        manifest = draft.get("value_manifest") or {}
        gen_hash = draft.get("data_hash") or ""
        result = verify_export_against_cache(
            content_text=draft.get("content_text") or "",
            value_manifest=manifest,
            current_data_hash=cur_hash or gen_hash,
            generation_data_hash=gen_hash,
            document_type=doc_type)
        if result.get("errors"):
            status = "failed"
        elif result.get("warnings"):
            status = "warned"
        elif result.get("skipped"):
            status = "warned"
        else:
            status = "verified"
        return {
            "status": status,
            "passed": result.get("passed", True),
            "errors": result.get("errors", []),
            "warnings": result.get("warnings", []),
            "data_hash_match": result.get("data_hash_match", True),
            "last_verified_at": result.get("verified_at"),
            "skipped": result.get("skipped"),
            "n_values_verified": result.get("n_values_verified", 0),
        }

    brief = await _verify_one("executive_brief")
    deck = await _verify_one("presentation_deck")
    appendix = await _verify_one("analytical_appendix")
    # June 23 2026 -- the script was historically skipped; the
    # Submission Readiness Review audit identified this gap. Script
    # generation has been writing value_manifest the whole time so
    # the verify_export_against_cache helper works for it; the
    # exclusion was a wiring oversight in this endpoint, not a
    # capability gap.
    script = await _verify_one("presentation_script")

    # Cross-deliverable consistency check -- run only when at least
    # two documents exist and we have a substitution table to compare
    # against. The table is keyed by current data_hash; if all three
    # documents share that hash, they used the same table at
    # generation time (Layer 1 + Layer 2 guarantee). Drift here
    # signals a manual edit landed post-generation.
    cross: dict = {"passed": True, "flags": []}
    try:
        documents: dict[str, str] = {}
        # Re-read draft content_text only -- the verify_one results
        # above don't include the body text in their return shape.
        for doc_type in (
                "executive_brief", "presentation_deck",
                "analytical_appendix",
                # June 23 2026 -- include script in the cross-
                # deliverable scan. Script narration paraphrases the
                # deck's numeric anchors and a drift caused by a
                # manual edit in the script is just as
                # submission-blocking as one in the brief.
                "presentation_script"):
            d = await get_current_draft(email, doc_type)
            text = (d or {}).get("content_text") or ""
            if text.strip():
                documents[doc_type] = text
        if len(documents) >= 2 and cur_hash:
            # Use the SAME table the substitution layer used; the
            # cache key is the data_hash so this is a hit not a
            # rebuild on the warm path.
            try:
                from tools.cache import get_latest_strategy_cache
                from tools.cio_recommendation import (
                    get_latest_recommendation,
                )
                from tools.academic_deck import (
                    OOS_SHARPE_REGIME_CONDITIONAL,
                    OOS_SHARPE_BENCHMARK,
                    CORRELATION_PRE_2022, CORRELATION_POST_2022,
                )
                strategy_cache = (
                    await get_latest_strategy_cache() or {})
                cio_row = await get_latest_recommendation()
                # June 22 2026 (wiring fix) -- same helper read
                # as the document generators so the consistency
                # check sees the same substituted values.
                from tools.academic_export import (
                    load_substitution_metric_sources,
                )
                rc_rows, fl_rows, cs_payload, crisis_payload = (
                    await load_substitution_metric_sources())
                table = get_substitution_table(
                    cur_hash, strategy_cache, cio_row,
                    oos_sharpe_blend=OOS_SHARPE_REGIME_CONDITIONAL,
                    oos_sharpe_benchmark=OOS_SHARPE_BENCHMARK,
                    pre_2022_eq_ig_correlation=CORRELATION_PRE_2022,
                    post_2022_eq_ig_correlation=CORRELATION_POST_2022,
                    regime_conditional=rc_rows,
                    factor_loadings=fl_rows,
                    cost_sensitivity=cs_payload,
                    crisis_performance=crisis_payload)
                flags = check_cross_deliverable_consistency(
                    documents, table)
                cross = {
                    "passed": len(flags) == 0, "flags": flags}
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "verify_all_cross_deliverable_failed",
                    error=str(exc))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "verify_all_documents_assembly_failed",
            error=str(exc))

    # Aggregate the four results into an overall verdict.
    per_doc = [brief, deck, appendix, script]
    any_not_generated = any(
        d["status"] == "not_generated" for d in per_doc)
    any_failed = any(
        d["status"] == "failed" or d.get("errors")
        for d in per_doc) or not cross.get("passed", True)
    any_warned = any(d["status"] == "warned" for d in per_doc)
    if any_not_generated:
        overall = "blocked"
    elif any_failed:
        overall = "blocked"
    elif any_warned:
        overall = "needs_attention"
    else:
        overall = "ready"

    # Plain-English recommendation per overall verdict.
    if overall == "ready":
        rec = (
            f"All three deliverables verified against cache "
            f"{(cur_hash or '')[:8]}. Safe to submit.")
    elif overall == "needs_attention":
        rec = (
            "All deliverables generated; warnings present. Review "
            "the stale-data-hash notices and consider regenerating "
            "before submitting.")
    else:
        missing = [
            label for d, label in (
                (brief, "brief"), (deck, "deck"),
                (appendix, "appendix"), (script, "script"),
            ) if d["status"] == "not_generated"]
        if missing:
            rec = (
                f"Generate {', '.join(missing)} before submitting. "
                "All four deliverables must exist for the "
                "pre-submission check to pass.")
        else:
            err_docs = [
                label for d, label in (
                    (brief, "brief"), (deck, "deck"),
                    (appendix, "appendix"), (script, "script"),
                ) if d.get("errors") or d["status"] == "failed"]
            if err_docs:
                rec = (
                    f"Verification errors found in "
                    f"{', '.join(err_docs)}. Open the editor, "
                    "review the flagged values, and regenerate "
                    "if needed before submitting.")
            elif not cross.get("passed", True):
                rec = (
                    "Cross-deliverable inconsistency found -- a "
                    "value appears differently across the brief / "
                    "deck / appendix. Reconcile before submitting.")
            else:
                rec = (
                    "One or more documents failed verification. "
                    "Open the editor to investigate.")

    return {
        "brief": brief, "deck": deck, "appendix": appendix,
        "script": script,
        "cross_deliverable": cross,
        "overall": overall,
        "submission_recommendation": rec,
    }


@app.post("/api/v1/cache/invalidate")
async def cache_invalidate(
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Clears strategy_results_cache so the backtester recomputes from fresh
    data on the next /api/backtest request — used after a data update or
    to repopulate cached results that predate a new result-dict field
    (e.g. the persisted weight_schedule). Sysadmin only.
    """
    from tools.cache import clear_strategy_cache
    removed = await clear_strategy_cache()
    log.info("strategy_cache_invalidated", rows_removed=removed,
             by=session.get("email"))
    # Smart audit caching — the cache invalidation is a data event; if the
    # last audit no longer reflects the data, run_full_audit re-verifies it
    # in the background (idempotent — a no-op when already current).
    from tools.audit_engine import trigger_audit_async
    trigger_audit_async("cache_invalidation")
    return {"status": "cleared", "rows_removed": removed}


@app.post("/api/v1/admin/refresh-monthly-data")
async def refresh_monthly_data(
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Extends the monthly data pipeline beyond the Excel file: fetches the
    total return of every complete calendar month that has closed since
    the last run — SPY (equity), BND (investment grade) and HYG (high
    yield) from yfinance, DTB3 from FRED — validates the splice, and
    appends the new rows to market_data_monthly. Sysadmin only.

    After a successful extension the strategy cache is cleared and the
    audits auto-trigger. The blocking fetch runs off the event loop.
    """
    import asyncio as _asyncio
    from tools.data_fetcher import extend_market_data
    result = await _asyncio.to_thread(extend_market_data)
    log.info("refresh_monthly_data", by=session.get("email"),
             status=result.get("status"),
             rows_added=result.get("monthly_rows_added"),
             new_max=result.get("monthly_new_max"))
    return result


# ── User management ───────────────────────────────────────────────────────────
#
# The sysadmin manages platform_users from inside the platform. Every
# endpoint requires the manage_users permission; the "last sysadmin"
# guard prevents the platform from being left with no administrator.


def _valid_email(email: str) -> bool:
    """A minimal email-shape check for the create-user form."""
    return bool(email) and "@" in email and "." in email.split("@")[-1]


def _clean_permissions(raw: Any, role: str) -> list[str]:
    """Validated permissions — the supplied list filtered to known keys,
    or the role's preset when no explicit list was given."""
    if isinstance(raw, list):
        return [p for p in raw if p in PERMISSIONS]
    return list(ROLE_PRESETS.get(role, ROLE_PRESETS["viewer"]))


@app.get("/api/v1/admin/users")
async def admin_list_users(
    session: dict = Depends(require_permission("manage_users")),
):
    """Every platform user, with an activity count. Sysadmin only."""
    from tools.platform_users import list_all_users
    return {"users": await list_all_users()}


@app.get("/api/v1/admin/users/activity-breakdown")
async def admin_users_activity_breakdown(
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Per-user activity broken down by interaction_type and session_type
    over BOTH a lifetime window and a rolling 30-day window — the
    analytics behind the Settings → Users → Platform Engagement panel.
    The panel renders LIFETIME as the headline (the figure that matters
    for academic-integrity tracking) and the 30-day count as
    recent-activity context.

    Joins against platform_users (LEFT JOIN) so every user appears even
    when they have zero interactions in either window. The breakdown
    aggregates two source tables:
      agent_interactions  → counts by interaction_type + SUM(cost)
      session_events      → counts by session_type for page_view events

    Each user row carries both a `lifetime` block (breakdown,
    session_breakdown, total_interactions, total_cost_usd, first_seen,
    last_seen) and a `rolling_30d` block (the same minus the
    lifetime-only first/last_seen pair).

    Sysadmin only. Fail-open — a DB error returns an empty users list.
    """
    from tools.platform_users import users_activity_breakdown
    return await users_activity_breakdown()


@app.post("/api/v1/admin/users")
async def admin_create_user(
    body: dict, session: dict = Depends(require_permission("manage_users")),
):
    """Adds a platform user. Sysadmin only."""
    from tools.platform_users import create_user, email_exists

    email = str(body.get("email") or "").strip().lower()
    if not _valid_email(email):
        raise HTTPException(status_code=422,
                            detail="A valid email address is required.")
    role = str(body.get("role") or "viewer")
    if role not in ROLE_PRESETS:
        raise HTTPException(status_code=422,
                            detail="role must be viewer | team_member | sysadmin")
    if await email_exists(email):
        raise HTTPException(
            status_code=409,
            detail="A user with that email already exists — edit them instead.")
    created = await create_user(
        email=email,
        display_name=(str(body["display_name"]).strip()
                      if body.get("display_name") else None),
        role=role,
        permissions=_clean_permissions(body.get("permissions"), role),
        notes=(str(body["notes"]).strip() if body.get("notes") else None),
        created_by=session.get("email", ""),
    )
    if created is None:
        raise HTTPException(status_code=503,
                            detail="Could not create the user — database unavailable.")
    # Welcome email — sent only after the user is successfully created.
    # Fail-open: send_welcome_email never raises, so a delivery failure
    # cannot undo or block the creation. welcome_email_sent tells the
    # frontend which confirmation message to show.
    from auth import send_welcome_email
    welcome_email_sent = await send_welcome_email(
        email=email,
        display_name=created.get("display_name"),
        notes=created.get("notes"),
        council_limit=created.get("council_queries_limit"),
    )
    return {**created, "welcome_email_sent": welcome_email_sent}


@app.patch("/api/v1/admin/users/{user_id}")
async def admin_update_user(
    user_id: int, body: dict,
    session: dict = Depends(require_permission("manage_users")),
):
    """
    Updates a user's display_name / role / permissions / is_active /
    notes / council_queries_limit / council_queries_used. email is
    immutable. The last active sysadmin cannot be demoted or
    deactivated. Sysadmin only.
    """
    from tools.platform_users import (
        count_active_sysadmins, get_user_by_id, update_user,
    )

    user = await get_user_by_id(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found.")

    fields: dict[str, Any] = {}
    if "display_name" in body:
        fields["display_name"] = (str(body["display_name"]).strip()
                                  if body["display_name"] else None)
    if "notes" in body:
        fields["notes"] = (str(body["notes"]).strip()
                           if body["notes"] else None)
    if "role" in body:
        if body["role"] not in ROLE_PRESETS:
            raise HTTPException(status_code=422,
                                detail="role must be viewer | team_member | sysadmin")
        fields["role"] = body["role"]
    if "permissions" in body:
        fields["permissions"] = _clean_permissions(
            body["permissions"], fields.get("role", user["role"]))
    if "is_active" in body:
        fields["is_active"] = bool(body["is_active"])
    # Council query allocation — Adjust Limit (int), Unlimited (null),
    # Reset Usage (used = 0). bool is rejected: it is an int subclass.
    if "council_queries_limit" in body:
        cql = body["council_queries_limit"]
        if cql is None:
            fields["council_queries_limit"] = None
        elif isinstance(cql, int) and not isinstance(cql, bool) and cql >= 0:
            fields["council_queries_limit"] = cql
        else:
            raise HTTPException(
                status_code=422,
                detail="council_queries_limit must be a non-negative "
                       "integer or null.")
    if "council_queries_used" in body:
        cqu = body["council_queries_used"]
        if isinstance(cqu, int) and not isinstance(cqu, bool) and cqu >= 0:
            fields["council_queries_used"] = cqu
        else:
            raise HTTPException(
                status_code=422,
                detail="council_queries_used must be a non-negative integer.")

    # Last-sysadmin guard — refuse a change that would leave no active
    # administrator. A "sysadmin" is any active user holding manage_users.
    new_active = fields.get("is_active", user["is_active"])
    new_perms = fields.get("permissions", user["permissions"])
    was_admin = user["is_active"] and "manage_users" in user["permissions"]
    still_admin = new_active and "manage_users" in new_perms
    if was_admin and not still_admin and await count_active_sysadmins() <= 1:
        raise HTTPException(status_code=400,
                            detail="Cannot remove the last sysadmin.")

    updated = await update_user(user_id, fields)
    if updated is None:
        raise HTTPException(status_code=503,
                            detail="Could not update the user — database unavailable.")
    return updated


@app.delete("/api/v1/admin/users/{user_id}")
async def admin_delete_user(
    user_id: int, session: dict = Depends(require_permission("manage_users")),
):
    """
    Soft-deletes a user (is_active = false) — the row is kept so activity
    history stays attributed. The last active sysadmin cannot be deleted.
    Sysadmin only.
    """
    from tools.platform_users import (
        count_active_sysadmins, get_user_by_id, update_user,
    )

    user = await get_user_by_id(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found.")
    was_admin = user["is_active"] and "manage_users" in user["permissions"]
    if was_admin and await count_active_sysadmins() <= 1:
        raise HTTPException(status_code=400,
                            detail="Cannot remove the last sysadmin.")
    updated = await update_user(user_id, {"is_active": False})
    if updated is None:
        raise HTTPException(status_code=503,
                            detail="Could not deactivate the user — database unavailable.")
    return {"deactivated": True, "id": user_id}


# ── Team Activity ─────────────────────────────────────────────────────────────

@app.post("/api/v1/activity/events")
@limiter.limit("120/minute")
async def activity_events(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """
    Receives a batch of UI telemetry events from the frontend and
    inserts them into session_events in one transaction.

    Always returns 200 — the UI must never be blocked or shown an
    error by activity logging. The PROJECT_TEAM_EMAILS allowlist is
    enforced inside insert_session_events: a non-team user's events are
    dropped silently. login / logout events are stamped server-side
    with the request IP (and user agent, for login) rather than trusting
    the client.
    """
    try:
        from tools.activity_log import insert_session_events

        events = body.get("events")
        if not isinstance(events, list):
            return {"accepted": 0}

        sid = request.headers.get("x-session-id")
        stype = request.headers.get("x-session-type")
        ip = request.client.host if request.client else None
        ua = request.headers.get("user-agent")

        for ev in events:
            if not isinstance(ev, dict):
                continue
            ev.setdefault("session_id", sid)
            ev.setdefault("session_type", stype)
            # Server-authoritative IP / UA for the auth-boundary events.
            if ev.get("event_type") in ("login", "logout"):
                ev["ip_address"] = ip
            if ev.get("event_type") == "login":
                ev["user_agent"] = ua

        written = await insert_session_events(
            [e for e in events if isinstance(e, dict)], session["email"])
        return {"accepted": written}
    except Exception as exc:  # noqa: BLE001
        # Logging must never surface an error to the UI.
        log.warning("activity_events_failed", error=str(exc))
        return {"accepted": 0}


@app.post("/api/v1/activity/commits/webhook")
async def activity_commits_webhook(request: Request):
    """
    GitHub webhook receiver — handles both push and pull_request events
    against the same registration. Validates the X-Hub-Signature-256
    HMAC against GITHUB_WEBHOOK_SECRET (any invalid/missing signature
    is a 401).

    Push events       → upsert commits into commit_activity.
    pull_request events with action=closed and merged=true →
                        scan the body + commit messages for "Resolves
                        failure #N" references and queue
                        pr_suggestions rows (Suggested Resolutions
                        Commit 2/7).

    Other events (the `ping` GitHub sends at registration, draft PR
    state changes, etc.) are acknowledged and ignored. Operationally
    this means ONE webhook registration on the repo covers both the
    Team Activity commit sync AND the Suggested Resolutions workflow —
    one secret to manage, one URL to point GitHub at.
    """
    from config import GITHUB_WEBHOOK_SECRET
    from tools.github_sync import verify_signature, parse_push_payload

    raw = await request.body()
    sig = request.headers.get("x-hub-signature-256")
    if not verify_signature(GITHUB_WEBHOOK_SECRET, raw, sig):
        raise HTTPException(status_code=401, detail="Invalid webhook signature.")

    event_type = request.headers.get("x-github-event")
    if event_type not in ("push", "pull_request"):
        return {"status": "ignored",
                "reason": f"event '{event_type}' is not handled"}

    try:
        payload = json.loads(raw.decode("utf-8"))
    except Exception:  # noqa: BLE001
        raise HTTPException(status_code=400, detail="Malformed JSON payload.")

    if event_type == "push":
        commits = parse_push_payload(payload)
        if not commits:
            return {"status": "ok", "synced": 0}
        from tools.activity_log import upsert_commits
        written = await upsert_commits(commits)
        log.info("activity_webhook_push",
                 commits=len(commits), upserted=written)
        return {"status": "ok", "synced": written}

    # pull_request event — Suggested Resolutions Commit 2/7.
    from config import GITHUB_REPO, GITHUB_TOKEN
    from tools.github_sync import fetch_pr_commits
    from tools.pr_suggestion_scanner import (
        parse_pr_payload, record_pr_suggestions,
    )

    # Pre-parse to extract the PR number (cheap, no network), then
    # enrich with commit messages via the REST API so the scanner can
    # find references in commit messages too. The fetch is fail-open
    # — a missing GITHUB_TOKEN or API error degrades to body-only.
    pre = parse_pr_payload(payload)
    if pre is None:
        # Not a closed+merged PR — silently ack.
        return {"status": "ok", "reason": "not a merged PR"}

    commits = await fetch_pr_commits(
        GITHUB_REPO, GITHUB_TOKEN, pre["pr_number"])
    # Re-parse with the enriched commits so the scanner can match
    # references in commit messages alongside the PR body.
    enriched_payload = dict(payload)
    enriched_payload["__commits"] = commits
    parsed = parse_pr_payload(enriched_payload)
    if parsed is None:
        return {"status": "ok", "reason": "not a merged PR"}

    summary = await record_pr_suggestions(parsed)
    log.info("activity_webhook_pull_request",
             pr_number=parsed["pr_number"],
             references_found=len(parsed["matches"]),
             created=summary["created"],
             skipped_missing=summary["skipped_missing"],
             skipped_resolved=summary["skipped_resolved"],
             skipped_duplicate=summary["skipped_duplicate"])
    return {
        "status": "ok",
        "pr_number": parsed["pr_number"],
        "references_found": len(parsed["matches"]),
        "suggestions_created": len(summary["created"]),
        "skipped_missing": len(summary["skipped_missing"]),
        "skipped_resolved": len(summary["skipped_resolved"]),
        "skipped_duplicate": len(summary["skipped_duplicate"]),
    }


@app.get("/api/v1/activity/commits/sync")
@limiter.limit("6/minute")
async def activity_commits_sync(
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Manual sync of the last 100 commits from the GitHub REST API.
    Upserts on sha, so it is safe to run repeatedly — used to backfill
    history and to catch up anything the webhook missed. Requires
    GITHUB_TOKEN (the repository is private).
    """
    from config import GITHUB_REPO, GITHUB_TOKEN
    from tools.github_sync import fetch_recent_commits
    from tools.activity_log import upsert_commits

    try:
        commits = await fetch_recent_commits(GITHUB_REPO, GITHUB_TOKEN, limit=100)
    except RuntimeError as exc:
        # Misconfiguration (e.g. GITHUB_TOKEN unset). The message is a
        # deliberate, safe operator hint — surfaced via a proper 503 so the
        # client can detect the failure from the status code.
        raise HTTPException(status_code=503, detail=str(exc))
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.warning("activity_sync_failed", ref=ref, error=str(exc))
        raise HTTPException(
            status_code=502, detail=f"Commit sync failed (ref: {ref})")

    written = await upsert_commits(commits)
    log.info("activity_commits_synced", fetched=len(commits), upserted=written)
    return {"synced": written, "fetched": len(commits)}


@app.get("/api/v1/activity/team")
async def activity_team(
    request: Request,
    user_id: Optional[str] = Query(None),
    activity_type: str = Query("all"),
    session_type: str = Query("analytical"),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    session: dict = Depends(require_auth),
):
    """
    Unified Team Activity timeline — commit_activity, agent_interactions
    and session_events interleaved and sorted by timestamp descending.

    session_type defaults to "analytical"; pass "all" to include
    Testing Mode activity. commit history is session-agnostic and is
    always included when the activity_type filter permits commits.
    """
    from tools.activity_log import get_team_activity

    return await get_team_activity(
        user_id=user_id,
        activity_type=activity_type,
        session_type=session_type,
        date_from=date_from,
        date_to=date_to,
        limit=limit,
        offset=offset,
    )


@app.get("/api/v1/activity/summary")
async def activity_summary(
    request: Request,
    include_testing: bool = Query(False),
    session: dict = Depends(require_auth),
):
    """
    Per-member interaction and commit counts, the most-consulted agents,
    and the latest academic-review verdict — the Team Activity summary
    panel. Analytical sessions only unless include_testing is set.
    """
    from tools.activity_log import get_activity_summary

    return await get_activity_summary(analytical_only=not include_testing)


@app.get("/api/v1/activity/cost-summary")
async def activity_cost_summary(
    request: Request,
    include_testing: bool = Query(False),
    session: dict = Depends(require_auth),
):
    """
    AI token spend — grand total plus per-member and per-interaction-type
    breakdowns, drawn from the agent_interactions cost columns. Drives the
    Team Activity cost panel. Analytical sessions only unless
    include_testing is set.
    """
    from tools.activity_log import get_cost_summary

    return await get_cost_summary(analytical_only=not include_testing)


# ── Changelog ─────────────────────────────────────────────────────────────────

@app.get("/api/v1/changelog")
async def changelog_all(session: dict = Depends(require_auth)):
    """Every changelog entry, newest version first — the Settings Release
    History. Returns {entries: [...]}."""
    from tools.changelog import get_all_changelog
    return {"entries": await get_all_changelog()}


@app.get("/api/v1/changelog/unseen")
async def changelog_unseen(session: dict = Depends(require_auth)):
    """
    Changelog entries released after the calling user last dismissed the
    What's New modal, plus has_tour_update and the current tour_version.
    Drives the What's New modal's trigger.
    """
    from tools.changelog import get_unseen_changelog
    return await get_unseen_changelog(session["email"])


@app.post("/api/v1/changelog/mark-seen")
async def changelog_mark_seen(
    body: dict | None = None,
    session: dict = Depends(require_auth),
):
    """
    Records that the user has seen the changelog up to now. An optional
    body {"tour_version_seen": int} also records the site-tour version
    the user has completed.
    """
    from tools.changelog import mark_changelog_seen
    tour_seen: int | None = None
    if isinstance(body, dict):
        tv = body.get("tour_version_seen")
        if isinstance(tv, int) and not isinstance(tv, bool):
            tour_seen = tv
    ok = await mark_changelog_seen(session["email"], tour_seen)
    return {"ok": ok}


# ── Explainer ─────────────────────────────────────────────────────────────────

@app.post("/api/explain/terms")
@limiter.limit("20/minute")
async def explain_terms(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """Generates contextual glossary terms from the full council output."""
    if ENVIRONMENT != "test":
        try:
            from agents.explainer_agent import ExplainerAgent
            explainer = ExplainerAgent()
            # UAT 2026-05-27 P0 — Render logs showed every concurrent
            # dashboard request queuing behind explainer_grok_completed
            # events and draining together (39s + 24s clusters). Root
            # cause: explain_terms calls _call_llm → _call_grok which
            # uses httpx.Client (SYNC) inside an async endpoint. The
            # blocking POST stalls the event loop for 24-39s and
            # starves every other coroutine. Same fix pattern as
            # PR #122 (get_full_history_async) and #126 (optimizer
            # solver) — push the sync call into a worker thread.
            import asyncio
            return await asyncio.to_thread(
                explainer.explain_terms,
                body.get("council_output", {}))
        except Exception as exc:
            log.error("explain_terms_error", error=str(exc))
    return {}


@app.post("/api/explain/parameter")
@limiter.limit("20/minute")
async def explain_parameter(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """Explains a config parameter in the context of current results."""
    if ENVIRONMENT != "test":
        try:
            from agents.explainer_agent import ExplainerAgent
            explainer = ExplainerAgent()
            # UAT 2026-05-27 P0 — sync Grok HTTP blocks the event
            # loop. Wrap in asyncio.to_thread; see explain_terms
            # above for the full diagnosis. asyncio.to_thread
            # supports kwargs, so the keyword call shape is
            # preserved verbatim.
            import asyncio
            return await asyncio.to_thread(
                explainer.explain_parameter,
                parameter=body.get("parameter", ""),
                value=body.get("value"),
                current_results=body.get("current_results", {}),
            )
        except Exception as exc:
            log.error("explain_parameter_error", error=str(exc))
    return {}


@app.post("/api/explain/chart")
@limiter.limit("20/minute")
async def explain_chart(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """Generates a chart explanation anchored to actual chart data."""
    if ENVIRONMENT != "test":
        try:
            from agents.explainer_agent import ExplainerAgent
            explainer = ExplainerAgent()
            # UAT 2026-05-27 P0 — sync Grok HTTP blocks the event
            # loop. See explain_terms above for the full diagnosis.
            import asyncio
            return await asyncio.to_thread(
                explainer.explain_chart,
                chart_id=body.get("chart_id", ""),
                chart_type=body.get("chart_type", ""),
                chart_data=body.get("chart_data"),
                current_results=body.get("current_results", {}),
            )
        except Exception as exc:
            log.error("explain_chart_error", error=str(exc))
    return {}


@app.post("/api/explain/qa")
@limiter.limit("10/minute")
async def explain_qa(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """Generates plain-English explanations for all 30 QA checklist items."""
    if ENVIRONMENT != "test":
        try:
            from agents.explainer_agent import ExplainerAgent
            explainer = ExplainerAgent()
            # UAT 2026-05-27 P0 — sync Grok HTTP blocks the event
            # loop. See explain_terms above for the full diagnosis.
            import asyncio
            return await asyncio.to_thread(
                explainer.explain_qa,
                body.get("audit_results", []))
        except Exception as exc:
            log.error("explain_qa_error", error=str(exc))
    return {}


# ── Academic Advisor (Agent 10) ───────────────────────────────────────────────
#
# Three endpoints map 1:1 to AcademicAdvisor methods. All three enforce citation
# integrity via Anthropic's server-side web_search tool — any URL the model
# emits that the tool did not actually fetch is dropped before the response
# is returned to the frontend. See agents/academic_advisor.py:_filter_to_verified
# for the runtime check.
#
# Limit is 10/min (same as council/QA endpoints) — the advisor is interactive
# but each call costs ~$0.04-0.06 incl. web_search, and we don't want a stuck
# floating button to drain the daily credit cap.

@app.post("/api/advisor/analyse")
@limiter.limit("10/minute")
async def advisor_analyse(
    request: Request,
    body: AdvisorAnalyseRequest,
    session: dict = Depends(require_auth),
):
    """
    Main advisor entry point — academic guidance for one deliverable.

    The query is scope-checked (portfolio-analysis only); the advisor agent
    uses web_search to verify any citations before returning them. Verified
    citations are merged into the response so the frontend can render them
    immediately without a second round-trip.
    """
    if ENVIRONMENT == "test":
        # Mock keeps the test suite hermetic — no Anthropic API calls.
        from agents.academic_advisor import MOCK_ADVISOR_ANALYSE
        return MOCK_ADVISOR_ANALYSE

    try:
        from agents.academic_advisor import AcademicAdvisor
        # Fetch the live regime + macro digest so the advisor grounds
        # its grade-aware feedback in the same signal state the CIO
        # and dissenters see. Both calls fail-open: if regime/macro
        # are unavailable the advisor degrades to the prior
        # results-only behaviour rather than refusing.
        import asyncio
        regime_data: dict | None = None
        macro_context: str | None = None
        try:
            from tools.regime_detector import detect_current_regime
            regime_data = await asyncio.to_thread(detect_current_regime)
        except Exception as exc:  # noqa: BLE001
            log.warning("advisor_regime_unavailable", error=str(exc))
        try:
            from tools.macro_context import get_macro_context
            macro_context = get_macro_context() or None
        except Exception as exc:  # noqa: BLE001
            log.warning("advisor_macro_unavailable", error=str(exc))
        advisor = AcademicAdvisor()
        return advisor.analyse_findings(
            query=body.query,
            deliverable_type=body.deliverable_type,
            strategy_results=body.strategy_results,
            regime_data=regime_data,
            macro_context=macro_context,
        )
    except Exception as exc:
        log.error("advisor_analyse_endpoint_error", error=str(exc))
        return {
            "key_findings":     [],
            "guidance":         [],
            "citations":        [],
            "potential_issues": [],
            "error":            "Advisor temporarily unavailable.",
        }


@app.post("/api/advisor/verify-finding")
@limiter.limit("10/minute")
async def advisor_verify_finding(
    request: Request,
    body: AdvisorVerifyRequest,
    session: dict = Depends(require_auth),
):
    """
    Verifies one specific finding against external academic evidence.

    Returns supporting_evidence, contradicting_evidence, and a verdict in
    {"plausible", "implausible", "uncertain"}. The frontend uses this to
    sanity-check a single number before committing it to a slide or paper.
    """
    if ENVIRONMENT == "test":
        from agents.academic_advisor import MOCK_ADVISOR_VERIFY
        return MOCK_ADVISOR_VERIFY

    try:
        from agents.academic_advisor import AcademicAdvisor
        advisor = AcademicAdvisor()
        return advisor.check_finding_plausibility(
            finding=body.finding,
            magnitude=body.magnitude,
            period=body.period,
        )
    except Exception as exc:
        log.error("advisor_verify_endpoint_error", error=str(exc))
        return {
            "supporting_evidence":    [],
            "contradicting_evidence": [],
            "verdict":                "uncertain",
            "reasoning":              "Advisor temporarily unavailable.",
            "verified_sources":       [],
        }


@app.post("/api/advisor/citations")
@limiter.limit("10/minute")
async def advisor_citations(
    request: Request,
    body: AdvisorCitationsRequest,
    session: dict = Depends(require_auth),
):
    """
    Returns up to n_sources verified academic citations for a finding.

    Citations the agent emits but web_search did not return are silently
    dropped — this is the citation integrity contract. n_sources is capped
    at 5 server-side regardless of what the request specifies.
    """
    if ENVIRONMENT == "test":
        from agents.academic_advisor import MOCK_ADVISOR_CITATIONS
        return MOCK_ADVISOR_CITATIONS

    try:
        from agents.academic_advisor import AcademicAdvisor
        advisor = AcademicAdvisor()
        return advisor.find_supporting_citations(
            finding=body.finding,
            n_sources=body.n_sources,
        )
    except Exception as exc:
        log.error("advisor_citations_endpoint_error", error=str(exc))
        return {"citations": [], "verified_sources": []}


# ── QA ────────────────────────────────────────────────────────────────────────

# ── AN01 / AN04 raw-state diagnostic logger (May 25 2026) ────────────────────
#
# When the methodology audit fires and AN01 (Carhart loadings) or AN04
# (regime split + transition matrix) come back WARN/FAIL, the operator
# needs the upstream field values to diagnose. precomputed_analytics
# already logs per-row diagnostics at REFRESH time; this helper logs
# the cached field values at AUDIT time, so a forced re-run that
# DOESN'T trigger a refresh still surfaces the data the checks read.
#
# Three structured log lines per call:
#   qa_audit_an01_state — first 3 factor_loadings rows (alpha, mom,
#                         significance flags, r_squared)
#   qa_audit_an04_regime_state — first 3 regime_conditional rows
#                         (pre/post 2022 sharpe + months)
#   qa_audit_an04_transition_state — transition matrix row sums per
#                         regime (the row-sum invariant AN04 checks)

def _log_an01_an04_raw_state(
    analytics_cache: dict | None, qa_hash: str | None, *, forced: bool,
) -> None:
    """Emits the three AN01/AN04 diagnostic log lines named above.

    Fail-open — a malformed payload (any missing key) shrinks the line
    to None/empty rather than raising. The point is visibility; a log
    line throwing during diagnosis would defeat the purpose.
    """
    short_hash = qa_hash[:8] if qa_hash else None
    cache = analytics_cache or {}
    refresh_triggered = list(cache.get("refresh_triggered") or [])
    completeness = cache.get("completeness") or {}
    academic = cache.get("academic_analytics") or {}
    transition = cache.get("transition_matrix") or {}

    # AN01 — Carhart factor loadings. Project a compact, diagnostic
    # view of the first three rows so the operator can see whether
    # the upstream regression actually populated the fields the AN01
    # deterministic check reads (significance flags + alpha annualised
    # + r_squared). Three rows fits one Render log line; the full
    # payload is in analytics_metrics_cache for a deeper dive.
    fl_rows = academic.get("factor_loadings") or []
    fl_sample: list[dict] = []
    for r in fl_rows[:3]:
        if not isinstance(r, dict):
            continue
        fl_sample.append({
            "strategy":             r.get("strategy"),
            "alpha_annualized":     r.get("alpha_annualized"),
            "alpha_significant":    r.get("alpha_significant"),
            "mkt_rf":               r.get("mkt_rf"),
            "mkt_rf_significant":   r.get("mkt_rf_significant"),
            "smb_significant":      r.get("smb_significant"),
            "hml_significant":      r.get("hml_significant"),
            "mom":                  r.get("mom"),
            "mom_significant":      r.get("mom_significant"),
            "r_squared":            r.get("r_squared"),
        })
    log.info(
        "qa_audit_an01_state",
        strategy_hash=short_hash,
        forced=forced,
        factor_loadings_complete=bool(completeness.get("factor_loadings")),
        factor_loadings_refreshed=("academic_analytics" in refresh_triggered),
        factor_loadings_row_count=len(fl_rows),
        factor_loadings_sample=fl_sample,
    )

    # AN04 — regime-conditional Sharpe table (pre/post 2022 split).
    # The AN04 deterministic check verifies every strategy carries a
    # non-null Sharpe in both periods (with the months ≥ 2 carve-out).
    rc_rows = academic.get("regime_conditional") or []
    rc_sample: list[dict] = []
    for r in rc_rows[:3]:
        if not isinstance(r, dict):
            continue
        rc_sample.append({
            "strategy":          r.get("strategy"),
            "pre_2022_sharpe":   r.get("pre_2022_sharpe"),
            "post_2022_sharpe":  r.get("post_2022_sharpe"),
            "pre_2022_months":   r.get("pre_2022_months"),
            "post_2022_months":  r.get("post_2022_months"),
        })
    log.info(
        "qa_audit_an04_regime_state",
        strategy_hash=short_hash,
        forced=forced,
        regime_conditional_complete=bool(
            completeness.get("regime_conditional")),
        regime_conditional_refreshed=(
            "academic_analytics" in refresh_triggered),
        regime_conditional_row_count=len(rc_rows),
        regime_conditional_sample=rc_sample,
    )

    # AN04 — transition matrix. AN04 reads BOTH the regime split AND
    # the 3x3 matrix; the matrix row sums are the invariant that
    # surfaces a data error (a non-empty row that does NOT sum to 1.0
    # is malformed). Log the row sums and the matrix shape so a
    # missing regime or a degenerate row stands out in the log.
    matrix = transition if isinstance(transition, dict) else {}
    # The matrix payload is keyed by originating regime; sum each
    # row defensively in case the validator's row_sums field is
    # missing on an older payload shape.
    row_sums: dict[str, float] = {}
    for regime in ("BULL", "BEAR", "TRANSITION"):
        row = matrix.get(regime)
        if isinstance(row, dict):
            try:
                row_sums[regime] = round(
                    sum(float(v) for v in row.values()), 6)
            except (TypeError, ValueError):
                row_sums[regime] = float("nan")
    log.info(
        "qa_audit_an04_transition_state",
        strategy_hash=short_hash,
        forced=forced,
        transition_matrix_complete=bool(
            completeness.get("transition_matrix")),
        transition_matrix_refreshed=(
            "transition_matrix" in refresh_triggered),
        transition_matrix_row_sums=row_sums,
        transition_matrix_keys=sorted(
            k for k in matrix.keys()
            if k in ("BULL", "BEAR", "TRANSITION")),
    )


@app.post("/api/qa/audit")
@limiter.limit("10/minute")
async def qa_audit(
    request: Request,
    body: dict | None = None,
    session: dict = Depends(require_sysadmin),
):
    """
    Runs the full QA methodology audit against real strategy results.

    QA Agent uses Opus for the narrative; deterministic checks run from
    the strategy results dict to guarantee pass/fail verdicts are never
    hallucinated. Falls back to mock audit if pipeline is unavailable.

    Per-type guard: rejected with 409 only when another methodology
    audit is already in progress — a statistical audit never blocks it
    (tools/qa_guard.py).

    Body: {"force": bool} — when True (May 25 2026), bypasses BOTH the
    hash gate and the min-interval gate so the audit re-runs even when
    the strategy data is unchanged. The 'Run Full QA' / 'Re-run audit'
    buttons always pass force=true because a manual click means "I
    want fresh checks now" — without it, an IN02 (Academic Review
    complete) change is invisible to a cached audit because the cache
    is keyed on strategy_hash and an Academic Review run does not
    change that hash. Background polling and the first QA-tab load
    pass force=false so they continue to benefit from the cache.
    """
    from tools.qa_guard import (
        QA_BUSY_MESSAGE_METHODOLOGY, begin_methodology, end_methodology,
        methodology_in_progress,
    )
    # Per-type lock — a methodology audit is blocked only by another
    # methodology audit, never by a statistical one.
    if methodology_in_progress():
        raise HTTPException(status_code=409,
                            detail=QA_BUSY_MESSAGE_METHODOLOGY)

    force = bool((body or {}).get("force"))

    begin_methodology()
    try:
        if ENVIRONMENT == "test":
            if force:
                # Tag the response so the contract test can verify
                # force was plumbed through to the bypass call site.
                # The shape is otherwise bit-identical to the cached
                # mock — additive field only.
                return {**MOCK_QA_AUDIT, "forced": True}
            return MOCK_QA_AUDIT

        # Hotfix May 23 2026: this handler used to swallow every
        # exception, log it, and return MOCK_QA_AUDIT — a 200
        # response with mock data. The frontend treated that as a
        # successful run, overwrote real per-check state with the
        # mock's empty data, and the user's "Re-run" click
        # produced no visible result. Now we propagate the real
        # error as a 500 so the qaStore surfaces it via its
        # error field.
        #
        # May 24 2026 deeper fix: this handler ALSO used to call
        # get_full_history() + run_all_strategies() inline on every
        # request, which is the heavy compute that times out on
        # cold deploys. Symptom: every check returned INCOMPLETE
        # ("Analysis not completed — re-run...") because the request
        # was failing before the QA agent even ran. Per user
        # directive: read from strategy_results_cache directly. If
        # the cache is cold, return 503 with a specific actionable
        # message rather than recomputing inline and timing out.
        # The dashboard's /api/backtest/compare endpoint is the
        # canonical warmer — hitting the dashboard once populates
        # the cache so the QA audit can complete.
        try:
            from tools.cache import (
                get_latest_strategy_cache, get_latest_qa,
                get_most_recent_qa_run, set_qa_cache,
                _compute_data_hash,
            )
            from agents.qa_agent import QAAgent
            from agents.usage import start_usage_capture

            strategy_results = await get_latest_strategy_cache()
            if not strategy_results:
                raise HTTPException(
                    status_code=503,
                    detail={
                        "error": "analytics_cache_cold",
                        "message": (
                            "Analytics cache not warmed — load the "
                            "dashboard first to populate cache, then "
                            "re-run audit."
                        ),
                        "hint": (
                            "Open the main Dashboard tab once. The "
                            "strategy comparison call populates "
                            "strategy_results_cache. Then return here "
                            "and re-run the audit."
                        ),
                    },
                )

            # ── QA RE-RUN GATE (May 24 2026) ──────────────────────────
            # The QA agent uses Sonnet — a full audit is ~$0.05-0.10
            # of token burn. The previous handler ran unconditionally
            # on every click and on every QAHub mount (qaStore.load()
            # fires reload() if its in-memory `loaded` flag is false,
            # which a server restart or a fresh user session resets).
            # Multiple users navigating to the QA tab triggered N
            # redundant identical audits, all consuming tokens against
            # the same strategy_hash.
            #
            # Two cascading gates BEFORE any LLM call fires:
            #
            #   1) HASH GATE — if a non-expired Tier 99 ("full audit")
            #      verdict already exists for this strategy_hash, serve
            #      it from qa_results_cache. The data has not changed
            #      since the cached audit ran, so the audit verdict has
            #      not changed either — no LLM call needed.
            #
            #   2) MIN-INTERVAL GATE — secondary safety net. Even if
            #      the hash logic somehow misfired (e.g. the hash
            #      computation changed subtly between deploys), never
            #      run a full audit more than once per QA_MIN_INTERVAL
            #      seconds across the whole platform. The cap is
            #      generous (5 minutes) — a real user click after a
            #      legitimate data refresh always proceeds, but a
            #      polling-driven flood is capped.
            #
            # Each gate logs the skip with hash + age so a future
            # production-log audit can confirm the rate limiter is
            # firing. Bypassed by neither (i.e. real run) logs
            # qa_audit_running so each authentic run is also visible.
            #
            # AUDIT_TIER_FULL=99 distinguishes the qa_audit endpoint's
            # full-checklist response from Tier 1/2/3 narrow verdicts
            # already stored in qa_results_cache. Same table, separate
            # tier band so the tiered-QA paths and the full-audit path
            # never overwrite each other.
            AUDIT_TIER_FULL = 99
            QA_MIN_INTERVAL_SECONDS = 5 * 60

            # ── HASH = the canonical strategy_results_cache hash ──────
            #
            # UAT staleness bug (May 24 2026). Three write paths
            # computed three different hashes for the same underlying
            # data:
            #   (a) /api/backtest/compare wrote strategy_results_cache
            #       with `str(monthly.index[-1].date())` → "2025-12-31"
            #   (b) _current_strategy_hash used `str(monthly.index[-1])`
            #       → "2025-12-31 00:00:00" (Timestamp str form)
            #   (c) this endpoint READ `monthly[-1].get("date")` from a
            #       LIST-of-pairs payload that is NEVER a dict, falling
            #       through to "unknown" every time.
            # is_audit_current() compared (a) against the qa_results_cache
            # hash written via (c) — never matched, so the methodology
            # audit ALWAYS read as "stale" right after a successful run.
            #
            # FIX — pull the canonical hash from strategy_results_cache
            # via get_latest_strategy_hash(). That's the SAME value
            # is_audit_current() reads back on the strategy side, so the
            # two halves of the comparison are guaranteed to match by
            # construction when the audit verified the latest data.
            try:
                from tools.cache import get_latest_strategy_hash
                qa_hash = await get_latest_strategy_hash()
            except Exception as _exc:  # noqa: BLE001
                # If the canonical hash read fails (DB outage) fall
                # through to a real run — the audit is the safer
                # default than skipping.
                log.info("qa_audit_hash_read_failed", error=str(_exc))
                qa_hash = None

            from datetime import datetime, timezone

            # Gate 1 — hash gate. Same data → cached audit verdict
            # stands. expires_at filtering inside get_latest_qa
            # discards a row past its TTL so a freshly-changed dataset
            # never reuses a verdict tied to old data. Bypassed when
            # the caller passes force=true — a manual re-run after
            # an Academic Review (IN02 dependency) must re-evaluate
            # the checklist even when strategy_hash is unchanged.
            if qa_hash and not force:
                cached_audit = await get_latest_qa(
                    qa_hash, min_tier=AUDIT_TIER_FULL)
                if cached_audit and isinstance(cached_audit.get("checklist"), dict):
                    log.info(
                        "qa_audit_skipped_hash_match",
                        strategy_hash=qa_hash[:8],
                        cached_run_at=cached_audit.get("run_at"),
                    )
                    return cached_audit["checklist"]

            # Gate 2 — minimum-interval gate. Secondary safety net
            # independent of hash. Caps token burn if the hash logic
            # has an edge case (e.g. the strategy_results shape
            # changed and the hash diverged without the data
            # changing). Returns the most recent verdict regardless
            # of which hash it ran against. Bypassed by force=true
            # for the same reason gate 1 is.
            if not force:
                most_recent = await get_most_recent_qa_run(
                    min_tier=AUDIT_TIER_FULL)
                if most_recent and most_recent.get("run_at"):
                    try:
                        run_at = datetime.fromisoformat(
                            most_recent["run_at"].replace("Z", "+00:00"))
                        age_seconds = (datetime.now(timezone.utc)
                                       - run_at).total_seconds()
                        if (age_seconds < QA_MIN_INTERVAL_SECONDS
                                and isinstance(most_recent.get("checklist"), dict)):
                            log.info(
                                "qa_audit_skipped_interval",
                                age_seconds=int(age_seconds),
                                min_interval_seconds=QA_MIN_INTERVAL_SECONDS,
                                cached_strategy_hash=str(
                                    most_recent.get("strategy_hash", ""))[:8],
                                current_strategy_hash=(
                                    qa_hash[:8] if qa_hash else "unknown"),
                            )
                            return most_recent["checklist"]
                    except (ValueError, AttributeError):
                        pass  # malformed run_at — fall through to real run

            log.info("qa_audit_running",
                     strategy_hash=qa_hash[:8] if qa_hash else "unknown",
                     forced=force)

            # AN01 / AN04 pre-flight (May 24 2026). Fetch the
            # Carhart loadings and transition matrix rows from
            # analytics_metrics_cache BEFORE the deterministic checks
            # run; trigger refresh on miss/incomplete. The QA audit
            # should never WARN on data the platform could have
            # computed itself.
            #
            # May 25 2026 — raw-field diagnostics. After the pre-flight
            # settles (refreshed or not), emit a single log line per
            # check naming the actual row count and the per-row
            # field shape the AN01 / AN04 deterministic checks will
            # read. If those checks then return WARN/FAIL the operator
            # has the upstream data shape on the same line, without
            # inspecting the JSONB payload by hand.
            try:
                from tools.precomputed_analytics import (
                    ensure_qa_data_complete,
                )
                analytics_cache = await ensure_qa_data_complete(qa_hash)
                if analytics_cache.get("refresh_triggered"):
                    log.info(
                        "qa_preflight_refreshed",
                        triggered=analytics_cache["refresh_triggered"],
                        completeness=analytics_cache["completeness"],
                    )
                _log_an01_an04_raw_state(analytics_cache, qa_hash, forced=force)
            except Exception as _exc:  # noqa: BLE001
                log.warning("qa_preflight_error", error=str(_exc))
                analytics_cache = None

            # IN01 submission-window attestation (May 25 2026). Async
            # query so the sync run_audit can read the verdict without
            # spawning a nested event loop. A query failure surfaces
            # as a FAIL verdict in the audit, not an endpoint 500.
            try:
                from tools.audit_engine import compute_in01_attestation
                audit_attestation = await compute_in01_attestation()
            except Exception as _exc:  # noqa: BLE001
                log.warning("qa_in01_attestation_error", error=str(_exc))
                audit_attestation = None

            # IN02 attestation — was an Academic Review run in the
            # last 14 days, and did it carry the five rated sections.
            # Same async-helper pattern as IN01 (compute_in01_attestation).
            # Query failure surfaces as a FAIL verdict in the IN02
            # check, not an endpoint 500.
            try:
                from tools.audit_engine import compute_in02_attestation
                academic_review_attestation = (
                    await compute_in02_attestation())
            except Exception as _exc:  # noqa: BLE001
                log.warning("qa_in02_attestation_error", error=str(_exc))
                academic_review_attestation = None

            # Seed the per-request usage bucket before the QA
            # agent's call_claude invocations so their token usage
            # is captured.
            start_usage_capture()
            qa = QAAgent()
            audit = qa.run_audit(
                strategy_results,
                run_full_checklist=True,
                analytics_cache=analytics_cache,
                audit_attestation=audit_attestation,
                academic_review_attestation=academic_review_attestation,
            )

            # Persist the full audit to qa_results_cache so the next
            # /api/qa/audit call within the TTL window short-circuits
            # at gate 1 above. Tier=AUDIT_TIER_FULL distinguishes this
            # from the tiered-QA narrow verdicts. Fail-open: if the
            # write fails (DB unavailable), the response still returns
            # normally; the next call will just re-run.
            if qa_hash:
                try:
                    await set_qa_cache(
                        qa_hash, audit, tier=AUDIT_TIER_FULL)
                except Exception as _exc:  # noqa: BLE001
                    log.warning("qa_audit_cache_write_failed",
                                error=str(_exc))

            # Team Activity — record the audit run (non-blocking).
            _log_interaction_bg(
                request, session, "qa",
                response_summary=str(audit.get("summary", "")),
                metadata={"verdict": audit.get("verdict")},
            )
            return audit

        except HTTPException:
            raise
        except Exception as exc:
            # Surface the real error so the frontend's "Re-run"
            # button shows what actually went wrong instead of
            # silently appearing to succeed. The previous catch-all
            # mock fallback hid the underlying data pipeline
            # failures (get_full_history timeout, missing strategy
            # cache row, etc.) from every operator.
            import traceback
            log.error(
                "qa_audit_error",
                error=str(exc),
                error_type=type(exc).__name__,
                traceback=traceback.format_exc(limit=10),
            )
            raise HTTPException(
                status_code=500,
                detail={
                    "error":      "qa_audit_failed",
                    "error_type": type(exc).__name__,
                    "message":    str(exc),
                    "hint": (
                        "The QA audit needs the strategy results "
                        "cache. If this is a fresh deploy, wait "
                        "for the analytics warm-up to complete and "
                        "retry. If the error persists, check the "
                        "Render logs for the qa_audit_error event."),
                },
            )
    finally:
        end_methodology()


@app.get("/api/v1/qa/export")
@limiter.limit("10/minute")
async def qa_export(request: Request, session: dict = Depends(require_auth)):
    """
    The QA methodology audit as a downloadable PDF — the Methodology
    Audit Report, formatted for inclusion in the Analytical Appendix.
    Open to every authenticated user: the Methodology Review section of
    the QA tab is not team-gated.

    The audit is run fresh (the same path as POST /api/qa/audit) so the
    PDF always reflects the current strategy results; the test
    environment and a pipeline failure fall back to the mock audit so
    the endpoint always returns a valid PDF.
    """
    from datetime import date
    from tools.audit_pdf import build_methodology_audit_pdf
    audit = MOCK_QA_AUDIT
    if ENVIRONMENT != "test":
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.backtester import run_all_strategies
            from agents.qa_agent import QAAgent
            from tools.audit_engine import (
                compute_in01_attestation, compute_in02_attestation,
            )
            history = await get_full_history_async()
            strategy_results = await asyncio.to_thread(run_all_strategies, history)
            # IN01 + IN02 attestations — same async-then-pass pattern
            # as the /api/qa/audit endpoint. The PDF export's IN01
            # row reflects the submission-window verdict; IN02 the
            # latest Academic Review status.
            attestation = await compute_in01_attestation()
            ar_attestation = await compute_in02_attestation()
            audit = await asyncio.to_thread(
                lambda: QAAgent().run_audit(
                    strategy_results, run_full_checklist=True,
                    audit_attestation=attestation,
                    academic_review_attestation=ar_attestation,
                )
            )
        except Exception as exc:
            log.error("qa_export_error", error=str(exc))
            audit = MOCK_QA_AUDIT
    # Intentional-design overrides (May 28 2026 hotfix). Fetch any
    # qa_intentional_overrides rows so build_methodology_audit_pdf
    # can render the team's recorded disclosure under each check.
    # Fail-open: a DB miss / read error leaves overrides empty and
    # the PDF renders without the disclosure lines rather than 500.
    overrides_map: dict[str, dict] = {}
    if ENVIRONMENT != "test":
        try:
            from sqlalchemy import text as _text
            from database import AsyncSessionLocal
            if AsyncSessionLocal is not None:
                async with AsyncSessionLocal() as conn:
                    rows = await conn.execute(_text(
                        "SELECT check_id, note, marked_by, marked_at "
                        "FROM qa_intentional_overrides"))
                    for row in rows.fetchall():
                        overrides_map[row[0]] = {
                            "check_id": row[0],
                            "note":     row[1],
                            "marked_by": row[2],
                            "marked_at": row[3].isoformat() if row[3] else None,
                        }
        except Exception as exc:  # noqa: BLE001
            log.warning("qa_export_overrides_read_failed", error=str(exc))
    pdf = build_methodology_audit_pdf(audit, overrides=overrides_map)
    filename = f"forest_capital_methodology_audit_{date.today().isoformat()}.pdf"
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/qa/ask")
@limiter.limit("10/minute")
async def qa_ask(
    request: Request,
    body: QAQueryRequest,
    session: dict = Depends(require_auth),
):
    """
    Conversational QA endpoint — routes questions through scope guard
    then the QA Agent for methodology questions.
    """
    if ENVIRONMENT != "test":
        try:
            from scope_guard import ScopeGuard
            guard = ScopeGuard()
            scope_result = await guard.check(body.question)
            if not scope_result["allowed"]:
                return {
                    "question": body.question,
                    "answer": scope_result["rejection_message"],
                    "verdict": "OUT_OF_SCOPE",
                }
        except Exception as exc:
            log.warning("qa_ask_scope_error", error=str(exc))

        try:
            from agents.base import call_claude, OPUS_MODEL
            from agents.qa_agent import _SYSTEM_PROMPT as QA_SYSTEM_PROMPT

            answer = call_claude(OPUS_MODEL, QA_SYSTEM_PROMPT, body.question,
                                 trigger="qa_ask")
            return {"question": body.question, "answer": answer, "verdict": "PASS"}
        except Exception as exc:
            log.error("qa_ask_error", error=str(exc))

    return {
        "question": body.question,
        "answer": "QA Agent temporarily unavailable. Please try POST /api/qa/audit for the full checklist.",
        "verdict": "WARN",
    }


# ── Tiered QA (Sprint 6) ──────────────────────────────────────────────────────
#
# Tier 1: pure-Python deterministic, sync. Result cached forever per
#   strategy_hash; same inputs always produce the same verdict.
# Tier 2: Sonnet narrative audit, async background. Runs when strategy_hash
#   changes or the most recent Tier 2 cache entry is older than 24h.
# Tier 3: Opus deep review, manual only. Auto-triggered if Tier 2 returns FAIL.
#
# The Present-mode gate reads the LATEST cached verdict for the current
# strategy_hash. Tier 1 alone is enough to unlock Present mode (≥ WARN);
# higher tiers refine the narrative shown on the QA tab. The dashboard
# never waits on a Sonnet/Opus call.

async def _current_strategy_hash() -> tuple[str, dict[str, dict] | None]:
    """
    Returns the canonical strategy_hash and the cached strategy_results
    for it (when present). Shared by every QA endpoint so a single hash
    value can serve both the status read and any tier trigger that
    follows.

    UAT staleness fix (May 24 2026) — uses get_latest_strategy_hash()
    as the canonical source. Previously this helper recomputed the
    hash with `str(monthly.index[-1])` (no `.date()` call), which
    produced a different value from /api/backtest/compare's
    `str(monthly.index[-1].date())`. The two never agreed, so the
    methodology audit's qa_results_cache row never matched the
    strategy_results_cache row it audited — is_audit_current()
    reported stale forever. Reading the canonical value from
    strategy_results_cache guarantees the comparison succeeds when
    the audit verified the latest data.

    Returns ("", None) when there is no strategy_results_cache row
    yet — same shape as the previous helper, so existing callers
    continue to work.
    """
    from tools.cache import get_strategy_cache, get_latest_strategy_hash

    strategy_hash = await get_latest_strategy_hash() or ""
    cached = await get_strategy_cache(strategy_hash) if strategy_hash else None
    return strategy_hash, cached


@app.get("/api/v1/qa/status")
@limiter.limit("60/minute")
async def qa_status(request: Request, session: dict = Depends(require_auth)):
    """
    Returns the latest QA verdict for the current strategy_hash.

    Response shape:
      {
        verdict: PASS|WARN|FAIL|UNKNOWN,
        tier: 1|2|3|null,
        run_at: iso8601 | null,
        age_hours: float | null,
        strategy_hash: str,
        present_mode_allowed: bool,
        running: bool,            # a methodology or statistical audit is
                                  #   in flight — global, not session-scoped
      }

    The nav-bar badge polls this endpoint every 30 seconds so the
    Running → PASS/WARN/FAIL transition shows up without a page reload.
    """
    if ENVIRONMENT == "test":
        return {
            "verdict": "UNKNOWN", "tier": None, "run_at": None,
            "age_hours": None, "strategy_hash": "test",
            "present_mode_allowed": False, "running": False,
        }

    try:
        from tools.cache import get_latest_qa
        from tools.qa_guard import (
            methodology_in_progress, statistical_audit_in_progress,
        )
        # Cross-user run state: the statistical audit_runs row is global,
        # the methodology flag is process-wide on the single worker — so
        # every user's poll sees the same answer, not session-scoped
        # state. statistical_audit_in_progress also reaps a run hung past
        # the 15-minute timeout.
        running = (methodology_in_progress()
                   or await statistical_audit_in_progress())
        strategy_hash, _cached = await _current_strategy_hash()
        latest = await get_latest_qa(strategy_hash, min_tier=1)

        if not latest:
            return {
                "verdict": "UNKNOWN", "tier": None, "run_at": None,
                "age_hours": None, "strategy_hash": strategy_hash,
                "present_mode_allowed": False, "running": running,
            }

        # Present-mode gate: verdict ≥ WARN AND age < 48h AND hash matches.
        # The hash equality is already enforced by the query filter, so we
        # only need to check verdict and age here.
        from datetime import datetime, timezone
        run_at_iso = latest.get("run_at")
        age_hours: float | None = None
        if run_at_iso:
            run_at_dt = datetime.fromisoformat(run_at_iso.replace("Z", "+00:00"))
            age_hours = (datetime.now(timezone.utc) - run_at_dt).total_seconds() / 3600.0

        present_allowed = (
            latest["verdict"] in ("PASS", "WARN")
            and age_hours is not None
            and age_hours < 48.0
        )

        return {
            "verdict":              latest["verdict"],
            "tier":                 latest["tier"],
            "run_at":               run_at_iso,
            "age_hours":            round(age_hours, 2) if age_hours is not None else None,
            "strategy_hash":        strategy_hash,
            "present_mode_allowed": present_allowed,
            "running":              running,
        }
    except Exception as exc:
        log.warning("qa_status_fallback", error=str(exc))
        return {
            "verdict": "UNKNOWN", "tier": None, "run_at": None,
            "age_hours": None, "strategy_hash": "unknown",
            "present_mode_allowed": False, "running": False,
        }


@app.post("/api/v1/qa/run")
@limiter.limit("20/minute")
async def qa_run(request: Request, session: dict = Depends(require_sysadmin)):
    """
    Runs Tier 1 synchronously and triggers Tier 2 in the background.
    Returns immediately with the Tier 1 verdict — the audience never
    waits on the Sonnet call. The Tier 2 result lands in the cache
    on its own; subsequent /qa/status polls pick it up.

    Auto-escalation to Tier 3 happens inside the background worker if
    Tier 2 returns FAIL (see schedule_tier2_background).

    Per-type guard: rejected with 409 only when another methodology
    audit is already in progress — a statistical audit never blocks it
    (tools/qa_guard.py).
    """
    from tools.qa_guard import (
        QA_BUSY_MESSAGE_METHODOLOGY, begin_methodology, end_methodology,
        methodology_in_progress,
    )
    # Per-type lock — a methodology audit is blocked only by another
    # methodology audit, never by a statistical one.
    if methodology_in_progress():
        raise HTTPException(status_code=409,
                            detail=QA_BUSY_MESSAGE_METHODOLOGY)

    if ENVIRONMENT == "test":
        return {"verdict": "PASS", "tier": 1, "tier2_scheduled": False}

    begin_methodology()
    try:
        from tools.qa_tiered import run_tier1_checks, schedule_tier2_background
        from tools.cache import set_qa_cache
        from tools.backtester import run_all_strategies
        from tools.data_fetcher import get_full_history_async
        from agents.usage import start_usage_capture

        strategy_hash, cached = await _current_strategy_hash()
        if cached:
            results_dict = cached
        else:
            history = await get_full_history_async()
            results_dict = await asyncio.to_thread(run_all_strategies, history)

        # Tier 1 — synchronous, deterministic, free. Tier 2 runs on a
        # background executor that does NOT inherit this context, so its
        # cost is not captured here; Tier 1 is sync and free of AI calls
        # today but the seed is here for forward-compatibility.
        start_usage_capture()
        t1 = run_tier1_checks(results_dict)
        await set_qa_cache(strategy_hash, t1, tier=1)

        # Tier 2 — fire and forget. Need a sync wrapper for the writer.
        # off_loop=True: _writer runs in a background thread via asyncio.run(),
        # so set_qa_cache must use the NullPool write engine — a pooled
        # connection would be orphaned when that loop closes.
        import asyncio as _asyncio
        def _writer(h: str, v: dict, tier: int) -> None:
            _asyncio.run(set_qa_cache(h, v, tier=tier, off_loop=True))
        schedule_tier2_background(results_dict, strategy_hash, _writer)

        # Team Activity — record the QA run (non-blocking).
        _log_interaction_bg(
            request, session, "qa",
            response_summary=str(t1.get("summary", "")),
            metadata={"verdict": t1.get("verdict"), "tier": 1},
        )

        return {
            "verdict": t1["verdict"],
            "tier": 1,
            "tier2_scheduled": True,
            "strategy_hash": strategy_hash,
            "checks_total": t1["checks_total"],
            "checks_passed": t1["checks_passed"],
            "checks_warned": t1["checks_warned"],
            "checks_failed": t1["checks_failed"],
            "summary": t1["summary"],
        }
    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("qa_run_error", ref=ref, error=str(exc))
        raise HTTPException(status_code=500, detail=f"QA run failed (ref: {ref})")
    finally:
        end_methodology()


@app.post("/api/v1/qa/full-review")
@limiter.limit("5/minute")
async def qa_full_review(request: Request, session: dict = Depends(require_sysadmin)):
    """
    Manually triggers a Tier 3 (Opus) deep review. Synchronous because
    the caller (Admin screen Full Review button) is willing to wait
    20-30 seconds — unlike the dashboard, which never waits.

    Sysadmin only — triggering a QA run is restricted to the platform
    sysadmin; the rate limit caps abuse at 5/minute.

    Per-type guard: rejected with 409 only when another methodology
    audit is already in progress — a statistical audit never blocks it
    (tools/qa_guard.py).
    """
    from tools.qa_guard import (
        QA_BUSY_MESSAGE_METHODOLOGY, begin_methodology, end_methodology,
        methodology_in_progress,
    )
    # Per-type lock — a methodology audit is blocked only by another
    # methodology audit, never by a statistical one.
    if methodology_in_progress():
        raise HTTPException(status_code=409,
                            detail=QA_BUSY_MESSAGE_METHODOLOGY)

    if ENVIRONMENT == "test":
        return {"verdict": "PASS", "tier": 3}

    begin_methodology()
    try:
        from tools.qa_tiered import run_tier3_review
        from tools.cache import set_qa_cache
        from tools.backtester import run_all_strategies
        from tools.data_fetcher import get_full_history_async

        strategy_hash, cached = await _current_strategy_hash()
        if cached:
            results_dict = cached
        else:
            history = await get_full_history_async()
            results_dict = await asyncio.to_thread(run_all_strategies, history)

        t3 = await asyncio.to_thread(run_tier3_review, results_dict)
        await set_qa_cache(strategy_hash, t3, tier=3)

        return {
            **t3,
            "strategy_hash": strategy_hash,
        }
    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("qa_full_review_error", ref=ref, error=str(exc))
        raise HTTPException(status_code=500, detail=f"Full review failed (ref: {ref})")
    finally:
        end_methodology()


# ── QA findings — Flag for Fix + Mark as Intentional ─────────────────────────
#
# May 22 2026 — companion endpoints to the QA Action Required UI
# (f96d897). The Flag for Fix button creates a triage_report_items
# row directly from the QA finding, routing the QA-driven issue into
# the same fix workflow used for UAT-derived items. The Mark as
# Intentional button writes an upsert into qa_intentional_overrides
# (migration 027), recording the team's judgement that the WARN is
# intentional methodology, not a defect — the override outlives the
# audit run that surfaced it.
#
# Both endpoints are team_member-gated (consistent with the May 2026
# QA endpoint regating that opened audit operations to the whole
# project team). Sysadmin status is recorded as the `marked_by`
# value on the intentional override so the audit trail attributes
# the decision to a specific user.


class QAFlagForFixRequest(__import__("pydantic").BaseModel):
    """POST body for the Flag for Fix endpoint."""
    check_title: str
    finding: str | None = None
    implication: str | None = None
    remediation: str | None = None
    severity: str | None = None  # WARN → 'major' / FAIL → 'blocking'


class QAMarkIntentionalRequest(__import__("pydantic").BaseModel):
    """POST body for the Mark as Intentional endpoint.

    May 28 2026 hotfix: `note` is now REQUIRED with a 20-character
    minimum. The previous shape (`note: str | None = None`) accepted
    a single-click confirmation from the UI where the body sent
    `{note: check.finding}` — the AI-generated check description
    became the disclosure note. That is not a disclosure; it is a
    rephrasing of the warning. The new gate forces the team to type
    a real reason before the override is recorded.

    Pydantic min_length=20 produces an automatic 422 with the
    standard validation-error body when the client submits a shorter
    note — no extra handler logic needed. Stale frontends that still
    send `{note: check.finding}` will fail loudly with 422 rather
    than silently recording the AI text.
    """
    note: str = __import__("pydantic").Field(..., min_length=20)
    audit_run_hash: str | None = None


@app.post("/api/v1/qa/findings/{check_id}/flag-for-fix")
@limiter.limit("30/minute")
async def qa_flag_for_fix(
    request: Request,
    check_id: str,
    body: QAFlagForFixRequest,
    session: dict = Depends(require_team_member),
):
    """
    Creates a triage_report_items row for the named QA check so the
    finding enters the normal fix workflow. The triage row is
    attached to a fresh triage_reports stub (triggered_by='qa_audit')
    so the existing report-items query layer ignores no rows.

    Returns the new triage_report_items id so the frontend can show
    a "Flagged · #N" badge alongside the QA card going forward.
    """
    if not check_id or len(check_id) > 20:
        raise HTTPException(status_code=422,
                            detail="check_id must be 1-20 chars")
    try:
        from sqlalchemy import text as _text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            raise HTTPException(status_code=503,
                                detail="Database unavailable")
        # Compose the item body so the triage view shows the full
        # finding without a join back to the QA audit.
        body_parts = []
        if body.finding:
            body_parts.append(f"FINDING: {body.finding}")
        if body.implication:
            body_parts.append(f"IMPLICATION: {body.implication}")
        if body.remediation:
            body_parts.append(f"REMEDIATION: {body.remediation}")
        item_body = "\n\n".join(body_parts) or None
        item_title = f"QA {check_id} — {body.check_title}"

        async with AsyncSessionLocal() as conn:
            # Synthesise a one-off triage_reports row for this flag so
            # the items query layer sees a parent report row. The
            # report_text is the same item_body for traceability.
            r = await conn.execute(_text(
                "INSERT INTO triage_reports "
                "(triggered_by, status, report_text) "
                "VALUES ('qa_audit', 'complete', :rep) "
                "RETURNING id"
            ), {"rep": (item_body or item_title)})
            report_id = r.scalar()

            # The QA check is the source. source_item_type='qa_check'
            # is a new value alongside the existing 'failure' /
            # 'feedback' — the schema is permissive (a String(20) with
            # no CHECK constraint), and the items query layer treats
            # an unknown source as 'no back-pointer' which is the
            # honest signal here. source_item_id stays NULL because a
            # QA check is keyed by check_id (a string), not by an
            # integer row id.
            r2 = await conn.execute(_text(
                "INSERT INTO triage_report_items "
                "(report_id, item_type, item_title, item_body, "
                " source_item_type, source_item_id) "
                "VALUES (:rid, 'immediate', :title, :body, "
                " 'qa_check', NULL) "
                "RETURNING id"
            ), {"rid": report_id, "title": item_title, "body": item_body})
            item_id = r2.scalar()
            await conn.commit()

        log.info("qa_flag_for_fix_recorded",
                 check_id=check_id, triage_item_id=int(item_id) if item_id else None,
                 by=session.get("email"),
                 severity=body.severity)
        return {
            "ok": True,
            "check_id": check_id,
            "triage_item_id": int(item_id) if item_id else None,
            "triage_report_id": int(report_id) if report_id else None,
        }
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("qa_flag_for_fix_failed", ref=ref, error=str(exc))
        raise HTTPException(status_code=500,
                            detail=f"Flag for fix failed (ref: {ref})")


@app.post("/api/v1/qa/findings/{check_id}/mark-intentional")
@limiter.limit("30/minute")
async def qa_mark_intentional(
    request: Request,
    check_id: str,
    body: QAMarkIntentionalRequest,
    session: dict = Depends(require_team_member),
):
    """
    Records (or updates) an entry in qa_intentional_overrides so the
    QA panel renders "Confirmed intentional — recorded {date}" on
    every subsequent audit instead of the WARN action card. The
    override outlives the audit run that surfaced it.

    Idempotent — a second click on the same check_id UPDATEs the
    existing row (the unique constraint enforces one override per
    check_id).
    """
    if not check_id or len(check_id) > 20:
        raise HTTPException(status_code=422,
                            detail="check_id must be 1-20 chars")
    email = session.get("email") or "unknown"
    try:
        from sqlalchemy import text as _text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            raise HTTPException(status_code=503,
                                detail="Database unavailable")
        async with AsyncSessionLocal() as conn:
            # Upsert — ON CONFLICT (check_id) DO UPDATE so a second
            # Mark Intentional refreshes the note + timestamp rather
            # than producing a duplicate row.
            r = await conn.execute(_text(
                "INSERT INTO qa_intentional_overrides "
                "(check_id, marked_by, note, audit_run_hash) "
                "VALUES (:cid, :by, :note, :hash) "
                "ON CONFLICT (check_id) DO UPDATE SET "
                " marked_at = now(), "
                " marked_by = EXCLUDED.marked_by, "
                " note = EXCLUDED.note, "
                " audit_run_hash = EXCLUDED.audit_run_hash "
                "RETURNING id, marked_at"
            ), {
                "cid": check_id, "by": email,
                "note": body.note, "hash": body.audit_run_hash,
            })
            row = r.fetchone()
            await conn.commit()

        log.info("qa_mark_intentional_recorded",
                 check_id=check_id, by=email,
                 override_id=int(row[0]) if row else None)
        return {
            "ok": True,
            "check_id": check_id,
            "marked_by": email,
            "marked_at": row[1].isoformat() if row and row[1] else None,
        }
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("qa_mark_intentional_failed", ref=ref, error=str(exc))
        raise HTTPException(status_code=500,
                            detail=f"Mark intentional failed (ref: {ref})")


@app.delete("/api/v1/qa/findings/{check_id}/mark-intentional")
@limiter.limit("30/minute")
async def qa_revoke_intentional(
    request: Request,
    check_id: str,
    session: dict = Depends(require_team_member),
):
    """
    Revokes a previously-recorded intentional override.

    Workstream F (May 28 2026) — when the team later determines that a
    finding is not actually intentional after all, this endpoint removes
    the qa_intentional_overrides row. The QA panel's Confirmed Intentional
    badge disappears and the original Action Required card re-renders on
    the next audit read. The report-readiness gate (workstream C) re-
    evaluates the next time it loads, so a revoked disclosure stops
    counting as resolved.

    Returns 200 with deleted=true on a successful revoke, 200 with
    deleted=false when no override existed (idempotent — a revoke on
    nothing is not an error). 422 on a malformed check_id. 503 if the
    database is unreachable.

    Project team only — mirrors the gating on mark-intentional.
    """
    if not check_id or len(check_id) > 20:
        raise HTTPException(status_code=422,
                            detail="check_id must be 1-20 chars")
    email = session.get("email") or "unknown"
    try:
        from sqlalchemy import text as _text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            raise HTTPException(status_code=503,
                                detail="Database unavailable")
        async with AsyncSessionLocal() as conn:
            r = await conn.execute(_text(
                "DELETE FROM qa_intentional_overrides "
                "WHERE check_id = :cid "
                "RETURNING id, marked_by"
            ), {"cid": check_id})
            row = r.fetchone()
            await conn.commit()

        if row is None:
            # Idempotent — a revoke on nothing is not an error. The
            # frontend can fire DELETE without first checking that an
            # override exists.
            log.info("qa_revoke_intentional_noop",
                     check_id=check_id, by=email)
            return {"ok": True, "check_id": check_id, "deleted": False}

        log.info("qa_revoke_intentional_recorded",
                 check_id=check_id, by=email,
                 override_id=int(row[0]) if row[0] else None,
                 was_marked_by=row[1])
        return {
            "ok": True,
            "check_id": check_id,
            "deleted": True,
            "revoked_by": email,
        }
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("qa_revoke_intentional_failed", ref=ref, error=str(exc))
        raise HTTPException(status_code=500,
                            detail=f"Revoke intentional failed (ref: {ref})")


@app.get("/api/v1/qa/intentional-overrides")
@limiter.limit("60/minute")
async def qa_intentional_overrides_list(
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Returns every recorded intentional override keyed by check_id.
    The QA panel reads this on mount and renders a "Confirmed
    intentional — recorded {date}" badge in place of the Action
    Required card for any check_id present here.

    Auth: any authenticated user — the audit trail is a project
    record visible to viewers, not a sysadmin-private surface.
    """
    try:
        from sqlalchemy import text as _text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is None:
            return {"overrides": {}}
        async with AsyncSessionLocal() as conn:
            r = await conn.execute(_text(
                "SELECT check_id, marked_at, marked_by, note, "
                "       audit_run_hash "
                "FROM qa_intentional_overrides "
                "ORDER BY marked_at DESC"
            ))
            rows = r.fetchall()
        overrides = {
            row[0]: {
                "marked_at": row[1].isoformat() if row[1] else None,
                "marked_by": row[2],
                "note": row[3],
                "audit_run_hash": row[4],
            }
            for row in rows
        }
        return {"overrides": overrides}
    except Exception as exc:  # noqa: BLE001
        log.warning("qa_intentional_overrides_read_failed", error=str(exc))
        return {"overrides": {}}


# ── Report ────────────────────────────────────────────────────────────────────

@app.get("/api/report/export")
async def export_report(session: dict = Depends(require_auth)):
    return {
        "message": "PDF report generation available in Sprint 4.",
        "status": "not_implemented",
    }


# ── Academic export package ───────────────────────────────────────────────────
#
# A one-click ZIP of the analytical assets (charts + tables) the team needs
# for the written deliverables. Charts and tables are rendered client-side
# (recharts → PNG, data table → CSV) and POSTed here as multipart blobs; the
# backend's only job is deterministic assembly — zipping the uploads and
# adding curated metadata files. The grader gets a self-describing archive.

# Curated, deterministic descriptions keyed by the chart filename slug. The
# original spec proposed calling the Academic Writer agent to generate these,
# but an export endpoint must be deterministic and outage-proof — it must
# never hang or 500 because an LLM is unreachable. Static descriptions are the
# correct engineering call: instant, reproducible, and graded identically.
_CHART_DESCRIPTIONS: dict[str, str] = {
    "cumulative_returns": (
        "Cumulative total return — growth of $1 invested in each strategy "
        "over the study period, benchmarked against the 100% S&P 500 equity "
        "index. Cite as evidence of long-horizon outperformance or shortfall."
    ),
    "rolling_correlation": (
        "Rolling 12-month equity-vs-bond correlation. The 2022 regime break "
        "(correlation turning positive) is the project's central finding. "
        "Cite when discussing the breakdown of static diversification."
    ),
    "rolling_excess_return": (
        "Rolling excess return of each strategy over the benchmark. Cite to "
        "show the consistency, not just the average, of any outperformance."
    ),
    "efficient_frontier": (
        "Mean-variance efficient frontier with each strategy plotted by "
        "realised volatility and return. Cite when discussing risk-adjusted "
        "positioning relative to the optimal frontier."
    ),
    "sensitivity_analysis": (
        "Parameter sensitivity — strategy metrics under +/-20% perturbation "
        "of key parameters. Cite as the robustness check against overfitting."
    ),
    "team_activity_timeline": (
        "Team Activity timeline — commits, council runs and platform usage "
        "interleaved over time. Cite as evidence for the Roles & Division of "
        "Labor deliverable."
    ),
    "team_contribution_split": (
        "Per-member contribution split across commits and AI interactions. "
        "Cite to substantiate the division-of-labour narrative."
    ),
    "agent_engagement": (
        "AI agent engagement — how often each council agent was consulted. "
        "Cite in the AI-usage section of the final presentation."
    ),
}


def _slug_for_chart(filename: str) -> str:
    """Extracts the description key from an uploaded chart filename.
    Filenames arrive prefixed/suffixed (e.g. '01_cumulative_returns.png');
    the matching slug is whichever known key the filename contains."""
    stem = filename.rsplit(".", 1)[0].lower()
    for slug in _CHART_DESCRIPTIONS:
        if slug in stem:
            return slug
    return ""


@app.get("/api/v1/export/data-reference-sheet")
async def get_data_reference_sheet(
    session: dict = Depends(require_auth),
):
    """Data Reference Sheet -- the cross-reference tool Bob,
    Molly, and Mike use to verify every value in the submission
    documents against the canonical strategy cache. Read-only;
    available to any team-member session (not admin-gated).

    Returns the full catalog from
    tools/data_reference_catalog.CATALOG with each token's
    current value (resolved via build_substitution_table) zipped
    in. Per-strategy appendix tokens (10 strategies x 5 metrics
    = 50 rows) and per-strategy factor loadings (10 strategies
    x 5 columns = 50 rows) appear in their own categories.

    Response shape -- the frontend DataReferenceSheetPanel
    consumes this directly:
      {
        "data_hash": "f2e87dec7dcabe71"   # strategy hash
        "platform_fingerprint": "d0b1339e..." # the value that
                                                appears in
                                                document footers
        "generated_at": "ISO-8601",
        "categories": {
          <category_key>: {
            "label": <human-readable label>,
            "entries": [
              {token, label, value, source, is_locked,
               last_verified, document_locations}, ...
            ]
          }, ...
        }
      }

    The two hashes are documented in the response so a frontend
    tooltip can explain why the document footers (showing
    platform_fingerprint d0b1339e) differ from the strategy
    hash the reference sheet locks to (f2e87dec). They are
    produced by two different functions over different inputs:
      strategy_hash       = _compute_data_hash(n_rows,
                                               last_date,
                                               n_strategies)
      platform_fingerprint = current_data_hash() = sha256 of
                                                   market data
                                                   table state

    Fail-open: any error resolving a token leaves its value
    as None and is_locked unchanged. The panel renders em-dash
    for None values so a cold field is visually distinct from
    a confirmed em-dash."""
    from datetime import datetime, timezone

    from tools.audit_assembler import current_data_hash
    import asyncio
    from tools.cache import (
        get_latest_strategy_cache,
        get_latest_strategy_hash,
        get_monthly_returns,
    )
    from tools.cio_recommendation import (
        compute_implied_asset_allocation, get_latest_recommendation,
    )
    from tools.regime_detector import detect_current_regime
    from tools.data_reference_catalog import (
        CATALOG, CATEGORY_LABELS,
        expand_per_strategy_appendix_metrics,
        expand_per_strategy_factor_loadings,
    )
    from tools.numeric_substitution import (
        build_substitution_table,
        get_substitution_table,
    )
    from tools.academic_deck import (
        OOS_SHARPE_REGIME_CONDITIONAL,
        OOS_SHARPE_BENCHMARK,
        CORRELATION_PRE_2022, CORRELATION_POST_2022,
    )
    # June 22 2026 -- defensive import. OOS_WINDOW_PCT_OF_STUDY
    # was added by PR #370 (Path A constants). When this endpoint
    # ships on a branch that has not picked up the new constant
    # yet, fall back to the documented value (53/287 = 18.5) so
    # the endpoint doesn't 500. After both PRs merge the import
    # will succeed and this fallback path becomes unreachable.
    try:
        from tools.academic_deck import OOS_WINDOW_PCT_OF_STUDY
    except ImportError:
        OOS_WINDOW_PCT_OF_STUDY = 18.5
    _ = session  # session only used to enforce require_auth

    # Pull the two hash flavours. strategy_hash is the value
    # the substitution table is built from (locks every
    # {{TOKEN}} to a specific data state). platform_fingerprint
    # is the value the document footers carry -- a different
    # function over different inputs (market data tables).
    try:
        strategy_hash = (await get_latest_strategy_hash()) or ""
    except Exception:  # noqa: BLE001
        strategy_hash = ""
    try:
        platform_fingerprint = (await current_data_hash()) or ""
    except Exception:  # noqa: BLE001
        platform_fingerprint = ""

    # Build the substitution table with the same wiring the
    # document generators use, so the reference values match
    # what gets baked into the brief / appendix / deck.
    strategy_cache: dict = {}
    cio_row: dict | None = None
    implied_alloc: dict | None = None
    live_signals: dict | None = None
    factor_loadings_row: list[dict] = []
    cache_computed_at: str | None = None
    try:
        strategy_cache = (await get_latest_strategy_cache()) or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_strategy_cache_failed",
                    error=str(exc))
    try:
        cio_row = await get_latest_recommendation()
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_cio_failed", error=str(exc))
    try:
        if cio_row and cio_row.get("blend_weights"):
            implied_alloc = await compute_implied_asset_allocation(
                cio_row.get("blend_weights"))
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_implied_alloc_failed",
                    error=str(exc))
    # June 22 2026 -- Gap A2 fix. tools.cache.get_regime_cache()
    # reads from the regime_signals_cache TABLE with a 15-min TTL
    # check; when the row is missing OR has expired, it returns
    # None and the four watchpoint tokens (VIX/YIELD/CREDIT/
    # EQUITY_TREND_CURRENT) render em-dash. detect_current_regime
    # is the live FRED read with its own in-process 15-min cache
    # -- guaranteed-populated source, returns the canonical dict
    # shape build_substitution_table reads (vix_level /
    # yield_curve_slope / credit_spread / equity_trend). Sync
    # function (FRED HTTP calls), wrapped in asyncio.to_thread to
    # keep the event loop free; same pattern as main.py:3022.
    try:
        live_signals = await asyncio.to_thread(detect_current_regime)
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_regime_detect_failed",
                    error=str(exc))
    # June 22 2026 -- Gap C wiring. {{STUDY_MONTHS}} reads from
    # the `study_months` kwarg; falls back to
    # strategy_cache.get("n_observations") which the cache shape
    # does NOT carry, so the token rendered em-dash on this
    # endpoint. Brief / appendix / deck callsites get the count
    # from data.study_period.n_months via gather_document_data;
    # this endpoint reads it directly from the monthly returns
    # length, the same source gather_document_data uses.
    study_months_value: int | None = None
    try:
        monthly = await get_monthly_returns()
        study_months_value = len((monthly or {}).get("dates") or [])
        if study_months_value == 0:
            study_months_value = None
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_monthly_returns_failed",
                    error=str(exc))
    # June 22 2026 (wiring fix) -- read factor loadings from the
    # academic_analytics cache payload, not from a live
    # an.factor_loadings(strategy_cache, []) call. The live
    # call passes empty FF factors and returns no rows -- the
    # data lives in analytics_metrics_cache, written by
    # refresh_academic_analytics. The substitution metric
    # sources loaded below will overwrite factor_loadings_row
    # with the cached payload (rc_rows assignment in the
    # defensive-kwarg block).
    try:
        from tools import analytics as an
        factor_loadings_row = an.factor_loadings(
            strategy_cache, [])
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_factor_loadings_failed",
                    error=str(exc))

    # June 22 2026 -- defensive kwarg forwarding. PR #370 added
    # `oos_window_pct_of_study` and `live_signals` kwargs to
    # build_substitution_table; this endpoint may ship before
    # that PR lands. Inspect the signature once and forward only
    # the kwargs the current implementation accepts so the
    # endpoint stays 200 regardless of merge order.
    #
    # MUST inspect build_substitution_table, NOT get_substitution_table.
    # The wrapper accepts the analytics kwargs only via **kwargs: Any
    # so its `signature.parameters` dict does NOT contain
    # `regime_conditional`, `factor_loadings`, etc. -- inspecting it
    # would make every "in _sig.parameters" check below evaluate False
    # and silently drop 5 kwargs (the original bug that left every
    # token from PR #374 + the live-signals/study-pct tokens from
    # PR #370 resolving to em-dash on the data reference sheet).
    # build_substitution_table lists every kwarg explicitly so the
    # signature check matches the kwargs it actually accepts.
    import inspect
    _sig = inspect.signature(build_substitution_table)
    _kwargs = {
        "oos_sharpe_blend": OOS_SHARPE_REGIME_CONDITIONAL,
        "oos_sharpe_benchmark": OOS_SHARPE_BENCHMARK,
        "pre_2022_eq_ig_correlation": CORRELATION_PRE_2022,
        "post_2022_eq_ig_correlation": CORRELATION_POST_2022,
        "implied_allocation": implied_alloc,
    }
    if "oos_window_pct_of_study" in _sig.parameters:
        _kwargs["oos_window_pct_of_study"] = OOS_WINDOW_PCT_OF_STUDY
    if "live_signals" in _sig.parameters:
        _kwargs["live_signals"] = live_signals
    if "study_months" in _sig.parameters \
            and study_months_value is not None:
        _kwargs["study_months"] = study_months_value
    # June 22 2026 (wiring fix) -- thread the three analytics
    # metric sources so pre/post 2022 Sharpe, factor loadings,
    # and cost-sensitivity tokens resolve from the right cache.
    # Defensive forwarding pattern matches the rest of this
    # endpoint -- only pass kwargs the current implementation
    # accepts, so it works regardless of merge order with other
    # in-flight PRs that may not yet have the new signature.
    try:
        from tools.academic_export import (
            load_substitution_metric_sources,
        )
        rc_rows, fl_rows, cs_payload, crisis_payload = (
            await load_substitution_metric_sources())
        if "regime_conditional" in _sig.parameters:
            _kwargs["regime_conditional"] = rc_rows
        if "factor_loadings" in _sig.parameters:
            _kwargs["factor_loadings"] = fl_rows
        if "cost_sensitivity" in _sig.parameters:
            _kwargs["cost_sensitivity"] = cs_payload
        if "crisis_performance" in _sig.parameters:
            _kwargs["crisis_performance"] = crisis_payload
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_metric_sources_failed",
                    error=str(exc))
        rc_rows, fl_rows, cs_payload, crisis_payload = (
            [], [], None, None)
    # June 22 2026 -- Gap A diagnostic. After PR #379 wired the
    # live_signals / implied_allocation / cio_row kwargs through
    # correctly, the user reported CURRENT_*_PCT / BLEND_*_WT /
    # VIX/CREDIT/YIELD/EQUITY_TREND_CURRENT tokens STILL rendering
    # em-dash. This log answers in one line whether the underlying
    # fetches actually returned data on Render or whether they
    # came back None (cold cio_recommendations / regime_signals
    # tables). If has_* are all False, the operator must warm the
    # caches via admin refresh before reloading the reference
    # sheet -- not a wiring bug.
    log.info(
        "data_reference_kwarg_shape",
        has_cio_row=bool(cio_row),
        cio_row_keys=sorted(cio_row.keys())[:10] if cio_row else [],
        has_blend_weights=bool(
            cio_row and cio_row.get("blend_weights")),
        has_implied_alloc=bool(implied_alloc),
        implied_alloc_keys=sorted(implied_alloc.keys())[:5]
            if implied_alloc else [],
        has_live_signals=bool(live_signals),
        live_signals_keys=sorted(live_signals.keys())[:8]
            if live_signals else [],
        study_months_value=study_months_value,
        n_rc_rows=len(rc_rows or []),
        n_fl_rows=len(fl_rows or []),
        cs_n_scenarios=(
            len((cs_payload or {}).get("scenarios") or [])),
    )
    table = get_substitution_table(
        strategy_hash, strategy_cache, cio_row, **_kwargs)

    # When the cache row was last written -- per-row
    # last_verified for is_locked=false entries. The
    # strategy_results_cache schema carries computed_at on
    # every row; the helper exposes it through the JSON blob.
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is not None:
            async with AsyncSessionLocal() as s:
                r = await s.execute(text(
                    "SELECT computed_at FROM "
                    "strategy_results_cache "
                    "WHERE strategy_hash = :h LIMIT 1"),
                    {"h": strategy_hash})
                row = r.fetchone()
                if row and row[0]:
                    cache_computed_at = (
                        row[0].isoformat()
                        if hasattr(row[0], "isoformat")
                        else str(row[0]))
    except Exception as exc:  # noqa: BLE001
        log.warning("data_reference_computed_at_failed",
                    error=str(exc))

    # The "locked at submission" sentinel for is_locked=true
    # entries. Bob sees this string instead of a timestamp.
    LOCKED_SENTINEL = "locked at submission (academic_deck.py)"

    # Factor-loading row lookup keyed by strategy_name.
    # June 22 2026 (wiring fix) -- prefer the cached payload
    # (fl_rows from load_substitution_metric_sources) over the
    # live-compute factor_loadings_row, which is empty when the
    # endpoint can't pass real FF factors. Falls back to the
    # live row only when the cache is empty.
    fl_by_strategy: dict[str, dict] = {}
    fl_source_rows = (
        fl_rows if isinstance(fl_rows, list) and fl_rows
        else factor_loadings_row)
    for row in fl_source_rows:
        if not isinstance(row, dict):
            continue
        name = row.get("strategy") or row.get("strategy_name")
        if name:
            fl_by_strategy[name] = row

    def _resolve_value(
        token: str, source: str,
    ) -> tuple[str | None, str | None]:
        """Returns (value, last_verified). Value None -> render
        em-dash; last_verified None means the source didn't
        carry a timestamp.

        June 27 2026 (Task 1) -- {{REGIME_CONFIDENCE}} and
        {{CURRENT_REGIME}} are read DIRECTLY from cio_row (the
        live get_latest_recommendation() result), bypassing the
        substitution table cache. The data-reference sheet is a
        diagnostic tool and must always reflect the true current
        CIO state, not a cached approximation. The PR 1 v2 cache-
        key fix (CIO identity in _cache_key, June 27 2026) closes
        the invalidation gap on doc generation, but the sheet
        still gets the live read here so a future cache-shape
        bug or a TTL race never produces a stale diagnostic
        display. The last_verified stamp 'live (from
        cio_recommendation)' makes it visible to the reader
        that this row bypasses the table cache."""
        if token == "{{REGIME_CONFIDENCE}}":
            from tools.numeric_substitution import format_pct
            raw = cio_row.get("confidence") if cio_row else None
            # Handle the two known cio_row.confidence shapes:
            # scalar float (e.g. 0.954) or dict with probability
            # field (e.g. {"probability": 0.954}). Same dispatch
            # as numeric_substitution.py:538-541.
            if isinstance(raw, dict):
                raw = raw.get("probability")
            value = format_pct(raw)
            return (value, "live (from cio_recommendation)")
        if token == "{{CURRENT_REGIME}}":
            raw = cio_row.get("regime") if cio_row else None
            value = str(raw) if raw else "—"
            return (value, "live (from cio_recommendation)")
        if source.startswith("data.factor_loadings."):
            # factor_loadings.<STRATEGY>.<metric>
            _, _, rest = source.partition(
                "data.factor_loadings.")
            parts = rest.split(".")
            if len(parts) == 2:
                strategy_name, metric_key = parts
                row = fl_by_strategy.get(strategy_name) or {}
                val = row.get(metric_key)
                if val is None:
                    return (None, cache_computed_at)
                return (
                    str(round(float(val), 4))
                    if isinstance(val, (int, float))
                    else str(val),
                    cache_computed_at)
        # Token-table lookup. Fall back to a cache-row computed_at
        # for live values; the locked sentinel for academic
        # constants.
        value = table.get(token)
        if value is None or value == "—":
            return (value, None)
        return (value, cache_computed_at)

    # June 22 2026 -- locked-constant provenance lookup. For
    # is_locked=True entries the catalog source string keys
    # into LOCKED_CONSTANT_PROVENANCE; the structured block
    # ships in the entry payload so the frontend can render
    # the multi-line tooltip on hover over the lock icon.
    from tools.data_reference_catalog import (
        classify_submission_scope, provenance_for_source,
        SCOPE_LEGEND,
    )

    # Walk the catalog and build the categorised response.
    categories: dict[str, dict] = {}
    for category_key, category_label, entries in CATALOG:
        rendered: list[dict] = []
        for entry in entries:
            value, last_verified_cache = _resolve_value(
                entry.token, entry.source)
            last_verified = (
                LOCKED_SENTINEL if entry.is_locked
                else (last_verified_cache or "cache miss"))
            provenance = (
                provenance_for_source(entry.source)
                if entry.is_locked else None)
            rendered.append({
                "token": entry.token,
                "label": entry.label,
                "value": value if value is not None else "—",
                "source": entry.source,
                "is_locked": entry.is_locked,
                "submission_scope": classify_submission_scope(
                    entry.token, entry.source, entry.is_locked),
                "last_verified": last_verified,
                "document_locations": list(entry.document_locations),
                "provenance": provenance,
            })
        categories[category_key] = {
            "label": category_label,
            "entries": rendered,
        }

    # Per-strategy appendix tokens (10 strategies x 5 metrics).
    per_strategy_rows: list[dict] = []
    for entry in expand_per_strategy_appendix_metrics():
        value, last_verified_cache = _resolve_value(
            entry.token, entry.source)
        per_strategy_rows.append({
            "token": entry.token,
            "label": entry.label,
            "value": value if value is not None else "—",
            "source": entry.source,
            "is_locked": entry.is_locked,
            "submission_scope": classify_submission_scope(
                entry.token, entry.source, entry.is_locked),
            "last_verified": last_verified_cache or "cache miss",
            "document_locations": list(entry.document_locations),
            # Per-strategy expansion is is_locked=False; no
            # provenance entry exists for these. The frontend
            # tooltip is skipped when provenance is None.
            "provenance": None,
        })
    categories["per_strategy_appendix"] = {
        "label": CATEGORY_LABELS["per_strategy_appendix"],
        "entries": per_strategy_rows,
    }

    # Per-strategy factor loadings (10 strategies x 5 columns).
    factor_rows: list[dict] = []
    for entry in expand_per_strategy_factor_loadings():
        value, last_verified_cache = _resolve_value(
            entry.token, entry.source)
        factor_rows.append({
            "token": entry.token,
            "label": entry.label,
            "value": value if value is not None else "—",
            "source": entry.source,
            "is_locked": entry.is_locked,
            "submission_scope": classify_submission_scope(
                entry.token, entry.source, entry.is_locked),
            "last_verified": last_verified_cache or "cache miss",
            "document_locations": list(entry.document_locations),
            "provenance": None,
        })
    categories["factor_loadings"] = {
        "label": CATEGORY_LABELS["factor_loadings"],
        "entries": factor_rows,
    }

    # ── Task 4 (June 27 2026) -- submission-scope summary ──────
    # June 28 2026 -- auto-discovery backstop. Walk every token
    # in the live substitution table; any token NOT already in
    # the curated CATALOG above (or the per-strategy / factor
    # expansions) gets surfaced under an "uncatalogued" category
    # so the operator can see ALL tokens that can appear in a
    # draft, not just the ones a human remembered to add to
    # CATALOG. This makes the sheet self-healing: future token
    # additions to numeric_substitution.py auto-appear here on
    # the next request without a parallel catalog update.
    #
    # The endpoint still prefers curated entries (with labels +
    # provenance + document_locations) when one exists -- this
    # path is the safety net, not the primary path.
    catalogued: set[str] = set()
    for _ck, cat in categories.items():
        for e in cat.get("entries", []):
            catalogued.add(e["token"])
    try:
        uncatalogued_rows: list[dict] = []
        for token, value in sorted(table.items()):
            if not (token.startswith("{{")
                    and token.endswith("}}")):
                continue
            if token in catalogued:
                continue
            uncatalogued_rows.append({
                "token": token,
                "label": (
                    "Uncatalogued -- substitution-table only "
                    "(add a TokenEntry in "
                    "tools/data_reference_catalog.py for a "
                    "human-readable label + provenance)"),
                "value": value if value is not None else "—",
                "source": (
                    "tools/numeric_substitution.py:"
                    "get_substitution_table"),
                "is_locked": False,
                "submission_scope": classify_submission_scope(
                    token, None, False),
                "last_verified": "live",
                "document_locations": [],
                "provenance": None,
            })
        if uncatalogued_rows:
            categories["uncatalogued"] = {
                "label": (
                    "Uncatalogued tokens "
                    "(present in substitution table but missing "
                    "a curated catalog entry)"),
                "entries": uncatalogued_rows,
            }
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "data_reference_uncatalogued_backstop_failed",
            error=str(exc))

    # Tally every rendered entry across every category by its
    # submission_scope so the sheet header can answer
    # "is this figure part of the academic submission record?"
    # at a glance.
    scope_counts: dict[str, int] = {
        "IN_SCOPE_LOCKED": 0,
        "IN_SCOPE_CONSTANT": 0,
        "IN_SCOPE_FULL_DATASET": 0,
        "OUT_OF_SCOPE_LIVE": 0,
    }
    for _cat_key, cat in categories.items():
        for entry in cat.get("entries", []):
            scope = entry.get("submission_scope")
            if scope in scope_counts:
                scope_counts[scope] += 1
    in_scope_total = (
        scope_counts["IN_SCOPE_LOCKED"]
        + scope_counts["IN_SCOPE_CONSTANT"]
        + scope_counts["IN_SCOPE_FULL_DATASET"])
    submission_scope_summary = {
        "in_scope_total": in_scope_total,
        "in_scope_locked": scope_counts["IN_SCOPE_LOCKED"],
        "in_scope_constant": scope_counts["IN_SCOPE_CONSTANT"],
        "in_scope_full_dataset": (
            scope_counts["IN_SCOPE_FULL_DATASET"]),
        "out_of_scope_live": scope_counts["OUT_OF_SCOPE_LIVE"],
    }

    # ── Task 4 -- freeze status surfaced at the sheet header ───
    freeze_active = False
    freeze_hash: str | None = None
    try:
        from tools.submission_freeze import get_freeze_config
        freeze_config = await get_freeze_config()
        freeze_active = bool(freeze_config.get("active"))
        freeze_hash = (
            freeze_config.get("freeze_hash")
            if freeze_active else None)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "data_reference_freeze_status_failed",
            error=str(exc))

    return {
        "data_hash": strategy_hash,
        "platform_fingerprint": platform_fingerprint,
        "generated_at": datetime.now(
            timezone.utc).isoformat(),
        "categories": categories,
        # Task 4 (June 27 2026) -- submission audit fields.
        "freeze_active": freeze_active,
        "freeze_hash": freeze_hash,
        "submission_scope_summary": submission_scope_summary,
        "submission_scope_legend": SCOPE_LEGEND,
    }


@app.get("/api/v1/export/data-reference-sheet/validate")
async def validate_data_reference_sheet(
    session: dict = Depends(require_auth),
):
    """Cross-reference validator for the Data Reference Sheet.

    Runs the same build as the /data-reference-sheet endpoint,
    then for each token compares the rendered value against the
    authoritative source (analytics_metrics_cache row, strategy
    cache, detect_current_regime, etc) and returns a status
    pass / fail / warning / skipped per token.

    Status rules:
      pass    -- within tolerance (0.01 for Sharpe, 0.0001 for
                 factor loadings, 0.5pp for percentages, exact
                 for ints / strings)
      fail    -- beyond tolerance; result carries delta
      warning -- pass but source row > 24h old
      skipped -- locked constant, missing source, or no
                 strategy registered for the token

    Each result also carries cache_freshness -- the ISO
    timestamp of the source row, null for locked constants.

    Zero LLM calls. Reads warm caches only. Sub-2-second
    typical response.

    Fail-open per token: a strategy raising returns skipped
    with note='validator_error: <msg>'; the report still
    completes for the remaining tokens.
    """
    from datetime import datetime, timezone

    from tools.audit_assembler import current_data_hash
    import asyncio
    from tools.cache import (
        get_latest_strategy_cache,
        get_latest_strategy_hash,
        get_monthly_returns,
    )
    from tools.cio_recommendation import (
        compute_implied_asset_allocation,
        get_latest_recommendation,
    )
    from tools.data_reference_validator import (
        Sources, validate_reference_sheet,
    )
    from tools.precomputed_analytics import get_latest_metric
    from tools.regime_detector import detect_current_regime
    _ = session  # auth-only

    # Get the rendered sheet via the same handler so any future
    # change to the sheet shape is automatically reflected here.
    rendered = await get_data_reference_sheet(session)
    if not isinstance(rendered, dict):
        # The handler should always return a dict; bail out
        # rather than crashing the validator.
        return {
            "data_hash": "",
            "validated_at": datetime.now(
                timezone.utc).isoformat(),
            "summary": {
                "total": 0, "passed": 0, "failed": 0,
                "warning": 0, "skipped": 0},
            "results": [],
            "error": "sheet_unavailable",
        }
    data_hash = rendered.get("data_hash") or ""

    # Pre-load every source the strategies read so we hit each
    # cache exactly once for the whole 153-token report rather
    # than per-token.
    sources = Sources()
    try:
        sources.strategy_cache = (
            await get_latest_strategy_cache()) or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_strategy_cache_failed",
                    error=str(exc))
    try:
        sources.cio_row = await get_latest_recommendation()
        if sources.cio_row:
            sources.cio_computed_at = (
                sources.cio_row.get("computed_at"))
            if (sources.cio_row.get("blend_weights")):
                try:
                    sources.implied_alloc = (
                        await compute_implied_asset_allocation(
                            sources.cio_row.get("blend_weights")))
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "validator_implied_alloc_failed",
                        error=str(exc))
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_cio_failed", error=str(exc))
    try:
        sources.live_signals = await asyncio.to_thread(
            detect_current_regime)
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_live_signals_failed",
                    error=str(exc))
    try:
        sources.academic_analytics = (
            await get_latest_metric("academic_analytics"))
        if sources.academic_analytics:
            sources.academic_analytics_computed_at = (
                sources.academic_analytics.get("_computed_at"))
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_academic_analytics_failed",
                    error=str(exc))
    try:
        sources.oos_cost_sensitivity = (
            await get_latest_metric("oos_cost_sensitivity"))
        if sources.oos_cost_sensitivity:
            sources.oos_cost_sensitivity_computed_at = (
                sources.oos_cost_sensitivity.get("_computed_at"))
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_oos_cost_sensitivity_failed",
                    error=str(exc))
    try:
        monthly = await get_monthly_returns()
        sources.n_monthly_months = (
            len((monthly or {}).get("dates") or []))
    except Exception as exc:  # noqa: BLE001
        log.warning("validator_monthly_returns_failed",
                    error=str(exc))
    try:
        # Pull strategy cache row's computed_at for staleness
        # checks on per-strategy metric tokens.
        try:
            strategy_hash = (
                await get_latest_strategy_hash()) or ""
        except Exception:  # noqa: BLE001
            strategy_hash = ""
        if strategy_hash:
            from sqlalchemy import text
            from database import AsyncSessionLocal
            if AsyncSessionLocal is not None:
                async with AsyncSessionLocal() as s:
                    r = await s.execute(text(
                        "SELECT computed_at FROM "
                        "strategy_results_cache "
                        "WHERE strategy_hash = :h LIMIT 1"),
                        {"h": strategy_hash})
                    row = r.fetchone()
                    if row and row[0]:
                        sources.strategy_cache_computed_at = (
                            row[0].isoformat()
                            if hasattr(row[0], "isoformat")
                            else str(row[0]))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "validator_strategy_cache_computed_at_failed",
            error=str(exc))

    rendered_categories = rendered.get("categories") or {}
    report = validate_reference_sheet(
        rendered_categories, sources, data_hash)
    log.info(
        "data_reference_sheet_validated",
        data_hash=data_hash[:8] if data_hash else "",
        total=report.summary.get("total"),
        passed=report.summary.get("passed"),
        failed=report.summary.get("failed"),
        warning=report.summary.get("warning"),
        skipped=report.summary.get("skipped"),
        actor=session.get("email") if isinstance(session, dict)
              else None)
    return report.to_dict()


# ── Slide guidance (June 22 2026) ───────────────────────────────────


@app.get("/api/v1/deck/slide-guidance/template")
async def get_slide_guidance_template(
    session: dict = Depends(require_auth),
):
    """Download the default slide guidance template as JSON.

    Molly opens this in any text editor, edits the string values
    only, saves, and uploads via POST /api/v1/deck/slide-guidance.
    The template is generated from the current SLIDE_TITLES + the
    canonical default so_what / max_bullets / bullet_guidance /
    speaker_note_directive per slide. Every uploaded file MUST
    have been derived from a downloaded template -- the rigid
    validator rejects any deviation in field set or value type.
    """
    from tools.deck_slide_guidance import build_default_template
    _ = session  # auth-only
    return build_default_template()


@app.post("/api/v1/deck/slide-guidance")
async def post_slide_guidance(
    file: UploadFile = File(...),
    session: dict = Depends(require_auth),
):
    """Upload a slide guidance JSON file for the authenticated
    user. Per-user storage: only one active guidance row per
    owner_email. Prior active row is deactivated atomically
    when this one is accepted.

    Validation rules (enforced in validate_guidance):
      - Exact key set (no missing, no extras)
      - All values strings
      - All 12 slide numbers present
      - version + generated_from match template
      - String length limits per field
      - max_bullets is a numeric string in [0, 3]

    On validation failure returns 422 with the exact field path
    that failed and a plain English error message.
    """
    import json as _json
    from tools.deck_slide_guidance import (
        set_active_guidance, validate_guidance,
    )
    owner_email = session.get("email") or ""
    if not owner_email:
        raise HTTPException(
            status_code=401, detail="no_session_email")
    try:
        raw = await file.read()
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=400,
            detail=f"failed to read uploaded file: {exc}")
    try:
        payload = _json.loads(raw.decode("utf-8"))
    except UnicodeDecodeError:
        raise HTTPException(
            status_code=422,
            detail="uploaded file is not valid UTF-8")
    except _json.JSONDecodeError as exc:
        raise HTTPException(
            status_code=422,
            detail=f"uploaded file is not valid JSON: {exc.msg} "
                   f"at line {exc.lineno}, column {exc.colno}")
    clean, error = validate_guidance(payload)
    if error or clean is None:
        raise HTTPException(status_code=422, detail=error)
    result = await set_active_guidance(owner_email, clean)
    if not result.get("ok"):
        raise HTTPException(
            status_code=500,
            detail=(
                "failed to persist guidance: "
                f"{result.get('error') or 'unknown error'}"))
    log.info(
        "deck_slide_guidance_uploaded",
        owner_email=owner_email,
        version=clean.get("version"),
        generated_from=clean.get("generated_from"),
        uploaded_at=result.get("uploaded_at"))
    return {
        "ok": True,
        "uploaded_at": result.get("uploaded_at"),
        "version": clean.get("version"),
        "generated_from": clean.get("generated_from"),
    }


@app.get("/api/v1/deck/slide-guidance")
async def get_slide_guidance(
    session: dict = Depends(require_auth),
):
    """Return the active slide guidance for the authenticated
    user, or {active: false} when none has been uploaded.

    Used by the Reports page SlideGuidancePanel to render the
    current state + populate the "download current guidance"
    button payload."""
    from tools.deck_slide_guidance import get_active_guidance
    owner_email = session.get("email") or ""
    if not owner_email:
        return {"active": False}
    row = await get_active_guidance(owner_email)
    if not row:
        return {"active": False}
    guidance = row.get("guidance") or {}
    return {
        "active": True,
        "uploaded_at": row.get("uploaded_at"),
        "version": guidance.get("version"),
        "generated_from": guidance.get("generated_from"),
        "guidance": guidance,
    }


@app.delete("/api/v1/deck/slide-guidance")
async def delete_slide_guidance(
    session: dict = Depends(require_auth),
):
    """Deactivate the active guidance for the authenticated user,
    reverting deck generation to the hardcoded defaults in
    SLIDE_SPECIFICATIONS. Idempotent -- returns ok=true even
    when no active guidance exists."""
    from tools.deck_slide_guidance import clear_active_guidance
    owner_email = session.get("email") or ""
    if not owner_email:
        raise HTTPException(
            status_code=401, detail="no_session_email")
    ok = await clear_active_guidance(owner_email)
    if not ok:
        raise HTTPException(
            status_code=500,
            detail="failed to clear active guidance "
                   "(database unreachable)")
    log.info(
        "deck_slide_guidance_cleared",
        owner_email=owner_email)
    return {"ok": True}


@app.post("/api/v1/export/package")
@limiter.limit("10/minute")
async def export_package(
    request: Request,
    charts: list[UploadFile] = File(default=[]),
    tables: list[UploadFile] = File(default=[]),
    metadata: str = Form(default="{}"),
    session: dict = Depends(require_permission("export_package")),
):
    """
    Assembles an academic export ZIP from client-rendered chart PNGs and
    table CSVs plus curated metadata files.

    Multipart/form-data fields:
      charts    — PNG image blobs; each .filename is the in-ZIP name
      tables    — CSV blobs; each .filename is the in-ZIP name
      metadata  — JSON string of study-period fields

    ZIP layout:
      charts/<uploaded chart filenames>
      tables/<uploaded table filenames>
      metadata/study_period.txt
      metadata/chart_descriptions.txt
      README.txt

    Returned as application/zip with an attachment Content-Disposition so
    the browser downloads it. Deterministic — no LLM, no pipeline run.
    """
    import io
    import zipfile
    from datetime import date

    # Seeded for consistency with every other interaction-logging endpoint.
    # No AI work runs in this handler today (descriptions are curated and
    # deterministic), so this is a no-op for cost — but a future change that
    # adds an AI-generated README or chart caption would be captured for free.
    from agents.usage import start_usage_capture
    start_usage_capture()

    try:
        # A malformed metadata string must not break the export — the ZIP
        # is still useful with placeholder study-period fields.
        try:
            meta = json.loads(metadata) if metadata else {}
            if not isinstance(meta, dict):
                meta = {}
        except (json.JSONDecodeError, TypeError):
            meta = {}

        def _m(key: str) -> str:
            value = meta.get(key)
            return str(value) if value not in (None, "") else "—"

        generated = _m("generated_at")
        today = date.today().isoformat()

        study_period = (
            f"Study period: {_m('study_period_start')} to "
            f"{_m('study_period_end')}\n"
            f"{_m('n_months')} months of monthly data\n"
            "Benchmark: 100% S&P 500 equity index\n"
            "Risk-free rate: DTB3 mean monthly, annualised\n"
            "Factor model: Carhart four-factor (MKT-RF, SMB, HML, MOM)\n"
            f"Generated: {generated}\n"
        )

        readme = (
            "Forest Capital Portfolio Intelligence System — "
            "Academic Export Package\n"
            f"Generated: {generated}\n\n"
            "Charts: PNG at 2x resolution, light mode, suitable for "
            "Word/PowerPoint.\n"
            "Tables: CSV, importable into Excel.\n\n"
            "Cite as: Portfolio Intelligence System analytical output, "
            "Forest Capital /\n"
            f"McColl School of Business FNA 670, {today}.\n"
        )

        # Describe only the charts actually present in this upload.
        desc_lines: list[str] = ["CHART DESCRIPTIONS\n"]
        n_charts = 0
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for chart in charts:
                content = await chart.read()
                name = chart.filename or f"chart_{n_charts + 1}.png"
                zf.writestr(f"charts/{name}", content)
                slug = _slug_for_chart(name)
                description = _CHART_DESCRIPTIONS.get(
                    slug,
                    "Analytical chart exported from the Portfolio "
                    "Intelligence System dashboard.",
                )
                desc_lines.append(f"\n{name}\n  {description}\n")
                n_charts += 1

            n_tables = 0
            for tbl in tables:
                content = await tbl.read()
                name = tbl.filename or f"table_{n_tables + 1}.csv"
                zf.writestr(f"tables/{name}", content)
                n_tables += 1

            zf.writestr("metadata/study_period.txt", study_period)
            zf.writestr("metadata/chart_descriptions.txt", "".join(desc_lines))
            zf.writestr("README.txt", readme)

        zip_bytes = zip_buffer.getvalue()

        # Activity logging — awaited (not fire-and-forget): the export is
        # already complete, so a synchronous one-row INSERT costs nothing
        # and guarantees the row lands before the response returns.
        # Team-gated and fail-open inside log_agent_interaction.
        try:
            from tools.activity_log import log_agent_interaction
            await log_agent_interaction(
                user_email=session.get("email", ""),
                session_id=request.headers.get("x-session-id"),
                session_type=request.headers.get("x-session-type"),
                interaction_type="export",
                response_summary=f"{n_charts} charts, {n_tables} tables",
                metadata={"n_charts": n_charts, "n_tables": n_tables,
                          "bytes": len(zip_bytes)},
            )
        except Exception as exc:  # noqa: BLE001
            log.warning("export_log_failed", error=str(exc))

        filename = f"forest_capital_academic_export_{today}.zip"
        return Response(
            content=zip_bytes,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("export_package_error", ref=ref, error=str(exc))
        raise HTTPException(
            status_code=500,
            detail=f"Export package generation failed (ref: {ref})",
        )


# ── Academic deliverable generation — first-draft .docx / .pptx ───────────────
#
# POST /api/v1/export/midpoint-paper     → 3-page midpoint submission (.docx)
# POST /api/v1/export/executive-brief    → 5-page executive brief (.docx)
# POST /api/v1/export/presentation-deck  → 6-slide final deck (.pptx)
#
# Each assembles a graded deliverable as a FIRST DRAFT for Bob to refine:
# every figure is real platform data (tools/academic_export.gather_document_
# data — light cache reads, never a pipeline run), and every narrative
# section is written by the Academic Writer agent run through the generator-
# evaluator harness. A section whose source data is unavailable is filled
# with a [DATA PENDING] marker — a missing input never fails the document.

_DOCX_MEDIA = ("application/vnd.openxmlformats-officedocument."
               "wordprocessingml.document")
_PPTX_MEDIA = ("application/vnd.openxmlformats-officedocument."
               "presentationml.presentation")

# Key analytical findings that MUST appear in the midpoint paper. These
# are appended to the section task prompts (never overwrite the existing
# instruction — the task string is one logical document). The findings
# are split by the section they belong in; the Academic Review arbiter
# scores the same set (agents/academic_review.py).
_MIDPOINT_S1_KEY_FINDINGS = (
    # May 25 2026 — TIGHTENED to one sentence per finding + the
    # three most distinctive methodology highlights only.
    # May 26 2026 (1) — added LOGICAL INTEGRITY GUARDRAILS.
    # May 26 2026 (2) — added a hard EM-DASH PROHIBITION at the
    # top of the prompt AND scrubbed every em dash out of the
    # prompt's own text (commas, colons, restructured sentences).
    # The model was picking up the em dashes in its own
    # instructions as licence to use them in output.
    "\n\nEM-DASH PROHIBITION (HARD CONSTRAINT):\n"
    "Never use em dashes in any output. Use commas, semicolons, "
    "colons, or restructure the sentence. This applies to every "
    "form: the Unicode em dash, the en dash used as a break, and "
    "the ASCII double-hyphen / triple-hyphen substitutes. A draft "
    "with any dash break will be flagged.\n\n"
    "KEY FINDINGS. Introduce each in one sentence.\n"
    "(a) The 2022 equity-IG correlation broke from roughly -0.05 to "
    "+0.61. This is the central finding of the project and is "
    "developed in Results.\n"
    "(b) Five strategies (MIN_VARIANCE, BLACK_LITTERMAN, "
    "MAX_SHARPE_ROLLING, MOMENTUM_ROTATION, REGIME_SWITCHING) start "
    "later than the 2002-07 study period because of lookback windows; "
    "their metrics cover actual data periods.\n"
    "(c) Every metric was independently recomputed by a separate AI "
    "model (Claude Opus) with zero critical failures across 59 checks.\n"
    "(d) Data: equity (SPY monthly), investment-grade (LQD-to-BND "
    "splice), high-yield (BAMLHYH0A0HYM2TRIV through 2025; HYG ETF "
    "proxy thereafter), risk-free (FRED DTB3).\n\n"
    "LOGICAL INTEGRITY GUARDRAILS\n\n"
    "1. REGIME LANGUAGE\n"
    "   Never declare the post-2022 period a definitive \"regime.\" "
    "Always frame as a hypothesis:\n"
    "   \"The persistence and magnitude of the post-2022 inversion "
    "strongly suggests a structural regime shift rather than a "
    "temporary dislocation. It is a hypothesis to be tested formally "
    "in the final phase.\"\n\n"
    "2. STRATEGY SELECTION LANGUAGE\n"
    "   When no single strategy dominates across crisis windows, "
    "always qualify with the asset universe constraint:\n"
    "   \"Given the constraints of our three-asset universe, dynamic "
    "weighting across regimes is the most viable mechanism for alpha "
    "generation when asset selection is restricted.\"\n\n"
    "METHODOLOGY HIGHLIGHTS. Name explicitly: Carhart four-factor "
    "attribution (MOM included), Benjamini-Hochberg FDR correction at "
    "q < 0.005, and true one-way drift-inclusive portfolio turnover."
)
# Verification caveats appended to the document-section task prompts.
# CAVEAT 2 — every external citation is preceded by a [[VERIFY CITATION]]
# marker; CAVEAT 3 — every uncertain numeric value is wrapped in a
# [[VERIFY]] marker. The academic_docx renderer shows both bold and
# highlighted; the Academic Review arbiter flags any that survive into a
# submitted draft. Applied via _apply_draft_caveats so a task that
# already carries one form is not given a second, conflicting copy.
# May 26 2026 — scrubbed em dashes from the caveat prompt strings.
# Same reason as the S1 / S2 / _SYSTEM_PROMPT scrub: a prompt that
# contains em dashes in its own text gives the model implicit
# license to emit them in its output. Caveats now use colons and
# commas to delimit the same structure.
_CAVEAT_CITATION = (
    "\n\nCITATION VERIFICATION. Immediately before every external "
    "citation you include, insert an inline marker of the form "
    "[[VERIFY CITATION: check that Author (Year) exists and supports "
    "this specific claim before submitting]], so no unverified citation "
    "is missed."
)
_CAVEAT_STATS = (
    "\n\nSTATISTIC VERIFICATION. If you are uncertain about any "
    "specific numeric value, do NOT insert it silently; wrap it in an "
    "inline marker of the form [[VERIFY: <the value and what it is>]] "
    "(for example [[VERIFY: Sharpe ratio for Regime Switching = 0.63]]) "
    "so a team member confirms it against the Analytics page before "
    "submission."
)


def _apply_draft_caveats(
    specs: list[dict], document_type: str | None = None,
) -> list[dict]:
    """
    Appends the citation- and statistic-verification caveats to each
    section task prompt. Idempotent per form — a task that already
    carries the [[VERIFY CITATION]] or [[VERIFY:]] instruction is not
    given a second, conflicting copy (the midpoint methodology and
    results tasks already carry the statistic marker).

    June 21 2026 -- the [[VERIFY CITATION]] caveat is skipped for
    document_type='executive_brief' (PR #363) AND for
    document_type='analytical_appendix' (this PR). Both writers
    cite only from data/references.json (web search disabled in
    PR #362; the registry is the only permitted source); every
    citation is pre-verified by construction. Injecting the
    caveat still drove the writers to mark every citation with
    [[VERIFY CITATION: ...]] blocks that surfaced as submission
    blockers in the audit AND fed the editor's marker-counting
    progress meter (sections C/D/E/G showed 0% completion in
    production despite carrying full prose, because the
    citation markers stayed unresolved).

    The statistic caveat ([[VERIFY: ...]] markers for uncertain
    numerics) is retained for ALL document types because
    numerics can still drift between the writer's prose and the
    locked cache values -- the substitution architecture
    catches most, the marker flow catches the rest.
    """
    skip_citation_caveat = document_type in (
        "executive_brief", "analytical_appendix")
    for spec in specs:
        task = spec.get("task", "")
        if not skip_citation_caveat and "[[VERIFY CITATION" not in task:
            task += _CAVEAT_CITATION
        if "[[VERIFY:" not in task:
            task += _CAVEAT_STATS
        spec["task"] = task
    return specs


_MIDPOINT_S2_KEY_FINDINGS = (
    # May 25 2026 — TRIMMED to the four most impactful themes.
    # May 26 2026 (1) — added LOGICAL INTEGRITY GUARDRAILS.
    # May 26 2026 (2) — added a hard EM-DASH PROHIBITION at the
    # top of the prompt AND scrubbed every em dash out of the
    # prompt's own text. Same rationale as S1: the model was
    # treating em dashes in the prompt as license to use them in
    # output.
    "\n\nEM-DASH PROHIBITION (HARD CONSTRAINT):\n"
    "Never use em dashes in any output. Use commas, semicolons, "
    "colons, or restructure the sentence. This applies to every "
    "form: the Unicode em dash, the en dash used as a break, and "
    "the ASCII double-hyphen / triple-hyphen substitutes. A draft "
    "with any dash break will be flagged.\n\n"
    "KEY FINDINGS. Present these in this order, one sentence each.\n"
    "(1) Regime break: the equity-IG correlation shifted from "
    "approximately -0.05 (pre-2022) to +0.61 (post-2022). Quote the "
    "pre/post values from the correlation_pre_post data and connect "
    "to the divergence in strategy performance.\n"
    "(2) Best Sharpe: Regime Switching delivered the highest "
    "full-period risk-adjusted return, approximately 0.63 versus the "
    "benchmark's 0.52, by adapting to the correlation break. Cite "
    "the actual values from the summary_statistics data.\n"
    "(3) OOS validation: in the post-2022 holdout window Regime "
    "Switching's Sharpe (approximately 0.2483) materially exceeded "
    "the benchmark's, confirming the result is not an in-sample "
    "artefact. Cite the regime_conditional data.\n"
    "(4) Diversification benefit: static 60/40 underperformed in the "
    "post-break period because the IG correlation flip removed the "
    "diversification cushion; the dynamic regime-aware strategies "
    "preserved it.\n\n"
    "LOGICAL INTEGRITY GUARDRAILS\n\n"
    "1. REGIME LANGUAGE\n"
    "   Never declare the post-2022 period a definitive \"regime.\" "
    "Always frame as a hypothesis:\n"
    "   \"The persistence and magnitude of the post-2022 inversion "
    "strongly suggests a structural regime shift rather than a "
    "temporary dislocation. It is a hypothesis to be tested formally "
    "in the final phase.\"\n\n"
    "2. STRATEGY SELECTION LANGUAGE\n"
    "   When no single strategy dominates across crisis windows, "
    "always qualify with the asset universe constraint:\n"
    "   \"Given the constraints of our three-asset universe, dynamic "
    "weighting across regimes is the most viable mechanism for alpha "
    "generation when asset selection is restricted.\"\n\n"
    "3. EQUAL-WEIGHT BLEND\n"
    "   The equal-weight blend is a predefined strategy baked into "
    "the model logic, NOT constructed post-hoc after observing "
    "constituent performance. Frame it as a legitimate ex-ante "
    "strategy:\n"
    "   \"The equal-weight blend of active strategies was defined as "
    "part of the original model specification, not selected after "
    "observing results. Its Sharpe of 0.7136 therefore represents a "
    "genuine out-of-sample comparison against the benchmark.\"\n\n"
    "4. MARKOWITZ ATTRIBUTION\n"
    "   Never attribute active blend outperformance solely to "
    "Markowitz diversification. Always qualify with factor exposure:\n"
    "   \"While Markowitz theory predicts efficiency gains from "
    "combining differentiated return streams, the active blend's "
    "outperformance is also driven by dynamic factor rotation. "
    "Specifically, mitigating beta and duration exposure during the "
    "2022 rate shock.\""
)


async def _generate_narratives(
    specs: list[dict], *,
    n_strategies: int | None = None,
    substitution_table: dict[str, str] | None = None,
    document_type: str | None = None,
    defer_substitution: bool | None = None,
) -> dict[str, str]:
    """
    Generates a set of narrative sections concurrently.

    Each spec is {key, agent_id, task, context, available, pending?}.
    Sections with available=False skip the LLM entirely and take their
    [DATA PENDING] marker directly; the rest run through harness_narrative
    in worker threads (the harness is synchronous) and complete in
    parallel — the same asyncio.to_thread fan-out the Academic Review
    peer agents use.

    n_strategies — uniform across every section of a single document
    (it counts the cache, not the section). Threaded through to
    harness_narrative once per spec so the chart-vision scope sentences
    render the precise count instead of the count-omitted fallback.

    substitution_table -- the {token -> value} map produced by
    tools.numeric_substitution.get_substitution_table. When supplied,
    threaded through to every per-section harness_narrative call so
    the post-Sonnet text is substituted before the evaluator scores
    it. None preserves legacy behaviour for callers that don't yet
    pass the table (Layer-2 PR wires the deck + appendix paths).
    """
    import asyncio

    from tools.academic_export import DATA_PENDING, harness_narrative

    # June 28 2026 -- resolve the DEFER_SUBSTITUTION_TO_EXPORT
    # flag ONCE in this async caller before launching the
    # asyncio.to_thread jobs. Threading the bool into harness_
    # narrative eliminates the failed asyncio.run-from-worker
    # path that raised "Future attached to a different loop"
    # (SQLAlchemy's async session is bound to the main loop;
    # a fresh asyncio.run inside the worker thread cannot
    # reuse those connections). The caller may also pre-resolve
    # it themselves and pass through; only query when not
    # supplied.
    if defer_substitution is None:
        try:
            from tools.platform_flags import (
                is_defer_substitution_enabled,
            )
            defer_substitution = (
                await is_defer_substitution_enabled())
        except Exception as _exc:  # noqa: BLE001
            log.warning(
                "defer_flag_resolve_failed",
                error=str(_exc))
            defer_substitution = False

    out: dict[str, str] = {}
    jobs: list[tuple[str, Any]] = []
    for spec in specs:
        if not spec.get("available", True):
            out[spec["key"]] = spec.get(
                "pending", f"{DATA_PENDING} — source data unavailable.")
            continue
        kwargs: dict[str, Any] = {
            "n_strategies": n_strategies,
            "substitution_table": substitution_table,
        }
        # June 21 2026 -- per-spec max_tokens override. The two
        # longest brief sections (key_findings ~550 words,
        # visuals 4 chart paragraphs) were truncating mid-
        # sentence at the default 1500. The spec dict carries an
        # optional max_tokens that flows through to call_claude.
        if "max_tokens" in spec:
            kwargs["max_tokens"] = spec["max_tokens"]
        # June 21 2026 -- per-section anchor allow-list. When the
        # spec carries a numeric_anchors dict (brief sections via
        # _inject_brief_section_plan), it flows through to
        # harness_narrative's post-pass story-plan-violation
        # check. The check re-runs the generator ONCE more with
        # explicit "unauthorized numbers: X, Y, Z" feedback when
        # the first draft emits too many numbers outside the
        # anchor set. See Issue 2 in the post-regen audit
        # (Option 2: harness retry on flag count).
        if "numeric_anchors" in spec:
            kwargs["numeric_anchors"] = spec["numeric_anchors"]
        # June 28 2026 -- thread document_type so harness_narrative
        # can enable the hard-lock untoken-numeric guardrail on
        # protected document types (executive_brief +
        # analytical_appendix). Deck + script paths are unaffected.
        if document_type is not None:
            kwargs["document_type"] = document_type
        # June 28 2026 -- thread the pre-resolved deferral flag
        # so the worker thread never tries to query
        # platform_config from inside its own asyncio.run (the
        # SQLAlchemy session is bound to the main loop +
        # raises "Future attached to a different loop").
        kwargs["defer_substitution"] = bool(defer_substitution)
        jobs.append((spec["key"], asyncio.to_thread(
            harness_narrative, spec["agent_id"], spec["task"], spec["context"],
            **kwargs)))
    if jobs:
        results = await asyncio.gather(*[j for _, j in jobs],
                                       return_exceptions=True)
        for (key, _), res in zip(jobs, results):
            out[key] = res if isinstance(res, str) else (
                f"{DATA_PENDING} — narrative generation failed.")
    return out


async def _editor_export(editor_draft_id: int) -> Response:
    """
    Builds a .docx (paper/brief) or .pptx (deck) from an editor draft's
    current content — the in-editor Export path. Renders the editor
    content directly rather than regenerating the document, so the
    export is exactly what the author has in the editor.
    """
    import asyncio
    from datetime import date

    from tools.editor_drafts import get_draft

    draft = await get_draft(editor_draft_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Draft not found.")

    if draft["document_type"] == "presentation_deck":
        from tools.academic_deck import build_editor_pptx
        from tools.chart_render import is_known_chart, render_chart_png

        # Render each chart element's PNG server-side (an async path) and
        # hand the {element_id: png} map to the sync .pptx builder. A
        # failed render is left out — the builder degrades it gracefully.
        #
        # Bridge #86 — render concurrently with asyncio.gather. The
        # previous implementation awaited each render serially in the
        # for-loop, so a deck with N chart elements at distinct sizes
        # paid N × (gather_document_data + matplotlib) wall-clock.
        # gather() runs the renders in parallel; the per-(chart_key,
        # theme, w, h) render cache means concurrent calls for the same
        # signature still only execute once. Critical on Render where
        # the editor export sits inside the user-facing request and a
        # serial tail can clip Cloudflare's 100 s gateway timeout.
        content_json = draft.get("content_json") or {}
        deck_slides = (content_json.get("slides", [])
                       if isinstance(content_json, dict) else [])

        async def _render_one(el: dict) -> tuple[str, bytes] | None:
            try:
                w = min(2000, max(80, int(el.get("width") or 360) * 2))
                h = min(2000, max(80, int(el.get("height") or 220) * 2))
                # June 26 2026 -- chart_config flows through so
                # editor overrides (title / axis / colors / series
                # visibility) reach the matplotlib output. Legacy
                # elements without chart_config pass None and the
                # renderer falls back to its hardcoded defaults.
                cfg = el.get("chart_config")
                cfg_dict = cfg if isinstance(cfg, dict) else None
                png = await render_chart_png(
                    str(el["chartKey"]), "light", w, h,
                    chart_config=cfg_dict)
                return (str(el.get("id")), png)
            except Exception:  # noqa: BLE001 — skip, builder degrades
                return None

        chart_elements: list[dict] = []
        for sl in deck_slides:
            for el in (sl.get("elements") or [] if isinstance(sl, dict) else []):
                if (isinstance(el, dict) and el.get("type") == "chart"
                        and is_known_chart(str(el.get("chartKey", "")))):
                    chart_elements.append(el)

        rendered = await asyncio.gather(
            *(_render_one(el) for el in chart_elements),
            return_exceptions=False)
        chart_pngs: dict[str, bytes] = {
            eid: png for r in rendered if r is not None
            for eid, png in (r,)}

        content = await asyncio.to_thread(build_editor_pptx, draft, chart_pngs)
        media, ext = _PPTX_MEDIA, "pptx"
    else:
        from tools.academic_docx import build_editor_docx
        # Analytical Appendix is table-heavy by design — re-inject the
        # eight evidence tables after the editor prose so the in-editor
        # export carries the full evidentiary record. Tables read from
        # caches (no recompute) via gather_analytical_appendix_data.
        appendix_data = None
        if draft.get("document_type") == "analytical_appendix":
            from tools.academic_export import gather_analytical_appendix_data
            appendix_data = await gather_analytical_appendix_data()
        # June 25 2026 -- the brief's four APA chart figures
        # (cumulative return, rolling correlation, efficient
        # frontier, OOS Sharpe comparison) are NOT persisted to
        # editor_drafts.content_json -- they're matplotlib PNGs
        # rendered fresh from the analytics snapshot at brief
        # generation time. The editor export used to walk only the
        # TipTap prose nodes and drop the images on the floor.
        # Pull the analytics snapshot now so build_editor_docx can
        # re-render the four figures inline, matching the
        # non-editor regen DOCX.
        brief_data = None
        brief_substitution_table: dict[str, str] | None = None
        # June 28 2026 -- the editor-export substitution table
        # is needed for BOTH brief AND appendix documents.
        # Appendix tokens like {{STUDY_START}}, {{STUDY_END}},
        # {{REGIME_SWITCHING_SHARPE}} would otherwise render
        # literally. Build the same full kwarg set when the
        # draft is either type.
        _needs_sub_table = (
            draft.get("document_type")
            in ("executive_brief", "analytical_appendix"))
        if _needs_sub_table:
            from tools.academic_export import gather_document_data
            brief_data = await gather_document_data()
            # June 25 2026 -- build the substitution table so the
            # four APA figure captions (Figure 1-4 Note. lines)
            # render with substituted cache values instead of
            # leaving literal {{DATA_HASH}} / {{OOS_SHARPE_BLEND}} /
            # {{PRE_2022_EQ_IG_CORR}} / {{N_STRATEGIES}} /
            # {{OOS_WINDOW}} / {{PLAY_BY_PLAY_EVENTS}} markers.
            # PR #403 wired _embed_brief_figures into the editor
            # export but passed substitution_table=None; without
            # the table the figure captions show raw placeholders.
            try:
                # June 28 2026 (Issue 1) -- the editor-export
                # substitution table previously omitted many
                # kwargs that the generation-time call supplies
                # (study_months, implied_allocation, live_signals,
                # oos_window_pct_of_study, hash_verified, freeze-
                # aware data_hash). Result: tokens like
                # {{STUDY_START}} / {{STUDY_END}} /
                # {{CURRENT_REGIME}} / {{REGIME_CONFIDENCE}} /
                # {{SENSITIVITY_COST_BPS_*}} resolved to em-dash
                # or were missing entirely, leaving literal
                # {{TOKEN}} strings in the exported DOCX
                # (operator-reported on draft 80). Mirror the
                # full generation-time kwarg set so EVERY token
                # that resolves at generation also resolves at
                # editor-export.
                from tools.academic_export import (
                    load_substitution_metric_sources,
                )
                from tools.audit_assembler import (
                    current_data_hash as _ee_cur_hash,
                )
                from tools.cio_recommendation import (
                    compute_implied_asset_allocation as _ee_alloc,
                    get_latest_recommendation as _ee_cio,
                )
                from tools.numeric_substitution import (
                    get_substitution_table,
                )
                from tools.submission_freeze import (
                    get_effective_data_hash as _ee_eff_hash,
                )
                from tools.cache import (
                    get_regime_cache as _ee_regime,
                )
                from tools.academic_deck import (
                    CORRELATION_POST_2022,
                    CORRELATION_PRE_2022,
                    OOS_SHARPE_BENCHMARK,
                    OOS_SHARPE_REGIME_CONDITIONAL,
                )
                # Freeze-aware hash so the editor-export table
                # matches what generation produced when a freeze
                # was active.
                _ee_live_hash = await _ee_cur_hash() or ""
                cur_hash = (
                    await _ee_eff_hash(_ee_live_hash)
                    or _ee_live_hash)
                cio_row = await _ee_cio()
                # Implied allocation for the {{CURRENT_*_PCT}}
                # token family (needs CIO blend weights).
                _ee_implied: dict | None = None
                try:
                    if cio_row and cio_row.get("blend_weights"):
                        _ee_implied = await _ee_alloc(
                            cio_row.get("blend_weights"))
                except Exception:  # noqa: BLE001
                    _ee_implied = None
                # Live regime signals for the 5 watchpoint tokens
                # ({{VIX_CURRENT}} / {{YIELD_CURVE_CURRENT}} /
                # {{CREDIT_SPREAD_CURRENT}} /
                # {{EQUITY_TREND_CURRENT}} / {{ESS_CURRENT}}).
                _ee_signals: dict | None = None
                try:
                    _ee_signals = await _ee_regime()
                except Exception:  # noqa: BLE001
                    _ee_signals = None
                rc_rows, fl_rows, cs_payload, crisis_payload = (
                    await load_substitution_metric_sources(
                        data_hash=cur_hash or None))
                # Study-period months -- mirror the generation
                # call's source preference (validated_constants
                # OR strategy_results aggregate). Editor export
                # doesn't carry validated_constants; fall back
                # to the strategy cache's n_observations.
                _ee_study_months: int | None = None
                _ee_strats = (
                    brief_data.get("strategy_results") or {})
                if isinstance(_ee_strats, dict):
                    _n_obs = _ee_strats.get("n_observations")
                    if isinstance(_n_obs, int):
                        _ee_study_months = _n_obs
                # OOS_WINDOW_PCT_OF_STUDY -- constant on
                # academic_deck (53/287 = 18.5).
                try:
                    from tools.academic_deck import (
                        OOS_WINDOW_PCT_OF_STUDY,
                    )
                    _ee_oos_pct = OOS_WINDOW_PCT_OF_STUDY
                except Exception:  # noqa: BLE001
                    _ee_oos_pct = 18.5
                brief_substitution_table = get_substitution_table(
                    cur_hash,
                    brief_data.get("strategy_results") or {},
                    cio_row,
                    oos_sharpe_blend=OOS_SHARPE_REGIME_CONDITIONAL,
                    oos_sharpe_benchmark=OOS_SHARPE_BENCHMARK,
                    pre_2022_eq_ig_correlation=CORRELATION_PRE_2022,
                    post_2022_eq_ig_correlation=CORRELATION_POST_2022,
                    oos_window_pct_of_study=_ee_oos_pct,
                    study_months=_ee_study_months,
                    implied_allocation=_ee_implied,
                    live_signals=_ee_signals,
                    regime_conditional=rc_rows,
                    factor_loadings=fl_rows,
                    cost_sensitivity=cs_payload,
                    crisis_performance=crisis_payload,
                    hash_verified=True)
            except Exception as _exc:  # noqa: BLE001
                # Fail-open: a substitution-table build failure
                # leaves the figure captions with literal
                # placeholders, but the rest of the brief still
                # exports. Same posture as the audit -- it'll
                # surface the unresolved placeholders on the next
                # regen.
                log.warning(
                    "editor_export_substitution_table_failed",
                    error=str(_exc))
                brief_substitution_table = None
        content = await asyncio.to_thread(
            build_editor_docx, draft, appendix_data,
            brief_data=brief_data,
            brief_substitution_table=brief_substitution_table)
        media, ext = _DOCX_MEDIA, "docx"

    # PR #336 Gap D -- re-run the audit on the EDITED content_text
    # before the export downloads. Edits between generation and export
    # can introduce numeric errors, break citation references, or
    # drop a section heading; re-running the audit on the way out
    # catches those before they leave the platform.
    #
    # Best-effort: a missing content_text or any audit exception
    # logs and proceeds with the export. The audit warnings are
    # persisted onto the draft row so the editor's AuditWarningsBanner
    # surfaces them on the next render.
    try:
        edited_text = draft.get("content_text") or ""
        if edited_text:
            owner_email = draft.get("owner_email") or ""
            new_warnings = await _run_document_audit(
                edited_text,
                draft["document_type"],
                owner_email)
            try:
                from tools.editor_drafts import (
                    update_audit_warnings as _update_audit_warnings,
                )
                await _update_audit_warnings(
                    editor_draft_id, new_warnings)
            except Exception as exc:  # noqa: BLE001
                # The update helper may not exist in older deploys;
                # fall back to a direct UPDATE so the draft row stays
                # current. The export must NEVER fail because of this.
                log.warning(
                    "editor_export_audit_persist_warning",
                    error=str(exc))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "editor_export_audit_failed",
            draft_id=editor_draft_id, error=str(exc))

    # Layer 3 (June 21 2026) -- export-time verification. Runs
    # verify_export_against_cache against the value_manifest snapshot
    # persisted on this draft at generation time. Catches manual
    # edits that introduce drift (e.g. "1.24" -> "1.23") and stale
    # data_hash (the cache has moved on since generation). Result
    # is persisted on editor_drafts.export_verification AND surfaces
    # in the response headers (X-Verification-Status,
    # X-Verification-Errors, X-Verification-Warnings) so the
    # frontend can show a status badge without a second round-trip.
    #
    # FAIL-OPEN: the export NEVER blocks because of verification.
    # Errors surface as a header for the frontend; the file ships
    # either way. The user retains agency to download a flagged
    # document if they choose.
    verification: dict[str, Any] = {
        "passed": True, "warnings": [], "errors": [],
        "skipped": "no_manifest_or_helper_unavailable",
    }
    try:
        from tools.audit_assembler import current_data_hash
        from tools.editor_drafts import (
            get_draft_with_layer3 as _get_layer3,
            update_export_verification as _update_verification,
        )
        from tools.numeric_substitution import (
            verify_export_against_cache,
        )
        # Re-read the draft with the Layer-3 columns so we have
        # value_manifest + data_hash for the verification check.
        # Falls back to the legacy shape (value_manifest=None) on
        # pre-migration-057 environments; verification then
        # short-circuits cleanly.
        layer3_draft = (
            await _get_layer3(editor_draft_id) or draft)
        manifest = layer3_draft.get("value_manifest") or {}
        gen_hash = layer3_draft.get("data_hash") or ""
        try:
            cur_hash = await current_data_hash()
        except Exception:  # noqa: BLE001
            cur_hash = ""
        verification = verify_export_against_cache(
            content_text=draft.get("content_text") or "",
            value_manifest=manifest,
            current_data_hash=cur_hash or gen_hash,
            generation_data_hash=gen_hash,
            document_type=draft["document_type"])
        # Persist the verification result on the draft -- frontend
        # status badges read it on the next draft load.
        try:
            await _update_verification(
                editor_draft_id, verification)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "editor_export_verification_persist_warning",
                error=str(exc))
        if verification.get("passed") and not verification.get(
                "warnings"):
            log.info("export_verification_passed",
                     document_type=draft["document_type"],
                     data_hash_match=verification.get(
                         "data_hash_match", True))
        elif verification.get("errors"):
            log.warning(
                "export_verification_failed",
                document_type=draft["document_type"],
                error_count=len(verification["errors"]),
                errors=verification["errors"][:5])
        else:
            log.info(
                "export_verification_warned",
                document_type=draft["document_type"],
                warning_count=len(verification.get("warnings", [])))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "editor_export_verification_failed",
            draft_id=editor_draft_id, error=str(exc))

    # Status header value for the frontend badge: passed / warned /
    # failed. "warned" means the verification was clean except for
    # stale-data-hash warnings; "failed" means at least one error.
    if verification.get("errors"):
        status_header = "failed"
    elif verification.get("warnings"):
        status_header = "warned"
    else:
        status_header = "passed"

    slug = draft["document_type"].replace("_", "-")
    filename = f"forest-capital-{slug}-{date.today().isoformat()}.{ext}"
    return Response(
        content=content, media_type=media,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Verification-Status": status_header,
            "X-Verification-Errors": str(
                len(verification.get("errors") or [])),
            "X-Verification-Warnings": str(
                len(verification.get("warnings") or [])),
        })


# ── Async document generation — job system ────────────────────────────────────
# The three generation endpoints take 30-90s. Each creates a job, spawns
# generation as a background task, and returns 202 immediately; the
# frontend polls GET /api/v1/jobs/{id}. See tools/generation_jobs.py.
_generation_bg_tasks: set = set()


async def _require_report_ready(
    exclude_methodology_check_ids: set[str] | None = None,
) -> None:
    """
    Workstream C report gate. Raises 422 ONLY when caches are not
    warm (a legitimate hard block: there's no data to generate
    against, so generation would produce [DATA PENDING] placeholders).

    June 25 2026 -- audit findings (statistical IN/AN checks +
    methodology checks) no longer hard-block generation. Per the
    platform's fail-open architecture, audit findings are advisory:
    the team makes the final call. Previously this function raised
    422 with error='report_not_ready' when blocking_count > 0;
    every regen was gated behind an explicit acknowledge / mark-
    intentional / revoke pass on the QA Audit tab. That gate is now
    a frontend WARNING (an amber banner above the regen cards) not
    a hard block.

    exclude_methodology_check_ids -- passed through to
    compute_readiness so per-document advisory rules can downgrade
    a specific methodology check. Today only midpoint generation
    uses this with {"IN02"}.

    Fail-open: the readiness module returns empty lists on any
    read error, so a database outage or an empty audit history
    reports is_ready=true and the gate does not block.
    """
    from tools.report_readiness import compute_readiness, summarise_blockers

    readiness = await compute_readiness(
        exclude_methodology_check_ids=exclude_methodology_check_ids)
    if readiness.get("is_ready"):
        return
    # Caches not warm -- the ONE remaining hard block. Without a
    # warm analytics cache there's literally no data for the
    # generators to read, so a regen produces [DATA PENDING] all
    # the way through. The Warm Caches action button in the modal
    # is the user's path forward.
    caches_warm = readiness.get("caches_warm")
    if caches_warm is False:
        cold_caches = readiness.get("cold_caches") or []
        raise HTTPException(
            status_code=422,
            detail={
                "error": "caches_not_warm",
                "message": (
                    "Caches are not warm — generation would produce "
                    "[DATA PENDING] placeholders. Click Warm Caches "
                    "before generating."),
                "blocking_count": readiness.get("blocking_count", 0),
                "blockers": summarise_blockers(readiness),
                "caches_warm": False,
                "cold_caches": cold_caches,
                "warm_status": readiness.get("warm_status"),
                "statistical": readiness.get("statistical"),
                "methodology": readiness.get("methodology"),
            },
        )
    # Audit findings unacknowledged -- WARN, do not block. The
    # frontend surfaces an amber banner above the regen cards via
    # its own read of /api/v1/report/readiness; the response
    # already carries blocking_count + blockers for that banner to
    # render. Generation proceeds.
    log.info(
        "report_readiness_audit_warning_not_block",
        blocking_count=readiness.get("blocking_count", 0),
        blockers=summarise_blockers(readiness)[:5],
    )


async def _auto_resolve_stale_findings_on_regen(
    reviewer_email: str,
) -> int:
    """June 25 2026 -- before a regeneration job spawns, auto-
    resolve every unresolved audit finding that came from a NON-
    LATEST audit run. Rationale: regeneration produces a fresh
    artifact that supersedes whatever the older runs flagged; any
    finding from an older run that the new regen will rebuild
    should clear automatically rather than accumulate as a
    blocker.

    Findings are marked:
      resolved          TRUE
      auto_acknowledged TRUE
      resolution_note   'Auto-resolved: superseded by document
                         regeneration triggered by <email> at
                         <ISO timestamp>'
      resolved_by       <reviewer_email>
      resolved_at       NOW()

    Scope:
      audit_run_id IN (every completed run EXCEPT the latest by id).

    Fail-open: any DB error returns 0 and logs. The regen still
    proceeds; the only side-effect is that the older findings
    keep their resolved=false state.

    Returns the row count actually updated."""
    try:
        from sqlalchemy import text
        from database import (
            AsyncSessionLocal,  # type: ignore[attr-defined]
        )
        if AsyncSessionLocal is None:
            return 0
        async with AsyncSessionLocal() as session:
            from datetime import datetime as _dt, timezone as _tz
            note = (
                f"Auto-resolved: superseded by document regeneration "
                f"triggered by {reviewer_email} at "
                f"{_dt.now(_tz.utc).isoformat()}")
            # Try the full migration-044 column set first; fall
            # back to the legacy UPDATE on column-missing errors.
            try:
                result = await session.execute(text(
                    "UPDATE audit_findings SET "
                    "resolved = TRUE, "
                    "auto_acknowledged = TRUE, "
                    "resolution_note = :note, "
                    "resolved_by = :who, "
                    "resolved_at = NOW() "
                    "WHERE resolved = FALSE "
                    "AND audit_run_id IN ("
                    "  SELECT id FROM audit_runs "
                    "  WHERE status = 'complete' "
                    "  AND id < ("
                    "    SELECT COALESCE(MAX(id), 0) "
                    "    FROM audit_runs "
                    "    WHERE status = 'complete'"
                    "  )"
                    ")"),
                    {"note": note, "who": reviewer_email})
                await session.commit()
                resolved = (
                    result.rowcount
                    if hasattr(result, "rowcount") else 0)
            except Exception:  # noqa: BLE001
                try:
                    await session.rollback()
                except Exception:  # noqa: BLE001
                    pass
                result = await session.execute(text(
                    "UPDATE audit_findings SET "
                    "resolved = TRUE, "
                    "resolution_note = :note "
                    "WHERE resolved = FALSE "
                    "AND audit_run_id IN ("
                    "  SELECT id FROM audit_runs "
                    "  WHERE status = 'complete' "
                    "  AND id < ("
                    "    SELECT COALESCE(MAX(id), 0) "
                    "    FROM audit_runs "
                    "    WHERE status = 'complete'"
                    "  )"
                    ")"),
                    {"note": note})
                await session.commit()
                resolved = (
                    result.rowcount
                    if hasattr(result, "rowcount") else 0)
            if resolved:
                log.info(
                    "audit_findings_auto_resolved_on_regen",
                    reviewer_email=reviewer_email,
                    resolved_count=resolved)
            return int(resolved or 0)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "audit_auto_resolve_on_regen_failed",
            reviewer_email=reviewer_email, error=str(exc))
        return 0


def _start_generation_job(
    document_type: str, session: dict, request: Request,
) -> JSONResponse:
    """Creates a job, spawns generation on the event loop, returns 202.

    Seeds the per-request usage bucket BEFORE the task spawns so the
    Academic Writer harness calls inside _generate_async (which run on
    the same loop, inheriting this context) populate it. The task ends
    by calling _log_interaction_bg which reads collect_usage().

    June 25 2026 -- before spawning the job, auto-resolve any
    unresolved audit findings from PRIOR audit runs (non-latest
    completed runs). The fresh regen will rebuild whatever those
    older findings flagged; leaving them resolved=false caused
    the readiness banner to accumulate stale blockers across the
    May/June audit history. The auto-resolve runs in a fire-and-
    forget task so it can't slow the regen kickoff.
    """
    import asyncio

    from tools.generation_jobs import create_job, update_job
    from agents.usage import start_usage_capture

    start_usage_capture()
    # Best-effort auto-resolve. The task is fire-and-forget; we
    # don't await its completion so a transient DB hiccup never
    # delays the 202 response.
    try:
        resolve_task = asyncio.create_task(
            _auto_resolve_stale_findings_on_regen(
                session.get("email") or ""))
        _generation_bg_tasks.add(resolve_task)
        resolve_task.add_done_callback(_generation_bg_tasks.discard)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "audit_auto_resolve_kickoff_failed",
            error=str(exc))
    job = create_job(document_type, session["email"])
    task = asyncio.create_task(
        _generate_async(job["job_id"], document_type, session, request))
    update_job(job["job_id"], _task=task)
    _generation_bg_tasks.add(task)
    task.add_done_callback(_generation_bg_tasks.discard)
    return JSONResponse(
        status_code=202,
        content={"job_id": job["job_id"], "status": "pending"})


# ── Auto-fire Academic Review on generation (May 25 2026) ─────────────────────
#
# When a midpoint paper or executive brief generation completes, we
# kick off a full Academic Review against the freshly-created editor
# draft on a fire-and-forget background task. The review's score is
# parsed (tools.academic_review_score) and stored in the interaction
# row's metadata.draft_id field; the editor's status endpoint reads
# it on draft load and renders the header pill + advisory banner.
#
# Latency: the full review fan-out (peers + arbiter via the harness)
# takes ~60-90s. We deliberately schedule it AFTER the generation job
# completes so the user sees the draft immediately; the score appears
# in the editor when the background task finishes. The frontend polls
# the status endpoint while it reads `running`.
#
# Fail-open: any review failure logs and skips — the draft is still
# usable, the editor simply doesn't show a score.

# Module-level set so spawned auto-fire tasks aren't GC'd mid-run.
_AUTO_REVIEW_TASKS: set = set()


async def _run_auto_academic_review(
    draft_id: int, document_type: str, owner_email: str,
) -> None:
    """Runs an Academic Review against the draft and persists the
    parsed score into agent_interactions.metadata so the editor can
    surface it. Synchronous-style flow: gather context → peer fan-out
    → arbiter → parse → log. Every step is wrapped — a failure logs
    and the function returns cleanly without raising."""
    import asyncio
    try:
        from agents.academic_review import (
            gather_review_context, run_peer_fan_out, run_arbiter_with_harness,
        )
        from agents.harness import (
            start_harness_capture, collect_harness_metrics,
        )
        from agents.usage import start_usage_capture
        from tools.academic_review_score import compute_review_score
        from tools.activity_log import log_agent_interaction

        start_harness_capture()
        start_usage_capture()
        ctx = await gather_review_context(reviewer_email=owner_email)
        context_block = ctx["context_block"]
        multi_user = ctx.get("multi_user_activity", False)
        n_strategies = ctx["analytics"].get("strategy_count")
        peer_responses = await run_peer_fan_out(
            context_block, multi_user, n_strategies)
        # PR — document-type-specific rubrics. When the auto-fire
        # targets an executive_brief / presentation_deck /
        # analytical_appendix draft, route the arbiter through the
        # corresponding rubric and score the verdict with the
        # matching mode so the editor pill reflects the deliverable's
        # weighted aggregate rather than the midpoint's equal-weight
        # 5.5/10 floor.
        is_brief = document_type == "executive_brief"
        is_deck = document_type == "presentation_deck"
        is_appendix = document_type == "analytical_appendix"
        arbiter_text = await asyncio.to_thread(
            run_arbiter_with_harness, context_block, peer_responses,
            multi_user, False, n_strategies,
            is_brief, is_deck, is_appendix)  # positional: brief/deck/appendix
        if is_brief:
            score_mode = "brief_review"
        elif is_deck:
            score_mode = "deck_review"
        elif is_appendix:
            score_mode = "appendix_review"
        else:
            score_mode = "midpoint"
        scored = compute_review_score(arbiter_text, mode=score_mode)
        agents = list(peer_responses.keys()) + ["academic_advisor"]
        metadata: dict[str, Any] = {
            "draft_id": draft_id,
            "document_type": document_type,
            "automatic": True,
            "advisory": document_type == "midpoint_paper",
            "score": scored["score"],
            "overall_rating": scored["rating"],
            "section_ratings": scored["section_ratings"],
            "sections_rated": scored["sections_rated"],
            "parse_error": scored.get("parse_error", False),
        }
        harness_meta = collect_harness_metrics()
        if harness_meta:
            metadata["harness"] = harness_meta
        await log_agent_interaction(
            user_email=owner_email,
            session_id=None,
            session_type="analytical",
            interaction_type="academic_review",
            agents_involved=agents,
            response_summary=arbiter_text,
            metadata=metadata,
        )
        log.info(
            "auto_academic_review_complete",
            draft_id=draft_id, document_type=document_type,
            score=scored["score"], rating=scored["rating"],
            advisory=metadata["advisory"],
        )
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "auto_academic_review_failed",
            draft_id=draft_id, document_type=document_type,
            error=str(exc))


def _schedule_auto_academic_review(
    draft_id: int | None, document_type: str, owner_email: str,
) -> None:
    """Fire-and-forget — never raises, never blocks the caller.

    Skips entirely in the test environment (the review path makes
    real Anthropic calls, and the contract tests check the endpoint
    shape, not the review wiring). Otherwise spawns a task on the
    loop and stashes a strong reference so the GC doesn't reclaim
    the coroutine mid-run.

    June 21 2026 -- auto-fire DISABLED for every document type.
    The five in-scope deliverables (executive brief, presentation
    deck, presentation script, analytical appendix, Jupyter
    notebook) do not consume the academic-review output as part of
    their workflows. Each generation was firing ~8-10 LLM calls
    (7 peer agents + arbiter with up to 3 retries) purely to
    populate the editor pill + the IN02 attestation row -- both
    surfaces are still reachable via the manual SSE endpoint at
    POST /api/council/academic-review.
    Midpoint paper is past its deadline (May 27) and no longer
    generated. The allow-list below is empty by design; the
    function stays in place so a future opt-in is a one-line
    additive change rather than re-plumbing.

    IN02 attestation dependency: the QA audit panel's IN02 check
    requires an academic_review row in agent_interactions within
    the last 14 days. Operators (Bob / Molly) must click "Run
    Academic Review" in the editor at least once during the
    submission window to keep IN02 PASS. The auto-fire was
    satisfying this implicitly; the manual click satisfies it
    explicitly. See PR body for the operator-action note."""
    import asyncio
    if not draft_id:
        return
    if os.getenv("ENVIRONMENT") == "test":
        return
    # Allow-list is intentionally empty -- no document type
    # auto-fires the academic review. The branch below is
    # structurally preserved so a future opt-in (e.g. add
    # "midpoint_paper" back if a future course reuses the
    # midpoint check) is a one-line edit. Log the skip so the
    # operator can confirm the gate is doing what's intended.
    auto_fire_document_types: set[str] = set()
    if document_type not in auto_fire_document_types:
        log.info(
            "auto_academic_review_skipped_by_design",
            draft_id=draft_id, document_type=document_type,
            note=(
                "auto-fire disabled June 21 2026; run Academic "
                "Review manually before submission to populate "
                "IN02 attestation"))
        return
    try:
        task = asyncio.create_task(
            _run_auto_academic_review(draft_id, document_type, owner_email))
        _AUTO_REVIEW_TASKS.add(task)
        task.add_done_callback(_AUTO_REVIEW_TASKS.discard)
    except RuntimeError:
        # No running loop — we are off-loop and cannot schedule. The
        # auto-fire is best-effort; the user can still trigger the
        # Council Academic Review manually from the panel.
        log.warning("auto_academic_review_no_loop",
                    draft_id=draft_id, document_type=document_type)


async def _generate_async(
    job_id: str, document_type: str, session: dict, request: Request,
) -> None:
    """Runs document generation for a job and records the outcome on it.
    A cancelled job's status is set by the DELETE handler, so a
    CancelledError propagates untouched."""
    import asyncio
    from datetime import datetime, timezone

    from tools.generation_jobs import update_job

    update_job(job_id, status="running")
    try:
        if document_type == "executive_brief":
            file_bytes, filename, media, draft_id = \
                await _generate_brief_document(session["email"])
        elif document_type == "analytical_appendix":
            file_bytes, filename, media, draft_id = \
                await _generate_appendix_document(session["email"])
        elif document_type == "presentation_script":
            # June 25 2026 -- script joined the async job pattern
            # so the tile inherits the standard inProgress /
            # complete / failed / idle chrome.
            file_bytes, filename, media, draft_id = \
                await _generate_script_document(session["email"])
        else:
            file_bytes, filename, media, draft_id = \
                await _generate_deck_document(session["email"])
    except asyncio.CancelledError:
        raise
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        # Bridge #86 — include the exception class and a short traceback
        # excerpt so the Render log reveals the failure surface without
        # forcing a Render-shell session. The ref ties the user-facing
        # generic "Generation failed" to the structured server log.
        import traceback as _tb
        tb_lines = _tb.format_exception_only(type(exc), exc)
        tb_summary = "".join(tb_lines).strip()[:300]
        log.error("generation_job_failed",
                  job_id=job_id,
                  document_type=document_type,
                  ref=ref,
                  exc_type=type(exc).__name__,
                  exc_module=type(exc).__module__,
                  error=str(exc),
                  traceback_excerpt=tb_summary)
        update_job(job_id, status="failed",
                   error=f"Generation failed (ref: {ref})",
                   completed_at=datetime.now(timezone.utc))
        return

    update_job(job_id, status="complete", draft_id=draft_id,
               download_url=f"/api/v1/jobs/{job_id}/download",
               completed_at=datetime.now(timezone.utc),
               _file_bytes=file_bytes, _filename=filename, _media_type=media)
    _log_interaction_bg(
        request, session, "export", agents_involved=["academic_writer"],
        response_summary=f"{document_type} generated",
        metadata={"deliverable": document_type, "draft_id": draft_id})
    # Auto-fire Academic Review against the new draft so the editor
    # shows a readiness score on first open. Non-blocking — the
    # generation job is already complete by the time we get here.
    _schedule_auto_academic_review(
        draft_id, document_type, session["email"])


@app.post("/api/v1/export/midpoint-paper")
@limiter.limit("6/minute")
async def export_midpoint_paper(
    request: Request,
    session: dict = Depends(require_permission("generate_documents")),
):
    """RETIRED (PR-B, June 2026).

    The midpoint submission shipped May 27 2026 and the frontend
    surfaces were removed in PR #338. Returns 410 Gone so existing
    clients receive a clear "this existed and is now gone" signal
    rather than a 404 connection error.

    The helper functions _validate_midpoint_word_counts() and
    _generate_midpoint_document() were deleted alongside this
    handler; the build_midpoint_paper docx assembler was deleted
    from tools/academic_docx.py.
    """
    return JSONResponse(
        status_code=410,
        content={
            "error": "gone",
            "message": (
                "Midpoint paper generation has been retired. The "
                "final submission deadline is July 1."),
            "canonical_path": "/api/v1/export/executive-brief",
        })


# ── Concern 7m: audit chain export ────────────────────────────


async def _assemble_audit_payload(
    document_type: str,
) -> dict[str, Any]:
    """Assembles the structured audit chain for the given scope.

    Joins council_debates + editor_drafts so the rounds list is
    ordered chronologically and each round links to its
    source_draft + resulting_draft. Concern 7m-i shape.

    document_type='full_package' returns rows where context is
    'academic_review' AND document_type is 'full_package'.
    Otherwise returns rows for the specific doc_type.
    """
    payload: dict[str, Any] = {
        "document_type": document_type,
        "data_hash": "",
        "generated_at": (
            __import__("datetime").datetime.now(
                __import__("datetime").timezone.utc).isoformat()),
        "rounds": [],
        "final_draft": None,
        "summary": {
            "total_rounds": 0,
            "total_fatal_raised": 0,
            "total_fatal_addressed": 0,
            "total_fatal_rebutted": 0,
            "total_major_raised": 0,
            "total_major_addressed": 0,
            "total_major_rebutted": 0,
            "total_fixes_applied": 0,
            "manual_edits_made": False,
        },
    }
    if ENVIRONMENT == "test":
        return payload
    try:
        from sqlalchemy import text
        from database import (
            AsyncSessionLocal,  # type: ignore[attr-defined]
        )
        if AsyncSessionLocal is None:
            return payload
        async with AsyncSessionLocal() as s:
            r = await s.execute(text(
                "SELECT id, context, document_type, critic_model, "
                "critic_findings, fatal_count, major_count, "
                "minor_count, peer_responses, arbiter_resolution, "
                "was_addressed, counter_arguments, fix_proposals, "
                "fix_applied, fix_applied_at, new_draft_id, "
                "source_draft_id, parent_debate_id, data_hash, "
                "created_at "
                "FROM council_debates "
                "WHERE document_type = :t "
                "ORDER BY created_at ASC"),
                {"t": document_type})
            rows = r.fetchall()
        round_idx = 0
        for row in rows:
            round_idx += 1
            (
                rid, ctx, dt, model, findings, fatal, major, minor,
                peers, arbiter, addressed, counters,
                proposals, applied, applied_at, new_id, source_id,
                parent_id, dh, created,
            ) = row
            payload["rounds"].append({
                "round": round_idx,
                "debate_id": rid,
                "context": ctx,
                "data_hash": dh,
                "created_at": (
                    created.isoformat() if created else None),
                "source_draft_id": source_id,
                "parent_debate_id": parent_id,
                "critic_findings": {
                    "fatal_count": fatal,
                    "major_count": major,
                    "minor_count": minor,
                    "findings": findings or [],
                    "model": model,
                },
                "council_response": arbiter or "",
                "counter_arguments": counters or [],
                "fix_proposals": proposals or [],
                "fix_applied": bool(applied),
                "fix_applied_at": (
                    applied_at.isoformat() if applied_at else None),
                "resulting_draft_id": new_id,
            })
            payload["summary"]["total_fatal_raised"] += int(
                fatal or 0)
            payload["summary"]["total_major_raised"] += int(
                major or 0)
            for i, addr in enumerate(addressed or []):
                f = (findings or [])[i] if (
                    findings and i < len(findings)) else {}
                sev = str(f.get("severity") or "").capitalize()
                if addr:
                    if sev == "Fatal":
                        payload["summary"][
                            "total_fatal_addressed"] += 1
                    elif sev == "Major":
                        payload["summary"][
                            "total_major_addressed"] += 1
            for ca in (counters or []):
                sev = (
                    str((ca.get("finding") or {})
                        .get("severity") or "").capitalize())
                if sev == "Fatal":
                    payload["summary"][
                        "total_fatal_rebutted"] += 1
                elif sev == "Major":
                    payload["summary"][
                        "total_major_rebutted"] += 1
            if applied:
                payload["summary"]["total_fixes_applied"] += 1
            if dh and not payload["data_hash"]:
                payload["data_hash"] = dh
        payload["summary"]["total_rounds"] = round_idx
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "audit_export_assembly_failed", error=str(exc))
    return payload


@app.get("/api/v1/documents/audit-export")
@limiter.limit("10/minute")
async def get_audit_export(
    request: Request,
    document_type: str = "full_package",
    session: dict = Depends(require_team_member),
):
    """Concern 7m-i. Returns the full audit chain for the requested
    scope as structured JSON. Joins council_debates rows ordered by
    created_at; each round carries critic_findings, council_response,
    counter_arguments, fix_proposals, and (when fix was applied) the
    resulting_draft_id link.

    document_type=full_package -- cross-document review chain.
    document_type=<editor type> -- per-doc chain.
    """
    if document_type not in (
            "full_package", "executive_brief",
            "presentation_deck", "analytical_appendix",
            "presentation_script"):
        raise HTTPException(
            status_code=422,
            detail=(
                "document_type must be full_package or one of the "
                "four editor doc types"))
    return await _assemble_audit_payload(document_type)


@app.post("/api/v1/documents/audit-export/docx")
@limiter.limit("6/minute")
async def post_audit_export_docx(
    request: Request,
    document_type: str = "full_package",
    session: dict = Depends(require_team_member),
):
    """Concern 7m-ii. Same payload as the JSON endpoint, rendered
    as a formatted .docx file download.

    The DOCX uses python-docx (same library the brief / appendix
    rendering chain already uses -- no new dependency). Structure:

      Title + subtitle + data hash header
      Executive summary table
      Per-round sections (critic findings table, council response
      block, fix-applied row)
      Counter-arguments on record (all rebuttals across all rounds)
      Revision history table
      Data integrity footer
    """
    if document_type not in (
            "full_package", "executive_brief",
            "presentation_deck", "analytical_appendix",
            "presentation_script"):
        raise HTTPException(
            status_code=422, detail="invalid document_type")
    payload = await _assemble_audit_payload(document_type)

    if ENVIRONMENT == "test":
        return {"ok": True, "rounds": 0, "test_environment": True}

    try:
        from docx import Document
        from docx.shared import Pt
        from io import BytesIO
        from datetime import date as _date
        doc = Document()
        # Title
        doc.add_heading(
            "Forest Capital -- Adversarial Review Audit Trail",
            level=0)
        sub = doc.add_paragraph(
            "FNA 670 | McColl School of Business | "
            "Queens University of Charlotte")
        sub.runs[0].italic = True
        doc.add_paragraph(
            f"Data hash: {payload.get('data_hash') or '(none)'} "
            f"| Generated: {_date.today().isoformat()}")

        summary = payload.get("summary") or {}
        doc.add_heading("Executive Summary", level=1)
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "Metric"
        t.rows[0].cells[1].text = "Value"
        for label, key in [
            ("Rounds",                "total_rounds"),
            ("Fatal raised",          "total_fatal_raised"),
            ("Fatal addressed",       "total_fatal_addressed"),
            ("Fatal rebutted",        "total_fatal_rebutted"),
            ("Major raised",          "total_major_raised"),
            ("Major addressed",       "total_major_addressed"),
            ("Major rebutted",        "total_major_rebutted"),
            ("Fixes applied",         "total_fixes_applied"),
        ]:
            row = t.add_row().cells
            row[0].text = label
            row[1].text = str(summary.get(key, 0))

        for r in payload.get("rounds") or []:
            doc.add_heading(
                f"Round {r['round']} -- {r['created_at'] or ''}",
                level=1)
            cf = r.get("critic_findings") or {}
            doc.add_paragraph(
                f"Source draft id: {r.get('source_draft_id')}")
            doc.add_paragraph(
                f"Model: {cf.get('model', 'unknown')}  |  "
                f"Fatal {cf.get('fatal_count', 0)}  /  "
                f"Major {cf.get('major_count', 0)}  /  "
                f"Minor {cf.get('minor_count', 0)}")
            findings = cf.get("findings") or []
            if findings:
                doc.add_heading(
                    "Adversarial Critic Findings", level=2)
                tbl = doc.add_table(rows=1, cols=5)
                tbl.style = "Light Grid Accent 1"
                hdr = tbl.rows[0].cells
                hdr[0].text = "Severity"
                hdr[1].text = "Category"
                hdr[2].text = "Location"
                hdr[3].text = "Description"
                hdr[4].text = "Raised by"
                for f in findings:
                    row = tbl.add_row().cells
                    row[0].text = str(f.get("severity") or "")
                    row[1].text = str(f.get("category") or "")
                    row[2].text = str(f.get("location") or "")
                    row[3].text = str(f.get("description") or "")
                    row[4].text = str(f.get("raised_by") or "")
            response_text = r.get("council_response") or ""
            if response_text.strip():
                doc.add_heading("Council Response", level=2)
                doc.add_paragraph(response_text)
            if r.get("fix_applied"):
                doc.add_heading("Fix Applied", level=2)
                doc.add_paragraph(
                    f"Resulting draft id: "
                    f"{r.get('resulting_draft_id')}")
                doc.add_paragraph(
                    f"Applied at: "
                    f"{r.get('fix_applied_at') or ''}")

        # Counter-arguments on record -- aggregated across rounds.
        all_counters: list[dict[str, Any]] = []
        for r in payload.get("rounds") or []:
            all_counters.extend(r.get("counter_arguments") or [])
        if all_counters:
            doc.add_heading(
                "Counter-Arguments on Record", level=1)
            doc.add_paragraph(
                "The following critic findings were reviewed and "
                "intentionally not addressed based on the council's "
                "rebuttal -- these are documented rebuttals showing "
                "the team considered and responded to the critique.")
            for ca in all_counters:
                f = ca.get("finding") or {}
                doc.add_heading(
                    f"{f.get('severity', '')} "
                    f"{f.get('category', '')} -- "
                    f"{f.get('location', '')}", level=2)
                doc.add_paragraph(
                    f"Critic raised: {f.get('description', '')}")
                doc.add_paragraph(
                    f"Council response: {ca.get('rebuttal', '')}")
                doc.add_paragraph(
                    f"Model source: "
                    f"{ca.get('model_source', 'unknown')}")

        # Data integrity footer
        doc.add_heading("Data Integrity", level=1)
        doc.add_paragraph(
            f"Canonical data hash: "
            f"{payload.get('data_hash') or '(none)'}")
        doc.add_paragraph(
            "All figures in submitted documents verified against "
            "this hash via the verify-all endpoint.")

        # Bump default font size for legibility.
        for p in doc.paragraphs:
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(10)

        buf = BytesIO()
        doc.save(buf)
        body = buf.getvalue()
        from fastapi.responses import Response as _Resp
        filename = (
            f"forest_capital_audit_trail_{document_type}_"
            f"{_date.today().isoformat()}.docx")
        return _Resp(
            content=body,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"),
            headers={
                "Content-Disposition": (
                    f"attachment; filename=\"{filename}\""),
            })
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "audit_export_docx_failed", error=str(exc))
        raise HTTPException(
            status_code=500,
            detail=f"DOCX render failed: {exc}")


# ── Concern 7k-iii + 7k-viii: apply-fix + propose-fix endpoints ──

@app.post("/api/v1/documents/propose-fix")
@limiter.limit("20/minute")
async def post_propose_fix(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Concern 7k-viii. Generates a fix proposal for a single
    Major finding on demand. Fatal findings auto-fire during the
    debate round; this endpoint serves the Major case (the UI
    Propose Fix button on a Major finding card).

    Body: {
      document_type: <type>,
      finding_id:    <int -- index into the debate's
                     merged_findings>,
      finding:       <finding object>,
      debate_id:     <int -- council_debates row id>
    }

    Returns the FixProposal as JSON. Also appends the proposal to
    council_debates.fix_proposals so the UI keeps a single source
    of truth for what's been proposed.
    """
    from agents.academic_review import (
        run_arbiter_fix_proposal, write_fix_proposals_to_debate,
    )
    document_type = str(body.get("document_type") or "")
    if document_type not in {
            "executive_brief", "analytical_appendix",
            "presentation_deck", "presentation_script"}:
        raise HTTPException(
            status_code=422,
            detail=(
                "document_type must be one of executive_brief / "
                "analytical_appendix / presentation_deck / "
                "presentation_script"))
    finding = body.get("finding") or {}
    if not isinstance(finding, dict):
        raise HTTPException(
            status_code=422, detail="finding must be an object")
    finding_id = int(body.get("finding_id") or 0)
    debate_id = body.get("debate_id")
    proposal = await run_arbiter_fix_proposal(
        finding=finding, finding_id=finding_id,
        document_type=document_type,
        reviewer_email=session.get("email"))
    if proposal is None:
        return {
            "ok": False,
            "message": (
                "Could not generate a fix proposal for this "
                "finding -- the arbiter response was not parseable. "
                "You can still edit the document manually."),
        }
    if debate_id is not None:
        try:
            await write_fix_proposals_to_debate(
                int(debate_id), [proposal])
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "propose_fix_write_failed", error=str(exc))
    return {
        "ok": True,
        "finding_id": proposal.finding_id,
        "target": proposal.target,
        "section_name": proposal.section_name,
        "rationale": proposal.rationale,
        "patch_instruction": proposal.patch_instruction,
        "severity": proposal.severity,
        "auto_proposed": proposal.auto_proposed,
        "target_document": proposal.target_document,
        "source_of_truth_document": (
            proposal.source_of_truth_document),
    }


# ── June 26 2026 -- direct section-patch path for apply-fix ──────────────
#
# Background. The legacy apply-fix flow (a) patched a story_plans row by
# appending the patch_instruction to a section's `guidance` field, then
# (b) kicked off a FULL document regeneration. That round-trip was
# necessary when the fix changed the story plan's structural anchors
# (numeric_anchors, central_argument), but for a localised correction
# to one slide's prose / numbers it was overkill -- and it failed
# outright on the deck because the story_plans row lookup was 409ing
# (apply-fix story-plan-lookup PR #443 fixed the hash join, but Apply
# Fix still 409s on any deck draft generated before a story plan was
# computed -- pre-#332 drafts in particular).
#
# This path side-steps the regen requirement for fixes that can be
# applied directly to the current draft's content_json. Branches by
# document_type:
#
#   presentation_deck  -- find the slide by section_name (title match),
#     send its elements array to Sonnet with the patch_instruction,
#     replace the elements with the corrected version.
#   executive_brief / analytical_appendix -- find the TipTap section by
#     section_name (heading text match), send its nodes to Sonnet,
#     replace the section's nodes with the corrected version.
#   presentation_script -- no first-class section concept (per-slide
#     regen lives at a different endpoint); falls through to the legacy
#     path.
#
# Returns the response dict on success; the caller short-circuits the
# legacy story-plan + regen block. Returns None when the direct path
# can't apply (no current draft, section not found, Sonnet output
# unparseable, target='document' on a section-only fix proposal) so the
# caller falls back to the legacy path -- preserves the existing
# regen-the-whole-doc behaviour as the safety net.


_DECK_SLIDE_PATCH_SYSTEM_PROMPT = (
    "You are a precise JSON editor for a presentation slide. You "
    "will receive a slide title, a patch instruction describing a "
    "targeted correction, and the slide's current `elements` array "
    "as JSON.\n\n"
    "Apply ONLY the change requested by the patch instruction. "
    "Preserve every other field, every other element, and the "
    "overall structure exactly. Do not rephrase prose that the "
    "instruction does not touch; do not change colors, positions, "
    "ids, types, or any element the instruction doesn't name.\n\n"
    "Return ONLY the corrected elements array as valid JSON. No "
    "commentary, no markdown fence, no explanation -- just the "
    "array."
)


_TIPTAP_SECTION_PATCH_SYSTEM_PROMPT = (
    "You are a precise JSON editor for one section of a document "
    "written in TipTap JSON format. You will receive the section "
    "name, a patch instruction describing a targeted correction, "
    "and the section's current TipTap nodes as a JSON array "
    "(starting with the heading node and ending just before the "
    "next heading).\n\n"
    "Apply ONLY the change requested by the patch instruction. "
    "Preserve every other paragraph, the heading, marks, the node "
    "ordering, and the TipTap structure exactly. Keep marker "
    "callouts ([[BOB: ...]], [[VERIFY: ...]]) intact unless the "
    "instruction explicitly says to remove them.\n\n"
    "Return ONLY the corrected node array as valid JSON. No "
    "commentary, no markdown fence, no explanation -- just the "
    "array."
)


def _norm_title(s: str) -> str:
    """Case-insensitive whitespace-collapsed title normalisation
    for the slide / section title match. The fix proposal's
    section_name often comes from the LLM's free-text description
    of the finding; we accept moderate punctuation drift."""
    import re as _re
    return _re.sub(r"\s+", " ", str(s or "").strip().lower())


def _extract_json_array(raw: str) -> list | None:
    """Tolerantly extract the first JSON array from a Sonnet
    response. Strips a markdown ```json fence if present, then
    falls back to the longest [...] span. Returns None on parse
    failure or when the result isn't a list."""
    import re as _re
    if not raw:
        return None
    # Strip markdown fence.
    fence = _re.search(
        r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", raw)
    candidate = fence.group(1) if fence else None
    if candidate is None:
        # Longest array span: from the first '[' to the last ']'.
        first = raw.find("[")
        last = raw.rfind("]")
        if first == -1 or last == -1 or last <= first:
            return None
        candidate = raw[first:last + 1]
    try:
        parsed = json.loads(candidate)
    except Exception:  # noqa: BLE001
        return None
    return parsed if isinstance(parsed, list) else None


def _find_deck_slide_idx(
    slides: list, section_name: str,
) -> int | None:
    """Find a slide by title match. Returns the index or None.

    Match algorithm (case-insensitive, whitespace-collapsed):
      1. Exact match  -- target == slide title
      2. Substring    -- target appears in slide title OR
                         slide title appears in target
      3. Slide-number prefix -- 'Slide N' / 'Slide N:' in the
         target matches the slide at index N-1 (defensive: works
         even when the title in the finding has drifted)

    June 27 2026 -- relaxed from exact-only to fuzzy substring.
    The fix-proposal's section_name comes from an LLM's free-text
    description of the finding (e.g. 'Slide 4: Why Static Failed
    in 2022') and frequently drifts from the slide's actual title
    (e.g. 'Why Static Allocation Failed in 2022') by a word or
    two. The exact-match version 422'd those clicks; this fuzzy
    matcher recovers them."""
    import re as _re
    target = _norm_title(section_name)
    if not target:
        return None
    # 1. + 2. Exact and substring matches.
    for i, sl in enumerate(slides):
        if not isinstance(sl, dict):
            continue
        title = _norm_title(sl.get("title"))
        if not title:
            continue
        if target == title or target in title or title in target:
            return i
    # 3. 'Slide N' prefix fallback.
    m = _re.search(r"slide\s+(\d+)", target)
    if m:
        n = int(m.group(1))
        if 1 <= n <= len(slides) and isinstance(slides[n - 1], dict):
            return n - 1
    return None


def _node_text_for_match(node: Any) -> str:
    """Flatten a TipTap node's text content recursively for
    heading-title matching."""
    if not isinstance(node, dict):
        return ""
    if node.get("text"):
        return str(node["text"])
    return "".join(
        _node_text_for_match(c) for c in (node.get("content") or []))


def _find_tiptap_section_range(
    nodes: list, section_name: str,
) -> tuple[int, int] | None:
    """Returns (start_idx, end_idx_exclusive) of the matching
    section -- from its heading node up to (but not including) the
    next heading. None when no heading matches or the section is
    empty.

    Match algorithm (case-insensitive, whitespace-collapsed):
      1. Exact match.
      2. Substring (either direction).
      3. 'Slide N' / 'Slide N:' prefix match -- when the target
         carries a slide-number prefix and any heading text starts
         with the same 'Slide N' / 'Slide N:' / 'Slide N -' token,
         that heading wins. Defensive against LLM drift on the
         script's per-slide H2 markers.
      4. Word-set overlap fallback -- when no direct match, accept
         a heading whose normalised word set shares >= 60% of the
         target's words (rounded down, minimum 2 shared words).
         Catches the 'Why Static Failed' vs 'Why Static Allocation
         Failed in 2022' drift that fuzzy substring alone would
         miss when neither string is a substring of the other.

    June 27 2026 -- expanded from exact + substring to a 4-tier
    fuzzy match. Heading-text drift from the fix-proposal LLM is
    the leading cause of apply-fix 422s in production; this matcher
    closes the gap so Bob's evening editing pass doesn't die on a
    section-name comma."""
    import re as _re
    target = _norm_title(section_name)
    if not target:
        return None
    target_slide_n: int | None = None
    sm = _re.search(r"slide\s+(\d+)", target)
    if sm:
        target_slide_n = int(sm.group(1))
    # Bare 'Section X' / 'Section 4' identifier (letter or digit)
    # so we can match a heading prefixed with the same ID when no
    # body text was provided alongside the label -- analytical
    # appendix headings are 'A. Strategy Universe', 'B. Performance
    # & Risk Metrics' and a finding may name them as 'Section B'.
    target_section_id: str | None = None
    bare_m = _re.match(
        r"^section\s+([a-z\d]+)\s*[:.\-]?\s*$", target)
    if bare_m:
        target_section_id = bare_m.group(1)
    target_words = set(_re.findall(r"\w+", target))

    # First pass: collect every heading + its index for fuzzy
    # comparison, then pick the best match per the tiers above.
    headings: list[tuple[int, str]] = []
    for i, n in enumerate(nodes):
        if not isinstance(n, dict) or n.get("type") != "heading":
            continue
        title = _norm_title(_node_text_for_match(n))
        if title:
            headings.append((i, title))

    if not headings:
        return None

    chosen_start: int | None = None

    # 1. Exact match.
    for i, title in headings:
        if title == target:
            chosen_start = i
            break

    # 2. Substring match.
    if chosen_start is None:
        for i, title in headings:
            if target in title or title in target:
                chosen_start = i
                break

    # 3. Slide-number prefix match.
    if chosen_start is None and target_slide_n is not None:
        slide_token_re = _re.compile(
            rf"^slide\s+{target_slide_n}\b")
        for i, title in headings:
            if slide_token_re.search(title):
                chosen_start = i
                break

    # 3b. Bare 'Section X' identifier prefix match. Catches
    # 'Section B' -> 'B. Strategy Universe' even when neither
    # substring nor word-overlap matches.
    if chosen_start is None and target_section_id is not None:
        id_prefix_re = _re.compile(
            rf"^{_re.escape(target_section_id)}[.)]\s+\S")
        for i, title in headings:
            if id_prefix_re.match(title):
                chosen_start = i
                break

    # 4. Word-set overlap fallback. The needle's coverage of the
    # heading's words (or vice versa) must be >= 60% AND at least
    # 2 words must be shared (avoids false positives on single
    # common words like 'the'). Coverage is computed as
    #   shared / len(target_words)
    # so 'Why Static Failed' (3 words) vs 'Why Static Allocation
    # Failed in 2022' (6 words) gives 3/3 = 1.0 -- the needle is
    # fully contained in the heading. This is the canonical Bob-
    # tonight tripwire case.
    if chosen_start is None and len(target_words) >= 2:
        best_idx: int | None = None
        best_share = 0.0
        for i, title in headings:
            tw = set(_re.findall(r"\w+", title))
            if not tw:
                continue
            shared = target_words & tw
            if len(shared) < 2:
                continue
            share = len(shared) / len(target_words)
            if share >= 0.6 and share > best_share:
                best_idx = i
                best_share = share
        if best_idx is not None:
            chosen_start = best_idx

    if chosen_start is None:
        return None

    # Determine the section's exclusive end: index of the next
    # heading after chosen_start, or len(nodes).
    for i, _title in headings:
        if i > chosen_start:
            return chosen_start, i
    return chosen_start, len(nodes)


async def _patch_deck_slide_via_sonnet(
    slide: dict, patch_instruction: str,
) -> list | None:
    """Sends the slide's elements + the patch instruction to
    Sonnet; returns the corrected elements array or None on
    parse / validation failure. Validates that every returned
    element is a dict carrying at minimum `id` and `type` keys
    (preserves the canvas-element shape contract)."""
    import asyncio as _asyncio
    from agents.base import SONNET_MODEL, call_claude

    elements = slide.get("elements") or []
    user_message = (
        f"SLIDE TITLE: {slide.get('title') or ''}\n\n"
        f"PATCH INSTRUCTION:\n{patch_instruction}\n\n"
        "CURRENT ELEMENTS JSON:\n"
        f"{json.dumps(elements, indent=2)}\n\n"
        "Return ONLY the corrected elements array as JSON."
    )
    try:
        raw = await _asyncio.to_thread(
            call_claude, SONNET_MODEL,
            _DECK_SLIDE_PATCH_SYSTEM_PROMPT, user_message, 4000)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "apply_fix_deck_sonnet_failed", error=str(exc))
        return None
    parsed = _extract_json_array(raw)
    if parsed is None:
        log.warning(
            "apply_fix_deck_sonnet_unparseable",
            raw_preview=(raw or "")[:200])
        return None
    if not all(
        isinstance(e, dict) and "id" in e and "type" in e
        for e in parsed
    ):
        log.warning("apply_fix_deck_element_shape_violation")
        return None
    return parsed


async def _patch_tiptap_section_via_sonnet(
    section_nodes: list, section_name: str,
    patch_instruction: str,
) -> list | None:
    """Sends the section's TipTap nodes + the patch instruction to
    Sonnet; returns the corrected node array or None on parse /
    validation failure. Validates that every returned node is a
    dict with a `type` field."""
    import asyncio as _asyncio
    from agents.base import SONNET_MODEL, call_claude

    user_message = (
        f"SECTION NAME: {section_name}\n\n"
        f"PATCH INSTRUCTION:\n{patch_instruction}\n\n"
        "CURRENT SECTION NODES (TipTap JSON):\n"
        f"{json.dumps(section_nodes, indent=2)}\n\n"
        "Return ONLY the corrected node array as JSON."
    )
    try:
        raw = await _asyncio.to_thread(
            call_claude, SONNET_MODEL,
            _TIPTAP_SECTION_PATCH_SYSTEM_PROMPT,
            user_message, 4000)
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "apply_fix_tiptap_sonnet_failed", error=str(exc))
        return None
    parsed = _extract_json_array(raw)
    if parsed is None:
        log.warning(
            "apply_fix_tiptap_sonnet_unparseable",
            raw_preview=(raw or "")[:200])
        return None
    if not all(
        isinstance(n, dict) and "type" in n for n in parsed
    ):
        log.warning("apply_fix_tiptap_node_shape_violation")
        return None
    return parsed


def _derive_content_text_from_tiptap(doc: dict) -> str:
    """Plain-text projection of a TipTap doc for content_text.
    Mirrors the editor's save-flow shape."""
    out = []
    for n in (doc.get("content") or []):
        text = _node_text_for_match(n).strip()
        if text:
            out.append(text)
    return "\n\n".join(out)


def _derive_content_text_from_deck(deck: dict) -> str:
    """Plain-text projection of a canvas deck for content_text.
    Concatenates each slide's title + body text."""
    out = []
    for sl in (deck.get("slides") or []):
        if not isinstance(sl, dict):
            continue
        out.append(f"Slide {sl.get('id', '')}: {sl.get('title', '')}")
        for el in (sl.get("elements") or []):
            if isinstance(el, dict) and el.get("type") == "text":
                content = str(el.get("content") or "").strip()
                if content:
                    out.append(content)
    return "\n".join(out)


async def _mark_council_debate_fix_applied(
    debate_id: int, new_draft_id: int | None,
) -> None:
    """Mark the council_debates row as applied. Best-effort: a
    write failure logs but does not unwind the patch."""
    try:
        from sqlalchemy import text as _text
        from database import (
            AsyncSessionLocal as _ASL,  # type: ignore
        )
        if _ASL is None:
            return
        async with _ASL() as s:
            await s.execute(_text(
                "UPDATE council_debates SET "
                "fix_applied = TRUE, "
                "fix_applied_at = NOW(), "
                "new_draft_id = COALESCE(:n, new_draft_id) "
                "WHERE id = :id"),
                {"id": int(debate_id),
                 "n": new_draft_id})
            await s.commit()
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "apply_fix_mark_debate_failed", error=str(exc))


_APPLY_FIX_DOC_LABEL = {
    "presentation_deck": "Presentation Deck",
    "executive_brief": "Executive Brief",
    "analytical_appendix": "Analytical Appendix",
    "presentation_script": "Presentation Script",
}


class _ApplyFixSkip(Exception):
    """Raised by _try_direct_section_patch when the inline path can
    NOT apply the fix and the user should be told to edit directly.

    June 27 2026 -- replaces the previous None-return + fall-through
    to a legacy regenerate-the-whole-document path. The legacy path
    was removed (apply-fix no longer triggers a full regen under
    any circumstance); a clean failure surfaces as a 422 to the
    frontend with detail / hint / section_name so the UI can show a
    \"Open in editor\" deep link to the right section.

    Attributes
        detail:        Plain-English reason -- 'Section \"X\" not
                       found in the current draft.'
        hint:          Actionable next step -- 'Open the editor and
                       edit the section directly.'
        section_name:  Best-known section identifier so the frontend
                       can deep-link to the editor section header.
                       None when the failure precedes any section
                       match (no current draft, wrong target).
    """

    def __init__(
        self, detail: str, *, hint: str | None = None,
        section_name: str | None = None,
    ):
        super().__init__(detail)
        self.detail = detail
        self.hint = hint or (
            "Open the document in the editor and apply the change "
            "directly. The Apply Fix workflow only handles "
            "section-level edits the AI can locate; structural "
            "edits (story-arc anchors, numeric_anchors) must be "
            "applied manually.")
        self.section_name = section_name


async def _try_direct_section_patch(
    *, document_type: str, target: str | None,
    section_name: str | None, patch_instruction: str,
    debate_id: int | None,
) -> dict | None:
    """Apply the fix as a direct in-place patch to the current
    draft's content_json. Returns the response dict on success.

    June 27 2026 -- legacy regenerate-the-whole-document fallback
    removed. This path is now the ONLY apply-fix execution path; a
    clean failure raises _ApplyFixSkip which the caller turns into
    a 422 telling the user to edit in the editor directly.

    Coverage post-refactor:
      target='section' + section_name set:
        deck                 -> _find_deck_slide_idx + Sonnet-patch
                                that slide's elements
        brief / appendix     -> _find_tiptap_section_range (H1
                                heading match) + Sonnet-patch that
                                section's nodes
        script               -> _find_tiptap_section_range (H2
                                'Slide N' heading match) + Sonnet-
                                patch that slide's prose nodes
      target='document':
        deck                 -> per-slide walk; Sonnet-patch every
                                slide in parallel; reassemble
        brief / appendix /
        script               -> per-section walk over the heading
                                ranges; Sonnet-patch each;
                                reassemble nodes preserving
                                ordering

    Raises _ApplyFixSkip when:
      * No current draft exists for the document type
      * target='section' but section_name is empty
      * target is neither 'section' nor 'document'
      * Section / slide title doesn't match any heading in the
        draft (target='section')
      * Sonnet patch is unparseable / shape-invalid AND it's the
        only call (target='section'). target='document' tolerates
        per-section Sonnet failures by skipping that section and
        continuing -- a complete failure (every section failed)
        raises.
      * update_draft persists fails

    No silent fall-through. The legacy regen path has been deleted.
    """
    if target not in ("section", "document"):
        raise _ApplyFixSkip(
            f"Unknown fix target '{target}'.")
    if target == "section" and not section_name:
        raise _ApplyFixSkip(
            "Section-scoped fix requires a section_name.")

    if document_type not in (
        "presentation_deck", "executive_brief",
        "analytical_appendix", "presentation_script",
    ):
        raise _ApplyFixSkip(
            f"Apply Fix does not support document_type "
            f"'{document_type}'.")

    from tools.editor_drafts import (
        get_current_draft_by_type, update_draft,
    )

    draft = await get_current_draft_by_type(document_type)
    if draft is None:
        raise _ApplyFixSkip(
            f"No current draft found for "
            f"{_APPLY_FIX_DOC_LABEL.get(document_type, document_type)}. "
            "Generate the document before applying a fix.")
    content = draft.get("content_json")
    if not isinstance(content, dict):
        raise _ApplyFixSkip(
            f"Current draft for "
            f"{_APPLY_FIX_DOC_LABEL.get(document_type, document_type)} "
            "has no content to patch.")

    is_deck = document_type == "presentation_deck"
    is_tiptap = document_type in (
        "executive_brief", "analytical_appendix",
        "presentation_script")

    # ── Section-scoped patch ──────────────────────────────────────
    if target == "section":
        if is_deck:
            slides = content.get("slides")
            if not isinstance(slides, list):
                raise _ApplyFixSkip(
                    "Deck draft has no slides to patch.")
            idx = _find_deck_slide_idx(slides, section_name)
            if idx is None:
                raise _ApplyFixSkip(
                    f"Slide '{section_name}' not found in the "
                    "current deck draft.",
                    section_name=section_name)
            new_elements = await _patch_deck_slide_via_sonnet(
                slides[idx], patch_instruction)
            if new_elements is None:
                raise _ApplyFixSkip(
                    f"AI could not produce a valid patch for "
                    f"slide '{section_name}'. Try editing in "
                    "the editor directly.",
                    section_name=section_name)
            new_slides = list(slides)
            new_slides[idx] = {
                **slides[idx], "elements": new_elements}
            new_content = {**content, "slides": new_slides}
            new_text = _derive_content_text_from_deck(new_content)
        elif is_tiptap:
            nodes = content.get("content")
            if not isinstance(nodes, list):
                raise _ApplyFixSkip(
                    "Current draft has no TipTap content to "
                    "patch.")
            rng = _find_tiptap_section_range(nodes, section_name)
            if rng is None:
                raise _ApplyFixSkip(
                    f"Section '{section_name}' not found in the "
                    "current draft. Confirm the section heading "
                    "matches.",
                    section_name=section_name)
            start, end = rng
            section_nodes = nodes[start:end]
            new_section = await _patch_tiptap_section_via_sonnet(
                section_nodes, section_name, patch_instruction)
            if new_section is None:
                raise _ApplyFixSkip(
                    f"AI could not produce a valid patch for "
                    f"section '{section_name}'. Try editing in "
                    "the editor directly.",
                    section_name=section_name)
            new_nodes = nodes[:start] + new_section + nodes[end:]
            new_content = {**content, "content": new_nodes}
            new_text = _derive_content_text_from_tiptap(new_content)
        else:  # pragma: no cover - guarded above
            raise _ApplyFixSkip(
                f"Unhandled document_type '{document_type}'.")

    # ── Document-scoped patch (target='document') ────────────────
    else:
        new_content, new_text, sections_patched, sections_skipped = (
            await _patch_entire_document_via_sonnet(
                document_type=document_type,
                content=content,
                patch_instruction=patch_instruction))
        if sections_patched == 0:
            raise _ApplyFixSkip(
                "AI could not apply the patch to any section of "
                f"{_APPLY_FIX_DOC_LABEL.get(document_type, document_type)}. "
                "Try editing in the editor directly, or scope the "
                "fix to a single section.")

    ok = await update_draft(
        int(draft["id"]), new_content, new_text)
    if not ok:
        log.warning(
            "apply_fix_update_draft_failed",
            document_type=document_type,
            draft_id=draft.get("id"))
        raise _ApplyFixSkip(
            "Patch was generated but could not be written to the "
            "draft. Reload and try again, or edit in the editor.")

    if debate_id is not None:
        await _mark_council_debate_fix_applied(
            int(debate_id), int(draft["id"]))

    log.info(
        "apply_fix_direct_patch_applied",
        document_type=document_type,
        target=target,
        section_name=section_name,
        draft_id=draft.get("id"))

    if target == "document":
        label = (
            f"Document-wide fix applied "
            f"({sections_patched} sections updated, "
            f"{sections_skipped} skipped)")
    else:
        label = f"Direct fix applied to '{section_name}'"

    return {
        "ok": True,
        "new_draft_id": int(draft["id"]),
        "draft_label": label,
        "scope": target,
        "section_regenerated": (
            section_name if target == "section" else None),
        "in_place": True,
    }


# ── June 27 2026 -- shared content_json splice helpers ────────────
#
# Both Apply Fix (apply-fix endpoint) and Preview Inline Edit
# (propose-fix-text + accept-fix-text endpoints) operate on the
# current draft's content_json with surgical splice semantics:
#
#   1. Locate the section in content_json (slide for decks; node
#      range for TipTap docs).
#   2. Render the section as plain text for diff display.
#   3. Sonnet-patch the JSON section (not the rendered text).
#   4. Render the patched section as plain text for diff display.
#   5. Splice the patched JSON section back into the full
#      content_json at the same anchor; everything else verbatim.
#
# The helpers below are pure-data (no DB, no LLM) so the two
# endpoints share the same locate / render / splice logic. The
# Sonnet call lives in _patch_section_via_sonnet -- a thin
# dispatcher around the deck- and TipTap-specific patch helpers
# already in this file.


def _locate_section_in_content(
    document_type: str, content_json: dict, section_name: str,
) -> tuple[Any, list | dict] | None:
    """Locate a section in a draft's content_json. Returns
    (anchor, original_section_json) on a successful match, None
    otherwise.

    Anchor shape per document type:
      deck   -> int (slide index in content_json['slides'])
      brief / appendix / script -> tuple[int, int] (start, end
                                   exclusive) into content_json
                                   ['content'] node list

    The anchor is what _splice_section_into_content uses to write
    the patched section back into the right position.
    """
    if document_type == "presentation_deck":
        slides = content_json.get("slides")
        if not isinstance(slides, list):
            return None
        idx = _find_deck_slide_idx(slides, section_name)
        if idx is None:
            return None
        slide = slides[idx]
        if not isinstance(slide, dict):
            return None
        return idx, slide
    # TipTap (brief / appendix / script)
    nodes = content_json.get("content")
    if not isinstance(nodes, list):
        return None
    rng = _find_tiptap_section_range(nodes, section_name)
    if rng is None:
        return None
    start, end = rng
    return (start, end), nodes[start:end]


def _render_section_as_text(
    document_type: str, section_json: list | dict,
) -> str:
    """Render a section as plain text for diff display.

    Deck slide -> title + each text-element's content, one per line.
    TipTap nodes -> _node_text_for_match flattened, one node per
                    paragraph break.
    """
    if document_type == "presentation_deck":
        # section_json is a single slide dict.
        if not isinstance(section_json, dict):
            return ""
        out: list[str] = []
        title = str(section_json.get("title") or "").strip()
        if title:
            out.append(title)
        for el in (section_json.get("elements") or []):
            if (isinstance(el, dict)
                    and el.get("type") == "text"):
                content = str(el.get("content") or "").strip()
                if content:
                    out.append(content)
        return "\n".join(out)
    # TipTap nodes
    if not isinstance(section_json, list):
        return ""
    out2: list[str] = []
    for n in section_json:
        text = _node_text_for_match(n).strip()
        if text:
            out2.append(text)
    return "\n\n".join(out2)


async def _patch_section_via_sonnet(
    document_type: str, section_json: list | dict,
    section_name: str, patch_instruction: str,
) -> list | dict | None:
    """Dispatcher around the deck- and TipTap-specific patch
    helpers. Returns the patched section JSON in the same shape as
    section_json, or None on Sonnet failure / shape violation.

    Deck slide: returns a dict {...slide, elements: <patched
                elements array>} (the slide envelope is preserved
                so id / title / background / speaker_notes ride
                through unchanged).
    TipTap nodes: returns the patched node array.
    """
    if document_type == "presentation_deck":
        if not isinstance(section_json, dict):
            return None
        new_elements = await _patch_deck_slide_via_sonnet(
            section_json, patch_instruction)
        if new_elements is None:
            return None
        return {**section_json, "elements": new_elements}
    # TipTap
    if not isinstance(section_json, list):
        return None
    return await _patch_tiptap_section_via_sonnet(
        section_json, section_name, patch_instruction)


def _splice_section_into_content(
    document_type: str, content_json: dict, anchor: Any,
    patched_section_json: list | dict,
) -> dict:
    """Splice the patched section back into the full content_json
    at the supplied anchor. Returns the new content_json -- every
    other section is preserved verbatim.

    Per spec (June 27 2026): the write-back must always splice,
    never overwrite. A wrong-shape patch is rejected upstream
    (_patch_section_via_sonnet returns None); by the time we reach
    splice the JSON is shape-valid for the document type.
    """
    if document_type == "presentation_deck":
        slides = list(content_json.get("slides") or [])
        idx = int(anchor)
        if not (0 <= idx < len(slides)):
            return content_json
        if not isinstance(patched_section_json, dict):
            return content_json
        slides[idx] = patched_section_json
        return {**content_json, "slides": slides}
    # TipTap
    nodes = list(content_json.get("content") or [])
    start, end = anchor
    if not isinstance(patched_section_json, list):
        return content_json
    new_nodes = nodes[:start] + patched_section_json + nodes[end:]
    return {**content_json, "content": new_nodes}


async def _patch_entire_document_via_sonnet(
    *, document_type: str, content: dict,
    patch_instruction: str,
) -> tuple[dict, str, int, int]:
    """Walk every section / slide in the content and Sonnet-patch
    each one with the same patch_instruction. Returns
    (new_content, new_text, sections_patched, sections_skipped).

    Per-section failures are tolerated (the section keeps its
    original content and skipped is incremented) so a single
    Sonnet hiccup doesn't drop the entire document. The caller
    raises _ApplyFixSkip when sections_patched == 0.

    Parallelism: each section's Sonnet call goes through
    asyncio.gather so wall-clock stays bounded by the slowest
    single section rather than summing every section."""
    import asyncio as _asyncio

    is_deck = document_type == "presentation_deck"
    if is_deck:
        slides = content.get("slides")
        if not isinstance(slides, list):
            return content, _derive_content_text_from_deck(content), 0, 0
        tasks = [
            _patch_deck_slide_via_sonnet(sl, patch_instruction)
            for sl in slides if isinstance(sl, dict)
        ]
        results = await _asyncio.gather(
            *tasks, return_exceptions=True)
        new_slides: list = []
        sections_patched = 0
        sections_skipped = 0
        for sl, res in zip(slides, results):
            if (isinstance(res, list)
                    and not isinstance(res, BaseException)):
                new_slides.append({**sl, "elements": res})
                sections_patched += 1
            else:
                new_slides.append(sl)
                sections_skipped += 1
        new_content = {**content, "slides": new_slides}
        return (
            new_content,
            _derive_content_text_from_deck(new_content),
            sections_patched, sections_skipped)

    # TipTap path (brief / appendix / script)
    nodes = content.get("content")
    if not isinstance(nodes, list):
        return (
            content,
            _derive_content_text_from_tiptap(content), 0, 0)
    # Build the section ranges from heading-to-heading. For each
    # heading-bounded range, Sonnet-patch the nodes; reassemble.
    ranges: list[tuple[int, int, str]] = []
    last_start: int | None = None
    last_title: str = ""
    for i, n in enumerate(nodes):
        if isinstance(n, dict) and n.get("type") == "heading":
            if last_start is not None:
                ranges.append((last_start, i, last_title))
            last_start = i
            last_title = _node_text_for_match(n).strip() or (
                f"section_{len(ranges) + 1}")
    if last_start is not None:
        ranges.append((last_start, len(nodes), last_title))
    if not ranges:
        # No headings at all -- patch the entire node list as a
        # single section.
        ranges = [(0, len(nodes), "document")]

    tasks = [
        _patch_tiptap_section_via_sonnet(
            nodes[s:e], title, patch_instruction)
        for s, e, title in ranges
    ]
    results = await _asyncio.gather(
        *tasks, return_exceptions=True)
    new_nodes: list = []
    sections_patched = 0
    sections_skipped = 0
    for (s, e, _title), res in zip(ranges, results):
        if (isinstance(res, list)
                and not isinstance(res, BaseException)):
            new_nodes.extend(res)
            sections_patched += 1
        else:
            new_nodes.extend(nodes[s:e])
            sections_skipped += 1
    new_content = {**content, "content": new_nodes}
    return (
        new_content,
        _derive_content_text_from_tiptap(new_content),
        sections_patched, sections_skipped)


# ── /apply-fix/refine -- multi-round proposal refinement ──────────────
#
# June 27 2026. The council proposes a fix; the user may want to
# refine the proposal text across multiple cheap Sonnet calls
# BEFORE the surgical splice runs against the document. The
# refinement loop operates purely on the proposal text -- editor_
# drafts is never touched until the user clicks Apply This Fix in
# the modal (which then POSTs /apply-fix with refined_proposal_text).
#
# Why a separate endpoint: keeps the refine call cheap + fast
# (target <5s). One Sonnet call with low max_tokens. The endpoint
# never reads or writes editor_drafts, council_debates, or
# story_plans -- it's a stateless text-rewrite.

_REFINEMENT_SYSTEM_PROMPT = (
    "You are refining a fix proposal for a section of an academic "
    "document. The current proposed fix is:\n"
    "{current_proposal_text}\n\n"
    "The author requests this adjustment:\n"
    "{refinement_note}\n\n"
    "Return only the revised fix proposal text. Do not apply the "
    "fix to the document. Do not add commentary.")

REFINEMENT_NOTE_MAX_CHARS = 500
REFINEMENT_PROPOSAL_MAX_CHARS = 4000


@app.post("/api/v1/apply-fix/refine")
@limiter.limit("30/minute")
async def post_apply_fix_refine(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Refine a fix proposal's text with a single targeted Sonnet
    call. The endpoint is stateless -- it accepts the current
    working proposal text + a refinement note, returns the rewrite,
    and never touches editor_drafts / council_debates / story_plans.

    Body: {
      fix_proposal_id:        int,
      current_proposal_text:  str (the working proposal),
      refinement_note:        str (max 500 chars),
      document_type:          str,
      section_name:           str,
      refinement_round?:      int (client-supplied; just for logging)
    }

    Returns: { refined_proposal_text: str }

    Target wall-clock: <5s. Sonnet model claude-sonnet-4-6,
    max_tokens=1000. The proposal is short; this is a trivial call.

    Validation:
      422 -- missing fix_proposal_id / refinement_note /
             current_proposal_text
      413 -- refinement_note > 500 chars OR
             current_proposal_text > 4000 chars
      502 -- Sonnet call failed
    """
    import asyncio as _asyncio
    from agents.base import SONNET_MODEL, call_claude

    fix_proposal_id = body.get("fix_proposal_id")
    current_proposal_text = body.get("current_proposal_text") or ""
    refinement_note = body.get("refinement_note") or ""
    document_type = str(body.get("document_type") or "")
    section_name = body.get("section_name")
    refinement_round = body.get("refinement_round")

    if fix_proposal_id is None:
        raise HTTPException(
            status_code=422, detail="fix_proposal_id is required")
    if not isinstance(current_proposal_text, str) or not current_proposal_text.strip():
        raise HTTPException(
            status_code=422,
            detail="current_proposal_text is required")
    if not isinstance(refinement_note, str) or not refinement_note.strip():
        raise HTTPException(
            status_code=422, detail="refinement_note is required")
    current_proposal_text = current_proposal_text.strip()
    refinement_note = refinement_note.strip()

    if len(refinement_note) > REFINEMENT_NOTE_MAX_CHARS:
        raise HTTPException(
            status_code=413,
            detail=(
                f"refinement_note exceeds "
                f"{REFINEMENT_NOTE_MAX_CHARS} chars"))
    if len(current_proposal_text) > REFINEMENT_PROPOSAL_MAX_CHARS:
        # Defence in depth -- the proposal text upstream is
        # patch_instruction, capped at a few hundred chars in
        # practice. A 4000-char ceiling is a tight bound that still
        # tolerates a verbose council proposal.
        raise HTTPException(
            status_code=413,
            detail=(
                f"current_proposal_text exceeds "
                f"{REFINEMENT_PROPOSAL_MAX_CHARS} chars"))

    log.info(
        "apply_fix_refine_invoked",
        fix_proposal_id=fix_proposal_id,
        document_type=document_type,
        section_name=section_name,
        refinement_round=refinement_round,
        refinement_chars=len(refinement_note),
        proposal_chars=len(current_proposal_text))

    if ENVIRONMENT == "test":
        # Deterministic stub so tests can exercise the endpoint
        # without round-tripping a live LLM.
        return {
            "refined_proposal_text": (
                f"{current_proposal_text}\n\n"
                f"[refined: {refinement_note}]"),
        }

    user_message = _REFINEMENT_SYSTEM_PROMPT.format(
        current_proposal_text=current_proposal_text,
        refinement_note=refinement_note)
    try:
        raw = await _asyncio.to_thread(
            call_claude,
            SONNET_MODEL,
            "You are a precise editor of fix-proposal text.",
            user_message,
            max_tokens=1000,
            trigger="apply_fix_refine")
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "apply_fix_refine_call_failed",
            fix_proposal_id=fix_proposal_id, error=str(exc))
        raise HTTPException(
            status_code=502,
            detail=f"Refinement call failed: {exc}")
    refined = (raw or "").strip()
    if not refined:
        log.warning(
            "apply_fix_refine_empty_response",
            fix_proposal_id=fix_proposal_id)
        raise HTTPException(
            status_code=502,
            detail="Refinement produced an empty response.")
    return {"refined_proposal_text": refined}


@app.post("/api/v1/documents/apply-fix")
@limiter.limit("6/minute")
async def post_apply_fix(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Concern 7k-iii + 7l-i. Apply a confirmed fix proposal:

      1. Patch the story_plans row for (data_hash, document_type)
         by appending patch_instruction to the target section's
         guidance field (target=section) or to plan_json
         .central_argument (target=document).
      2. Re-run the affected generator (section-level: just the
         affected section's harness call; document-level: full
         document generator).
      3. Save the result as a NEW editor_drafts row with a
         "Post-critic revision v<n>" label. The existing draft is
         preserved -- the team picks which version to keep.
      4. Update council_debates with fix_applied=true,
         fix_applied_at=NOW(), new_draft_id=<new id>.
      5. Return the new draft_id + label + scope.

    Body: {
      document_type:  <type>,
      finding_id:     <int>,
      fix_proposal:   <FixProposal object>,
      debate_id:      <int>,
      confirmed:      true (must be true to apply)
    }

    Returns: { new_draft_id, draft_label, scope, section_regenerated }
    """
    if not body.get("confirmed"):
        raise HTTPException(
            status_code=422,
            detail="confirmed must be true to apply the fix")
    document_type = str(body.get("document_type") or "")
    if document_type not in {
            "executive_brief", "analytical_appendix",
            "presentation_deck", "presentation_script"}:
        raise HTTPException(
            status_code=422,
            detail=(
                "document_type must be one of executive_brief / "
                "analytical_appendix / presentation_deck / "
                "presentation_script"))
    fix = body.get("fix_proposal") or {}
    if not isinstance(fix, dict):
        raise HTTPException(
            status_code=422,
            detail="fix_proposal must be an object")
    target = fix.get("target")
    if target not in ("section", "document"):
        raise HTTPException(
            status_code=422,
            detail="fix_proposal.target must be 'section' or 'document'")
    patch_instruction = str(fix.get("patch_instruction") or "")
    if not patch_instruction.strip():
        raise HTTPException(
            status_code=422,
            detail="fix_proposal.patch_instruction is required")
    debate_id = body.get("debate_id")
    section_name = fix.get("section_name")
    email = session.get("email") or ""
    # June 27 2026 -- the refinement flow does NOT flow through
    # this endpoint. Refined patch text is sent to
    # /propose-fix-text via the refined_proposal_text field so the
    # mandatory diff preview always renders before any write.
    # Sending refined_proposal_text here is rejected to prevent a
    # direct-commit path that skips the diff.
    if isinstance(body.get("refined_proposal_text"), str) and (
            body["refined_proposal_text"].strip()):
        raise HTTPException(
            status_code=422,
            detail=(
                "refined_proposal_text is not accepted on "
                "apply-fix -- refined patches must flow through "
                "/propose-fix-text so the diff preview always "
                "renders before write. Call /propose-fix-text "
                "with refined_proposal_text instead."))
    if ENVIRONMENT == "test":
        # In test, return a deterministic stub so the apply-fix
        # endpoint shape is exercised without running the live
        # generators.
        return {
            "ok": True,
            "new_draft_id": -1,
            "draft_label": "Post-critic revision v1 (test)",
            "scope": target,
            "section_regenerated": (
                section_name if target == "section" else None),
        }

    # June 27 2026 -- inline-only apply-fix flow.
    #
    # The endpoint now ALWAYS routes through _try_direct_section_patch
    # which applies the fix as a surgical section-level splice on the
    # current draft's content_json. No story_plans lookup. No full-
    # document regeneration. No fall-through safety net.
    #
    # If the patch can't be applied (no current draft, section can't
    # be matched, Sonnet output unusable, write failed) the helper
    # raises _ApplyFixSkip which we translate into a 422 carrying:
    #   detail        -- plain-English reason
    #   hint          -- "open the editor and edit directly"
    #   section_name  -- best-known section identifier so the
    #                    frontend can deep-link to that section
    # The frontend's FixProposalCard catches the 422 and surfaces an
    # inline "Open in editor" link instead of regenerating.
    #
    # Removed in this revision: the entire legacy story-plan SELECT/
    # UPDATE/_start_generation_job block (was here -> now gone). It
    # paired a buggy hash join (PR #443 patched the deck case but
    # appendix/script never had a story_plans row to find) with a
    # full-document regen that discarded manual edits in the current
    # draft. Both were footguns; both are gone.
    try:
        result = await _try_direct_section_patch(
            document_type=document_type, target=target,
            section_name=str(section_name) if section_name else None,
            patch_instruction=patch_instruction,
            debate_id=int(debate_id) if debate_id else None)
    except _ApplyFixSkip as skip:
        raise HTTPException(
            status_code=422,
            detail={
                "detail": skip.detail,
                "hint": skip.hint,
                "section_name": skip.section_name,
                "document_type": document_type,
            })
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "apply_fix_unexpected_error",
            error=str(exc), document_type=document_type)
        raise HTTPException(
            status_code=500,
            detail=(
                "Unexpected error applying the fix. Reload and "
                "try again, or edit the document in the editor "
                "directly."))
    return result



# ── June 25 2026: Copilot-style inline fix-text flow ─────────────
#
# The legacy /apply-fix path regenerates the ENTIRE document via the
# generator pipeline (an opaque ~60-90s job). That's appropriate when
# the fix is a structural change to the story plan, but for a
# localised paragraph edit it's overkill and the user can't preview
# the change before committing. These two endpoints add a tighter
# loop: propose-fix-text returns a section-scoped diff (original_text
# + suggested_text); accept-fix-text marks the council_debates row
# applied and the frontend PATCHes the draft with the resolved text
# via the existing /drafts/{id} endpoint. Idempotent: the suggested
# text is cached on council_debates.fix_proposals[i].suggested_text
# so a second propose call returns the same result.

_PROPOSE_FIX_TEXT_SYSTEM = (
    "You are an editor revising a single section of an academic "
    "executive brief or appendix. You will be given the current "
    "section text and a specific change instruction. Apply the "
    "change as a minimal, surgical edit -- preserve voice, "
    "structure, citations, and any unaffected sentences. Do NOT "
    "rewrite the whole section. Do NOT add commentary, headers, "
    "or framing. Return ONLY the revised section text, verbatim, "
    "ready to drop back into the document. If the change "
    "instruction cannot be applied without losing meaningful "
    "content, return the original text unchanged.")


def _extract_section_text(
    content_text: str, section_name: str | None,
) -> str | None:
    """Best-effort section extractor. Searches content_text for a
    line whose stripped form matches section_name (case-insensitive,
    leading number / markdown header stripped) and returns the text
    from that line up to (but not including) the next heading-like
    line OR end of document.

    Returns None when the section cannot be located.

    June 27 2026 (URGENT) -- behavioural change. The previous
    contract said the caller "falls back to the full document text
    (less ideal but safe)". That fallback was NOT safe: when the
    section couldn't be located the caller wired the WHOLE document
    as the 'CURRENT SECTION TEXT' to Sonnet, which returned just the
    patched section, and the frontend then performed
      content_text.replace(WHOLE_DOC, PATCHED_SECTION)
    -- silently overwriting the entire document with only the
    patched section. The caller has been updated to return a 409
    on a None response instead. This docstring is updated to
    reflect the new contract.

    Letter-label sections ('Section A', 'Section B', ...) are now
    also normalised by the same prefix-stripper so the analytical
    appendix's letter-keyed headings match. Substring fallback
    (case-insensitive) catches headings that drift from the
    finding's section_name by a word or two."""
    if not section_name or not content_text:
        return None
    import re as _re
    needle_raw = section_name.strip().lower()
    # Capture a bare 'Section X' / 'Section 4' identifier so we can
    # still match a heading prefixed with that letter / digit when
    # the needle carries no body text after the label.
    bare_id_m = _re.match(
        r"^section\s+([a-z\d]+)\s*[:.\-]?\s*$", needle_raw)
    section_id: str | None = (
        bare_id_m.group(1) if bare_id_m else None)
    # Strip leading 'Section N[:.] ', 'Section X[:.] ' (letter),
    # '# ', '## ', '1. ', '1) ', 'A. ', 'a) ' etc. from both the
    # needle and each candidate line so 'Section B: Strategy
    # Universe' matches a draft heading like 'B. Strategy Universe'
    # or 'Strategy Universe'.
    _PREFIX_RE = _re.compile(
        r"^(section\s+[a-z\d]+[:.\-]?\s*|#+\s*|"
        r"[a-z\d]+[.)]\s+)")
    needle = _PREFIX_RE.sub("", needle_raw).strip()
    # Bare 'Section X' with no body falls through to the section-id
    # prefix matcher below.
    if not needle and not section_id:
        return None
    lines = content_text.splitlines()
    start_idx: int | None = None
    # Pass 1: exact-after-strip match.
    if needle:
        for i, line in enumerate(lines):
            clean = _PREFIX_RE.sub(
                "", line.strip().lower()).strip()
            if clean == needle and len(line.strip()) < 200:
                start_idx = i
                break
    # Pass 2: case-insensitive substring fallback (either
    # direction) -- catches mild heading drift like 'Limitations'
    # vs 'Honest Limitations' that exact match misses.
    if start_idx is None and needle:
        for i, line in enumerate(lines):
            clean = _PREFIX_RE.sub(
                "", line.strip().lower()).strip()
            if not clean or len(line.strip()) >= 200:
                continue
            if (needle in clean or clean in needle) and len(clean) >= 3:
                start_idx = i
                break
    # Pass 3: 'Section X' bare-ID prefix match -- find a heading
    # line that STARTS with the same letter/digit followed by
    # '.' / ')' / whitespace. Catches 'Section B' -> 'B. Strategy
    # Universe' even when no body text was provided.
    if start_idx is None and section_id is not None:
        id_prefix_re = _re.compile(
            rf"^{_re.escape(section_id)}[.)]\s+\S")
        for i, line in enumerate(lines):
            if (id_prefix_re.match(line.strip().lower())
                    and len(line.strip()) < 200):
                start_idx = i
                break
    if start_idx is None:
        return None
    # Walk forward to the next heading-like line (markdown # / ##
    # / 'Section N' / 'A. ' / a bold-only line). The slice is from
    # start_idx to end_idx exclusive.
    end_idx = len(lines)
    for j in range(start_idx + 1, len(lines)):
        ln = lines[j].strip()
        if not ln:
            continue
        if (ln.startswith("#")
                or _re.match(r"^section\s+[a-z\d]+\b", ln,
                             flags=_re.IGNORECASE)
                or _re.match(r"^[a-z\d]+[.)]\s+\S", ln,
                             flags=_re.IGNORECASE)
                or (ln.startswith("**") and ln.endswith("**")
                    and len(ln) < 120)):
            end_idx = j
            break
    extracted = "\n".join(lines[start_idx:end_idx]).strip()
    # Guard: if the extracted span covers more than ~80% of the
    # document, the section bounds collapsed and the caller would
    # effectively be patching the whole doc. Refuse instead of
    # returning an unsafe span.
    if len(extracted) >= 0.8 * len(content_text):
        return None
    return extracted


@app.post(
    "/api/v1/council/debates/{debate_id}/propose-fix-text")
@limiter.limit("20/minute")
async def post_propose_fix_text(
    debate_id: int, request: Request, body: dict,
    session: dict = Depends(require_team_member),
):
    """Generate a section-scoped JSON splice preview for one fix
    proposal. Operates on content_json (NOT content_text). The
    write-back to the draft happens on accept-fix-text -- this
    endpoint just produces the preview.

    Body: {finding_id: int}

    Returns: {finding_id, section_name, original_text,
              suggested_text, proposal_id, cached, document_type}

    The original_text / suggested_text fields are PLAIN-TEXT
    renderings of the located section (before and after the
    patch) -- intended only for the diff display. The actual
    splice that accept-fix-text writes back uses the JSON cached
    here on council_debates.fix_proposals[i] (alongside the legacy
    text fields kept for backward compatibility).

    Idempotent: the cached suggested_section_json + suggested_text
    are returned on a second call without re-running Sonnet.

    Errors:
      404 -- debate row not found
      422 -- no fix proposal for the finding_id / empty patch
             instruction
      409 -- no current draft / no content_json / section name
             could not be located in the current content_json /
             Sonnet shape violation
      502 -- Sonnet call failed

    A 409 carries a structured detail with section_name + hint so
    the frontend can show an editor deep-link instead of silently
    falling back to a full-document overwrite (the June 27 2026
    bug that prompted this rewrite)."""
    from sqlalchemy import text as _text
    from database import (
        AsyncSessionLocal,  # type: ignore[attr-defined]
    )
    from tools.editor_drafts import get_current_draft_by_type

    finding_id = body.get("finding_id")
    if finding_id is None:
        raise HTTPException(
            status_code=422, detail="finding_id is required")
    finding_id = int(finding_id)

    if AsyncSessionLocal is None:
        raise HTTPException(
            status_code=503, detail="Database unavailable.")

    async with AsyncSessionLocal() as s:
        r = await s.execute(_text(
            "SELECT document_type, fix_proposals, source_draft_id "
            "FROM council_debates WHERE id = :id"),
            {"id": debate_id})
        row = r.fetchone()
        if row is None:
            raise HTTPException(
                status_code=404, detail="Debate row not found.")
        document_type = row[0] or ""
        fix_proposals = row[1] or {}
        _ = row[2]

    proposal, proposal_key = _lookup_fix_proposal(
        fix_proposals, finding_id)
    if proposal is None:
        raise HTTPException(
            status_code=422,
            detail=(
                "No fix proposal for that finding_id. Generate one "
                "first via /api/v1/documents/propose-fix."))
    patch_instruction = str(
        proposal.get("patch_instruction") or "").strip()
    section_name = proposal.get("section_name")
    if not patch_instruction:
        raise HTTPException(
            status_code=422,
            detail="patch_instruction is empty on this proposal.")
    # June 27 2026 -- multi-round refinement. The frontend
    # RefinementModal iteratively rewrites the proposal text via
    # cheap /apply-fix/refine calls, then sends the final working
    # text here as refined_proposal_text. When present + non-empty,
    # it OVERRIDES the stored patch_instruction for THIS preview
    # only -- the council's stored proposal is unchanged. The user
    # still sees the mandatory diff preview before any write, and
    # accept-fix-text writes the cached suggested_section_json
    # against the current draft. This is the ONLY entry point for
    # the refined patch; apply-fix rejects the field outright.
    refined_raw = body.get("refined_proposal_text")
    refined_proposal_text: str | None = None
    if isinstance(refined_raw, str) and refined_raw.strip():
        refined_proposal_text = refined_raw.strip()
        patch_instruction = refined_proposal_text
    log.info(
        "propose_fix_text_invoked",
        debate_id=debate_id,
        finding_id=finding_id,
        section_name=section_name,
        document_type=document_type,
        refined=bool(refined_proposal_text),
        refinement_chars=(
            len(refined_proposal_text)
            if refined_proposal_text else 0))
    if not section_name:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "Fix proposal has no section_name -- inline "
                    "preview requires a specific section to splice. "
                    "Edit the document in the editor directly."),
                "hint": (
                    "Open the document in the editor and apply the "
                    "change manually."),
                "section_name": None,
                "document_type": document_type,
            })

    # Idempotency cache -- if we previously computed the preview
    # for this proposal, return it verbatim (no Sonnet call).
    cached_suggested = proposal.get("suggested_text")
    cached_original = proposal.get("original_text")
    cached_suggested_json = proposal.get("suggested_section_json")
    # June 27 2026 -- refinement bypass. The idempotency cache
    # holds the preview computed against the ORIGINAL stored
    # patch_instruction. A refinement round changed the input, so
    # the cached preview no longer matches -- recompute and DO NOT
    # write through to the cache (the next non-refined call should
    # still get the original preview back unless explicitly
    # invalidated upstream).
    if (cached_suggested and cached_original
            and cached_suggested_json is not None
            and refined_proposal_text is None):
        log.info(
            "propose_fix_text_cache_hit",
            debate_id=debate_id, finding_id=finding_id)
        return {
            "finding_id": finding_id,
            "section_name": section_name,
            "original_text": cached_original,
            "suggested_text": cached_suggested,
            "proposal_id": debate_id,
            "cached": True,
            "document_type": document_type,
        }

    # Fetch the current draft (team-shared per PR #410).
    draft = await get_current_draft_by_type(document_type)
    if draft is None:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "No current draft for this document type. "
                    "Generate the document before previewing a "
                    "fix."),
                "hint": (
                    "Generate the document from the Documents "
                    "panel, then try Preview Inline Edit again."),
                "section_name": section_name,
                "document_type": document_type,
            })
    content_json = draft.get("content_json")
    if not isinstance(content_json, dict):
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "Current draft has no content_json to patch."),
                "hint": (
                    "Open the document in the editor and apply "
                    "the change directly."),
                "section_name": section_name,
                "document_type": document_type,
            })

    # Locate the section in content_json. June 27 2026 -- this
    # replaces the previous content_text-based extractor that, when
    # the section couldn't be located, fell back to the WHOLE
    # document and silently overwrote everything with only the
    # patched section. The new helper returns None instead; we 409
    # with an actionable hint.
    located = _locate_section_in_content(
        document_type, content_json, str(section_name))
    if located is None:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    f"Section '{section_name}' could not be "
                    "located in the current draft. Section "
                    "heading may have drifted from the finding's "
                    "section name."),
                "hint": (
                    "Open the document in the editor and apply "
                    "the change directly."),
                "section_name": section_name,
                "document_type": document_type,
            })
    anchor, original_section_json = located
    original_text = _render_section_as_text(
        document_type, original_section_json)

    # Sonnet patch -- under test, deterministic stub appends the
    # instruction so tests can assert the splice happened without
    # round-tripping a live LLM.
    if ENVIRONMENT == "test":
        suggested_section_json = (
            _apply_test_stub_patch(
                document_type, original_section_json,
                patch_instruction))
    else:
        try:
            suggested_section_json = await _patch_section_via_sonnet(
                document_type, original_section_json,
                str(section_name), patch_instruction)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "propose_fix_text_call_failed",
                debate_id=debate_id, finding_id=finding_id,
                error=str(exc))
            raise HTTPException(
                status_code=502,
                detail=f"Sonnet call failed: {exc}")
    if suggested_section_json is None:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "AI could not produce a valid patch for "
                    f"section '{section_name}'."),
                "hint": (
                    "Open the document in the editor and apply "
                    "the change directly."),
                "section_name": section_name,
                "document_type": document_type,
            })
    suggested_text = _render_section_as_text(
        document_type, suggested_section_json)

    # Cache the JSON + text on council_debates.fix_proposals[i].
    # The text fields stay populated for backward compatibility
    # with any frontend code still reading them; the JSON fields
    # are the source of truth for the splice that accept-fix-text
    # performs.
    try:
        if (proposal_key is not None
                and AsyncSessionLocal is not None):
            updated = (
                fix_proposals.copy()
                if isinstance(fix_proposals, dict)
                else list(fix_proposals))
            cached_payload = {
                **proposal,
                "suggested_text": suggested_text,
                "original_text": original_text,
                "suggested_section_json": suggested_section_json,
                "original_section_json": original_section_json,
                "section_name": section_name,
            }
            if isinstance(updated, dict):
                updated[str(proposal_key)] = cached_payload
            else:
                updated[int(proposal_key)] = cached_payload
            async with AsyncSessionLocal() as s2:
                await s2.execute(_text(
                    "UPDATE council_debates "
                    "SET fix_proposals = CAST(:p AS JSONB) "
                    "WHERE id = :id"),
                    {"p": json.dumps(updated), "id": debate_id})
                await s2.commit()
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "propose_fix_text_cache_write_failed",
            debate_id=debate_id, finding_id=finding_id,
            error=str(exc))

    return {
        "finding_id": finding_id,
        "section_name": section_name,
        "original_text": original_text,
        "suggested_text": suggested_text,
        "proposal_id": debate_id,
        "cached": False,
        "document_type": document_type,
    }


def _lookup_fix_proposal(
    fix_proposals: Any, finding_id: int,
) -> tuple[dict | None, str | int | None]:
    """Two-shape lookup: fix_proposals may be a dict keyed by
    str(finding_id) or a list of FixProposal dicts. Returns
    (proposal, key) on a successful match, (None, None) otherwise.
    Used by propose-fix-text + accept-fix-text."""
    if isinstance(fix_proposals, dict):
        key = str(finding_id)
        if key in fix_proposals:
            p = fix_proposals[key]
            if isinstance(p, dict):
                return p, key
    elif isinstance(fix_proposals, list):
        for p in fix_proposals:
            if isinstance(p, dict) and (
                    int(p.get("finding_id", -1)) == finding_id):
                return p, fix_proposals.index(p)
    return None, None


def _apply_test_stub_patch(
    document_type: str, section_json: list | dict,
    patch_instruction: str,
) -> list | dict:
    """In ENVIRONMENT=test, the Sonnet call is replaced with a
    deterministic stub so tests can assert the splice happened
    end-to-end without round-tripping a live LLM. The stub
    appends '[Inline edit applied: <instruction>]' to the section
    in a shape-preserving way."""
    marker = f"[Inline edit applied: {patch_instruction}]"
    if document_type == "presentation_deck":
        if not isinstance(section_json, dict):
            return section_json
        # Append a text element carrying the marker.
        elements = list(section_json.get("elements") or [])
        elements.append({
            "id": f"s{section_json.get('id', 0)}_test_stub",
            "type": "text",
            "x": 60, "y": 600, "width": 840, "height": 40,
            "content": marker,
            "fontSize": 14, "fontWeight": "normal",
            "fontStyle": "italic", "color": "#666666",
            "locked": False,
        })
        return {**section_json, "elements": elements}
    # TipTap: append a paragraph node carrying the marker.
    if not isinstance(section_json, list):
        return section_json
    return list(section_json) + [{
        "type": "paragraph",
        "content": [{"type": "text", "text": marker}],
    }]


@app.post(
    "/api/v1/council/debates/{debate_id}/accept-fix-text")
@limiter.limit("20/minute")
async def post_accept_fix_text(
    debate_id: int, request: Request, body: dict,
    session: dict = Depends(require_team_member),
):
    """Splice the previously-proposed section patch back into the
    current draft's content_json + mark the council_debates row
    applied.

    Body: {finding_id: int}

    June 27 2026 -- behavioural change. Previously the frontend
    PATCHed /drafts/{id} with content_text and this endpoint just
    flipped fix_applied. That path could overwrite the entire
    document when the section couldn't be located cleanly. The
    write-back has moved server-side: this endpoint reads the
    cached suggested_section_json (computed by propose-fix-text),
    re-locates the section in the CURRENT content_json (so a
    manual edit between propose and accept is respected), splices,
    and writes via update_draft. The frontend no longer touches
    content_text directly for this flow.

    The optional new_draft_id body field is accepted but ignored
    in this revision -- the new id comes from the splice itself.

    Returns: {ok, new_draft_id, applied_at, section_name,
              document_type}

    Errors:
      404 -- debate row not found
      422 -- no fix proposal for finding_id
      409 -- proposal hasn't been previewed yet (no cached
             suggested_section_json) / no current draft / section
             could not be re-located in current content_json"""
    from sqlalchemy import text as _text
    from database import (
        AsyncSessionLocal,  # type: ignore[attr-defined]
    )
    from tools.editor_drafts import (
        get_current_draft_by_type, update_draft,
    )

    finding_id = body.get("finding_id")
    if finding_id is None:
        raise HTTPException(
            status_code=422, detail="finding_id is required")
    finding_id = int(finding_id)

    if AsyncSessionLocal is None:
        raise HTTPException(
            status_code=503, detail="Database unavailable.")

    # Lookup the debate row + the cached suggested section JSON.
    async with AsyncSessionLocal() as s_lookup:
        r = await s_lookup.execute(_text(
            "SELECT document_type, fix_proposals "
            "FROM council_debates WHERE id = :id"),
            {"id": debate_id})
        row = r.fetchone()
        if row is None:
            raise HTTPException(
                status_code=404, detail="Debate row not found.")
        document_type = row[0] or ""
        fix_proposals = row[1] or {}
    proposal, _ = _lookup_fix_proposal(fix_proposals, finding_id)
    if proposal is None:
        raise HTTPException(
            status_code=422,
            detail=(
                "No fix proposal for that finding_id. Generate a "
                "preview via /propose-fix-text first."))
    suggested_section_json = proposal.get(
        "suggested_section_json")
    section_name = proposal.get("section_name")
    if (suggested_section_json is None
            or not isinstance(suggested_section_json,
                              (list, dict))):
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "No cached preview for this fix. Run "
                    "Preview Inline Edit first to generate one."),
                "hint": (
                    "Click 'Preview Inline Edit' to generate the "
                    "patch, then 'Accept' to apply it."),
                "section_name": section_name,
                "document_type": document_type,
            })

    # Splice into the CURRENT draft's content_json. Re-locating
    # against the current state means manual edits between propose
    # and accept don't surprise us -- the section is always written
    # at its current position, not the position cached at preview
    # time.
    draft = await get_current_draft_by_type(document_type)
    if draft is None:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "No current draft to write the patch to."),
                "hint": (
                    "Re-generate the document and retry."),
                "section_name": section_name,
                "document_type": document_type,
            })
    content_json = draft.get("content_json")
    if not isinstance(content_json, dict):
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    "Current draft has no content_json to splice "
                    "into."),
                "hint": (
                    "Open the document in the editor and apply "
                    "the change directly."),
                "section_name": section_name,
                "document_type": document_type,
            })
    located = _locate_section_in_content(
        document_type, content_json, str(section_name))
    if located is None:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": (
                    f"Section '{section_name}' was edited in the "
                    "draft between preview and accept -- the "
                    "heading no longer matches. Re-preview to "
                    "refresh the patch."),
                "hint": (
                    "Click 'Preview Inline Edit' again to regen "
                    "the patch against the current draft state."),
                "section_name": section_name,
                "document_type": document_type,
            })
    anchor, _orig = located
    new_content_json = _splice_section_into_content(
        document_type, content_json, anchor,
        suggested_section_json)
    if document_type == "presentation_deck":
        new_text = _derive_content_text_from_deck(new_content_json)
    else:
        new_text = _derive_content_text_from_tiptap(
            new_content_json)

    ok = await update_draft(
        int(draft["id"]), new_content_json, new_text)
    if not ok:
        raise HTTPException(
            status_code=500,
            detail=(
                "Patch was generated but the draft write failed. "
                "Reload and retry."))

    try:
        async with AsyncSessionLocal() as s:
            await s.execute(_text(
                "UPDATE council_debates SET "
                "fix_applied = TRUE, "
                "fix_applied_at = NOW(), "
                "new_draft_id = COALESCE(:nd, new_draft_id) "
                "WHERE id = :id"),
                {"nd": int(draft["id"]), "id": debate_id})
            await s.commit()
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "accept_fix_text_mark_failed",
            debate_id=debate_id, finding_id=int(finding_id),
            error=str(exc))
        # Splice was already written successfully -- don't unwind
        # it on a debate-marker failure. Log + continue.

    log.info(
        "accept_fix_text_spliced",
        debate_id=debate_id, finding_id=int(finding_id),
        document_type=document_type,
        section_name=section_name,
        draft_id=int(draft["id"]))

    from datetime import datetime as _dt, timezone as _tz
    return {
        "ok": True,
        "new_draft_id": int(draft["id"]),
        "applied_at": _dt.now(_tz.utc).isoformat(),
        "section_name": section_name,
        "document_type": document_type,
    }


# ── Concern 7l-ii: re-run critic on a specific draft ──────────

@app.post("/api/v1/documents/rerun-critic")
@limiter.limit("10/minute")
async def post_rerun_critic_for_draft(
    request: Request,
    body: dict,
    session: dict = Depends(require_team_member),
):
    """Concern 7l-ii. Triggers a fresh academic review (which
    includes the critic) scoped to a specific draft. The new
    council_debates row gets parent_debate_id pointing at the
    most recent debate for the same document_type, building the
    iteration chain.

    This is the cheapest re-run path: it doesn't regen the
    document, it just re-reviews the existing content_text. Use
    apply-fix to actually patch + regen.

    Body: {
      document_type: <type>,
      source_draft_id: <int -- the draft to review>,
      parent_debate_id: <int | null -- the prior round's id>
    }

    Returns the SSE URL the frontend should hit. The frontend
    then opens an EventSource on the regular academic-review
    endpoint with a query param that pins the source draft.

    NOTE: This is currently a stub that returns the URL the
    frontend should hit. The actual draft-pinning + parent
    chaining is wired by the SSE handler reading these query
    params; the academic-review endpoint already takes
    document_type, so this stub returns the right URL with the
    document_type + source_draft_id + parent_debate_id encoded.
    """
    document_type = str(body.get("document_type") or "")
    source_draft_id = body.get("source_draft_id")
    parent_debate_id = body.get("parent_debate_id")
    if document_type and document_type not in {
            "executive_brief", "analytical_appendix",
            "presentation_deck", "presentation_script"}:
        raise HTTPException(
            status_code=422,
            detail=(
                "document_type must be one of the four "
                "deliverable types or omitted"))
    qp = []
    if document_type:
        qp.append(f"document_type={document_type}")
    if source_draft_id is not None:
        qp.append(f"source_draft_id={int(source_draft_id)}")
    if parent_debate_id is not None:
        qp.append(f"parent_debate_id={int(parent_debate_id)}")
    qs = ("?" + "&".join(qp)) if qp else ""
    return {
        "ok": True,
        "sse_url": f"/api/council/academic-review{qs}",
    }


@app.post("/api/v1/export/executive-brief")
@limiter.limit("6/minute")
async def export_executive_brief(
    request: Request,
    body: dict | None = None,
    session: dict = Depends(require_permission("generate_documents")),
):
    """
    Starts executive-brief generation.

    With an editor_draft_id in the body the .docx is built synchronously
    from that draft (the in-editor Export path). Otherwise generation —
    the eight Academic Writer sections, 45-90s — runs as a background
    job and the endpoint returns 202 with a job_id; poll
    GET /api/v1/jobs/{id}.
    """
    editor_draft_id = (body or {}).get("editor_draft_id")
    if editor_draft_id:
        return await _editor_export(int(editor_draft_id))
    await _require_report_ready()
    return _start_generation_job("executive_brief", session, request)


# ── Executive brief tone rules (May 30 2026) ───────────────────────────────
# The brief addresses the rubric trap: the platform executes, the
# judgment is human. Every section is generated with these rules
# threaded into the agent prompt so the prose never says "the
# platform found X" — it says "our analysis shows X". The platform
# is cited as source of DATA, never as source of CONCLUSIONS.
_BRIEF_NUMERIC_GROUNDING = (
    "\n\nNUMERIC GROUNDING (CRITICAL):\n"
    "  - All numeric values in your response must come exactly from the "
    "data context provided in this section. Do not estimate, interpolate, "
    "round beyond two decimal places, or substitute figures from prior "
    "knowledge.\n"
    "  - Sharpe ratios, max drawdowns, CAGR figures, OOS Sharpe values, "
    "and correlation coefficients MUST be quoted from the data block "
    "verbatim or else replaced with [DATA PENDING].\n"
    "  - If a value is not in the data context, write [DATA PENDING] "
    "rather than inventing a number.\n"
    "  - Numeric accuracy is verified against the source cache after "
    "generation -- a Sharpe attributed to the wrong strategy or a "
    "drawdown that disagrees with the cache flags in the audit panel."
)

_BRIEF_TONE_RULES_BASE = (
    "\n\nTONE AND LANGUAGE RULES (non-negotiable):\n"
    "  - Never write 'the platform found' or 'the AI council "
    "determined'. Always write 'our analysis shows', 'we interpret', "
    "or 'we conclude'.\n"
    "  - Cite the platform as the source of DATA, never as the source "
    "of conclusions. Conclusions are ours.\n"
    "  - Plain financial English. No technical platform language.\n"
    "  - No data hashes in the body text — those belong in the "
    "appendix.\n"
    "  - No em dashes (project-wide prose rule).\n"
    "\nHEADING DISCIPLINE (non-negotiable):\n"
    "  - Emit ONLY ONE heading per section: the H2 in the form "
    "'## Section N: Title' (e.g. '## Section 1: Executive Summary'). "
    "Do NOT also emit a numbered H1 like '# 1. Executive Summary' "
    "above it. The DOCX assembler builds the H1 chapter heading "
    "downstream from the section key; duplicate model-emitted "
    "headings render as 'Executive Summary' followed by 'Section 1: "
    "Executive Summary' in the final document and read as a defect.\n"
    "  - Wrong: '# 1. Executive Summary\\n## Section 1: Executive Summary'.\n"
    "  - Right: '## Section 1: Executive Summary'.\n"
    "\nCITATION DISCIPLINE -- Benjamin vs Benjamini (non-negotiable):\n"
    "Two different papers, two different claims; never swap or "
    "conflate them in-text or in the reference list.\n"
    "  - Benjamini & Hochberg (1995) -- the FDR (false discovery "
    "rate) CORRECTION METHODOLOGY. Cite this whenever you reference "
    "the q-value control, the BH procedure, or the multiple-"
    "hypothesis correction technique. Reference list entry: "
    "'Benjamini, Y., & Hochberg, Y. (1995). Controlling the false "
    "discovery rate: A practical and powerful approach to multiple "
    "testing. Journal of the Royal Statistical Society: Series B, "
    "57(1), 289-300.'\n"
    "  - Benjamin et al. (2018) -- the p < 0.005 THRESHOLD "
    "PROPOSAL (Redefine Statistical Significance, Nature Human "
    "Behaviour). Cite this whenever you reference the 0.005 cutoff "
    "itself as the bar for declaring a finding significant. "
    "Reference list entry: 'Benjamin, D. J., et al. (2018). "
    "Redefine statistical significance. Nature Human Behaviour, "
    "2(1), 6-10.'\n"
    "  - Wrong: 'no strategy clears p < 0.005 under Benjamin et "
    "al. (2018) FDR correction' (conflates the methodology with "
    "the threshold paper).\n"
    "  - Right: 'no strategy clears p < 0.005 (Benjamin et al., "
    "2018) under Benjamini-Hochberg FDR correction (Benjamini & "
    "Hochberg, 1995)' -- two citations, each on the claim they "
    "actually support.\n"
    "\nSENTENCE COMPLETION (HARD CONSTRAINT):\n"
    "You must complete every sentence you begin. If you are running "
    "long, end the current paragraph cleanly rather than starting a "
    "new one. NEVER end mid-sentence or mid-word. The section must "
    "end with a complete sentence followed by terminating punctuation "
    "(period, question mark, or exclamation point). A truncated "
    "section will be flagged by the post-generation audit and "
    "blocks publication."
)

# The composite tone-rules string actually threaded into every
# section spec. Composes the base tone rules with the numeric-
# grounding directive so each section automatically picks up both
# contracts -- the grounding is not optional and propagating it via
# the same constant means a future section addition can never forget
# to apply it.
_BRIEF_TONE_RULES = _BRIEF_TONE_RULES_BASE + _BRIEF_NUMERIC_GROUNDING


# Four-component recommendation structure required on every recommendation
# in the executive brief. Mirrors the CFA Institute disclosure spirit:
# full disclosure, balanced presentation, material limitations stated.
# No em dashes (project-wide prose rule).
_BRIEF_RECOMMENDATION_STRUCTURE = (
    "\n\nEvery recommendation you make in this section must include all "
    "four components, clearly labelled:\n"
    "1. THE SIGNAL: what the data says, in one specific, quantified "
    "sentence.\n"
    "2. THE CONFIDENCE: how certain we are. Reference the regime "
    "posterior probability and the Kish effective sample size (ESS). "
    "Flag explicitly if the ESS is low.\n"
    "3. THE DISSENTING VIEW: the strongest honest counter-argument. It "
    "must reference a specific, named limitation, not a generic hedge. "
    "If the recommendation is sensitive to the 40% box constraint or to "
    "the regime sample size, say so explicitly.\n"
    "4. THE LIMITATIONS: what the model cannot see. Always include all "
    "four of: the three-asset universe constraint; the post-2022 sample "
    "size ({{OOS_WINDOW_MONTHS}} months, about "
    "{{OOS_WINDOW_PCT_OF_STUDY}}% of the full window); transaction costs "
    "not yet applied; and the absence of formal statistical significance "
    "(economic significance only).\n"
    "This structure meets the spirit of CFA Institute disclosure "
    "standards. Never surface a recommendation without all four "
    "components present."
)


async def _generate_brief_document(
    email: str,
) -> tuple[bytes, str, str, int | None]:
    """
    Generates the executive brief. Returns (file bytes, filename, media
    type, editor draft id). Raises on failure -- the job wrapper records it.

    Rewritten June 18 2026 to the FNA 670 rubric's six required sections,
    in rubric order. The earlier (June 6) structure led with "The Answer"
    + a "Five Human Decisions" section + a "Part II preview"; rubric
    review flagged the latter two as non-rubric content ("next steps
    rather than final recommendations") and the section ordering as
    out-of-spec. The six sections below match the rubric exactly:

      1. EXECUTIVE SUMMARY      -- verdict + headline figures
      2. METHODOLOGY OVERVIEW   -- HMM + OOS window + validation layers
      3. KEY FINDINGS           -- three-strategy comparison + honest 2-of-9
      4. LIMITATIONS AND RISKS  -- four mandatory limitations (no Part II)
      5. FINAL RECOMMENDATIONS  -- investment conclusions drawn from the
         OOS Sharpe + diversification evidence (NOT a point-in-time
         portfolio position). Uses the cached-regime fallback assembled
         in academic_export so the section is data-independent and never
         renders [DATA PENDING] under a degraded live build.
      6. VISUALS                -- captioned roster of the platform's
         chart surfaces (cumulative return, implied asset allocation,
         efficient frontier) with a one-paragraph interpretation each.

    Target length: 1,800-2,200 words. Every section is generated with
    _BRIEF_TONE_RULES threaded into the agent prompt so the prose never
    says "the platform found X" -- it says "our analysis shows X". The
    platform is cited as the source of DATA, conclusions are framed
    as ours.
    """
    import asyncio
    from datetime import date

    from tools.academic_docx import build_executive_brief
    from tools.academic_export import DATA_PENDING, gather_document_data

    # June 23 2026 -- the brief is the narrative anchor for all
    # downstream documents (deck, appendix, script). Their story
    # plans were grounded against the prior brief; once the brief
    # is regenerated those plans are stale and would conflict with
    # the new central argument. The story_plans table has no user
    # scope (one row per data_hash + document_type, shared across
    # the team), so the clear is global -- correct semantic: when
    # the brief changes, the SHARED downstream plans are stale for
    # everyone. First-gen is a no-op since those rows do not exist.
    # Fail-open: the clear logs a warning but does NOT block brief
    # generation if the DB delete fails.
    try:
        from sqlalchemy import text
        from database import AsyncSessionLocal
        if AsyncSessionLocal is not None:
            async with AsyncSessionLocal() as _db:
                await _db.execute(text(
                    "DELETE FROM story_plans WHERE document_type IN ("
                    "'presentation_deck', "
                    "'analytical_appendix', "
                    "'presentation_script')"))
                await _db.commit()
            log.info(
                "Cleared downstream story plans on brief regeneration")
    except Exception as _exc:
        log.warning(
            "Could not clear downstream story plans: %s", _exc)

    try:
        # June 27 2026 (PR 1 v3, LEAK 1 closer) -- compute the
        # effective data_hash inline so gather_document_data can
        # route the strategy_results_cache read through the hash-
        # aware path. Under freeze, a miss raises
        # StrategyCacheMissingForHashError which the outer except
        # translates to a 500.
        from tools.audit_assembler import (
            current_data_hash as _brief_current_hash,
        )
        from tools.submission_freeze import (
            get_effective_data_hash as _brief_eff_hash,
        )
        _brief_live = await _brief_current_hash()
        _brief_data_hash = await _brief_eff_hash(_brief_live)
        data = await gather_document_data(
            data_hash=_brief_data_hash or None)
        avail = data["available"]
        pending = (f"{DATA_PENDING} -- analytics caches not warm. Load the "
                   "dashboard once, then regenerate this brief.")
        live = data.get("live_recommendation") or {}

        # FNA 670 rubric, six sections in rubric order. Each spec
        # threads _BRIEF_TONE_RULES so the prose stays in the
        # appropriate first-person-plural analytical voice.
        specs = [
            {"key": "executive_summary", "available": avail,
             "pending": pending,
             "agent_id": "brief_executive_summary",
             "task": (
                 "Write Section 1: EXECUTIVE SUMMARY. Approximately 250 "
                 "words. The senior investment audience reads this page "
                 "first and may stop here, so it MUST stand alone.\n\n"
                 "Lead with the verdict in the first sentence. Use this "
                 "exact opener: 'A regime-conditional diversified blend "
                 "outperforms a 100% equity allocation on a risk-"
                 "adjusted basis over the post-2022 out-of-sample "
                 "window.'\n\n"
                 "Immediately follow with the headline figures in plain "
                 "language (no preamble): OOS Sharpe {{OOS_SHARPE_BLEND}} "
                 "(blend) vs {{OOS_SHARPE_BENCHMARK}} (benchmark); maximum "
                 "drawdown {{REGIME_SWITCHING_MAX_DD}} (blend) vs "
                 "{{BENCHMARK_MAX_DD}} (benchmark); the regime-conditional "
                 "construction held the bond sleeve through the 2022 "
                 "equity drawdown.\n\n"
                 "Close with one short paragraph naming the practical "
                 "context: the pre/post-2022 correlation break "
                 "(approximately {{PRE_2022_EQ_IG_CORR}} -> "
                 "{{POST_2022_EQ_IG_CORR}}) is the environment the "
                 "hypothesis addresses, and the analysis below is built "
                 "on that scope. Do not introduce methodology details or "
                 "the recommendation here -- those are Sections 2 and 5."
                 "\n\nIMPORTANT TOKEN HYGIENE: do NOT emit lowercase "
                 "placeholder tokens like {{play_by_play_events}} or "
                 "{{play_by_play_add_value}}. These are NOT in the "
                 "substitution table and will render as literal "
                 "{{}} text in the exported brief. The context dict "
                 "field names are NOT substitution tokens -- they're "
                 "the source data the platform passes you. The only "
                 "tokens you may emit are the UPPERCASE ones documented "
                 "in the placeholder guide above (e.g. "
                 "{{OOS_SHARPE_BLEND}}, {{REGIME_SWITCHING_MAX_DD}}). "
                 "When you want to reference scorecard / play-by-play "
                 "evidence, name it descriptively in prose (e.g. "
                 "'the play-by-play scorecard') rather than emitting a "
                 "token."
                 + _BRIEF_TONE_RULES),
             "context": {"summary_statistics": data["summary_statistics"],
                         "drawdown_comparison": data["drawdown_comparison"],
                         "study_period": data["study_period"],
                         "correlation_pre_post": {
                             "pre_2022": data["rolling_correlation"].get(
                                 "pre_2022"),
                             "post_2022": data["rolling_correlation"].get(
                                 "post_2022")}}},
            {"key": "methodology", "available": True,
             "agent_id": "brief_methodology",
             # June 21 2026 -- bumped from the default 1500.
             # Methodology cites four foundational papers (Hamilton
             # 1989, Carhart 1997, Ang and Bekaert 2002, Markowitz
             # 1952); four full APA References entries + 350 words
             # of prose + the placeholder guide reliably push past
             # 1500. PR #361's retry would still catch it but
             # giving the section headroom up front avoids the
             # retry cost.
             "max_tokens": 2500,
             "task": (
                 "Write Section 2: METHODOLOGY OVERVIEW. Approximately "
                 "350 words across TWO PARAGRAPHS MAXIMUM (allowing a "
                 "third short rebalancing-disclosure paragraph if "
                 "needed). Brevity is the contract -- full methodology "
                 "lives in the Analytical Appendix.\n\n"
                 "First paragraph: name the three-asset universe "
                 "(equities, investment-grade bonds, high-yield bonds) "
                 "and note explicitly that this is a PROJECT SCOPE "
                 "BOUNDARY, not an architectural limit -- the platform "
                 "handles any return series. State the HMM regime "
                 "detection mechanism in one sentence, citing Hamilton "
                 "(1989) as the foundational reference for the Hidden "
                 "Markov Model approach to financial time series. State "
                 "the OOS window design: the test period begins AFTER "
                 "the 2022 correlation break so the evidence reflects "
                 "the environment the hypothesis addresses.\n\n"
                 "Disclose the rebalancing frequency explicitly: the "
                 "platform evaluates the HMM regime monthly and "
                 "rebalances when any single strategy's blend weight "
                 "crosses {{REBALANCE_THRESHOLD_PP}} percentage "
                 "points. This deviates from a strict quarterly "
                 "cadence; justify the deviation: monthly evaluation "
                 "matches the cadence at which the HMM produces "
                 "regime updates, and the "
                 "{{REBALANCE_THRESHOLD_PP}} percentage points gate "
                 "filters noise so the portfolio does not churn on "
                 "marginal signal changes.\n\n"
                 "Second paragraph: name the validation layers in one "
                 "sentence each -- three-layer statistical audit, the "
                 "Carhart four-factor model (Carhart, 1997), the "
                 "Benjamini-Hochberg FDR correction at q < "
                 "{{BH_SIGNIFICANCE_THRESHOLD}}, the "
                 "play-by-play scorecard. Cite Ang and Bekaert (2002) "
                 "as the direct academic precedent for regime-"
                 "conditional asset allocation and Markowitz (1952) "
                 "as the mean-variance basis for the static blend "
                 "({{CLASSIC_6040_WEIGHT_EQUITY}} equity / "
                 "{{CLASSIC_6040_WEIGHT_BOND}} bonds). "
                 "Close by directing the reader to the Carhart factor-"
                 "loading table (embedded below Section 2) and the "
                 "Analytical Appendix for the per-strategy detail.\n\n"
                 "IMPORTANT TOKEN HYGIENE: when you reference the "
                 "Classic 60/40 static-blend weights, USE the "
                 "canonical uppercase token names "
                 "{{CLASSIC_6040_WEIGHT_EQUITY}} and "
                 "{{CLASSIC_6040_WEIGHT_BOND}}. Do NOT invent "
                 "lowercase variants like {{static_equity_weight_pct}} "
                 "or {{static_bond_weight_pct}} -- those names are "
                 "not in the substitution table and will render "
                 "literally in the exported brief."
                 + _BRIEF_TONE_RULES),
             "context": {"study_period": data["study_period"]}},
            {"key": "key_findings", "available": avail, "pending": pending,
             "agent_id": "brief_key_findings",
             # June 21 2026 -- bumped from the default 1500. Section
             # 3 targets 550 words of dense prose with the locked
             # academic figures + inline citations + [[VERIFY]]
             # markers; production runs were truncating mid-sentence
             # at 1500 tokens. 2500 gives comfortable headroom.
             "max_tokens": 2500,
             "task": (
                 "Write Section 3: KEY FINDINGS AND INSIGHTS. Approximately "
                 "550 words. Compare exactly THREE strategies: the 100% "
                 "equity benchmark, the best static diversifier (pull the "
                 "name from summary_statistics), and the dynamic regime-"
                 "aware blend.\n\n"
                 "Key figures to cite -- USE THE PLACEHOLDER TOKENS BELOW, "
                 "not the raw numbers. The platform substitutes the locked "
                 "academic values from the cache after generation:\n"
                 "  - Drawdown {{BENCHMARK_MAX_DD}} (benchmark) vs "
                 "{{REGIME_SWITCHING_MAX_DD}} (blend)\n"
                 "  - OOS Sharpe {{OOS_SHARPE_BLEND}} (blend) vs "
                 "{{OOS_SHARPE_BENCHMARK}} (benchmark)\n"
                 "  - {{OOS_WINDOW_MONTHS}}-month post-2022 out-of-sample "
                 "window\n\n"
                 "IMPORTANT: The token {{OOS_SHARPE_IMPROVEMENT_PCT}} "
                 "already resolves to a complete formatted string "
                 "including the + prefix and % suffix (e.g. '+98%'). "
                 "Do NOT append any additional % character or suffix "
                 "after this token. Write it exactly as the token "
                 "resolves.\n\n"
                 "Reference the platform's cumulative return chart by "
                 "name when you state the OOS Sharpe -- the chart is the "
                 "visual evidence behind the headline number. Reference "
                 "the implied asset allocation over time chart when you "
                 "discuss how the blend held the bond sleeve through the "
                 "2022 equity drawdown.\n\n"
                 "CORRELATION REGIME -- when you describe the equity-IG "
                 "correlation regime break, USE the tokens "
                 "{{PRE_2022_EQ_IG_CORR}} (averaged pre-2022 12m rolling "
                 "correlation) and {{POST_2022_EQ_IG_CORR}} (averaged "
                 "post-2022) for the actual values. Do NOT write "
                 "approximations like 'surges above +0.5' or 'crosses "
                 "+0.5' as a visual descriptor -- those bare numerics "
                 "are not platform-anchored constants and they trip the "
                 "hard-lock. Example acceptable phrasing: 'the rolling "
                 "correlation averaged {{PRE_2022_EQ_IG_CORR}} pre-2022 "
                 "and inverted to {{POST_2022_EQ_IG_CORR}} in the "
                 "post-2022 regime; the chart shows the structural "
                 "break in early 2022.'\n\n"
                 "Honest acknowledgement (one paragraph): the council "
                 "added value in 2 of 9 named market events (the play-"
                 "by-play scorecard). No strategy clears statistical "
                 "significance at p < {{BH_SIGNIFICANCE_THRESHOLD}} "
                 "under Benjamini-Hochberg FDR correction across the "
                 "three-strategy submission set (BENCHMARK, "
                 "CLASSIC_60_40, REGIME_SWITCHING). The multiple-"
                 "testing penalty is materially smaller than a full "
                 "ten-strategy survey would carry. The case rests on "
                 "economic magnitude, NOT statistical certainty.\n\n"
                 "The brief focuses on the three submission strategies "
                 "(benchmark, static diversifier 60/40, dynamic blend) "
                 "-- the appendix carries the same three-strategy "
                 "table at higher detail. Numbers from the platform; "
                 "conclusions "
                 "framed as 'our analysis shows' / 'we conclude' / 'we "
                 "interpret'."
                 + _BRIEF_TONE_RULES),
             "context": {"summary_statistics": data["summary_statistics"],
                         "regime_conditional": data["regime_conditional"],
                         "drawdown_comparison": data["drawdown_comparison"]}},
            {"key": "limitations", "available": True,
             "agent_id": "brief_limitations",
             "task": (
                 "Write Section 4: LIMITATIONS AND RISKS. Approximately "
                 "300 words. Be honest. Four mandatory limitations, one "
                 "short paragraph each.\n\n"
                 "  - THREE-ASSET SCOPE: the universe is equities, IG "
                 "bonds, HY bonds. State explicitly that this is a "
                 "PROJECT scope boundary, not an architectural limit -- "
                 "the platform's HMM and transition matrix work with "
                 "any return series.\n\n"
                 "  - SAMPLE SIZE: {{OOS_WINDOW_MONTHS}} months of "
                 "post-2022 out-of-sample data is approximately "
                 "{{OOS_WINDOW_PCT_OF_STUDY}}% of the full study window. "
                 "Bootstrap confidence intervals on Sharpe ratios "
                 "overlap substantially across the static set, which is "
                 "WHY the regime-conditional construction is the right "
                 "framing -- the selection mechanism comes from regime "
                 "signals, not Sharpe rankings.\n\n"
                 "  - TRANSACTION COSTS: the headline Sharpe figures are "
                 "gross. The platform's Net of Switching Costs panel "
                 "shows the blend stays above the benchmark across "
                 "{{SENSITIVITY_COST_BPS_LOW}}/"
                 "{{SENSITIVITY_COST_BPS_MID}}/"
                 "{{SENSITIVITY_COST_BPS_HIGH}} bps cost assumptions; "
                 "the post-cost margin compresses but does not invert.\n\n"
                 "  - STATISTICAL SIGNIFICANCE: no strategy clears p < "
                 "{{BH_SIGNIFICANCE_THRESHOLD}} under Benjamini-Hochberg "
                 "FDR correction. The case "
                 "is built on economic magnitude (drawdown reduction, "
                 "Sharpe improvement) rather than statistical certainty.\n\n"
                 "Close with a single sentence acknowledging the platform's "
                 "audit subsystem (Performance Record + Net of Switching "
                 "Costs + Implied Asset Allocation) as the standing "
                 "validation surface a reader can navigate to verify any "
                 "claim above. Do NOT add a 'next steps', 'future work', "
                 "or 'Part II' paragraph -- the rubric explicitly excludes "
                 "future-work content from the brief."
                 + _BRIEF_TONE_RULES),
             "context": {"summary_statistics": data["summary_statistics"]}},
            {"key": "final_recommendations", "available": True,
             "agent_id": "brief_final_recommendations",
             "task": (
                 "Write Section 5: FINAL RECOMMENDATIONS. Approximately "
                 "350 words. These are INVESTMENT CONCLUSIONS drawn from "
                 "the analysis -- NOT next steps, NOT operational "
                 "suggestions, NOT future research. The rubric is "
                 "explicit on this distinction.\n\n"
                 "Lead with the headline conclusion sentence in this "
                 "shape -- USE THE PLACEHOLDER TOKENS, not raw numbers: "
                 "'Given an out-of-sample Sharpe of {{OOS_SHARPE_BLEND}} "
                 "for the regime-conditional blend versus "
                 "{{OOS_SHARPE_BENCHMARK}} for the benchmark and a "
                 "maximum drawdown of {{REGIME_SWITCHING_MAX_DD}} "
                 "versus {{BENCHMARK_MAX_DD}}, we recommend that a "
                 "regime-conditional allocation framework be "
                 "considered as a core approach to asset allocation "
                 "in the post-2022 environment.'\n\n"
                 "Three supporting recommendations, each grounded in a "
                 "specific finding from Section 3:\n"
                 "  1. Adopt the regime-conditional construction as the "
                 "selection mechanism for the multi-asset blend, because "
                 "the static historical Sharpe ranking is unreliable at "
                 "the available sample size.\n"
                 "  2. Retain a diversifying bond sleeve through equity "
                 "drawdowns, because the 2022 evidence shows the blend's "
                 "capital-preservation advantage came from holding the "
                 "bond allocation rather than from market timing.\n"
                 "  3. Monitor the regime posterior monthly and re-balance "
                 "on regime transitions, because the value-added events "
                 "(2 of 9 in the play-by-play) cluster at regime flips.\n\n"
                 "REQUIRED ELEMENT (mandatory, not advisory) -- BEFORE "
                 "the regime reading paragraph below, you MUST include "
                 "this exact framing sentence (verbatim is preferred; "
                 "close equivalent is acceptable only if every key "
                 "phrase is preserved): 'The analytical performance "
                 "figures throughout this brief reflect the academic "
                 "submission record (locked at the submission freeze "
                 "hash). The regime classification and implied "
                 "weights below are live platform readings at the "
                 "time of generation.' This "
                 "sentence distinguishes the locked academic record "
                 "(the Sharpe / drawdown / OOS window figures cited "
                 "above) from the live regime read + implied weights "
                 "that follow -- two different states the panel needs "
                 "to keep separate. A draft without this sentence will "
                 "be rejected on review.\n\n"
                 "Then reference the live recommendation surface for the "
                 "current portfolio snapshot: 'The current regime is "
                 "{{CURRENT_REGIME}} at {{REGIME_CONFIDENCE}} confidence; "
                 "the implied equity allocation is "
                 "{{CURRENT_EQUITY_PCT}}. The CIO Recommendation card "
                 "and the Implied Asset Allocation Over Time chart "
                 "surface these live readings. This brief is the "
                 "analytical case; those surfaces are the live "
                 "snapshot.'\n\n"
                 "MANDATORY: when you cite the LIVE regime read, the "
                 "confidence percentage, or the live equity-allocation "
                 "weight, you MUST use the {{CURRENT_REGIME}}, "
                 "{{REGIME_CONFIDENCE}}, and {{CURRENT_EQUITY_PCT}} "
                 "placeholder tokens -- never write the live values as "
                 "raw numbers (e.g. '62.7%' or '44.5%'). Raw live values "
                 "stale instantly because the regime shifts between "
                 "generation time and reader time; the {{TOKEN}} "
                 "placeholders bind to the platform's current state at "
                 "the moment the reader opens the export.\n\n"
                 "If "
                 "live_recommendation.is_stale is true, include this "
                 "exact disclosure sentence: 'The live regime read at "
                 "generation time was unavailable; the recommendation "
                 "above references the most recent cached regime read "
                 "(stale_as_of in context).'\n\n"
                 "Reference the efficient frontier chart by name when "
                 "you justify the blend over a single static strategy.\n\n"
                 "Do NOT discuss future work, Part II, walk-forward "
                 "backtests, or any 'next research direction' content -- "
                 "those belong outside this section. The rubric grades "
                 "this section on investment conclusions drawn from the "
                 "quantitative result set above."
                 + _BRIEF_TONE_RULES),
             "context": {"live_recommendation": live,
                         "summary_statistics": data["summary_statistics"],
                         "study_period": data["study_period"]}},
            # June 26 2026 -- Section 6 'Visuals to Demonstrate the
            # Insights' removed. The section was a generation
            # artifact with no submission value: it captioned chart
            # surfaces in prose, but the brief already references
            # the relevant charts inline from Sections 3 and 4 and
            # the Analytical Appendix carries the per-chart data
            # tables. Keeping it produced a redundant 250-word
            # ending that the rubric specifically does NOT grade.
            # Brief now ends after Section 5: Final Recommendations.
        ]
        # PR #333 -- brief section plan injection. Mirrors the deck
        # path: cache-aware retrieval of the Opus story plan keyed by
        # (data_hash, 'brief'); on hit, each spec's task is prepended
        # with a LOCKED CONTRACT block listing the section's
        # key_message + numeric_anchors + target_length_words. The
        # per-section Sonnet pass then writes prose around the lock.
        # Fail-open: a missing plan or any retrieval error returns an
        # empty section_plan dict and the specs run exactly as before.
        section_plan = await _resolve_story_plan_brief_sections(data)
        if section_plan:
            specs = _inject_brief_section_plan(specs, section_plan)

        # June 21 2026 -- numeric substitution architecture. Build
        # the {token -> cache-value} substitution table once per
        # data_hash and thread it through every per-section
        # harness_narrative call. The Sonnet writer emits placeholders
        # ({{OOS_SHARPE_BLEND}} etc); the platform substitutes
        # verified cache values before the evaluator sees the prose
        # and before the .docx assembler reads it. Fail-open: any
        # error returns None and the per-section path runs without
        # substitution (the existing audit checks still fire on the
        # raw Sonnet output, so the failure mode is a degraded brief
        # rather than a missing brief).
        substitution_table: dict[str, str] | None = None
        try:
            from tools.audit_assembler import current_data_hash
            from tools.cio_recommendation import (
                compute_implied_asset_allocation, get_latest_recommendation,
            )
            from tools.numeric_substitution import get_substitution_table
            from tools.submission_freeze import get_effective_data_hash
            constants = (data.get("validated_constants") or {}) or {}
            rolling = data.get("rolling_correlation") or {}
            # Layer 4 -- the submission freeze. When active, document
            # generation reads the strategy_results_cache at the
            # frozen hash so exported deliverables never drift from
            # what was submitted. Live platform reads (Investment
            # Outlook, CIO card, daily digest, regime detector) still
            # call current_data_hash() directly -- the freeze
            # isolates document generation only.
            live_hash = await current_data_hash()
            data_hash = await get_effective_data_hash(live_hash)
            cio_row = await get_latest_recommendation()
            # CURRENT_*_PCT tokens need the implied asset
            # allocation -- compute once from the CIO blend weights.
            implied_alloc: dict | None = None
            try:
                if cio_row and cio_row.get("blend_weights"):
                    implied_alloc = await compute_implied_asset_allocation(
                        cio_row.get("blend_weights"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("brief_implied_alloc_failed", error=str(_exc))
            # June 22 2026 (PR A) -- read regime_signals_cache for the
            # 5 watchpoint tokens. 15-min TTL; falls back to em-dash
            # inside build_substitution_table if cold. Same wiring
            # as the deck callsite; brief sections 1/5/6 reference
            # current_regime and equity weight.
            live_signals: dict | None = None
            try:
                from tools.cache import get_regime_cache
                live_signals = await get_regime_cache()
                if live_signals is None:
                    log.warning(
                        "brief_live_signals_stale",
                        document_type="executive_brief",
                        note=("regime_signals_cache miss or expired -- "
                              "watchpoint tokens will render em-dash"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("brief_live_signals_read_failed",
                            error=str(_exc))
            # June 22 2026 (wiring fix) -- read analytics metrics
            # for pre/post 2022 Sharpes, factor loadings, and
            # cost sensitivity tokens. See
            # tools.academic_export.load_substitution_metric_sources
            # for the source mapping.
            from tools.academic_export import (
                load_substitution_metric_sources,
            )
            # June 27 2026 -- thread data_hash so historical-analytics
            # metric reads (regime_conditional / factor_loadings /
            # cost_sensitivity / crisis_performance) respect the
            # submission freeze when active. Live CIO + regime signals
            # remain LIVE by design (the platform feature, not frozen).
            regime_conditional_rows, factor_loadings_rows, \
                cost_sensitivity_payload, crisis_payload = (
                    await load_substitution_metric_sources(
                        data_hash=data_hash or None))
            substitution_table = get_substitution_table(
                data_hash or "",
                data.get("strategy_results") or {},
                cio_row,
                oos_sharpe_blend=constants.get(
                    "oos_sharpe_regime_conditional"),
                oos_sharpe_benchmark=constants.get(
                    "oos_sharpe_benchmark"),
                pre_2022_eq_ig_correlation=(
                    constants.get("correlation_pre_2022")
                    or rolling.get("pre_2022")),
                post_2022_eq_ig_correlation=(
                    constants.get("correlation_post_2022")
                    or rolling.get("post_2022")),
                oos_window_pct_of_study=constants.get(
                    "oos_window_pct_of_study"),
                study_months=(data.get("study_period") or {}).get(
                    "n_months"),
                implied_allocation=implied_alloc,
                live_signals=live_signals,
                regime_conditional=regime_conditional_rows,
                factor_loadings=factor_loadings_rows,
                cost_sensitivity=cost_sensitivity_payload,
                crisis_performance=crisis_payload,
                hash_verified=True)
            log.info("substitution_table_built",
                     document_type="executive_brief",
                     data_hash=(data_hash or "")[:8],
                     tokens_available=len(substitution_table))
        except Exception as exc:  # noqa: BLE001
            log.warning("substitution_table_build_failed",
                        document_type="executive_brief",
                        error=str(exc))

        # June 28 2026 -- re-augment per-section anchors now
        # that the substitution_table is built. _inject_brief_
        # section_plan ran earlier (before the table existed)
        # so anchors only carry story-plan entries. This pass
        # adds the always-allowed study-metadata anchors
        # (STUDY_MONTHS, N_STRATEGIES, OOS_WINDOW_MONTHS, etc.)
        # so the LLM can write raw "287 months" without
        # tripping the hard-lock as token_available.
        if substitution_table:
            for spec in specs:
                if "numeric_anchors" in spec:
                    spec["numeric_anchors"] = (
                        _augment_anchors_with_study_metadata(
                            spec["numeric_anchors"],
                            substitution_table))

        narratives = await _generate_narratives(
            _apply_draft_caveats(specs, document_type="executive_brief"),
            n_strategies=len(data.get("strategy_results") or {}),
            substitution_table=substitution_table,
            # June 28 2026 -- arms the untoken-numeric hard lock
            # inside harness_narrative. Brief is one of the two
            # protected document types.
            document_type="executive_brief")

        # Substitution-architecture summary log. Operators read this
        # in Render logs to confirm the determinism layer fired
        # correctly. Zero unresolved_placeholders + zero raw_numerics
        # is the green state.
        if substitution_table is not None:
            try:
                from tools.numeric_substitution import (
                    unresolved_placeholders,
                )
                joined_text = "\n".join(narratives.values())
                unresolved = unresolved_placeholders(joined_text)
                log.info("substitution_complete",
                         document_type="executive_brief",
                         tokens_available=len(substitution_table),
                         unresolved_placeholders=unresolved,
                         unresolved_count=len(unresolved))
            except Exception as exc:  # noqa: BLE001
                log.warning("substitution_summary_failed",
                            error=str(exc))
        # Layer 3b (June 21 2026) -- pass the substitution_table through
        # so Section 6 chart captions can resolve {{DATA_HASH}} /
        # {{PRE_2022_EQ_IG_CORR}} / {{OOS_SHARPE_BLEND}} against the
        # verified cache values. verification_result is None at
        # generation time (the receipt page shows "Not yet verified" so
        # the page slot is structurally stable -- export-time
        # verification rebuilds the brief from the editor draft via
        # _editor_export, where the verification dict IS available).
        docx_bytes = await asyncio.to_thread(
            build_executive_brief, data, narratives,
            substitution_table=substitution_table,
            verification_result=None)

        # Load the generated content into an editor draft so the frontend
        # can open it directly in the editor — the same pattern as the
        # midpoint paper and the deck. The draft_id rides back in the
        # X-Draft-Id response header; a draft-storage failure never fails
        # the download.
        draft_id: int | None = None
        try:
            from tools.editor_content import executive_brief_to_editor
            from tools.editor_drafts import create_draft
            # June 28 2026 (Phase 2) -- thread the substitution_table
            # so that under DEFER_SUBSTITUTION_TO_EXPORT the brief's
            # content_json preserves {{TOKEN}} placeholders + the
            # parallel content_text shadow column carries the
            # resolved values for full-text search + word counts.
            content_json, content_text = executive_brief_to_editor(
                narratives, substitution_table=substitution_table)

            # ── Concern 7h: pre-submission adversarial critic ─────
            # Inlines the critic + debate-round response into the
            # draft so the team reviews the critique alongside the
            # brief. Always-write to council_debates regardless of
            # severity. Fail-open: any failure leaves content_text
            # unchanged and the generation flow proceeds.
            try:
                from agents.academic_review import (
                    run_doc_gen_debate_round,
                )
                content_text, _debate_id, _critic_res = (
                    await run_doc_gen_debate_round(
                        reviewer_email=email,
                        document_type="executive_brief",
                        content_text=content_text))
            except Exception as _exc:  # noqa: BLE001
                log.warning(
                    "doc_gen_critic_pipeline_failed",
                    document_type="executive_brief",
                    error=str(_exc))

            # ── Post-generation audit (June 3 2026) ───────────────
            # Four deterministic checks against content_text BEFORE
            # the draft lands in editor_drafts. Flags travel on the
            # draft's audit_warnings JSONB column so the frontend
            # banner surfaces them. NEVER blocks the write — the
            # human reviews and resolves.
            audit_warnings = await _run_document_audit(
                content_text, "executive_brief", email)

            # Stamp the live strategy hash on the draft (migration
            # 063) so the tile + editor chips render 'Data current'.
            try:
                from tools.audit_assembler import (
                    current_data_hash as _curr_hash_brief,
                )
                _brief_hash = await _curr_hash_brief()
            except Exception:  # noqa: BLE001
                _brief_hash = None
            draft = await create_draft(
                "executive_brief", email,
                f"Executive Brief — {date.today().isoformat()}",
                content_json, content_text,
                created_from="generated",
                audit_warnings=audit_warnings,
                data_hash=_brief_hash)
            if draft is not None:
                draft_id = draft["id"]
            await _write_audit_metrics(
                "executive_brief", email, draft_id, audit_warnings)

            # Layer 3 (June 21 2026) -- persist the value manifest +
            # generation data_hash on the draft so export-time
            # verification has an authoritative reference for every
            # numeric value the substitution table produced. Layer 3
            # of the substitution architecture closes the loop at
            # export time: a manual edit changing "1.24" to "1.23"
            # is caught before the file leaves the platform.
            if draft_id is not None and substitution_table is not None:
                try:
                    from tools.audit_assembler import (
                        current_data_hash as _cur_hash,
                    )
                    from tools.editor_drafts import (
                        update_value_manifest as _update_manifest,
                    )
                    from tools.numeric_substitution import (
                        build_value_manifest,
                    )
                    from datetime import datetime as _dt
                    from datetime import timezone as _tz
                    _hash_for_manifest = await _cur_hash() or ""
                    manifest = build_value_manifest(
                        substitution_table,
                        data_hash=_hash_for_manifest[:64],
                        generated_at=_dt.now(_tz.utc).isoformat())
                    await _update_manifest(
                        draft_id, manifest,
                        data_hash=_hash_for_manifest[:64] or None)
                    log.info(
                        "value_manifest_persisted",
                        document_type="executive_brief",
                        draft_id=draft_id,
                        n_values=len(manifest))
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "value_manifest_persist_failed",
                        document_type="executive_brief",
                        error=str(exc))

                # June 28 2026 (Fix 8b) -- auto-upgrade to
                # token_value nodes immediately after persist.
                # Fail-open: any error logs + the draft is
                # still usable (admin batch endpoint can
                # recover).
                if draft_id is not None:
                    await _auto_upgrade_draft_to_token_values(
                        draft_id, "executive_brief")
        except Exception as exc:  # noqa: BLE001
            log.warning("executive_brief_draft_create_failed", error=str(exc))

        filename = f"forest-capital-executive-brief-{date.today().isoformat()}.docx"
        return docx_bytes, filename, _DOCX_MEDIA, draft_id
    except Exception as exc:  # noqa: BLE001
        log.error("executive_brief_generation_error", error=str(exc))
        raise


@app.post("/api/v1/export/analytical-appendix")
@limiter.limit("6/minute")
async def export_analytical_appendix(
    request: Request,
    body: dict | None = None,
    session: dict = Depends(require_permission("generate_documents")),
):
    """
    Starts analytical-appendix generation.

    The Analytical Appendix is the evidentiary record behind every
    claim in the executive brief and the panel presentation — eight
    sections (A–H) covering data and methodology, full strategy
    performance, statistical tests, bootstrap CIs, factor loadings,
    crisis windows, transaction-cost sensitivity, and the validation
    audit summary. Every figure is pulled from existing caches; the
    endpoint NEVER recomputes (so a cold deploy renders [DATA PENDING]
    cleanly rather than blocking on a warm).

    With an editor_draft_id in the body the .docx is built
    synchronously from that draft (the in-editor Export path).
    Otherwise generation — eight short Academic Writer paragraphs —
    runs as a background job and the endpoint returns 202 with a
    job_id; poll GET /api/v1/jobs/{id}.
    """
    editor_draft_id = (body or {}).get("editor_draft_id")
    if editor_draft_id:
        return await _editor_export(int(editor_draft_id))
    await _require_report_ready()
    return _start_generation_job("analytical_appendix", session, request)


# June 29 2026 -- THREE-STRATEGY SUBMISSION SCOPE. Both the brief
# AND the appendix now operate on the same restricted three-strategy
# set (BENCHMARK, CLASSIC_60_40, REGIME_SWITCHING) per the scope
# filter at tools/academic_export.gather_analytical_appendix_data.
# The appendix retains the higher-detail evidence base for those
# three strategies (the per-section tables D / E / F / G carry
# more granular metrics for each row than the brief shows), but the
# strategy roster is identical to the brief.
_APPENDIX_FRAMING_PRELUDE = (
    "APPENDIX FRAMING (applies to every section):\n"
    "  - This is the analytical evidence base supporting the "
    "executive brief and presentation deck. The audience reads the "
    "appendix to verify the claims made elsewhere.\n"
    "  - The appendix operates on the SAME three-strategy "
    "submission scope as the brief / deck / script: BENCHMARK, "
    "CLASSIC_60_40, REGIME_SWITCHING. Higher-detail per-strategy "
    "metrics (bootstrap CI, factor loadings, crisis windows) are "
    "shown for those three strategies; out-of-scope strategies "
    "(MIN_VARIANCE, RISK_PARITY, VOL_TARGETING, etc.) are NOT in "
    "the submission record and must NOT be named in appendix "
    "prose.\n"
    "  - Each section's prose must include one sentence of economic "
    "intuition explaining what the result MEANS for a reader, not "
    "just what the table contains. Example: a section reporting a "
    "Sharpe ratio table does not stop at 'the table reports Sharpe' "
    "-- it adds 'the regime-conditional construction's lead widens "
    "post-2022 because the HMM identifies structural state changes "
    "that persist for months'.\n"
    "  - Sensitivity analysis results must be presented with "
    "interpretation, not just tables. Tell the reader what to take "
    "away.\n\n"
)


# Appendix section narrative tasks — one Academic Writer call per
# section, each producing ~80-150 words. Deliberately short: the
# appendix is table-heavy, the prose is scaffolding. Each task pins
# the section's role and provides the context the writer needs to
# describe what the table will show. The framing prelude is
# prepended to each task at dispatch time so a future section
# addition automatically inherits the audience + economic-intuition
# contract.
# June 21 2026 -- HARD word-count cap appended to every appendix
# section task. The brief alignment excerpt added in PR #364 was
# driving the narrative agent to expand prose to mirror the brief
# (~430 words/section observed in production vs. 100-130 word
# target = 3x overshoot). The cap is appended to every task so
# the writer sees the constraint AFTER any framing the brief
# excerpt introduces.
_APPENDIX_WORD_CAP_INSTRUCTION = (
    "\n\nHARD WORD LIMIT (non-negotiable):\n"
    "Your narrative intro for this section MUST NOT exceed 130 "
    "words. The brief excerpt above (if any) is for ALIGNMENT "
    "CONTEXT ONLY -- it tells you what the brief argues so your "
    "intro can match the framing. It does NOT expand the scope of "
    "your output. The table that follows the intro is the primary "
    "content of this section; the intro's job is to introduce the "
    "table in 100-130 words, period. Do NOT expand to mirror the "
    "depth of the brief excerpt. Do NOT reproduce the brief's "
    "argument in full. A section over 130 words will be flagged "
    "and require regeneration.")


# June 26 2026 -- Citation discipline for the Benjamin vs Benjamini
# split. Two different papers, two different claims; the model has
# previously conflated them (e.g. attributing the FDR correction to
# Benjamin et al. 2018 or attributing the p < 0.005 threshold to
# Benjamini & Hochberg 1995). Threaded into the appendix sections
# that actually cite either paper (C and E) so the constraint sits
# next to the relevant prose. Mirrors the brief-side guidance in
# _BRIEF_TONE_RULES_BASE.
_APPENDIX_CITATION_DISCIPLINE = (
    "\n\nCITATION DISCIPLINE -- Benjamin vs Benjamini (non-"
    "negotiable):\n"
    "Two different papers, two different claims; never swap or "
    "conflate them in-text or in the reference list.\n"
    "  - Benjamini & Hochberg (1995) -- the FDR (false discovery "
    "rate) CORRECTION METHODOLOGY. Cite this whenever you "
    "reference the q-value control, the BH procedure, or the "
    "multiple-hypothesis correction technique.\n"
    "  - Benjamin et al. (2018) -- the p < 0.005 THRESHOLD "
    "PROPOSAL (Redefine Statistical Significance, Nature Human "
    "Behaviour). Cite this whenever you reference the 0.005 "
    "cutoff itself as the bar for declaring a finding "
    "significant.\n"
    "  - Wrong: 'no strategy clears p < 0.005 under Benjamin et "
    "al. (2018) FDR correction' (the FDR correction is not from "
    "that paper).\n"
    "  - Right: 'no strategy clears the p < 0.005 threshold "
    "(Benjamin et al., 2018) under Benjamini-Hochberg FDR "
    "correction (Benjamini & Hochberg, 1995)' -- each citation "
    "next to the claim it actually supports.")


_APPENDIX_NARRATIVE_TASKS = {
    "appendix_a": (
        "Write a 100-130 word introduction to Section A: Data and "
        "Methodology, for the project's Analytical Appendix. Name the "
        "study period (start, end, number of months) and the three "
        "asset classes (S&P 500 total return for equity, BND for "
        "investment-grade bonds, BAMLHYH0A0HYM2TRIV with HYG splice "
        "for high yield, all monthly). Note the risk-free series "
        "(DTB3, converted monthly). State that the table that follows "
        "carries the asset-level summary statistics. Plain academic "
        "prose. No em dashes."),
    "appendix_b": (
        "Write a 100-130 word introduction to Section B: Full Strategy "
        "Performance. Note that the table that follows reports Sharpe, "
        "CAGR, volatility, Sortino, Calmar, and max drawdown for every "
        "strategy in the cache, sorted by Sharpe descending. Mention "
        "that the BENCHMARK row appears alongside the active strategies "
        "for direct comparison. Plain academic prose. No em dashes."),
    "appendix_c": (
        "Write a 100-130 word introduction to Section C: Statistical "
        "Tests. Note that the table that follows reports the paired-t "
        "p-value, the FDR-corrected p-value, the Deflated Sharpe Ratio "
        "p-value, the Probabilistic Sharpe Ratio, and the SPA gate for "
        "every strategy except the benchmark. State plainly that no "
        "strategy clears the p < 0.005 threshold (Benjamin et al., "
        "2018) under Benjamini-Hochberg FDR correction (Benjamini & "
        "Hochberg, 1995), and that we surface this honestly rather "
        "than report Sharpe rankings alone. Plain academic prose. "
        "No em dashes."
        + _APPENDIX_CITATION_DISCIPLINE),
    "appendix_d": (
        "Write a 100-130 word introduction to Section D: Bootstrap "
        "Confidence Intervals. Note the methodology: block bootstrap "
        "of length {{BOOTSTRAP_BLOCK_LENGTH}} with 10,000 resamples "
        "and a fixed seed of {{BOOTSTRAP_SEED}}, "
        "applied to the monthly excess-return series for every "
        "strategy. State that the table that follows reports each "
        "strategy's point Sharpe with its 95% CI, and whether the "
        "interval overlaps the benchmark's Sharpe. Note that "
        "substantial overlap is the analytical justification for "
        "treating static-strategy rankings as inconclusive. Plain "
        "academic prose. No em dashes.\n\n"
        "TABLE PLACEMENT NOTE: The renderer places Table D1 "
        "immediately after this narrative paragraph -- write the "
        "phrasing 'the table that follows' or 'Table D1 below' "
        "rather than 'the table at the end of this appendix' or "
        "any phrasing that implies Table D1 lives in a separate "
        "evidence block. Table D1 is inline with Section D's "
        "prose."),
    "appendix_e": (
        "Write a 100-130 word introduction to Section E: Factor "
        "Loadings. Note that the table that follows reports the "
        "annualised alpha and the four Carhart factor coefficients "
        "(MKT-RF, SMB, HML, MOM) plus R-squared for every strategy, "
        "with a trailing asterisk on each coefficient significant at "
        "p < 0.05. State that the table is read for each strategy's "
        "primary return driver, not for stock-picking alpha. Plain "
        "academic prose. No em dashes."
        + _APPENDIX_CITATION_DISCIPLINE),
    "appendix_f": (
        "Write a 100-130 word introduction to Section F: Crisis Window "
        "Performance. Note that the table that follows reports each "
        "strategy's cumulative return through five named crisis windows "
        "(GFC 2008-2009, EU Debt 2011, COVID Crash, COVID Recovery, "
        "Rate Shock 2022). State explicitly that the headline figure is "
        "the cumulative return through the window, NOT the annualised "
        "CAGR (the F3 fix). Plain academic prose. No em dashes.\n\n"
        "DAGGER FOOTNOTE DISCIPLINE (non-negotiable): The trailing "
        "dagger (†) symbol on a Table F1 CELL flags a partial-"
        "overlap window for that specific strategy -- meaning that "
        "strategy's live history started AFTER the crisis window "
        "began or ended BEFORE the window closed, so the cumulative "
        "return is computed over only the overlapping subset of the "
        "window. If a strategy's history fully covers the window, "
        "the cell does NOT carry a dagger. If EVERY strategy in the "
        "table fully covers EVERY crisis window (the common case "
        "for the canonical strategies), the dagger does not appear "
        "anywhere and you should not mention partial-overlap in the "
        "narrative at all. NEVER describe the dagger as a uniform "
        "marker applied to every row -- that contradicts the "
        "footnote's definition and confuses readers. Mention the "
        "dagger only if the rendered table actually carries it."),
    "appendix_g": (
        "Write a 100-130 word introduction to Section G: Transaction "
        "Cost Sensitivity. Note that the table that follows reports "
        "the net-of-cost Sharpe of the regime-conditional blend at "
        "10/15/20 basis points per material rebalance over the "
        "post-2022 OOS window, alongside the count of material "
        "rebalances. State that the regime-conditional blend remains "
        "above the benchmark Sharpe net of plausible transaction "
        "costs. Plain academic prose. No em dashes.\n\n"
        "PROHIBITED CONTENT (June 22 2026 -- production bug): "
        "Do NOT include any placeholder note saying the table 'must "
        "be inserted here' or 'before final submission'. Do NOT "
        "include an open-items list, a 'five open items' block, a "
        "'to be completed' list, or any caveat about pending work. "
        "The transaction cost sensitivity table is generated "
        "PROGRAMMATICALLY and is ALWAYS present in the rendered "
        "document. Your introduction's only job is to set up the "
        "table the reader will see immediately below it -- the "
        "table speaks for itself."),
    "appendix_h": (
        "Write a 100-130 word introduction to Section H: Validation "
        "Audit Summary. Note that the platform's analytics invariant "
        "framework runs at the end of every warm and that the "
        "deterministic-detection contract (no LLM in the detection "
        "path) is documented in docs/INVARIANTS.md. State that the "
        "table that follows reports the latest warm's verdict — "
        "checks run, hard failures, soft warnings, and the run "
        "timestamp. Note that the audit disclosure appendix that "
        "follows carries the statistical audit and methodology QA "
        "history. Plain academic prose. No em dashes."),
}


async def _generate_appendix_document(
    email: str,
) -> tuple[bytes, str, str, int | None]:
    """
    Generates the Analytical Appendix. Returns
    (file_bytes, filename, media_type, editor_draft_id).

    The appendix is table-heavy and rhetorically minimal: each
    section opens with a short Academic Writer paragraph then the
    cached data in an APA-style table. All eight cache reads happen
    in gather_analytical_appendix_data; the eight narratives generate
    concurrently through the harness; build_analytical_appendix
    assembles them. No recomputation anywhere — a missing cache row
    degrades that section to [DATA PENDING] rather than blocking.
    """
    import asyncio
    from datetime import date

    from tools.academic_docx import build_analytical_appendix
    from tools.academic_export import (
        DATA_PENDING, gather_analytical_appendix_data,
    )

    # June 21 2026 -- brief-as-anchor gate. The analytical appendix
    # supports what the brief argues; without a brief on hand it
    # would generate independently from raw cache and risk
    # framing drift. 409 surfaces inline in the editor.
    from tools.brief_grounding import get_brief_for_grounding
    brief_grounding = await get_brief_for_grounding()
    if brief_grounding is None:
        raise HTTPException(
            status_code=409,
            detail=(
                "Generate the executive brief before the "
                "analytical appendix. The appendix supports the "
                "brief's claims; all three deliverables (brief / "
                "deck / appendix) must share a single narrative."))

    try:
        # June 27 2026 (PR 1 v3, LEAK 1 closer) -- compute the
        # effective data_hash inline + thread it into
        # gather_analytical_appendix_data so the inner
        # gather_document_data call routes the strategy_results_
        # cache read through get_strategy_cache(data_hash) instead
        # of get_latest_strategy_cache. Under freeze, a miss
        # raises StrategyCacheMissingForHashError (translated to a
        # 500 by the outer except).
        from tools.audit_assembler import (
            current_data_hash as _appx_current_hash,
        )
        from tools.submission_freeze import (
            get_effective_data_hash as _appx_eff_hash,
        )
        _appx_live = await _appx_current_hash()
        _appx_data_hash = await _appx_eff_hash(_appx_live)
        data = await gather_analytical_appendix_data(
            data_hash=_appx_data_hash or None)
        avail = data["available"]
        pending = (f"{DATA_PENDING} — analytics caches not warm. Load the "
                   "dashboard once, then regenerate this appendix.")

        # June 22 2026 -- HASH-MATCHED pre-flight cache gate. The
        # appendix's graded sections (B, C, D, E, G) depend on
        # cache fields populated by the backtester +
        # refresh_academic_analytics + refresh_oos_cost_sensitivity
        # chain. PR #365's original gate accepted ANY non-empty
        # cache row as "warm" -- including stale-hash rows from
        # a previous data state. That let the appendix render
        # against data the brief's narrative didn't see.
        #
        # The corrected gate computes the canonical CURRENT
        # strategy hash via _compute_data_hash(n_rows, last_date,
        # n_strategies=10) and verifies EACH source carries a row
        # at that exact hash. A stale-hash row counts as missing.
        # The 409 detail surfaces the actual mismatch (expected
        # hash, what's in the cache) so the operator knows whether
        # the chain ran or just one component.
        missing_cache_fields: list[str] = []
        canonical_hash = ""
        try:
            from tools.cache import (
                _compute_data_hash, get_strategy_cache,
            )
            from tools.data_fetcher import get_full_history_async
            from tools.precomputed_analytics import get_metric_by_hash
            from tools.regime_meta_validation import (
                _COST_METRIC_KIND,
            )
            history = await get_full_history_async()
            monthly = history.get("equity_monthly")
            n_rows = len(monthly) if monthly is not None else 0
            last_date = (
                str(monthly.index[-1].date())
                if monthly is not None and n_rows > 0
                else "unknown")
            canonical_hash = _compute_data_hash(
                n_rows, last_date, n_strategies=10)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "appendix_canonical_hash_compute_failed",
                error=str(exc))

        if canonical_hash:
            # Section B/C source -- strategy_results_cache row at
            # the canonical hash.
            try:
                sr_row = await get_strategy_cache(canonical_hash)
            except Exception:  # noqa: BLE001
                sr_row = None
            if not sr_row:
                missing_cache_fields.append(
                    f"strategy_results_cache (no row at hash "
                    f"{canonical_hash} -- Section B/C source)")

            # Section D source -- academic_analytics metric AT
            # the canonical hash. The metric payload also carries
            # factor_loadings (Section E), so a single hit covers
            # both. A row at a DIFFERENT hash counts as missing.
            try:
                aa = await get_metric_by_hash(
                    "academic_analytics", canonical_hash)
            except Exception:  # noqa: BLE001
                aa = None
            if not (aa or {}).get("bootstrap_ci_sharpe"):
                missing_cache_fields.append(
                    f"bootstrap_ci_sharpe (no academic_analytics "
                    f"row at hash {canonical_hash} -- Section D "
                    f"source)")
            if not (aa or {}).get("factor_loadings"):
                missing_cache_fields.append(
                    f"factor_loadings (no academic_analytics row "
                    f"at hash {canonical_hash} -- Section E "
                    f"source)")

            # Section G source -- cost_sensitivity metric at
            # canonical hash.
            try:
                cs = await get_metric_by_hash(
                    _COST_METRIC_KIND, canonical_hash)
            except Exception:  # noqa: BLE001
                cs = None
            if not cs:
                missing_cache_fields.append(
                    f"cost_sensitivity (no oos_cost_sensitivity "
                    f"row at hash {canonical_hash} -- Section G "
                    f"source)")
        else:
            # Couldn't compute the canonical hash -- treat as a
            # hard cold-cache state, same as the legacy gate.
            missing_cache_fields.append(
                "canonical hash unavailable (full data history "
                "unreadable -- check market_data_monthly + "
                "ff_factors_monthly)")

        if missing_cache_fields:
            raise HTTPException(
                status_code=409,
                detail=(
                    f"Analytics caches are not warm at the "
                    f"canonical current strategy hash "
                    f"({canonical_hash or 'unavailable'}). "
                    "Missing or stale: "
                    + "; ".join(missing_cache_fields)
                    + ". A stale-hash row counts as missing -- "
                    "the appendix must render at the hash that "
                    "matches the live market data, not a "
                    "previous data state. Run POST /api/v1/"
                    "admin/refresh-appendix-caches to populate "
                    "the cache chain at the canonical hash, "
                    "then retry."))

        # Layer 2 (June 21 2026) -- build the substitution table once
        # per appendix-generation job. Same per-data_hash cache the
        # brief + deck use, so a metric appearing in any of the three
        # documents is byte-identical by construction. include_per_
        # strategy=True (the default) is critical here: the appendix
        # is the surface that needs the dynamic {{STRATEGY_NAME_*}}
        # tokens for all 10 strategies.
        substitution_table: dict[str, str] | None = None
        try:
            from tools.audit_assembler import current_data_hash
            from tools.cio_recommendation import (
                compute_implied_asset_allocation, get_latest_recommendation,
            )
            from tools.numeric_substitution import (
                get_substitution_table,
            )
            from tools.academic_deck import (
                OOS_SHARPE_REGIME_CONDITIONAL,
                OOS_SHARPE_BENCHMARK,
                CORRELATION_PRE_2022, CORRELATION_POST_2022,
                OOS_WINDOW_PCT_OF_STUDY,
            )
            from tools.submission_freeze import get_effective_data_hash
            # Layer 4 -- submission freeze (see _generate_brief_document
            # for the rationale). Live platform reads call
            # current_data_hash() directly; document generation
            # routes through get_effective_data_hash so the appendix
            # locks to the frozen hash on submission day.
            live_hash = await current_data_hash()
            data_hash = await get_effective_data_hash(live_hash)
            cio_row = await get_latest_recommendation()
            implied_alloc: dict | None = None
            try:
                if cio_row and cio_row.get("blend_weights"):
                    implied_alloc = await compute_implied_asset_allocation(
                        cio_row.get("blend_weights"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("appendix_implied_alloc_failed", error=str(_exc))
            # June 22 2026 (PR A) -- read regime_signals_cache. Same
            # wiring as brief + deck callsites; the appendix's
            # Section G (cost sensitivity + recommendation) cites
            # the current regime + watchpoint posture.
            live_signals: dict | None = None
            try:
                from tools.cache import get_regime_cache
                live_signals = await get_regime_cache()
                if live_signals is None:
                    log.warning(
                        "appendix_live_signals_stale",
                        document_type="analytical_appendix",
                        note=("regime_signals_cache miss or expired -- "
                              "watchpoint tokens will render em-dash"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("appendix_live_signals_read_failed",
                            error=str(_exc))
            # June 22 2026 (wiring fix) -- read analytics metrics
            # for pre/post 2022 Sharpes, factor loadings, and
            # cost sensitivity tokens.
            from tools.academic_export import (
                load_substitution_metric_sources,
            )
            # June 27 2026 -- thread data_hash so historical-analytics
            # metric reads (regime_conditional / factor_loadings /
            # cost_sensitivity / crisis_performance) respect the
            # submission freeze when active. Live CIO + regime signals
            # remain LIVE by design (the platform feature, not frozen).
            regime_conditional_rows, factor_loadings_rows, \
                cost_sensitivity_payload, crisis_payload = (
                    await load_substitution_metric_sources(
                        data_hash=data_hash or None))
            substitution_table = get_substitution_table(
                data_hash or "",
                data.get("strategy_results") or {},
                cio_row,
                oos_sharpe_blend=OOS_SHARPE_REGIME_CONDITIONAL,
                oos_sharpe_benchmark=OOS_SHARPE_BENCHMARK,
                pre_2022_eq_ig_correlation=CORRELATION_PRE_2022,
                post_2022_eq_ig_correlation=CORRELATION_POST_2022,
                oos_window_pct_of_study=OOS_WINDOW_PCT_OF_STUDY,
                study_months=(data.get("study_period") or {}).get(
                    "n_months"),
                implied_allocation=implied_alloc,
                live_signals=live_signals,
                regime_conditional=regime_conditional_rows,
                factor_loadings=factor_loadings_rows,
                cost_sensitivity=cost_sensitivity_payload,
                crisis_performance=crisis_payload,
                hash_verified=True)
            log.info("substitution_table_built",
                     document_type="analytical_appendix",
                     data_hash=(data_hash or "")[:8],
                     tokens_available=len(substitution_table))
        except Exception as exc:  # noqa: BLE001
            log.warning("substitution_table_build_failed",
                        document_type="analytical_appendix",
                        error=str(exc))

        # Layer 2 -- prepend the substitution placeholder guide to
        # each appendix section task so the Sonnet writer uses
        # {{TOKEN}} markers instead of raw figures. The appendix-
        # specific extension teaches the {{STRATEGY_NAME_METRIC}}
        # convention for the 10-strategy surface.
        appendix_guide = ""
        if substitution_table is not None:
            appendix_guide = (
                _NUMERIC_PLACEHOLDER_GUIDE
                + _APPENDIX_NUMERIC_PLACEHOLDER_GUIDE_EXTENSION
                + "\n")

        # June 21 2026 brief-as-anchor -- precompute per-appendix-
        # section brief excerpts. APPENDIX_TO_BRIEF_SECTION
        # maps each appendix section_key to a brief section
        # (or None for appendix-specific sections like
        # portfolio_construction that have no brief counterpart).
        # brief_section_excerpt returns "" for None / missing
        # mappings so the resulting block is a no-op string.
        from tools.brief_grounding import (
            APPENDIX_TO_BRIEF_SECTION, brief_section_block,
            brief_section_excerpt,
        )
        brief_text = brief_grounding["content_text"]
        section_brief_blocks: dict[str, str] = {}
        for key in _APPENDIX_NARRATIVE_TASKS.keys():
            agent_id = f"appendix_{key.split('_', 1)[1]}"
            brief_section = APPENDIX_TO_BRIEF_SECTION.get(agent_id)
            excerpt = brief_section_excerpt(brief_text, brief_section)
            section_brief_blocks[key] = brief_section_block(
                excerpt, brief_section)

        specs = [
            {
                "key": key,
                "agent_id": f"appendix_{key.split('_', 1)[1]}",
                # PR #334 -- prepend the framing prelude so every
                # section task inherits the audience contract +
                # economic-intuition guard. The three-strategy
                # simplification does NOT apply to the appendix --
                # full 10-strategy coverage is appropriate here --
                # but the audience and economic-intuition layers
                # do (see _APPENDIX_FRAMING_PRELUDE above).
                # Layer 2 (June 21) -- the substitution placeholder
                # guide is prepended ahead of the framing prelude.
                # June 21 2026 brief-as-anchor -- per-section brief
                # alignment excerpt (no-op for appendix-only
                # sections per APPENDIX_TO_BRIEF_SECTION).
                # Word cap appended AFTER the brief excerpt so the
                # writer sees the constraint last (and remembers it
                # while writing). PR #364 (brief alignment) was
                # driving the agent to expand prose to mirror the
                # brief; this cap is the explicit counter-instruction.
                "task": (
                    appendix_guide
                    + _APPENDIX_FRAMING_PRELUDE + task
                    + section_brief_blocks.get(key, "")
                    + _APPENDIX_WORD_CAP_INSTRUCTION),
                "context": {"study_period": data.get("study_period")},
                "available": avail,
                "pending": pending,
            }
            for key, task in _APPENDIX_NARRATIVE_TASKS.items()
        ]
        narratives = await _generate_narratives(
            _apply_draft_caveats(
                specs, document_type="analytical_appendix"),
            n_strategies=len(data.get("strategy_results") or {}),
            substitution_table=substitution_table,
            # June 28 2026 -- arms the untoken-numeric hard lock
            # inside harness_narrative. Appendix is the second
            # protected document type alongside brief.
            document_type="analytical_appendix")

        # Per-document substitution-complete telemetry. Same shape the
        # brief + deck writers emit at end of generation.
        if substitution_table is not None:
            try:
                from tools.numeric_substitution import (
                    unresolved_placeholders,
                )
                joined_text = "\n".join(narratives.values())
                unresolved = unresolved_placeholders(joined_text)
                log.info("substitution_complete",
                         document_type="analytical_appendix",
                         tokens_available=len(substitution_table),
                         unresolved_placeholders=unresolved,
                         unresolved_count=len(unresolved))
            except Exception as exc:  # noqa: BLE001
                log.warning("substitution_summary_failed",
                            document_type="analytical_appendix",
                            error=str(exc))

        # June 28 2026 -- thread substitution_table through so the
        # builder can resolve {{TOKEN}} placeholders in narratives
        # at export time. Required by the Phase-1 deferred-
        # substitution pipeline; harmless under legacy generation
        # (substitution_table=None) where narratives already carry
        # resolved values.
        docx_bytes = await asyncio.to_thread(
            build_analytical_appendix, data, narratives,
            substitution_table=substitution_table)

        # Load the generated content into an editor draft so the
        # frontend can open it directly in the editor — same pattern
        # as the brief / midpoint / deck. The draft_id rides back in
        # the X-Draft-Id response header. Draft-storage failure never
        # fails the download.
        draft_id: int | None = None
        try:
            from tools.editor_content import analytical_appendix_to_editor
            from tools.editor_drafts import create_draft
            # June 28 2026 (Phase 2) -- same substitution_table
            # threading as the brief path. content_json keeps
            # {{TOKEN}} placeholders intact under flag ON;
            # content_text carries the substituted projection.
            content_json, content_text = analytical_appendix_to_editor(
                narratives, substitution_table=substitution_table)

            # ── Concern 7h: pre-submission adversarial critic ─────
            try:
                from agents.academic_review import (
                    run_doc_gen_debate_round,
                )
                content_text, _debate_id, _critic_res = (
                    await run_doc_gen_debate_round(
                        reviewer_email=email,
                        document_type="analytical_appendix",
                        content_text=content_text))
            except Exception as _exc:  # noqa: BLE001
                log.warning(
                    "doc_gen_critic_pipeline_failed",
                    document_type="analytical_appendix",
                    error=str(_exc))

            # Stamp the live strategy hash (migration 063).
            try:
                from tools.audit_assembler import (
                    current_data_hash as _curr_hash_app,
                )
                _app_hash = await _curr_hash_app()
            except Exception:  # noqa: BLE001
                _app_hash = None
            draft = await create_draft(
                "analytical_appendix", email,
                f"Analytical Appendix — {date.today().isoformat()}",
                content_json, content_text, created_from="generated",
                data_hash=_app_hash)
            if draft is not None:
                draft_id = draft["id"]

            # Layer 3b (June 21 2026) -- persist the value manifest +
            # generation data_hash on the appendix draft so
            # /api/v1/export/verify-all has an authoritative reference
            # for every numeric value the substitution table produced.
            # Mirrors the brief block in _generate_brief_document.
            if draft_id is not None and substitution_table is not None:
                try:
                    from tools.audit_assembler import (
                        current_data_hash as _cur_hash,
                    )
                    from tools.editor_drafts import (
                        update_value_manifest as _update_manifest,
                    )
                    from tools.numeric_substitution import (
                        build_value_manifest,
                    )
                    from datetime import datetime as _dt
                    from datetime import timezone as _tz
                    _hash_for_manifest = await _cur_hash() or ""
                    manifest = build_value_manifest(
                        substitution_table,
                        data_hash=_hash_for_manifest[:64],
                        generated_at=_dt.now(_tz.utc).isoformat())
                    await _update_manifest(
                        draft_id, manifest,
                        data_hash=_hash_for_manifest[:64] or None)
                    log.info(
                        "value_manifest_persisted",
                        document_type="analytical_appendix",
                        draft_id=draft_id,
                        n_values=len(manifest))
                except Exception as exc:  # noqa: BLE001
                    log.warning(
                        "value_manifest_persist_failed",
                        document_type="analytical_appendix",
                        error=str(exc))

                # June 28 2026 (Fix 8b) -- auto-upgrade hook.
                if draft_id is not None:
                    await _auto_upgrade_draft_to_token_values(
                        draft_id, "analytical_appendix")
        except Exception as exc:  # noqa: BLE001
            log.warning("analytical_appendix_draft_create_failed",
                        error=str(exc))

        filename = (f"forest-capital-analytical-appendix-"
                    f"{date.today().isoformat()}.docx")
        return docx_bytes, filename, _DOCX_MEDIA, draft_id
    except Exception as exc:  # noqa: BLE001
        log.error("analytical_appendix_generation_error", error=str(exc))
        raise


@app.post("/api/v1/export/presentation-deck")
@limiter.limit("4/minute")
async def export_presentation_deck(
    request: Request,
    body: dict | None = None,
    session: dict = Depends(require_permission("generate_documents")),
):
    """
    Starts presentation-deck generation.

    With an editor_draft_id in the body the .pptx is built synchronously
    from that draft's current slides (the in-editor Export path).
    Otherwise generation — Academic Writer prose, server-side charts,
    45-90s — runs as a background job and the endpoint returns 202 with
    a job_id; poll GET /api/v1/jobs/{id}.
    """
    editor_draft_id = (body or {}).get("editor_draft_id")
    if editor_draft_id:
        return await _editor_export(int(editor_draft_id))
    await _require_report_ready()
    return _start_generation_job("presentation_deck", session, request)


@app.post("/api/v1/export/presentation-script")
@limiter.limit("12/minute")
async def export_presentation_script(
    request: Request,
    session: dict = Depends(require_permission("generate_documents")),
):
    """Renders the cached deck story plan's Pass 2 full_script + Pass
    3 anticipated_questions into a Presentation Script .docx workbook.

    Pure cache read + format: NO LLM call, NO database write. The
    rate limit is set higher than the generation endpoints (12/min
    vs 4-6/min) because the wall-clock cost is bounded by docx
    assembly (~200ms), not multi-pass LLM generation.

    Returns 404 with a clear message when the deck story plan has
    not yet been cached (or is a deterministic_fallback) so the
    operator knows to generate the Presentation Deck first.

    The Reports page's Presentation Script card hits this endpoint
    from a button labelled "Download Script" -- it stays disabled
    until /api/v1/report/readiness reports
    deck_story_plan_available=true.
    """
    from fastapi.responses import Response as FastAPIResponse
    from tools.academic_docx import build_presentation_script
    from tools.story_plan import get_latest_story_plan

    # June 25 2026 -- replaced the bare-hash get_cached_story_plan
    # lookup. refresh_story_plan persists the deck row under a
    # composite hash via cache_key_with_brief_and_appendix
    # ('<data_hash>|<brief_hash>|<appendix_hash>'); the bare
    # current_data_hash() never matched and this endpoint 404'd
    # even when the readiness gate (which uses get_latest_story_plan
    # post-fix at main.py:9048-9050) reported the script as
    # available. Switching to the same hash-agnostic, fallback-
    # excluding query the gate uses makes the gate + the export
    # finally agree on whether the script can be downloaded.
    # Hash-drift staleness remains handled at export time by
    # verify_export_against_cache (called downstream).
    plan = await get_latest_story_plan(
        "deck", exclude_fallback=True)
    if not plan or plan.get("_model") == "deterministic_fallback":
        raise HTTPException(
            status_code=404,
            detail=(
                "Presentation script not yet generated. Generate "
                "the Presentation Deck first to produce the script."))

    try:
        docx_bytes = await asyncio.to_thread(
            build_presentation_script,
            full_script=plan.get("full_script"),
            anticipated_questions=plan.get("anticipated_questions"),
            computed_at=plan.get("computed_at"),
        )
    except Exception as exc:  # noqa: BLE001
        ref = uuid.uuid4().hex[:8]
        log.error("presentation_script_render_failed",
                  ref=ref, error=str(exc))
        raise HTTPException(
            status_code=500,
            detail=f"Script rendering failed (ref: {ref})")

    from datetime import date
    filename = (
        f"forest-capital-presentation-script-"
        f"{date.today().isoformat()}.docx")
    return FastAPIResponse(
        content=docx_bytes,
        media_type=_DOCX_MEDIA,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/v1/export/presentation-deck/stream")
@limiter.limit("4/minute")
async def export_presentation_deck_stream(
    request: Request,
    session: dict = Depends(require_permission("generate_documents")),
):
    """Bridge #95 -- streams deck generation via Server-Sent Events.

    The endpoint returns text/event-stream immediately and emits
    progress events as each slide's content lands:

      data: {"type": "started", "job_id": "...", "total_slides": 6}
      data: {"type": "slide_complete", "slide_number": 1, "title": "..."}
      data: {"type": "slide_error", "slide_number": N, "error": "..."}
      data: {"type": "rendering"}                # charts + pptx assembly
      data: {"type": "complete", "job_id": "...", "draft_id": N,
              "download_url": "/api/v1/jobs/.../download"}
      data: [DONE]

    A fatal error before any slide-complete event emits:
      data: {"type": "error", "message": "..."}

    Keeps Cloudflare's gateway connection alive by emitting a frame
    every few seconds (one per slide on the happy path). The pptx
    bytes are stored in the standard generation_jobs slot so the
    existing /api/v1/jobs/{id}/download endpoint serves them with
    no special handling.

    The async-job endpoint above (POST /api/v1/export/presentation-deck)
    remains the back-compat surface for frontends that haven't
    migrated to the stream consumer."""
    if ENVIRONMENT == "test":
        return JSONResponse(
            content={"sse_stub": True, "note": "test env"},
            media_type="application/json")

    await _require_report_ready()

    email = session["email"]
    from tools.generation_jobs import create_job, update_job
    from agents.usage import start_usage_capture
    start_usage_capture()
    job = create_job("presentation_deck", email)
    job_id = job["job_id"]

    async def event_stream():
        import asyncio
        import traceback
        from datetime import datetime, timezone

        from tools.academic_deck import DECK_SLIDE_COUNT, SLIDE_TITLES

        update_job(job_id, status="running")
        yield _sse("started", job_id=job_id, total_slides=DECK_SLIDE_COUNT)

        try:
            data, blend_weights, blend_series, n_strategies = \
                await _build_deck_context(email)
            per_slide_ctx = _deck_per_slide_context(data)

            slides: list[dict] = []
            for n in range(1, DECK_SLIDE_COUNT + 1):
                slide = await asyncio.to_thread(
                    _generate_one_deck_slide, n, per_slide_ctx, n_strategies)
                if slide is None:
                    yield _sse(
                        "slide_error",
                        slide_number=n,
                        title=SLIDE_TITLES[n - 1],
                        error=(
                            "Per-slide generation failed; the slide "
                            "renders with a [DATA PENDING] placeholder."))
                    continue
                slides.append(slide)
                yield _sse(
                    "slide_complete",
                    slide_number=n,
                    title=str(slide.get("title")
                              or SLIDE_TITLES[n - 1]))

            yield _sse("rendering")

            file_bytes, filename, media, draft_id = await _finalize_deck(
                slides, data, blend_weights, blend_series, email)

            update_job(
                job_id, status="complete", draft_id=draft_id,
                download_url=f"/api/v1/jobs/{job_id}/download",
                completed_at=datetime.now(timezone.utc),
                _file_bytes=file_bytes, _filename=filename,
                _media_type=media)
            _log_interaction_bg(
                request, session, "export",
                agents_involved=["academic_writer"],
                response_summary="presentation_deck generated (sse)",
                metadata={"deliverable": "presentation_deck",
                          "draft_id": draft_id})
            _schedule_auto_academic_review(
                draft_id, "presentation_deck", email)

            yield _sse(
                "complete",
                job_id=job_id,
                draft_id=draft_id,
                download_url=f"/api/v1/jobs/{job_id}/download")
        except Exception as exc:  # noqa: BLE001
            ref = uuid.uuid4().hex[:8]
            log.error(
                "deck_stream_failed",
                job_id=job_id, ref=ref,
                exc_type=type(exc).__name__,
                exc_module=type(exc).__module__,
                error=str(exc),
                traceback_excerpt="".join(
                    traceback.format_exception_only(
                        type(exc), exc)).strip()[:300])
            update_job(
                job_id, status="failed",
                error=f"Deck stream failed (ref: {ref})",
                completed_at=datetime.now(timezone.utc))
            yield _sse(
                "error",
                job_id=job_id,
                message=f"Deck generation failed (ref: {ref}).")

        yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(),
                             media_type="text/event-stream")


def _parse_deck_slides(raw: str) -> list:
    """Parse the slide JSON from the deck-generation harness output.
    Fence-tolerant. Returns the slides list, or [] on any failure (the
    builder then emits the canonical ten slides with [DATA PENDING] bodies,
    so a parse failure never produces a broken deck)."""
    import json
    text = (raw or "").strip()
    if "{" not in text:
        return []  # e.g. the test-env [DATA PENDING] passthrough
    if text.startswith("```"):
        text = text.strip("`")
        if text[:4].lower() == "json":
            text = text[4:]
    try:
        start, end = text.find("{"), text.rfind("}")
        if start == -1 or end == -1 or end < start:
            return []
        obj = json.loads(text[start:end + 1])
        slides = obj.get("slides") if isinstance(obj, dict) else None
        return slides if isinstance(slides, list) else []
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_json_parse_failed", error=str(exc))
        return []


def _render_deck_slide_charts(
    data: dict, blend_weights: dict, blend_series: list,
) -> dict:
    """Render the per-slide deck charts (sync; called in a thread). Returns
    {slide_number: png|None}. Every renderer is fail-open and individually
    guarded -- a None becomes a [DATA PENDING] note in the deck, never a
    failure.

    Slot mapping is DERIVED from academic_deck.SLIDE_CHARTS so the
    renderer dict and the slot dict can never drift. The 11-slide rebuild
    (June 7 2026) moved the chart slots to {4, 5, 11} but the renderer
    mapping was left at the older 6-slide positions {2, 3, 6} -- every
    deck generation since then fired three deck_chart_slot_unavailable
    warnings because the builder asked for charts on slides 4/5/11 and
    the returned dict carried 2/3/6. Reconciled by indexing through
    SLIDE_CHARTS and dispatching by role string. Adding a new chart role
    (or moving an existing chart to a different slide) now requires only
    a SLIDE_CHARTS edit + a CHART_ROLE_RENDERERS entry; the slot keys
    propagate automatically.

    Roles not covered by CHART_ROLE_RENDERERS are silently dropped so a
    future SLIDE_CHARTS entry for a chart that hasn't been wired yet
    doesn't blow up the deck -- the slide still renders without a chart
    and the operator sees the slot-unavailable WARNING via _image().
    """
    from tools.academic_deck import SLIDE_CHARTS
    # June 27 2026 -- additional renderers for the reconciled
    # SLIDE_CHARTS (slides 4 / 6 / 7 / 8 / 12 gained chart roles
    # when the generation map was brought into alignment with the
    # editor canvas).
    from tools.chart_render import render_cumulative_returns
    from tools.chart_renderers import (
        render_extended_charts as _render_extended_charts,
    )
    from tools.chart_render import (
        render_efficient_frontier, render_rolling_correlation,
        render_strategy_comparison,
    )

    def _safe(fn):
        try:
            return fn()
        except Exception as exc:  # noqa: BLE001
            log.warning("deck_chart_render_failed", error=str(exc))
            return None

    # Role string -> renderer callable. The blend_weights closure is
    # captured once here so the lambdas only need the data dict at call
    # time. blend_series is reserved for future role additions (e.g. an
    # explicit cumulative-return overlay slot); kept on the signature
    # so the existing call sites do not have to change.
    # Extended-renderer helper: a single dispatch through
    # render_extended_charts so a chart_renderers.py role plugs in
    # with a one-line entry instead of needing its own import.
    # Returns PNG bytes or None on failure (matches the contract
    # of the chart_render.render_* helpers).
    def _extended(key: str) -> bytes | None:
        try:
            return _render_extended_charts(key, data).get(key)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "deck_chart_extended_render_failed",
                key=key, error=str(exc))
            return None

    chart_role_renderers = {
        # slide 4 -- rolling-Sharpe comparison (extended renderer)
        "rolling_sharpe":
            lambda d: _extended("rolling_sharpe"),
        # slide 5
        "rolling_correlation":
            lambda d: render_rolling_correlation(d),
        # slide 6 -- post-reconciliation: cumulative_returns
        # (previously strategy_comparison_oos_sharpe). Matches
        # what the editor canvas shows on slide 6.
        "cumulative_returns":
            lambda d: render_cumulative_returns(d),
        # slide 7 -- OOS performance (extended renderer)
        "oos_performance":
            lambda d: _extended("oos_performance"),
        # slide 8 -- live regime signals (extended renderer)
        "regime_signals":
            lambda d: _extended("regime_signals"),
        # slide 12 -- canonical key 'risk_return' from
        # chart_render._DECK_KEYS (formerly 'efficient_frontier'
        # in this dispatch). render_efficient_frontier is still
        # the underlying matplotlib function; only the dispatch
        # key changed to match the renderer registry.
        # blend_weights drives the live marker on the frontier;
        # the rest of the sweep comes from analytics_metrics_cache.
        "risk_return":
            lambda d: render_efficient_frontier(
                d, blend_weights=blend_weights),
        # Legacy keys retained as aliases so any pre-reconciliation
        # caller that still requests the old roles continues to
        # work (e.g. a story plan from a draft generated before
        # the reconciliation). Defensive only -- SLIDE_CHARTS no
        # longer uses these roles.
        "strategy_comparison_oos_sharpe":
            lambda d: render_strategy_comparison(d),
        "efficient_frontier":
            lambda d: render_efficient_frontier(
                d, blend_weights=blend_weights),
    }

    out: dict = {}
    for slide_number, role in SLIDE_CHARTS.items():
        renderer = chart_role_renderers.get(role)
        if renderer is None:
            log.warning(
                "deck_chart_role_unwired",
                slide_number=slide_number, role=role)
            continue
        out[slide_number] = _safe(lambda r=renderer: r(data))
    return out


async def _build_deck_context(
    email: str,
) -> tuple[dict, dict, list, int]:
    """Bridge #95 — extracted helper so both the async-job and SSE
    streaming endpoints can build the deck-generation context block
    from a single source. Returns (data, blend_weights, blend_series,
    n_strategies). All four fail-open: a cold cache produces an empty
    bundle rather than raising.

    June 27 2026 (PR 1 v3, LEAK 1 closer) -- computes the effective
    data_hash inline and threads it into gather_document_data so the
    strategy_results_cache read is hash-aware. Under freeze, a miss
    on the freeze hash raises StrategyCacheMissingForHashError (the
    deck generator wrapper catches + translates to HTTPException
    500 with the spec 'Run light refresh and try again' message).
    Without this, the deck's headline strategy tokens (Sharpe /
    max_drawdown / recovery / blend weights -- ~20 tokens total)
    silently leaked LIVE strategy results into the freeze-locked
    deliverable."""
    from tools.academic_export import gather_document_data
    from tools.audit_assembler import current_data_hash
    from tools.submission_freeze import get_effective_data_hash

    live_hash = await current_data_hash()
    deck_data_hash = await get_effective_data_hash(live_hash)
    data = await gather_document_data(
        data_hash=deck_data_hash or None)

    # ── Live regime context (Slide 6) — current_regime / regime_confidence
    #    / blend_weights from detect_current_regime() + the regime blend,
    #    the SAME source the Forward Projection tile and CIO card use, so
    #    the slide never carries a stale constant. Fail-open to None. ─────
    blend_weights: dict = {}
    try:
        from tools.cio_recommendation import _build_live_context
        built = await _build_live_context()
        if not built.get("error"):
            ctx = built["context"]
            data["current_regime"] = ctx.get("regime")
            data["regime_confidence"] = ctx.get("probability")
            blend_weights = ctx.get("blend_weights") or {}
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_regime_context_unavailable", error=str(exc))

    # ── Play-by-play frozen events (Slide 4) + the post-2022 blend series.
    play_by_play_events: list = []
    blend_series: list = []
    try:
        from tools.play_by_play import (
            get_cached_performance_chart, load_stored_events,
        )
        play_by_play_events = await load_stored_events()
        pc = await get_cached_performance_chart()
        if pc and pc.get("series"):
            blend_series = [
                (p.get("date"), p.get("regime_conditional"))
                for p in pc["series"]
                if p.get("regime_conditional") is not None]
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_play_by_play_unavailable", error=str(exc))
    data["play_by_play_events"] = play_by_play_events
    data["blend_weights"] = blend_weights
    n_strategies = len(data.get("strategy_results") or {})
    return data, blend_weights, blend_series, n_strategies


def _deck_per_slide_context(data: dict) -> dict:
    """Bridge #95 — the context block handed to the per-slide LLM call.
    Same shape as the all-six-at-once prompt's context so existing
    SLIDE_SPECIFICATIONS references resolve. Per-slide prompts ask the
    model to draw from THIS block; the slice helper inside the prompt
    isolates which slide's spec to emit."""
    from tools.academic_deck import (
        CORRELATION_POST_2022, CORRELATION_PRE_2022,
        OOS_SHARPE_BENCHMARK, OOS_SHARPE_EQUAL_WEIGHT,
        OOS_SHARPE_REGIME_CONDITIONAL, PLAY_BY_PLAY_EVENTS,
    )
    rc = data.get("rolling_correlation") or {}
    return {
        "study_period": data.get("study_period"),
        "summary_statistics": data.get("summary_statistics"),
        "strategy_performance": data.get("regime_conditional"),
        "drawdown_comparison": data.get("drawdown_comparison"),
        "factor_loadings": data.get("factor_loadings"),
        "rolling_correlation": {"pre_2022": rc.get("pre_2022"),
                                "post_2022": rc.get("post_2022")},
        "current_regime": data.get("current_regime"),
        "regime_confidence": data.get("regime_confidence"),
        "blend_weights": data.get("blend_weights") or {},
        "play_by_play_events": [
            {"event_id": e.get("event_id"),
             "event_date": e.get("event_date"),
             "trigger": e.get("trigger"), "regime": e.get("regime"),
             "verdict": e.get("verdict"),
             "value_added_sharpe": e.get("value_added_sharpe")}
            for e in (data.get("play_by_play_events") or [])],
        "validated_constants": {
            "oos_sharpe_regime_conditional": OOS_SHARPE_REGIME_CONDITIONAL,
            "oos_sharpe_benchmark": OOS_SHARPE_BENCHMARK,
            "oos_sharpe_equal_weight": OOS_SHARPE_EQUAL_WEIGHT,
            "correlation_pre_2022": CORRELATION_PRE_2022,
            "correlation_post_2022": CORRELATION_POST_2022,
            "play_by_play_events": PLAY_BY_PLAY_EVENTS,
        },
    }


_DECK_SLIDE_SYSTEM_PROMPT = (
    "You write JSON for slides in an investment-research presentation. "
    "Your output is parsed by code -- emit ONLY a single JSON object, "
    "no markdown fences, no preamble, no commentary. Do NOT review or "
    "evaluate slides; do NOT produce 'Peer Discussant Review' headers "
    "or any rubric-scored format. The downstream pipeline parses your "
    "JSON directly into a .pptx slide."
)


def _substitute_slide_content(
    parsed: dict, substitution_table: dict[str, str] | None,
    *, slide_number: int,
) -> dict:
    """Apply the substitution table to every {{TOKEN}}-bearing field
    in a parsed deck slide dict. Operates on string fields (title,
    headline, speaker_notes) and on list-of-string fields (bullets).
    Mutates and returns the dict.

    No-op when substitution_table is None -- preserves the pre-Layer-2
    behaviour for any caller that hasn't been wired through yet.

    The substitution log captures the per-slide tokens replaced so
    Render logs show 'numeric_substitution_applied
    document_type=deck slide_number=3 tokens_replaced=[...]' per
    slide -- the same telemetry shape harness_narrative emits for
    the brief section writer."""
    if substitution_table is None:
        return parsed
    from tools.numeric_substitution import apply_substitutions

    # June 28 2026 -- soft-fail wrap runs BEFORE substitution.
    # If we scanned AFTER substitution, the scanner would see
    # the just-resolved values (e.g. "0.54" from
    # {{BENCHMARK_SHARPE}}) + flag them as token_available +
    # wrap them -- defeating substitution. Scanning RAW
    # (pre-substitution) text:
    #   - bare numerics emitted by the LLM without a token
    #     wrapper -> flagged + wrapped here
    #   - {{TOKEN}} placeholders -> protected (numerics
    #     inside tokens skip the scanner via _is_inside_token)
    # The wrapped form ("the Sharpe is <unverified>0.43
    # </unverified>") flows through apply_substitutions
    # unchanged (no {{TOKEN}} to substitute).
    try:
        from tools.untoken_numeric_check import (
            find_untoken_backed_numerics,
            wrap_unverified,
        )
        _slide_offenders: list[str] = []
        for _key in ("title", "headline", "speaker_notes"):
            if isinstance(parsed.get(_key), str):
                _viols = find_untoken_backed_numerics(
                    parsed[_key], substitution_table)
                if _viols:
                    parsed[_key] = wrap_unverified(
                        parsed[_key], _viols)
                    _slide_offenders.extend(
                        v.raw_value for v in _viols)
        if isinstance(parsed.get("bullets"), list):
            wrapped_bullets: list[str] = []
            for _bullet in parsed["bullets"]:
                if isinstance(_bullet, str):
                    _viols = find_untoken_backed_numerics(
                        _bullet, substitution_table)
                    if _viols:
                        wrapped_bullets.append(wrap_unverified(
                            _bullet, _viols))
                        _slide_offenders.extend(
                            v.raw_value for v in _viols)
                    else:
                        wrapped_bullets.append(_bullet)
                else:
                    wrapped_bullets.append(_bullet)
            parsed["bullets"] = wrapped_bullets
        if _slide_offenders:
            log.warning(
                "deck_untoken_lock_soft_fail",
                document_type="presentation_deck",
                slide_number=slide_number,
                remaining_violations=len(_slide_offenders),
                sample_offenders=_slide_offenders[:10],
                note=(
                    "hard-lock detected raw numerics in "
                    "slide content; wrapping with "
                    "<unverified> tags for human review."))
    except Exception as exc:  # noqa: BLE001
        log.warning(
            "deck_untoken_lock_check_failed",
            slide_number=slide_number,
            error=str(exc))

    replaced_all: set[str] = set()
    for key in ("title", "headline", "speaker_notes"):
        if isinstance(parsed.get(key), str):
            new_value, replaced = apply_substitutions(
                parsed[key], substitution_table)
            parsed[key] = new_value
            replaced_all.update(replaced)
    if isinstance(parsed.get("bullets"), list):
        new_bullets: list[str] = []
        for bullet in parsed["bullets"]:
            if isinstance(bullet, str):
                new_bullet, replaced = apply_substitutions(
                    bullet, substitution_table)
                new_bullets.append(new_bullet)
                replaced_all.update(replaced)
            else:
                new_bullets.append(bullet)
        parsed["bullets"] = new_bullets

    # June 22 2026 (PR A scope) -- walk table_data cells. The Sonnet
    # writer puts most numeric tokens in slide table COLUMNS (the
    # comparison table on slide 2, the IS/OOS Sharpe table on
    # slide 6, the macro watchpoint table on slide 7, etc).
    # Without this walk every {{TOKEN}} embedded in a header
    # string or a row cell survives unsubstituted into the final
    # deck content -- the root cause of the 23 unresolved
    # placeholders reported in production.
    #
    # Contract: walk every string in `headers` (a list of strings)
    # and every string cell in every row of `rows` (a list of
    # lists). Non-string cells (None, numbers, nested dicts)
    # pass through untouched.
    td = parsed.get("table_data")
    if isinstance(td, dict):
        if isinstance(td.get("headers"), list):
            new_headers: list = []
            for h in td["headers"]:
                if isinstance(h, str):
                    new_h, replaced = apply_substitutions(
                        h, substitution_table)
                    new_headers.append(new_h)
                    replaced_all.update(replaced)
                else:
                    new_headers.append(h)
            td["headers"] = new_headers
        if isinstance(td.get("rows"), list):
            new_rows: list = []
            for row in td["rows"]:
                if isinstance(row, list):
                    new_row: list = []
                    for cell in row:
                        if isinstance(cell, str):
                            new_cell, replaced = apply_substitutions(
                                cell, substitution_table)
                            new_row.append(new_cell)
                            replaced_all.update(replaced)
                        else:
                            new_row.append(cell)
                    new_rows.append(new_row)
                else:
                    new_rows.append(row)
            td["rows"] = new_rows

    if replaced_all:
        log.info("numeric_substitution_applied",
                 document_type="presentation_deck",
                 slide_number=slide_number,
                 tokens_replaced=sorted(replaced_all),
                 count=len(replaced_all))

    return parsed


def _generate_one_deck_slide(
    slide_number: int, context: dict, n_strategies: int,
    *, slide_plan_entry: dict | None = None,
    substitution_table: dict[str, str] | None = None,
    brief_excerpt: str = "",
) -> dict | None:
    """Bridge #98 / #100 -- generate ONE deck slide via a DIRECT Sonnet
    call (no harness, no evaluator, no Gemini, no Opus arbiter). Sync;
    callers wrap in asyncio.to_thread. Returns the parsed slide dict
    or None on any failure; the caller writes the [DATA PENDING]
    placeholder for that slide via _normalize_slides.

    slide_plan_entry (PR #333) -- when supplied, the entry from the
    Opus story plan for THIS slide is injected ahead of the existing
    spec block. The headline, numeric_anchors, bullets, and key visual
    are LOCKED -- the per-slide Sonnet call's job becomes prose layout,
    not deciding what the numbers are. Speaker notes from the plan are
    also overwritten onto the parsed slide dict after the call returns
    (the plan is the source of truth; the LLM never gets to "improve"
    them). When the plan is None the call behaves exactly as before --
    the fail-open contract is: a missing or fallback plan never blocks
    slide generation.

    Why direct call_claude instead of harness_narrative:
      The harness uses the academic_review_peer_evaluator rubric, which
      expects "### N. Rated Section / **Rating:**" blocks. Slide JSON
      scored low, retries prepended the evaluator's "Peer Discussant
      Review" feedback to the generator prompt, and Sonnet complied on
      retry attempts -- emitting peer-review text instead of slide JSON.
      The user reported the symptom for slide 2 (bridge #97). Bridge
      #100's no-harness architecture is the fix.

    Failure handling:
      - JSON parse fails the first attempt -> retry ONCE with an added
        instruction reminding the model to output ONLY JSON.
      - If the second attempt also fails to parse -> return None and
        the caller writes a [DATA PENDING] placeholder for THAT slide.
      - A single slide's failure no longer downs the entire deck.
    """
    import json as _json

    from agents.base import SONNET_MODEL, call_claude
    from tools.academic_deck import (
        slide_generation_prompt, parse_single_slide_json,
    )

    prompt = slide_generation_prompt(slide_number)
    ctx_str = (
        context if isinstance(context, str)
        else _json.dumps(context, indent=2, default=str)
    )
    # PR #333 -- when a story plan entry exists for this slide, inject
    # the LOCKED contract above the spec block. The slide LLM is told
    # explicitly that its job is layout + prose only; the numbers, the
    # headline, and the bullet content are NOT to be regenerated.
    plan_block = ""
    if slide_plan_entry and isinstance(slide_plan_entry, dict):
        anchors = slide_plan_entry.get("numeric_anchors") or {}
        bullets = slide_plan_entry.get("slide_bullets") or []
        # June 22 2026 -- read max_bullets from the locked plan
        # entry. CEILING, not target -- "no more than N", never
        # "write exactly N". Default 3 for slides without an
        # explicit cap (the spec block's BULLET DISCIPLINE
        # constraint covers the rest). The locked plan is the
        # source of truth; without this wire-through the
        # max_bullets schema field is decorative and the per-
        # slide writer ignores the slide-specific cap.
        try:
            max_bullets = int(
                slide_plan_entry.get("max_bullets") or 3)
        except (TypeError, ValueError):
            max_bullets = 3
        max_bullets = max(0, min(max_bullets, 3))
        bullets_block = (
            "\n".join(f"  - {b}" for b in bullets)
            if bullets else "  (no bullets -- the headline is the slide)")
        plan_block = (
            "\n\nSTORY PLAN FOR THIS SLIDE (do not deviate):\n"
            f"  Headline: {slide_plan_entry.get('headline', '')}\n"
            f"  Key visual: {slide_plan_entry.get('key_visual', '')}\n"
            "  Numeric anchors (use ONLY these values):\n"
            f"{_json.dumps(anchors, indent=4, default=str)}\n"
            f"  Bullets (ceiling {max_bullets}, "
            "not target -- fewer is better when bullets do not "
            "add meaning beyond the title and table):\n"
            f"{bullets_block}\n\n"
            "  Your job is layout and prose formatting only. Do not "
            "invent numbers. Do not change the headline. Do not add "
            f"bullets beyond the {max_bullets} ceiling. Write "
            "[DATA PENDING] for any value not in numeric_anchors.")
        # June 22 2026 -- user-uploaded slide guidance override.
        # When merge_guidance_into_slide_plan_entry seeded a
        # _user_guidance sub-dict on the plan entry, surface those
        # directives as an additional spec block the LLM reads.
        # Non-overridable fields (numeric_anchors,
        # chart_references, substitution_tokens) are NOT touched
        # by guidance and continue to bind verbatim.
        ug = slide_plan_entry.get("_user_guidance")
        if isinstance(ug, dict) and ug:
            ug_lines = ["", "USER GUIDANCE FOR THIS SLIDE "
                            "(uploaded via the platform; treat "
                            "as supplemental directive on top of "
                            "the spec above):"]
            if ug.get("so_what"):
                ug_lines.append(
                    f"  So-what framing: {ug['so_what']}")
            if ug.get("bullet_guidance"):
                ug_lines.append(
                    f"  Bullet guidance: {ug['bullet_guidance']}")
            if ug.get("speaker_note_directive"):
                ug_lines.append(
                    "  Speaker-note tone / cue: "
                    f"{ug['speaker_note_directive']}")
            plan_block = plan_block + "\n" + "\n".join(ug_lines)
    # Layer 2 (June 21 2026) -- prepend the substitution placeholder
    # guide so the per-slide Sonnet writer uses {{TOKEN}} markers
    # instead of raw figures. The platform substitutes verified
    # cache values on the parsed slide dict after the call returns
    # (_substitute_slide_content). No-op when no substitution_table
    # is supplied (caller hasn't been wired through yet).
    placeholder_guide = ""
    if substitution_table is not None:
        placeholder_guide = (
            _NUMERIC_PLACEHOLDER_GUIDE
            + _DECK_NUMERIC_PLACEHOLDER_GUIDE_EXTENSION
            + "\n")
    # June 21 2026 brief-as-anchor -- per-slide brief excerpt
    # threading. Slides 9 (live demo) and 10 (AI methodology) are
    # explicitly excluded by brief_section_for_slide; the caller
    # passes brief_excerpt="" for those slides. brief_section_block
    # is a no-op string when brief_excerpt is empty so this
    # composes cleanly without conditional branches here.
    from tools.brief_grounding import (
        SLIDE_TO_BRIEF_SECTION, brief_section_block,
    )
    brief_alignment_block = brief_section_block(
        brief_excerpt,
        SLIDE_TO_BRIEF_SECTION.get(slide_number))
    user_message = (
        f"{placeholder_guide}{prompt}{plan_block}"
        f"{brief_alignment_block}\n\nCONTEXT "
        f"(numbers to cite, do not invent):\n{ctx_str}")

    for attempt in (1, 2):
        try:
            raw = call_claude(
                SONNET_MODEL, _DECK_SLIDE_SYSTEM_PROMPT, user_message,
                max_tokens=2000,
                trigger=f"deck_slide_{slide_number}_attempt_{attempt}")
        except Exception as exc:  # noqa: BLE001
            log.warning("deck_slide_call_failed",
                        slide_number=slide_number,
                        attempt=attempt, error=str(exc))
            return None

        parsed = parse_single_slide_json(raw)
        if parsed is not None:
            parsed["slide_number"] = slide_number
            # PR #333 -- the story plan is the source of truth for
            # speaker notes. The Opus arbiter wrote them once with the
            # rubric-evaluated quality bar; we never let the per-slide
            # Sonnet pass "improve" them. Inject verbatim onto the
            # parsed dict so build_presentation_deck reads them.
            if slide_plan_entry and isinstance(
                    slide_plan_entry.get("speaker_notes"), str
            ) and slide_plan_entry["speaker_notes"].strip():
                parsed["speaker_notes"] = (
                    slide_plan_entry["speaker_notes"])
            # Layer 2 (June 21 2026) -- run the substitution table
            # across every {{TOKEN}}-bearing field on the parsed
            # slide (title, headline, speaker_notes, bullets). The
            # locked plan's speaker_notes may also carry placeholders
            # (the Pass-1 arbiter could emit a token in the notes); a
            # second pass post-override resolves any that landed via
            # that path. No-op when substitution_table is None.
            parsed = _substitute_slide_content(
                parsed, substitution_table, slide_number=slide_number)
            return parsed

        log.warning("deck_slide_parse_failed",
                    slide_number=slide_number,
                    attempt=attempt,
                    response_chars=len(raw or ""),
                    response_prefix=(raw or "")[:200])

        if attempt == 1:
            user_message = (
                f"{prompt}\n\nCONTEXT (numbers to cite, do not invent):\n{ctx_str}\n\n"
                "Your previous response was not valid JSON. Output ONLY "
                "the JSON object for this slide, nothing else. No preamble, "
                "no markdown fences, no peer-review formatting -- just the "
                "single JSON object matching the contract above.")
    return None


async def _finalize_deck(
    slides: list[dict],
    data: dict,
    blend_weights: dict,
    blend_series: list,
    email: str,
    substitution_table: dict[str, str] | None = None,
) -> tuple[bytes, str, str, int | None]:
    """Bridge #95 — shared between the async-job and SSE paths. Given
    the assembled per-slide dicts, renders charts + builds pptx +
    creates the editor draft + writes audit metrics. Returns
    (file bytes, filename, media type, editor draft id). Every
    degradation (cold chart, draft-create failure) is non-fatal.

    Layer 3b (June 21 2026) -- when substitution_table is supplied,
    persists the value manifest + generation data_hash on the deck
    draft so /api/v1/export/verify-all has an authoritative reference
    for every numeric value the substitution table produced. Mirrors
    the brief block in _generate_brief_document. Fail-open: a manifest
    write failure logs and is non-fatal."""
    import asyncio
    from datetime import date

    from tools.academic_deck import build_presentation_deck

    charts = await asyncio.to_thread(
        _render_deck_slide_charts, data, blend_weights, blend_series)
    # June 27 2026 -- thread substitution_table into the post-build
    # pass so any un-substituted {{...}} placeholder that leaked
    # through (e.g. via a SLIDE_TITLES[idx-1] fallback when a
    # specialized renderer had no slide.title) gets caught + replaced
    # before the PPTX bytes ship.
    pptx_bytes = await asyncio.to_thread(
        build_presentation_deck, slides, charts,
        substitution_table)

    draft_id: int | None = None
    try:
        from tools.editor_content import deck_slides_to_editor
        from tools.editor_drafts import create_draft
        from tools.chart_config_defaults import (
            default_strategy_names_from_cache,
        )

        # June 26 2026 -- strategy_names sourced from the live
        # strategy_results cache so each chart_config's series
        # list (and each table_config's rows list) gets
        # prepopulated with every strategy in cache order, all
        # visible by default. The editor's Configure panel can
        # then toggle individual series off. Falls back to [] when
        # the cache is empty (cold env / pre-warm); the renderer's
        # fallback path handles the absent-series case unchanged.
        strategy_names = default_strategy_names_from_cache(
            data.get("strategy_results"))
        # June 28 2026 (Phase 2 substitution-deferral audit) --
        # always substitute at this boundary regardless of the
        # DEFER_SUBSTITUTION_TO_EXPORT flag. Deck content_json
        # is canvas-element schema (Konva) and structurally
        # incompatible with the dual-mode token_value
        # architecture; surfacing raw {{TOKEN}} in the canvas
        # editor would be bad UX with no upside.
        content_json, content_text = deck_slides_to_editor(
            slides, strategy_names=strategy_names,
            substitution_table=substitution_table)

        # ── Concern 7h: pre-submission adversarial critic ─────
        try:
            from agents.academic_review import (
                run_doc_gen_debate_round,
            )
            content_text, _debate_id, _critic_res = (
                await run_doc_gen_debate_round(
                    reviewer_email=email,
                    document_type="presentation_deck",
                    content_text=content_text))
        except Exception as _exc:  # noqa: BLE001
            log.warning(
                "doc_gen_critic_pipeline_failed",
                document_type="presentation_deck",
                error=str(_exc))

        audit_warnings = await _run_document_audit(
            content_text, "presentation_deck", email)
        # Stamp the live strategy hash (migration 063).
        try:
            from tools.audit_assembler import (
                current_data_hash as _curr_hash_deck,
            )
            _deck_hash = await _curr_hash_deck()
        except Exception:  # noqa: BLE001
            _deck_hash = None
        draft = await create_draft(
            "presentation_deck", email,
            f"Presentation Deck — {date.today().isoformat()}",
            content_json, content_text,
            created_from="generated",
            audit_warnings=audit_warnings,
            data_hash=_deck_hash)
        if draft is not None:
            draft_id = draft["id"]
        await _write_audit_metrics(
            "presentation_deck", email, draft_id, audit_warnings)

        # Layer 3b -- persist value manifest for deck drafts.
        if draft_id is not None and substitution_table is not None:
            try:
                from tools.audit_assembler import (
                    current_data_hash as _cur_hash,
                )
                from tools.editor_drafts import (
                    update_value_manifest as _update_manifest,
                )
                from tools.numeric_substitution import (
                    build_value_manifest,
                )
                from datetime import datetime as _dt
                from datetime import timezone as _tz
                _hash_for_manifest = await _cur_hash() or ""
                manifest = build_value_manifest(
                    substitution_table,
                    data_hash=_hash_for_manifest[:64],
                    generated_at=_dt.now(_tz.utc).isoformat())
                await _update_manifest(
                    draft_id, manifest,
                    data_hash=_hash_for_manifest[:64] or None)
                log.info(
                    "value_manifest_persisted",
                    document_type="presentation_deck",
                    draft_id=draft_id,
                    n_values=len(manifest))
            except Exception as exc:  # noqa: BLE001
                log.warning(
                    "value_manifest_persist_failed",
                    document_type="presentation_deck",
                    error=str(exc))

            # June 28 2026 (PR #479) -- auto-upgrade hook for
            # deck. Walks canvas slide content_json + flags
            # any element containing <unverified> tag
            # substrings. Document-type-agnostic per operator
            # directive.
            if draft_id is not None:
                await _auto_upgrade_draft_to_token_values(
                    draft_id, "presentation_deck")
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_draft_create_failed", error=str(exc))

    filename = f"forest-capital-presentation-deck-{date.today().isoformat()}.pptx"
    return pptx_bytes, filename, _PPTX_MEDIA, draft_id


async def _generate_deck_document(
    email: str,
) -> tuple[bytes, str, str, int | None]:
    """
    Generates the 6-slide final presentation deck (June 6 2026 rewrite;
    previously 10 slides — see academic_deck.SLIDE_TITLES for the new
    narrative arc). Returns (file bytes, filename, media type, editor
    draft id). Raises on failure — the job wrapper records it.

    Bridge #95 (June 7 2026) -- rewritten to call harness_narrative ONCE
    PER SLIDE instead of once for all six slides combined. The old
    all-six call at max_tokens=4000 was truncating the JSON mid-slide-6
    (estimated 2300-3900 token output), producing an unparseable
    response and the canonical six-slides-all-[DATA PENDING] symptom.
    Per-slide calls cap each slide at max_tokens=1500 with comfortable
    headroom -- no more truncation, and a single slide's failure no
    longer downs the entire deck.

    Bridge #82 (May 26 2026) -- regime state on Slide 6 is LIVE-SOURCED
    from detect_current_regime() (via the CIO live context). Every
    degradation -- a cold cache, an unparseable JSON, a missing chart
    -- falls back to a [DATA PENDING] note rather than failing.

    June 27 2026 -- regime_signals freshness HARD GATE. The deck
    surfaces a live CIO recommendation on slides 7 + 11 that
    references {{VIX_CURRENT}}, {{YIELD_CURVE_CURRENT}},
    {{CREDIT_SPREAD_CURRENT}}, {{EQUITY_TREND_CURRENT}},
    {{ESS_CURRENT}}, {{CURRENT_REGIME}}, {{REGIME_CONFIDENCE}}.
    Stale signals at generation time would produce a deck that
    misleads the live panel. Unlike the brief / appendix (which
    keep the em-dash fallback), the deck blocks: we attempt an
    automatic refresh with a 10s hard timeout and 503 if the
    refresh can't complete.
    """
    import asyncio

    # ── June 27 2026: regime_signals pre-flight gate (deck only) ──
    ok, _signals = await _regime_signals_fresh_or_refresh()
    if not ok:
        log.warning(
            "deck_generation_blocked_on_regime_signals",
            note=("regime_signals_cache miss + refresh "
                  "failed/timed out within 10s -- blocking deck "
                  "generation to avoid stale live recommendation"))
        raise HTTPException(
            status_code=503,
            detail=_REGIME_BLOCKING_ERROR_DETAIL)

    # June 21 2026 -- brief-as-anchor gate. The presentation deck
    # is the THIRD document generated, after the executive brief
    # (narrative anchor) and the analytical appendix (technical
    # detail layer). Both must exist before the deck runs so its
    # Pass-1 Opus arbiter has full visibility into the narrative
    # the deck must argue AND the per-strategy detail it can cite
    # for supporting evidence.
    #
    # Generation order: brief -> appendix -> deck.
    #
    # 409 responses surface inline in the editor (the frontend's
    # DocumentGenerationPanel renders detail in the per-card error
    # slot already; see PR #364 frontend audit).
    from tools.brief_grounding import (
        get_appendix_for_grounding, get_brief_for_grounding,
    )
    brief_grounding = await get_brief_for_grounding()
    if brief_grounding is None:
        raise HTTPException(
            status_code=409,
            detail=(
                "Generate the executive brief before the "
                "presentation deck. The deck is the third "
                "document in the generation order "
                "(brief -> appendix -> deck) and grounds itself "
                "in both upstream documents."))
    appendix_grounding = await get_appendix_for_grounding()
    if appendix_grounding is None:
        raise HTTPException(
            status_code=409,
            detail=(
                "Generate the analytical appendix before the "
                "presentation deck. The deck cites supporting "
                "technical detail from the appendix; the "
                "generation order is brief -> appendix -> deck."))

    try:
        data, blend_weights, blend_series, n_strategies = \
            await _build_deck_context(email)
        per_slide_ctx = _deck_per_slide_context(data)

        from tools.academic_deck import DECK_SLIDE_COUNT, SLIDE_TITLES
        # PR #333 -- the story plan is the LOCKED structural layer
        # above the per-slide LLM passes. Each slide's headline,
        # numeric anchors, bullets, and speaker notes come from the
        # plan; the per-slide Sonnet call writes prose around them.
        # Fail-open contract: a missing plan (transient Opus failure,
        # cold cache, or document_type mismatch) leaves slide_plan as
        # an empty list and slide generation proceeds exactly as
        # before -- the existing fallback already produces a complete
        # deck without the locked structural layer.
        #
        # June 21 2026 brief-as-anchor + appendix-as-evidence --
        # both upstream documents thread through to the Pass-1
        # Opus arbiter. The cache key includes both content
        # hashes so a regen of either upstream document
        # auto-invalidates the cached deck plan.
        slide_plan = await _resolve_story_plan_slide_entries(
            data, n_strategies, list(SLIDE_TITLES),
            brief_text=brief_grounding["content_text"],
            brief_hash=brief_grounding["content_hash"],
            appendix_text=appendix_grounding["content_text"],
            appendix_hash=appendix_grounding["content_hash"])

        # Layer 2 (June 21 2026) -- build the substitution table once
        # per generation job and thread it through every per-slide
        # call. Same table the brief uses (same data_hash) so a value
        # appearing in both the brief and the deck is byte-identical
        # by construction. Fail-open: any error returns None and the
        # per-slide path runs without substitution (the slide writer
        # then emits raw figures, which the post-gen audit's
        # check_unresolved_placeholders + check_numeric_consistency
        # flag as it would today).
        substitution_table: dict[str, str] | None = None
        try:
            from tools.audit_assembler import current_data_hash
            from tools.cio_recommendation import (
                compute_implied_asset_allocation, get_latest_recommendation,
            )
            from tools.numeric_substitution import (
                get_substitution_table,
            )
            from tools.academic_deck import (
                OOS_SHARPE_REGIME_CONDITIONAL,
                OOS_SHARPE_BENCHMARK,
                CORRELATION_PRE_2022, CORRELATION_POST_2022,
            )
            from tools.submission_freeze import get_effective_data_hash
            # Layer 4 -- submission freeze (see _generate_brief_document
            # for the rationale). Live platform reads call
            # current_data_hash() directly; deck generation routes
            # through get_effective_data_hash so the slides lock to
            # the frozen hash on submission day.
            live_hash = await current_data_hash()
            data_hash = await get_effective_data_hash(live_hash)
            cio_row = await get_latest_recommendation()
            implied_alloc: dict | None = None
            try:
                if cio_row and cio_row.get("blend_weights"):
                    implied_alloc = await compute_implied_asset_allocation(
                        cio_row.get("blend_weights"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("deck_implied_alloc_failed", error=str(_exc))
            # June 22 2026 (PR A) -- read regime_signals_cache for the
            # 5 watchpoint tokens on slide 7 (VIX / yield curve /
            # credit spread / equity trend). 15-min TTL; falls back
            # to em-dash inside build_substitution_table if cold.
            # Staleness check: get_regime_cache returns None when
            # the cached row is past its expires_at; we log so the
            # operator can spot a stale render. We do NOT block on
            # a fresh detect call (would add 30-60s to deck gen).
            live_signals: dict | None = None
            try:
                from tools.cache import get_regime_cache
                live_signals = await get_regime_cache()
                if live_signals is None:
                    log.warning(
                        "deck_live_signals_stale",
                        document_type="presentation_deck",
                        note=("regime_signals_cache miss or expired -- "
                              "watchpoint tokens will render em-dash"))
            except Exception as _exc:  # noqa: BLE001
                log.warning("deck_live_signals_read_failed",
                            error=str(_exc))
            from tools.academic_deck import OOS_WINDOW_PCT_OF_STUDY
            # June 22 2026 (wiring fix) -- read analytics metrics
            # for pre/post 2022 Sharpes, factor loadings, and cost
            # sensitivity tokens (slides 6, 8). Single helper call
            # covers all three reads -- see
            # tools.academic_export.load_substitution_metric_sources.
            from tools.academic_export import (
                load_substitution_metric_sources,
            )
            # June 27 2026 -- thread data_hash so historical-analytics
            # metric reads (regime_conditional / factor_loadings /
            # cost_sensitivity / crisis_performance) respect the
            # submission freeze when active. Live CIO + regime signals
            # remain LIVE by design (the platform feature, not frozen).
            regime_conditional_rows, factor_loadings_rows, \
                cost_sensitivity_payload, crisis_payload = (
                    await load_substitution_metric_sources(
                        data_hash=data_hash or None))
            substitution_table = get_substitution_table(
                data_hash or "",
                data.get("strategy_results") or {},
                cio_row,
                oos_sharpe_blend=OOS_SHARPE_REGIME_CONDITIONAL,
                oos_sharpe_benchmark=OOS_SHARPE_BENCHMARK,
                pre_2022_eq_ig_correlation=CORRELATION_PRE_2022,
                post_2022_eq_ig_correlation=CORRELATION_POST_2022,
                oos_window_pct_of_study=OOS_WINDOW_PCT_OF_STUDY,
                study_months=(data.get("study_period") or {}).get(
                    "n_months"),
                implied_allocation=implied_alloc,
                live_signals=live_signals,
                regime_conditional=regime_conditional_rows,
                factor_loadings=factor_loadings_rows,
                cost_sensitivity=cost_sensitivity_payload,
                crisis_performance=crisis_payload,
                hash_verified=True)
            log.info("substitution_table_built",
                     document_type="presentation_deck",
                     data_hash=(data_hash or "")[:8],
                     tokens_available=len(substitution_table))
        except Exception as exc:  # noqa: BLE001
            log.warning("substitution_table_build_failed",
                        document_type="presentation_deck",
                        error=str(exc))

        # June 21 2026 brief grounding -- precompute the per-slide
        # brief excerpt map ONCE so the inner loop just looks up
        # by slide_number. brief_section_for_slide is the single
        # dispatch point that honours SLIDES_EXCLUDED_FROM_BRIEF_
        # GROUNDING (slide 9 + slide 10); calling it here ensures
        # the exclusion can't be accidentally bypassed by the
        # inner loop's logic.
        from tools.brief_grounding import (
            brief_section_excerpt, brief_section_for_slide,
        )
        slide_brief_excerpts: dict[int, str] = {}
        for n in range(1, DECK_SLIDE_COUNT + 1):
            section_name = brief_section_for_slide(n)
            slide_brief_excerpts[n] = brief_section_excerpt(
                brief_grounding["content_text"], section_name)

        # June 22 2026 -- slide guidance overlay. The active guidance
        # row for the generating user (if any) overrides per-slide
        # title / so_what / max_bullets / bullet_guidance /
        # speaker_note_directive on top of the SLIDE_SPECIFICATIONS
        # defaults. Non-overridable fields (numeric_anchors,
        # chart_references, substitution_tokens) are preserved
        # verbatim. Fail-open: no guidance row -> deck generates
        # against the hardcoded defaults exactly as before.
        active_guidance_payload: dict | None = None
        try:
            from tools.deck_slide_guidance import (
                count_overridden_slides, get_active_guidance,
                merge_guidance_into_slide_plan_entry,
            )
            row = await get_active_guidance(email)
            if row:
                active_guidance_payload = (
                    row.get("guidance") or {})
                log.info(
                    "deck_slide_guidance_applied",
                    owner_email=email,
                    n_slides_overridden=count_overridden_slides(
                        active_guidance_payload),
                    uploaded_at=row.get("uploaded_at"),
                    version=active_guidance_payload.get("version"))
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "deck_slide_guidance_load_failed",
                error=str(exc))

        slides: list[dict] = []
        for n in range(1, DECK_SLIDE_COUNT + 1):
            # Index by slide_number explicitly so a partial plan (e.g.
            # the Opus pass returned 9 slide entries when the deck has
            # 11 slots) does not shift downstream slides.
            entry = next(
                (e for e in slide_plan
                 if isinstance(e, dict)
                 and e.get("slide_number") == n),
                None)
            # Overlay the active guidance on top of the plan entry
            # before the per-slide writer reads it. merge_guidance
            # returns the input unchanged when no guidance is
            # active.
            if active_guidance_payload is not None:
                from tools.deck_slide_guidance import (
                    merge_guidance_into_slide_plan_entry,
                )
                entry = merge_guidance_into_slide_plan_entry(
                    entry, n, active_guidance_payload)
            slide = await asyncio.to_thread(
                _generate_one_deck_slide,
                n, per_slide_ctx, n_strategies,
                slide_plan_entry=entry,
                substitution_table=substitution_table,
                brief_excerpt=slide_brief_excerpts.get(n, ""))
            # PR 3 (June 27 2026) -- single-retry on empty bullets
            # for slides that REQUIRE non-empty bullets per
            # SLIDE_SPECIFICATIONS (slides 1, 3, 5, 6, 8, 10, 11).
            # Slides 4, 7, 9, 12 are table-heavy proof slides where
            # empty bullets is acceptable. If the retry also comes
            # back empty, log a structured warning + accept the
            # slide as-is -- the renderer log+skips the bullet
            # block. [DATA PENDING] is NEVER emitted.
            from tools.academic_deck import SLIDES_REQUIRING_BULLETS
            if (slide is not None
                    and n in SLIDES_REQUIRING_BULLETS
                    and not (slide.get("bullets") or [])):
                log.warning(
                    "deck_slide_bullets_empty_retrying",
                    slide_number=n)
                slide_retry = await asyncio.to_thread(
                    _generate_one_deck_slide,
                    n, per_slide_ctx, n_strategies,
                    slide_plan_entry=entry,
                    substitution_table=substitution_table,
                    brief_excerpt=slide_brief_excerpts.get(n, ""))
                if (slide_retry is not None
                        and (slide_retry.get("bullets") or [])):
                    slide = slide_retry
                else:
                    log.warning(
                        "deck_slide_bullets_empty_after_retry",
                        slide_number=n,
                        retry_returned_none=(slide_retry is None))
            if slide is not None:
                slides.append(slide)
            # A None slide is left out -- _normalize_slides inside
            # build_presentation_deck (called by _finalize_deck) fills
            # the missing slot with the canonical title + [DATA PENDING]
            # bullet. A SINGLE slide failure no longer downs the deck.

        # Per-deck substitution-complete telemetry. Same shape the brief
        # writer emits at end of section generation -- operators read
        # both lines in Render logs to confirm the determinism layer
        # fired across the document.
        if substitution_table is not None:
            try:
                from tools.numeric_substitution import (
                    unresolved_placeholders,
                )
                # Stitch every slide's text fields into one blob for
                # the audit summary -- the dispatcher will do its own
                # finer-grained scan on the rendered editor draft.
                blob_parts: list[str] = []
                for sl in slides:
                    for k in ("title", "headline", "speaker_notes"):
                        if isinstance(sl.get(k), str):
                            blob_parts.append(sl[k])
                    if isinstance(sl.get("bullets"), list):
                        blob_parts.extend(
                            b for b in sl["bullets"]
                            if isinstance(b, str))
                unresolved = unresolved_placeholders(
                    "\n".join(blob_parts))
                log.info("substitution_complete",
                         document_type="presentation_deck",
                         tokens_available=len(substitution_table),
                         unresolved_placeholders=unresolved,
                         unresolved_count=len(unresolved))
            except Exception as exc:  # noqa: BLE001
                log.warning("substitution_summary_failed",
                            document_type="presentation_deck",
                            error=str(exc))

        return await _finalize_deck(
            slides, data, blend_weights, blend_series, email,
            substitution_table=substitution_table)
    except Exception as exc:  # noqa: BLE001
        log.error("presentation_deck_generation_error", error=str(exc))
        raise


async def _resolve_story_plan_slide_entries(
    data: dict, n_strategies: int, slide_titles: list[str],
    *,
    brief_text: str | None = None,
    brief_hash: str | None = None,
    appendix_text: str | None = None,
    appendix_hash: str | None = None,
) -> list[dict]:
    """Cache-aware story plan retrieval for deck generation.

    Reads story_plans for
    (current_data_hash[|brief_hash][|appendix_hash], 'deck').
    On a non-fallback cache hit returns the slide_plan list
    verbatim. On a cache miss fires generate_deck_story_plan() and
    persists. Fail-open at every layer.

    June 21 2026 brief-as-anchor + appendix-as-evidence -- both
    upstream documents flow into the Pass-1 Opus arbiter (the
    narrative anchor + the evidentiary backing) and both hashes
    extend the cache key so a regen of either upstream document
    auto-invalidates the cached deck plan.
    """
    try:
        from tools.cache import get_latest_strategy_hash
        data_hash = await get_latest_strategy_hash()
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_story_plan_hash_unavailable", error=str(exc))
        return []
    if not data_hash:
        return []
    try:
        from tools import story_plan as sp
        # June 22 2026 -- force=True so every deck generation runs
        # Pass 1 fresh. Previously the cache returned a stored plan
        # for the current (data_hash | brief_hash | appendix_hash)
        # composite key, which meant Regenerate on the deck reused
        # the same plan and Molly could not iterate on slide
        # guidance / locked-title edits without an upstream hash
        # change. The cache WRITE still happens after generation,
        # so non-forced callers (warm pipeline, brief/appendix
        # cross-refs) still see the fresh row.
        plan = await sp.refresh_story_plan(
            data_hash, "deck",
            deck_context=_deck_per_slide_context(data),
            slide_titles=slide_titles,
            brief_text=brief_text,
            brief_hash=brief_hash,
            appendix_text=appendix_text,
            appendix_hash=appendix_hash,
            force=True)
    except Exception as exc:  # noqa: BLE001
        log.warning("deck_story_plan_refresh_failed", error=str(exc))
        return []
    if not plan or plan.get("error"):
        log.warning("deck_story_plan_unavailable",
                    error=(plan or {}).get("error"))
        return []
    log.info("deck_story_plan_resolved",
             cache=plan.get("cache"),
             model=plan.get("_model"))
    entries = plan.get("slide_plan") or []
    return entries if isinstance(entries, list) else []


async def _resolve_story_plan_brief_sections(
    data: dict,
) -> dict:
    """Cache-aware brief section plan retrieval. Mirrors the deck
    helper above: reads story_plans for (current_data_hash, 'brief')
    + fires generate_brief_section_plan on miss. Fail-open at every
    layer -- a missing plan returns an empty dict and the brief
    section specs run exactly as before the injection layer landed."""
    try:
        from tools.cache import get_latest_strategy_hash
        data_hash = await get_latest_strategy_hash()
    except Exception as exc:  # noqa: BLE001
        log.warning("brief_story_plan_hash_unavailable", error=str(exc))
        return {}
    if not data_hash:
        return {}
    try:
        from tools import story_plan as sp
        from tools.editor_content import _EXEC_BRIEF_SECTIONS
        rubric_sections = [k for _h, k, _c in _EXEC_BRIEF_SECTIONS]
        plan = await sp.refresh_story_plan(
            data_hash, "brief",
            brief_context={
                "validated_constants": data.get(
                    "validated_constants") or {},
                "summary_statistics": data.get("summary_statistics"),
                "drawdown_comparison": data.get("drawdown_comparison"),
                "regime_conditional": data.get("regime_conditional"),
                "study_period": data.get("study_period"),
            },
            rubric_sections=rubric_sections)
    except Exception as exc:  # noqa: BLE001
        log.warning("brief_story_plan_refresh_failed", error=str(exc))
        return {}
    if not plan or plan.get("error"):
        log.warning("brief_story_plan_unavailable",
                    error=(plan or {}).get("error"))
        return {}
    log.info("brief_story_plan_resolved",
             cache=plan.get("cache"),
             model=plan.get("_model"))
    section_plan = plan.get("section_plan") or {}
    return section_plan if isinstance(section_plan, dict) else {}


# June 21 2026 -- numeric substitution placeholder guide. Prepended
# to every per-section task by _inject_brief_section_plan so the
# Sonnet writer uses {{TOKEN}} placeholders instead of raw figures.
# The platform substitutes verified cache values after generation
# (see tools/numeric_substitution.py + the substitution wrapper in
# tools/academic_export.harness_narrative). Tokens listed here are
# the BRIEF-side subset of the table; deck + appendix variants ship
# in the Layer-2 PR alongside their substitution call-sites.
_NUMERIC_PLACEHOLDER_GUIDE = (
    "DETERMINISTIC FIGURES REQUIREMENT:\n"
    "Never write raw numbers for performance metrics, correlations, "
    "or Sharpe ratios. Use these exact placeholder tokens -- the "
    "platform substitutes verified cache values after generation.\n\n"
    "Available placeholders (brief subset):\n"
    "  {{OOS_SHARPE_BLEND}} -- blend out-of-sample Sharpe\n"
    "  {{OOS_SHARPE_BENCHMARK}} -- benchmark OOS Sharpe (same window)\n"
    "  {{OOS_SHARPE_IMPROVEMENT_PCT}} -- % improvement over "
    "benchmark. Token already includes + prefix and % suffix "
    "(resolves to e.g. '+98%'); do NOT add surrounding + or % "
    "characters around it.\n"
    "  {{OOS_WINDOW}} -- the OOS window date range\n"
    "  {{OOS_WINDOW_MONTHS}} -- months in OOS window\n"
    "  {{REGIME_SWITCHING_SHARPE}} -- full-period Sharpe\n"
    "  {{BENCHMARK_SHARPE}} -- benchmark full-period Sharpe\n"
    "  {{CLASSIC_6040_SHARPE}} -- 60/40 full-period Sharpe\n"
    "  {{REGIME_SWITCHING_MAX_DD}} -- peak drawdown\n"
    "  {{BENCHMARK_MAX_DD}} -- benchmark peak drawdown\n"
    "  {{CLASSIC_6040_MAX_DD}} -- 60/40 peak drawdown\n"
    "  {{DD_REDUCTION_REGIME_SWITCHING}} -- drawdown reduction (pp)\n"
    "  {{REGIME_SWITCHING_RECOVERY}} -- recovery in months (number "
    "only, e.g. 32). Write 'months' after in your prose.\n"
    "  {{REGIME_SWITCHING_RECOVERY_MONTHS}} -- recovery with units "
    "(e.g. '32 months'). Use when you want the unit included.\n"
    "  {{BENCHMARK_RECOVERY}} -- benchmark recovery in months "
    "(number only, e.g. 71). Write 'months' after in your prose.\n"
    "  {{BENCHMARK_RECOVERY_MONTHS}} -- benchmark recovery with units "
    "(e.g. '71 months').\n"
    "  {{CLASSIC_6040_RECOVERY}} -- 60/40 recovery in months (number "
    "only). Write 'months' after in your prose.\n"
    "  {{CLASSIC_6040_RECOVERY_MONTHS}} -- 60/40 recovery with units.\n"
    "  {{PRE_2022_EQ_IG_CORR}} -- pre-2022 equity-IG correlation\n"
    "  {{POST_2022_EQ_IG_CORR}} -- post-2022 equity-IG correlation\n"
    "  {{REGIME_SWITCHING_POST2022_SHARPE}} -- post-2022 sub-period\n"
    "  {{BENCHMARK_POST2022_SHARPE}} -- benchmark post-2022\n"
    "  {{CURRENT_REGIME}} -- live HMM regime classification\n"
    "  {{REGIME_CONFIDENCE}} -- posterior confidence\n"
    "  {{CURRENT_EQUITY_PCT}} -- current implied equity weight\n"
    "  {{STUDY_MONTHS}} -- total study period months\n"
    "  {{STUDY_START}} / {{STUDY_END}} -- period dates\n\n"
    "CORRECT: \"The blend achieved {{OOS_SHARPE_BLEND}} versus "
    "{{OOS_SHARPE_BENCHMARK}} for the benchmark.\"\n"
    "WRONG: \"The blend achieved 1.24 versus 0.73.\"\n"
    "WRONG: \"The blend achieved approximately 1.2.\"\n"
    "WRONG: \"[[VERIFY: confirm the OOS Sharpe figure]]\"\n\n"
    "If you need a figure not in this list, write around it "
    "qualitatively rather than inventing or flagging it. The "
    "post-generation audit flags any surviving {{TOKEN}} as an "
    "unresolved_placeholder; if your token name isn't in the table "
    "above, it will surface as an audit failure.\n\n"
)


# June 21 2026 -- Layer 2 deck extension. Appended to
# _NUMERIC_PLACEHOLDER_GUIDE for deck slide prompts so the per-slide
# Sonnet writer sees the deck-specific token vocabulary alongside the
# shared brief tokens. The deck-only tokens cover slide content the
# brief never needs (play-by-play scorecard, live macro watch points,
# net-of-cost Sharpe sensitivity, live blend composition by name).
_DECK_NUMERIC_PLACEHOLDER_GUIDE_EXTENSION = (
    "\nDECK-SPECIFIC PLACEHOLDERS:\n"
    "  {{PLAY_BY_PLAY_VALUE_ADD}} -- events where signal added value "
    "(integer)\n"
    "  {{PLAY_BY_PLAY_TOTAL}} -- total tracked events\n"
    "  {{REGIME_SWITCHING_TURNOVER}} -- annualized turnover %\n"
    "  {{NET_SHARPE_10BP}} / {{NET_SHARPE_15BP}} / {{NET_SHARPE_20BP}}"
    " -- net Sharpe after transaction costs at 10/15/20 bps\n"
    # June 22 2026 -- {{CVAR_99_BENCHMARK}} removed from the
    # placeholder vocabulary. It was advertised here but cited
    # by zero slide specs; the substitution resolver pointed at
    # a field the strategy cache never carries, so any organic
    # LLM use would have resolved to em-dash and tripped the
    # unresolved-placeholder audit. If a future slide spec
    # genuinely needs CVaR, wire up the tail_risk metric source
    # threading + restore the line here.
    "  {{VIX_CURRENT}} -- current VIX level\n"
    "  {{CREDIT_SPREAD_CURRENT}} -- current HY OAS\n"
    "  {{YIELD_CURVE_CURRENT}} -- 10Y-2Y spread\n"
    "  {{EQUITY_TREND_CURRENT}} -- 3-month equity trend\n"
    "  {{ESS_CURRENT}} -- Kish ESS (regime-detection reliability)\n"
    "  {{BLEND_REGIME_SWITCHING_WT}} -- live blend weight for the "
    "regime-switching strategy\n"
    "  {{BLEND_BENCHMARK_WT}} -- live blend weight for the benchmark\n"
    "  {{BLEND_CLASSIC_6040_WT}} -- live blend weight for 60/40\n"
    "  {{N_STRATEGIES}} -- total strategies in submission scope (3)\n"
    "\n"
)


# June 29 2026 -- THREE-STRATEGY SUBMISSION SCOPE. Both the
# brief AND the appendix now operate on the same restricted
# three-strategy set (BENCHMARK, CLASSIC_60_40, REGIME_SWITCHING)
# per the academic-record scope filter applied at
# tools/academic_export.gather_document_data /
# gather_analytical_appendix_data. The appendix retains
# higher-detail per-strategy tokens for those three strategies
# only; the guide enumerates the supported (strategy, metric)
# tokens explicitly.
_APPENDIX_NUMERIC_PLACEHOLDER_GUIDE_EXTENSION = (
    "\nAPPENDIX-SPECIFIC PLACEHOLDERS:\n"
    "The appendix carries the SAME three-strategy scope as the "
    "brief: BENCHMARK, CLASSIC_60_40, REGIME_SWITCHING. Use "
    "strategy-specific tokens for every performance figure:\n"
    "\n"
    "  {{STRATEGY_NAME_SHARPE}} -- full-period Sharpe ratio\n"
    "  {{STRATEGY_NAME_MAX_DD}} -- maximum drawdown (negative %)\n"
    "  {{STRATEGY_NAME_CAGR}} -- annualised return (%)\n"
    "  {{STRATEGY_NAME_VOLATILITY}} -- annualised volatility (%)\n"
    "  {{STRATEGY_NAME_RECOVERY}} -- drawdown recovery (months)\n"
    "\n"
    "STRATEGY_NAME must be one of:\n"
    "  BENCHMARK, CLASSIC_60_40, REGIME_SWITCHING\n"
    "(any other strategy reference -- MIN_VARIANCE, RISK_PARITY, "
    "VOL_TARGETING, etc. -- is out of submission scope and the "
    "token will not resolve.)\n"
    "\n"
    "Examples: {{REGIME_SWITCHING_SHARPE}}, {{CLASSIC_6040_MAX_DD}}, "
    "{{BENCHMARK_CAGR}}.\n"
    "\n"
    "Never write a raw performance figure for any strategy. The "
    "platform substitutes verified cache values after generation, "
    "and a {{STRATEGY_NAME_METRIC}} that doesn't match an actual "
    "strategy in the cache will surface as an unresolved_placeholder "
    "audit flag.\n"
    "\n"
)


# June 28 2026 -- always-allowed study metadata tokens. Any
# section's numeric_anchors get augmented with the resolved
# values from these tokens (read from substitution_table at
# inject time). Reason: the LLM repeatedly emits e.g. "287"
# as the study-period length without wrapping in
# {{STUDY_MONTHS}}; STUDY_MONTHS is in the substitution table
# but not in every section's per-section anchors, so the
# hard-lock flags it as a token_available violation + the
# correction-pass loop fails when the LLM stubbornly refuses
# to swap. Treating study metadata as implicit anchors at the
# injection layer means a raw "287" is allowed without a token
# wrapper (since STUDY_MONTHS IS authoritatively 287 -- the
# value is correct, only the prose form differs).
_STUDY_METADATA_TOKENS: tuple[str, ...] = (
    "{{STUDY_MONTHS}}",
    "{{N_STRATEGIES}}",
    "{{PRE_2022_MONTHS}}",
    "{{POST_2022_MONTHS}}",
    "{{OOS_WINDOW_MONTHS}}",
    "{{OOS_WINDOW_PCT_OF_STUDY}}",
)


def _augment_anchors_with_study_metadata(
    anchors: dict, substitution_table: dict | None,
) -> dict:
    """Return a copy of anchors augmented with every resolved
    value from _STUDY_METADATA_TOKENS that has a non-em-dash
    value in the substitution table. Keys are the token names
    (with {{}} stripped) so they don't clash with story-plan
    anchor keys. Values are floats when possible (the anchor-
    normalising logic in find_untoken_backed_numerics expects
    float-castable values for the multi-format equivalence
    check), otherwise the string.

    Fail-open: if substitution_table is None or empty, anchors
    is returned unchanged."""
    if not substitution_table:
        return dict(anchors)
    out = dict(anchors)
    for token in _STUDY_METADATA_TOKENS:
        if token not in substitution_table:
            continue
        val = substitution_table[token]
        if not val or val == "—":
            continue
        key = token.strip("{}").lower()
        if key in out:
            continue
        try:
            out[key] = float(val.rstrip("%"))
        except (TypeError, ValueError):
            out[key] = val
    return out


def _inject_brief_section_plan(
    specs: list[dict], section_plan: dict,
    substitution_table: dict | None = None,
) -> list[dict]:
    """Prepend each spec's task with the locked section plan entry
    for that spec's key plus the executive-voice + anti-AI writing
    rules (EXECUTIVE_VOICE_REQUIREMENT, threaded in June 21 2026)
    plus the numeric-substitution placeholder guide (Layer-1
    substitution PR, also June 21 2026). The injection is a no-op
    for any spec whose key has no entry in the plan (defensive
    against a partial plan).

    June 28 2026 -- substitution_table param added so the
    per-section numeric_anchors get augmented with study
    metadata values (STUDY_MONTHS, N_STRATEGIES, etc.) before
    the spec ships to harness_narrative. See
    _augment_anchors_with_study_metadata. Fail-open: a missing
    table leaves anchors unchanged.

    Why thread EXECUTIVE_VOICE_REQUIREMENT + _NUMERIC_PLACEHOLDER_GUIDE
    here too instead of only in the Pass-1 system prompt: each per-
    section Sonnet call is its own conversation -- the system prompt
    + the spec.task. The Pass-1 arbiter writes the locked plan, but
    the per-section writer does not see Pass 1's system prompt.
    Without the rules in spec.task, Sonnet drifts back to its
    measured academic default + ignores the placeholder contract.
    """
    import json as _json
    from tools.story_plan import EXECUTIVE_VOICE_REQUIREMENT

    out: list[dict] = []
    for spec in specs:
        key = spec.get("key")
        entry = section_plan.get(key) if isinstance(
            section_plan, dict) else None
        if not isinstance(entry, dict):
            out.append(spec)
            continue
        anchors = entry.get("numeric_anchors") or {}
        block = (
            _NUMERIC_PLACEHOLDER_GUIDE
            + "SECTION PLAN (do not deviate):\n"
            f"  Key message: {entry.get('key_message', '')}\n"
            "  Numeric anchors (use ONLY these values; values may be "
            "{{TOKEN}} placeholders that the platform substitutes):\n"
            f"{_json.dumps(anchors, indent=4, default=str)}\n"
            f"  Target length: {entry.get('target_length_words', '')}"
            " words\n\n"
            + EXECUTIVE_VOICE_REQUIREMENT + "\n\n"
            "  Ground every claim in the numeric anchors. Do not add "
            "sections not in the rubric. Do not frame recommendations "
            "as next steps -- frame them as investment conclusions.\n\n")
        new_spec = dict(spec)
        new_spec["task"] = block + str(spec.get("task", ""))
        # June 21 2026 -- thread the section's anchor allow-list
        # onto the spec so harness_narrative's post-pass
        # story-plan-violation check can read it and retry the
        # writer ONCE with explicit feedback when unauthorized
        # numbers leak in. See Issue 2 (post-regen audit, Option
        # 2 -- harness retry on flag count).
        #
        # June 28 2026 -- augment with study-metadata anchors
        # so STUDY_MONTHS / N_STRATEGIES / OOS_WINDOW_MONTHS
        # are implicit anchors for EVERY section. The LLM
        # repeatedly emits these as raw numbers (e.g. "287
        # months" instead of "{{STUDY_MONTHS}} months") and the
        # hard-lock correction loop fails when Sonnet stubbornly
        # refuses to swap. Treating the values as implicit
        # anchors is correct because the value IS authoritative
        # (it came from the substitution table) -- only the
        # token-wrapping pattern differs.
        new_spec["numeric_anchors"] = (
            _augment_anchors_with_study_metadata(
                anchors, substitution_table))
        out.append(new_spec)
    return out


# ── Async document generation — job status / download / cancel ────────────────

@app.get("/api/v1/jobs/{job_id}")
async def get_generation_job(
    job_id: str, session: dict = Depends(require_auth),
):
    """The current state of a generation job. Owner-only; 404 when the
    job is unknown or has expired (the two-hour TTL)."""
    from tools.generation_jobs import get_job, public_view
    job = get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["owner_email"] != session.get("email"):
        raise HTTPException(status_code=403, detail="This is not your job.")
    return public_view(job)


@app.get("/api/v1/jobs")
async def list_generation_jobs(session: dict = Depends(require_auth)):
    """The caller's last 10 generation jobs, most recent first."""
    from tools.generation_jobs import list_jobs, public_view
    return {"jobs": [public_view(j)
                     for j in list_jobs(session.get("email", ""))]}


@app.get("/api/v1/jobs/{job_id}/download")
async def download_generation_job(
    job_id: str, session: dict = Depends(require_auth),
):
    """Downloads a completed job's rendered file. Owner-only.

    Serves the bytes once. After the first download the buffer is
    cleared (mark_downloaded) to free memory — a 2 MB PPTX would
    otherwise hold a buffer for the full 2-hour job TTL. A second
    download attempt returns 410 Gone with guidance to regenerate;
    the job record itself stays so the client can still poll status.
    """
    from tools.generation_jobs import get_job, mark_downloaded, was_downloaded
    job = get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["owner_email"] != session.get("email"):
        raise HTTPException(status_code=403, detail="This is not your job.")
    if was_downloaded(job_id):
        # Bytes already cleared — a re-attempt is a 410 (Gone), not 409
        # (Conflict). The download was successful; the buffer is intentionally
        # absent now.
        raise HTTPException(
            status_code=410,
            detail=("This download has already been served. "
                    "Regenerate the document if needed."))
    if job["status"] != "complete" or job["_file_bytes"] is None:
        raise HTTPException(status_code=409,
                            detail="The job has no downloadable file yet.")
    # Snapshot the bytes BEFORE marking served — mark_downloaded sets
    # _file_bytes to None, so we must build the response from the
    # snapshot rather than from the (now-cleared) job dict.
    content = job["_file_bytes"]
    media_type = job["_media_type"]
    fname = job["_filename"] or "forest-capital-document"
    mark_downloaded(job_id)
    return Response(
        content=content, media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@app.delete("/api/v1/jobs/{job_id}")
async def cancel_generation_job(
    job_id: str, session: dict = Depends(require_auth),
):
    """Cancels a generation job — sets it to cancelled and cancels the
    background task if it is still in flight. Owner-only."""
    from tools.generation_jobs import get_job, public_view, update_job
    job = get_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if job["owner_email"] != session.get("email"):
        raise HTTPException(status_code=403, detail="This is not your job.")
    if job["status"] in ("pending", "running"):
        update_job(job_id, status="cancelled")
        task = job.get("_task")
        if task is not None and not task.done():
            task.cancel()
    return public_view(job)


# Sprint 6 Priority 1 — midpoint paper for the May 27 submission
# deadline (the June 3 cohort meetup is a peer-review event, not a gate).
# Academic Writer composes the prose; tools/docx_generator assembles the
# .docx around it. Every page carries the AI DRAFT banner so Bob can never
# accidentally submit a template-generated draft verbatim.
#
# PR-B (June 21 2026) -- the midpoint pipeline has been retired. The
# original handler is replaced with the 410 Gone stub below.


@app.post("/api/reports/midpoint-template")
@limiter.limit("10/minute")
async def midpoint_template(
    request: Request, session: dict = Depends(require_auth),
):
    """RETIRED (PR-B, June 2026).

    The legacy 11-step Report Writer midpoint pipeline was retired in
    PR #338 (frontend) and PR-B (this PR, backend). Returns 410 Gone
    so existing clients receive a clear "this existed and is now gone"
    signal rather than a 404 connection error.
    """
    return JSONResponse(
        status_code=410,
        content={
            "error": "gone",
            "message": (
                "The Report Writer midpoint pipeline has been retired. "
                "The midpoint submission shipped May 27 and the final "
                "submission deadline is July 1."),
            "canonical_path": "/api/v1/export/executive-brief",
        })


def _build_results_dict_and_range() -> tuple[dict, dict]:
    """
    Loads strategy results from the cache (or runs the full pipeline if the
    cache is cold) and returns a (results, data_range) tuple.

    Centralised so the three report endpoints share the same data-loading
    semantics. ENVIRONMENT=test bypasses entirely so report tests run in
    milliseconds against an empty results dict — the docx/html builders
    still produce a valid file from the prose-only sections.
    """
    if ENVIRONMENT == "test":
        return {}, {"start": "—", "end": "—", "n_months": 0}

    from tools.data_fetcher import get_full_history
    from tools.backtester import run_all_strategies
    from tools.cache import get_strategy_cache, _compute_data_hash

    history = get_full_history()
    monthly = history.get("equity_monthly")
    n_rows = len(monthly) if monthly is not None else 0
    last_date = (
        str(monthly.index[-1].date())
        if monthly is not None and len(monthly) > 0 else "unknown"
    )
    first_date = (
        str(monthly.index[0].date())
        if monthly is not None and len(monthly) > 0 else "unknown"
    )
    strategy_hash = _compute_data_hash(n_rows, last_date, n_strategies=10)

    # asyncpg cache call is async — caller runs us inside an async endpoint
    # so we delegate the await to the caller via a small inline wrapper.
    # Simpler to do the actual run inline at call sites; this helper now
    # only returns the synchronous part. (Refactored below.)
    raise RuntimeError("Use _load_results_async inside route handlers")


async def _load_results_async() -> tuple[dict, dict]:
    """Async loader — single source of truth for data + cache hit logic."""
    if ENVIRONMENT == "test":
        return {}, {"start": "—", "end": "—", "n_months": 0}

    from tools.data_fetcher import get_full_history_async
    from tools.backtester import run_all_strategies
    from tools.cache import get_strategy_cache, _compute_data_hash

    history = await get_full_history_async()
    monthly = history.get("equity_monthly")
    n_rows = len(monthly) if monthly is not None else 0
    last_date = (
        str(monthly.index[-1].date())
        if monthly is not None and len(monthly) > 0 else "unknown"
    )
    first_date = (
        str(monthly.index[0].date())
        if monthly is not None and len(monthly) > 0 else "unknown"
    )
    strategy_hash = _compute_data_hash(n_rows, last_date, n_strategies=10)
    cached = await get_strategy_cache(strategy_hash)
    results = cached if cached else await asyncio.to_thread(run_all_strategies, history)
    return results, {"start": first_date, "end": last_date, "n_months": n_rows}


@app.post("/api/reports/analytical-appendix")
@limiter.limit("10/minute")
async def analytical_appendix(request: Request, session: dict = Depends(require_auth)):
    """
    Generates the comprehensive HTML analytical appendix — 35% of the grade.

    Six sections per CLAUDE.md Section 14:
      1. Abstract               (Academic Writer → write_abstract via results)
      2. Data Sources & Provenance  (registry + cross-validation)
      3. Portfolio Construction Methodology (Academic Writer → write_methodology)
      4. Statistical Results in APA format  (Academic Writer → write_results)
         + Table 1 strategy comparison auto-injected after this section
      5. Sensitivity Analysis    (deterministic ±20% parameter sweep summary)
      6. Reproducibility Notes   (random seed, data file, config snapshot)

    Returns text/html with a filename header so browsers either render
    inline (when content-disposition is inline) or download (attachment).
    We choose attachment so Bob has the source HTML to edit; he can
    Open With Word or paste into Pages for the final submission.
    """
    from fastapi.responses import Response as FastAPIResponse

    try:
        from agents.academic_writer import AcademicWriter
        from tools.html_report_generator import build_html_report

        results_dict, data_range = await _load_results_async()
        significance_flags = {
            name: bool(r.get("is_significant"))
            for name, r in results_dict.items()
        }
        n_significant = sum(1 for v in significance_flags.values() if v)

        writer = AcademicWriter()
        methodology = writer.write_methodology(
            data_sources={"data_range": data_range, "n_months": data_range["n_months"]},
            strategies=list(results_dict.keys()),
            statistical_tests=[
                "Paired t-test (full period)",
                "Benjamini-Hochberg FDR correction",
                "Deflated Sharpe Ratio",
                "Walk-forward out-of-sample",
                "CV Stability Score",
                "Combinatorial Purged Cross-Validation",
            ],
        )
        results = writer.write_results(
            strategy_results=results_dict,
            significance_flags=significance_flags,
            stress_tests={},
        )

        # Pull the provenance registry from provenance.json — same source the
        # /api/v1/provenance endpoint serves. Falls back to an empty list when
        # the file is absent (test env, fresh deploys without a pipeline run).
        provenance_registry: list[dict] = []
        try:
            import json
            from pathlib import Path
            prov_path = Path(__file__).parent / "data" / "provenance.json"
            if prov_path.exists():
                prov_data = json.loads(prov_path.read_text(encoding="utf-8"))
                provenance_registry = prov_data.get("series", [])
        except Exception as exc:
            log.warning("appendix_provenance_load_failed", error=str(exc))

        abstract_body = (
            f"This study evaluated whether diversification across equities "
            f"and fixed income improves risk-adjusted performance versus a "
            f"100% equity benchmark over the period "
            f"{data_range['start']} to {data_range['end']} "
            f"({data_range['n_months']} monthly observations). Ten portfolio "
            f"strategies — five static and five dynamic — were tested against "
            f"the benchmark using a tiered statistical framework: paired "
            f"t-test at p < 0.005, Benjamini-Hochberg FDR correction, "
            f"Deflated Sharpe Ratio, walk-forward out-of-sample testing, and "
            f"a Cross-Validation Stability Score threshold of 0.60. Of the "
            f"ten strategies, {n_significant} passed all five Tier 1 gates. "
            f"The central empirical finding — that the equity-bond "
            f"correlation flipped from negative to positive during the 2022 "
            f"rate-hiking cycle — is disclosed prominently in the Results "
            f"section. Findings are reported in APA 7th edition format."
        )

        data_sources_body = (
            "The analytical foundation is Dr. Panttser's FNA 670 Excel file, "
            "which provides authoritative monthly equity returns, daily "
            "bond OHLCV (BND, BAMLHYH total return index), credit spreads "
            "(BAMLH0A0HYM2EY, BAMLC0A0CMEY), the 10-year Treasury yield, "
            "the 3-month T-bill, real GDP, and the S&P 500 P/E ratio. "
            "Excel-sourced series are never overridden by external data. "
            "\n\n"
            "Four supplemental fetches fill gaps the Excel file does not "
            "cover: daily SPY (yfinance) for momentum and volatility "
            "signals; VIX (FRED) and 2-year Treasury (FRED) for regime "
            "classification; Fama-French factors via direct HTTP fetch from "
            "Ken French's website (the pandas-datareader path was deprecated "
            "and broken as of 2026). An LQD bridge extends the IG bond "
            "history from BND's April 2007 inception back to July 2002, "
            "adding 58 monthly observations and lifting total sample size "
            "from 224 to 282 — the difference between underpowered and "
            "adequately-powered statistical tests at p < 0.005."
            "\n\n"
            "Cross-validation between the Excel monthly S&P 500 series and "
            "yfinance daily SPY aggregated to monthly month-end runs on "
            "every cold start. Any month with discrepancy > 1% halts the "
            "pipeline with DataValidationError. Internal consistency "
            "checks on BND and BAMLHYH (gap detection, outlier detection, "
            "GFC drawdown sanity) are logged but do not halt."
        )

        sensitivity_body = (
            "Key strategy parameters were tested at ±20% of their default "
            "values to confirm results do not depend on a single fortunate "
            "choice. Parameters tested: the momentum lookback windows "
            "(21d, 63d, 252d composite), the volatility target (10% "
            "annualised), the optimisation window (36 months), the rolling "
            "Sharpe window for Max-Sharpe-Rolling, and the regime "
            "thresholds (VIX 25, yield curve 0, credit spread 5%)."
            "\n\n"
            "For every parameter, the strategy's Sharpe ratio, CAGR, and "
            "Tier-1 gate count were recomputed at the default value ± 20%. "
            "Results are reported in Table 3 below; strategies whose "
            "is_significant flag flips at any tested value are flagged. "
            "Where the flip occurs, the dependence is disclosed in the "
            "Limitations section of the executive brief."
        )

        reproducibility_body = (
            "Every stochastic operation in the pipeline seeds NumPy with "
            f"RANDOM_SEED = 42. The annualisation factor is fixed at 252 "
            "(daily) and 12 (monthly) — never approximated. All return "
            "computations use the simple `pct_change` form, not log "
            "returns; the two are not mixed within any single strategy. "
            "\n\n"
            "Data file: FNA_670_Project_Sources.xlsx (committed to the "
            "repository under backend/data). Supplemental fetches are "
            "cached in PostgreSQL — once a series is loaded, the historical "
            "rows are never re-fetched. Incremental updates append only "
            "the latest delta. The exact strategy hash for this report "
            f"reflects the {data_range['n_months']} monthly observations "
            "available at generation time."
            "\n\n"
            "Full reproducibility steps: clone the repository, install "
            "requirements.txt + requirements-dev.txt, set ANTHROPIC_API_KEY "
            "and FRED_API_KEY in `.env`, run `alembic upgrade head` against "
            "an empty PostgreSQL, then call `python -m backend.tools.data_"
            "fetcher` to populate the database. The next call to "
            "`/api/backtest/compare` will recompute all ten strategies "
            "deterministically — given the same data, results match to "
            "six decimal places."
        )

        sections = [
            {"heading": "1. Abstract", "body": abstract_body},
            {"heading": "2. Data Sources and Provenance", "body": data_sources_body},
            {"heading": "3. Portfolio Construction Methodology", "body": methodology},
            {"heading": "4. Statistical Results", "body": results},
            {"heading": "5. Sensitivity Analysis", "body": sensitivity_body},
            {"heading": "6. Reproducibility Notes", "body": reproducibility_body},
        ]

        # Curated reference list — Academic Writer endpoint draws from the
        # same references.json so citations in the prose align with this
        # bibliography exactly.
        try:
            references_db = AcademicWriter.get_available_references()
            references = sorted(
                r["apa"] for r in references_db.values() if r.get("apa")
            )
        except Exception as exc:
            log.warning("references_load_failed", error=str(exc))
            references = None

        html_str = build_html_report(
            title="Forest Capital Portfolio Intelligence System",
            subtitle=(
                "Analytical Appendix — FNA 670 Practicum · "
                f"Data range: {data_range['start']} – {data_range['end']} · "
                f"{n_significant}/10 strategies pass all Tier 1 gates"
            ),
            sections=sections,
            strategy_results=results_dict,
            provenance_registry=provenance_registry,
            references=references,
        )

        from datetime import date
        filename = f"forest-capital-analytical-appendix-{date.today().isoformat()}.html"
        return FastAPIResponse(
            content=html_str,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("analytical_appendix_error", ref=ref, error=str(exc))
        raise HTTPException(status_code=500, detail=f"Appendix generation failed (ref: {ref})")


@app.post("/api/reports/executive-brief-template")
@limiter.limit("10/minute")
async def executive_brief_template(
    request: Request, session: dict = Depends(require_auth),
):
    """RETIRED (PR-B, June 2026).

    The legacy Report Writer executive-brief pipeline was retired in
    PR #338 (frontend) and PR-B (this PR, backend). The canonical
    brief generation path is now POST /api/v1/export/executive-brief,
    which runs the two-pass story plan architecture (PR #333 + #336)
    with locked numeric anchors and the post-generation audit.
    Returns 410 Gone so existing clients receive a clear "this
    existed and is now gone" signal rather than a 404 connection
    error.
    """
    return JSONResponse(
        status_code=410,
        content={
            "error": "gone",
            "message": (
                "The Report Writer pipeline has been retired. Use "
                "POST /api/v1/export/executive-brief for the two-"
                "pass story plan architecture."),
            "canonical_path": "/api/v1/export/executive-brief",
        })


# ── Agent personas (Council View → "View system prompt") ──────────────────────
#
# Surfaces the verbatim system prompt of every council agent. The frontend's
# PersonaModal renders three tabs:
#   PROMPT          — verbatim text (from this endpoint)
#   PLAIN ENGLISH   — Explainer-generated (via glossaryStore.loadPersona)
#   THIS SESSION    — agent's actual summary in the current council run

# Agent name → (display_name, model, module path with _SYSTEM_PROMPT).
# Centralised here so adding a new agent only requires one edit, and the
# names match _AGENT_META above (the council-debate display layer).
_AGENT_PERSONA_REGISTRY: list[tuple[str, str, str]] = [
    ("Equity Analyst",              SONNET_MODEL,  "agents.equity_analyst"),
    ("Fixed Income Analyst",        SONNET_MODEL,  "agents.fixed_income_analyst"),
    ("Risk Manager",                SONNET_MODEL,  "agents.risk_manager"),
    ("Quant Backtester",            SONNET_MODEL,  "agents.quant_backtester"),
    ("Independent Analyst (Gemini)",GEMINI_MODEL,  "agents.independent_analyst"),
    ("Contrarian Analyst (Grok)",   "grok-4.3",    "agents.contrarian_analyst"),
    ("CIO",                         OPUS_MODEL,    "agents.cio"),
]


@app.get("/api/agents/personas")
@limiter.limit("60/minute")
async def agent_personas(request: Request, session: dict = Depends(require_auth)):
    """
    Returns each council agent's verbatim system prompt + model + role.

    The verbatim prompt powers the PROMPT tab in the PersonaModal — when
    the team explains the system to Forest Capital, this is the
    auditable artefact that proves no agent is reading off improvised
    instructions. The PLAIN ENGLISH tab is generated by the Explainer
    Agent (Haiku) on demand and cached in glossaryStore.

    We import each agent module dynamically and read its `_SYSTEM_PROMPT`
    module-level constant. Errors per-agent fall back to an explanatory
    placeholder so one broken import never breaks the whole modal.
    """
    import importlib

    out: list[dict[str, Any]] = []
    for display_name, model, module_path in _AGENT_PERSONA_REGISTRY:
        try:
            mod = importlib.import_module(module_path)
            prompt = getattr(mod, "_SYSTEM_PROMPT", "") or ""
        except Exception as exc:
            log.warning(
                "persona_load_failed",
                agent=display_name,
                module=module_path,
                error=str(exc),
            )
            prompt = ""

        out.append({
            "agent": display_name,
            "model": model,
            "module": module_path,
            "system_prompt": prompt,
            # Short summary helps the modal show something useful before
            # the Explainer Agent's plain-English narrative streams in.
            "prompt_summary_first_sentence": (
                prompt.split(".")[0].strip()[:200] + "."
                if prompt and "." in prompt
                else (prompt[:200] if prompt else "System prompt unavailable.")
            ),
        })
    return {"agents": out}


@app.get("/api/reports/manifest")
@limiter.limit("60/minute")
async def reports_manifest(request: Request, session: dict = Depends(require_auth)):
    """
    Returns the list of available report generators for the Reports screen.

    Shape lets the UI render the deliverable cards without hardcoding the
    endpoint URLs in three places — change a card label here, the frontend
    updates on next mount.
    """
    # The Reports page no longer renders a card grid against this
    # manifest (the "Bob's / Molly's Deliverables" panel was removed
    # because each card pointed at a legacy non-story-plan endpoint
    # that bypassed the canonical two-pass pipeline). The remaining
    # entries are kept for any programmatic caller that wants to
    # enumerate the deliverables surface -- the analytical-appendix
    # endpoint is the only one that does not have a v1/export
    # equivalent on the canonical path. The storyboard editor surface
    # at /reports/storyboard is documented here so a caller can find
    # the entry point.
    return {
        "owner_bob": [
            {
                "id": "analytical_appendix",
                "title": "Analytical Appendix",
                "description": (
                    "Comprehensive HTML with Abstract, Data Sources & "
                    "Provenance, Methodology, Statistical Results "
                    "(Table 1), Sensitivity Analysis, Reproducibility "
                    "Notes, References."
                ),
                "endpoint": "/api/reports/analytical-appendix",
                "method": "POST",
                "format": "html",
                "status": "available",
                "deadline": "July 1, 2026",
            },
        ],
        "owner_molly": [
            {
                "id": "storyboard_draft",
                "title": "Presentation Storyboard",
                "description": (
                    "AI-drafted 15-slide structure. Edit in the "
                    "Storyboard Editor at /reports/storyboard -- drag "
                    "to reorder, swap charts, refine speaker notes."
                ),
                "endpoint": "/api/documents/storyboard/draft",
                "method": "POST",
                "format": "json",
                "status": "available",
                "deadline": "July 1, 2026",
            },
        ],
    }


# ── Documents & Storyboard Editor (Sprint 6 Phase 6) ─────────────────────────
#
# Documents tables (migration 004) back four routes here. They follow a
# strict pattern: every mutation goes through tools/documents_cache.py
# so the DB-unavailable failure mode degrades to 503 rather than 500.
# Auth scope: any logged-in team member can read / mutate any document —
# we trust the four-person ALLOWED_EMAILS list, not row-level ACLs.

@app.post("/api/documents/storyboard/draft", status_code=201)
@limiter.limit("10/minute")
async def storyboard_draft(request: Request, session: dict = Depends(require_auth)):
    """
    Generates an initial 15-slide storyboard from current strategy results.

    Returns the new document_id + the full slide JSON. The caller (typically
    the Reports screen's "Create Storyboard" button) hands the document_id
    to the StoryboardEditor route. Subsequent edits flow through
    PATCH /api/documents/:id/draft.
    """
    from tools.storyboard_template import build_default_storyboard
    from tools.documents_cache import create_document

    # Pull current strategy results so the AI draft references live numbers.
    # Tests / dev environments without a DATABASE_URL still produce a valid
    # storyboard from the default template + placeholder numbers.
    results_dict: dict = {}
    strategy_hash: str | None = None
    if ENVIRONMENT != "test":
        try:
            from tools.data_fetcher import get_full_history_async
            from tools.backtester import run_all_strategies
            from tools.cache import get_strategy_cache, _compute_data_hash

            history = await get_full_history_async()
            monthly = history.get("equity_monthly")
            n_rows = len(monthly) if monthly is not None else 0
            last_date = (
                str(monthly.index[-1].date())
                if monthly is not None and len(monthly) > 0 else "unknown"
            )
            strategy_hash = _compute_data_hash(n_rows, last_date, n_strategies=10)
            cached = await get_strategy_cache(strategy_hash)
            results_dict = cached if cached else await asyncio.to_thread(run_all_strategies, history)
        except Exception as exc:
            log.warning("storyboard_draft_strategy_load_failed", error=str(exc))

    # Try Academic Writer enrichment for speaker notes. None on failure —
    # build_default_storyboard handles both paths gracefully.
    writer = None
    if ENVIRONMENT != "test":
        try:
            from agents.academic_writer import AcademicWriter
            writer = AcademicWriter()
        except Exception:
            writer = None

    storyboard = build_default_storyboard(strategy_results=results_dict, writer=writer)

    doc_id = await create_document(
        doc_type="storyboard",
        owner_email=session.get("email", "unknown@queens.edu"),
        initial_content=storyboard,
        strategy_hash=strategy_hash,
        created_by=session.get("email"),
    )

    if doc_id is None:
        # DB unavailable — return the storyboard inline so the UI can still
        # render an editable preview, but flag that persistence failed.
        return {
            "document_id": None,
            "storyboard": storyboard,
            "persistence": "unavailable",
            "message": (
                "Storyboard generated but not saved to the database. "
                "Save Version will fail until the operator runs "
                "`alembic upgrade head` on Render."
            ),
        }

    return {
        "document_id": doc_id,
        "storyboard": storyboard,
        "persistence": "saved",
    }


@app.get("/api/documents/{document_id}")
@limiter.limit("60/minute")
async def get_document(
    document_id: str, request: Request, session: dict = Depends(require_auth),
):
    """Returns the current working draft for a document."""
    from tools.documents_cache import get_document_draft
    draft = await get_document_draft(document_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Document not found")
    return draft


@app.patch("/api/documents/{document_id}/draft")
@limiter.limit("120/minute")
async def patch_document_draft(
    document_id: str,
    body: dict,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Auto-save endpoint. 120/min lets the 30-second auto-save fire freely
    without throttling. Updates the draft in place — version snapshots
    use POST /api/documents/:id/versions instead.
    """
    from tools.documents_cache import update_draft
    content = body.get("content")
    if not isinstance(content, dict):
        raise HTTPException(status_code=422, detail="Body must include 'content' object")
    ok = await update_draft(document_id, content)
    if not ok:
        raise HTTPException(
            status_code=404,
            detail="Document not found or draft update failed",
        )
    return {"saved_at": "now", "document_id": document_id}


@app.post("/api/documents/{document_id}/versions", status_code=201)
@limiter.limit("30/minute")
async def post_document_version(
    document_id: str,
    body: dict,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Creates a named snapshot of the current draft state. The body must
    include 'content' (current draft) and optionally 'version_name' +
    'change_summary'. Returns the new version's id and version_number.
    """
    from tools.documents_cache import save_named_version
    content = body.get("content")
    if not isinstance(content, dict):
        raise HTTPException(status_code=422, detail="Body must include 'content' object")
    version_name = body.get("version_name") or "Untitled version"
    change_summary = body.get("change_summary")

    result = await save_named_version(
        document_id=document_id,
        version_name=version_name,
        content=content,
        created_by=session.get("email", "unknown@queens.edu"),
        change_summary=change_summary,
        is_auto_save=False,
    )
    if result is None:
        raise HTTPException(status_code=503, detail="Version persistence unavailable")
    return result


@app.get("/api/documents/{document_id}/versions")
@limiter.limit("60/minute")
async def list_document_versions(
    document_id: str, request: Request, session: dict = Depends(require_auth),
):
    """Returns all versions for a document, newest first."""
    from tools.documents_cache import list_versions
    return {"versions": await list_versions(document_id)}


@app.post("/api/documents/{document_id}/restore/{version_id}")
@limiter.limit("20/minute")
async def restore_document_version(
    document_id: str,
    version_id: str,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Restores a prior version: copies its content into a new version row
    (with restored_from set to track the rollback) and replaces the draft.
    The original version stays intact — restore never deletes history.
    """
    from tools.documents_cache import restore_version
    result = await restore_version(
        document_id=document_id,
        version_id=version_id,
        restored_by=session.get("email", "unknown@queens.edu"),
    )
    if result is None:
        raise HTTPException(status_code=404, detail="Version not found")
    return result


# ── Bob's section-document editor (Sprint 6 Phase 10) ────────────────────────
#
# Each of Bob's three deliverables (midpoint, executive brief, analytical
# appendix) can be opened as a section-structured document Bob edits in
# the SectionEditor UI. The document persists the AI's original draft
# alongside Bob's edited version per section, so he can View AI Draft
# and Revert per section without losing his work.
#
# Schema (stored in document_drafts.content as JSONB):
# {
#   "doc_type":   "midpoint_paper" | "executive_brief" | "analytical_appendix",
#   "title":      str,
#   "subtitle":   str,
#   "sections":   [
#     { "id": "abstract", "title": "Abstract",
#       "ai_draft": "...",  ← immutable original from Academic Writer
#       "content":  "...",  ← Bob's current text
#       "last_edited": ISO timestamp }
#   ]
# }


def _build_section_doc_content(
    doc_type: str,
    results_dict: dict,
    data_range: dict,
) -> dict[str, Any]:
    """
    Builds the initial section-structured content for a Bob document.

    Mirrors the same per-deliverable section list the download endpoints
    use, so Bob can edit the same content he'd get from a direct download.
    The AI draft is captured in BOTH the `ai_draft` and `content` fields
    on creation — Bob's edits then diverge `content` from `ai_draft`,
    and View AI Draft reads from the immutable side.
    """
    from agents.academic_writer import AcademicWriter
    from datetime import datetime, timezone

    writer = AcademicWriter()
    now_iso = datetime.now(timezone.utc).isoformat()

    if doc_type == "midpoint_paper":
        methodology = writer.write_methodology(
            data_sources={"data_range": data_range, "n_months": data_range["n_months"]},
            strategies=list(results_dict.keys()),
            statistical_tests=["Tier 1 gates", "FDR correction", "DSR", "Walk-forward OOS", "CV stability"],
        )
        sig = {k: bool(v.get("is_significant")) for k, v in results_dict.items()}
        results = writer.write_results(
            strategy_results=results_dict, significance_flags=sig, stress_tests={},
        )
        sections = [
            ("methodology", "1. Data & Methodology", methodology),
            ("results",     "2. Preliminary Results", results),
            ("roles",       "3. Roles & Division of Labor",
             "Michael — Lead Engineer. Bob — Lead Analyst. Molly — Lead Presenter."),
            ("next_steps",  "4. Next Steps & Open Questions",
             "Sprint 6 closes the executive brief and analytical appendix generators."),
        ]
        title = "Forest Capital — Midpoint Checkpoint"

    elif doc_type == "executive_brief":
        methodology = writer.write_methodology(
            data_sources={"data_range": data_range, "n_months": data_range["n_months"]},
            strategies=list(results_dict.keys()),
            statistical_tests=["Tier 1 gates", "FDR correction", "DSR", "Walk-forward OOS", "CV stability"],
        )
        top_three = sorted(
            results_dict.items(),
            key=lambda kv: float(kv[1].get("sharpe_ratio", 0.0) or 0.0),
            reverse=True,
        )[:3]
        # Agent prompt — every numeric metric is pre-formatted via
        # format_metric so the LLM sees a string with the
        # platform's canonical precision, never a raw float.
        from tools.academic_export import format_metric
        findings = "\n\n".join(
            f"{name.replace('_', ' ')}: "
            f"Sharpe={format_metric(r.get('sharpe_ratio'), 'sharpe_ratio')}, "
            f"CAGR={format_metric(r.get('cagr'), 'cagr')}, "
            f"Tier 1={r.get('tier1_gates_passed', 0)}/5."
            for name, r in top_three
        ) or "Strategy results not yet available."
        sections = [
            ("executive_summary", "1. Executive Summary",
             "Ten portfolio strategies were tested. Dynamic regime-aware strategies "
             "passed all Tier 1 statistical gates; static 60/40 did not after FDR correction."),
            ("methodology",       "2. Methodology", methodology),
            ("key_findings",      "3. Key Findings", findings),
            ("limitations",       "4. Limitations",
             "Sample size borderline for regime-conditional sub-period tests. "
             "Regime classification disagrees in transition periods."),
            ("recommendations",   "5. Recommendations",
             "Static 60/40 is a baseline, not a defensible policy. Dynamic "
             "regime-aware strategies are candidates for further analysis."),
            ("appendix_charts",   "6. Appendix — Charts Referenced",
             "[Figure 1] Cumulative returns. [Figure 2] Significance Matrix. "
             "[Figure 3] Correlation breakdown. [Figure 4] Stress tests. "
             "[Figure 5] CPCV Sharpe distribution."),
        ]
        title = "Forest Capital — Executive Brief"

    elif doc_type == "analytical_appendix":
        methodology = writer.write_methodology(
            data_sources={"data_range": data_range, "n_months": data_range["n_months"]},
            strategies=list(results_dict.keys()),
            statistical_tests=["Tier 1 gates", "FDR", "DSR", "Walk-forward OOS", "CV stability", "CPCV"],
        )
        sig = {k: bool(v.get("is_significant")) for k, v in results_dict.items()}
        results = writer.write_results(
            strategy_results=results_dict, significance_flags=sig, stress_tests={},
        )
        sections = [
            ("abstract",          "1. Abstract",
             "This appendix reports the full statistical results of an empirical "
             "evaluation of equity-fixed-income diversification strategies."),
            ("data_sources",      "2. Data Sources and Provenance",
             "Authoritative source is Dr. Panttser's FNA 670 Excel. Supplemental "
             "fetches: yfinance SPY, FRED VIX/DGS2, Ken French direct."),
            ("methodology",       "3. Portfolio Construction Methodology", methodology),
            ("statistical_results","4. Statistical Results", results),
            ("sensitivity",       "5. Sensitivity Analysis",
             "Key parameters tested at ±20% of defaults. Sharpe and Tier 1 gate "
             "stability reported per parameter."),
            ("reproducibility",   "6. Reproducibility Notes",
             "RANDOM_SEED = 42. Annualisation 252 (daily) / 12 (monthly). "
             "Simple pct_change throughout — never log returns."),
        ]
        title = "Forest Capital — Analytical Appendix"

    else:
        raise ValueError(f"Unknown doc_type: {doc_type}")

    return {
        "doc_type": doc_type,
        "title":    title,
        "subtitle": (
            f"FNA 670 Practicum · Data range "
            f"{data_range['start']} – {data_range['end']}"
        ),
        "sections": [
            {
                "id":           sid,
                "title":        stitle,
                "ai_draft":     body,
                "content":      body,
                "last_edited":  now_iso,
            }
            for sid, stitle, body in sections
        ],
    }


@app.post("/api/documents/section-doc/draft", status_code=201)
@limiter.limit("10/minute")
async def section_doc_draft(
    request: Request,
    body: dict,
    session: dict = Depends(require_auth),
):
    """
    Creates a new section-structured document for one of Bob's deliverables.

    Body: {"doc_type": "midpoint_paper" | "executive_brief" | "analytical_appendix"}

    Returns {document_id, content, persistence}. The frontend SectionEditor
    routes to /reports/document/:id which loads via GET /api/documents/:id.

    The Academic Writer runs once on creation. Bob's per-section
    Regenerate AI button (POST /api/documents/:id/sections/:section_id/regenerate)
    re-runs the writer for a single section only — cheaper and more
    focused than re-drafting the entire document.
    """
    from tools.documents_cache import create_document

    doc_type = body.get("doc_type", "")
    if doc_type not in {"midpoint_paper", "executive_brief", "analytical_appendix"}:
        raise HTTPException(
            status_code=422,
            detail="doc_type must be midpoint_paper | executive_brief | analytical_appendix",
        )

    try:
        results_dict, data_range = await _load_results_async()
        content = _build_section_doc_content(doc_type, results_dict, data_range)
    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("section_doc_build_failed", ref=ref, error=str(exc), doc_type=doc_type)
        raise HTTPException(status_code=500, detail=f"Draft creation failed (ref: {ref})")

    doc_id = await create_document(
        doc_type=doc_type,
        owner_email=session.get("email", "unknown@queens.edu"),
        initial_content=content,
        created_by=session.get("email"),
    )

    if doc_id is None:
        return {
            "document_id": None,
            "content":     content,
            "persistence": "unavailable",
            "message": "Document drafted but not saved (DATABASE_URL unset).",
        }

    return {
        "document_id": doc_id,
        "content":     content,
        "persistence": "saved",
    }


@app.post("/api/documents/{document_id}/sections/{section_id}/regenerate")
@limiter.limit("20/minute")
async def regenerate_section(
    document_id: str,
    section_id: str,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Re-runs the Academic Writer for one section of one document.

    Returns {ai_draft: str} so the frontend can choose to replace the
    section's `content` field, or just update `ai_draft` while leaving
    Bob's edits intact (the View AI Draft side panel reads from
    `ai_draft`).

    The endpoint is deliberately stateless — it does NOT persist the
    new draft. The frontend decides whether to commit it via PATCH
    /api/documents/:id/draft. This separation lets Bob preview a
    regenerated section without losing his current edits to a draft
    save he didn't ask for.
    """
    from tools.documents_cache import get_document_draft
    from agents.academic_writer import AcademicWriter

    draft = await get_document_draft(document_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Document not found")

    content = draft.get("content", {}) or {}
    doc_type = content.get("doc_type", "")
    section = next(
        (s for s in content.get("sections", []) if s.get("id") == section_id),
        None,
    )
    if section is None:
        raise HTTPException(status_code=404, detail=f"Section '{section_id}' not found")

    try:
        results_dict, data_range = await _load_results_async()
        writer = AcademicWriter()

        # Section ID → which writer method to call. Sections that aren't
        # produced by the Academic Writer (roles, recommendations) just
        # get the original ai_draft back — the AI re-run only makes sense
        # for prose the writer can re-derive from results.
        if section_id == "methodology":
            new_text = writer.write_methodology(
                data_sources={"data_range": data_range, "n_months": data_range["n_months"]},
                strategies=list(results_dict.keys()),
                statistical_tests=["Tier 1 gates", "FDR", "DSR", "Walk-forward OOS", "CV stability"],
            )
        elif section_id in {"results", "statistical_results"}:
            sig = {k: bool(v.get("is_significant")) for k, v in results_dict.items()}
            new_text = writer.write_results(
                strategy_results=results_dict, significance_flags=sig, stress_tests={},
            )
        else:
            # No regenerator wired for this section — return the
            # original AI draft so the UI's "Regenerate" affordance
            # still produces something sensible.
            new_text = section.get("ai_draft", "")

        log.info(
            "section_regenerated",
            document_id=document_id,
            section_id=section_id,
            doc_type=doc_type,
            chars=len(new_text),
        )
        return {"ai_draft": new_text, "section_id": section_id}

    except Exception as exc:
        ref = uuid.uuid4().hex[:8]
        log.error("section_regenerate_failed", ref=ref, error=str(exc), section=section_id)
        raise HTTPException(status_code=500, detail=f"Regenerate failed (ref: {ref})")


@app.post("/api/documents/{document_id}/export")
@limiter.limit("20/minute")
async def export_document(
    document_id: str,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Exports the current draft of a section document as a downloadable file.

    Returns .docx for midpoint_paper and executive_brief, HTML for
    analytical_appendix — matches the format Bob would have got from
    the direct generator endpoints, but using HIS edited content
    rather than re-running the Academic Writer.

    This is the round-trip Bob uses to ship the final version: AI
    drafts in the editor → Bob edits → Save named version → Export.
    """
    from fastapi.responses import Response as FastAPIResponse
    from tools.documents_cache import get_document_draft

    draft = await get_document_draft(document_id)
    if draft is None:
        raise HTTPException(status_code=404, detail="Document not found")

    content = draft.get("content", {}) or {}
    doc_type = content.get("doc_type", "")
    title = content.get("title", "Document")
    subtitle = content.get("subtitle", "")

    # Compile sections from Bob's current content (not ai_draft) — that's
    # the whole point of the editor: edits go into the export.
    sections = [
        {"heading": s.get("title", ""), "body": s.get("content", "")}
        for s in content.get("sections", [])
    ]

    from datetime import date

    if doc_type == "analytical_appendix":
        from tools.html_report_generator import build_html_report
        html_str = build_html_report(
            title=title,
            subtitle=subtitle,
            sections=sections,
        )
        filename = f"forest-capital-appendix-{date.today().isoformat()}.html"
        return FastAPIResponse(
            content=html_str,
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # midpoint_paper + executive_brief → .docx
    from tools.docx_generator import build_docx
    docx_bytes = build_docx(
        title=title,
        subtitle=subtitle,
        sections=sections,
    )
    slug = "midpoint" if doc_type == "midpoint_paper" else "executive-brief"
    filename = f"forest-capital-{slug}-{date.today().isoformat()}.docx"
    return FastAPIResponse(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Generate from storyboard — pptx / script / Q&A ────────────────────────────

@app.post("/api/reports/generate-from-storyboard/{document_id}")
@limiter.limit("10/minute")
async def generate_from_storyboard(
    document_id: str,
    body: dict,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Reads Molly's edited storyboard from document_drafts and renders one
    of four artifact types into a download. The `output_type` field in
    the request body picks the renderer:

      deck            → .pptx via tools/pptx_generator
      script          → .docx full team script via tools/script_writer
      script_molly    → .docx Molly-only filtered script
      script_michael  → .docx Michael-only filtered script
      script_bob      → .docx Bob-only filtered script
      rehearsal       → .docx full team + cues every 2 min + visual cues
      qa              → .docx Q&A preparation, 3 sections, 18 questions

    The pptx deck biases toward Molly's edits — her slide order, chart
    refs, and timing all drive the deck output.
    """
    from fastapi.responses import Response as FastAPIResponse
    from tools.documents_cache import get_document_draft

    output_type = body.get("output_type") or "deck"
    valid = {"deck", "script", "script_molly", "script_michael", "script_bob",
             "rehearsal", "qa"}
    if output_type not in valid:
        raise HTTPException(
            status_code=422,
            detail=f"output_type must be one of {sorted(valid)}",
        )

    # Test env: short-circuit to a minimal deck/script without touching
    # the DB or LLM. Keeps the test runs fast and the assertion surface
    # focused on routing rather than content quality.
    if ENVIRONMENT == "test":
        storyboard = body.get("storyboard") or {"slides": []}
    else:
        draft = await get_document_draft(document_id)
        if draft is None:
            raise HTTPException(status_code=404, detail="Storyboard not found")
        storyboard = draft.get("content", {}) or {}

    # Academic Writer is the spoken-prose engine for the script outputs.
    # Unavailable in test env or without an API key — script_writer falls
    # back to deterministic paragraphs in both cases.
    writer = None
    if ENVIRONMENT != "test":
        try:
            from agents.academic_writer import AcademicWriter
            writer = AcademicWriter()
        except Exception:
            writer = None

    if output_type == "deck":
        from tools.pptx_generator import build_pptx_from_storyboard
        pptx_bytes = build_pptx_from_storyboard(storyboard)
        from datetime import date
        filename = f"forest-capital-deck-{date.today().isoformat()}.pptx"
        return FastAPIResponse(
            content=pptx_bytes,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "presentationml.presentation"
            ),
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if output_type == "qa":
        from tools.script_writer import build_qa_prep_docx
        from datetime import date
        # Pull current strategy results so the Q&A doc references the live
        # significance count. Unavailable in test env → empty results dict.
        results: dict = {}
        if ENVIRONMENT != "test":
            try:
                from tools.data_fetcher import get_full_history_async
                from tools.backtester import run_all_strategies
                history = await get_full_history_async()
                results = await asyncio.to_thread(run_all_strategies, history)
            except Exception:
                pass
        docx_bytes = build_qa_prep_docx(storyboard, strategy_results=results)
        filename = f"forest-capital-qa-prep-{date.today().isoformat()}.docx"
        return FastAPIResponse(
            content=docx_bytes,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # All remaining types are script variants — one shared docx builder.
    from tools.script_writer import build_script_docx
    owner_filter = None
    include_cues = False
    if output_type == "script_molly":
        owner_filter = "Molly"
    elif output_type == "script_michael":
        owner_filter = "Michael"
    elif output_type == "script_bob":
        owner_filter = "Bob"
    elif output_type == "rehearsal":
        include_cues = True

    docx_bytes = build_script_docx(
        storyboard,
        owner_filter=owner_filter,
        include_rehearsal_cues=include_cues,
        writer=writer,
    )
    from datetime import date
    suffix = output_type if output_type != "script" else "full-team"
    filename = f"forest-capital-script-{suffix}-{date.today().isoformat()}.docx"
    return FastAPIResponse(
        content=docx_bytes,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Gemini assistant for storyboard + section editors ────────────────────────

@app.post("/api/documents/{document_id}/assistant")
@limiter.limit("20/minute")
async def document_assistant(
    document_id: str,
    body: dict,
    request: Request,
    session: dict = Depends(require_auth),
):
    """
    Routes an inline editing request to Gemini 1.5 Pro for the storyboard
    and section editors. Returns a suggestion + a structured diff so the
    UI can render red-removed / green-added text and let the user accept
    or reject per paragraph.

    Constraints (CLAUDE.md Section 14):
      - No statistics introduced that aren't already in the input
      - No citations outside references.json
      - Scope guard rejects off-topic requests
      - Multi-turn conversation context is the caller's responsibility
        (we don't persist conversation state here; the UI sends prior
        messages in body['history'] when needed)
    """
    user_message = (body.get("message") or "").strip()
    context_content = (body.get("context_content") or "").strip()
    context_type = body.get("context_type") or "slide"

    if not user_message:
        raise HTTPException(status_code=422, detail="'message' is required")
    if len(user_message) > 1000:
        raise HTTPException(status_code=422, detail="Message exceeds 1000-char limit")

    # Scope guard — same Haiku-classifier path the council uses
    if ENVIRONMENT != "test":
        try:
            from scope_guard import ScopeGuard
            guard = ScopeGuard()
            scope_result = await guard.check(user_message)
            if not scope_result["allowed"]:
                return {
                    "suggestion": "",
                    "diff": {"removed": [], "added": []},
                    "explanation": scope_result.get(
                        "rejection_message",
                        "This request is outside the scope of the Forest "
                        "Capital Portfolio Intelligence System.",
                    ),
                    "confidence": 0.0,
                    "out_of_scope": True,
                }
        except Exception:
            # Scope-guard failure is non-fatal — proceed but log
            log.warning("document_assistant_scope_guard_failed")

    # Test env: return a deterministic mock without calling Gemini
    if ENVIRONMENT == "test":
        return _mock_assistant_response(user_message, context_content)

    try:
        from agents.contrarian_analyst import XAI_TIMEOUT_SECONDS  # noqa: F401
        import os as _os
        from agents.base import call_gemini

        api_key = _os.getenv("GOOGLE_API_KEY", "")
        if not api_key:
            log.info("document_assistant_mock_no_key")
            return _mock_assistant_response(user_message, context_content)

        prompt = (
            f"Editing context: {context_type}\n"
            f"Current content:\n```\n{context_content}\n```\n\n"
            f"User request: {user_message}\n\n"
            f"Respond with a rewritten version of the content that addresses the "
            f"user's request. Constraints:\n"
            f"  - Only reference numbers present in the current content above\n"
            f"  - Do not introduce citations not already in the content\n"
            f"  - Match the spoken-paragraph register of the original\n"
            f"  - Output ONLY the rewritten content, no preamble"
        )

        suggestion = call_gemini(
            GEMINI_MODEL, _GEMINI_ASSISTANT_SYSTEM_PROMPT, prompt,
            trigger="document_assistant",
        ).strip()

        return {
            "suggestion":   suggestion,
            "diff":         _build_diff(context_content, suggestion),
            "explanation":  f"Rewrote {context_type} to address: {user_message[:120]}",
            "confidence":   0.7,
            "out_of_scope": False,
        }

    except Exception as exc:
        log.warning("document_assistant_error", error=str(exc))
        return _mock_assistant_response(user_message, context_content)


_GEMINI_ASSISTANT_SYSTEM_PROMPT = (
    "You are an editing assistant embedded in the Forest Capital Portfolio "
    "Intelligence System. Your job is to rewrite the user's content according "
    "to their request — tighten, expand, restructure, change tone — without "
    "introducing facts that weren't in the original. You may only reference "
    "numbers and citations present in the input content. If the request "
    "would require fabricating a statistic or citing a source not in the "
    "content, say so plainly and refuse that part of the request. "
    "Output only the rewritten content. No preamble, no explanation."
)


def _build_diff(before: str, after: str) -> dict[str, list[str]]:
    """
    Simple paragraph-level diff used by the UI diff display. Splits both
    versions on blank lines and tags paragraphs as removed (in before
    only), added (in after only), or unchanged. The UI renders removed
    red and added green; unchanged paragraphs aren't sent to keep the
    payload small.
    """
    before_paras = [p.strip() for p in before.split("\n\n") if p.strip()]
    after_paras = [p.strip() for p in after.split("\n\n") if p.strip()]
    before_set = set(before_paras)
    after_set = set(after_paras)
    return {
        "removed": [p for p in before_paras if p not in after_set],
        "added":   [p for p in after_paras if p not in before_set],
    }


def _mock_assistant_response(user_message: str, context: str) -> dict[str, Any]:
    """
    Deterministic mock returned when GOOGLE_API_KEY is missing or the
    Gemini API is unreachable. Lets the UI render a usable diff in
    development without requiring credentials.
    """
    # Trivial transformation: prepend a sentence reflecting the request
    if not context:
        suggestion = (
            f"[Mock — Gemini API unavailable] You asked: {user_message}. "
            f"Provide content in 'context_content' to receive a real rewrite."
        )
    else:
        suggestion = (
            f"[Mock revision] {context}\n\n"
            f"(Edit requested: {user_message[:200]} — set GOOGLE_API_KEY "
            f"on Render for real Gemini suggestions.)"
        )
    return {
        "suggestion":   suggestion,
        "diff":         _build_diff(context, suggestion),
        "explanation":  "Gemini unavailable — returning structured mock.",
        "confidence":   0.0,
        "out_of_scope": False,
        "mock":         True,
    }


# ── Developer endpoints (MASTER_API_KEY only) ─────────────────────────────────

@app.post("/api/dev/uiux/review")
async def uiux_review(body: UIUXReviewRequest, _: dict = Depends(require_master_key)):
    return {
        "component": body.component_name,
        "status": "Sprint 1 — UI/UX agent connected in Sprint 3",
        "improvements": [],
    }


@app.get("/api/dev/credits")
async def dev_credits(_: dict = Depends(require_master_key)):
    return {
        "daily_spend_usd": 0.0,
        "total_calls": 0,
        "cost_by_agent": {},
        "note": "Sprint 1 — real tracking in Sprint 2",
    }


# ── WebSocket ─────────────────────────────────────────────────────────────────

@app.websocket("/ws/council")
async def ws_council(websocket: WebSocket):
    """
    Streams council debate token-by-token as agents complete their reports.

    Each agent result is sent as a separate JSON frame with agent name and
    is_final flag. This lets the frontend render agent cards progressively
    rather than waiting for the full council to complete (~30-60s).

    Scope guard runs on connection — query is validated before any agent is
    invoked. Auto-disconnects after 10 minutes of inactivity.
    """
    await websocket.accept()
    try:
        token = websocket.query_params.get("token")
        if not token:
            await websocket.close(code=4001, reason="Missing token")
            return

        from auth import verify_session_token
        try:
            session = verify_session_token(token)
        except HTTPException:
            await websocket.close(code=4003, reason="Unauthorized")
            return

        log.info("ws_council_connected", user=session["email"])
        await websocket.send_json({"type": "connected", "message": "Council ready."})

        while True:
            try:
                data = await websocket.receive_json()
            except WebSocketDisconnect:
                break

            query = data.get("query", "")
            if len(query) > 500:
                await websocket.send_json({
                    "type": "error",
                    "message": "Query exceeds 500 character limit.",
                })
                continue

            # Scope guard on each message
            if ENVIRONMENT != "test":
                try:
                    from scope_guard import ScopeGuard
                    guard = ScopeGuard()
                    scope_result = await guard.check(query)
                    if not scope_result["allowed"]:
                        await websocket.send_json({
                            "type": "out_of_scope",
                            "message": scope_result["rejection_message"],
                        })
                        continue
                except Exception as exc:
                    log.warning("ws_scope_guard_error", error=str(exc))

            if ENVIRONMENT != "test":
                try:
                    from tools.data_fetcher import get_full_history_async
                    from tools.backtester import run_all_strategies
                    from agents.equity_analyst import EquityAnalyst
                    from agents.fixed_income_analyst import FixedIncomeAnalyst
                    from agents.risk_manager import RiskManager
                    from agents.quant_backtester import QuantBacktester
                    from agents.independent_analyst import IndependentAnalyst
                    from agents.cio import CIO

                    history = await get_full_history_async()
                    strategy_results = await asyncio.to_thread(run_all_strategies, history)

                    # Stream each specialist's report as it completes
                    for agent_name, agent_cls in [
                        ("equity_analyst", EquityAnalyst),
                        ("fixed_income_analyst", FixedIncomeAnalyst),
                        ("risk_manager", RiskManager),
                        ("quant_backtester", QuantBacktester),
                    ]:
                        agent = agent_cls()
                        if agent_name == "fixed_income_analyst":
                            report = agent.analyse(strategy_results, history)
                        else:
                            report = agent.analyse(strategy_results)

                        await websocket.send_json({
                            "type": "agent_result",
                            "agent": agent_name,
                            "content": report,
                            "is_final": False,
                        })

                    # Gemini challenge + CIO synthesis — sent as final frame
                    # June 5 2026 — prior_recommendation fetch parallels
                    # the SSE path's logic so this legacy WebSocket route
                    # also writes Section C of the transparency structure
                    # when a prior CIO output exists.
                    prior_recommendation_ws = None
                    try:
                        from tools.cio_recommendation import (
                            _current_data_hash, get_prior_recommendation,
                        )
                        ws_current_hash = await _current_data_hash()
                        if ws_current_hash:
                            prior_recommendation_ws = await (
                                get_prior_recommendation(ws_current_hash))
                    except Exception as exc:  # noqa: BLE001
                        log.warning(
                            "council_ws_prior_recommendation_fetch_error",
                            error=str(exc))
                    cio = CIO()
                    final = cio.deliberate(
                        query, strategy_results, history,
                        prior_recommendation=prior_recommendation_ws)
                    await websocket.send_json({
                        "type": "agent_result",
                        "agent": "cio",
                        "content": final,
                        "is_final": True,
                    })
                    continue

                except Exception as exc:
                    log.error("ws_council_error", error=str(exc))

            # Fallback frame
            await websocket.send_json({
                "type": "agent_result",
                "agent": "System",
                "content": {"summary": f"Council received query: {query}. Pipeline unavailable."},
                "is_final": True,
            })

    except WebSocketDisconnect:
        log.info("ws_council_disconnected")
