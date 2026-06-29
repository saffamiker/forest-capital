# DEPRECATED -- June 2026
# This agent was used for the FNA 670 midpoint peer review surface
# (POST /api/council/peer-review). That endpoint was retired in
# PR-B (June 2026) and the frontend surface was removed in PR #338.
#
# The defense-prep helpers in this module (build_defense_prep_context_block,
# render_defense_prep_context_block, run_defense_prep_with_harness) are
# still imported by main.py's /api/council/defense-prep endpoint, which
# is NOT in the PR-B retirement scope. The peer-review-specific helpers
# (build_peer_review_*_prompt, run_peer_review_with_harness,
# extract_peer_paper_text, mock_peer_review_verdict) are now unreachable
# through any documented endpoint.
#
# This file is preserved as reference material for the five-section
# critique pattern (sections A-E) which may inform future self-review
# surfaces. Do not import the peer-review-specific helpers in new code.
# Safe to delete the peer-review helpers after July 2026 if no reuse
# planned; the defense-prep helpers remain in active use.
"""agents/peer_review.py — Peer Review Assistant + Thesis Defense Prep.

Item 7 (May 23 2026). Two related features built on the
academic_review infrastructure:

  FEATURE A — Peer Review Assistant
    Bob, Michael, and Molly each must review another team's
    midpoint submission for the June 3 cohort meetup (3-4 minute
    critical review per student + 2-minute Q&A). Upload the other
    team's PDF / DOCX, the agent evaluates it against the four
    FNA 670 midpoint rubric dimensions, and produces a structured
    review script the reviewer can read aloud — observations per
    dimension, course-concept anchors, and 2-3 suggested Q&A
    questions. Tone: critical but professional.

  FEATURE B — Thesis Defense Prep
    Generates a mock panel Q&A sheet against the team's OWN
    most-recently-submitted midpoint draft. Anticipated questions
    across three categories — technical/methodological,
    academic/theoretical, governance/practical — each with a
    suggested response the team can rehearse against before the
    cohort meetup OR the July 1 panel.

Both flows reuse the existing peer fan-out + arbiter pattern from
academic_review — same SSE wire format, same generator-evaluator
harness, same Sonnet / Opus model assignments. The differences are
PURELY in the prompts:
  - Different rubric framing (reviewing another team vs anticipating
    questions about your own submission).
  - Different output structure (review script vs Q&A prep sheet).
  - Different context block (uploaded peer paper vs current draft).

The endpoint surface is in main.py — these helpers just produce
the prompt + context shape the streaming layer expects.
"""
from __future__ import annotations

import io
from typing import Any

from agents.academic_review import (
    chunk_arbiter_text,
    peer_agent_ids,
)


# ── Multi-format text extraction (peer papers) ──────────────────────────────


# Maximum bytes accepted for an uploaded peer paper. Tighter than
# the academic_documents 10MB ceiling because peer papers are
# 3-page midpoint submissions — anything beyond a few hundred KB
# is almost certainly an image-heavy PDF or a malformed upload.
MAX_PEER_PAPER_BYTES: int = 2 * 1024 * 1024  # 2 MB


def extract_peer_paper_text(filename: str, raw: bytes) -> str:
    """Extracts plain text from an uploaded peer-team midpoint
    submission. Routes by extension — PDF / .md / .docx are
    supported; everything else raises ValueError.

    Unlike tools.academic_context.extract_document_text (which is
    PDF-only by design), this extractor handles all three formats
    inline so the Peer Review endpoint can accept whatever Bob's
    cohort produces. The peer paper is NOT persisted to
    academic_documents — it lives in-memory for the duration of
    one review session and is discarded.

    Raises ValueError when the format is unsupported or no text
    can be extracted (e.g. a scanned-image PDF). The endpoint
    catches and returns 422.
    """
    if not raw:
        raise ValueError("Uploaded file is empty.")
    if len(raw) > MAX_PEER_PAPER_BYTES:
        raise ValueError(
            f"Uploaded file exceeds {MAX_PEER_PAPER_BYTES} bytes.")
    name = (filename or "").strip().lower()
    if name.endswith(".pdf"):
        return _extract_pdf_text(raw)
    if name.endswith(".md") or name.endswith(".markdown"):
        return _extract_md_text(raw)
    if name.endswith(".docx"):
        return _extract_docx_text(raw)
    raise ValueError(
        f"Unsupported file type for '{filename}'. Accepted formats: "
        f".pdf, .docx, .md.")


def _extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "PDF support unavailable — pypdf not installed.") from exc
    reader = PdfReader(io.BytesIO(raw))
    text = "\n".join(
        (page.extract_text() or "") for page in reader.pages).strip()
    if not text:
        raise ValueError(
            "No text could be extracted from the PDF — it may be a "
            "scanned image. Upload a text-based PDF.")
    return text


def _extract_md_text(raw: bytes) -> str:
    # Markdown is plain UTF-8; render as-is. Strip leading/trailing
    # whitespace so the peer-agent prompts don't blow up on empty
    # newlines.
    try:
        return raw.decode("utf-8").strip()
    except UnicodeDecodeError as exc:
        raise ValueError(
            "Markdown file is not valid UTF-8.") from exc


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError(
            "DOCX support unavailable — python-docx not installed."
        ) from exc
    doc = Document(io.BytesIO(raw))
    parts: list[str] = []
    # Walk paragraphs IN ORDER. Tables are walked separately —
    # python-docx's iteration is by body element kind, not source
    # order, so a paper with interleaved tables and prose can
    # render slightly out of order. Acceptable for a peer-review
    # input where the content matters more than the layout.
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(
            "No text could be extracted from the .docx — the "
            "document may be empty or image-only.")
    return text


# ── FEATURE A — Peer Review Assistant ───────────────────────────────────────


# The four FNA 670 midpoint rubric dimensions. The peer agents
# evaluate the uploaded paper against EACH dimension and emit
# observations, course-concept anchors, and suggested Q&A.
PEER_RUBRIC_DIMENSIONS: list[str] = [
    "Clarity and rigor of written submission",
    "Evidence of meaningful analytical progress",
    "Quality of preliminary results and interpretation",
    "Clear division of labor",
]


_PEER_REVIEW_PEER_QUESTION = """=== YOUR TASK — PEER REVIEW NOTES ===
You are reviewing ANOTHER team's midpoint submission (3 pages,
double-spaced) for the FNA 670 cohort meetup peer review. Each
student in our team will deliver a 3-4 minute critical review
followed by 2-minute Q&A. These peer notes will feed into the
arbiter's review script generator below.

The submitted document is included in the context block above
(under "Peer team submission"). Evaluate it against the FOUR
midpoint rubric dimensions verbatim:

  1. Clarity and rigor of written submission
  2. Evidence of meaningful analytical progress
  3. Quality of preliminary results and interpretation
  4. Clear division of labor

For each dimension, produce two short paragraphs from your
expert lens:

  OBSERVATION — what is strong, what is missing, what is unclear.
  Cite specific passages from the submission when possible (one
  short quote, not a paragraph). Reference course concepts from
  Markowitz, Carhart, Fama-French, regime-switching, factor
  investing, or the Forest Capital practicum brief where the
  submission's claims connect to (or contradict) them.

  Q&A SUGGESTIONS — 1-2 questions the reviewer could ask the
  presenters to probe their analytical reasoning. Specific to the
  submission, not generic. Format each as a single sentence ending
  in '?'.

TONE — critical but professional. The cohort is graded on the
quality of peer feedback so reviews are expected to find genuine
weaknesses. Avoid: hedging, padding, generic encouragement, or
restating what the paper said without commentary.

OUT OF SCOPE — do NOT compare the submission to Forest Capital's
own work. The reviewer evaluates the submission on its own merits
against the rubric, not against an alternative project. Avoid
references to "our project" or "Forest Capital found that …"."""


_PEER_REVIEW_ARBITER_INSTRUCTIONS = """=== YOUR TASK — REVIEW SCRIPT ===
You are the arbiter. The peer notes above evaluated the uploaded
peer submission against the four FNA 670 rubric dimensions from
multiple expert lenses. Synthesise them into a STRUCTURED REVIEW
SCRIPT a reviewer can deliver aloud in 3-4 minutes.

Output the script in this EXACT markdown format so the UI can
parse it:

**Overall verdict:** <Strong | Developing | Needs Work>
**Estimated delivery time:** <X-Y minutes at 130 wpm>

The verdict line is a single judgement aggregated across the four
dimensions. Use Needs Work if any dimension is materially missing;
Developing if every dimension is at least covered but some are
shallow; Strong if every dimension is well executed.

### 1. Clarity and rigor of written submission
**Dimension rating:** <Strong | Developing | Needs Work>
<2-3 paragraphs of OBSERVATIONS the reviewer reads aloud. Cite
specific passages from the submission when relevant. Reference
course concepts where applicable.>

**Suggested questions for Q&A:**
- <question 1>
- <question 2>

### 2. Evidence of meaningful analytical progress
**Dimension rating:** <Strong | Developing | Needs Work>
<2-3 paragraphs of observations.>

**Suggested questions for Q&A:**
- <question 1>
- <question 2>

### 3. Quality of preliminary results and interpretation
**Dimension rating:** <Strong | Developing | Needs Work>
<2-3 paragraphs of observations.>

**Suggested questions for Q&A:**
- <question 1>
- <question 2>

### 4. Clear division of labor
**Dimension rating:** <Strong | Developing | Needs Work>
<2-3 paragraphs of observations.>

**Suggested questions for Q&A:**
- <question 1>
- <question 2>

### Closing summary
<One paragraph wrap-up the reviewer reads at the end — surfaces
the single most important strength + the single most important
weakness so the presenters walk away with a clear takeaway.>

DELIVERY TIME — target 3-4 minutes at 130 words per minute.
Total prose across the four sections + closing should land around
400-520 words. Note the estimated delivery time at the top so the
reviewer can adjust pacing.

TONE GUARD — critical but professional throughout. The peer
review is part of the graded deliverable; generic encouragement
("nice work!") reads as low-effort feedback. Specific
observations grounded in the submission's text are what graders
look for."""


def build_peer_review_peer_prompt(context_block: str) -> str:
    """The user-message body sent to every peer agent for the
    Peer Review Assistant. The context_block is prepended by the
    caller (includes the peer team's submitted text)."""
    return context_block + "\n\n" + _PEER_REVIEW_PEER_QUESTION


def build_peer_review_arbiter_prompt(
    context_block: str, peer_responses: dict[str, str],
) -> str:
    """The arbiter prompt. Stitches the context, the peer notes,
    and the review-script generation instructions."""
    notes = []
    for agent_id, text in peer_responses.items():
        notes.append(f"## Peer agent — {agent_id}\n{text.strip()}")
    peer_block = "\n\n".join(notes) if notes else "(no peer notes)"
    return (
        context_block
        + "\n\n=== PEER REVIEW NOTES ===\n\n"
        + peer_block
        + "\n\n"
        + _PEER_REVIEW_ARBITER_INSTRUCTIONS
    )


# ── FEATURE B — Thesis Defense Prep ─────────────────────────────────────────


# The three Q&A categories the panel-prep sheet is organised
# around. Named verbatim per the May 23 2026 spec.
DEFENSE_CATEGORIES: list[str] = [
    "Technical / methodological",
    "Academic / theoretical",
    "Governance / practical",
]


_DEFENSE_PREP_PEER_QUESTION = """=== YOUR TASK — DEFENSE PREP NOTES ===
The submitted midpoint draft is in the context block above (under
"Team submission"). The team is preparing to present it at the
cohort meetup AND defend it at the July 1 panel. Your job is to
anticipate the questions an MSFA panel — Dr. Panttser, the Forest
Capital partners, and the broader graduate cohort — would ask the
team.

From your expert lens, propose 3-5 anticipated questions across
THE THREE CATEGORIES BELOW. For each question, also propose a
2-3 sentence response the team could rehearse against.

  TECHNICAL / METHODOLOGICAL
    Questions about the data, the backtest, the statistical
    framework, the cross-validation, the regime detection, the
    factor model. Examples of the SHAPE: "Why p < 0.005 and not
    p < 0.05?", "How did you handle the 2007 BND inception gap?",
    "What is the OOS window?", "How does CPCV differ from k-fold
    and why does it matter here?".

  ACADEMIC / THEORETICAL
    Questions linking the work to the published literature —
    Markowitz mean-variance, the Fama-French / Carhart factor
    model, regime-switching literature, the diversification
    decay hypothesis, the post-2022 correlation regime, the
    Deflated Sharpe Ratio framework. Examples of the SHAPE:
    "How does this extend Markowitz?", "Where does the 2022
    break sit in the regime-switching literature?", "Why use
    the Carhart 4-factor over Fama-French 3?".

  GOVERNANCE / PRACTICAL
    Questions about fiduciary responsibility, capital planning,
    implementation, transaction costs, real-world constraints,
    and what a Forest Capital partner would actually do with
    the recommendation. Examples of the SHAPE: "What happens
    when this strategy is deployed at $X AUM?", "How does this
    survive a flight-to-quality event?", "What does the
    capital plan look like under the recommended allocation?",
    "Have you stress-tested the recommendation against the
    Forest Capital mandate?".

For each anticipated question:
  QUESTION — the panel's question, phrased as the panel would
  ask it. Specific to THIS submission's findings, not generic.
  SUGGESTED RESPONSE — 2-3 sentences the team can use as the
  spine of their answer. Ground the response in the
  submission's actual numbers and findings (use the context
  block above). Flag honestly when the response should
  acknowledge a limitation rather than defend.

These peer notes will feed into the arbiter's Q&A prep sheet
generator below. Be SPECIFIC — generic questions ("how confident
are you?") are not useful. The shape "Given that you found X, why
not Y?" is more useful than "Why X?"."""


_DEFENSE_PREP_ARBITER_INSTRUCTIONS = """=== YOUR TASK — Q&A PREP SHEET ===
You are the arbiter. The peer notes above proposed anticipated
panel questions across three categories from multiple expert
lenses. Synthesise into a STRUCTURED Q&A PREP SHEET the team can
rehearse against.

Output in this EXACT markdown format so the UI can parse it:

**Mock panel — overall readiness:** <Strong | Developing | Needs Work>

The readiness line aggregates how well-prepared the submission
is to defend itself across the three categories. Strong when
every category is well covered; Developing when one category has
weak spots; Needs Work when the submission has material gaps a
panel would expose.

### 1. Technical / methodological
**Category readiness:** <Strong | Developing | Needs Work>

For each anticipated question in this category, render:

**Q:** <the panel's question, specific to this submission>
**Suggested response:** <2-3 sentences grounded in the
submission's findings. Flag limitations honestly where
warranted; do not invent strengths.>

Include 4-6 questions in this section.

### 2. Academic / theoretical
**Category readiness:** <Strong | Developing | Needs Work>

[Same Q/A format. Include 3-5 questions.]

### 3. Governance / practical
**Category readiness:** <Strong | Developing | Needs Work>

[Same Q/A format. Include 3-5 questions.]

### Rehearsal recommendations
<One paragraph naming the THREE highest-risk questions across
the three categories and which team member should rehearse each
one based on the submission's division of labour.>

TONE — mock-panel adversarial. The panel's questions are not
softballs; the prep sheet should sound like a sceptical reviewer
probing the submission's weakest points. Suggested responses
should be honest about limitations rather than defensive — the
graders reward honest acknowledgment of constraints over
performative confidence.

GROUND EVERY QUESTION IN THE SUBMISSION — the team's submission
is in the context block above. Specific questions about the
submission's actual numbers (the 2022 correlation shift, the
FDR result, the regime-switching strategy's Sharpe, the BND
splice) are more useful than generic methodology questions."""


def build_defense_prep_peer_prompt(context_block: str) -> str:
    """User-message body for every peer agent in the Defense
    Prep flow. context_block prepended by the caller (includes
    the team's own submitted draft text)."""
    return context_block + "\n\n" + _DEFENSE_PREP_PEER_QUESTION


def build_defense_prep_arbiter_prompt(
    context_block: str, peer_responses: dict[str, str],
) -> str:
    notes = []
    for agent_id, text in peer_responses.items():
        notes.append(f"## Peer agent — {agent_id}\n{text.strip()}")
    peer_block = "\n\n".join(notes) if notes else "(no peer notes)"
    return (
        context_block
        + "\n\n=== ANTICIPATED QUESTION NOTES ===\n\n"
        + peer_block
        + "\n\n"
        + _DEFENSE_PREP_ARBITER_INSTRUCTIONS
    )


# ── Orchestration — single-arbiter generator-evaluator runs ────────────────
#
# Both features are focused-output flows (a review script for one,
# a Q&A prep sheet for the other). A single thoughtful Opus arbiter
# call wrapped in the GeneratorEvaluatorHarness gives the quality
# the user asked for without paying for the 7-agent fan-out
# academic_review uses for its multi-rubric breadth. If a later
# iteration wants the orthogonal-perspective fan-out, the peer
# prompts above are ready — academic_review.run_peer_fan_out can
# be wired in by passing each peer a context_block already loaded
# with the task framing.


# Model the arbiter uses. Imported lazily so test environments
# without an Anthropic key never hit the model constant.
def _opus_model() -> str:
    from agents.base import OPUS_MODEL
    return OPUS_MODEL


def _sonnet_model() -> str:
    from agents.base import SONNET_MODEL
    return SONNET_MODEL


def _is_test_env() -> bool:
    import os
    return os.getenv("ENVIRONMENT", "").lower() == "test"


def _evaluator_prompt() -> str:
    """Evaluator system prompt for the harness's quality gate.
    Scores the generated output against five focused criteria so
    a draft that misses a rubric section or trails off mid-answer
    gets regenerated rather than streamed."""
    return (
        "You score peer-review and defense-prep outputs on five "
        "dimensions, each 1-10. Return ONLY a JSON object with "
        "this exact shape:\n"
        '{"scores":{"rubric_coverage":N,"specificity":N,'
        '"actionability":N,"tone":N,"completeness":N},'
        '"overall":N.NN,"passed":bool,"feedback":"..."}\n\n'
        "Criteria:\n"
        "  rubric_coverage — every required section / dimension "
        "is present and rated.\n"
        "  specificity — observations and questions are grounded "
        "in the submission's actual text, not generic.\n"
        "  actionability — suggestions / questions are concrete "
        "enough to act on.\n"
        "  tone — critical-but-professional; flag hedging, "
        "padding, or performative confidence.\n"
        "  completeness — output is complete; no mid-section "
        "truncation.\n\n"
        "passed = overall >= 7.0. Return ONLY the JSON."
    )


def _peer_review_system_prompt() -> str:
    return (
        "You are a senior FNA 670 reviewer evaluating ANOTHER "
        "team's midpoint submission for the cohort meetup peer "
        "review. Produce a critical-but-professional review "
        "script the reviewer can deliver in 3-4 minutes. Ground "
        "every observation in the submitted text — generic "
        "encouragement is not useful."
    )


# Full per-term primer text — kept out of the initial-generation system
# prompt to reduce per-call input tokens (and overall generation latency,
# the proximate cause of the May 30 Defense Prep timeout on Render). The
# initial system prompt lists the term INDEX only; the full text below is
# available to a future follow-up Q&A endpoint that detects a primer
# question and injects only the relevant per-term block.
_DEFENSE_PREP_FULL_PRIMERS = (
    "TECHNICAL PRIMERS — FULL DEFINITIONS.\n"
    "Each primer follows the same shape: what it measures or does, why "
    "it matters for this study, and an honest limitation. A primer never "
    "defines a term using the term itself.\n\n"

    "SHARPE RATIO. Measures average excess return per unit of total "
    "volatility — return per unit of risk taken. Why for this study: "
    "the headline yardstick the panel will compare the regime-conditional "
    "blend against the benchmark on. Limitation: assumes symmetric "
    "returns, so a strategy with frequent small gains and rare large "
    "losses looks better than it deserves.\n\n"

    "CVaR (Conditional Value-at-Risk). Measures the average loss across "
    "the worst-case slice of outcomes (the worst 1% of months at "
    "CVaR-99). Why: captures left-tail (downside) risk that variance and "
    "Sharpe both ignore. Limitation: estimated from a finite sample "
    "(~286 monthly observations here), so the deepest-tail figure "
    "carries real uncertainty.\n\n"

    "CAGR (Compound Annual Growth Rate). The constant annual growth "
    "rate that, compounded over the period, would produce the observed "
    "cumulative return. Why: a single, comparable, compounding-aware "
    "annual headline. Limitation: hides the path — a smooth journey and "
    "a violently volatile one can share the same CAGR.\n\n"

    "MAXIMUM DRAWDOWN. The largest peak-to-trough decline an investor "
    "in the strategy would have actually lived through. Why: the loss "
    "number that drives behavioural risk — whether the holder sticks "
    "with the strategy through the trough. Limitation: a single "
    "realised history, not a forecast of the next worst case.\n\n"

    "CORRELATION. Measures how two return series move together on a -1 "
    "to +1 scale. Why: the equity-bond figure shifted from around -0.05 "
    "pre-2022 to +0.57 post-2022, which is the central finding driving "
    "the rest of the study. Limitation: a linear, average measure — can "
    "miss state-dependent or tail-dependent comovement that only "
    "appears in stress.\n\n"

    "COVARIANCE. The joint variability of two return series in their "
    "native units (the building block of the variance of a weighted "
    "portfolio). Why: feeds the optimiser directly — Min Variance and "
    "Max Sharpe Rolling allocate against the rolling 36-month "
    "covariance matrix. Limitation: estimated covariance is noisy, "
    "especially for short windows.\n\n"

    "EFFICIENT FRONTIER. The set of portfolios that achieve the highest "
    "expected return for each level of risk (or the lowest risk for "
    "each expected return). Why: visualises whether the "
    "regime-conditional blend is doing useful work or being dominated "
    "by simpler static mixes. Limitation: a single-period mean-variance "
    "snapshot — no path awareness and no regime conditioning of its "
    "own.\n\n"

    "BETA and FACTOR EXPOSURE. Beta measures how sensitively a "
    "strategy's returns move with a benchmark; factor exposure "
    "generalises this to named risk drivers (market, size, value, "
    "momentum). Why: separates active skill from passive loadings on "
    "common factors. Limitation: regression coefficients drift through "
    "time, so a single beta or factor loading is an average that can "
    "mask regime-dependent behaviour.\n\n"

    "MOMENTUM STRATEGY. Rotates into assets that have outperformed over "
    "a recent lookback window, on the premise that short-horizon trends "
    "persist. Why: a long-documented anomaly and one of the dynamic "
    "strategies tested here. Limitation: vulnerable to sharp reversals "
    "when a trend breaks — the OOS window contains such an event.\n\n"

    "VOLATILITY TARGETING. Scales portfolio exposure up or down each "
    "month so realised volatility hits a fixed target. Why: delivers a "
    "more stable risk experience to the holder across calm and "
    "turbulent regimes. Limitation: necessarily lags realised "
    "volatility shifts — it tends to de-risk after the move it was "
    "meant to dampen has already happened.\n\n"

    "RISK PARITY. Allocates so each asset contributes EQUALLY to total "
    "portfolio risk — not equal dollar weight. Why: corrects 60/40's "
    "equity-risk dominance without leverage; bonds carry more weight "
    "than they would on a market-cap basis. Limitation: leans on the "
    "historical risk profile, which moved sharply in 2022 and changed "
    "the implied allocation.\n\n"

    "REGIME SWITCHING. Allocates differently in different market states "
    "(BULL, BEAR, TRANSITION) using a model to infer the current state "
    "from observable returns. Why: lets the portfolio adapt to "
    "structural breaks rather than holding one allocation through every "
    "environment. Limitation: only as good as the regime-inference "
    "model — a misclassified turning point costs real money.\n\n"

    "REBALANCING. Periodically returns the portfolio to its target "
    "weights as price moves drift it away. Why: enforces the strategy "
    "— without rebalancing, the winning asset slowly takes over and the "
    "design dissolves. Limitation: each rebalance carries a transaction "
    "cost (tested at 10/15/20 bps per event in the sensitivity "
    "analysis).\n\n"

    "BACKTESTING vs OUT-OF-SAMPLE TESTING. Backtesting runs the "
    "strategy on historical data; out-of-sample testing runs it on a "
    "separate period that was NOT used to design the strategy. Why: "
    "in-sample results can be overfit — the OOS window is the honest "
    "test of whether the rule generalises. Limitation: even an OOS "
    "result is one realised history, not a guarantee about the next "
    "period.\n\n"

    "STATISTICAL SIGNIFICANCE. A measure of how unlikely the observed "
    "effect would be if the strategy had no real edge and the result "
    "came from chance alone. Why: separates a genuine outperformance "
    "from a lucky draw. Limitation: depends heavily on sample size — "
    "with this study's window, the outperformance is real in magnitude "
    "but our sample is too small to prove it statistically, which we "
    "disclose explicitly.\n\n"

    "BLACK-LITTERMAN MODEL. Combines the market-implied equilibrium "
    "returns with the user's own views to produce more stable, "
    "optimisation-ready return estimates. Why: addresses mean-variance "
    "optimisation's extreme sensitivity to noisy mean estimates; one of "
    "the static blended strategies tested here. Limitation: the views "
    "themselves are subjective — without strong views the model "
    "collapses back toward market-cap weights."
)


def _defense_prep_system_prompt() -> str:
    return (
        # ── ROLE ────────────────────────────────────────────────────────
        "You are a mock panel of senior FNA 670 reviewers and Forest "
        "Capital partners stress-testing the team's draft (the document "
        "supplied in the user message) before the cohort meetup and the "
        "July 1 panel. Generate the sharpest realistic anticipated "
        "questions across three categories — technical, academic, "
        "governance — with rehearsable responses. Honest about "
        "limitations beats performative defence.\n\n"

        # ── TWO QUESTIONER LEVELS ───────────────────────────────────────
        "TWO QUESTIONER LEVELS. Read each question and choose the "
        "register before answering:\n"
        "  - PEER questions come from MSFA cohort-mates. Use the PEER "
        "FRAMING below — accessible language, business analogies, "
        "practical focus.\n"
        "  - PROFESSOR questions (Dr. Panttser's technical follow-ons) "
        "land on the specific topics listed under PROFESSOR FRAMING. "
        "Engage at full technical depth — do not simplify. Give the "
        "specification, the rationale for the choice, and the known "
        "limitation. Professor-level questions deserve "
        "professor-level answers.\n\n"

        # ── PEER FRAMING ────────────────────────────────────────────────
        "PEER FRAMING — audience and language.\n"
        "The peer audience is MSFA (Master of Science in Finance and "
        "Analytics) graduate students. They have a solid business and "
        "finance foundation and basic statistics awareness, but they "
        "are not quants, not economists, and not math majors. Answer "
        "as if explaining to a smart, finance-literate peer who:\n"
        "  - understands Sharpe ratio, drawdown, diversification, and "
        "portfolio theory at a conceptual level;\n"
        "  - does not need (and will not follow) mathematical "
        "derivations;\n"
        "  - responds well to business analogies and real-world "
        "framing;\n"
        "  - will push back on claims that seem too convenient or "
        "oversimplified;\n"
        "  - cares about practical implications for a capital planning "
        "mandate, not academic proofs.\n\n"

        "LANGUAGE RULES for peer answers:\n"
        "  - no mathematical notation;\n"
        "  - no p-value explanations beyond 'the outperformance is real "
        "but our sample is too small to prove it statistically -- we "
        "disclose that';\n"
        "  - replace 'CVaR 99% annualized' with 'in the worst "
        "scenarios, here is how much the portfolio loses';\n"
        "  - replace 'HMM posterior probability' with 'the model's "
        "confidence that we are in a particular market regime';\n"
        "  - replace 'factor exposure' with 'how sensitive the strategy "
        "is to broad market moves';\n"
        "  - explain the 2022 correlation break as: 'bonds stopped "
        "cushioning equity losses -- they started moving in the same "
        "direction'.\n\n"

        "TONE for peer answers:\n"
        "  - collegial, not professorial;\n"
        "  - direct answers before caveats;\n"
        "  - acknowledge limitations honestly without being defensive;\n"
        "  - never make the questioner feel like they asked a naive "
        "question.\n\n"

        # ── PROFESSOR FRAMING ───────────────────────────────────────────
        "PROFESSOR FRAMING — full technical depth on these topics.\n"
        "When a question touches any of the topics below, give the "
        "SPECIFICATION, the RATIONALE for the choice, and the KNOWN "
        "LIMITATION — in that order. Do not simplify the terminology.\n\n"

        "HMM (Hidden Markov Model).\n"
        "  WHAT IT IS: a statistical model that assumes the market "
        "moves through hidden states (BULL, BEAR, TRANSITION) that "
        "cannot be observed directly but can be inferred from "
        "observable returns. Each state has its own return distribution "
        "(mean and variance).\n"
        "  METHOD: GaussianHMM with 3 states, fitted by the Baum-Welch "
        "EM algorithm on daily equity returns. 500 iterations, "
        "tolerance 1e-5.\n"
        "  WHY 3 STATES: bull, bear, and a transition state that "
        "captures the regime-switching process itself.\n"
        "  KNOWN LIMITATION: the EM algorithm exhibits a non-monotonic "
        "step at the log-likelihood plateau on this dataset — a 0.01% "
        "wobble that does not affect posterior quality. Disclosed "
        "explicitly.\n\n"

        "CORRELATION.\n"
        "  WHAT WE MEASURE: rolling pairwise Pearson correlation "
        "between monthly equity and investment-grade bond returns. The "
        "central finding is the shift from -0.05 pre-2022 to +0.57 "
        "post-2022. Pearson correlation assumes linearity — "
        "appropriate for monthly return series at this horizon.\n"
        "  WHY IT MATTERS: negative correlation was the statistical "
        "foundation of 60/40 diversification. Its disappearance means "
        "the hedge is gone.\n\n"

        "COVARIANCE.\n"
        "  HOW WE USE IT: rolling 36-month covariance matrix for Min "
        "Variance and Max Sharpe Rolling strategies. The 36-month "
        "window balances responsiveness to regime changes against "
        "estimation noise.\n"
        "  KNOWN ISSUE: with only three assets, the covariance matrix "
        "is 3x3 and well-conditioned — no shrinkage estimator needed. "
        "In a larger universe, regularization would be required.\n\n"

        "SKEWNESS.\n"
        "  WHAT IT MEASURES: asymmetry of the return distribution. "
        "Negative skew (left tail heavier) is the norm for "
        "equity-heavy strategies — rare large losses, more frequent "
        "small gains.\n"
        "  WHY SHARPE MISSES IT: Sharpe assumes symmetric returns. A "
        "strategy with negative skew looks better on Sharpe than it "
        "deserves.\n"
        "  WHAT WE REPORT: CVaR captures the left tail directly. The "
        "skew differences across strategies are visible in the CVaR "
        "rankings.\n\n"

        "KURTOSIS.\n"
        "  WHAT IT MEASURES: the fatness of the tails relative to a "
        "normal distribution. Monthly financial returns have excess "
        "kurtosis — extreme events happen more often than a normal "
        "distribution predicts.\n"
        "  WHY IT MATTERS: Sharpe and standard deviation assume "
        "normality. CVaR via historical simulation does not — it reads "
        "actual tail observations, so fat tails are captured directly. "
        "The 286-observation sample means extreme tail estimates carry "
        "uncertainty, which is why we use 99% CVaR as an approximation "
        "rather than claiming precision.\n\n"

        # ── TECHNICAL PRIMERS — CONDENSED INDEX ────────────────────────
        # Initial-generation context kept lean: the index lists every term
        # the team can defend from first principles, but the full per-term
        # definitions live in _DEFENSE_PREP_FULL_PRIMERS (this module) and
        # are injected only on a follow-up question that explicitly asks
        # for a definition. Saves ~4kB of input on every initial run.
        "TECHNICAL PRIMERS — INDEX.\n"
        "The team is briefed on the following terms and can defend each "
        "from first principles when asked: Sharpe ratio, CVaR, CAGR, "
        "maximum drawdown, correlation, covariance, efficient frontier, "
        "beta and factor exposure, momentum strategy, volatility "
        "targeting, risk parity, regime switching, rebalancing, "
        "backtesting vs out-of-sample testing, statistical significance, "
        "Black-Litterman model. When a Q&A item touches one of these, "
        "anchor in plain English and the PEER FRAMING language rules "
        "above (or PROFESSOR FRAMING for the five named technical "
        "topics). Full definitions are available on demand for a "
        "follow-up question that explicitly asks for one.\n\n"

        # ── RESPONSE BALANCE ────────────────────────────────────────────
        "RESPONSE BALANCE. Answer at TWO levels simultaneously:\n"
        "  1. NARRATIVE — lead with the 'so what' and the reasoning "
        "behind each decision (why these assets, why this framework, "
        "why the 2022 break matters).\n"
        "  2. TECHNICAL DEFENCE — when a question is explicitly "
        "technical (statistical significance, CVaR methodology, "
        "validation approach), provide specific and accurate answers "
        "grounded in the uploaded document.\n\n"

        "Lead with the conceptual argument. Follow with technical "
        "specifics only if the question requires them or if precision "
        "would strengthen the response. The audience is academic peers "
        "who understand finance but may not know the platform's "
        "implementation details. If a question uses technical "
        "language, match that language; if it is conversational, stay "
        "conversational — read the register of the question and "
        "respond in kind.\n\n"

        "DO NOT:\n"
        "  - recite tables of numbers unprompted;\n"
        "  - answer purely technically when the question is "
        "conceptual;\n"
        "  - be vague when a specific methodological challenge is "
        "raised.\n\n"

        # ── THE HUMAN JUDGMENT QUESTION ─────────────────────────────────
        # Reviewers asking "where is the human analytical judgment given
        # you have an AI council and a platform?" expect a prepared,
        # specific answer naming the decisions that required human
        # judgment. The five named below are the load-bearing ones —
        # they are not platform outputs, not council outputs, they are
        # team interpretations and design choices. Answer in that
        # framing every time, in this order, without hedging.
        "THE HUMAN JUDGMENT QUESTION.\n"
        "If asked 'where is the human analytical judgment given you "
        "have an AI council and a platform?' the prepared response is:\n\n"

        "\"The platform executes. The judgment is in what we built "
        "it to do and why.\n\n"

        "Five decisions required human judgment that the AI council "
        "could not have made:\n\n"

        "One: interpreting the 2022 correlation shift as structural "
        "rather than cyclical. The data shows a shift. The "
        "interpretation — that it changes the nature of the "
        "portfolio construction problem — is ours.\n\n"

        "Two: choosing to span ten distinct signal mechanisms rather "
        "than optimise parameters on one approach. We could have run "
        "five hundred variations of momentum. We chose to ask whether "
        "different mechanisms produce different results. That is a "
        "research design decision.\n\n"

        "Three: disclosing that no strategy achieves statistical "
        "significance, when we could have reported Sharpe rankings "
        "and moved on. We chose honesty about the limits of the "
        "evidence.\n\n"

        "Four: designing dissent into the AI council. We could have "
        "asked the AI for a recommendation and taken it. We built a "
        "structure that forces the system to challenge itself before "
        "it advises us.\n\n"

        "Five: the constraint framework. Forty percent maximum per "
        "strategy and fifty percent maximum per asset class are not "
        "data-derived. They are fiduciary design decisions about "
        "what an institutional mandate can implement responsibly.\n\n"

        "The platform gives us the evidence. The interpretation, the "
        "design, and the governance framework are ours.\"\n\n"

        "Deliver this response verbatim when the question is asked in "
        "that form. Adapt the framing if a follow-up zooms in on one "
        "of the five — but always return to 'the platform executes, "
        "the judgment is ours.'"
    )


def run_peer_review_with_harness(context_block: str) -> str:
    """Generates the Peer Review Assistant verdict, harness-gated.
    Returns the accepted markdown. Test env returns a deterministic
    mock so SSE wire-format tests pass without an API key.
    """
    if _is_test_env():
        return mock_peer_review_verdict("(test environment)")
    from agents.base import call_claude
    from agents.harness import GeneratorEvaluatorHarness

    sys_prompt = _peer_review_system_prompt()
    generator_prompt = (
        _PEER_REVIEW_ARBITER_INSTRUCTIONS + "\n\n" + context_block)

    def _generate(prompt: str) -> str:
        # The harness threads feedback into `prompt` directly on
        # retries — we just forward it to the model.
        return call_claude(
            model=_opus_model(),
            system_prompt=sys_prompt,
            user_message=prompt,
            max_tokens=2200,
            trigger="peer_review_assistant",
        )

    harness = GeneratorEvaluatorHarness(
        evaluator_model=_sonnet_model(),
    )
    result = harness.run(
        generator_fn=_generate,
        evaluator_prompt=_evaluator_prompt(),
        generator_prompt=generator_prompt,
        context=context_block,
        agent_id="peer_review_assistant",
    )
    return result.response


def run_defense_prep_with_harness(context_block: str) -> str:
    """Generates the Thesis Defense Prep Q&A sheet, harness-gated.
    Test env returns a deterministic mock."""
    if _is_test_env():
        return mock_defense_prep_verdict("(test environment)")
    from agents.base import call_claude
    from agents.harness import GeneratorEvaluatorHarness

    sys_prompt = _defense_prep_system_prompt()
    generator_prompt = (
        _DEFENSE_PREP_ARBITER_INSTRUCTIONS + "\n\n" + context_block)

    def _generate(prompt: str) -> str:
        return call_claude(
            model=_opus_model(),
            system_prompt=sys_prompt,
            user_message=prompt,
            max_tokens=2400,
            trigger="defense_prep",
        )

    harness = GeneratorEvaluatorHarness(
        evaluator_model=_sonnet_model(),
    )
    result = harness.run(
        generator_fn=_generate,
        evaluator_prompt=_evaluator_prompt(),
        generator_prompt=generator_prompt,
        context=context_block,
        agent_id="thesis_defense_prep",
    )
    return result.response


# ── Shared utilities ───────────────────────────────────────────────────────


# Re-export so callers don't have to reach into academic_review
# for the chunker. Keeps the peer-review module the single import
# the endpoint needs.
__all__ = [
    "PEER_RUBRIC_DIMENSIONS",
    "DEFENSE_CATEGORIES",
    "MAX_PEER_PAPER_BYTES",
    "extract_peer_paper_text",
    "build_peer_review_peer_prompt",
    "build_peer_review_arbiter_prompt",
    "build_peer_review_context_block",
    "render_peer_review_context_block",
    "build_defense_prep_peer_prompt",
    "build_defense_prep_arbiter_prompt",
    "build_defense_prep_context_block",
    "render_defense_prep_context_block",
    "run_peer_review_with_harness",
    "run_defense_prep_with_harness",
    "mock_peer_review_verdict",
    "mock_defense_prep_verdict",
    "chunk_arbiter_text",
    "peer_agent_ids",
]


def mock_peer_review_verdict(submission_name: str) -> str:
    """Test-environment / cold-deploy stand-in for the Peer Review
    Assistant arbiter output. Returns a well-formed (parseable)
    markdown verdict with placeholder content so endpoint tests
    can assert wire-format conformance without an Anthropic key."""
    name = submission_name or "the submission"
    return (
        f"**Overall verdict:** Developing\n"
        f"**Estimated delivery time:** 3-4 minutes at 130 wpm\n\n"
        f"### 1. Clarity and rigor of written submission\n"
        f"**Dimension rating:** Developing\n"
        f"This is a deterministic mock review of {name} for the "
        f"test environment.\n\n"
        f"**Suggested questions for Q&A:**\n"
        f"- Mock question about clarity?\n"
        f"- Mock question about rigour?\n\n"
        f"### 2. Evidence of meaningful analytical progress\n"
        f"**Dimension rating:** Developing\nMock observations.\n\n"
        f"**Suggested questions for Q&A:**\n"
        f"- Mock progress question?\n\n"
        f"### 3. Quality of preliminary results and interpretation\n"
        f"**Dimension rating:** Developing\nMock observations.\n\n"
        f"**Suggested questions for Q&A:**\n"
        f"- Mock results question?\n\n"
        f"### 4. Clear division of labor\n"
        f"**Dimension rating:** Developing\nMock observations.\n\n"
        f"**Suggested questions for Q&A:**\n"
        f"- Mock division-of-labor question?\n\n"
        f"### Closing summary\n"
        f"This is a deterministic test-environment mock; no real "
        f"agent ran.\n"
    )


def mock_defense_prep_verdict(team_name: str) -> str:
    """Test-environment stand-in for the Defense Prep arbiter
    output. Same purpose as mock_peer_review_verdict — a parseable
    markdown sheet so the wire-format tests run without an API
    key."""
    return (
        f"**Mock panel — overall readiness:** Developing\n\n"
        f"### 1. Technical / methodological\n"
        f"**Category readiness:** Developing\n\n"
        f"**Q:** Mock technical question for {team_name}?\n"
        f"**Suggested response:** Mock response anchored to the "
        f"submission's findings.\n\n"
        f"### 2. Academic / theoretical\n"
        f"**Category readiness:** Developing\n\n"
        f"**Q:** Mock academic question?\n"
        f"**Suggested response:** Mock academic response.\n\n"
        f"### 3. Governance / practical\n"
        f"**Category readiness:** Developing\n\n"
        f"**Q:** Mock governance question?\n"
        f"**Suggested response:** Mock governance response.\n\n"
        f"### Rehearsal recommendations\n"
        f"Deterministic test-environment mock — no real agent ran.\n"
    )


# ── Context-block helpers (built up by main.py from the inputs) ────────────


def build_peer_review_context_block(
    submission_name: str,
    submission_text: str,
    rubric_dimensions: list[str] | None = None,
) -> dict[str, Any]:
    """Assembles the context block injected into every peer agent
    prompt for the Peer Review Assistant. Returns a dict the
    endpoint serialises into the prompt; tests assert on the
    presence of each field."""
    return {
        "submission_name":   submission_name,
        "submission_text":   submission_text,
        "rubric_dimensions": rubric_dimensions or PEER_RUBRIC_DIMENSIONS,
    }


def render_peer_review_context_block(
    ctx: dict[str, Any],
) -> str:
    """Renders the peer-review context dict into the leading
    string the peer / arbiter prompts prepend. Distinct from the
    builder so tests can exercise the rendering independently
    of the assembly."""
    lines: list[str] = [
        "=== PEER TEAM SUBMISSION ===",
        f"Name: {ctx.get('submission_name', '(unnamed)')}",
        "",
        "FNA 670 midpoint rubric dimensions evaluated:",
    ]
    for i, dim in enumerate(ctx.get("rubric_dimensions",
                                     PEER_RUBRIC_DIMENSIONS), 1):
        lines.append(f"  {i}. {dim}")
    lines.append("")
    lines.append("--- BEGIN SUBMITTED TEXT ---")
    lines.append(ctx.get("submission_text", "(no text extracted)"))
    lines.append("--- END SUBMITTED TEXT ---")
    return "\n".join(lines)


def build_defense_prep_context_block(
    team_name: str,
    draft_text: str,
    categories: list[str] | None = None,
    source_name: str | None = None,
) -> dict[str, Any]:
    """Context for the Thesis Defense Prep flow. The team's OWN submitted
    document is the primary input; the categories drive the Q&A structure.
    `source_name` is the uploaded filename (when supplied) and is surfaced
    in the labelled block header so the model knows exactly which document
    it's answering from."""
    return {
        "team_name":   team_name,
        "draft_text":  draft_text,
        "categories":  categories or DEFENSE_CATEGORIES,
        "source_name": source_name,
    }


def render_defense_prep_context_block(ctx: dict[str, Any]) -> str:
    """Render the context block with the SUBMITTED DOCUMENT first and
    clearly labelled as the primary source — the model must ground every
    Q&A answer in this text, not in cached project context. The team /
    category framing follows so the Q&A structure is consistent."""
    source = ctx.get("source_name")
    header = (
        f"=== SUBMITTED ACADEMIC DOCUMENT: {source} ==="
        if source else "=== SUBMITTED ACADEMIC DOCUMENT ===")
    lines: list[str] = [
        header,
        "The following is the submitted academic document. Answer all "
        "questions using this as the primary source.",
        "",
        "--- BEGIN SUBMITTED DOCUMENT ---",
        ctx.get("draft_text", "(no draft text available)"),
        "--- END SUBMITTED DOCUMENT ---",
        "",
        "=== TEAM SUBMISSION ===",
        f"Team: {ctx.get('team_name', '(unnamed)')}",
        "",
        "Q&A categories anticipated for the cohort meetup + July 1 panel:",
    ]
    for i, cat in enumerate(ctx.get("categories", DEFENSE_CATEGORIES), 1):
        lines.append(f"  {i}. {cat}")
    return "\n".join(lines)
