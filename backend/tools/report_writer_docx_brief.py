"""tools/report_writer_docx_brief.py — Executive Brief memo formatter.

May 22 2026 (item 12 commit D). Separate from the APA paper
formatter (tools/report_writer_docx.py) because the executive
brief is a fundamentally different document type: memo header
TO/FROM/DATE/RE, Calibri 11pt single-spaced body, ALL CAPS
left-aligned section headings, no title page, no references list,
no first-line indent.

  Two builders:
    build_brief_docx(paper_md, *, references_md=None) → bytes
    build_brief_appendix_docx(context) → bytes
       (Returns the same APA appendix as the midpoint paper. The
        executive brief is the human-facing summary; the appendix
        carries the full data record either way.)

  Dispatch is in tools/report_generator.render_paper_bytes which
  reads format_spec.memo_style from the template row and chooses
  this builder vs the APA paper builder accordingly.
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Inches, Pt


_BODY_FONT = "Calibri"
_BODY_SIZE = Pt(11)


def _new_brief_document() -> Document:
    """Memo-style document — 1-inch margins, Calibri 11pt single-
    spaced body. No header / footer (a 2-page memo doesn't need
    page numbers per the format spec)."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = _BODY_SIZE
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(6)
    return doc


def _add_para(
    doc: Document, text: str, *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    space_after_pt: int = 6,
) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.space_after = Pt(space_after_pt)
    for run_text, b_in, i_in in _split_inline(text):
        if not run_text:
            continue
        run = para.add_run(run_text)
        run.font.name = _BODY_FONT
        run.font.size = _BODY_SIZE
        run.bold = b_in or bold
        run.italic = i_in


def _add_section_heading(doc: Document, text: str) -> None:
    """Memo-style heading: bold, left-aligned, ALL CAPS, body font."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run((text or "").upper())
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = _BODY_SIZE


_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _split_inline(block: str) -> "list[tuple[str, bool, bool]]":
    out: list[tuple[str, bool, bool]] = []
    for part in _INLINE_RE.split(block):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True, False))
        elif part.startswith("`") and part.endswith("`"):
            out.append((part[1:-1], False, True))
        elif part.startswith("*") and part.endswith("*"):
            out.append((part[1:-1], False, True))
        else:
            out.append((part, False, False))
    return out


def _add_memo_header(doc: Document) -> None:
    """TO/FROM/DATE/RE header. Each line bold-label + body text,
    single-spaced, separated from the body by a divider rule."""
    for label, body in (
        ("TO:",   "Forest Capital Leadership"),
        ("FROM:", "FNA670 Industry Practicum Team"),
        ("DATE:", "May 27, 2026"),
        ("RE:",   "Multi-Strategy Portfolio Diversification — "
                  "Preliminary Findings"),
    ):
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_after = Pt(2)
        run_label = para.add_run(label + " ")
        run_label.bold = True
        run_label.font.name = _BODY_FONT
        run_label.font.size = _BODY_SIZE
        run_body = para.add_run(body)
        run_body.font.name = _BODY_FONT
        run_body.font.size = _BODY_SIZE

    # Divider rule — horizontal line via underscore characters.
    divider = doc.add_paragraph()
    divider.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    divider.paragraph_format.space_before = Pt(2)
    divider.paragraph_format.space_after = Pt(8)
    div_run = divider.add_run("_" * 80)
    div_run.font.name = _BODY_FONT
    div_run.font.size = _BODY_SIZE


# ── Section parser ──────────────────────────────────────────────────────────


_SECTION_RE = re.compile(
    r"(?m)^#{1,3}\s*(?:SECTION\s+)?(\d+)\.?\s*(.*?)$",
    re.IGNORECASE)

_BULLET_RE = re.compile(r"^\s*(?:[-*•]|\d+\.)\s+")


def _split_brief_sections(
    brief_md: str,
) -> list[tuple[str, str]]:
    """Splits the brief markdown body on H1/H2/H3 section markers.

    Returns a list of (heading, body) tuples. Headings retain the
    "1. Title" form so the memo renderer can extract just the title.
    Falls back to a single ("", whole-text) tuple when no headings
    are detected."""
    matches = list(_SECTION_RE.finditer(brief_md or ""))
    if not matches:
        return [("", (brief_md or "").strip())]
    out: list[tuple[str, str]] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = (matches[i + 1].start()
               if i + 1 < len(matches) else len(brief_md))
        body = brief_md[start:end].strip()
        title = m.group(2).strip() or f"Section {m.group(1)}"
        out.append((title, body))
    return out


def _add_brief_body_block(doc: Document, text: str) -> None:
    """Renders the brief's body block. Lines starting with -, *, or
    a numbered marker become bulleted items; other paragraphs render
    as plain memo prose."""
    for block in (text or "").strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Multi-line bullet block?
        lines = block.split("\n")
        if all(_BULLET_RE.match(ln) for ln in lines if ln.strip()):
            for ln in lines:
                if not ln.strip():
                    continue
                content = _BULLET_RE.sub("", ln).strip()
                _add_para(doc, f"• {content}", space_after_pt=2)
            continue
        _add_para(doc, block, space_after_pt=6)


def build_brief_docx(
    brief_md: str,
    *,
    references_md: str | None = None,  # noqa: ARG001 — accepted for parity
) -> bytes:
    """Renders an executive brief markdown body into memo-style .docx
    bytes. The references_md parameter is accepted for interface
    parity with build_paper_docx but is not rendered — the executive
    brief does not include an inline References list (the appendix
    carries the full citation record)."""
    doc = _new_brief_document()
    _add_memo_header(doc)

    sections = _split_brief_sections(brief_md)
    for heading, body in sections:
        if heading:
            # Strip the leading "N." numbering for the rendered title.
            clean = re.sub(r"^\d+\.\s*", "", heading).strip()
            _add_section_heading(doc, clean or heading)
        _add_brief_body_block(doc, body)

    # Footer line per the system prompt spec.
    footer = doc.add_paragraph()
    footer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    footer.paragraph_format.space_before = Pt(10)
    f_run = footer.add_run(
        "Full methodology and data validation documentation available "
        "on request. Platform: forest-capital.vercel.app")
    f_run.italic = True
    f_run.font.name = _BODY_FONT
    f_run.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Word count helpers (Bob's editor consults these client-side; the
#    server-side count is also exposed here for the harness's word-
#    budget enforcement). ───────────────────────────────────────────────


_BRIEF_SECTION_BUDGETS: dict[int, int] = {
    1: 60, 2: 180, 3: 80, 4: 80, 5: 90,
}
_BRIEF_TOTAL_BUDGET = 490


def get_section_budgets() -> dict[int, int]:
    """Read-only accessor for the executive brief section word
    budgets. Frontend mirrors this in lib/bobBlocks for the editor's
    Word Counts sidebar to render the correct totals when the brief
    template is selected. (Tested via test_report_writer_docx_brief.)"""
    return dict(_BRIEF_SECTION_BUDGETS)


def get_total_budget() -> int:
    return _BRIEF_TOTAL_BUDGET
