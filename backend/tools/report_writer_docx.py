"""tools/report_writer_docx.py — APA 7th edition docx assembly.

May 22 2026 (item 12 commit C). Renders the midpoint paper +
appendix as APA 7th edition formatted .docx files for FNA670
submission.

APA 7th EDITION CONFORMANCE:

  Document formatting:
    Font: Times New Roman 12pt
    Line spacing: double (1.0 in tables / data blocks)
    Margins: 1 inch all sides
    Page numbers: top-right header, starting at the title page
    Running head: NOT required for student papers in APA 7th
    First-line indent: 0.5 inch on every body paragraph

  Title page (additional page, not counted in the three-page limit):
    Paper title in title case — bold, centred, 3-4 lines down from
      the top
    Author names — centred below the title
    Institutional affiliation — Queens University · McColl School
      of Business — centred
    Course name + number — FNA670 Industry Practicum
    Instructor — Dr. Panttser
    Submission date — May 27 2026

  Headings:
    Level 1 (section): bold, centred, title case
    Level 2 (subsection): bold, left-aligned, title case
    (No heading at the document top — the title page replaces it.)

  References page:
    Starts on a fresh page after Section 4
    Title: References (bold, centred, Level 1)
    Hanging indent of 0.5 inch on every entry
    Double-spaced

  Appendix formatting:
    Each appendix starts on a new page
    Label: Appendix A / B / ... (bold, centred)
    Title below the label (bold, centred)
    Continued APA paragraph formatting

The paper has NO header text — APA 7th student papers do not use a
running head. Only the page number appears in the top-right header.

Build entry points (unchanged interface):
  build_paper_docx(paper_md, references_md=None) → bytes
  build_appendix_docx(context) → bytes
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# Citations are "verified-equivalent" if any of these four states.
# Mirrors template_pipeline.CITATION_VERIFIED_STATES; inlined here to
# avoid an import cycle with the pipeline module. May 26 2026 —
# centralised so References / Appendix A methodological references /
# Appendix D GIPS reference all use the same filter (the prior code
# checked only the literal 'verified' state, so adjudicated
# citations — human_verified / search_selected / manually_added —
# silently dropped out of all three sections).
_VERIFIED_STATES: frozenset[str] = frozenset({
    "verified",          # automatic trusted-domain hit
    "human_verified",    # reviewer accepted an untrusted-domain hit
    "search_selected",   # reviewer accepted one of the alternatives
    "manually_added",    # reviewer typed the citation in manually
})


def _is_verified(citation: dict | None) -> bool:
    """True when this citation row counts as verified for the purposes
    of References / inline citation rendering. Centralised so the
    four call sites stay in lockstep."""
    if not citation:
        return False
    return citation.get("verification_status") in _VERIFIED_STATES


_BODY_FONT = "Times New Roman"
_INDENT_FIRST_LINE = Inches(0.5)
_INDENT_HANGING = Inches(0.5)


# ── Title page constants ────────────────────────────────────────────────────


PAPER_TITLE = (
    "Multi-Strategy Portfolio Diversification: A Midpoint Analysis")
AUTHOR_LINE = "Michael Ruurds, Bob Thao, and Molly Murdock"
INSTITUTION_LINE = "Queens University — McColl School of Business"
COURSE_LINE = "FNA670 Industry Practicum"
INSTRUCTOR_LINE = "Instructor: Dr. Panttser"
SUBMISSION_DATE_LINE = "May 27, 2026"


# ── Low-level helpers ────────────────────────────────────────────────────────


def _set_run_color(run, hex_color: str) -> None:
    rgb = hex_color.lstrip("#")
    run.font.color.rgb = RGBColor(
        int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))


def _add_page_number(paragraph) -> None:
    """Live PAGE field via raw OOXML."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _add_page_break(doc: Document) -> None:
    """Forces the next paragraph onto a new page. APA 7th requires
    page breaks before the References section and before each
    Appendix label."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _new_apa_document(
    *, body_spacing: WD_LINE_SPACING = WD_LINE_SPACING.DOUBLE,
) -> Document:
    """APA 7th edition document — 1-inch margins, Times New Roman
    12pt double-spaced body, page number in the top-right header.
    No running head (student papers in APA 7th do not require it)."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = body_spacing
    normal.paragraph_format.space_after = Pt(0)

    # APA 7th student paper: page number only in the top-right header.
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number(header_para)
    for r in header_para.runs:
        r.font.name = _BODY_FONT
        r.font.size = Pt(12)

    return doc


def _add_title_page(doc: Document) -> None:
    """APA 7th student paper title page. The paper title appears 3-4
    double-spaced lines down from the top — implemented as three
    empty paragraphs above the title for a defensible approximation
    of the convention. Every line is centred, double-spaced, and in
    the body font.

    The page closes with a page break so Section 1 begins on a fresh
    page (also a standard APA convention — body text starts on
    page 2)."""
    # Spacer lines pushing the title block down the page.
    for _ in range(4):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title — bold, centred, title case.
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(PAPER_TITLE)
    title_run.bold = True
    title_run.font.name = _BODY_FONT
    title_run.font.size = Pt(12)

    # One empty line between the title and the author block.
    spacer = doc.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author + institution + course + instructor + date — each on
    # its own centred line.
    for line in (
        AUTHOR_LINE,
        INSTITUTION_LINE,
        COURSE_LINE,
        INSTRUCTOR_LINE,
        SUBMISSION_DATE_LINE,
    ):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.font.name = _BODY_FONT
        run.font.size = Pt(12)

    _add_page_break(doc)


def _add_section_heading(
    doc: Document, text: str, *,
    level: int = 1,
) -> None:
    """APA 7th heading.

      Level 1: bold, centred, title case  (section headings)
      Level 2: bold, left-aligned, title case  (subsection if used)

    Both use 12pt body font (NOT enlarged) per APA 7th convention.
    Single-spaced before/after to keep the body double spacing intact."""
    para = doc.add_paragraph()
    para.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if level <= 1
        else WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.name = _BODY_FONT
    run.font.size = Pt(12)


def _add_body_paragraph(doc: Document, text: str) -> None:
    """Adds one body paragraph with the APA 0.5-inch first-line
    indent and double spacing. Inline markdown bold/italic/code is
    preserved via _split_inline."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.first_line_indent = _INDENT_FIRST_LINE
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for run_text, bold, italic in _split_inline(text):
        if not run_text:
            continue
        run = para.add_run(run_text)
        run.font.name = _BODY_FONT
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic


def _add_reference_entry(doc: Document, text: str) -> None:
    """Adds one References-list entry with the APA hanging indent
    (the FIRST line starts at the left margin; every subsequent line
    indents 0.5 inch). Double-spaced like the body."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.left_indent = _INDENT_HANGING
    para.paragraph_format.first_line_indent = Inches(-0.5)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_after = Pt(0)
    for run_text, bold, italic in _split_inline(text):
        if not run_text:
            continue
        run = para.add_run(run_text)
        run.font.name = _BODY_FONT
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic


def _add_body_block(doc: Document, text: str) -> None:
    """Renders a markdown body block — paragraphs separated by blank
    lines become individual APA paragraphs. Empty input is a no-op."""
    for block in (text or "").strip().split("\n\n"):
        block = block.strip()
        if not block:
            continue
        _add_body_paragraph(doc, block)


_INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _split_inline(block: str) -> "list[tuple[str, bool, bool]]":
    out: list[tuple[str, bool, bool]] = []
    for part in _INLINE_RE.split(block):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True, False))
        elif part.startswith("`") and part.endswith("`"):
            out.append((part[1:-1], False, True))  # render code as italic
        elif part.startswith("*") and part.endswith("*"):
            out.append((part[1:-1], False, True))
        else:
            out.append((part, False, False))
    return out


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    *,
    column_widths_in: list[float] | None = None,
) -> None:
    """APA tables. Single-spaced, font matches the body. APA 7th uses
    minimal borders (no vertical lines); we keep the Table Grid style
    for legibility — the produced docx renders cleanly in Word."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(h)
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(str(val) if val is not None else "")
            run.font.name = _BODY_FONT
            run.font.size = Pt(11)
    if column_widths_in:
        for i, w in enumerate(column_widths_in):
            if i < len(table.columns):
                for cell in table.columns[i].cells:
                    cell.width = Inches(w)


# ── Paper builder ────────────────────────────────────────────────────────────


_PAPER_SECTION_RE = re.compile(
    r"(?m)^#{1,3}\s*(?:SECTION\s+)?(\d+)\.?\s*(.*?)$",
    re.IGNORECASE)


def _split_paper_sections(
    paper_md: str,
) -> list[tuple[str, str]]:
    """Splits a paper markdown body into (heading, body) tuples.
    Falls back to a single ("", body) tuple when no H2 headings are
    found. The heading is reformatted into APA title case — the
    section number is dropped from the rendered heading per APA 7th
    convention (numbered subsection labels are unusual)."""
    matches = list(_PAPER_SECTION_RE.finditer(paper_md or ""))
    if not matches:
        return [("", (paper_md or "").strip())]
    out: list[tuple[str, str]] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(paper_md)
        body = paper_md[start:end].strip()
        title = m.group(2).strip() or f"Section {m.group(1)}"
        out.append((_to_title_case(title), body))
    return out


_TITLE_CASE_LOWERS = {
    "a", "an", "the",
    "and", "but", "or", "nor", "for", "yet", "so",
    "at", "by", "in", "of", "on", "to", "up", "as", "via",
    "with", "from", "into", "onto", "over", "vs",
}


def _to_title_case(text: str) -> str:
    """APA title case: capitalise the first/last word + all major
    words; leave short prepositions / articles / conjunctions
    lowercase unless they're the first or last token. Already-cased
    input passes through largely unchanged."""
    raw = (text or "").strip()
    if not raw:
        return ""
    parts = raw.split()
    out: list[str] = []
    for i, w in enumerate(parts):
        bare = w.lower()
        if i != 0 and i != len(parts) - 1 and bare in _TITLE_CASE_LOWERS:
            out.append(bare)
        else:
            out.append(w[:1].upper() + w[1:] if w else w)
    return " ".join(out)


def build_paper_docx(
    paper_md: str,
    *,
    references_md: str | None = None,
) -> bytes:
    """Renders the main paper markdown into APA 7th edition .docx
    bytes.

    Layout:
      Page 1  Title page (centred title, authors, institution,
              course, instructor, date)
      Page 2+ Section 1 → 4 with the configured first-line indent
              and double-spacing; Level 1 headings centred + bold
      Final   References (when references_md is non-empty), starting
              on its own page with hanging-indent entries
    """
    doc = _new_apa_document()
    _add_title_page(doc)

    sections = _split_paper_sections(paper_md)
    for heading, body in sections:
        if heading:
            _add_section_heading(doc, heading, level=1)
        _add_body_block(doc, body)

    if references_md:
        _add_page_break(doc)
        _add_section_heading(doc, "References", level=1)
        for line in (references_md or "").split("\n\n"):
            line = line.strip()
            if line:
                _add_reference_entry(doc, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Appendix builder ────────────────────────────────────────────────────────


def _fmt_value(v: Any) -> str:
    if v is None:
        return "—"
    if isinstance(v, bool):
        return "Yes" if v else "No"
    if isinstance(v, float):
        return f"{v:.4f}"
    # May 23 2026 — a dict or list slipped through (e.g. an upstream
    # snapshot shape drift) used to render as its Python repr — "{}"
    # or "[]" — leaking into Appendix D table cells. Render the em
    # dash instead; the upstream caller is responsible for flattening
    # nested structures before they reach the table.
    if isinstance(v, (dict, list, tuple, set)):
        return "—"
    return str(v)


def _add_appendix_label(
    doc: Document, letter: str, title: str,
) -> None:
    """APA 7th appendix: label (Appendix A) on one centred bold
    line; title on the next centred bold line. Starts on a new page."""
    _add_page_break(doc)
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    label_para.paragraph_format.space_after = Pt(6)
    label_run = label_para.add_run(f"Appendix {letter}")
    label_run.bold = True
    label_run.font.name = _BODY_FONT
    label_run.font.size = Pt(12)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = _BODY_FONT
    title_run.font.size = Pt(12)


def _appendix_a(doc: Document, context: dict) -> None:
    """Appendix A — Platform Overview."""
    _add_appendix_label(doc, "A", "Platform Overview")
    _add_body_paragraph(doc, (
        "The Forest Capital Portfolio Intelligence Platform is a "
        "purpose-built portfolio intelligence system developed for "
        "this practicum. It provides live data pipelines, a ten-"
        "strategy backtesting engine, an AI council of specialists, "
        "independent three-layer data validation, and an integrated "
        "academic review workflow."))

    vd = context.get("verified_data") or {}
    activity = context.get("team_activity") or {}

    _add_section_heading(doc, "Key Platform Statistics", level=2)
    _add_table(
        doc,
        headers=["Statistic", "Value"],
        rows=[
            ["Strategies analyzed",                10],
            ["Data period start",                  _fmt_value(
                vd.get("study_period_start"))],
            ["Data period end",                    _fmt_value(
                vd.get("study_period_end"))],
            ["Monthly observations",               _fmt_value(
                vd.get("n_months"))],
            ["Independent validations passed",     _fmt_value(
                activity.get("team_total_audit_validations"))],
            ["Council sessions completed",         _fmt_value(
                activity.get("team_total_council_sessions"))],
            ["UAT test steps completed",           _fmt_value(
                activity.get("team_total_uat_steps"))],
        ],
        column_widths_in=[3.0, 2.0])

    _add_body_paragraph(doc, "Platform URL: forest-capital.vercel.app.")
    _add_body_paragraph(doc, (
        "Access has been granted to course faculty. Login uses the "
        "user's Queens University email address."))

    citations = context.get("citations_cache") or {}
    cited: list[str] = []
    for cid in ("cvar_coherent_risk", "four_factor_model",
                "portfolio_diversification", "regime_switching"):
        c = citations.get(cid) or {}
        # May 26 2026 — _is_verified accepts all four verified-equivalent
        # states (verified / human_verified / search_selected /
        # manually_added). Previously this checked the literal 'verified'
        # only, so Bob's adjudicated citations dropped out and the
        # [CITATION REQUIRED] placeholder rendered even when valid
        # references were in the cache.
        if _is_verified(c):
            cited.append(c.get("formatted") or "")
    if cited:
        _add_section_heading(doc, "Methodological References", level=2)
        _add_body_paragraph(doc, (
            "The platform's independent validation methodology draws "
            "on established practices in quantitative finance:"))
        for entry in cited:
            _add_body_paragraph(doc, f"• {entry}")
    else:
        _add_body_paragraph(
            doc, "[CITATION REQUIRED — methodological references]")


def _appendix_b(doc: Document, context: dict) -> None:
    """Appendix B — Full Analytical Findings."""
    _add_appendix_label(doc, "B", "Full Analytical Findings")
    _add_body_paragraph(doc, (
        "All findings pre-computed from live platform data. The "
        "findings below form the factual record alongside the main "
        "paper's evidence."))

    findings = context.get("ranked_findings") or []
    if not findings:
        _add_body_paragraph(doc, "[DATA REQUIRED — no findings staged]")
    for i, f in enumerate(findings, start=1):
        _add_section_heading(
            doc,
            f"F{i} — {f.get('title', '')} ({f.get('nugget_strength', 'LOW')})",
            level=2)
        _add_body_paragraph(doc, f"**FINDING:** {f.get('finding', '')}")
        ev = f.get("evidence") or []
        if ev:
            _add_body_paragraph(doc, "**EVIDENCE:**")
            for e in ev:
                _add_body_paragraph(doc, f"• {e}")
        _add_body_paragraph(
            doc, f"**IMPLICATION:** {f.get('implication', '')}")
        if f.get("surprise"):
            _add_body_paragraph(
                doc, f"**SURPRISE:** {f.get('surprise_reason') or 'yes'}")

    md = context.get("findings_metadata") or {}
    # May 26 2026 — submission fix. dict.get(key, default) returns the
    # value when the key exists, EVEN IF that value is None. So
    # md.get('data_hash', '—') rendered as 'None' when the underlying
    # findings_row carried a NULL data_hash. _fmt_value normalises
    # None / dict / etc. to an em dash. Same defensive read applied
    # to every interpolated field.
    _add_body_paragraph(doc, (
        f"All figures pulled live from platform analytics cache on "
        f"{_fmt_value(md.get('computed_at'))}. Data hash: "
        f"{_fmt_value(md.get('data_hash'))}. Independent validation: "
        f"three-layer audit {_fmt_value(md.get('audit_status'))}."))


def _appendix_c(doc: Document, context: dict) -> None:
    """Appendix C — Team Activity Log."""
    _add_appendix_label(doc, "C", "Team Activity Log")
    _add_body_paragraph(doc, (
        "Auditable record of team contributions throughout the "
        "project lifecycle. All counts pulled live from the "
        "platform's audit tables at generation time."))

    activity = context.get("team_activity") or {}
    rows: list[list[Any]] = []
    member_groups = [
        ("Michael Ruurds", [
            ("Commits to repository",  "michael_commits"),
            ("PRs merged",             "michael_prs_merged"),
            ("Migrations deployed",    "michael_migrations_deployed"),
            ("Failure reports resolved", "michael_failure_reports_resolved"),
            # UAT 2026-05-24 — Michael runs UAT Section 2 per the
            # project backlog. The row was missing from the rendered
            # table even though michael_uat_steps fed the cross-check
            # reconciliation, leaving the reader unable to see why
            # the platform total exceeded Bob + Molly.
            ("UAT test steps completed", "michael_uat_steps"),
        ]),
        ("Bob Thao", [
            ("UAT test steps completed", "bob_uat_steps"),
            ("Council sessions initiated", "bob_council_sessions"),
            ("Academic review runs",     "bob_academic_review_runs"),
            ("Report drafts generated",  "bob_report_drafts"),
        ]),
        ("Molly Murdock", [
            ("UAT test steps completed", "molly_uat_steps"),
            ("Failure reports filed",    "molly_failure_reports_filed"),
            ("Feedback items submitted", "molly_feedback_items"),
        ]),
    ]
    for member, fields in member_groups:
        for activity_label, field_key in fields:
            rows.append([
                member, activity_label, _fmt_value(activity.get(field_key)),
            ])
    rows.append([
        "Platform total", "UAT test steps",
        _fmt_value(activity.get("team_total_uat_steps")),
    ])
    rows.append([
        "Platform total", "Failure reports filed",
        _fmt_value(activity.get("team_total_failure_reports")),
    ])
    rows.append([
        "Platform total", "Failure reports resolved",
        _fmt_value(activity.get("team_total_failure_reports_resolved")),
    ])
    rows.append([
        "Platform total", "Council sessions",
        _fmt_value(activity.get("team_total_council_sessions")),
    ])
    rows.append([
        "Platform total", "Independent validations",
        _fmt_value(activity.get("team_total_audit_validations")),
    ])
    _add_table(
        doc,
        headers=["Member / Total", "Activity", "Count"],
        rows=rows,
        column_widths_in=[2.0, 3.0, 1.0])
    # May 26 2026 — footer was breaking mid-sentence in the docx.
    # The previous _add_body_paragraph call applied APA double
    # spacing + a 0.5-inch first-line indent — that combined with
    # a long ISO timestamp pushed the second sentence onto its own
    # line and (depending on table position) sometimes onto its
    # own page. A SHORT table footnote is the conventional shape
    # here; render as a single-spaced, non-indented caption with
    # `keep_together` so Word keeps it on the same page as the
    # table and never breaks mid-sentence.
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.SINGLE)
    footer_para.paragraph_format.first_line_indent = Inches(0)
    footer_para.paragraph_format.space_before = Pt(6)
    # Word's "keep lines together" — the footer's lines never break
    # across pages. python-docx exposes this via paragraph_format.
    footer_para.paragraph_format.keep_together = True
    footer_run = footer_para.add_run(
        f"Activity data pulled live from the platform's audit log "
        f"at {context.get('generated_at', '—')}. "
        "All timestamps recorded at point of action.")
    footer_run.font.name = _BODY_FONT
    footer_run.font.size = Pt(10)
    footer_run.italic = True


def _appendix_d(doc: Document, context: dict) -> None:
    """Appendix D — Independent Data Validation Summary."""
    _add_appendix_label(
        doc, "D", "Independent Data Validation Summary")
    validation = context.get("validation_summary") or {}

    _add_section_heading(doc, "Three-Layer Audit Results", level=2)
    _add_table(
        doc,
        headers=["Layer", "Status", "Checks", "Last Run"],
        rows=[
            ["1 — Raw data audit",
             _fmt_value(validation.get("layer1_status")),
             _fmt_value(validation.get("layer1_count")),
             _fmt_value(validation.get("layer1_date"))],
            ["2 — Calculation audit",
             _fmt_value(validation.get("layer2_status")),
             _fmt_value(validation.get("layer2_count")),
             _fmt_value(validation.get("layer2_date"))],
            ["3 — Consistency audit",
             _fmt_value(validation.get("layer3_status")),
             _fmt_value(validation.get("layer3_count")),
             _fmt_value(validation.get("layer3_date"))],
        ],
        column_widths_in=[2.5, 1.5, 1.0, 1.5])

    _add_section_heading(doc, "Validation Methodology", level=2)
    _add_body_paragraph(doc, (
        "Layer 1 (raw data audit) verifies source data integrity "
        "against published benchmarks. US equity returns benchmarked "
        "against published S&P 500 total return index data; risk-free "
        "rate benchmarked against published DTB3 Treasury bill data "
        "(Federal Reserve H.15 release)."))
    _add_body_paragraph(doc, (
        "Layer 2 (calculation audit) recomputes every metric from "
        "raw inputs in an independent model instance. Standard "
        "formulas: CAGR as the geometric mean of (1 + r) compounded "
        "monthly; Sharpe as (Rp − Rf) / σp annualized by √12; CVaR "
        "as the mean of returns below the VaR threshold via "
        "historical simulation; max drawdown as the maximum peak-"
        "to-trough decline in the cumulative return series."))
    _add_body_paragraph(doc, (
        "Layer 3 (consistency audit) cross-checks every metric "
        "across every surface where it appears. Tolerances: ratios "
        "0.001, percentages 0.01%, factor betas 0.001."))

    citations = context.get("citations_cache") or {}
    gips = citations.get("gips_verification") or {}
    if _is_verified(gips):
        _add_body_paragraph(doc, (
            f"This methodology is consistent with industry standards "
            f"for independent performance verification as described "
            f"in {gips.get('formatted')}"))
    else:
        _add_body_paragraph(doc, "[CITATION REQUIRED — GIPS]")


def _build_references_md(citations: dict) -> str:
    """Builds the consolidated References markdown from every
    verified-equivalent entry. Alphabetical by first-author surname.

    May 26 2026 — submission fix. The filter previously matched ONLY
    the literal 'verified' state, which is the automatic
    trusted-domain pass-1 outcome. Citations that Bob ADJUDICATED
    via the Citation Review panel (becoming human_verified /
    search_selected / manually_added) were silently excluded from
    References — making the section render empty even when Bob had
    just confirmed every citation. Broadened to match
    CITATION_VERIFIED_STATES (the same set citation_quality() uses).
    """
    # Inline the verified set so this module has no import-cycle risk
    # with template_pipeline.CITATION_VERIFIED_STATES. The two must
    # stay in lockstep; the inline copy carries a comment naming the
    # canonical source.
    _VERIFIED_STATES = frozenset({
        # Mirrors template_pipeline.CITATION_VERIFIED_STATES.
        "verified",          # automatic trusted-domain hit
        "human_verified",    # Bob accepted an untrusted-domain hit
        "search_selected",   # Bob accepted one of the alternatives
        "manually_added",    # Bob typed the citation in manually
    })
    verified = [
        c for c in (citations or {}).values()
        if c.get("verification_status") in _VERIFIED_STATES]
    if not verified:
        return ""
    verified.sort(key=lambda c: (c.get("author") or "").lower())
    lines = [
        c.get("formatted") or "(unformatted)"
        for c in verified]
    return "\n\n".join(lines)


def build_appendix_docx(context: dict) -> bytes:
    """Assembles the four-section appendix into APA 7th .docx bytes.

    Each appendix starts on a new page with a centred 'Appendix X'
    label and a centred bold title. The References list appears
    after Appendix D with hanging-indent APA entries.

    context schema:
      verified_data:       dict — for Appendix A statistics + footers
      ranked_findings:     list — for Appendix B
      team_activity:       dict — for Appendix C
      validation_summary:  dict — for Appendix D
      citations_cache:     dict — for inline references + References
      findings_metadata:   dict — computed_at, data_hash, audit_status
      generated_at:        str  — for Appendix C footer
    """
    doc = _new_apa_document()

    _appendix_a(doc, context)
    _appendix_b(doc, context)
    _appendix_c(doc, context)
    _appendix_d(doc, context)

    refs_md = _build_references_md(context.get("citations_cache") or {})
    _add_page_break(doc)
    _add_section_heading(doc, "References", level=1)
    if refs_md:
        for line in refs_md.split("\n\n"):
            line = line.strip()
            if line:
                _add_reference_entry(doc, line)
    else:
        _add_body_paragraph(
            doc,
            "[REFERENCES UNAVAILABLE — citations not yet sourced]")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
