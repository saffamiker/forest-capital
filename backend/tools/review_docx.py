"""
tools/review_docx.py -- June 25 2026.

Builds a downloadable DOCX of a completed academic review for one
draft. Sourced from:

  council_debates row     (peer_responses, critic_findings,
                           arbiter_resolution, fix_proposals,
                           counter_arguments)
  agent_interactions row  (response_summary = arbiter verdict text,
                           metadata = harness + critic counts +
                           overall_rating + section_ratings)
  editor draft            (title, document_type, version,
                           created_at, data_hash)

Six sections in the spec'd order:

  1. Cover block          Forest Capital -- FNA 670 banner +
                          per-draft metadata
  2. Peer Review Summary  Per-peer-agent sections rendering each
                          peer's prose output as Heading 2 +
                          paragraph runs. Falls back to a single
                          'Arbiter Verdict' section if peer
                          payload absent.
  3. Adversarial Critic   Fatal/Major/Minor counts + per-finding
                          block (severity | category | location,
                          description, evidence, recommendation)
  4. Council Response     The arbiter_resolution prose with
                          heading hierarchy preserved
  5. Fix Proposals        Target / severity / patch / rationale
  6. Independent Verdict  When present in metadata

Uses python-docx (same stack as build_editor_docx). Returns raw
bytes; the endpoint wraps with the standard Content-Disposition
download headers.
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


# ── Constants ────────────────────────────────────────────────────────────────

_DOC_TYPE_LABELS = {
    "executive_brief":      "Executive Brief",
    "presentation_deck":    "Final Presentation Deck",
    "analytical_appendix":  "Analytical Appendix",
    "presentation_script":  "Presentation Script",
    "midpoint_paper":       "Midpoint Paper",
}

_PEER_AGENT_LABELS = {
    "claude":    "Anthropic Claude (Opus)",
    "gpt":       "OpenAI GPT",
    "gpt-4":     "OpenAI GPT-4",
    "gemini":    "Google Gemini",
    "grok":      "xAI Grok",
    "perplexity": "Perplexity",
}

_SEVERITY_LABELS = {
    "fatal":  "Fatal",
    "major":  "Major",
    "minor":  "Minor",
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def _format_timestamp(iso: str | None) -> str:
    if not iso:
        return "—"
    try:
        from datetime import datetime
        # ISO strings may carry the +HH:MM offset; fromisoformat handles
        # post-3.11 input including Z (we strip Z for older runtimes).
        s = iso.replace("Z", "+00:00")
        dt = datetime.fromisoformat(s)
        return dt.strftime("%b %d, %Y %I:%M %p UTC")
    except Exception:  # noqa: BLE001
        return iso


def _peer_label(agent_id: str) -> str:
    return _PEER_AGENT_LABELS.get(agent_id.lower(), agent_id)


def _doc_label(document_type: str | None) -> str:
    if not document_type:
        return "Document"
    return _DOC_TYPE_LABELS.get(document_type, document_type)


def _add_heading(doc: Any, text: str, level: int) -> None:
    """Heading 1/2/3 with the docx built-in heading style. Level
    is clamped to [1, 4] so a stray nested heading still renders."""
    h = doc.add_heading(text, level=max(1, min(level, 4)))
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_para(
    doc: Any, text: str, bold: bool = False, italic: bool = False,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def _add_bullet(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)


def _add_numbered(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(11)


def _add_label_value(
    doc: Any, label: str, value: str,
) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    value_run = p.add_run(value)
    value_run.font.size = Pt(11)


def _render_markdown_block(doc: Any, text: str) -> None:
    """Renders a peer / arbiter Markdown-shaped string as DOCX
    headings + paragraphs + bullets. Handles:
      ### / #### / ##### -> Heading 3-5
      ## -> Heading 2 (treated as sub-section under the current
            agent header which is already H2)
      -  / *           -> bullet
      1. 2. 3.         -> numbered
      bold **xx**      -> bold run inside paragraph (one level)
      blank lines      -> paragraph breaks
    A line that isn't a heading or list becomes a paragraph; the
    fallback is forgiving so a peer response that goes off-script
    still surfaces something readable rather than nothing."""
    if not text:
        return
    lines = text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        # Heading detection.
        if line.startswith("##### "):
            _add_heading(doc, line[6:].strip(), level=4)
            continue
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), level=4)
            continue
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=3)
            continue
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=2)
            continue
        # Bullet.
        stripped = line.lstrip()
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            _render_inline(doc, bullet_match.group(1), bullet=True)
            continue
        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            _render_inline(doc, numbered_match.group(1), numbered=True)
            continue
        _render_inline(doc, line)


def _render_inline(
    doc: Any, text: str, bullet: bool = False, numbered: bool = False,
) -> None:
    """One paragraph with bold runs extracted from **bold** spans
    and italic runs from *italic*. Falls back to plain text when
    no markers are present."""
    style: str | None = None
    if bullet:
        style = "List Bullet"
    elif numbered:
        style = "List Number"
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    # Token sequence: split on bold first, then italic within each
    # plain span. Keeps the rendering layer trivial without dragging
    # a full Markdown parser in.
    bold_parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in bold_parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            continue
        # Process italic within plain spans.
        italic_parts = re.split(r"(?<!\*)(\*[^*]+\*)(?!\*)", part)
        for sub in italic_parts:
            if not sub:
                continue
            if sub.startswith("*") and sub.endswith("*"):
                run = p.add_run(sub[1:-1])
                run.italic = True
                run.font.size = Pt(11)
            else:
                run = p.add_run(sub)
                run.font.size = Pt(11)


def _add_severity_run(p: Any, sev: str) -> None:
    """Coloured severity tag for a finding header. Fatal=red,
    Major=orange, Minor=gray. Bold for all."""
    run = p.add_run(_SEVERITY_LABELS.get(sev.lower(), sev.upper()))
    run.bold = True
    run.font.size = Pt(11)
    lower = sev.lower()
    if lower == "fatal":
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    elif lower == "major":
        run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    else:
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)


def _add_section_divider(doc: Any) -> None:
    """Blank paragraph after a section block for visual breathing
    room. Six sections looks dense without one."""
    p = doc.add_paragraph()
    p.add_run("")


# ── Builders ─────────────────────────────────────────────────────────────────


def build_review_docx(
    *,
    draft: dict[str, Any],
    debate: dict[str, Any] | None,
    interaction: dict[str, Any] | None,
) -> bytes:
    """Assembles the review report from the three data sources.

    draft       -- editor_drafts row (id, document_type, title,
                   version, created_at, data_hash, ...).
    debate      -- council_debates row (peer_responses,
                   critic_findings, arbiter_resolution,
                   fix_proposals, counter_arguments, fatal/major/
                   minor counts, document_type, data_hash) or None.
    interaction -- agent_interactions row for the academic_review
                   entry (response_summary = arbiter prose,
                   metadata = harness + critic + overall_rating +
                   section_ratings) or None.

    At least one of debate/interaction should be present; both
    None still produces a valid DOCX with a cover block and a
    'no review found' notice so the download path doesn't error."""
    doc = Document()

    # ── 1. Cover block ─────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Forest Capital — FNA 670")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Academic Review Report")
    sub_run.bold = True
    sub_run.font.size = Pt(13)

    doc.add_paragraph()

    document_type = draft.get("document_type")
    doc_label = _doc_label(document_type)
    _add_label_value(doc, "Document", doc_label)
    review_date = (
        interaction.get("timestamp") if interaction else None
    ) or (debate.get("created_at") if debate else None)
    _add_label_value(doc, "Review date", _format_timestamp(review_date))
    draft_label = (
        f"#{draft.get('id', '—')} "
        f"v{draft.get('version', '—')} "
        f"({_format_timestamp(draft.get('created_at'))})"
    )
    _add_label_value(doc, "Draft", draft_label)
    data_hash = (
        draft.get("data_hash")
        or (debate.get("data_hash") if debate else None)
        or "—")
    _add_label_value(doc, "Data hash", str(data_hash))
    overall_rating: str | None = None
    score: float | None = None
    if interaction:
        meta = interaction.get("metadata") or {}
        overall_rating = meta.get("overall_rating")
        s = meta.get("score")
        if isinstance(s, (int, float)):
            score = float(s)
    if overall_rating:
        _add_label_value(doc, "Overall rating", str(overall_rating))
    if score is not None:
        _add_label_value(doc, "Arbiter score", f"{score:.1f} / 10")

    _add_section_divider(doc)

    # ── 2. Peer Review Summary ────────────────────────────────
    _add_heading(doc, "1. Peer Review Summary", level=1)
    peer_responses: dict[str, Any] = (
        (debate or {}).get("peer_responses") or {})
    if peer_responses:
        for agent_id in sorted(peer_responses.keys()):
            text = str(peer_responses[agent_id] or "").strip()
            if not text:
                continue
            _add_heading(doc, _peer_label(agent_id), level=2)
            _render_markdown_block(doc, text)
            _add_section_divider(doc)
    else:
        _add_para(
            doc,
            "No structured peer responses were stored for this "
            "review. Refer to the arbiter resolution in Section 3 "
            "for the synthesised verdict.",
            italic=True)
        _add_section_divider(doc)

    # ── 3. Adversarial Critic Findings ────────────────────────
    _add_heading(doc, "2. Adversarial Critic Findings", level=1)
    if debate is not None:
        _add_label_value(doc, "Fatal",
                         str(debate.get("fatal_count") or 0))
        _add_label_value(doc, "Major",
                         str(debate.get("major_count") or 0))
        _add_label_value(doc, "Minor",
                         str(debate.get("minor_count") or 0))
        _add_label_value(
            doc, "Critic models",
            str(debate.get("critic_model") or "—"))
        doc.add_paragraph()
        findings = debate.get("critic_findings") or []
        if isinstance(findings, list) and findings:
            for finding in findings:
                if not isinstance(finding, dict):
                    continue
                header = doc.add_paragraph()
                _add_severity_run(
                    header, str(finding.get("severity") or "Minor"))
                cat = finding.get("category") or "—"
                loc = finding.get("location") or "—"
                tail = header.add_run(f"  •  {cat}  •  {loc}")
                tail.bold = False
                tail.font.size = Pt(11)
                tail.font.color.rgb = RGBColor(0x55, 0x6B, 0x78)
                desc = (finding.get("description") or "").strip()
                if desc:
                    _add_label_value(doc, "Description", desc)
                ev = (finding.get("evidence") or "").strip()
                if ev:
                    _add_label_value(doc, "Evidence", ev)
                rec = (finding.get("recommendation") or "").strip()
                if rec:
                    _add_label_value(doc, "Recommendation", rec)
                doc.add_paragraph()
        else:
            _add_para(
                doc,
                "Critic returned no actionable findings.",
                italic=True)
    else:
        _add_para(
            doc,
            "Critic pipeline did not run for this review.",
            italic=True)
    _add_section_divider(doc)

    # ── 4. Council Response to Critic ─────────────────────────
    _add_heading(doc, "3. Council Response to Critic", level=1)
    arbiter_resolution = (
        (debate or {}).get("arbiter_resolution") or "").strip()
    if arbiter_resolution:
        _render_markdown_block(doc, arbiter_resolution)
    else:
        # Fall back to the agent_interactions response_summary --
        # that's the original arbiter verdict before any debate
        # round, but it's the best available text when no debate
        # round fired (the minor-only path skips the resolution).
        response_summary = (
            (interaction or {}).get("response_summary") or "").strip()
        if response_summary:
            _add_para(
                doc,
                "(Debate round did not fire — original arbiter "
                "verdict shown below.)",
                italic=True)
            _render_markdown_block(doc, response_summary)
        else:
            _add_para(
                doc, "No arbiter resolution recorded.",
                italic=True)
    _add_section_divider(doc)

    # ── 5. Fix Proposals ──────────────────────────────────────
    _add_heading(doc, "4. Fix Proposals", level=1)
    fix_proposals = (debate or {}).get("fix_proposals") or {}
    if isinstance(fix_proposals, dict) and fix_proposals:
        idx = 1
        for finding_idx in sorted(
                fix_proposals.keys(),
                key=lambda k: (
                    int(k) if str(k).isdigit() else 0)):
            proposal = fix_proposals[finding_idx]
            if not isinstance(proposal, dict):
                continue
            _add_heading(
                doc,
                f"Proposal {idx} — finding #{finding_idx}",
                level=2)
            target = proposal.get("target_section") or "—"
            sev = proposal.get("severity") or "—"
            patch = proposal.get("patch_instruction") or ""
            rationale = proposal.get("rationale") or ""
            _add_label_value(doc, "Target section", str(target))
            _add_label_value(doc, "Severity", str(sev))
            if patch:
                _add_label_value(doc, "Patch instruction", str(patch))
            if rationale:
                _add_label_value(doc, "Rationale", str(rationale))
            doc.add_paragraph()
            idx += 1
    elif isinstance(fix_proposals, list) and fix_proposals:
        for i, proposal in enumerate(fix_proposals, start=1):
            if not isinstance(proposal, dict):
                continue
            _add_heading(
                doc, f"Proposal {i}", level=2)
            target = proposal.get("target_section") or "—"
            sev = proposal.get("severity") or "—"
            patch = proposal.get("patch_instruction") or ""
            rationale = proposal.get("rationale") or ""
            _add_label_value(doc, "Target section", str(target))
            _add_label_value(doc, "Severity", str(sev))
            if patch:
                _add_label_value(doc, "Patch instruction", str(patch))
            if rationale:
                _add_label_value(doc, "Rationale", str(rationale))
            doc.add_paragraph()
    else:
        _add_para(
            doc,
            "No fix proposals generated. The critic findings either "
            "were not actionable, or the arbiter resolved the "
            "concerns without specific patches.",
            italic=True)
    _add_section_divider(doc)

    # ── 6. Independent Verdict ────────────────────────────────
    _add_heading(doc, "5. Independent Verdict", level=1)
    independent: dict[str, Any] | None = None
    if interaction:
        meta = interaction.get("metadata") or {}
        ind_block = meta.get("independent_review")
        if isinstance(ind_block, dict):
            independent = ind_block
    if independent:
        verdict = independent.get("verdict") or "—"
        reasoning = (independent.get("overall_reasoning")
                     or independent.get("reasoning") or "")
        _add_label_value(doc, "Verdict", str(verdict))
        if reasoning:
            _add_para(doc, str(reasoning))
    else:
        _add_para(
            doc,
            "No independent verdict was recorded for this review.",
            italic=True)

    # ── Footer ────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run(
        "Forest Capital — Forest Council adversarial review report. "
        "AI DRAFT — REQUIRES HUMAN REVIEW.")
    f_run.italic = True
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Page margins -- 1in everywhere; matches the editor export.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
