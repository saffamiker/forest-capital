"""
tools/script_docx.py

Renders a presentation_script editor draft to a Word (.docx) document —
either the master script (every speaker) or one speaker's individual
script (their slides only).

The draft's content_json is a TipTap document (see
tools/script_generation.script_to_tiptap): H2 headings are slide /
section headers, H3 headings carry the speaker label ("Speaker: Name"),
paragraphs are spoken delivery text, and blockquotes are the
slide-to-slide transitions. This builder groups those nodes into
sections and lays them out — slide header, a colour-coded SPEAKER
label, the delivery text, and italic indented transitions.

Each speaker is given a stable colour for the run of the document, so a
reader can scan the master script for one presenter's parts.
"""
from __future__ import annotations

import re
from datetime import date
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from tools.academic_docx import _add_page_number, _set_run_color
from tools.speaker_colours import get_speaker_colour

_BODY_FONT = "Times New Roman"
_AI_DRAFT_BANNER = "AI DRAFT — REQUIRES HUMAN REVIEW"

_SPEAKER_RE = re.compile(r"Speaker:\s*(.+)", re.IGNORECASE)


def _node_text(node: Any) -> str:
    """The concatenated plain text of a TipTap node and its descendants."""
    if not isinstance(node, dict):
        return ""
    if node.get("text"):
        return str(node["text"])
    return "".join(_node_text(c) for c in (node.get("content") or []))


class _Section:
    """One script section — a slide (or OPENING / CLOSING) and its blocks."""

    def __init__(self, heading: str) -> None:
        self.heading = heading
        self.speaker: str | None = None
        # (kind, text) — kind is 'para' or 'transition'.
        self.blocks: list[tuple[str, str]] = []


def parse_sections(content_json: Any) -> list[_Section]:
    """Groups a script's TipTap nodes into sections at each H2 heading."""
    nodes = (content_json.get("content", [])
             if isinstance(content_json, dict) else [])
    sections: list[_Section] = []
    current: _Section | None = None
    for node in nodes:
        if not isinstance(node, dict):
            continue
        ntype = node.get("type")
        level = (node.get("attrs") or {}).get("level", 1)
        text = _node_text(node).strip()
        if ntype == "heading" and level <= 2:
            current = _Section(text or "Section")
            sections.append(current)
        elif ntype == "heading" and level == 3:
            if current is not None:
                m = _SPEAKER_RE.search(text)
                current.speaker = (m.group(1).strip() if m else text) or None
        elif ntype in ("paragraph", "blockquote") and text:
            if current is None:
                current = _Section("")
                sections.append(current)
            kind = "transition" if ntype == "blockquote" else "para"
            current.blocks.append((kind, text))
    return sections


def script_speakers(content_json: Any) -> list[str]:
    """Unique speaker names in the script, in first-seen order."""
    out: list[str] = []
    for sec in parse_sections(content_json):
        if sec.speaker and sec.speaker not in out:
            out.append(sec.speaker)
    return out


def _is_slide_section(sec: _Section) -> bool:
    """A slide section carries a speaker; OPENING / CLOSING do not."""
    return sec.speaker is not None


def build_script_docx(
    draft: dict[str, Any], speaker: str | None = None,
) -> bytes:
    """
    Renders a presentation_script draft to .docx.

    speaker=None → the master script (every section, every speaker).
    speaker set  → that speaker's individual script — only the slide
    sections assigned to them, with a per-page header naming them.
    """
    content_json = draft.get("content_json") or {}
    sections = parse_sections(content_json)
    speakers = script_speakers(content_json)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1)
    sec.top_margin = sec.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Page header — the speaker's name for an individual script, the AI
    # DRAFT banner for the master.
    header_para = sec.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head_run = header_para.add_run(
        f"{speaker} — Forest Capital Presentation Script"
        if speaker else _AI_DRAFT_BANNER)
    head_run.bold = True
    head_run.font.name = _BODY_FONT
    head_run.font.size = Pt(10)
    _set_run_color(head_run, "#b45309")

    # Page number — centred in the footer.
    footer_para = sec.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_para)

    # ── Title page ──
    for i, line in enumerate([
        "Presentation Script",
        f"Individual Script — {speaker}" if speaker else "Master Script",
        "Forest Capital — FNA 670",
        "Queens University Charlotte",
        date.today().strftime("%B %d, %Y"),
    ]):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = para.add_run(line)
        run.font.name = _BODY_FONT
        run.bold = i == 0
        run.font.size = Pt(22 if i == 0 else 12)
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brun = banner.add_run(_AI_DRAFT_BANNER)
    brun.bold = True
    brun.font.name = _BODY_FONT
    brun.font.size = Pt(11)
    _set_run_color(brun, "#b45309")

    # ── Sections ──
    for section in sections:
        # An individual export carries only that speaker's slide sections.
        if speaker is not None:
            if not _is_slide_section(section) or section.speaker != speaker:
                continue

        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(14)
        hrun = head.add_run(section.heading or "Section")
        hrun.bold = True
        hrun.font.name = _BODY_FONT
        hrun.font.size = Pt(15)

        if section.speaker:
            sp = doc.add_paragraph()
            sprun = sp.add_run(f"SPEAKER: {section.speaker}")
            sprun.bold = True
            sprun.font.name = _BODY_FONT
            sprun.font.size = Pt(12)
            _set_run_color(sprun, get_speaker_colour(section.speaker, speakers))

        for kind, text in section.blocks:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = _BODY_FONT
            run.font.size = Pt(12)
            if kind == "transition":
                run.italic = True
                para.paragraph_format.left_indent = Inches(0.5)

    if not sections:
        doc.add_paragraph("[DATA PENDING] — the script draft is empty.")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
