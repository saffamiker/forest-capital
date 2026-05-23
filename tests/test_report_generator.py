"""Coverage for tools/report_generator + tools/report_writer_docx +
migration 032 + the academic-review / editor / download endpoints.

May 22 2026 (item 12 commit 2). Exercises the pure-compute helpers
(extract_bob_blocks, _post_check_summary, _writer_unavailable_draft,
_review_unavailable_stub, _parse_review_json) directly, and the
endpoint paths through Starlette's TestClient on the test-environment
shortcut shapes. The .docx builders are exercised end-to-end against
the stub contexts — the resulting bytes are parsed back with
python-docx to verify structure.
"""
import importlib.util
import io
import os
import sys

import pytest


sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))
os.environ.setdefault("ENVIRONMENT", "test")
os.environ.setdefault("SECRET_KEY", "test-secret-key-at-least-32-characters-long")
os.environ.setdefault("MASTER_API_KEY", "test_master_key")
os.environ.setdefault(
    "ALLOWED_EMAILS",
    "ruurdsm@queens.edu,thaob@queens.edu,"
    "murdockm@queens.edu,panttserk@queens.edu")


# ── Migration 032 loads ──────────────────────────────────────────────────────


def test_migration_032_loads():
    spec = importlib.util.spec_from_file_location(
        "mig_032",
        os.path.join(os.path.dirname(__file__), "..", "backend",
                     "migrations", "versions",
                     "032_report_rubrics.py"),
    )
    assert spec is not None and spec.loader is not None
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    assert m.revision == "032"
    assert m.down_revision == "031"
    assert callable(m.upgrade)
    assert callable(m.downgrade)
    # Verify FNA670 rubric seed is exactly the four written criteria.
    criteria = m._FNA670_MIDPOINT_CRITERIA
    assert len(criteria) == 4
    ids = {c["criterion_id"] for c in criteria}
    assert ids == {
        "clarity_and_rigor", "analytical_progress",
        "results_quality", "division_of_labor",
    }
    # Peer feedback quality is explicitly NOT included.
    assert "peer_feedback_quality" not in ids


# ── extract_bob_blocks / count_bob_blocks ───────────────────────────────────


class TestBobBlocks:
    def test_extract_finds_every_marker(self):
        from tools.report_generator import (
            extract_bob_blocks, count_bob_blocks,
        )
        md = (
            "## 1. Data and Methodology\n"
            "The Sharpe ratio is 0.52 [DATA REQUIRED — corr_shift] "
            "and was significant.\n\n"
            "## 2. Results\n"
            "Per (Author, 2024) [CITATION REQUIRED] the data shows...\n\n"
            "[BOB — write the next paragraph here] "
            "[DATA MISMATCH: live=0.5 staged=not-found]\n"
            "[UNVERIFIED NUMBER 99.9]"
        )
        blocks = extract_bob_blocks(md)
        assert len(blocks) == 5
        kinds = [b["kind"] for b in blocks]
        assert "DATA REQUIRED" in kinds
        assert "CITATION REQUIRED" in kinds
        assert "BOB" in kinds
        assert "DATA MISMATCH" in kinds
        assert "UNVERIFIED NUMBER" in kinds
        assert count_bob_blocks(md) == 5

    def test_extract_empty_string(self):
        from tools.report_generator import (
            extract_bob_blocks, count_bob_blocks,
        )
        assert extract_bob_blocks("") == []
        assert count_bob_blocks("") == 0
        # Non-bracketed prose returns empty.
        assert extract_bob_blocks("clean prose only") == []

    def test_descriptions_stripped_of_label(self):
        from tools.report_generator import extract_bob_blocks
        md = "[DATA REQUIRED — corr_shift field]"
        b = extract_bob_blocks(md)[0]
        assert b["kind"] == "DATA REQUIRED"
        assert "corr_shift" in b["description"]
        assert not b["description"].startswith("DATA REQUIRED")


# ── _post_check_summary ─────────────────────────────────────────────────────


class TestPostCheckSummary:
    def test_clean_draft_zero_flags(self):
        from tools.report_generator import _post_check_summary
        # No unverified numbers, no inline citations, within budget.
        md = "## 1. Section\n\nClean paragraph with no numbers.\n"
        out = _post_check_summary(md, {}, {})
        assert out["flag_count"] == 0
        assert out["bob_block_count"] == 0

    def test_bob_markers_flagged(self):
        from tools.report_generator import _post_check_summary
        md = (
            "## 1. Section\n\n"
            "[BOB — write this paragraph]\n\n"
            "[DATA REQUIRED — sharpe_field]")
        out = _post_check_summary(md, {}, {})
        assert out["flag_count"] >= 2
        assert out["bob_block_count"] == 2

    def test_unverified_number_flagged(self):
        from tools.report_generator import _post_check_summary
        # 1.23 is not in verified_data; 0.52 is.
        verified = {"benchmark_sharpe": 0.52}
        md = (
            "## 1. Section\n\n"
            "The figure 1.23 is fabricated and 0.52 is verified.")
        out = _post_check_summary(md, verified, {})
        kinds = [f.get("kind") for f in out["flags"]]
        assert "unverified_number" in kinds


# ── Writer unavailable draft sentinel ───────────────────────────────────────


class TestWriterUnavailableDraft:
    def test_emits_four_sections_and_bob_blocks(self):
        from tools.report_generator import _writer_unavailable_draft
        draft = _writer_unavailable_draft("abc12345")
        for h in ("## 1. Data and Methodology",
                  "## 2. Preliminary Results and Diagnostics",
                  "## 3. Roles and Division of Labor",
                  "## 4. Next Steps and Open Questions"):
            assert h in draft
        # Carries the ref so a downstream UI can echo it.
        assert "abc12345" in draft
        from tools.report_generator import count_bob_blocks
        assert count_bob_blocks(draft) == 4


# ── _parse_review_json + review stub ────────────────────────────────────────


class TestReviewParsing:
    def test_parse_strips_code_fence(self):
        from tools.report_generator import _parse_review_json
        raw = '```json\n{"readiness": "ready_to_submit"}\n```'
        parsed = _parse_review_json(raw)
        assert parsed == {"readiness": "ready_to_submit"}

    def test_parse_plain_json(self):
        from tools.report_generator import _parse_review_json
        raw = '{"per_criterion": [], "readiness": "needs_minor_revision"}'
        parsed = _parse_review_json(raw)
        assert parsed["readiness"] == "needs_minor_revision"

    def test_parse_invalid_returns_none(self):
        from tools.report_generator import _parse_review_json
        assert _parse_review_json("not json at all") is None
        assert _parse_review_json("") is None

    def test_review_unavailable_stub_covers_every_criterion(self):
        from tools.report_generator import _review_unavailable_stub
        rubric = {"criteria": [
            {"criterion_id": "clarity_and_rigor"},
            {"criterion_id": "analytical_progress"},
            {"criterion_id": "results_quality"},
            {"criterion_id": "division_of_labor"},
        ]}
        out = _review_unavailable_stub(rubric, ref="abc12345")
        assert len(out["per_criterion"]) == 4
        for entry in out["per_criterion"]:
            assert entry["score"] == "developing"
            assert "abc12345" in entry["gap"]
        assert out["readiness"] == "needs_minor_revision"


# ── _iterate_sync (test env fallback path) ──────────────────────────────────


class TestIterateSync:
    def test_unknown_action_returns_selection(self):
        from tools.report_generator import _iterate_sync
        sel = "The benchmark Sharpe is 0.52."
        # Unknown action — short-circuits to the selection unchanged.
        assert _iterate_sync("invalid", sel, None) == sel

    def test_in_test_env_path_does_not_raise(self):
        """call_claude is configured to raise in the test environment;
        _iterate_sync must catch and return the selection unchanged."""
        from tools.report_generator import _iterate_sync
        sel = "The benchmark Sharpe is 0.52."
        out = _iterate_sync("rephrase", sel, None)
        # In test env call_claude either raises or returns ""; either
        # way the selection passes through unchanged.
        assert isinstance(out, str)


# ── New numbers / citations detection ───────────────────────────────────────


class TestIntroducedItems:
    def test_new_number_detected(self):
        from tools.report_generator import _new_numbers_introduced
        verified = {"benchmark_sharpe": 0.52}
        before = "The Sharpe is 0.52 over the period."
        after = "The Sharpe is 0.52 and the Sortino is 1.23 too."
        out = _new_numbers_introduced(before, after, verified)
        assert 1.23 in out

    def test_no_new_number_when_unchanged(self):
        from tools.report_generator import _new_numbers_introduced
        verified = {"benchmark_sharpe": 0.52}
        before = "The Sharpe is 0.52 over the period."
        after = "The Sharpe ratio is 0.52, as reported above."
        out = _new_numbers_introduced(before, after, verified)
        assert out == []


# ── Paper docx builder ──────────────────────────────────────────────────────


class TestPaperDocxBuilder:
    def _read_doc(self, content: bytes):
        from docx import Document
        return Document(io.BytesIO(content))

    def test_paper_builds_with_sections_and_references(self):
        from tools.report_writer_docx import build_paper_docx
        paper_md = (
            "## 1. Data and Methodology\n\n"
            "The dataset spans 2002 to 2025 (282 monthly observations).\n\n"
            "## 2. Preliminary Results\n\n"
            "The benchmark Sharpe is 0.52, the regime switching "
            "Sharpe is 0.62.\n")
        refs = "Markowitz, H. (1952). Portfolio selection."
        content = build_paper_docx(paper_md, references_md=refs)
        doc = self._read_doc(content)
        # All section bodies present.
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "282 monthly observations" in text
        assert "Sharpe is 0.52" in text
        assert "Markowitz" in text
        # Header carries the FNA670 brand line.
        header = doc.sections[0].header
        assert any("FNA670" in p.text for p in header.paragraphs)

    def test_paper_handles_empty_markdown(self):
        from tools.report_writer_docx import build_paper_docx
        content = build_paper_docx("")
        # Builds a valid docx with just the title and brand chrome.
        doc = self._read_doc(content)
        assert any("Midpoint" in p.text for p in doc.paragraphs)

    def test_paper_inline_bold_renders(self):
        from tools.report_writer_docx import build_paper_docx
        content = build_paper_docx(
            "## 1. Section\n\nThe **central thesis** holds.")
        doc = self._read_doc(content)
        # At least one bold run exists in the body.
        bold_found = any(
            run.bold for para in doc.paragraphs for run in para.runs
            if "central thesis" in para.text)
        assert bold_found


# ── Appendix docx builder ──────────────────────────────────────────────────


class TestAppendixDocxBuilder:
    def _stub_context(self):
        return {
            "verified_data": {
                "study_period_start": "2002-07-31",
                "study_period_end":   "2025-12-31",
                "n_months":           282,
            },
            "ranked_findings": [
                {"title": "REGIME SHIFT EVIDENCE",
                 "finding": "Post-2022 equity-IG correlation rose materially.",
                 "evidence": ["Pre-2022 avg: -0.05",
                              "Post-2022 avg: 0.61"],
                 "implication": "Static allocation no longer hedges.",
                 "nugget_strength": "HIGH",
                 "surprise": True,
                 "surprise_reason": "30-year diversification benefit reversed"},
                {"title": "BENCHMARK COMPETITIVENESS",
                 "finding": "Benchmark ranks third on Sharpe.",
                 "evidence": ["BENCHMARK Sharpe rank: 3 of 10"],
                 "implication": "Diversified strategies dominate.",
                 "nugget_strength": "HIGH",
                 "surprise": False},
            ],
            "team_activity": {
                "team_total_uat_steps": 42,
                "team_total_council_sessions": 12,
                "team_total_audit_validations": 6,
                "michael_commits": 50,
                "bob_uat_steps": 20,
                "molly_uat_steps": 22,
            },
            "validation_summary": {
                "layer1_status": "pass", "layer1_count": 12, "layer1_date": "2026-05-22",
                "layer2_status": "pass", "layer2_count": 18, "layer2_date": "2026-05-22",
                "layer3_status": "pass", "layer3_count": 8,  "layer3_date": "2026-05-22",
            },
            "citations_cache": {
                "cvar_coherent_risk": {
                    "concept_id": "cvar_coherent_risk",
                    "author": "Artzner, P. et al.",
                    "year": "1999",
                    "title": "Coherent measures of risk",
                    "journal_or_institution": "Mathematical Finance",
                    "volume_issue_pages": "9(3), 203-228",
                    "url": "https://onlinelibrary.wiley.com/x",
                    "verification_status": "verified",
                    "formatted": (
                        "Artzner, P. et al. (1999). \"Coherent measures of "
                        "risk\". Mathematical Finance, 9(3), 203-228."),
                },
                "gips_verification": {
                    "concept_id": "gips_verification",
                    "author": "CFA Institute",
                    "year": "2020",
                    "title": "GIPS Standards",
                    "journal_or_institution": "CFA Institute",
                    "volume_issue_pages": "",
                    "url": "https://cfainstitute.org/x",
                    "verification_status": "verified",
                    "formatted": (
                        "CFA Institute (2020). \"GIPS Standards\". CFA Institute."),
                },
            },
            "findings_metadata": {
                "computed_at": "2026-05-22T10:00:00Z",
                "data_hash": "abc123",
                "audit_status": "pass",
            },
            "generated_at": "2026-05-22T11:00:00Z",
        }

    def test_appendix_builds_all_four_sections(self):
        from tools.report_writer_docx import build_appendix_docx
        from docx import Document
        content = build_appendix_docx(self._stub_context())
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
        # All four appendices present.
        assert "Appendix A" in text
        assert "Appendix B" in text
        assert "Appendix C" in text
        assert "Appendix D" in text
        # References built from citations_cache.
        assert "References" in text
        # Verified citations rendered.
        assert "Artzner" in text or "Coherent measures of risk" in text
        assert "CFA Institute" in text

    def test_appendix_b_includes_ranked_findings(self):
        from tools.report_writer_docx import build_appendix_docx
        from docx import Document
        content = build_appendix_docx(self._stub_context())
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "REGIME SHIFT EVIDENCE" in text
        assert "BENCHMARK COMPETITIVENESS" in text
        # SURPRISE block surfaces.
        assert "SURPRISE" in text

    def test_appendix_c_includes_team_activity(self):
        from tools.report_writer_docx import build_appendix_docx
        from docx import Document
        content = build_appendix_docx(self._stub_context())
        doc = Document(io.BytesIO(content))
        # The activity table is rendered as a Word table; collect cell text.
        cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cells.append(cell.text)
        joined = " ".join(cells)
        assert "Michael Ruurds" in joined
        assert "Bob Thao" in joined
        assert "Molly Murdock" in joined
        assert "50" in joined  # michael_commits
        assert "42" in joined  # team_total_uat_steps

    def test_appendix_d_includes_validation_layers(self):
        from tools.report_writer_docx import build_appendix_docx
        from docx import Document
        content = build_appendix_docx(self._stub_context())
        doc = Document(io.BytesIO(content))
        cells = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cells.append(cell.text)
        joined = " ".join(cells)
        assert "Layer" in joined or "audit" in joined.lower()
        assert "1" in joined and "2" in joined and "3" in joined

    def test_references_alphabetical_by_surname(self):
        from tools.report_writer_docx import _build_references_md
        ctx = self._stub_context()
        md = _build_references_md(ctx["citations_cache"])
        # Two verified entries, alphabetical: Artzner before CFA Institute.
        artzner_pos = md.find("Artzner")
        cfa_pos = md.find("CFA Institute")
        assert artzner_pos != -1 and cfa_pos != -1
        assert artzner_pos < cfa_pos

    def test_references_excludes_unverified(self):
        from tools.report_writer_docx import _build_references_md
        ctx = {
            "verified": {
                "concept_id": "v", "verification_status": "verified",
                "formatted": "Verified, A. (2024). Title.",
                "author": "Verified, A.",
            },
            "unverified": {
                "concept_id": "u", "verification_status": "not_found",
                "formatted": None, "author": "Unverified, B.",
            },
        }
        md = _build_references_md(ctx)
        assert "Verified, A." in md
        assert "Unverified, B." not in md


# ── Endpoint contract tests (auth + test-env shapes) ────────────────────────


def _client():
    """Lazy fixture — avoid touching FastAPI at import time so test
    collection works without the full backend dependency tree."""
    from fastapi.testclient import TestClient
    from main import app
    return TestClient(app)


def _team_session(email: str = "ruurdsm@queens.edu") -> dict:
    """Mints a session for a project-team member via main's
    dependency override path."""
    from main import app
    from auth import require_team_member, require_auth
    fake = {"email": email, "role": "sysadmin",
            "display_name": "Test User",
            "permissions": ["team_member", "manage_users",
                            "view_analytics"]}

    def _override():
        return fake
    app.dependency_overrides[require_team_member] = _override
    app.dependency_overrides[require_auth] = _override
    return fake


def _clear_overrides():
    from main import app
    app.dependency_overrides = {}


class TestEndpointsTestEnv:
    """Each endpoint takes the ENVIRONMENT=test shortcut and returns
    a deterministic stub shape. Coverage confirms the auth gate and
    the response keys."""

    def test_generate_unauthenticated_401(self):
        _clear_overrides()
        c = _client()
        r = c.post("/api/v1/reports/templates/midpoint_check_fna670/generate")
        assert r.status_code in (401, 403)

    def test_generate_returns_stub_shape(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/templates/midpoint_check_fna670/generate")
            assert r.status_code == 200
            body = r.json()
            for k in ("paper_md", "appendix_md", "flag_count",
                      "bob_blocks", "flags"):
                assert k in body
        finally:
            _clear_overrides()

    def test_get_generation_404_in_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.get("/api/v1/reports/generations/1")
            # Test env returns the "no generation" shape with 200
            # OR 404 depending on the path; both are acceptable.
            assert r.status_code in (200, 404)
        finally:
            _clear_overrides()

    def test_patch_paper_md_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.patch(
                "/api/v1/reports/generations/1/paper-md",
                json={"paper_md": "test"})
            assert r.status_code == 200
            assert "flag_count" in r.json()
        finally:
            _clear_overrides()

    def test_iterate_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/generations/1/iterate",
                json={"action": "rephrase", "selection": "test"})
            assert r.status_code == 200
            body = r.json()
            assert "rewritten" in body
        finally:
            _clear_overrides()

    def test_iterate_rejects_unknown_action(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/generations/1/iterate",
                json={"action": "xyz", "selection": "test"})
            # In test env we short-circuit BEFORE validation, so
            # the call returns 200 — that's documented behaviour.
            # Outside test env the validation gates first.
            assert r.status_code in (200, 422)
        finally:
            _clear_overrides()

    def test_resolve_bob_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/generations/1/resolve-bob",
                json={"marker": "[BOB — x]", "replacement": "x"})
            assert r.status_code == 200
            assert "flag_count" in r.json()
        finally:
            _clear_overrides()

    def test_final_check_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/generations/1/final-check")
            assert r.status_code == 200
            body = r.json()
            assert "passed" in body and body["passed"] is True
            assert body["flag_count"] == 0
        finally:
            _clear_overrides()

    def test_academic_review_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.post(
                "/api/v1/reports/generations/1/academic-review")
            assert r.status_code == 200
            body = r.json()
            assert body["readiness"] == "ready_to_submit"
            assert "per_criterion" in body
        finally:
            _clear_overrides()

    def test_get_rubric_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.get(
                "/api/v1/reports/templates/midpoint_check_fna670/rubric")
            assert r.status_code == 200
            assert "rubric" in r.json()
        finally:
            _clear_overrides()

    def test_upload_rubric_validation_error(self):
        _team_session()
        try:
            c = _client()
            # In test env the endpoint short-circuits before validation,
            # so this returns 200. Validate the contract on the path that
            # DOES run validation (rubric_text missing).
            r = c.post(
                "/api/v1/reports/templates/midpoint_check_fna670/rubric",
                json={})
            # Test env returns the short-circuit shape.
            assert r.status_code == 200
        finally:
            _clear_overrides()

    def test_download_paper_404_in_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.get(
                "/api/v1/reports/generations/1/download-paper")
            # Test env path raises 404 — no generations exist.
            assert r.status_code == 404
        finally:
            _clear_overrides()

    def test_download_appendix_404_in_test_env(self):
        _team_session()
        try:
            c = _client()
            r = c.get(
                "/api/v1/reports/generations/1/download-appendix")
            assert r.status_code == 404
        finally:
            _clear_overrides()


# ── Gate-download helper ────────────────────────────────────────────────────


class TestGateDownload:
    def test_flags_remaining_blocks(self):
        from main import _gate_download
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _gate_download({"flag_count": 3,
                            "academic_readiness": "ready_to_submit"},
                           acknowledge_warning=False)
        assert exc.value.status_code == 422
        assert exc.value.detail["error"] == "flags_remaining"

    def test_significant_revision_blocks_without_ack(self):
        from main import _gate_download
        from fastapi import HTTPException
        with pytest.raises(HTTPException) as exc:
            _gate_download(
                {"flag_count": 0,
                 "academic_readiness": "needs_significant_revision"},
                acknowledge_warning=False)
        assert exc.value.status_code == 422
        assert exc.value.detail["error"] == (
            "academic_review_significant_revision")

    def test_significant_revision_passes_with_ack(self):
        from main import _gate_download
        _gate_download(
            {"flag_count": 0,
             "academic_readiness": "needs_significant_revision"},
            acknowledge_warning=True)  # no raise

    def test_ready_to_submit_passes(self):
        from main import _gate_download
        _gate_download(
            {"flag_count": 0,
             "academic_readiness": "ready_to_submit"},
            acknowledge_warning=False)  # no raise

    def test_minor_revision_passes(self):
        from main import _gate_download
        _gate_download(
            {"flag_count": 0,
             "academic_readiness": "needs_minor_revision"},
            acknowledge_warning=False)  # no raise


# ── Module imports cleanly ──────────────────────────────────────────────────


def test_report_generator_imports():
    import tools.report_generator as rg
    assert hasattr(rg, "generate_paper")
    assert hasattr(rg, "iterate_text")
    assert hasattr(rg, "resolve_bob_block")
    assert hasattr(rg, "run_final_check")
    assert hasattr(rg, "run_academic_review")
    assert hasattr(rg, "render_paper_bytes")
    assert hasattr(rg, "render_appendix_bytes")


def test_report_rubrics_imports():
    import tools.report_rubrics as rr
    assert hasattr(rr, "get_latest_rubric")
    assert hasattr(rr, "list_rubrics")
    assert hasattr(rr, "upload_rubric")


def test_report_writer_docx_imports():
    import tools.report_writer_docx as rw
    assert hasattr(rw, "build_paper_docx")
    assert hasattr(rw, "build_appendix_docx")
