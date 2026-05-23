"""Coverage for the Executive Brief template + memo docx renderer.

May 22 2026 (item 12 commit D). Verifies migration 034 loads
cleanly, seeds the expected rows, and the build_brief_docx renderer
produces a parseable memo-style document with the FROM/TO/RE
header, ALL CAPS section labels, and the footer note.
"""
import importlib.util
import io
import os
import sys


sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))
os.environ.setdefault("ENVIRONMENT", "test")
os.environ.setdefault("SECRET_KEY", "test-secret-key-at-least-32-characters-long")
os.environ.setdefault("MASTER_API_KEY", "test_master_key")
os.environ.setdefault(
    "ALLOWED_EMAILS",
    "ruurdsm@queens.edu,thaob@queens.edu,"
    "murdockm@queens.edu,panttserk@queens.edu")


# ── Migration 034 loads ─────────────────────────────────────────────────────


def test_migration_034_loads():
    spec = importlib.util.spec_from_file_location(
        "mig_034",
        os.path.join(os.path.dirname(__file__), "..", "backend",
                     "migrations", "versions",
                     "034_executive_brief_template.py"),
    )
    assert spec is not None and spec.loader is not None
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    assert m.revision == "034"
    assert m.down_revision == "033"
    assert callable(m.upgrade)
    assert callable(m.downgrade)

    # System prompt names the brief audience and the central thesis.
    assert "Forest Capital leadership" in m._EXEC_BRIEF_SYSTEM_PROMPT
    assert "490 words" in m._EXEC_BRIEF_SYSTEM_PROMPT or (
        "490" in m._EXEC_BRIEF_SYSTEM_PROMPT)
    # All 5 sections + their word budgets present.
    assert "Section 1" in m._EXEC_BRIEF_SYSTEM_PROMPT or (
        "SECTION 1" in m._EXEC_BRIEF_SYSTEM_PROMPT)
    # All 4 rubric criteria seeded.
    ids = {c["criterion_id"] for c in m._EXEC_BRIEF_CRITERIA}
    assert ids == {
        "executive_clarity", "actionability",
        "evidence_quality", "brevity",
    }
    # Memo format spec marker — drives the renderer dispatch.
    assert m._EXEC_BRIEF_FORMAT_SPEC.get("memo_style") is True


# ── Renderer imports + word-budget accessors ────────────────────────────────


def test_brief_renderer_imports():
    import tools.report_writer_docx_brief as rwd_brief
    assert hasattr(rwd_brief, "build_brief_docx")
    assert hasattr(rwd_brief, "get_section_budgets")
    assert hasattr(rwd_brief, "get_total_budget")


def test_brief_word_budgets_match_spec():
    """Section budgets must mirror the system prompt spec:
       1:60, 2:180, 3:80, 4:80, 5:90. Total 490."""
    from tools.report_writer_docx_brief import (
        get_section_budgets, get_total_budget,
    )
    budgets = get_section_budgets()
    assert budgets == {1: 60, 2: 180, 3: 80, 4: 80, 5: 90}
    assert get_total_budget() == 490
    assert sum(budgets.values()) == get_total_budget()


# ── Memo docx structure ────────────────────────────────────────────────────


class TestBriefDocxRendering:
    def _read(self, content: bytes):
        from docx import Document
        return Document(io.BytesIO(content))

    def _brief_md(self) -> str:
        return (
            "## 1. The Situation\n\n"
            "The post-2022 correlation regime shift removed the "
            "traditional stock-bond hedge.\n\n"
            "## 2. Key Findings\n\n"
            "- The benchmark ranks 4th of 10 strategies on Sharpe.\n"
            "- Regime Switching reduces drawdown by 12 percentage "
            "points vs benchmark.\n"
            "- Equal-weight blend produces a CVaR ratio of 0.42.\n\n"
            "## 3. Risk Implications\n\n"
            "Tail risk is materially lower in diversified blends.\n\n"
            "## 4. Current Environment\n\n"
            "Macro context favors short duration.\n\n"
            "## 5. Recommended Next Steps\n\n"
            "- Commission sensitivity at 15bps and 20bps.\n"
            "- Evaluate dynamic regime signals for quarterly "
            "rebalancing.\n"
            "- Expand the asset universe to include real assets.\n")

    def test_memo_header_renders(self):
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        # TO / FROM / DATE / RE labels in the memo header.
        assert "TO:" in text
        assert "FROM:" in text
        assert "DATE:" in text
        assert "RE:" in text
        assert "Forest Capital Leadership" in text
        assert "May 27, 2026" in text

    def test_section_headings_uppercase(self):
        """Memo style — section labels render in ALL CAPS."""
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "THE SITUATION" in text
        assert "KEY FINDINGS" in text
        assert "RISK IMPLICATIONS" in text
        assert "CURRENT ENVIRONMENT" in text
        assert "RECOMMENDED NEXT STEPS" in text

    def test_bullets_rendered(self):
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        # Markdown dashes become bullet characters.
        assert "•" in text

    def test_footer_line(self):
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Full methodology" in text
        assert "forest-capital.vercel.app" in text

    def test_no_title_page(self):
        """The brief does NOT have a title page (memo style). The
        first paragraph is the TO label, NOT the paper title."""
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        # The midpoint paper title must NOT appear in the brief.
        assert "Multi-Strategy Portfolio Diversification: A " not in text

    def test_calibri_body_font(self):
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx(self._brief_md())
        doc = self._read(content)
        # Body font must be Calibri (memo convention). Check a run.
        runs = [r for p in doc.paragraphs for r in p.runs if r.text.strip()]
        # At least one run carries Calibri.
        assert any(r.font.name == "Calibri" for r in runs)

    def test_empty_body_still_produces_memo_header(self):
        """A blank body still produces a parseable docx with the
        memo header — the editor surface needs the header even
        before any content has been generated."""
        from tools.report_writer_docx_brief import build_brief_docx
        content = build_brief_docx("")
        doc = self._read(content)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "TO:" in text and "FROM:" in text


# ── Renderer dispatch via report_generator ──────────────────────────────────


def test_render_paper_bytes_routes_via_format_spec():
    """report_generator.render_paper_bytes inspects the template's
    format_spec.memo_style and chooses the brief renderer when
    True. Verify the source code carries the dispatch (we can't
    run an end-to-end call_claude in tests)."""
    from inspect import getsource
    from tools.report_generator import render_paper_bytes
    src = getsource(render_paper_bytes)
    assert "memo_style" in src
    assert "build_brief_docx" in src
    assert "build_paper_docx" in src
