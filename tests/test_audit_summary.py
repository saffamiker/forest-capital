"""
tests/test_audit_summary.py

Workstream D — deterministic audit disclosures for the generated
reports (May 28 2026). Two layers:

  1. The pure functions in tools/audit_summary — gather, summary
     sentence, body paragraph, table-row helpers. Pinned with
     monkeypatch / fixture input.

  2. End-to-end render: build_midpoint_paper and build_executive_brief
     emit the appendix and (in the brief's case) the summary sentence
     and body paragraph when the data bundle carries audit_disclosures.
"""
from __future__ import annotations

import asyncio
import io
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))
os.environ.setdefault("SECRET_KEY", "test-secret-key-at-least-32-characters-long")
os.environ.setdefault("ENVIRONMENT", "test")
os.environ.setdefault("MASTER_API_KEY", "test-master-key")


def _docx_text(blob: bytes) -> str:
    """Pulls the visible body text out of a .docx blob — every body
    paragraph and every cell. Used to assert sections / tables render."""
    from docx import Document  # type: ignore[import]
    doc = Document(io.BytesIO(blob))
    chunks: list[str] = []
    for para in doc.paragraphs:
        chunks.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def _disclosures(**over) -> dict:
    """Fixture bundle in the shape gather_audit_disclosures() returns."""
    base = {
        "available": True,
        "statistical": {
            "present": True, "run_id": 42, "completed_at": "2026-05-25",
            "total": 56, "passed": 50, "warnings": 6, "failures": 0,
            "acknowledged": [{
                "layer": 2, "check_name": "Sharpe verification",
                "metric": "sharpe_ratio", "strategy": "REGIME_SWITCHING",
                "status": "warning", "discrepancy": "0.4%",
                "resolved_by": "bob@queens.edu",
                "resolved_at": "2026-05-25",
                "resolution_note":
                    "ACKDISCLOSURETOKEN — reviewed and accepted.",
                "auto_acknowledged": False,
            }],
        },
        "methodology": {
            "present": True, "verdict": "WARN", "run_at": "2026-05-25",
            "total": 39, "passed": 32, "warnings": 7, "failures": 0,
            "intentional": [{
                "check_id": "P03",
                "check": "Transaction costs applied",
                "description": "Turnover sums |Δw|.",
                "category": "PORTFOLIO_MECHANICS",
                "marked_by": "ruurdsm@queens.edu",
                "marked_at": "2026-05-25",
                "note":
                    "INTENTIONALDISCLOSURETOKEN — capture is intentional.",
            }],
        },
        "generated_at": "2026-05-28T00:00:00+00:00",
    }
    base.update(over)
    return base


def _empty_disclosures() -> dict:
    return {
        "available": False,
        "statistical": {
            "present": False, "run_id": None, "completed_at": None,
            "total": 0, "passed": 0, "warnings": 0, "failures": 0,
            "acknowledged": [],
        },
        "methodology": {
            "present": False, "verdict": None, "run_at": None,
            "total": 0, "passed": 0, "warnings": 0, "failures": 0,
            "intentional": [],
        },
        "generated_at": "2026-05-28T00:00:00+00:00",
    }


# ── Pure functions ───────────────────────────────────────────────────────────

class TestSummarySentence:

    def test_names_audit_runs_and_disclosed_count(self):
        from tools.audit_summary import audit_summary_sentence

        out = audit_summary_sentence(_disclosures())
        # The sentence names both audit totals and the disclosed count.
        assert "56 checks" in out
        assert "39 checks" in out
        # 1 statistical + 1 methodology = 2 disclosed.
        assert "2 warnings" in out
        assert "disclosed in the appendix" in out

    def test_handles_singular_correctly(self):
        from tools.audit_summary import audit_summary_sentence

        d = _disclosures()
        # Drop the methodology disclosure — only 1 acknowledged warning total.
        d["methodology"]["intentional"] = []
        out = audit_summary_sentence(d)
        assert "1 warning has been reviewed" in out
        assert "warnings" not in out.replace("1 warning has", "")

    def test_no_audit_on_record(self):
        from tools.audit_summary import audit_summary_sentence

        out = audit_summary_sentence(_empty_disclosures())
        assert "No platform audit" in out
        assert "Appendix is empty" in out

    def test_zero_disclosures(self):
        from tools.audit_summary import audit_summary_sentence

        d = _disclosures()
        d["statistical"]["acknowledged"] = []
        d["methodology"]["intentional"] = []
        out = audit_summary_sentence(d)
        assert "no disclosures requiring acknowledgement" in out


class TestBodyParagraph:

    def test_names_both_audit_subsystems(self):
        from tools.audit_summary import audit_body_paragraph

        out = audit_body_paragraph(_disclosures())
        # Statistical audit framing.
        assert "claude-opus-4-7" in out
        assert "three layers" in out
        # Methodology framing.
        assert "checklist" in out
        # The Workstream C gate guarantee.
        assert "report-readiness" in out or "refuses to generate" in out

    def test_no_audit_on_record_falls_back(self):
        from tools.audit_summary import audit_body_paragraph

        out = audit_body_paragraph(_empty_disclosures())
        assert "No platform audit was on record" in out

    def test_omits_subsystem_block_when_absent(self):
        from tools.audit_summary import audit_body_paragraph

        d = _disclosures()
        d["methodology"]["present"] = False
        out = audit_body_paragraph(d)
        # Statistical block still present, methodology block gone.
        assert "statistical audit" in out
        # The methodology-specific phrasing should NOT appear.
        assert "methodology review evaluated" not in out


class TestTableRowHelpers:

    def test_statistical_rows_render_check_label_with_strategy(self):
        from tools.audit_summary import acknowledged_statistical_rows

        rows = acknowledged_statistical_rows(_disclosures())
        assert len(rows) == 1
        layer, check, by, when, note = rows[0]
        assert layer == "L2"
        # Strategy is appended to the check name.
        assert "REGIME_SWITCHING" in check
        assert by == "bob@queens.edu"
        assert when == "2026-05-25"
        assert "ACKDISCLOSURETOKEN" in note

    def test_methodology_rows_prefix_check_with_id(self):
        from tools.audit_summary import intentional_methodology_rows

        rows = intentional_methodology_rows(_disclosures())
        assert len(rows) == 1
        check, cat, by, when, note = rows[0]
        assert check.startswith("P03")
        assert cat == "PORTFOLIO_MECHANICS"
        assert by == "ruurdsm@queens.edu"
        assert when == "2026-05-25"
        assert "INTENTIONALDISCLOSURETOKEN" in note

    def test_empty_disclosures_returns_empty_lists(self):
        from tools.audit_summary import (
            acknowledged_statistical_rows, intentional_methodology_rows,
        )

        empty = _empty_disclosures()
        assert acknowledged_statistical_rows(empty) == []
        assert intentional_methodology_rows(empty) == []


# ── gather_audit_disclosures fail-open contract ──────────────────────────────

class TestGatherAuditDisclosures:

    def test_fail_open_with_no_history(self):
        """In the test environment there is no audit history reachable;
        gather still returns the empty-bundle shape rather than raising."""
        from tools.audit_summary import gather_audit_disclosures

        out = asyncio.run(gather_audit_disclosures())
        # available may be True or False depending on test-DB state, but
        # the shape is always complete.
        assert "statistical" in out
        assert "methodology" in out
        for key in ("present", "run_id", "completed_at", "total",
                    "passed", "warnings", "failures", "acknowledged"):
            assert key in out["statistical"]
        for key in ("present", "verdict", "run_at", "total", "passed",
                    "warnings", "failures", "intentional"):
            assert key in out["methodology"]
        assert "generated_at" in out


# ── docx end-to-end render ───────────────────────────────────────────────────

class TestExecutiveBriefAudit:
    """build_executive_brief renders the audit summary sentence in
    Executive Summary, the body paragraph in Methodology Overview, and
    the Audit Disclosure Appendix at the end."""

    def _data(self, **over) -> dict:
        base = {
            "study_period": {"start": "2002-07", "end": "2025-12",
                             "n_months": 282, "ff_factors_end": "2025-12"},
            "regime_conditional": [],
            "summary_statistics": [],
            "drawdown_comparison": [],
            "factor_loadings": [],
            "audit_disclosures": _disclosures(),
        }
        base.update(over)
        return base

    def test_renders_summary_sentence(self):
        from tools.academic_docx import build_executive_brief

        blob = build_executive_brief(self._data(), {})
        text = _docx_text(blob)
        assert "independent statistical audit" in text
        assert "56 checks" in text or "39 checks" in text

    def test_renders_methodology_body_paragraph(self):
        from tools.academic_docx import build_executive_brief

        blob = build_executive_brief(self._data(), {})
        text = _docx_text(blob)
        assert "claude-opus-4-7" in text

    def test_renders_appendix_at_end(self):
        from tools.academic_docx import build_executive_brief

        blob = build_executive_brief(self._data(), {})
        text = _docx_text(blob)
        # The headline appears once for the heading.
        assert text.count("Audit Disclosure Appendix") >= 1
        assert "ACKDISCLOSURETOKEN" in text
        assert "INTENTIONALDISCLOSURETOKEN" in text

    def test_empty_disclosures_still_render_a_summary_line(self):
        from tools.academic_docx import build_executive_brief

        blob = build_executive_brief(
            self._data(audit_disclosures=_empty_disclosures()), {})
        text = _docx_text(blob)
        # The audit_summary_sentence falls back to the no-audit text.
        assert "No platform audit" in text


# ── academic_export wiring ───────────────────────────────────────────────────

class TestAcademicExportBundleIncludesDisclosures:
    """gather_document_data must include the audit_disclosures key so
    the document builders can read it via data["audit_disclosures"]."""

    def test_bundle_has_audit_disclosures_key(self):
        from tools.academic_export import gather_document_data

        bundle = asyncio.run(gather_document_data())
        # In the test environment the key may be None (the early-return
        # path) — but the key MUST exist so the builders' .get() works.
        assert "audit_disclosures" in bundle
