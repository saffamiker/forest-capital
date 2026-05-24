"""tests/test_peer_review.py — Item 7 contract.

Pins the Peer Review Assistant + Thesis Defense Prep flows:
  - the four FNA 670 rubric dimensions surface in the peer-review
    arbiter instructions (a regression that drops one would
    quietly miss a rubric dimension on every review)
  - the three Q&A categories surface in the defense-prep arbiter
    instructions
  - the multi-format extractor handles PDF / MD / DOCX inputs and
    rejects everything else with ValueError
  - the two new endpoints are registered, team-gated, and 422 the
    common malformed inputs (no file, oversized file, bad format)
  - the SSE wire format matches the spec — submission_meta /
    arbiter_chunk / [DONE] for Feature A; draft_meta / arbiter_
    chunk / [DONE] for Feature B; error frames carry a message
  - Defense Prep emits an error frame (not a 500) when the
    caller has no current midpoint_paper draft
"""
from __future__ import annotations

import io
import os
import sys
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest


sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "backend"))
os.environ.setdefault("ENVIRONMENT", "test")
os.environ.setdefault(
    "SECRET_KEY", "test-secret-key-at-least-32-characters-long")
os.environ.setdefault("MASTER_API_KEY", "test_master_key")
os.environ.setdefault(
    "ALLOWED_EMAILS",
    "ruurdsm@queens.edu,thaob@queens.edu,"
    "murdockm@queens.edu,panttserk@queens.edu")


def _client():
    from fastapi.testclient import TestClient
    from main import app
    from auth import generate_session_token
    client = TestClient(app)
    team = {"X-API-Key": generate_session_token("thaob@queens.edu")}
    viewer = {"X-API-Key": generate_session_token(
        "panttserk@queens.edu")}
    return client, team, viewer


# ── 1. Prompt content — rubric dimensions + categories pinned ──────────────


class TestPeerReviewPrompts:

    def test_peer_review_arbiter_names_four_rubric_dimensions(self):
        # The four FNA 670 midpoint rubric dimensions must each be
        # named in the arbiter instructions verbatim — a regression
        # that drops one would quietly omit it from every review
        # script the agent produces.
        from agents.peer_review import (
            _PEER_REVIEW_ARBITER_INSTRUCTIONS,
            PEER_RUBRIC_DIMENSIONS,
        )
        assert len(PEER_RUBRIC_DIMENSIONS) == 4
        for dim in PEER_RUBRIC_DIMENSIONS:
            assert dim in _PEER_REVIEW_ARBITER_INSTRUCTIONS, (
                f"Rubric dimension '{dim}' missing from arbiter "
                f"instructions.")

    def test_peer_review_arbiter_requires_qa_per_dimension(self):
        # Every dimension's section must request 'Suggested
        # questions for Q&A' — the spec is explicit that each
        # reviewer comes with 2-3 Q&A questions per dimension.
        from agents.peer_review import _PEER_REVIEW_ARBITER_INSTRUCTIONS
        count = _PEER_REVIEW_ARBITER_INSTRUCTIONS.count(
            "Suggested questions for Q&A:")
        assert count >= 4, (
            f"Expected at least 4 Suggested-Q&A blocks "
            f"(one per dimension), got {count}.")

    def test_peer_review_arbiter_caps_delivery_time(self):
        from agents.peer_review import _PEER_REVIEW_ARBITER_INSTRUCTIONS
        # The spec is 3-4 minutes at 130 wpm. Pin both numbers
        # so a future copy edit can't silently bump them.
        assert "3-4 minutes" in _PEER_REVIEW_ARBITER_INSTRUCTIONS
        assert "130 words per minute" in _PEER_REVIEW_ARBITER_INSTRUCTIONS \
            or "130 wpm" in _PEER_REVIEW_ARBITER_INSTRUCTIONS

    def test_defense_prep_arbiter_names_three_categories(self):
        from agents.peer_review import (
            _DEFENSE_PREP_ARBITER_INSTRUCTIONS,
            DEFENSE_CATEGORIES,
        )
        assert len(DEFENSE_CATEGORIES) == 3
        for cat in DEFENSE_CATEGORIES:
            assert cat in _DEFENSE_PREP_ARBITER_INSTRUCTIONS, (
                f"Q&A category '{cat}' missing from arbiter "
                f"instructions.")

    def test_defense_prep_arbiter_requires_rehearsal_recommendations(self):
        from agents.peer_review import _DEFENSE_PREP_ARBITER_INSTRUCTIONS
        # The spec asks for a Rehearsal recommendations section
        # that names the three highest-risk questions and assigns
        # them to team members.
        assert "Rehearsal recommendations" in _DEFENSE_PREP_ARBITER_INSTRUCTIONS

    def test_mock_verdicts_are_parseable_markdown(self):
        # The test-env mocks must produce verdicts that have the
        # correct structural anchors (the headers the UI uses to
        # parse rating sections). Without this, the SSE wire-
        # format tests would pass but the frontend would render
        # garbage.
        from agents.peer_review import (
            mock_peer_review_verdict, mock_defense_prep_verdict,
        )
        peer = mock_peer_review_verdict("Test Team")
        assert peer.startswith("**Overall verdict:**")
        assert "### 1. Clarity and rigor" in peer
        assert "### 4. Clear division of labor" in peer

        defense = mock_defense_prep_verdict("Forest Capital")
        assert defense.startswith("**Mock panel — overall readiness:**")
        assert "### 1. Technical / methodological" in defense
        assert "### 3. Governance / practical" in defense


# ── 2. Context block rendering ──────────────────────────────────────────────


class TestContextBlocks:

    def test_peer_review_context_includes_submission_text(self):
        from agents.peer_review import (
            build_peer_review_context_block,
            render_peer_review_context_block,
        )
        ctx = build_peer_review_context_block(
            "Other Team Alpha", "This is the submitted text.")
        s = render_peer_review_context_block(ctx)
        assert "Other Team Alpha" in s
        assert "This is the submitted text." in s
        # All four rubric dimensions named in the block.
        from agents.peer_review import PEER_RUBRIC_DIMENSIONS
        for dim in PEER_RUBRIC_DIMENSIONS:
            assert dim in s

    def test_defense_prep_context_includes_draft_text(self):
        from agents.peer_review import (
            build_defense_prep_context_block,
            render_defense_prep_context_block,
        )
        ctx = build_defense_prep_context_block(
            "Forest Capital", "The team's midpoint draft text.")
        s = render_defense_prep_context_block(ctx)
        assert "Forest Capital" in s
        assert "The team's midpoint draft text." in s


# ── 3. Multi-format text extraction ────────────────────────────────────────


class TestExtractPeerPaperText:

    def test_extracts_markdown(self):
        from agents.peer_review import extract_peer_paper_text
        text = extract_peer_paper_text(
            "paper.md",
            b"# Title\n\nThis is the paper body.")
        assert "Title" in text
        assert "This is the paper body." in text

    def test_rejects_empty_upload(self):
        from agents.peer_review import extract_peer_paper_text
        with pytest.raises(ValueError):
            extract_peer_paper_text("paper.md", b"")

    def test_rejects_oversized_upload(self):
        from agents.peer_review import (
            extract_peer_paper_text, MAX_PEER_PAPER_BYTES,
        )
        big = b"x" * (MAX_PEER_PAPER_BYTES + 1)
        with pytest.raises(ValueError):
            extract_peer_paper_text("paper.md", big)

    def test_rejects_unsupported_extension(self):
        from agents.peer_review import extract_peer_paper_text
        with pytest.raises(ValueError):
            extract_peer_paper_text("paper.txt", b"data")
        with pytest.raises(ValueError):
            extract_peer_paper_text("paper.xyz", b"data")

    def test_extracts_docx(self):
        # Build a one-paragraph .docx in-memory so the test runs
        # without a fixture file. python-docx is already a project
        # dependency for the academic_docx writer.
        from docx import Document
        from agents.peer_review import extract_peer_paper_text
        doc = Document()
        doc.add_paragraph(
            "Section 1. The 2022 correlation regime break.")
        doc.add_paragraph(
            "Equity-IG correlation flipped from -0.05 to +0.61.")
        buf = io.BytesIO()
        doc.save(buf)
        text = extract_peer_paper_text("paper.docx", buf.getvalue())
        assert "2022 correlation regime break" in text
        assert "+0.61" in text

    def test_rejects_empty_docx(self):
        from docx import Document
        from agents.peer_review import extract_peer_paper_text
        doc = Document()  # no paragraphs added
        buf = io.BytesIO()
        doc.save(buf)
        with pytest.raises(ValueError):
            extract_peer_paper_text("empty.docx", buf.getvalue())


# ── 4. Endpoint registration + auth gates ───────────────────────────────────


class TestEndpointsAuth:

    def test_peer_review_route_registered(self):
        from main import app
        paths = {r.path for r in app.routes if hasattr(r, "path")}
        assert "/api/council/peer-review" in paths

    def test_defense_prep_route_registered(self):
        from main import app
        paths = {r.path for r in app.routes if hasattr(r, "path")}
        assert "/api/council/defense-prep" in paths

    def test_peer_review_requires_team_member(self):
        client, team, viewer = _client()
        # Unauthed — 401.
        r = client.post(
            "/api/council/peer-review",
            files={"file": ("p.md", b"hello", "text/markdown")})
        assert r.status_code == 401
        # Viewer — 403.
        r = client.post(
            "/api/council/peer-review",
            headers=viewer,
            files={"file": ("p.md", b"hello", "text/markdown")})
        assert r.status_code == 403

    def test_defense_prep_requires_team_member(self):
        client, team, viewer = _client()
        r = client.post("/api/council/defense-prep")
        assert r.status_code == 401
        r = client.post("/api/council/defense-prep",
                         headers=viewer)
        assert r.status_code == 403


# ── 5. Input-validation 422 paths ───────────────────────────────────────────


class TestPeerReviewValidation:

    def test_empty_file_rejected_422(self):
        client, team, _ = _client()
        r = client.post(
            "/api/council/peer-review",
            headers=team,
            files={"file": ("p.md", b"", "text/markdown")})
        assert r.status_code == 422

    def test_oversize_file_rejected_422(self):
        client, team, _ = _client()
        from agents.peer_review import MAX_PEER_PAPER_BYTES
        big = b"x" * (MAX_PEER_PAPER_BYTES + 1)
        r = client.post(
            "/api/council/peer-review",
            headers=team,
            files={"file": ("p.md", big, "text/markdown")})
        assert r.status_code == 422

    def test_unsupported_extension_rejected_422(self):
        client, team, _ = _client()
        r = client.post(
            "/api/council/peer-review",
            headers=team,
            files={"file": ("p.xyz", b"data", "text/plain")})
        assert r.status_code == 422


# ── 6. SSE wire-format contract (test environment streams the mock) ─────────


def _parse_sse_frames(response_text: str) -> list[dict]:
    """Walks an SSE response body and returns the parsed JSON
    payloads. The literal '[DONE]' sentinel is normalised into a
    dict so a single iteration covers every frame."""
    import json
    frames = []
    for line in response_text.split("\n"):
        line = line.strip()
        if not line.startswith("data:"):
            continue
        payload = line[len("data:"):].strip()
        if payload == "[DONE]":
            frames.append({"type": "__done__"})
            continue
        try:
            frames.append(json.loads(payload))
        except json.JSONDecodeError:
            continue
    return frames


class TestPeerReviewWireFormat:

    def test_peer_review_emits_submission_meta_then_chunks_then_done(self):
        # In test env, run_peer_review_with_harness returns the
        # deterministic mock so this exercises the wire format
        # without a model call.
        client, team, _ = _client()
        r = client.post(
            "/api/council/peer-review",
            headers=team,
            data={"submission_name": "Other Team Alpha"},
            files={"file": ("paper.md",
                            b"# Sample paper\n\nBody text.",
                            "text/markdown")})
        assert r.status_code == 200
        frames = _parse_sse_frames(r.text)
        # First non-empty frame must be submission_meta.
        assert frames, "no SSE frames emitted"
        assert frames[0]["type"] == "submission_meta"
        assert frames[0]["name"] == "Other Team Alpha"
        # At least one arbiter_chunk in the middle.
        chunks = [f for f in frames if f.get("type") == "arbiter_chunk"]
        assert len(chunks) >= 1
        # Final frame is [DONE].
        assert frames[-1]["type"] == "__done__"

    def test_submission_name_defaults_to_filename_stem(self):
        client, team, _ = _client()
        r = client.post(
            "/api/council/peer-review",
            headers=team,
            files={"file": ("MyTeamMidpoint.md",
                            b"# Sample\n\nBody.",
                            "text/markdown")})
        assert r.status_code == 200
        frames = _parse_sse_frames(r.text)
        assert frames[0]["type"] == "submission_meta"
        assert frames[0]["name"] == "MyTeamMidpoint"


class TestDefensePrepWireFormat:

    def test_defense_prep_no_draft_emits_error_frame(self):
        # No editor_drafts row exists for the test user, so the
        # endpoint emits an error frame + [DONE] rather than
        # running the agent against empty input.
        client, team, _ = _client()
        r = client.post("/api/council/defense-prep", headers=team)
        assert r.status_code == 200
        frames = _parse_sse_frames(r.text)
        assert frames, "no SSE frames emitted"
        # Must include an error frame that names the missing draft.
        error_frames = [f for f in frames if f.get("type") == "error"]
        assert len(error_frames) == 1
        assert "draft" in error_frames[0]["message"].lower()
        assert frames[-1]["type"] == "__done__"

    def test_defense_prep_with_draft_streams_chunks(self):
        # Patch get_current_draft so the endpoint sees a draft and
        # runs the (mocked) agent. The patch returns the shape
        # editor_drafts.get_current_draft would.
        client, team, _ = _client()
        fake_draft = {
            "id": 1,
            "title": "Midpoint draft v3",
            "content_text": "The 2022 correlation break in detail.",
            "word_count": 1200,
            "updated_at": "2026-05-23T10:00:00Z",
        }
        with patch(
            "tools.editor_drafts.get_current_draft",
            new=AsyncMock(return_value=fake_draft),
        ):
            r = client.post(
                "/api/council/defense-prep", headers=team)
        assert r.status_code == 200
        frames = _parse_sse_frames(r.text)
        # First frame is draft_meta carrying the title + word count.
        assert frames[0]["type"] == "draft_meta"
        assert frames[0]["title"] == "Midpoint draft v3"
        assert frames[0]["word_count"] == 1200
        # At least one arbiter_chunk in the middle.
        chunks = [f for f in frames if f.get("type") == "arbiter_chunk"]
        assert len(chunks) >= 1
        assert frames[-1]["type"] == "__done__"


# ── 7. Activity-log interaction types registered ───────────────────────────


class TestInteractionTypesRegistered:
    def test_peer_review_and_defense_prep_in_interaction_set(self):
        # The _log_interaction_bg call in the endpoint hardcodes
        # "peer_review" / "defense_prep" as the interaction_type.
        # If these aren't in the allowlist, every successful run
        # would silently drop its activity log entry.
        from tools.activity_log import _INTERACTION_TYPES
        assert "peer_review" in _INTERACTION_TYPES
        assert "defense_prep" in _INTERACTION_TYPES
