"""Pre-submission blockers in the Analytical Appendix (bridge #55).

Three fixes pinned here:

  FIX 1 -- table_bootstrap_ci read the wrong key for the Sharpe column
    (point_estimate vs sharpe), so the Sharpe column always rendered as
    "--" in Section D of the appendix. The producer
    bootstrap_ci_table emits `sharpe`; the consumer must read the same
    name.

  FIX 2 -- the overlaps_benchmark flag was never produced. The
    consumer rendered "--" for every row because the field was
    missing. The producer now compares each row's CI bounds to the
    BENCHMARK row's point Sharpe and emits overlaps_benchmark =
    bool(ci_low <= benchmark_sharpe <= ci_high). When no BENCHMARK row
    is present, the field stays None so the consumer's "--" branch
    fires.

  FIX 3 -- build_editor_docx silently dropped all eight data tables
    when the editor export path ran on an analytical_appendix draft.
    The appendix is table-heavy by design; the export of an editor
    draft must re-inject the eight evidence tables after the prose
    nodes. The builder now accepts an optional appendix_data dict and
    appends the tables via _add_appendix_tables when document_type ==
    "analytical_appendix". Other document types are unaffected -- the
    midpoint paper / executive brief / arbitrary editor draft still
    export with prose only and no table block.
"""
from __future__ import annotations

import io

import numpy as np
import pandas as pd

from tools.academic_export import table_bootstrap_ci
from tools.academic_docx import build_editor_docx
from tools.analytics import bootstrap_ci_table


# ── helpers ──────────────────────────────────────────────────────────────

def _strategy_result(name: str, returns: pd.Series) -> dict:
    pairs = [[d.strftime("%Y-%m-%d"), float(v)]
             for d, v in returns.items()]
    return {"strategy_name": name, "monthly_returns": pairs}


def _synthetic(
    n: int = 60, mean: float = 0.006, std: float = 0.04, seed: int = 17,
) -> pd.Series:
    rng = np.random.default_rng(seed=seed)
    return pd.Series(
        rng.normal(mean, std, n),
        index=pd.date_range("2018-01-31", periods=n, freq="ME"))


# ── FIX 1 -- consumer reads `sharpe`, not `point_estimate` ──────────────

def test_table_bootstrap_ci_renders_sharpe_from_sharpe_key():
    """The consumer must read the producer's `sharpe` field. With the
    previous bug it read point_estimate (absent) and the column
    rendered as '--' for every strategy. Regression on that exact
    branch."""
    rows_in = [
        {"strategy": "VOL_TARGETING", "sharpe": 0.86, "ci_low": 0.42,
         "ci_high": 1.34, "overlaps_benchmark": False},
        {"strategy": "BENCHMARK", "sharpe": 0.43, "ci_low": -0.10,
         "ci_high": 0.91, "overlaps_benchmark": True},
    ]
    headers, rows = table_bootstrap_ci(rows_in)
    assert headers[1] == "Sharpe"
    # Sharpe column NEVER renders as the em-dash for present values.
    assert rows[0][1] != "—"
    assert rows[1][1] != "—"
    # Spot check the formatted output -- format_metric on a sharpe_ratio
    # carries two decimal places.
    assert "0.86" in rows[0][1]
    assert "0.43" in rows[1][1]


def test_table_bootstrap_ci_renders_dash_when_sharpe_absent():
    """Missing sharpe key still renders cleanly as '--' rather than a
    NoneType crash. Defensive read on the consumer."""
    rows_in = [{"strategy": "X", "ci_low": 0.1, "ci_high": 0.5}]
    _, rows = table_bootstrap_ci(rows_in)
    assert rows[0][1] == "—"


# ── FIX 2 -- producer emits overlaps_benchmark with the right key ──────

def test_producer_emits_overlaps_benchmark_when_benchmark_row_present():
    """The producer adds the bool field on every row, computed against
    the BENCHMARK row's point Sharpe. Section D's 'Overlaps benchmark'
    column reads this directly."""
    strategies = {
        "BENCHMARK": _strategy_result("BENCHMARK", _synthetic(n=60, seed=1)),
        "ALPHA":     _strategy_result("ALPHA",     _synthetic(n=60, seed=2)),
        "BETA":      _strategy_result("BETA",      _synthetic(n=60, seed=3)),
    }
    rows = bootstrap_ci_table(strategies)
    assert all("overlaps_benchmark" in r for r in rows)
    # Every value is an actual bool -- not None, not a string.
    assert all(isinstance(r["overlaps_benchmark"], bool) for r in rows)
    # The BENCHMARK row's own CI trivially brackets its own point.
    bench = next(r for r in rows if r["strategy"] == "BENCHMARK")
    assert bench["overlaps_benchmark"] is True
    # The bracket logic agrees with the row's own bounds.
    for r in rows:
        if r["strategy"] != "BENCHMARK":
            in_band = r["ci_low"] <= bench["sharpe"] <= r["ci_high"]
            assert r["overlaps_benchmark"] is in_band


def test_producer_emits_none_when_no_benchmark_row():
    """Without a BENCHMARK row the producer cannot compute the bracket
    flag. It emits None so the consumer renders '--' rather than
    making up a False."""
    strategies = {
        "ALPHA": _strategy_result("ALPHA", _synthetic(n=60, seed=4)),
        "BETA":  _strategy_result("BETA",  _synthetic(n=60, seed=5)),
    }
    rows = bootstrap_ci_table(strategies)
    assert all(r["overlaps_benchmark"] is None for r in rows)


def test_producer_consumer_roundtrip_renders_yes_no_or_dash():
    """End-to-end: producer output threads into the consumer renderer
    and the Overlaps-benchmark column reads yes/no for a real bool and
    '--' for None."""
    strategies = {
        "BENCHMARK": _strategy_result("BENCHMARK", _synthetic(n=60, seed=6)),
        "ALPHA":     _strategy_result("ALPHA",     _synthetic(n=60, seed=7)),
    }
    rows = bootstrap_ci_table(strategies)
    _, rendered = table_bootstrap_ci(rows)
    overlap_column = [r[4] for r in rendered]
    assert all(cell in {"yes", "no"} for cell in overlap_column)

    # With no BENCHMARK -> the consumer should render '--'.
    no_bench = {"ALPHA": _strategy_result("ALPHA", _synthetic(n=60, seed=8))}
    rows = bootstrap_ci_table(no_bench)
    _, rendered = table_bootstrap_ci(rows)
    assert rendered[0][4] == "—"


# ── FIX 3 -- editor export re-injects the eight evidence tables ────────

def _docx_text(content: bytes) -> str:
    """Concatenate every paragraph + cell of a DOCX into a single
    string so tests can spot-check the rendered surface without
    walking the python-docx XML by hand."""
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


_MIN_APPENDIX_FIXTURE = {
    # Empty cache rows -- the table builders fail-open to a header-only
    # table per existing convention. Exercising the wiring, not the
    # numeric correctness of the appendix data.
    "summary_statistics":     [],
    "strategy_results":       {},
    "bootstrap_ci_sharpe":    [],
    "factor_loadings":        [],
    "crisis_performance":     None,
    "cost_sensitivity":       None,
    "invariant_summary":      None,
    "audit_disclosures":      {},
}


def test_editor_export_reinjects_appendix_tables_for_analytical_appendix():
    """The previous behaviour: build_editor_docx renders only the
    TipTap prose, dropping all eight evidence tables. When the document
    type is analytical_appendix and the caller passes appendix_data,
    the new code path appends the eight table headings after the
    prose."""
    draft = {
        "title":         "Analytical Appendix",
        "document_type": "analytical_appendix",
        "content_json": {
            "type": "doc",
            "content": [
                {"type": "heading",
                 "attrs": {"level": 1},
                 "content": [{"type": "text", "text": "Section A intro"}]},
                {"type": "paragraph",
                 "content": [{"type": "text",
                              "text": "The appendix opens with summary stats."}]},
            ],
        },
    }
    content = build_editor_docx(draft, appendix_data=_MIN_APPENDIX_FIXTURE)
    text = _docx_text(content)
    # All eight table titles must appear in the rendered DOCX.
    for marker in ("Table A1.", "Table B1.", "Table C1.", "Table D1.",
                   "Table E1.", "Table F1.", "Table G1.", "Table H1."):
        assert marker in text, f"missing table title: {marker}"
    # The author's prose is still rendered alongside the tables.
    assert "Section A intro" in text
    assert "The appendix opens with summary stats." in text


def test_editor_export_omits_tables_for_other_document_types():
    """Brief / midpoint paper / arbitrary editor drafts must NOT have
    the appendix table block appended -- only analytical_appendix
    triggers it."""
    draft = {
        "title":         "Executive Brief",
        "document_type": "executive_brief",
        "content_json": {
            "type": "doc",
            "content": [
                {"type": "paragraph",
                 "content": [{"type": "text",
                              "text": "Brief lead paragraph."}]},
            ],
        },
    }
    # Pass appendix_data deliberately -- it should be ignored on a
    # non-appendix document type.
    content = build_editor_docx(draft, appendix_data=_MIN_APPENDIX_FIXTURE)
    text = _docx_text(content)
    assert "Brief lead paragraph." in text
    # None of the eight table titles should appear.
    for marker in ("Table A1.", "Table B1.", "Table C1.", "Table D1.",
                   "Table E1.", "Table F1.", "Table G1.", "Table H1."):
        assert marker not in text, (
            f"unexpected table block in non-appendix export: {marker}")


def test_editor_export_omits_tables_when_appendix_data_absent():
    """When the document_type is analytical_appendix but no data dict
    is passed (e.g. an early-render path that cannot await the gather),
    the builder must NOT raise -- it just skips the table block."""
    draft = {
        "title":         "Analytical Appendix",
        "document_type": "analytical_appendix",
        "content_json": {
            "type": "doc",
            "content": [
                {"type": "paragraph",
                 "content": [{"type": "text", "text": "Prose only."}]},
            ],
        },
    }
    content = build_editor_docx(draft)
    text = _docx_text(content)
    assert "Prose only." in text
    for marker in ("Table A1.", "Table B1."):
        assert marker not in text
